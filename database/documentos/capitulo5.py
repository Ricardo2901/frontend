"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""
def capitulo5():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 5
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo V.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    ########################################################################################################################################################################
    # Capitulo 5
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.- DESCRIPCIÓN DE LAS CONDICIONES DEL ÁREA SUJETA A CAMBIO DE USO DE SUELO EN TERRENOS FORESTALES, QUE INCLUYA CLIMA, TIPOS DE SUELO, PENDIENTE MEDIA, RELIEVE, HIDROGRAFÍA Y TIPOS DE VEGETACIÓN Y DE FAUNA.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.1. Fines a que está destinado el área de cambio de uso de suelo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El área Sujeto a estudio se encuentra dentro de los terrenos rústicos conocidos como de ___________________________________________________________________________________________________________, esta área se encuentra como uso preferentemente ___________, y se localiza un tipo de vegetación perteneciente a ____________________________________________________________, como lo manifiesta en la carta de uso de suelo del INEGI, el suelo que ostenta es de tipo ________________________________________________________________. Por todos los factores bióticos y abióticos anteriores, la promovente reúne las características principales para el nuevo uso que se pretende dar al área, siendo esta, ____________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDentro del área de cambio de uso de suelo se cuenta con la biodiversidad vegetativa típica de este ecosistema al cual se encuentra representada en el sistema ambiental (SA) y su afectación no tendrá influencia y no se afectará la biodiversidad con el Cambio de Uso de Suelo. Dentro de las características abióticas del Área para el cambio de Uso de Suelo Forestal (ACUSTF), presenta pendientes que oscila entre _______________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nAnte las actividades que se desarrollarán en el área, respecto a la fauna es muy aislado el avistamiento de algunas especies, al menos en los recorridos realizados no se detectó ni se visualizaron físicamente la presencia de algunas especies que posiblemente encuentren enlistadas en alguna categoría, sin embargo, se pueden observar que se desplazarán dentro del mismo sistema ambiental. La función que se desarrollará con el Cambio de Uso del Suelo que se pretende _______________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.1. Fines a que está destinado el área de cambio de uso de suelo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Para la elaboración de la presente caracterización del área de Cambio de Uso de Suelo se utilizó la ___________________________________________________________________________________________ describiendo lo siguiente: El clima dominante en el área de estudio corresponde a climas _________________________________________________________________', de acuerdo con la clasificación climática de Köppen (1948), modificado por Enriqueta García (1964), a continuación se mencionan los tipos de climas que se encuentran en el área de estudio en mención. (Ver Mapa 5-1).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.1.- Tipos de climas del CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.2 ###
    #########################
    tabla5 = doc.add_table(rows=6, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nA continuación, se describe los tipos de climas presentes en el área del CUSTF.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=3, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(3):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'


    ########################################################################################################################################################################
    # Capitulo 5.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.2.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.2.1.- Temperatura registrada en el sitio del proyecto.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.2.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("El proyecto se ubica dentro del Municipio de _______________________, la estación meteorológica más cercana, se localiza aproximadamente a ___________________________________________ y pertenece a la Red Meteorológica de CONAGUA, la cual se encuentra en el Municipio de __________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.2.1 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=3, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(3):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.2.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nLa estación meteorológica tiene los siguientes registros desde el año ______ hasta el año _______, reporta una temperatura máxima promedio de ______ ºC, una temperatura media promedio de _______ °C y una temperatura mínima promedio de ______ °C. Los meses con temperatura más bajas ocurren predominantemente en la época de otoño-invierno en los meses de _______________________________________________________________________________________________________________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.2.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.2.- Temperatura promedio.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.2.1 ###
    #########################
    tabla5 = doc.add_table(rows=14, cols=4, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(14):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.2.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5 = doc.add_paragraph()
    imagenCapitulo5.text = ''
    imagenCapitulo5 = doc.add_picture('capitulo5/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo5.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo5.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 5.2.1 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('\nGrafica 5.1.- Temperatura promedio.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 5.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.2.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.2.2. Precipitación.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.2.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Las precipitaciones observadas se tiene un registro anual acumulado de _________ milímetros de lluvia durante _________, siendo los ______________ y agosto los meses más lluviosos con ___________________ y el mes donde se presentó menor cantidad de lluvia fue el mes de marzo con _______________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.2.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.3.- Precipitación promedio de los últimos 29 años.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.2.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=14, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(14):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.2.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5 = doc.add_paragraph()
    imagenCapitulo5.text = '\n'
    imagenCapitulo5 = doc.add_picture('capitulo5/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo5.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo5.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 5.2.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('\nGrafica 5.2.- Precipitación promedio.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 5.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.2.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.2.3. Evapotranspiración.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.2.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Los valores mensuales de evapotranspiración se calcularon de acuerdo al método de Thornthwaite (1948), este método es basado en la determinación de la evapotranspiración en función de la temperatura media correlacionada con la duración astronómica del día y el número de días. Por lo que cuando más alta es la temperatura, mayor es el valor de evapotranspiración. _______________________________________________________________________, la mayor concentración de valores de evapotranspiración se presentó en el mes de agosto, de acuerdo a la estación meteorológica que registra estos datos, a continuación, _____________________________________________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.2.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.4.- Evapotranspiración de 1981 al 2010.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.2.3 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=14, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(14):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.2.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5 = doc.add_paragraph()
    imagenCapitulo5.text = '\n'
    imagenCapitulo5 = doc.add_picture('capitulo5/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo5.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo5.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 5.2.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('\nGrafica 5.2.- Evapotranspiración.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Grafica del capitulo 5.2.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5 = doc.add_paragraph()
    imagenCapitulo5.text = '\n'
    imagenCapitulo5 = doc.add_picture('capitulo5/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo5.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo5.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 5.2.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('\nGrafica 5.4.- Climograma estación 5003 CONAGUA.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 5.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.2.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.2.4. Viento')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.2.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Para obtener la velocidad del viento se obtuvo de la página Meteored, en la cual pudimos acceder a los datos del ___________________________________________ del Municipio de __________________________________________________________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.2.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.5.- Velocidad de viento.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.2.4 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=14, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(14):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3. Suelo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La carta edafológica indica la distribución geográfica de los suelos del país, clasificados de acuerdo con las descripciones de unidades FAO/UNESCO 1968, modificadas por DETENAL en 1970. Estas modificaciones consisten básicamente, en agregar nuevas subunidades que se han encontrado en el país y que no se consideran en la clave original de la FAO, el INEGI realizó una serie de modificaciones aplicables en México, con base a esto, se toma de referencia la carta edafológica a escala 1: 250,000, __________________ donde indica que de acuerdo a las condiciones de clima, fisiografía y geología han determinado la ocurrencia y abundancia de diferentes tipos de suelo, donde __________________________________________________________, (Mapa 5-2). En la siguiente tabla se presentan los tipos de suelos existentes en el área del CUSTF.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.6.- Tipo de suelo en el ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=6, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nDescripción de las características de cada tipo de suelo presentes en las diferentes áreas:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f"\nCaracteristica {lista + 1}: Descripcion {lista + 1}")
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        descripcionCapitulo5.bold = True
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1. Tipos de erosión ')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Con respecto al grado de susceptibilidad a la erosión hídrica y antrópica, así como sus causas que lo originan en el área del ACUSTF se manifiesta lo siguiente: el área por su ubicación está catalogada según la Cartografía del INEGI _________________ en su cobertura de degradación de suelo de acuerdo a ello, la superficie del área presenta erosión de tipo _____________________________________, (Mapa 5-3) esto se debe a las características topográficas, principalmente la orografía y las pendientes donde se encuentran inmersas. INEGI lo cataloga sin erosión.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.7.- Tipos de erosión presentes en el CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=6, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nDescripción de los tipos de erosión presentes en el área del ACUSTF:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nEl grado de erosión se conoce midiendo la capa superficial que queda en una superficie después de un evento erosivo determinado ya sea lluvia (Hídrica), viento (Eólica) o por actividades humanas (Antrópica). Dentro del ACUSTF de acuerdo a la carta de INEGI, el área presenta erosión hídrica y antrópica, sin embargo, la mayoría de la superficie no presenta erosión.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    #descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f"\nCaracteristica {lista + 1}:")
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        descripcionCapitulo5.bold = True
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f"Descripcion {lista + 1}")
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        #descripcionCapitulo5.bold = True
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.-Metodología para determinar la erosión hídrica en el área de Cambio de Uso de Suelo.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La degradación del suelo se define como “un grupo de procesos que ocasionan el deterioro del recurso, los cuales provocan una disminución de la productividad biológica y la pérdida de la Biodiversidad”. En este sentido, el estado de degradación en que se encuentran los suelos de uso pecuario y forestal, se estima por medio de las pérdidas de suelo que ocurren en los terrenos, de modo que sea posible determinar si el uso que se está dando a estos es el correcto. Cuando la tasa de erosión es mayor que la tasa de formación del suelo, es señal de que el manejo está originando su degradación y se hace necesario realizar prácticas y obras de conservación, para de esa forma contribuir al desarrollo sostenible de los recursos naturales. Para estimar la erosión de los suelos se ha utilizado la Ecuación Universal de Pérdida de Suelo (EUPS), un Modelo que permite estimar la erosión actual en campo y la potencial de dicho recurso. Esta ecuación constituye un instrumento de planeación para establecer las prácticas y obras de conservación para que hagan que la erosión actual sea menor que la tasa máxima permisible de erosión. La tasa máxima permisible de pérdidas de suelo es de 10 T/ha (toneladas por hectárea); siendo que mayores pérdidas significan degradación.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nPara estimar la erosión del suelo se puede utilizar la siguiente ecuación:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nE = R * K * LS * C * P")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nDonde:"
                                       '\nE = Erosión del suelo (T/ha año).'
                                       '\nR = Erosividad de la lluvia (Mj/ha mm/hr).'
                                       '\nK = Erosionabilidad del suelo.'
                                       '\nLS = Longitud y grado de pendiente.'
                                       '\nC = Factor de vegetación.'
                                       '\nP = Factor de prácticas mecánicas.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    #di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para explicar la utilización de este Modelo en forma práctica para el área del ACUSTF, se utilizarán algunos resultados que se han obtenido de la investigación en México y que han permitido, a nivel nacional, hacer un uso adecuado de este modelo predictivo. La erosión potencial se estima con la siguiente ecuación (los factores se consideran como inmodificables):')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nEP = R*K*LS")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nLa erosión actual se estima utilizando la ecuación anterior, que considera los factores inmodificables “R”, “K” y “LS”. Los factores de protección, como son la vegetación y las prácticas y obras de manejo para reducir las pérdidas de suelo se pueden modificar.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara utilizar este Modelo, se han propuesto diferentes metodologías para estimar cada una de las variables; sin embargo, la aplicación de algunas de ellas en el campo es difícil de realizar por no contar con la información necesaria. Para evitar estos problemas, en este apartado se presentará una metodología simplificada y adecuada para utilizarse en nuestro país.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nErosividad (R). - La estimación de “R” se puede realizar conociendo la energía cinética de la lluvia y la velocidad de caída de las gotas de lluvia, utilizando la ecuación de Ec = (m*v2) /2; donde “m” es la masa de la lluvia y “v” la velocidad de caída de las gotas de lluvia. Considerando lo complejo de hacer esta estimación, se propuso que un mejor estimador de la agresividad de la lluvia sería este valor de erosividad (R).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara estimar “R” se obtiene el valor de energía cinética por evento, para lo que hay que conocer la intensidad de la lluvia, y obtener el valor de “Ec” y multiplicarlo por la intensidad máxima de la lluvia en 30 minutos. La suma de estos valores en un año da el valor de “R”.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEste procedimiento es complicado cuando no se cuenta con datos de intensidad de la lluvia; por esta razón se buscó correlacionar los datos de precipitación anual con los valores de “R” estimados en el país, utilizando la información de intensidad de la lluvia disponible (Cortés y Figueroa, 1991).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDe acuerdo con este procedimiento, se elaboraron modelos de regresión donde, a partir de datos de precipitación anual (P) se puede estimar el valor de “R” de la EUPS, estos modelos de regresión son aplicados para 14 diferentes regiones del país.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEcuaciones para estimar la Erosividad de la lluvia “R” en la República Mexicana para estimar “R” en el ámbito regional, se puede utilizar la precipitación anual y con un modelo lineal muy simple estimarlo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRegiones de erosividad de la lluvia en México')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 5.3.1.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    mapaCapitulo5 = doc.add_paragraph()
    mapaCapitulo5.text = ''
    mapaCapitulo5 = doc.add_picture('capitulo5/capitulo5311/capitulo5.png')  # Nombre del archivo, tiene que estar en la parte donde se encuentra el documento
    mapaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    mapaCapitulo5.width = Cm(14.18)  # Ancho de la imagen en centimetros
    mapaCapitulo5.height = Cm(8.68)  # Alto de la imagen en centimetros
    mapaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Tabla del capitulo 5.3.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=15, cols=3, style='Table Grid')

    for filasRomanos in range(14):
        numero = filasRomanos + 1
        romano = entero_a_romano(numero)
        cell = tabla5.cell(filasRomanos + 1, 0)
        t5 = cell.paragraphs[0].add_run(romano)
        t5.font.size = Pt(12)
        t5.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ecuacionRegiones = [
        'R = 1.2078P + 0.002276P2',
        'R = 3.4555P + 0.006470P2',
        'R = 3.6752P - 0.001720P2',
        'R = 2.8559P + 0.002983P2',
        'R = 3.4880P - 0.00088P2',
        'R = 6.6847P + 0.001680P2',
        'R = -0.0334P + 0.006661P2 ',
        'R = 1.9967P + 0.003270P2',
        'R = 7.0458P - 0.002096P2',
        'R = 6.8938P + 0.000442P2',
        'R = 3.7745P + 0.004540P2',
        'R = 2.4619P + 0.006067P2',
        'R = 10.7427P - 0.00108P2',
        'R = 1.5005P + 0.002640P2',
    ]

    ecuaciones = range(len(ecuacionRegiones))

    for ecuacion in ecuaciones:
        cell = tabla5.cell(ecuacion + 1, 1)
        t5 = cell.paragraphs[0].add_run(f'{ecuacionRegiones[ecuacion]}')
        t5.font.size = Pt(12)
        t5.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    radioCuadrado = [
        0.92,
        0.93,
        0.94,
        0.92,
        0.94,
        0.9,
        0.98,
        0.98,
        0.97,
        0.95,
        0.98,
        0.96,
        0.97,
        0.95,
    ]

    radio = range(len(radioCuadrado))

    for cuadrado in radio:
        cell = tabla5.cell(cuadrado + 1, 2)
        t5 = cell.paragraphs[0].add_run(f'{radioCuadrado[cuadrado]}')
        t5.font.size = Pt(12)
        t5.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.3.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Respecto a la estimación de la erosión del suelo hídrica que se presenta en el ACUSTF.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nErosividad (K). - La susceptibilidad de los suelos a erosionarse depende del tamaño de las partículas, del contenido de materia orgánica, así como de la estructura, en especial del tamaño de los agregados y de la permeabilidad. Para su estimación se utilizan fórmulas complicadas; para condiciones de campo se recomienda el uso de la textura de los suelos y contenido de materia orgánica, se estime el valor de Erosividad (K). Es importante destacar que a medida que el valor de “K” aumenta, se incrementa la susceptibilidad del suelo a erosionarse.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nValores del Factor K")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.3.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo5 = doc.add_paragraph()
    formulaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo5 = formulaCapitulo5.add_run()
    fCapitulo5.add_picture('capitulo5/capitulo5311/tabla5531.png', width=Cm(9.27), height=Cm(9.36))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Tabla del capitulo 5.3.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla5.cell(0, 0)
    t5 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(0, 1)
    t5 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(1, 0)
    t5 = cell.paragraphs[0].add_run('A')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(2, 0)
    t5 = cell.paragraphs[0].add_run('B')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(3, 0)
    t5 = cell.paragraphs[0].add_run('C')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(1, 1)
    t5 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y loess poco compactados')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(2, 1)
    t5 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(3, 1)
    t5 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien arcillas')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for width in range(4):
        tabla5.cell(width, 0).width = Cm(4.1)

    for width in range(4):
        tabla5.cell(width, 1).width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 5.3.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Fuente: NOM-011-CNA-2000")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nTomando en consideración la condición del suelo presente en el área se determinó que es un tipo de suelo ___ y de acuerdo a la información recabada en campo se cuenta con una cobertura vegetal de ______________________ de materia orgánica, por lo que nos da un factor de K de _____")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nLongitud y grado de pendiente (LS). Este factor considera la longitud y el grado de pendiente. La pendiente media del terreno se obtiene dividiendo la diferencia de elevación del punto más alto del terreno al más bajo entre la longitud del mismo. Esto es:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 5.3.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo5 = doc.add_paragraph()
    formulaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo5 = formulaCapitulo5.add_run()
    fCapitulo5.add_picture('capitulo5/capitulo5311/formula_1.png', width=Cm(5.7), height=Cm(1.59))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripción de la Fórmula del capítulo 5.3.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        'Donde:'
        '\nS = Pendiente media del terreno (%)'
        '\nHf = Altura más alta del terreno (m).'
        '\nHi = Altura más baja del terreno (m).'
        '\nL = Longitud del terreno (m).\n'
    )

    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para calcular LS (el factor de longitud y grado de la pendiente) se puede utilizar la siguiente fórmula:')

    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del capitulo 5.3.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo5 = doc.add_paragraph()
    formulaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo5 = formulaCapitulo5.add_run()
    fCapitulo5.add_picture('capitulo5/capitulo5311/formula_2.png', width=Cm(7.91), height=Cm(1.43))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripción de la Fórmula del capítulo 5.3.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        'Donde:'
        '\nLS = Factor de longitud y grado de la pendiente.'
        '\nλ = Longitud de la pendiente.'
        '\nS = Pendiente media del terreno.'
        '\nm = Parámetro cuyo valor es 0.5.\n'
    )

    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 5.3.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.1.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.1.- Estimación de la erosión potencial derivada de la realización del proyecto')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Teniendo en consideración lo anterior, a continuación, se presenta la estimación de la erosión potencial para el polígono que involucra al proyecto con base a las características físicas y topográficas que presenta en la actualidad. De acuerdo a los datos climatológicos registrados en la estación 5003 de CONAGUA que se encuentra ubicada en el municipio de Arteaga a 5.5 kilómetros del área del CUSTF con dirección sur denominada “Arteaga (OBS)”, la precipitación en la región fue de 285.2 mm anuales esta precipitación es promedio de 29 años de registro de 1981 hasta el año 2010, y considerando que el Estado de Coahuila se localiza en la región IV de acuerdo a el mapa de regiones de erosividad de la lluvia en México por lo que el valor de R para el proyecto sería:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Poner las formulas")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nPor otra parte, a continuación, se presenta el valor de K para la sección de terreno involucrada.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nEstimación de Erosión del Suelo.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.8.- Porcentaje de la cubierta vegetal en el área CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=4, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nDe la misma manera, para estimar el valor de LS se hace necesario tomar en cuenta las características topográficas del polígono de afectación.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nEstimación del valor del Factor de longitud y grado de la pendiente del ACUSTF")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.9.- Valor de longitud y grado de la pendiente del CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFinalmente, se estima la Erosión Potencial (Ep) sustituyendo estos valores en la ecuación, obtenidos en la fórmula: Ep = R*K*LS*C. Los resultados se presentan en la tabla siguiente:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.10.- Erosión potencial para el Cambio de Uso de Suelo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.11.- Valores del Factor C que se pueden utilizar para estimar pérdidas de suelo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.1.1 ###
    #########################
    tabla5 = doc.add_table(rows=15, cols=4, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(15):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFuente SAGARPA.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.12.- Valor P.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 5.3.1.1.1 ###
    #########################
    tablaCapitulo5 = doc.add_table(cols=2, rows=9, style='Table Grid')

    practicaValorP_5 = [
        'Surcado al contorno',
        'Surcos rectos',
        'Franjas al contorno',
        'Terrazas (2-7% de pendiente)',
        'Terrazas (7-13% de pendiente)',
        'Terrazas mayores de 13%',
        'Terrazas de banco',
        'Terrazas de banco en contrapendiente',
    ]

    valorP_5 = [
        '0.75-0.90',
        '0.80-0.95',
        '0.60-0.80',
        '0.5',
        '0.6',
        '0.8',
        '0.1',
        '0.05',
    ]

    encabezados_5 = ['Práctica', 'Valor de P']

    # Insertar encabezados
    for col in range(len(encabezados_5)):
        cell = tablaCapitulo5.cell(0, col)
        run = cell.paragraphs[0].add_run(encabezados_5[col])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insertar prácticas
    for i in range(len(practicaValorP_5)):
        cell = tablaCapitulo5.cell(i + 1, 0)
        run = cell.paragraphs[0].add_run(practicaValorP_5[i])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Insertar valores
    for i in range(len(valorP_5)):
        cell = tablaCapitulo5.cell(i + 1, 1)
        run = cell.paragraphs[0].add_run(valorP_5[i])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Pintar celda específica de verde
    celdaVerde_5 = tablaCapitulo5.cell(1, 1)
    cell_background_color(celdaVerde_5, '92D050')

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFuente SAGARPA.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.3.1.1.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.2.- Erosión del suelo en la condición actual en el área para ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La Erosión Potencial (Ep) sustituyendo estos valores en la ecuación, obtenidos en la fórmula: Ep = R*K*LS*C, dado lo siguiente el valor de R sería de _________ de acuerdo a la tabla de ecuaciones para estimar la erosión de la lluvia y el mapa de regiones de la erosión de la lluvia, en la tabla de valores del factor K de acuerdo al tipo de vegetación y el porcentaje de la cobertura de la misma arroja un valor de ______ y el factor de longitud de grado de pendiente que es igual ____________ resultado tenemos que la erosión potencial sería de ________________________________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Poner las formulas")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ########################################################################################################################################################################
    # Capitulo 5.3.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.1.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.3.- Erosión potencial con el cambio de uso de suelo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.1.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Para calcular la pérdida de Suelo se aplicará la ecuación potencial de acuerdo a la siguiente fórmula utilizando los valores obtenidos de las variables R, K, LS, quedando como sigue: Ep = R*K*LS")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Poner descripcion de la formula")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.1.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.4.- Resultados obtenidos en el área del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.13.- Erosión hídrica con y sin proyecto en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.1.4 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.1.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("En las condiciones actuales por efecto de la lluvia se pueden tener pérdidas de _____________________, con la implementación del proyecto al quedar desnudo el suelo incrementa una pérdida hasta ______________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.1.4.- Resultados obtenidos en el área del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La metodología que se empleó para obtener dichos resultados es la tomada por SAGARPA, la cual es la siguiente:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo5 = di5.add_run("Predicción de la erosión eólica: La predicción de erosión eólica se puede llevar a cabo por la ecuación desarrollada por Chepil (1963) similar a la propuesta por Wischmeier (1968).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("X a = (F, G, R, W, V)")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                       '\nXa = Promedio potencial de erosión anual.'
                                       '\nF = Erosión del suelo'
                                       '\nG = Factor local geográfico para la erosión por viento.'
                                       '\nR = Rugosidad de la superficie del suelo.'
                                       '\nW = Ancho equivalente del campo.'
                                       '\nV = Cantidad equivalente de cubierta vegetal.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    #di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFactores considerados.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La velocidad del viento es calculada con la suposición de que la velocidad es superior a la necesaria para mover una partícula del suelo. La humedad del suelo es considerada tratando de encontrar que la erosión del suelo por viento es una función de las fuerzas de cohesión del agua alrededor de las partículas. La máxima erosividad por viento se presenta en suelos que contienen menos de 1/3 de la humedad al punto de marchitamiento permanente (PMP), se considera como un suelo secado al aire) sobre este contenido de humedad la erosión decrece hasta el contenido de PMP, hasta cierto punto en donde la erosión decrece al máximo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    #descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La ecuación usa agregados mayores a 0.84 mm obtenidos por tamizado en suelo seco.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    #descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFactor climatico G.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('C = 1/100 i=1∑12 (Ve3 / 100) (((PET - P) / PET) * n) ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                       '\nG = Promedio de la erosión eólica anual.'
                                       '\nV = Velocidad media mensual a 2 metros de altura, m/s.'
                                       '\nP = Precipitación pluvial, mm.'
                                       '\nPET = Evapotranspiración potencial, mm.'
                                       '\nn = Número de día del mes cuando se tomó la velocidad del viento')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEl número de días sobre el cual la erosión ocurre es asumido que sea proporcional a (PET - P) / PET por el número de días total al mes.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('PET, puede ser estimado por Penman, Thornthwaite, Blanney, etc.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para estimar la erosión eólica del área de CUSTF se utilizaron los siguientes factores.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5. 14.- Valor de factores para cálculo de erosión en el área de CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=2, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nLos valores mencionados anteriormente se obtienen de la siguiente manera:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nErosividad (K o F). - Para conocer el valor de K que se obtiene del siguiente cuadro:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo5 = doc.add_paragraph()
    formulaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo5 = formulaCapitulo5.add_run()
    fCapitulo5.add_picture('capitulo5/capitulo5311/tabla5531.png', width=Cm(9.27), height=Cm(9.36))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla5.cell(0, 0)
    t5 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(0, 1)
    t5 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(1, 0)
    t5 = cell.paragraphs[0].add_run('A')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(2, 0)
    t5 = cell.paragraphs[0].add_run('B')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(3, 0)
    t5 = cell.paragraphs[0].add_run('C')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(1, 1)
    t5 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y loess poco compactados')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(2, 1)
    t5 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(3, 1)
    t5 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien arcillas')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for width in range(4):
        tabla5.cell(width, 0).width = Cm(4.1)

    for width in range(4):
        tabla5.cell(width, 1).width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nValor del Factor K = 0.20")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFACTOR G:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Para obtener el factor G se utilizó información meteorológica de CONAGUA, en su estación, ____________________________________________, para el valor del viento se obtuvo de la página Meteored _____________ insertar el url _________________, en la cual arroja datos del año _____, los cuales son datos recogidos por las estaciones meteorológicas más cercanas al __________________________________________________________________________________________________________________:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nG.C = 1/100 i=1∑12 (Ve3 / 100) (((PET - P) / PET) * n)")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    descripcionCapitulo5.font.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.15.- Valores utilizados en esta ecuación.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=3, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(3):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nPoner las formulas aqui")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFACTOR R:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Este considera la rugosidad del terreno la cual está, influenciado por el tipo de suelo específicamente en el tamaño granular de las partículas, sabiendo que el tipo de suelo presente en el área del ACUSTF, es tipo Cambisol presenta un tamaño promedio de 0.5 mm")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFACTOR W:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Este factor contempla la distancia de afectación del área (ancho del terreno en estudio).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.16.- Valores del Factor C que se pueden utilizar para estimar pérdidas de suelo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    tabla5 = doc.add_table(rows=20, cols=4, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Fuente SAGARPA.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.1.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.17.- Valor P.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 5.3.1.1.1 ###
    #########################
    tablaCapitulo5 = doc.add_table(cols=2, rows=9, style='Table Grid')

    practicaValorP_5 = [
        'Surcado al contorno',
        'Surcos rectos',
        'Franjas al contorno',
        'Terrazas (2-7% de pendiente)',
        'Terrazas (7-13% de pendiente)',
        'Terrazas mayores de 13%',
        'Terrazas de banco',
        'Terrazas de banco en contrapendiente',
    ]

    valorP_5 = [
        '0.75-0.90',
        '0.80-0.95',
        '0.60-0.80',
        '0.5',
        '0.6',
        '0.8',
        '0.1',
        '0.05',
    ]

    encabezados_5 = ['Práctica', 'Valor de P']

    # Insertar encabezados
    for col in range(len(encabezados_5)):
        cell = tablaCapitulo5.cell(0, col)
        run = cell.paragraphs[0].add_run(encabezados_5[col])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insertar prácticas
    for i in range(len(practicaValorP_5)):
        cell = tablaCapitulo5.cell(i + 1, 0)
        run = cell.paragraphs[0].add_run(practicaValorP_5[i])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Insertar valores
    for i in range(len(valorP_5)):
        cell = tablaCapitulo5.cell(i + 1, 1)
        run = cell.paragraphs[0].add_run(valorP_5[i])
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Pintar celda específica de verde
    celdaVerde_5 = tablaCapitulo5.cell(1, 1)
    cell_background_color(celdaVerde_5, '92D050')

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nFuente SAGARPA.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nCabe hacer mención que, para estimar la erosión eólica, para escenarios con proyecto y con medidas de mitigación uno de los factores que influyen en los resultados es la velocidad del viento y el factor práctica de manejo.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nUtilizando la información anterior y la ecuación se tiene lo siguiente:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.2.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.2.1- Erosión actual en el área de cambio de uso de suelo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.2.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Xa = (F, G, R, W, V) ")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Poner el resto de las formulas")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.2.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.2.2.- Erosión con la implementación del proyecto')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.3.1.2.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Xa = (F, G, R, W)")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Poner el resto de las formulas")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.3.1.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.3.1.2.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.3.1.2.3.- Resultados obtenidos en el ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.3.1.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.18.- Erosión eólica para en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.3.1.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.3.1.2.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("En las condiciones actuales por efecto del viento se tiene una pérdida de suelo de _____________________s, con la implementación del proyecto al quedar desnudo el suelo se incrementa hasta ____________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.4.- Geología')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Se utilizó la Carta Geológica __________________ del INEGI y el conjunto de datos vectoriales del continuo nacional de efectos geológicos escala 1: 250,000, en formato digital, encontrando que los tipos de roca existentes en el ACUSTF pertenecen al tipo _______________________________________________________________________________________________________ (Mapa 5-4) las rocas se enlistan a continuación:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.19.- Tipos de rocas en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.4 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("A continuación, se describen cada uno de los tipos de roca encontrados dentro del ACUSTF:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di5 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo5 = di5.add_run("Tipo: Descripcion")
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.5.- Topografía.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.5.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.5.1.- Pendiente media')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.5.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("En el ACUSTF presenta pendientes que oscilan entre los _______, debido a que encuentran en topoformas de ___________________, por ello los grados de inclinación son de _________________________________________________________________________ (Mapa 5-5). A continuación, se presentan las pendientes registradas.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.5.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.20.- Pendientes dentro del ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.5.1 ###
    #########################
    tabla5 = doc.add_table(rows=8, cols=6, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(8):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.5.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.5.2.- Exposición del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.5.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("De acuerdo a la orografía que se presenta en el ACUSTF, donde exposición _________________________________________________________________________________________________________ (Mapa 5-6).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.5.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.21.- Exposición del CUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.5.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.5.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.5.3.- Elevación del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.5.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La elevación que presenta el área de cambio de uso de suelo oscila entre ____________________________________________________________________________, ver mapa 5.7, a continuación, se presenta el cuadro de elevaciones:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.5.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.22.- Elevación del CUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.5.3 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.5.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.5.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.5.4.- Relieve ')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.5.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("El área se encuentra ubicada en la ......... (Mapa 5-8).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.5.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.23.- Tipo de topo formas en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.5.4 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=4, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6. Hidrografía ')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    

    #########################
    ### Descripcion del capitulo 5.6 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("El ACUSTF se encuentra inmersa dentro de La ______________________________________________________________________________________________ de acuerdo al simulador de flujo de aguas de cuencas hidrográficas INEGI, ______________________________. El uso más importante del agua es para recarga de los acuíferos ______________________________________________.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("De acuerdo a los datos arrojados por el SIATL (simulador de flujo) del INEGI en el Área de Cambio de Uso de Suelo, en el predio se presenta escurrimientos ____________________________, donde se tiene que estos pequeños Escurrimientos o corrientes superficiales  tiene una intensidad de lluvia de ____________ con una precipitación promedio de __________ anuales, una duración de concentración de esta agua de ___________ con un período de retorno de 1 años y un coeficiente de escorrentía de _______________________________________________________, con una elevación máxima de _________ y elevación mínima de _________ con una pendiente media de _____% y una longitud del caudal de ______, así como su área de drenada de _____ km2 , la época donde mayor precipitación se registra de acuerdo al a estación meteorológica son los meses ___________________. (Mapa 5-9).")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

    ########################################################################################################################################################################
    # Capitulo 5.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6.1.- Permeabilidad del área en estudio.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.6.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("La Permeabilidad es la propiedad que tiene el suelo de transmitir el agua y el aire y es una de las cualidades más importantes que han de considerarse para el desarrollo de cualquier tipo de proyecto, mientras más permeable sea el suelo, mayor será la filtración. El área de cambio de uso de suelo presenta una ______________________________, ver mapa a continuación se presenta.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.6.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.24.- Permeabilidad en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.6.1 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.6.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6.2.- Metodología para el cálculo de infiltración dentro del área en estudio.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.6.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Para la cuantificación del volumen medio anual de escurrimiento natural se determinó indirectamente, mediante la siguiente expresión:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        '\nVolumen Anual de Escurrimiento = Precipitación Anual * Área Total * Coeficiente de Escurrimiento'
    )
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.italic = True
    descripcionCapitulo5.font.size = Pt(11)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Area Total
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nÁrea total")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    descripcionCapitulo5.underline = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        'Área total del estudio (ha) * 10000 = área en metros cuadrados'
    )
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.italic = True
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Coeficiente de Escurrimiento
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("\nCoeficiente de Escurrimiento")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    descripcionCapitulo5.underline = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El cual se calcula mediante las fórmulas siguientes:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del Capitulo 5.6.2 ###
    #########################
    tabla5 = doc.add_table(cols=2, rows=3, style='Table Grid')

    cell = tabla5.cell(0, 0)
    t5 = cell.paragraphs[0].add_run('COEFICIENTE DE ESCURRIMIENTO ANUAL (Ce)')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(0, 1)
    t5 = cell.paragraphs[0].add_run('K: PARAMETRO QUE DEPENDE DEL TIPO Y USO DE SUELO')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(1, 0)
    t5 = cell.paragraphs[0].add_run('Ce = K(P-250) / 2000')
    t5.font.size = Pt(12)
    t5.font.name = 'Times New Roman'
    t5.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(1, 1)
    t5 = cell.paragraphs[0].add_run('Si K resulta menor o igual que 0.15')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(2, 0)
    t5 = cell.paragraphs[0].add_run('Ce = (K(P-250) / 2000) + (K - 0.15) / 1.5')
    t5.font.size = Pt(12)
    t5.font.name = 'Times New Roman'
    t5.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(2, 1)
    t5 = cell.paragraphs[0].add_run('Si K es mayor que 0.15')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '4F81BD')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            cell.height = Cm(1.22)
            cell.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    #########################
    ### Descripcion del capitulo 5.6.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        "\nDónde:"
        '\nCe = Coeficiente de escurrimiento para diferentes superficies'
        '\nP = Precipitación media anual'
        '\nK = Factor que depende de la cobertura arbolada y del tipo de suelo la cual se describe en el siguiente cuadro:'
        )
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Valores del Factor K')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.6.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo5 = doc.add_paragraph()
    formulaCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo5 = formulaCapitulo5.add_run()
    fCapitulo5.add_picture('capitulo5/capitulo5311/tabla5531.png', width=Cm(9.27), height=Cm(9.36))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Tabla del capitulo 5.6.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla5.cell(0, 0)
    t5 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(0, 1)
    t5 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    t5.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla5.cell(1, 0)
    t5 = cell.paragraphs[0].add_run('A')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(2, 0)
    t5 = cell.paragraphs[0].add_run('B')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(3, 0)
    t5 = cell.paragraphs[0].add_run('C')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla5.cell(1, 1)
    t5 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y loess poco compactados')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(2, 1)
    t5 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla5.cell(3, 1)
    t5 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien arcillas')
    t5.font.size = Pt(12)
    t5.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for width in range(4):
        tabla5.cell(width, 0).width = Cm(4.1)

    for width in range(4):
        tabla5.cell(width, 1).width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 5.6.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Fuente: NOM-011-CNA-2000')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Tomando en consideración la condición del suelo presente en el área y de acuerdo a INEGI se determinó que es un tipo de suelo _____ y de acuerdo a la información recabada en campo se cuenta con una cobertura vegetal de _________________________________________ de materia orgánica por lo que nos da un factor de K de ___________ por lo anterior que el Coeficiente de Escurrimiento Anual (Ce) se determinará a través de la siguiente fórmula:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nCe= K ((P-250)/2000)) + (K-0.15)1.5')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescribir el resto')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPor lo anterior el volumen medio anual de escurrimiento natural se determinó mediante el método indirecto, mediante la siguiente expresión:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nVe= (P) (At) (Ce)')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Evapotranspiración por el método de Coutagne
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEvapotranspiración por el método de Coutagne')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nETR = P-xP2')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run(
        'Donde'
        '\nETR= Evapotranspiración m/año'
        '\nP= Precipitación en m/año'
        '\nX= 1/ (0.8 + 0.14 t)'
        '\nDescribir el resto'
        )
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nCon los datos necesarios calculados se podrá obtener el grado de infiltración en el área sujeta a Cambio de Uso del Suelo desde tres escenarios tal y como se manifiesta a continuación.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.6.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6.2.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6.2.1.- Situación actual hidrológica sin proyecto en el área para ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.6.2.1 ###
    #########################

    #########################
    # Infiltracion
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("INFILTRACION")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Infiltración = P – ETR – Ve")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Donde:")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.6.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6.2.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6.2.2.- Con la implementación del proyecto en el área sujeta a cambio de uso de suelo.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.6.2.2 ###
    #########################
    
    #########################
    # Infiltracion
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("INFILTRACION")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Infiltración = P – ETR – Ve")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Times New Roman'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.italic = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Describir el resto")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.6.2.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.25.-	Volumen de escurrimiento en el ACUSTF')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.6.2.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=6, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Considerando la información antes señalada, se interrumpe un volumen de escurrimiento de ___________. La cual se puede capturar con la implementación de obras de conservación.")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.6.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.6.2.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.6.2.3.- Resultados obtenidos de la Infiltración.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.6.2.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.26.- Infiltración en el ACUSTF para los tres escenarios. ')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.6.2.3 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.6.2.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Derivado del análisis se concluye que en la condición actual con la cobertura que posee, se tiene una infiltración normal de _________ anuales, con la implementación del proyecto al quedar sin vegetación esto aumenta la evapotranspiración por lo cual se dejará de captar agua reduciendo su infiltración a ________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 5.7 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7. Tipos de vegetación')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run("Describir los tipos de vegetación =) =)")
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5. 27.-4Tipos de vegetación en el área del CUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presenta la descripción del tipo de vegetación del área de cambio de uso de suelo:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run('Tipo de vegetacion')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.1. Tipos generales de vegetación ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El proyecto contempla la modificación de la cubierta vegetal, con motivo de las actividades de cambio de uso de suelo, donde la vegetación que es de tipo __________________________________________ de acuerdo a la carta Uso de Suelo y Vegetación, su serie _____, a escala 1:250,000 _________________, del Instituto Nacional de Estadística, Geografía e Informática (INEGI), su estado de conservación se encuentra ________________________________________________________________________________. (Ver anexo mapa 5.10.- vegetación del área).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.- Metodología para el estudio de las Comunidades vegetales.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para los métodos de medición de especies se utilizó la diversidad alfa, el cual consiste en saber la diversidad dentro de un hábitat o comunidad, correspondiente a la riqueza de especies que hay en una unidad o hábitat determinado.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Los métodos propuestos se refieren a la medición de la diversidad dentro de comunidades, para diferenciarlos en función de las variables biológicas que miden, se dividen en dos grandes grupos: ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista572 = [
        'Métodos basados en la cuantificación del número de especies presentes (riqueza específica).',
        'Métodos basados en la estructura de la comunidad, es decir, la distribución proporcional del valor de importancia de cada especie (abundancia relativa de los individuos, su biomasa, cobertura, productividad, etc.).'
    ]

    listaCapitulo572 = range(len(lista572))

    for lista in listaCapitulo572:
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f'{lista + 1}) {lista572[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara cada estrato se evaluó lo siguiente:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Densidad Absoluta. Está dada por el número de individuos de una especie o de todas las especies dividido por el número de sitios muestreados.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del Capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_1.png', width=Cm(2.01), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del Capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nD = Densidad'
                                    '\nN = Número de individuos muestreados por especie'
                                    '\nA = número de sitios muestreados o superficie muestrea según sea (x sito, ha o ACUSTF)')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDensidad relativa. Está dada por el resultado de la densidad absoluta entre el número total de todos los individuos muestreados expresados en porcentajes ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del Capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_2.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del Capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nDer = Densidad Relativa'
                                    '\nNi = Número de individuos de la especie'
                                    '\nNt = Número total de individuos de todas las especies')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDominancia absoluta. Se define como el porcentaje de biomasa (área basal o superficie horizontal) que aporta una especie.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del Capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_3.png', width=Cm(2.75), height=Cm(1.5))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nDa = Densidad absoluta'
                                    '\nABi = Área basal de una especie'
                                    '\nA = Área muestreada (sitios muestreados)')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nLa dominancia relativa. Se calcula como la proporción de una especie en el área total evaluada, expresada en porcentaje.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_4.png', width=Cm(4.79), height=Cm(1.50))
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nDor = Densidad relativa'
                                    '\nDai = Densidad absoluta de una especie'
                                    '\nDat= Densidad absoluta total de todas las especies')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFrecuencia absoluta. Permite conocer las veces que se repite una especie en cada sitio de muestreo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_5.png', width=Cm(3.89), height=Cm(1.50))
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nFa = Frecuencia absoluta'
                                    '\nnsi = sumatoria del número de veces que una especie se observa dentro de todos los sitios de muestreo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFrecuencia relativa. Es el resultado de dividir la frecuencia absoluta de cada especie entre el número total de esas especies expresadas en porcentajes.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_6.png', width=Cm(4.74), height=Cm(1.50))
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nFr = Frecuencia relativa'
                                    '\nFai = Frecuencia absoluta de cada especie'
                                    '\nFat = Frecuencia absoluta de todas las especies')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia (IVI). El índice de valor de importancia define cuáles de las especies presentes contribuyen en el carácter y estructura de una Comunidad. Este valor se obtiene mediante la sumatoria de la frecuencia relativa, la densidad relativa y la dominancia relativa.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_7.png', width=Cm(4.99), height=Cm(1.20))
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nIVI = Índice de Valor de Importancia'
                                    '\nDer = Densidad relativa'
                                    '\nDor = Dominancia relativa'
                                    '\nFr = Frecuencia relativa')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Shannon-Wiener (H’). Tiene en cuenta la riqueza de especies y su abundancia. Este índice relaciona el número de especies con la proporción de individuos pertenecientes a cada una de ellas presente en la muestra. Además, mide la uniformidad de la distribución de los individuos entre las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    
    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_8.png', width=Cm(4.89), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nH’ = Índice de Shannon'
                                    '\nS = Número de especies'
                                    '\nPi = Proporción de individuos de la especie entre todas las especies. A mayor valor de H’ mayor diversidad de especies.'
                                    '\nLn = Logaritmo natural')

    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                    '\nDmg = Índice de Margalef'
                                    '\nS = Número de especies.'
                                    '\nN = Número total de individuos'
                                    '\nD = Densidad'
                                    '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                    '\nLn = Logaritmo natural')

    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. El índice de Simpson representa la probabilidad de que dos individuos, dentro de un hábitat, seleccionados al azar pertenezcan a la misma especie. Es decir, cuanto más se acerca el valor de este índice a la unidad, existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 5.7.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                        '\nƛ = índice de dominancia se Simpson'
                                        '\nID = índice de diversidad'
                                        '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, Que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                        '\nDMn = índice de Menhinick'
                                        '\nS= Número total de especies'
                                        '\nN = Numero de total de todos los individuos de todas las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEl índice de Pielou: se expresa como el grado de uniformidad en la distribución de individuos entre especies. Se puede medir comparando la diversidad observada en una Comunidad contra la diversidad máxima posible de una Comunidad hipotética con el mismo número de especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_12.png', width=Cm(4.72), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 5.7.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                        '\nê = índice de Pielou'
                                        '\n∑ = es la sumatoria de la proporción de individuos (pi) por la sumatoria del logaritmo natura de la proporción de individuos (lnpi), o el Índice de Shannon – Wiener '
                                        '\nS = es el número de especies presentes')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    #########################
    ### Formula del capitulo 5.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo572/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde:'
                                        '\nNmax = Es el número de individuos en la especie más abundante.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRango de escala de 0 - 1')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Donde las escalas para la interpretación de los rangos de 0-1 son las siguientes:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo5 = di5.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo5 = di5.add_run('De 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo5 = di5.add_run('Mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.1.- Muestreo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para determinar la homogeneidad de las especies presentes en el área de cambio de uso de suelo, se realizó un inventario con un esfuerzo de muestreo del _____% para arbóreas y el _____% para arbustivas y suculentas en donde se levantaron un total de 22 sitios, de igual forma se levantaron __________ para los estratos gramíneo y herbáceo, sin embargo, el muestreo se realizó, de acuerdo a los tipos de vegetación que se encontraron dentro del área siendo estos, ___________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.2.-Diseño e intensidad de muestreo utilizado.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el diseño de acuerdo al tipo de vegetación y la superficie del área de cambio de uso de suelo se realizó un muestreo sistemático, con una separación de ______________________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 5.7.2.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo5722/cap_5722.png', width=Cm(13.23), height=Cm(4.73))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    ########################################################################################################################################################################
    # Capitulo 5.7.2.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.3.- Número de sitios de muestreo y su distribución en función de las características que presenta cada polígono o polígonos')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo a la superficie de muestreo (__________) se calculó el número de sitios de muestreo donde se obtuvo lo siguiente: ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 5.7.2.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo5723/formula_1.png', width=Cm(9.91), height=Cm(1.48))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 5.7.2.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara el estrato arbóreo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFavor de poner el resto de las formulas')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara el estrato arbustivo y suculento')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nFavor de poner el resto de las formulas')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nUna vez que se identificaron el número de sitios de muestreo se realizó una malla de puntos en forma sistemática para la cual nos arroja una distancia entre cada sitio de muestreo de ____m, ____________________ que para el caso de las arbóreas corresponde a ____% de la superficie muestreada mientras que para los estratos arbustivo y suculento corresponde a ___%.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.4. - Formas de los sitios de muestreo')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para determinar la homogeneidad de las especies presentes en el ACUSTF y realizar un esfuerzo de muestreo menor en función del número de especies, se realizó un muestreo sistemático, en donde se levantaron _____ sitios de muestreo de forma circular con una separación de _______________________________________________________________________________________________________________________. (Ver anexo mapa 5.11.- muestreo del área).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capitulo 5.7.2.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo5_parrafo = doc.add_paragraph()
    formulaCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo5_run = formulaCapitulo5_parrafo.add_run('')
    imagen = formulaCapitulo5_run.add_picture('capitulo5/capitulo5724/cap_5724.png', width=Cm(6.43), height=Cm(5.98))

    # Opcional: espacio después del párrafo
    formulaCapitulo5_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 5.7.2.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara realizar el muestreo de vegetación en el ACUSTF se determinó que para cada sitio se llevaría a cabo el levantamiento de datos mediante el registro de plantas representativas, realizando la medición de altura y cobertura en cm en cada estrato, con lo cual se contabilizó el número de individuos tal y como se manifiestan los datos en las tablas; así mismo se declara que el diseño de muestreo fue sistemático en el cual se levantaron _____ sitios en el ACUSTF esto fue ante la homogeneidad de la vegetación, para el estrato herbáceo y gramíneo que no siempre es continuo si no que depende de la estación del año y por las precipitaciones que se hayan presentado, la condición sigue siendo de tipo anual. El levantamiento de la información de llevo a cabo el __________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nPara este procedimiento se utilizaron los siguientes materiales:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista5724 = [
        'Cinta métrica de 20 m.',
        'Cinta métrica de 3 m. ',
        'Cuerda compensada.',
        'Estacas para señalamiento de sitios.',
        'Mazo',
        'Marcador permanente.',
        'GPS.',
        'Prensa para recolección de muestras',
        'Tabla para registro de especies y dimensiones.',
        'Cuadrante de 1 m2 de PVC.',
        'Cámara fotográfica.'
    ]

    listaCapitulo5724 = range(len(lista5724))

    for lista in listaCapitulo5724:
        di5 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo5 = di5.add_run(f'{lista5724[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15
        descripcionCapitulo5_format.space_after = 0
        descripcionCapitulo5_format.space_before = 0

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.5.-Tamaño de los sitios expresados en m2.s')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Con una superficie de 1,000 m2, para arbóreas, y 250 m2 para arbustivas, suculentas y de 1 m2 para herbáceas y gramíneas')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.2.6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.2.6 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.2.6.- - Variables dasométricas (Diámetro normal, altura, total etc.)')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.2.6 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo a la composición estructural de las especies de zonas áridas, que no presentan un fuste bien definido en el caso de las arbóreas, si no que presentan varias ramas o tallos a ras del suelo, no se puede medir un diámetro normal a la altura del pecho (1.30 m), para realizar una estimación del volumen de las arbóreas se midió el diámetro de la rama principal (céntrica) de la parte media de la rama, para su posterior estimación, así también se midieron de las demás especies, con la altura, cobertura y número de individuos.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.- Resultados del inventario de Matorral Desértico Micrófilo (MDM) en el ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.3.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.7.3.1.- Coordenadas de los sitios de muestreo.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presenta la tabla de coordenadas de los sitios de muestreo del área de cambio de uso de suelo del __________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.28.- Coordenadas de los sitios de muestreo del MDM.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=20, cols=6, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.2.- Resultados del inventario del __________________________________ en el ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presentan los datos obtenidos del resultado del levantamiento por sitio de muestreos del Matorral Desértico Micrófilo, del área de cambio de uso de suelo se observó la condición de la vegetación de acuerdo a las siguientes tablas.A continuación, se presentan los datos obtenidos del resultado del levantamiento por sitio de muestreos del ________________________, del área de cambio de uso de suelo se observó la condición de la vegetación de acuerdo a las siguientes tablas.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.29.- Resultado del inventario del ACUSTF en el ____.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=60, cols=8, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(60):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.3.- Estatus de la vegetación encontrada en el ____ del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presenta el listado de la vegetación encontrada en el área a cambio de uso de suelo, además de su clasificación en estatus de protección por estrato de vegetación.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.30.- Categoría de las especies del MDM.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=60, cols=5, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(60):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.4- Análisis de la información en el _____ del ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.7.3.4.1.- Análisis de la información del estrato de arbóreo del _____.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.4.2.- Análisis de la información del estrato de arbustivo del ____.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.31.- Categoría de las especies del _____.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=40, cols=9, style='Table Grid')
    tabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.3.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5 = doc.add_paragraph()
    imagenCapitulo5.text = '\n'
    imagenCapitulo5 = doc.add_picture('capitulo5/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo5.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo5.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo5.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo5.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('\nGrafica 5.5.- Densidad del estrato Arbustivo del MDM.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El Índice de diversidad es un parámetro que permite conocer la abundancia de una especie o una clase de plantas. Describir el resto.................')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capítulo 5.7.3.4.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 5.7.3.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.32.- Valor de Importancia de las arbustivas ___ en el SA.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=40, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 5.7.3.4.2 ###
    #########################
    """
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """

    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Título de la gráfica del capítulo 5.7.3.4.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Gráfica 5.6.- Valor de Importancia Estrato Arbustivo ___.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEl valor de importancia es un parámetro que mide el valor de las especies, con base a los parámetros de dominancia, densidad y frecuencia, es la suma de estos tres parámetros, sobre tres. Descripcion del capitulo ...')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capítulo 5.7.3.4.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Abundancia #

    #########################
    ### Subtítulo del capítulo 5.7.3.4.2 ###
    #########################
    di5_abundancia = doc.add_paragraph()
    run_abundancia = di5_abundancia.add_run('ABUNDANCIA')
    run_abundancia_format = di5_abundancia.paragraph_format
    run_abundancia_format.line_spacing = 1.15
    run_abundancia_format.space_after = 0
    run_abundancia_format.space_before = 0

    run_abundancia.font.name = 'Arial'
    run_abundancia.font.size = Pt(12)
    run_abundancia.bold = True
    di5_abundancia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 5.7.3.4.2 ###
    #########################
    tituloTabla5b = doc.add_paragraph()
    dti5b = tituloTabla5b.add_run('\nTabla 5.33.- Valor de Importancia de las arbustivas ___ en el Sistema Ambiental.')
    dti5b_format = tituloTabla5b.paragraph_format
    dti5b_format.line_spacing = 1.15
    dti5b_format.space_after = 0

    dti5b.font.name = 'Bookman Old Style'
    dti5b.font.size = Pt(12)
    tituloTabla5b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 5.7.3.4.2 ###
    #########################
    tabla5b = doc.add_table(rows=40, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla5b.cell(rows, cols)
            t5b = cell.paragraphs[0].add_run(' ')
            t5b.font.size = Pt(12)
            t5b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 5.7.3.4.2 ###
    #########################
    """
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """

    imagenCapitulo5b_parrafo = doc.add_paragraph()
    imagenCapitulo5b_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5b_run = imagenCapitulo5b_parrafo.add_run('')
    imagenCapitulo5b_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Título de la gráfica del capítulo 5.7.3.4.2 ###
    #########################
    tituloGrafico5b = doc.add_paragraph()
    dgi5b = tituloGrafico5b.add_run('Gráfica 5.2.- Valor de abundancia absoluta ___.')
    dgi5b_format = tituloGrafico5b.paragraph_format
    dgi5b_format.line_spacing = 1.15
    dgi5b_format.space_after = 0

    dgi5b.font.name = 'Bookman Old Style'
    dgi5b.font.size = Pt(12)
    tituloGrafico5b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 5.7.3.4.2 ###
    #########################
    di5_desc = doc.add_paragraph()
    descripcionCapitulo5b = di5_desc.add_run('\nLa abundancia relativa expresa la representatividad de una especie dentro del conjunto de especies en el área, Descripción del capítulo..................')
    descripcionCapitulo5b_format = di5_desc.paragraph_format
    descripcionCapitulo5b_format.line_spacing = 1.15
    descripcionCapitulo5b_format.space_after = 0
    descripcionCapitulo5b_format.space_before = 0

    descripcionCapitulo5b.font.name = 'Arial'
    descripcionCapitulo5b.font.size = Pt(12)
    di5_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Riqueza de Especie #

    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA DE ESPECIE')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nEl índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nLa riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Dominancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Shannon')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del Capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.34.- Índice de diversidad de estrato de las arbbustivas ___')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.8.- Densidad de Estrato Herbaceo ___.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.3.4.4 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.35.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.10.- Valor de importancia del estrato de las herbáceas en el _____.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.3.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.36.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 4.10.- Abundancia del estrato Herbáceas ____')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.3.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.4.5.- Análisis de la información del estrato suculento en el _____.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.5 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.37.- Densidad del estrato Suculentos en el MDM.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.5 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.11.- Densidad del estrato Suculento en el MDM.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.3.4.5 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.5 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.38.- Valor de importancia del estrato suculento en el ____.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.5 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.12.- Valor de importancia del estrato suculento del MDM.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.3.4.5 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.5 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.39.- Abundancia del estrato suculento del MDM.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.3.4.5 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.13.- Abundancia del estrato suculento del MDM.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.3.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.3.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.3.4.6
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.3.4.6 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.4.6.- Análisis de la información del MDM en el ACUSTF por estrato.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.3.4.6 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Describir el resto del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.3.4.6 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.40.- Rangos y valores resultados de los índices, de _____.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.3.4.6 ###
    #########################
    tabla5 = doc.add_table(rows=16, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(16):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.3.4.6 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De manera general, la vegetación presente en el área del cambio de uso de suelo en el ___________________________________________________________________________________________________________________________________. Para el estrato suculento se considera y se propone hacer un programa de rescate y reubicación como medida de mitigación, ya que cuenta con especies de importancia ecológica, que de acuerdo a su fisionomía y fisiología son de lento crecimiento y difícil regeneración, las cuales fueron muestreadas y registradas, con lo cual se evita el riesgo de pérdida de biodiversidad. En general en el área el estado de conservación de la vegetación va de ____________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.4.- Resultados del inventario de Matorral Desértico Rosetófilo (MDR) en el ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.4.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.7.4.1.- Coordenadas de los sitios de muestreo.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presenta la tabla de coordenadas de los sitios de muestreo del área de cambio de uso de suelo del ___________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.41.- Coordenadas de los sitios de muestreo del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.4.2.- Resultados del inventario del MDR, en el ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presentan los datos obtenidos del resultado del levantamiento por sitio de muestreos en el ________________________________, del área de cambio de uso de suelo se observó la condición de la vegetación de acuerdo a las siguientes tablas.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.42.- Resultado del inventario del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=60, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(60):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.4.3.- Clasificación del estatus de la vegetación encontrada en el MDR, del ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('A continuación, se presenta el listado de la vegetación encontrada en el área a cambio de uso de suelo, además de su clasificación en estatus de protección por estrato de vegetación.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.42.- Resultado del inventario del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=60, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(60):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.4.4.- Análisis de la información del MDR del ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.44.- Densidad del estrato arbustivo del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.4.4.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.1 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.14.- Densidad del estrato Arbustivo del MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.4.4.1 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.45.- Índice de valor de importancia del estrato arbustivo del ____.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.1 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.15.- Valor de importancia del estrato Arbustivo del _____.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.4.4.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.46.- Valor de abundancia del estrato arbustivo del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.1 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Descripcion del capitulo')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.1 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Densidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDensidad')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.47.- Densidad del estrato gramíneo del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.4.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.17.- Densidad del estrato gramíneo del MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.4.4.2 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.48.- Valor de importancia del estrato gramíneo del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.18.- Valor de importancia del estrato gramíneo del MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.4.4.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.49.- Valor de abundancia del estrato gramíneo en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.19.- Valor de abundancia del estrato gramíneo del MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del Capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.2 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Densidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDensidad')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.50.- Densidad del estrato del herbáceo en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.4.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.20.- Densidad del estrato de herbáceo en el MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.4.4.3 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.51.- Valor de importancia del estrato herbáceo en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.21.- Valor de importancia del estrato de herbáceo en el MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.4.4.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.52.- Valor de abundancia del estrato herbáceo en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.19.- Valor de abundancia del estrato gramíneo del MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del Capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\n')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Densidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDensidad')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.53.- Valor de densidad del estrato suculento en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.7.4.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.23.- Valor de densidad del estrato de suculento en el ____.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #######################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 5.7.4.4.4 ###
    #######################################################################################################
    """
        ===================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ===================================================================================================
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    # Índice de valor de importancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de valor de importancia')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.54.- Valor de importancia del estrato de las suculento en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.24.- Valor de importancia del estrato de suculento en el MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break()  # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 5.7.4.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Abundancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nABUNDANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.55.- Valor de abundancia del estrato de las suculento en el MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 5.7.4.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 5.7.4.4.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.25.- Valor de abundancia del estrato de suculento en el MDR.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 5.7.4.4.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nDescripcion del Capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    # Riqueza Especifica #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.font.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Margalef')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Menhinick')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de dominancia #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Simpson')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.3 ###
    #########################
    tabla5 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Indice de Equidad #
    #########################
    ### Descripcion del capitulo 5.7.4.4.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    #descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Índice de Pielou')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de Pielou para el área estudiada da un resultado de _______ el cual indica que _________. Esto se debe a que los valores oscilan entre 0 y 1, donde valores cercanos a 1 indican una distribución equitativa de las especies.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.5
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.7.4.4.4 ###
    #########################
    tabla5 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.7.4.4.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.7.4.4.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.7.3.4.6.- Análisis de la información del MDM en el ACUSTF por estrato.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.7.4.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Describir el resto del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.7.4.4.5 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.56.- Rangos y valores resultados de los índices del MDR.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.7.4.4.5 ###
    #########################
    tabla5 = doc.add_table(rows=16, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(16):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.7.4.4.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De manera general, la vegetación presente en el área del cambio de uso de suelo en el ___________________________________________________________________________________________________________________________________. Para el estrato suculento se considera y se propone hacer un programa de rescate y reubicación como medida de mitigación, ya que cuenta con especies de importancia ecológica, que de acuerdo a su fisionomía y fisiología son de lento crecimiento y difícil regeneración, las cuales fueron muestreadas y registradas, con lo cual se evita el riesgo de pérdida de biodiversidad. En general en el área el estado de conservación de la vegetación va de ____________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.- Fauna Silvestre.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('La metodología utilizada para le evaluación de Fauna Silvestre consistió en definir la forma de análisis de trabajo en el área de estudio, implementando diversas técnicas de  muestreo, utilizando transectos y estaciones olfativas y de escucha, considerando los grupos de vertebrados terrestres representados por Aves, Mamíferos, Reptiles y Anfibios, para este último grupo siempre y cuando se presenten las condiciones adecuadas, para esto se determinaron sus hábitats, se efectuaron observaciones (a simple vista o con binoculares) realizándose de la siguiente forma y orden, todo esto por la cantidad de área que manejaremos en nuestro muestreo. Para los análisis estadísticos se utilizaron el número de individuos observados durante el monitoreo en el área, de igual manera como información para cada grupo faunístico se determinó el número de individuos por área de muestreo y el número de individuos extrapolados al área CUSTF.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el caso de las especies de los lepidópteros se implementaron transectos de muestreo en franja, registrando todas aquellas especies de lepidópteros dentro de la franja de muestreo del ancho del transecto, así mismo en el área se tiene registrada distribución de mariposa monarca, por lo cual los esfuerzos de muestreo son direccionados para detectar a la especie dentro del área.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.- Metodología para el muestreo de fauna en el área del ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.1.- Aves.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.1.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el caso de muestreo de aves se utilizó el método en transectos de franja fija, el cual permite estimar la riqueza específica y la abundancia relativa de las especies de fauna silvestre correspondientes a este grupo, el procedimiento en el cual se basó este muestreo que consta de las siguientes etapas:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista5811 = [
        'Elección del transecto: El punto de partida quedo definido por el tipo de hábitat y tipo de especie, potencialmente presente en este caso, dado que la longitud del tramo o sección del área a estudiar es relativamente amplia, se realizaron transectos que cubrieron un porcentaje de muestreo de la totalidad del área.',
        'Longitud del transecto: El transecto fue lineal y con una extensión de ______ m la cual también puede ser determinada por el observador y franjas de ___ m de ancho a cada lado eje central del transecto de muestreo.',
        'Muestreo: El transecto en su totalidad se recorrió a pie, se registraron todos los individuos avistados dentro de la franja y a cada lado del eje del transecto, mediante binoculares y observación directa.',
        'Análisis de datos: como resultado, se confeccionó una lista de especies presentes por sitio, con sus respectivas estimaciones lo cual permitió estimar la riqueza específica y la abundancia relativa de las especies de fauna silvestre correspondientes. (Nº de individuos por área).'
    ]

    listaCapitulo5811 = range(len(lista5811))

    for lista in listaCapitulo5811:
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f'{lista + 1}) {lista5811[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.2.- Mamíferos.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.1.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el monitoreo de mamíferos se recurrió al empleo de transectos de muestreo de ancho variable, así como el uso de técnicas de identificación indirectas como la localización e identificación de heces fecales, huellas, sitios de alimentación, madrigueras y restos óseos, entre otros y, eventualmente, la observación directa de ejemplares, esto por la dificultad para avistarlos, el proceso metodológico para este muestreo comprende las siguientes actividades:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista5812 = [
        'Elección del transecto: El punto de partida quedo definido por el tipo de hábitat y tipo de especies estableciendo así un transecto de muestreo lineal de ancho variable, además se establecieron cinco cámaras de foto trampeo destituidas en el área de estudio a una distancia de 200 mts entre cámara. Así mismo se colocaron trampas de Sherman las cuales consisten en pequeñas cajas metálicas con una puerta de acceso que se activan al encontrarse algún animal de talla pequeña dentro de ellas como roedores, como atrayente se utilizó una mezcla de avena con crema de cacahuate y vainilla.',
        'Muestreo: El recorrido en transecto se realizó a pie en donde el o los observadores caminan en una línea recta observando a las especies que se avisten dentro del ancho de transecto establecido el cual fue de 25 mts para cada eje del transecto, así mismo durante el recorrido se revisaron las trampas para verificar la captura de algún individuo, además el recorrido se registraron huellas, excretas, restos óseos, pelaje que puedan representar alguna especie de mamífero en el área.',
        'Análisis de datos: como resultado, se confeccionó una lista de especies presentes, con sus respectivas estimaciones lo cual permitió estimar la riqueza específica y la abundancia relativa de las especies de fauna silvestre correspondientes. (Nº de individuos por área).'
    ]

    listaCapitulo5812 = range(len(lista5812))

    for lista in listaCapitulo5812:
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f'{lista + 1}) {lista5812[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.3.- Quirópteros.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.1.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el grupo de los quirópteros a diferencia de los muestreos que comúnmente se realizan para especies de mamíferos terrestres, los murciélagos son especies que se caracterizan por tener actividades nocturnas y crepusculares por lo cual los métodos convencionales tales como captura en trampas metálicas; por observación directa, etc., por mencionar algunas, no son tan efectivos, siendo las técnicas más utilizadas tales como los registros por monitoreo acústico, captura mediante redes, captura mediante trampa de arpa, redes de golpeo, cámaras IR, cámaras térmicas; sin embargo cada técnica para su buen funcionamiento está determinada por las condiciones ecológicas y paisajísticas del área de estudio de tal manera que para el área de estudio al ser un área _______________________________________________________________________________________________________________________________________________, las cuales se colocaron de manera estratégica en los lugares con mayor posibilidad de captura de especies de quirópteros.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Las redes se revisaron en intervalos de tiempo para verificar o descartar la captura de individuos, así como también evitar el estrés de los individuos, cada sitio de red se visitó en un tiempo de ____________________. En caso de ser capturado algún individuo se procedía a su identificación mediante guías de mamíferos, se les tomaría datos biométricos y se les marcaria en la parte interna del ala un código de identificación o numeración para evitar un sobrestimar a los individuos, con estas características y preferencias de la especie aunado al monitoreo de quirópteros implementado en el área se eligen las áreas con estas características para establecer estos sitios de muestreo y así aumentar las posibilidades de captura de la especie en caso de que se encuentre en el área de estudio, dicha especie no se registró dentro del área de cambio de uso de suelo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.4.- Reptiles.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.1.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Se utilizó el método de muestreo en transectos, que es el que permite estimar la riqueza específica y la abundancia relativa, el procedimiento se fue desarrollando de acuerdo a las siguientes etapas:')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista5812 = [
        'Elección del transecto: El punto de partida quedo definido por el tipo de hábitat y tipo de especie, potencialmente presente, en este caso, dado que los hábitats por sitio de estudio, son relativamente homogéneos, el punto de inicio fue seleccionado arbitrariamente, sin embargo, todos los transectos siguieron paralelos al curso transecto inicial.',
        'Longitud del transecto: Cada transecto se realizó en forma lineal y en una extensión de _______ de longitud recorrida y una anchura de _____ (____ a cada lado del transecto).',
        'Muestreo: Cada transecto se recorrió a pie, en un tiempo estandarizado para todos los transectos, se registraron todos los individuos avistados en una franja de 6 metros a cada lado del eje del transecto. Se realizó una exhaustiva revisión del área circundante (dentro de la franja) especialmente bajo piedras, remoción somera de sustratos y cerca de las madrigueras anotando en formatos de campo toda especie correspondiente a este grupo. ',
        'Análisis de datos: Como resultado, se confeccionó una lista de especies presentes por sitio, con sus respectivas estimaciones de densidad y abundancia (Nº de individuos por área).'
    ]

    listaCapitulo5812 = range(len(lista5812))

    for lista in listaCapitulo5812:
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f'{lista + 1}) {lista5812[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.1.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.1.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.1.5.- Lepidópteros.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.1.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el grupo de los lepidópteros se utilizó el método de muestreo en transectos, que es el que permite estimar la riqueza específica y la abundancia relativa, el procedimiento se fue desarrollando de acuerdo a las siguientes etapas: ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista5815 = [
        'Elección del transecto: La elección del método y del transecto quedo definido por el tipo de hábitat y de la especie en cuestión a monitoreo, tomando en cuenta lo anterior, el área de estudio es una área abierta de poca vegetación, esta característica permite al o los observadores tener una visión del área más extensa generando la oportunidad de registrar el mayor número de especies posibles dentro de nuestra superficie de muestreo, una vez analizadas estas variables se optó por implementar el monitoreo estableciendo transectos de franja o de banda, métodos que son adecuados para este tipo de hábitat.',
        'Longitud del transecto: Los transectos se establecieron de manera lineal con una 1,119 mts extensión de m de longitud por 20 m de ancho (10 m a cada lado del transecto), en áreas donde se observaron fauna, las dimensiones del transecto de muestreo son determinadas por el observador dependiendo el tipo de hábitat y la superficie del mismo.',
        'Muestreo: el transecto se recorrió a pie, en un tiempo estandarizado para todos los transectos, durante el recorrido se busca registrar todos los individuos avistados en una franja de 10 metros a cada lado del eje del transecto, se realizó una exhaustiva revisión del área circundante (dentro de la franja), para registrar todos los ejemplares que se encuentren dentro de la superficie del transecto, con el transecto establecido es más fácil registrar más fácilmente aquellas especies sedentarias, territoriales y las de vuelo corto así como también permite la identificación rápida al vuelo o la captura en caso necesario para una mejor identificación, un ejemplo de transecto de muestreo se muestra en la imagen siguiente. Transecto de muestreo de franja, tiene como objetivo registrar a todas las especies que se encuentren dentro de la superficie de muestreo del transecto.',
        'Análisis de datos: Como resultado del recorrido en caso de observaron especies se confecciona una lista de especies presentes, con su identificación y el número de individuos observados para posteriormente realizar los análisis estadísticos utilizando índices de diversidad y riqueza, así como también determinar la densidad de las especies por la superficie de muestreo y la abundancia relativa de las mismas (Nº de individuos por área.).'
    ]

    listaCapitulo5815 = range(len(lista5815))

    for lista in listaCapitulo5815:
        di5 = doc.add_paragraph()
        descripcionCapitulo5 = di5.add_run(f'{lista + 1}) {lista5815[lista]}')
        descripcionCapitulo5_format = di5.paragraph_format
        descripcionCapitulo5_format.line_spacing = 1.15

        descripcionCapitulo5.font.name = 'Arial'
        descripcionCapitulo5.font.size = Pt(12)
        di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('\nLos métodos que se utilizaron en el muestreo de los diferentes tipos de fauna silvestre en el área de estudio son una herramienta básica, que permite al analista por medio de los estudios pertinentes y sus distintos métodos obtener una idea de las especies que pudieran existir en el área y poder hacer una extrapolación a la superficie que se desee, las tomas muéstrales son sencillamente un procedimiento que empleamos para extraer tan solo una pequeña parte de una población dentro de una área a esto lo llamaremos espacio muestral dentro de una área determinada.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.2.- Resultados encontrados en el área sujeta a cambio de uso de suelo ACUSTF')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.57.- Fauna presente en área de cambio de uso de suelo. ')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.2 ###
    #########################
    tabla5 = doc.add_table(rows=16, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(16):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    # Coordenadas de los transectos de muestreo en el ACUSTF. #
    #########################
    ### Descripcion del capitulo 5.8.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run()
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    descripcionCapitulo5.bold = True
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.58.- Coordenadas de los transectos en el ACUSTF, VI: Vértice inicial, VF: Vértice final.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.2 ###
    #########################
    tabla5 = doc.add_table(rows=16, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(16):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.8.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.3.- Resultado de especies faunísticas en el área ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 5.8.3.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3.1 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'V.8.3.1.- Análisis de información del grupo de las aves en el área del ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el análisis de la información del grupo de las aves en el área ACUSTF el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra  el número de individuos por superficie muestreada y el número de individuos extrapolados a la superficie correspondiente al ACUSTF, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.59.- Número de individuos del grupo de aves presentes en el área ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.60.- Listado de las especies observadas en el área ACUSTF con su categoría de riesgo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Sociabilidad (SOCI.); abundancia (ABU.); residencia (RES.); alimentación (ALIM.) y el tipo de observación (OBS.); Sc: Sociabilidad, R: Residente; C: Común, SL: Solitario, GR: Gregario, PJ: Pareja; Sc: Sin categoría, Pr: Sujeta a protección especial; A: Amenazada; P: En peligro de extinción; E: Extinta en medio silvestre.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(10)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.61.- Análisis estadístico por índices de diversidad para el grupo de las aves.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Los índices de diversidad de las ___ especies del grupo de las aves presentes en el área del ACUSTF muestran que para el índice de Shannon  tenemos una diversidad de __________ lo cual quiere decir que para este grupo los valores resultantes se encuentran ________ ya que los rangos de valores para este índice van de 0 a 1.35 para valores bajos, 1.36 a 3.5 para valores medios y para los valores 3.5 en adelante se son aquellos considerados de alta diversidad; para el índice de Simpson resulta una diversidad media de 0.612 y una dominancia de las especies media de 0.388, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad muestra valores bajos de 1.5417 ya que los valores de medida considerados para una baja biodiversidad son para valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores a 5. La especie más representativa fue el ____________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.1 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.62.- Análisis estadístico por índices de riqueza de especies, frecuencia y abundancia relativa para el grupo de las aves en el ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de diversidad para el grupo de las aves de las ____ especies presentes en el área del ACUSTF presenta un índice de diversidad de _______, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de _______, para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del ACUSTF en estudio nos indica la dominancia de la _____________________ como la más representativa, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie lo cual para este índice resulta que la especie más representativa fue ________________________, tal como se puede observar en las siguiente gráfica.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.8.3.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(12.24), height=Cm(8.18))

    #########################
    ### Titulo de la grafica del capitulo 5.8.3.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.26.- Frecuencia y abundancia relativa del grupo de las aves en el área del ACUSTF.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.1 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('En el grupo de las aves dentro del área del acustf, poseé una riqueza específica de ___ especies las cuales tienen una equidad de ___________, con lo cual se puede afirmar que la mayoria de las especies son equitativas. La máxima diversidad que se puede alcanzar en el sistema ambiental de este grupo es de ________ y la diversidad calculada es de __________ que indica que este grupo está cerca de alcanzar la máxima diversidad y posee una distribución equitativa, la especie mas representativa para este grupo _________________________________________ en comparación con las demás especies observadas en el área de estudio, considerando que el grupo tendra un porcentaje de desplazamiento ____% en el área ACUSTF, por lo tanto el grupo de las aves no se vera afectado ya que las espcies que se cuentren dentro del area de estudio se podran desplazar hacia el area del sistema ambiental sin ningun inconveniente.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo con los datos que anteceden por las caracteristicas del area ACUSTF el grupo de las aves se presenta en condiciones de _____________ en cuanto a riqueza y equidad de especies, en cuanto a la dominancia de especies se obtuvieron valores de ________________________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.8.3.1 ###
    #########################
    tabla5 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.8.3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3.2 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.3.2.- Análisis de Información del grupo de los mamíferos en el área CUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el análisis de la información del grupo de los mamíferos en el área del ACUSTF el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al ACUSTF, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.63.- Número de especies del grupo de los mamíferos en el área ACUSTF. ')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.64.- Listado de las especies de mamíferos observadas en el área ACUSTF con su categoría de riesgo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Sociabilidad (SOCI.); abundancia (ABU.); residencia (RES.); alimentación (ALIM.) y el tipo de observación (OBS.); Sc: Sociabilidad, R: Residente; C: Común, SL: Solitario, GR: Gregario, PJ: Pareja; Sc: Sin categoría, Pr: Sujeta a protección especial; A: Amenazada; P: En peligro de extinción; E: Extinta en medio silvestre.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(10)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.65.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef, para el grupo de los mamíferos en el área del ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Los índices de diversidad de las ___ especies del grupo de las aves presentes en el área del ACUSTF muestran que para el índice de Shannon  tenemos una diversidad de __________ lo cual quiere decir que para este grupo los valores resultantes se encuentran ________ ya que los rangos de valores para este índice van de 0 a 1.35 para valores bajos, 1.36 a 3.5 para valores medios y para los valores 3.5 en adelante se son aquellos considerados de alta diversidad; para el índice de Simpson resulta una diversidad media de 0.612 y una dominancia de las especies media de 0.388, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad muestra valores bajos de 1.5417 ya que los valores de medida considerados para una baja biodiversidad son para valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores a 5. La especie más representativa fue el ____________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.2 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.66.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para el grupo de los mamíferos en el área del ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de diversidad para el grupo de los mamíferos de las _____ especies presentes en el área del ACUSTF presenta un índice de _______, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de ________; para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del ACUSTF en estudio nos indica la dominancia de la ____________________- como la más representativa, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie más representativa ________________________, tal como se puede observar en la siguiente gráfica.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.8.3.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(12.24), height=Cm(8.18))

    #########################
    ### Titulo de la grafica del capitulo 5.8.3.2 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.27.- Frecuencia y abundancia relativa del grupo de mamíferos en el área ACUSTF.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.2 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El grupo de los mamíferos posee una riqueza específica de 3 especies las cuales tienen una distribución de 0.9464 con lo cual se puede afirmar que la equidad de especies es alta. La máxima diversidad que este grupo adquiere dentro del sistema ambiental es de _________________________________________________________________________________, por lo cual las especies registradas se podrán desplazar hacia el área del sistema ambiental sin ningún inconveniente. ')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo con los datos que anteceden por las caracteristicas del area Acustf el grupo de las mamiferos se presenta en condiciones _________________________________________________________________________________, acontinuacion se muestra en el cuadro de rangos de valor para el grupo de los mamiferos en el area Acustf.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.8.3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3.3 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.3.3.- Análisis de Información del grupo de los reptiles ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.3.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el análisis de la información del grupo de los reptiles en el área del ACUSTF se muestra el número de individuos (ni) los cuales fueron aquellos individuos observados en campo por la metodología aplicada para este grupo, así como también se muestra el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al ACUSTF, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.67.- Número de individuos del grupo de reptiles del área ACUSTF. ')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.3 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.68.- Listado de las especies de reptiles observadas en el área ACUSTF con su categoría de riesgo.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.3 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Sociabilidad (SOCI.); abundancia (ABU.); residencia (RES.); alimentación (ALIM.) y el tipo de observación (OBS.); Sc: Sociabilidad, R: Residente; C: Común, SL: Solitario, GR: Gregario, PJ: Pareja; Sc: Sin categoría, Pr: Sujeta a protección especial; A: Amenazada; P: En peligro de extinción; E: Extinta en medio silvestre.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(10)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.69.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef, para el grupo de los reptiles en el área del ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Los índices de diversidad de las _____ especies del grupo de los reptiles presentes en el área del ACUSTF muestran que para el índice de Shannon tenemos una diversidad ___________ lo cual quiere decir que para este grupo los valores resultantes se encuentran bajos ya que los rangos de valores para este índice van de 0 a 1.35 para _________________, 1.36 a 3.5 para valores medios y para los valores 3.5 en adelante se son aquellos considerados de alta diversidad; para el índice de Simpson resulta una diversidad ________________________________________________________________, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad muestra _______________________ que los valores de medida considerados para una baja biodoversidad son para aquellos valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores a 5, la espcie representativa para este gruo fue ______________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.3 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.70.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para el grupo de los reptiles en el área del ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.2 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de diversidad para el grupo de los mamíferos de las ____ especies presentes en el área del ACUSTF presenta un índice de ____, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de ________, para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del ACUSTF en estudio nos indica que la especie más representativa fue _______________________, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie, la especie más representativa fue _______________________ tal como se muestra en la siguiente gráfica, cabe señalar que para los quirópteros estos no se tiene indicios o registro durante los recorridos de los sitios de muestreo.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.8.3.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(12.24), height=Cm(8.18))

    #########################
    ### Titulo de la grafica del capitulo 5.8.3.3 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.28.- Frecuencia y abundancia relativa del grupo de los reptiles en el área CUSTF.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.3 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.3 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El grupo de los reptiles posee una riqueza específica de __ especies las cuales tienen una distribución de _______ con lo cual se puede afirmar que la equitatividad de especies es _____. La máxima diversidad que este grupo adquiere dentro del sistema ambiental es de ___________________ para la diversidad calculada lo que quiere decir que este grupo se encuentra muy cerca de llegar a su máxima diversidad y pose una distribución equitativa, lo que equivale a que este grupo se tenga un desplazamiento del ____% de las especies que se encuentren en el área hacia el área del sistema ambiental.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo con los datos que anteceden por las caracteristicas del area Acustf el grupo de los reptiles se presenta en condiciones de _______________ en cuanto a riqueza y equidad de especies, para la dominancia de especies se presenta de una _______________, ______________________________________________________, acontinuacion se muestra en el cuadro de rangos de valor para el grupo de los reptiles en el area Acustf.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.8.3.3 ###
    #########################
    tabla5 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.8.3.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3.4 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.3.4.- Análisis de Información del grupo de los lepidópteros en el ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.3.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Para el análisis de la información del grupo de los lepidópteros así como se otras especies de invertebrados que se observaron en el área del sistema ambiental se plasma la siguiente información en la cual se muestra el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra  el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al sistema ambiental, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.71.- Listado de especies de lepidópteros y otras especies de insectos observados en el ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.72.- Listado de especies de lepidópteros y otros insectos con su estatus de categoría por especies del sistema ambiental.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Sociabilidad (SOCI.); abundancia (ABU.); residencia (RES.); alimentación (ALIM.) y el tipo de observación (OBS.); Sc: Sociabilidad, R: Residente; C: Común, SL: Solitario, GR: Gregario, PJ: Pareja; Sc: Sin categoría, Pr: Sujeta a protección especial; A: Amenazada; P: En peligro de extinción; E: Extinta en medio silvestre.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(10)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.73.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef para las especies de lepidópteros y otros insectos observados en el ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Los índices de diversidad de las ____ especies de insectos presentes en el área CUSTF muestran que para el índice de Shannon  tenemos una equidad ___________ lo cual quiere decir que para este grupo los valores resultantes se encuentran a niveles _______ ya que los rangos de valores para este índice van de 0 a 1.35 para valores bajos, 1.36 a 3.5 para valores medios y para los valores 3.5 en adelante se son aquellos considerados de alta diversidad; para el índice de Simpson resulta una diversidad _____________ y una dominancia de las especies _____________, lo cual quiere decir que _____________________________, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad muestra ________________________ ya que los valores de medida considerados para una baja biodiversidad son para valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores 5; ______________________________________________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 5.8.3.4 ###
    #########################
    tituloTabla5 = doc.add_paragraph()
    dti5 = tituloTabla5.add_run('\nTabla 5.74.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para las especies de lepidópteros y otros insectos observados en el área del ACUSTF.')
    dti5_format = tituloTabla5.paragraph_format
    dti5_format.line_spacing = 1.15
    dti5_format.space_after = 0

    dti5.font.name = 'Bookman Old Style'
    dti5.font.size = Pt(12)
    tituloTabla5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=7, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('El índice de diversidad para el grupo de los insectos conformados por ____ especies presentes en el ACUSTF presenta un índice __________, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza ____________, para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área CUSTF en estudio nos indica que la especie más representativa es ________________, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie lo cual para este índice resulta como especie más representativa ____________________________ de representatividad, tal como se puede observar en las siguiente gráfica, Frecuencia y abundancia relativa de los reptiles en el Sistema ambiental.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 5.8.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo5/grafico.jpg', width=Cm(12.24), height=Cm(8.18))

    #########################
    ### Titulo de la grafica del capitulo 5.8.3.4 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 5.29.- Frecuencia y abundancia relativa del grupo de los lepidópteros en el área CUSTF.')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 5.8.3.4 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('En el grupo de insectos dentro del área del sistema ambiental posee una riqueza específica de ___ especies, con una distribución de __________ lo que equivale a que la equidad en las especies __________, la máxima diversidad que se puede alcanzar en este grupo es de _________ y la diversidad calculada es de _______________________________________________________________________________________, la especie más representativa fue ________________________________, considerando que el grupo tendrá un desplazamiento del ________% hacia el área del sistema ambiental.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('De acuerdo con los datos que anteceden por las caracteristicas del area del sistema ambiental el grupo de los insectos se presenta en condiciones de __________________ para la riqueza y dominancia de las especies, sin embargo para la equidad de especies los valores resultantes fueron _________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 5.8.3.4 ###
    #########################
    tabla5 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla5.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla5.cell(rows, cols)
            t5 = cell.paragraphs[0].add_run(' ')
            t5.font.size = Pt(12)
            t5.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 5.8.3.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 5.8.3.5 ###
    #########################
    capitulo5 = doc.add_paragraph()
    i5 = capitulo5.add_run(f'\nV.8.3.5.- Análisis de la información de la fauna en el ACUSTF.')
    i5_format = capitulo5.paragraph_format
    i5_format.line_spacing = 1.15

    i5.font.name = 'Arial'
    i5.font.size = Pt(12)
    i5.font.bold = True
    capitulo5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 5.8.3.5 ###
    #########################
    di5 = doc.add_paragraph()
    descripcionCapitulo5 = di5.add_run('Descripcion del capitulo')
    descripcionCapitulo5_format = di5.paragraph_format
    descripcionCapitulo5_format.line_spacing = 1.15
    descripcionCapitulo5_format.space_after = 0
    descripcionCapitulo5_format.space_before = 0

    descripcionCapitulo5.font.name = 'Arial'
    descripcionCapitulo5.font.size = Pt(12)
    di5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO X DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO X DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO X DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 5 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo5() # Crear el documento