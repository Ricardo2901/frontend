"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes, celdas, etc; en pulgadas y en centrimetros
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo16():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 16
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XVI.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 16
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'XV.-	Datos de Inscripción en el Registro del Prestador de Servicios Forestales que Haya Elaborado el Estudio y, del Que Estará a Cargo de la Ejecución del Cambio De Uso De Suelo.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'XVI.1.- Programa de Ordenamiento Ecológico del Territorio (POET)')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El Ordenamiento ecológico se define, jurídicamente, como \"El instrumento de política ambiental cuyo objeto es regular o inducir el uso del suelo y las actividades productivas, con el fin de lograr la protección del medio ambiente y la preservación y el aprovechamiento sustentable de los recursos naturales, a partir del análisis de las tendencias de deterioro y las potencialidades de aprovechamiento de los mismos\" (Ley General del Equilibrio Ecológico y la Protección al Ambiente, Titulo Primero, Art.3 fracción XXIV). Con lo que se establece un marco básico de gestión integral del territorio y sus recursos, siendo además una herramienta estratégica para la convergencia entre Estado y Sociedad."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "La Ley General del Equilibrio Ecológico y la Protección al Ambiente (LGEEPA) establece que el Ordenamiento ecológico es un instrumento que se deberá incorporar en la planeación nacional del desarrollo (Artículo 17). Señala, además, cuáles son los criterios que deben considerarse para la formulación de este (Artículo 19), cuáles son sus modalidades (Artículo 19 Bis), y describe cuáles son las instancias y los órdenes de gobierno a quienes corresponde la formulación de las diferentes modalidades del Ordenamiento Ecológico, lo mismo que los alcances de dichos programas (Artículos 20 al 20 Bis 7)."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "En el Reglamento de la LGEEPA, en materia de Ordenamiento ecológico, se definen las competencias de la SEMARNAT, así como la participación de las dependencias y entidades de la Administración Pública Federal en la formulación, expedición, ejecución, asesoría, evaluación, validación y vigilancia de los ordenamientos ecológicos de competencia federal; la participación en la formulación de los programas de Ordenamiento ecológico regional de interés de la Federación y en la participación en la elaboración y en su caso, la aprobación de los programas de Ordenamiento ecológico local. Finalmente, cada Entidad Federativa tiene atribuciones particulares en materia de Ordenamiento ecológico, establecidas en su respectiva legislación local."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El ordenamiento ecológico se puede formular según las modalidades: a) Programa de Ordenamiento Ecológico General del Territorio; b) Programa de Ordenamiento Ecológico Marino; c) Programa de Ordenamiento Ecológico Regional. Para fines del presente proyecto nos enfocaremos a la primera."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Programa de Ordenamiento Ecológico General del Territorio. Tiene como objetivo fundamental, vincular las acciones y programas de la Administración Pública Federal cuyas actividades inciden en el patrón de ocupación del territorio; en particular, se puede destacar que con el Programa se busca llevar a cabo la regionalización ecológica del territorio nacional y de las zonas sobre las que la nación ejerce su soberanía y jurisdicción, identificando áreas de atención prioritaria. Su formulación deberá atender a lo establecido en los artículos 20 y 20 bis de la LGEEPA y el capítulo tercero de su Reglamento."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            1.	Regionalización Ecológica.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n1.	Regionalización Ecológica.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Primer párrafo
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "La base para la regionalización ecológica, comprende unidades territoriales sintéticas que se integran a partir de los principales factores del medio biofísico: clima, relieve, vegetación y suelo. "
        "La interacción de estos factores determina la homogeneidad relativa del territorio hacia el interior de cada unidad y la heterogeneidad con el resto de las unidades. Con este principio se obtuvo como resultado "
        "la diferenciación del territorio nacional en 145 unidades denominadas Unidades Ambientales Biofísicas (UAB), representadas a escala 1:2, 000,000, empleadas como base para el análisis de las etapas de diagnóstico "
        "y pronóstico, y para construir la propuesta del POEGT. Así, las regiones ecológicas se integran por un conjunto de UAB que comparten la misma prioridad de atención, de aptitud sectorial y de política ambiental. "
        "Con base en lo anterior, a cada UAB le fueron asignados lineamientos y estrategias ecológicas específicas, de la misma manera que ocurre con las Unidades de Gestión Ambiental (UGA) previstas en los Programas "
        "de Ordenamiento Ecológico Regionales y Locales. Cabe señalar que, aun cuando las UAB y las UGA comparten el objetivo de orientar la toma de decisiones sobre la ubicación de las actividades productivas y los "
        "asentamientos humanos en el territorio, así como fomentar el mantenimiento de los bienes y servicios ambientales; dichas Unidades difieren en el proceso de construcción, toda vez que las UGA se construyen "
        "originalmente como unidades de síntesis que concentran, en su caso, lineamientos, criterios y estrategias ecológicas, en tanto que las UAB, considerando la extensión y complejidad del territorio sujeto a "
        "ordenamiento, se construyeron en la etapa de diagnóstico como unidades de análisis, mismas que fueron empleadas en la etapa de propuesta, como unidades de síntesis para concentrar lineamientos y estrategias "
        "ecológicas aplicables en dichas Unidades y, por ende, a las regiones ecológicas de las que formen parte."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo párrafo
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Las áreas de atención prioritaria de un territorio, son aquellas donde se presentan o se puedan potencialmente presentar, conflictos ambientales o que por sus características ambientales requieren de atención "
        "inmediata para su preservación, conservación, protección, restauración o la mitigación de impactos ambientales adversos. El resultado del análisis de estos aspectos permitió aportar la información útil para "
        "generar un consenso en la forma como deben guiarse los sectores, de tal manera que se transite hacia el desarrollo sustentable. Se establecieron 5 niveles de prioridad que son: Muy alta, Alta, Media, Baja y Muy baja. "
        "Dentro de éstos el muy alto se aplicó a aquellas UAB que requieren de atención urgente porque su estado ambiental es crítico y porque presentan muy alto o alto nivel de conflicto ambiental, por otro lado, el nivel muy "
        "bajo se aplicó a las UAB que presentan un estado del medio ambiente estable a medianamente estable y conflictos ambientales de medio a muy bajo."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Tercer párrafo
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Las políticas ambientales (aprovechamiento, restauración, protección y preservación) son las disposiciones y medidas generales que coadyuvan al desarrollo sustentable. Su aplicación promueve que los sectores del "
        "Gobierno Federal actúen y contribuyan en cada UAB hacia este modelo de desarrollo. Como resultado de la combinación de las cuatro políticas ambientales principales, para este Programa se definieron 18 grupos, "
        "los cuales fueron tomados en consideración para las propuestas sectoriales y finalmente para establecer las estrategias y acciones ecológicas en función de la complejidad interior de la UAB, de su extensión territorial "
        "y de la escala. El orden en la construcción de la política ambiental refleja la importancia y rumbo de desarrollo que se desea inducir en cada UAB. Tomando como base la política ambiental asignada para cada una "
        "de las 145 UAB, los sectores rectores del desarrollo que resultaron de la definición de los niveles de corresponsabilidad sectorial, y la prioridad de atención que los diferentes sectores deberán considerar para el "
        "desarrollo sustentable del territorio nacional, se realizó una síntesis que dio como resultado las 80 regiones ecológicas, que finalmente se emplearon en la propuesta del POEGT."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            2.	Lineamientos y estrategias ecológicas. 
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n2.	Lineamientos y estrategias ecológicas. ')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        'Los 10 lineamientos ecológicos que se formularon para este Programa, mismos que reflejan el estado deseable de una región ecológica o unidad biofísica ambiental, se instrumentan a través de las directrices generales que en lo ambiental, social y económico se deberán promover para alcanzar el estado deseable del territorio nacional. Por su parte, las estrategias ecológicas, definidas como los objetivos específicos, las acciones, los proyectos, los programas y los responsables de su realización dirigidas al logro de los lineamientos ecológicos aplicables en el territorio nacional, fueron construidas a partir de los diagnósticos, objetivos y metas comprendidos en los programas sectoriales, emitidos respectivamente por las dependencias que integran el Grupo de Trabajo Intersecretarial. Las estrategias se implementarán a partir de una serie de acciones que cada uno de los sectores en coordinación con otros sectores deberán llevar a cabo, con base en lo establecido en sus programas sectoriales o el compromiso que asuman dentro del Grupo de Trabajo Intersecretarial para dar cumplimiento a los objetivos de este POEGT. En este sentido, se definieron tres grandes grupos de estrategias: las dirigidas a lograr la sustentabilidad ambiental del territorio, las dirigidas al mejoramiento del sistema social e infraestructura urbana y las dirigidas al fortalecimiento de la gestión y la coordinación institucional. Los lineamientos ecológicos a cumplir son los siguientes:'
        )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Lista de objetivos
    objetivos = [
        "Proteger y usar responsablemente el patrimonio natural y cultural del territorio, consolidando la aplicación y el cumplimiento de la normatividad en materia ambiental, desarrollo rural y ordenamiento ecológico del territorio.",
        "Mejorar la planeación y coordinación existente entre las distintas instancias y sectores económicos que intervienen en la instrumentación del programa de ordenamiento ecológico general del territorio, con la activa participación de la sociedad en las acciones en esta área.",
        "Contar con una población con conciencia ambiental y responsable del uso sustentable del territorio, fomentando la educación ambiental a través de los medios de comunicación y sistemas de educación y salud.",
        "Contar con mecanismos de coordinación y responsabilidad compartida entre los diferentes niveles de gobierno para la protección, conservación y restauración del capital natural.",
        "Preservar la Flora y la Fauna, tanto en su espacio terrestre como en los sistemas hídricos a través de las acciones coordinadas entre las instituciones y la sociedad civil.",
        "Promover la conservación de los recursos naturales y la biodiversidad, mediante formas de utilización y aprovechamiento sustentable que beneficien a los habitantes locales y eviten la disminución del capital natural.",
        "Brindar información actualizada y confiable para la toma de decisiones en la instrumentación del ordenamiento ecológico territorial y la planeación sectorial.",
        "Fomentar la coordinación intersectorial a fin de fortalecer y hacer más eficiente al sistema económico.",
        "Incorporar al SINAP las áreas prioritarias para la preservación, bajo esquemas de preservación y manejo sustentable.",
        "Reducir las tendencias de degradación ambiental, consideradas en el escenario tendencial del pronóstico, a través de la observación de las políticas del Ordenamiento Ecológico General del Territorio."
    ]

    # Bucle para agregarlos al documento
    for i, objetivo in enumerate(objetivos, start=1):
        di16 = doc.add_paragraph()
        descripcionCapitulo16 = di16.add_run(f"{i}. {objetivo}")
        descripcionCapitulo16_format = di16.paragraph_format
        descripcionCapitulo16_format.line_spacing = 1.15
        descripcionCapitulo16.font.name = 'Arial'
        descripcionCapitulo16.font.size = Pt(12)
        di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Respecto a este programa de Ordenamiento Ecológico General del Territorio (POEGT) con referencia al predio sujeto de estudio se describe la información acorde a la ubicación."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Al respecto se incluyen las estrategias y políticas ambientales con fundamento en el Programa de Ordenamiento Ecológico General del Territorio (POEGT) en base a lo siguiente:"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    articulos = [
        "ARTÍCULO PRIMERO. -Se expide el Programa de Ordenamiento Ecológico General del Territorio en términos del documento adjunto al presente Acuerdo.",
        "ARTÍCULO SEGUNDO. -En términos del Artículo 19 del Reglamento de la Ley General del Equilibrio Ecológico y la Protección al Ambiente en Materia de Ordenamiento Ecológico, el Programa de Ordenamiento Ecológico General del Territorio será de observancia obligatoria en todo el territorio nacional y vinculará las acciones y programas de la Administración Pública Federal y las entidades paraestatales en el marco del Sistema Nacional de Planeación Democrática.",
        "ARTÍCULO TERCERO. -De conformidad con el Artículo 34 del Reglamento de la Ley General del Equilibrio Ecológico y la Protección al Ambiente en Materia de Ordenamiento Ecológico, las Dependencias y Entidades de la Administración Pública Federal deberán observar el Programa de Ordenamiento Ecológico General del Territorio en sus programas operativos anuales, en sus proyectos de presupuestos de egresos y en sus programas de obra pública.",
        "ARTÍCULO CUARTO. -La Secretaría de Medio Ambiente y Recursos Naturales tendrá a su cargo la etapa de ejecución y evaluación del Programa de Ordenamiento Ecológico General del Territorio, de conformidad con las disposiciones aplicables de la Ley General del Equilibrio Ecológico y la Protección al Ambiente, así como del Reglamento de la Ley General del Equilibrio Ecológico y la Protección al Ambiente en Materia de Ordenamiento Ecológico."
    ]

    for articulo in articulos:
        di16 = doc.add_paragraph()
        descripcionCapitulo16 = di16.add_run()
        
        # Separar el título del artículo del resto del texto
        partes = articulo.split(" -", 1)
        titulo = partes[0]  # Ejemplo: ARTÍCULO PRIMERO.
        cuerpo = partes[1] if len(partes) > 1 else ""  # El resto del texto
        
        # Agregar el título en negritas
        run_titulo = di16.add_run(titulo + " -")
        run_titulo.bold = True
        
        # Agregar el resto del texto normal
        run_cuerpo = di16.add_run(" " + cuerpo)
        
        # Formato
        descripcionCapitulo16_format = di16.paragraph_format
        descripcionCapitulo16_format.line_spacing = 1.15
        di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Fuente para ambos runs
        run_titulo.font.name = 'Arial'
        run_titulo.font.size = Pt(12)
        run_cuerpo.font.name = 'Arial'
        run_cuerpo.font.size = Pt(12)

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Así, las regiones ecológicas se integran por un conjunto de UAB que comparten la misma prioridad de atención, de aptitud sectorial y de política ambiental. Con base en lo anterior, a cada UAB le fueron asignados lineamientos y estrategias ecológicas específicas, de la misma manera que ocurre con las Unidades de Gestión Ambiental (UGA) previstas en los Programas de Ordenamiento Ecológico Regionales y Locales.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('A continuación, se hace la vinculación del proyecto con las estrategias y políticas ambientales con fundamento en el Programa de Ordenamiento Ecológico General del Territorio (POEGT) con base a las Regiones Ecológicas: __________________________________________________________. En las que se hará énfasis en las estrategias y acciones que contempla esta área. (Anexo Mapa 16-1 Ubicación en Poete Federal).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.1.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\n')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 16.1.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo16 = doc.add_paragraph()
    imagenCapitulo16.text = '\n'
    imagenCapitulo16 = doc.add_picture('capitulo16/mapa.png')  # Ancho de la imagen en centimetros
    imagenCapitulo16.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo16.width = Cm(17.09)  # Ancho de la imagen en centimetros
    imagenCapitulo16.height = Cm(11.06)  # Alto de la imagen en centimetros
    imagenCapitulo16.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del capitulo 16.1.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nREGION ECOLOGICA: _______'
                                         '\nUnidad Ambiental Biofisica que la compone:'
                                         '\n__________________________________________'
                                         '\nLocalización:'
                                         '\n__________________________________________')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Superficie en km\u2082',
        'Población por UAB',
        'Poblacion Indigena',
    ]

    filas = 2
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(len(encabezados)):
        cell = tabla16b.cell(0, cols)
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        #cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Estado Actual del Medio Ambiente 2008, para el área de estudio donde se encuentra el proyecto solo se menciona la Unidad Ambiental Biofísica',
    ]

    filas = 2
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(len(encabezados)):
        cell = tabla16b.cell(0, cols)
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run('\n\n\n')
            t16b.font.size = Pt(12)

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Escenario al 2033',
        'Política Ambiental',
        'Prioridad de Atención',
    ]

    datos_tabla = [
        ' ',
        'Aprovechamiento Sustentable y Restauración',
        'Muy Baja',
    ]

    filas = len(encabezados)
    columnas = 2
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for rows in range(len(encabezados)):
        cell = tabla16b.cell(rows, 0)
        t16b = cell.paragraphs[0].add_run(f'{encabezados[rows]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for rows in range(len(datos_tabla)):
        cell = tabla16b.cell(rows, 1)
        t16b = cell.paragraphs[0].add_run(f'{datos_tabla[rows]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 16.1.1 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\n')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'UAB',
        'Rectores del Desarrollo',
        'Coadyuvantes del desarrollo',
        'Asociados del Desarrollo',
        'Estrategias Sectoriales',
    ]

    filas = 2
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(len(encabezados)):
        cell = tabla16b.cell(0, cols)
        t16b = cell.paragraphs[0].add_run(encabezados[cols])
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(10)
        t16b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 16.1.1 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\nTabla XVI.1 Estrategias UAB _________________________')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Polítca',
        'Estrategia/Acciones',
        'Viculación',
    ]

    filas = 15
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    
    """
        *************************************
        Celda fusionada
        *************************************
    """
    # ✅ Celda fusionada
    row = tabla16b.rows[0]
    merged_cell = row.cells[1].merge(row.cells[1].merge(row.cells[2]))

    # Agregar texto a la celda fusionada
    t16b = merged_cell.paragraphs[0].add_run('Grupo x. Bla bla Bla bla bla bla')
    t16b.font.name = 'Arial'
    t16b.font.size = Pt(12)
    t16b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla16b.cell(0, 0)
    t16b = cell.paragraphs[0].add_run('Lineamiento')
    t16b.font.name = 'Arial'
    t16b.font.size = Pt(12)
    t16b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        cell = tabla16b.cell(1, cols)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 16.1.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Conclusión: el presente proyecto el cual tiene como objetivo principal la ________________, dicha área se ubica en las regiones ecológicas ____________________________, las cuales tiene indicadores de media degradación de suelo, baja degradación de vegetación y baja degradación por desertificación, así como, alta importancia de la actividad minera y ganadera Por ello, para para lograr el equilibrio ecológico y la sostenibilidad ambiental, el proyecto contempla la aplicación de un programa de rescate de flora y fauna para conservación de las biodiversidades, asi mismo se contempla la aplicación de las medidas de preventivas y de mitigación mediante el control de ruido, control y remediación en caso de algún derrame accidental al suelo tomando en cuenta el mantenimiento preventivo y correctivo de la maquinaria a usar. ________________________________________________________________________.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.1.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.1.2.- Programa de Ordenamiento Ecológico del Territorio del Estado de Coahuila.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.1.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El Programa de Ordenamiento Ecológico es un instrumento de política ambiental que promueve el aprovechamiento de los recursos naturales, sin hacer a un lado, la protección del medio ambiente y la preservación de los recursos naturales en la Planeación del desarrollo."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El 28 de noviembre del 2017 se acuerda por el cual se expide el Programa de Ordenamiento Ecológico Regional del Territorio del Estado de Coahuila de Zaragoza, conforme a la cual el gobierno estatal, los municipios y la comunidad en general, participaran en la planeación, ordenamiento y regulación de todas las acciones en materia de ordenamiento ecológico, así como la responsabilidad y alcances del mismo programa. Se presentó el decreto del Ordenamiento Ecológico del estado de Coahuila, así como, sus lineamientos, mapas de cada municipio con localización de las Unidades de Gestión Ambiental (UGA´s) y las fichas de las UGA´s, donde su realización debe entenderse como un proceso planificado de la naturaleza política, técnica y administrativa que plantea el análisis de un sistema socio espacial, conducente a organizar y administrar el uso y ocupación de ese espacio, en conformidad con las condiciones naturales y de los recursos naturales, la dinámica social, la estructura productiva, los asentamientos humanos y la infraestructura de servicios, para prever los efectos que provocan las actividades socioeconómicas en esa realidad espacial y establecer las acciones a ser instrumentadas con miras a que se cumplan los objetivos de bienestar social, manejo adecuado de las reservas naturales y calidad de vida, es decir, con miras al desarrollo sostenible."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Las Unidades de Gestión Ambiental (UGA): Son áreas del territorio relativamente homogéneas a las que se les asignan los lineamientos y las estrategias ecológicas. El estado deseable de cada UGA se refleja en la asignación de la política ambiental y el lineamiento ecológico que le corresponde. Debido a su extensión y complejidad territorial, el modelo de ordenamiento ecológico para el estado contiene 468 tipos diferentes de UGA."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Para el presente proyecto de “_______________________”, de acuerdo con dicho programa se ubica dentro de las Unidades de Gestión Ambiental (UGA). (Ver anexo Mapa 16.2.- ubicación de la UGA)."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.1.2 ###
    #########################
    encabezados = [
        'UGA',
        'No',
        'Superficie',
        'Ubicación',
        'Compatible',
    ]

    filas = 5
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, '0070C0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run('  ')
            t16b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 16.1.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n___________')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Descripcion xxxxxxxx')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 16.1.1 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\nTabla XVI.3. Criterios de regulación ecológica _________________________')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Clave',
        'Estrategias',
        'Aplicacion de los criterios',
    ]

    filas = 25
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, "#0095FF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 16.1.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n___________')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Descripcion xxxxxxxx')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 16.1.1 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\nTabla XVI.4. Criterios de regulación ecológica _________________________')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.1 ###
    #########################
    encabezados = [
        'Clave',
        'Estrategias',
        'Aplicacion de los criterios',
    ]

    filas = 25
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, "#0095FF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 16.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.1.3 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.1.3.-Programa de Ordenamiento Ecológico de la Región Cuenca de Burgos del Estado de Coahuila.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.1.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('El Programa de Ordenamiento Ecológico de la Región Cuenca de Burgos, el cual es de carácter regional, conforme a la fracción II del Artículo 19 Bis de la Ley General del Equilibrio Ecológico y la Protección al Ambiente. El Programa Regional de Ordenamiento Ecológico "Cuenca de Burgos" fue formulado por la Federación, por conducto de la Secretaría de Medio Ambiente y Recursos Naturales, por los Gobiernos de los Estados y de los Municipios que más adelante se señalan, de conformidad con los convenios de coordinación celebrados al efecto y con fundamento en los Artículos 20 BIS 1 y 20 BIS 2 de la Ley General del Equilibrio Ecológico y la Protección al Ambiente.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Para la formulación de este ordenamiento ecológico, se redefinió la Región Cuenca de Burgos, tomando como punto de partida el criterio de cuenca e identificando las principales cuencas con influencia en la Cuenca Gasífera de Burgos. De esta manera, el área que abarca este ordenamiento ecológico involucra a las 7 cuencas más importantes, de acuerdo con la regionalización hidrológica de la Comisión Nacional del Agua. Estas son: Presa Falcón-Río Salado, Río Bravo-Matamoros-Reynosa, Río Bravo-Nuevo Laredo, Río Bravo-San Juan, Río Bravo-Sosa, Río San Fernando y Laguna Madre. Administrativamente, esta área involucra en su totalidad la superficie de 31 municipios del Estado de Coahuila, 48 de Nuevo León y 19 de Tamaulipas, lo que da como resultado una superficie total de 208,805 km\u2082. ')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 16.1.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo16 = doc.add_paragraph()
    imagenCapitulo16.text = '\n'
    imagenCapitulo16 = doc.add_picture('capitulo16/mapa_1.png')  # Ancho de la imagen en centimetros
    imagenCapitulo16.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo16.width = Cm(13.25)  # Ancho de la imagen en centimetros
    imagenCapitulo16.height = Cm(10.19)  # Alto de la imagen en centimetros
    imagenCapitulo16.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 16.1.3 ###
    #########################
    diMap16 = doc.add_paragraph()
    descripcionCapituloMapa16 = diMap16.add_run('Mapa. - La Región Cuenca de Burgos considerada para el Ordenamiento Ecológico.')
    descripcionCapituloMapa16_format = diMap16.paragraph_format
    descripcionCapituloMapa16_format.line_spacing = 1.15
    descripcionCapituloMapa16.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa16.font.name = 'Arial'
    descripcionCapituloMapa16.font.size = Pt(12)
    descripcionCapituloMapa16.font.italic = True
    diMap16.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    #########################
    ### Descripcion del capitulo 16.1.3 ###
    #########################

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Municipios que componen la Región Cuenca de Burgos en el Estado Coahuila.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nMunicipios que componen la Región Cuenca de Burgos en el Estado Coahuila.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Abasolo, Frontera, Morelos, Sabinas, Acuña, General Cepeda, Múzquiz, Sacramento, Allende, Guerrero, Nadadores, Saltillo, Arteaga, Hidalgo, Nava, San Buenaventura, Candela, Jiménez, Ocampo, San Juan de Sabinas, Castaños, Juárez Piedras Negras, Villa Unión, Cuatrociénegas, Lamadrid, Progreso, Zaragoza.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nLa Región Cuenca de Burgos es poseedora de enormes recursos naturales no renovables y renovables, como es el caso de las reservas de gas natural, una rica y variada vida silvestre y recursos pesqueros. Dentro de la región se ubica la Laguna Madre, considerada como una zona de gran valor, por ser hábitat natural y de reproducción de varias especies de aves residentes y migratorias, así como de algunas especies marinas. De igual importancia están las poblaciones de Fauna cinegética localizadas dentro del matorral espinoso tamaulipeco. Los tipos de vegetación más representativos en la Región Cuenca de Burgos son el matorral espinoso tamaulipeco, el mezquital, el pastizal y la vegetación halófila, que resultan ser más abundantes en la Planicie Costera del Golfo, región fisiográfica donde se localiza esta región.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('El Programa de Ordenamiento Ecológico de la Región Cuenca de Burgos es un instrumento de política ambiental que promueve el aprovechamiento de los recursos naturales, sin hacer a un lado, la protección del medio ambiente y la preservación de los recursos naturales en la Planeación del desarrollo. Su objetivo es inducir el desarrollo de las actividades productivas en la región, siempre considerando la conservación y protección de los recursos naturales. De esta manera, este ordenamiento ecológico pretende ser el instrumento que le permita al Gobierno Federal, Estatal y Municipal hacer una mayor y mejor gestión de los recursos naturales en beneficio de la sociedad y del medio ambiente.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Programa de ordenamiento ecológico de la región cuenca de burgos.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Programa de ordenamiento ecológico de la región cuenca de burgos.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('De acuerdo con el Reglamento en materia de Ordenamiento Ecológico de la Ley General del Equilibrio Ecológico y la Protección al Ambiente (DOF agosto 2003), un programa de ordenamiento ecológico debe contener un modelo de ordenamiento ecológico y las estrategias ecológicas aplicables. A su vez, el modelo de ordenamiento ecológico contiene la regionalización o la determinación de las zonas ecológicas, según corresponda, y los lineamientos ecológicos aplicables. Por su parte, las estrategias ecológicas son el resultado de la integración de objetivos específicos, acciones, proyectos, programas y responsables de su realización y están dirigidas al logro de los lineamientos ecológicos aplicables.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Las Unidades de Gestión Ambiental (UGA).
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Las Unidades de Gestión Ambiental (UGA).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Son áreas del territorio relativamente homogéneas a las que se les asignan los lineamientos y las estrategias ecológicas. El Estado deseable de cada UGA se refleja en la asignación de la política ambiental y el lineamiento ecológico que le corresponde.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Debido a su extensión y complejidad territorial, el modelo de ordenamiento ecológico para la Región Cuenca de Burgos en el Estado contiene 398 tipos diferentes de UGA. ')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 16.1.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo16_parrafo = doc.add_paragraph()
    imagenCapitulo16_run = imagenCapitulo16_parrafo.add_run('\n')
    imagenCapitulo16_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo16_run.add_picture('capitulo16/mapa_2.png', width=Cm(10.53), height=Cm(8.25))

    #########################
    ### Descripcion del Mapa del capitulo 16.1.3 ###
    #########################
    diMap16 = doc.add_paragraph()
    descripcionCapituloMapa16 = diMap16.add_run('Mapa. - Modelo de Ordenamiento Ecológico con las Unidades de Gestión Ambiental por Política Ambiental.')
    descripcionCapituloMapa16_format = diMap16.paragraph_format
    descripcionCapituloMapa16_format.line_spacing = 1.15
    descripcionCapituloMapa16.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa16.font.name = 'Arial'
    descripcionCapituloMapa16.font.size = Pt(12)
    descripcionCapituloMapa16.font.italic = True
    diMap16.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    #########################
    ### Descripcion del capitulo 16.1.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nMás adelante se hace la vinculación con las Unidades de Gestión Ambiental y Estrategias Ecológicas aplicables al Proyecto denominado ________________________, en el municipio de ___________, Estado de Coahuila.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Las Unidades de Gestión Ambiental y Estrategias Ecológicas aplicables al Proyecto se muestran en la tabla siguiente, (ver anexo Mapa 16-3 Ubicación en Cuenca de Burgos).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.1.3 ###
    #########################
    filas = 3
    columnas = 2
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, "#BFBFBF")

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 16.1.3 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\nTabla XVI.6. Criterios de regulación ecológica Cuenca de Burgos UGA _________________________')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.1.3 ###
    #########################
    encabezados = [
        'Clave',
        'Estrategias',
        'Aplicacion de los criterios',
    ]

    filas = 100
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, "#0095FF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 16.1.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nDado que el área del proyecto se encuentra inmersa en las UGA de ______________________________________________________, con los lineamientos de: ___________________________________________________________________. para ello, considerando que el objetivo es el ____________________________, se desarrollará con la aplicación de las medidas preventivas (control de polvo y ruido) y de mitigación (extracción de suelo contaminado en caso de algún derrame), realizando previamente actividades de rescate y reubicación de especies de flora y fauna silvestre (especies aptas a rescate de lento crecimiento y difícil regeneración, así como, las que se encuentran en la NOM-059-SEMARNAT-2010). Dichas actividades no pondrán en riesgo los acuíferos y se utilizará la infraestructura necesaria para contrarrestar los efectos del cambio climático, mismo que se hará uso medido en todas las etapas del proyecto el recurso agua. Además, se realizará bajo los lineamientos Ordenamiento Territorial y Desarrollo Urbano del Estado de Coahuila de Zaragoza.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.2.- Planes o Programas de Desarrollo Urbano')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.2.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.2.1.- ')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.2.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "A continuación, se describe lo establecido en el Plan Nacional de Desarrollo del periodo 2019-2024, de acuerdo con lo publicado en el Diario Oficial de la Federación con fecha 12 de julio de 2019."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "De acuerdo con la Constitución del Estado mexicano, y con el fin de ver los procedimientos de participación y consulta popular en el sistema nacional de planeación democrática, y los criterios para la formulación, instrumentación, control y evaluación del plan y los programas de desarrollo\", el Plan Nacional de Desarrollo (PND) es, un instrumento para enunciar los problemas nacionales y enumerar las soluciones en una proyección sexenal."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Por lo anterior, los lineamientos en los que se enmarca el Plan Nacional de Desarrollo 2019-2024 se basa en los siguientes principios rectores:"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Honradez y honestidad, No al gobierno rico con pueblo pobre, Al margen de la ley, nada; por encima de la ley, nadie, Economía para el bienestar, El mercado no sustituye al Estado, Por el bien de todos, primero los pobres, No dejar a nadie atrás, no dejar a nadie fuera, No puede haber paz sin justicia, El respeto al derecho ajeno es la paz, No más migración por hambre o por violencia, Democracia significa el poder del pueblo, Ética, libertad, confianza."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Visión
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nVisión')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Construir un México justo, pacífico, libre, solidario, democrático, próspero y feliz. Para ello, la Presidencia de la República y el gobierno federal en su conjunto trabajarán sin descanso para articular los esfuerzos sociales para lograr ese objetivo. El Ejecutivo Federal tiene ante sí la responsabilidad de operar una transformación mayor en el aparato administrativo y de reorientar las políticas públicas, las prioridades gubernamentales y los presupuestos para ser el eje rector de la Cuarta Transformación. El fortalecimiento de los principios éticos irá acompañado de un desarrollo económico que habrá alcanzado para entonces una tasa de crecimiento de 6 por ciento, con un promedio sexenal de 4 por ciento. La economía deberá haber crecido para entonces más del doble que el crecimiento demográfico. De tal manera, en 2024 el país habrá alcanzado el objetivo de crear empleos suficientes para absorber la demanda de los jóvenes que se estén incorporando al mercado laboral. Los programas de creación de empleos y de becas para los jóvenes habrán surtido su efecto y el desempleo será mínimo; la nación contará con una fuerza laboral mejor capacitada y con un mayor grado de especialización."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "En 2021 deberá cumplirse la meta de alcanzar la autosuficiencia en maíz y frijol y tres años más tarde, en arroz, carne de res, cerdo, aves y huevos; las importaciones de leche habrán disminuido considerablemente, la producción agropecuaria en general habrá alcanzado niveles históricos y la balanza comercial del sector dejará de ser deficitaria. Se habrá garantizado la preservación integral de la flora y de la fauna, se habrá reforestado buena parte del territorio nacional y ríos, arroyos y lagunas estarán recuperados y saneados; el tratamiento de aguas negras y el manejo adecuado de los desechos serán prácticas generalizadas en el territorio nacional y se habrá expandido en la sociedad la conciencia ambiental y la convicción del cuidado del entorno."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "En el último año del sexenio habrá cesado la emigración de mexicanos al exterior por causas de necesidad laboral, inseguridad y falta de perspectivas, la población crecerá de manera mejor distribuida en el territorio nacional y millones de mexicanas y mexicanos encontrarán bienestar, trabajo y horizontes de realización personal en sus sitios de origen, desarrollando su vida al lado de sus familias, arraigados en sus entornos culturales y ambientales. La delincuencia organizada estará reducida y en retirada, los índices delictivos de homicidios dolosos, secuestros, robo de vehículos, robo a casa habitación, asalto en las calles y en el transporte público y otros se habrán reducido en 50 por ciento en comparación con los de 2018 y México habrá dejado de ser la dolorosa y vergonzosa referencia internacional como tierra de violencia, desaparecidos y violaciones a los derechos humanos."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "La delincuencia de cuello blanco habrá desaparecido y la corrupción política y la impunidad que han prevalecido como norma hasta 2018 habrán quedado reducidas a casos excepcionales, individuales e inmediatamente investigados y sancionados. Las instituciones estarán al servicio de las necesidades del pueblo y de los intereses nacionales, el principio de la separación de poderes y el respeto al pacto federal serán la norma y no la excepción, el acatamiento de las leyes regirá el comportamiento de los servidores públicos y el fraude electoral, la compra de voto y todas las formas de adulteración de la voluntad popular serán sólo un recuerdo. En los procesos electorales que se realicen en el curso del presente sexenio habrá quedado demostrado con hechos que es posible, deseable y obligatorio respetar el sufragio, hacer cumplir la legalidad democrática y sancionar las prácticas fraudulentas. Se habrán incorporado a la vida pública del país las distintas prácticas de la democracia participativa y el principio del gobierno del pueblo y para el pueblo será una realidad. En el último año del presente sexenio, en suma, el país habrá llevado a cabo lo sustancial de su cuarta transformación histórica, tanto en el ámbito económico, social y político, como en el de la ética para la convivencia: se habrá consumado la revolución de las conciencias y la aplicación de sus principios honradez, respeto a la legalidad y a la veracidad, solidaridad con los semejantes, preservación de la paz será la principal garantía para impedir un retorno de la corrupción, la simulación, la opresión, la discriminación y el predomino del lucro sobre la dignidad."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Para esto, se basará en tres Ejes rectores para el desarrollo del país:"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Lista de ejes rectores
    ejes_rectores = [
        "Política y Gobierno",
        "Política social",
        "Economía"
    ]

    for eje in ejes_rectores:
        di16 = doc.add_paragraph()
        descripcionCapitulo16 = di16.add_run(f" {ejes_rectores.index(eje)+1}. {eje}")
        descripcionCapitulo16_format = di16.paragraph_format
        descripcionCapitulo16_format.line_spacing = 1.15
        descripcionCapitulo16.font.name = 'Arial'
        descripcionCapitulo16.font.size = Pt(12)
        di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "A continuación, se hace la vinculación de las actividades que conlleva el desarrollo del proyecto con los ejes rectores y estrategias del plan nacional de desarrollo."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Política Social
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\n2.    Política Social')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo16 = di16.add_run('Desarrollo sostenible')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('El gobierno de México está comprometido a impulsar el desarrollo sostenible, que en la época presente se ha evidenciado como un factor indispensable del bienestar. Se le define como la satisfacción de las necesidades de la generación presente sin comprometer la capacidad de las generaciones futuras para satisfacer sus propias necesidades. Esta fórmula resume insoslayables mandatos éticos, sociales, ambientales y económicos que deben ser aplicados en el presente para garantizar un futuro mínimamente habitable y armónico. El hacer caso omiso de este paradigma no sólo conduce a la gestación de desequilibrios de toda suerte en el corto plazo, sino que conlleva una severa violación a los derechos de quienes no han nacido. Por ello, el Ejecutivo Federal considerará en toda circunstancia los impactos que tendrán sus políticas y programas en el tejido social, en la ecología y en los horizontes políticos y económicos del país. Además, se guiará por una idea de desarrollo que subsane las injusticias sociales e impulse el crecimiento económico sin provocar afectaciones a la convivencia pacífica, a los lazos de solidaridad, a la diversidad cultural ni al entorno.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nVinculación: El promovente llevará a cabo las actividades para el desarrollo del proyecto para lograr el objeto de poder aprovechar el área donde se pretende realizar la __________________, esto en apego a las normas ambientales vigentes, aplicando las medidas preventivas y de mitigación para minimizar efectos adversos al medio ambiente y así garantizar la salud de los habitantes de la región y con visión a la sostenibilidad ambiental.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.2.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.2.2.- Plan Estatal de Desarrollo 2017-2023.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.2.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "En cumplimiento a lo establecido en la Ley de Planeación para el Desarrollo del Estado de Coahuila de Zaragoza, y la definición de las prioridades con respecto a objetivos y las estrategias necesarias para alcanzarlos; con políticas públicas modernas, el uso ordenado y transparente de los recursos públicos, funcionarios públicos eficientes y la participación permanente de la sociedad en todos los ámbitos de gobierno, lograremos hacer que nuestro estado sea una referencia de seguridad y bienestar para todas las familias. Estas directrices surgen de la propia voz ciudadana, expresada a través de las demandas sociales, las contribuciones de académicos y expertos en los diversos ámbitos que son competencia del Estado, así como de personas de los distintos sectores de todas las regiones. Partimos de un diagnóstico preciso sobre la situación de la entidad al iniciar nuestra gestión; el enfoque adoptado consiste en identificar las capacidades y ventajas competitivas, así como las nuevas oportunidades de desarrollo que se prevén en un futuro inmediato."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Por otra parte, se analizaron los retos derivados tanto del entorno nacional como el internacional que habremos de enfrentar y superar. A partir de los elementos de información del Diagnóstico General, se establece una Visión de lo que el Gobierno del Estado desea para Coahuila hacia el término del sexenio. Con base en esta, se traza una Estrategia General de Desarrollo, dividida en cuatro ejes rectores, los cuales dieron nombre a los grandes capítulos del Plan; en cada uno se plantean objetivos generales y específicos, así como estrategias."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El plan estatal de desarrollo 2017-2023 fue elaborado mediante un largo proceso de consulta ciudadana, en la que participaron miles de coahuilenses en diferentes etapas. El proceso inició con la presentación de la Plataforma Política ante el Instituto Electoral de Coahuila, que fue un ejercicio de consulta pública a través de oros ciudadanos encabezados por especialistas en diversos temas relevantes para el diseño del proyecto de gobierno. Durante la campaña política, el entonces candidato a Gobernador del Estado enriqueció su oferta política a través de 21 foros de consulta, en los que se recibieron 2,313 propuestas que sirvieron como base para conformar el proyecto de este documento."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Coahuila es actualmente una de las entidades federativas más desarrolladas del país de acuerdo con indicadores como el ingreso per cápita la incidencia de la pobreza y la pobreza extrema el acceso a los servicios de salud y la cobertura de los distintos niveles educativos. Sin embargo, para que el estado mantenga su senda de crecimiento es preciso que se aprovechen las oportunidades y se atiendan las amenazas que se presentan desde fuera que se consoliden las fortalezas con que cuenta y se superen las debilidades que persisten internamente. Entre las fortalezas de Coahuila destaca su alta competitividad, entendida como la capacidad de atraer inversiones. Esto es muy importante puesto que significa que el estado se encuentra en una situación ventajosa respecto a una gran variedad de factores relacionados con la inclusión social, educación, salud, seguridad, sostenibilidad, entre otros. Entre las ramas de actividad más importantes del estado destacan la minería, siderurgia, agricultura, ganadería, así como las industrias automotriz y maquiladora de exportación."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Al concluir esta administración se habrán satisfecho las demandas más importantes de la sociedad en materia de combate a la impunidad y la corrupción. Su economía crecerá de manera sostenida por su grado de competitividad, así como por contar con un gobierno eficaz y moderno, por la seguridad pública, la calidad de su capital humano, la infraestructura estratégica, su desarrollo tecnológico y el aprovechamiento sostenible de sus recursos naturales. Este crecimiento le permitirá generar los empleos productivos que demanda la población, y que serán el mejor medio para erradicar la pobreza. Todos los habitantes del estado tendrán acceso equitativo a servicios públicos de calidad."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "La administración estatal tendrá como misión desempeñarse en forma honesta y transparente. Será austera, responsable y cuidadosa en el ejercicio de los recursos públicos. Los trabajadores al servicio del Estado ejercerán sus funciones escuchando a los ciudadanos y observando un código de ética y conducta que asegure la integridad del gobierno, el logro de sus objetivos y la calidad en la prestación de los servicios públicos. Adoptará un modelo de planeación estratégica que le permita a Coahuila consolidar su desarrollo mediante la focalización de esfuerzos de todas las dependencias en cuatro ejes rectores: Integridad y Buen Gobierno; Seguridad y Justicia; Desarrollo Económico Sustentable, y Desarrollo Social Incluyente y Participativo"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "El proyecto al cual se hace referencia este documento podrá vincular o estar relacionado a lo siguiente:"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Eje Rector 1. Integridad y buen gobierno.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nEje Rector 1. Integridad y buen gobierno.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 1
    di16_1 = doc.add_paragraph()
    descripcionCapitulo16_1 = di16_1.add_run(
        "Si bien en Coahuila se ha integrado el Sistema Estatal Anticorrupción, como lo requiere la reforma constitucional que dio vida al Sistema Nacional, "
        "y además se cuenta con una ley de transparencia de avanzada, es evidente que hay vulnerabilidades importantes en la administración pública estatal "
        "que llaman a abordar el tema de la corrupción con una perspectiva integral."
    )
    descripcionCapitulo16_format_1 = di16_1.paragraph_format
    descripcionCapitulo16_format_1.line_spacing = 1.15
    descripcionCapitulo16_1.font.name = 'Arial'
    descripcionCapitulo16_1.font.size = Pt(12)
    di16_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 2
    di16_2 = doc.add_paragraph()
    descripcionCapitulo16_2 = di16_2.add_run(
        "Lo anterior se hace evidente en la percepción de la ciudadanía. De acuerdo con la Encuesta Nacional de Calidad e Impacto Gubernamental (INEGI, 2015), "
        "44% de la población piensa que la corrupción es muy frecuente en el gobierno estatal, mientras que 38% piensa lo mismo para los gobiernos municipales. "
        "Reformas recientes a la Ley de Adquisiciones, Arrendamientos y Contratación de Servicios y a la Ley de Obras Públicas y Servicios Relacionados con las Mismas "
        "establecieron el Manifiesto de No Conflicto de Intereses y el Código de Conducta para proveedores potenciales, no obstante, es necesario también desarrollar "
        "protocolos para que los funcionarios públicos identifiquen y gestionen situaciones de dilemas éticos."
    )
    descripcionCapitulo16_format_2 = di16_2.paragraph_format
    descripcionCapitulo16_format_2.line_spacing = 1.15
    descripcionCapitulo16_2.font.name = 'Arial'
    descripcionCapitulo16_2.font.size = Pt(12)
    di16_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 3
    di16_3 = doc.add_paragraph()
    descripcionCapitulo16_3 = di16_3.add_run("Objetivo 1.2 Manejo de los recursos públicos y prevención de la corrupción")
    descripcionCapitulo16_format_3 = di16_3.paragraph_format
    descripcionCapitulo16_format_3.line_spacing = 1.15
    descripcionCapitulo16_3.font.name = 'Arial'
    descripcionCapitulo16_3.font.size = Pt(12)
    di16_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 4
    di16_4 = doc.add_paragraph()
    descripcionCapitulo16_4 = di16_4.add_run("Lograr un manejo responsable y honesto de los recursos públicos y prevenir la corrupción.")
    descripcionCapitulo16_format_4 = di16_4.paragraph_format
    descripcionCapitulo16_format_4.line_spacing = 1.15
    descripcionCapitulo16_4.font.name = 'Arial'
    descripcionCapitulo16_4.font.size = Pt(12)
    di16_4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 5
    di16_5 = doc.add_paragraph()
    descripcionCapitulo16_5 = di16_5.add_run("Estrategia 1.2.4 Fortalecer la coordinación con los órganos de fiscalización federales para combatir la corrupción.")
    descripcionCapitulo16_format_5 = di16_5.paragraph_format
    descripcionCapitulo16_format_5.line_spacing = 1.15
    descripcionCapitulo16_5.font.name = 'Arial'
    descripcionCapitulo16_5.font.size = Pt(12)
    di16_5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 6
    di16_6 = doc.add_paragraph()
    descripcionCapitulo16_6 = di16_6.add_run("Objetivo 1.7 servicios registrales")
    descripcionCapitulo16_format_6 = di16_6.paragraph_format
    descripcionCapitulo16_format_6.line_spacing = 1.15
    descripcionCapitulo16_6.font.name = 'Arial'
    descripcionCapitulo16_6.font.size = Pt(12)
    di16_6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 7
    di16_7 = doc.add_paragraph()
    descripcionCapitulo16_7 = di16_7.add_run("Modernizar los servicios registrales en materia civil, catastral, de la propiedad y el comercio para mejorar la calidad de los trámites.")
    descripcionCapitulo16_format_7 = di16_7.paragraph_format
    descripcionCapitulo16_format_7.line_spacing = 1.15
    descripcionCapitulo16_7.font.name = 'Arial'
    descripcionCapitulo16_7.font.size = Pt(12)
    di16_7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 8
    di16_8 = doc.add_paragraph()
    descripcionCapitulo16_8 = di16_8.add_run("Estrategia 1.7.1 Crear el Instituto Registral y Catastral de Coahuila, para asegurar la certeza patrimonial y responder de manera expedita a las solicitudes de trámites, tanto de individuos como de empresas.")
    descripcionCapitulo16_format_8 = di16_8.paragraph_format
    descripcionCapitulo16_format_8.line_spacing = 1.15
    descripcionCapitulo16_8.font.name = 'Arial'
    descripcionCapitulo16_8.font.size = Pt(12)
    di16_8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 9
    di16_9 = doc.add_paragraph()
    descripcionCapitulo16_9 = di16_9.add_run("Estrategia 1.7.2 Apoyar la modernización de los catastros municipales, a fin de fortalecer los ingresos propios y favorecer la certeza jurídica de los propietarios de predios.")
    descripcionCapitulo16_format_9 = di16_9.paragraph_format
    descripcionCapitulo16_format_9.line_spacing = 1.15
    descripcionCapitulo16_9.font.name = 'Arial'
    descripcionCapitulo16_9.font.size = Pt(12)
    di16_9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Eje rector 3. Desarrollo económico sustentable
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('\nEje rector 3. Desarrollo económico sustentable.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Coahuila se ubica como la séptima economía a nivel nacional por su contribución al producto interno bruto (PIB) con 3.7% y en el mismo lugar en competitividad. "
        "El PIB per cápita de cerca de 195 mil pesos por año, es el quinto más elevado en el país. El estado ocupa el segundo lugar nacional en participación de las exportaciones; "
        "en las cuales destaca el sector automotriz, cuya producción es la tercera más grande de México. Posee una economía relativamente diversificada, lo que reduce su vulnerabilidad respecto de las variaciones económicas externas e internas"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Respecto a la estructura de la producción, la industria es, por mucho, el sector más importante, pues representa poco más de la mitad del PIB estatal. "
        "Le siguen en importancia el de los servicios (34%) y el comercio (13%); el sector agropecuario representa 2% de la producción total. Por su parte, la estructura del empleo se distribuye de una manera un tanto distinta por diferencias en la productividad entre los sectores. "
        "Así, el sector terciario, que incluye comercio y servicios, absorbe 57% de la población ocupada, mientras que la industria emplea 40% y el sector agropecuario al restante 3%. "
        "Se observa una escasa integración de cadenas productivas en sectores de la economía, que ha limitado el desarrollo de proveedores locales, especialmente de micro, pequeñas y medianas empresas. "
        "En materia de infraestructura, los principales centros de población están bien conectados con vías carreteras y ferroviarias, con centros de población de otros estados y puertos fronterizos y marítimos. "
        "El 90% de los habitantes de la entidad vive en zonas urbanas, lo que facilita dotar de servicios básicos a la población; las coberturas de los servicios de agua potable, drenaje y electricidad son cercanas a 100%. "
        "Uno de los sectores con mayor potencial en la entidad es el turismo. Actualmente, la entidad ocupa el 16.º lugar nacional con mayor turismo internacional con más de 510 mil visitantes, y el 23.º lugar en turismo nacional con más de 3.6 millones de visitas. "
        "Se cuenta con seis pueblos mágicos: Arteaga, Candela, Cuatro Ciénegas, Guerrero, Parras y Viesca, sin embargo, hay oportunidades de desarrollo en varios municipios más para este segmento. "
        "Un reto importante es fortalecer la infraestructura logística, crear las condiciones de competitividad de las empresas del sector y promover a nivel nacional e internacional los lugares turísticos."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "En Coahuila, 99% de las empresas son micro, pequeñas y medianas (MiPyMEs). Estas desempeñan un papel importante en el proceso de transformación porque complementan las cadenas de valor; "
        "son flexibles para adoptar nuevos procesos productivos, tecnologías e innovaciones. Además, muchos de los servicios modernos pueden ser suministrados por este tipo de empresas. "
        "El fomento a las MiPyMEs representa un reto muy importante; es necesario consolidar una política que permita favorecer el desarrollo de proveedores, adoptar medidas para apoyar la organización de las empresas de sectores prioritarios para propiciar su crecimiento y mejorar su competitividad."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "La extensión del territorio, el clima, la topografía y variedad de ecosistemas, aunado a la vocación industrial, minera y energética, plantean un verdadero reto para consolidar a Coahuila como un estado sustentable, donde el derecho de sus habitantes a disfrutar de un medio ambiente sano sea garantizado. "
        "Vale la pena destacar que 17% de la superficie del estado se encuentra bajo un esquema de protección a la biodiversidad, tercero a nivel nacional en extensión de áreas naturales protegidas; de esta manera, cumplimos la meta federal en esta materia. "
        "Ocupa el segundo lugar nacional en superficie bajo manejo sustentable mediante Unidades de Manejo para la Conservación y el Aprovechamiento Sustentable de la Vida Silvestre, con 23.9% del territorio; esto representa una actividad relevante sobre todo para los municipios del norte. "
        "Se cuenta con un Sistema de Monitoreo de la Calidad del Aire de Coahuila conformado por cuatro estaciones, en Monclova, Piedras Negras, Saltillo y Torreón. El 67% de la población vive en ciudades donde opera este monitoreo; es decir, quinto lugar a nivel nacional."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Objetivo general: Orientar la estructura productiva hacia los sectores más competitivos, en un marco de crecimiento económico sostenido y de respeto a los derechos laborales y al medio ambiente."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Objetivo 3.10 Medio ambiente: "
        "Asegurar el derecho de los coahuilenses a un medio ambiente sano, mediante política pública que garanticen el uso sustentable de los recursos naturales, así como la regulación de actividades que impacten el medio ambiente."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.1 Elaborar la Estrategia Estatal de Biodiversidad, en la que se enmarquen los esfuerzo e iniciativa para la conservación y recuperación del capital natural del estado."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.2 Promover el uso sostenible de los recursos naturales, como un mecanismo que garantice la conservación de las especies, los ecosistemas y el paisaje, con una visión a largo plazo."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.3 Ampliar la cobertura de la disposición de residuos sólidos urbanos en rellenos sanitarios en cabeceras municipales y poblaciones de más de cinco mil habitantes."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.4 Impulsar una política integral de gestión del agua, con el propósito de ordenar la distribución y uso de este recurso; mejorar el manejo de la conservación y recuperación de las cuencas hídricas y los cuerpos de agua, especialmente ríos y arroyos, así como prevenir la contaminación."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.14 Poner en marcha una campaña permanente de cultura del cuidado del agua."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.15 Impulsar proyectos regionales de aprovechamiento ambiental, en materia de residuos y energía limpia."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Estrategia 3.10.17 Poner en marcha un programa permanente de reforestación de áreas verdes."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Vinculación. El presente proyecto se apegará a las estrategias mencionadas anteriormente en el plan estatal de desarrollo, a través de diferentes acciones apegándose a los planes de ordenamiento ecológico y a fortalecer los mecanismos cuyos propósitos estén enfocados a la protección ambiental, así mismo a la generación de empleos que garantice una mejor calidad de vida de los habitantes de esta zona."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.2.3 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.2.3.- ')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.2.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('Descripcion del capitulo')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('A continuación, se hace la vinculación del desarrollo del proyecto con algunos de los ejes temáticos mencionado anteriormente:')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.3.- Otros instrumentos a Considerar')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'XVI.5.1.- Constitución Política de los Estados Unidos Mexicanos')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.3.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el Diario Oficial de la Federación el 5 de febrero de 1917, con última reforma del día 28 de mayo de 2021).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.3.1 ###
    #########################
    encabezados = [
        'Artículos relacionados con el proyecto',
        'Vinculación',
    ]

    capitulo_seccion = [
        'Capitulo I. De los Derechos Humanos y sus Garantías',
        'Sección III. De las Facultades del Congreso'
    ]

    datos_tabla = [
        [
            'Artículo 4º. párrafo quinto: Toda persona tiene derecho a un medio ambiente sano para su desarrollo y bienestar. El Estado garantizará el respeto a este derecho. El daño y deterioro ambiental generará responsabilidad para quien lo provoque en términos de lo dispuesto por la ley.\n\nArtículo 27.- párrafo tercero: La nación tendrá en todo tiempo el derecho de imponer a la propiedad privada las modalidades que dicte el interés público, así como el de regular, en beneficio social, el aprovechamiento de los elementos naturales susceptibles de apropiación, con objeto de hacer una distribución equitativa de la riqueza pública, cuidar de su conservación, lograr el desarrollo equilibrado del país y el mejoramiento de las condiciones de vida de la población rural y urbana. En consecuencia, se dictarán las medidas necesarias para ordenar los asentamientos humanos y establecer adecuadas provisiones, usos, reservas y destinos de tierras, aguas y bosques, a efecto de ejecutar obras públicas y de planear y regular la fundación, conservación, mejoramiento y crecimiento de los centros de población; para preservar y restaurar el equilibrio ecológico; para el fraccionamiento de los latifundios; para disponer, en los términos de la ley reglamentaria, la organización y explotación colectiva de los ejidos y comunidades; para el desarrollo de la pequeña propiedad rural; para el fomento de la agricultura, de la ganadería, de la silvicultura y de las demás actividades económicas en el medio rural, y para evitar la destrucción de los elementos naturales y los daños que la propiedad pueda sufrir en perjuicio de la sociedad.',
            'El proyecto contempla aplicar las medidas preventivas y de mitigación en todas sus etapas de desarrollo para no alterar el medio ambiente y así garantizar el desarrollo y bienestar de las personas en la zona de influencia, buscando beneficiar a la población y sin causar daños irreversibles al medio ambiente.',
        ],
        [
            'Artículo 73.- fracción XXIX-G: Para expedir leyes que establezcan la concurrencia del Gobierno Federal, de los gobiernos de las entidades federativas, de los Municipios y, en su caso, de las demarcaciones territoriales de la Ciudad de México, en el ámbito de sus respectivas competencias, en materia de protección al ambiente y de preservación y restauración del equilibrio ecológico.',
            'La promovente se apegará a las normas y leyes ambientales con la aplicación de medidas preventivas para el desarrollo de las actividades del proyecto para conservar el equilibrio ecológico.'
        ]
    ]

    filas = 5
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla):
        celda = (i * 2) + 2

        for cols in range(2):
            cell = tabla16b.cell(celda, cols)
            t16b = cell.paragraphs[0].add_run(f'{rows[cols]}')
            t16b.font.size = Pt(12)
            t16b.font.name = 'Arial'
            t16b.bold = False

    # Rellenar las celdas fusionadas
    if tabla16b.cell(1, 0) and tabla16b.cell(3, 0):
        for i in range(2):
            celda_fusionada = (i * 2) + 1
            row = tabla16b.rows[celda_fusionada]
            merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
            cell_background_color(merged_cell, '#DBE5F1')

            # Agregar texto a la celda fusionada
            t16b = merged_cell.paragraphs[0].add_run(f'{capitulo_seccion[i]}')
            t16b.font.name = 'Arial'
            t16b.font.size = Pt(12)
            t16b.bold = True
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # Celdas de encabezados
    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#007AD1")

    ########################################################################################################################################################################
    # Capitulo 16.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.2.- ')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.3.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el DOF el 28 de enero de 1988). Ultima reforma publicada en el DOF 24-01-2024')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 16.3.2 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\n')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.3.2 ###
    #########################
    filas = 40
    columnas = 2
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 16.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3.3 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.3.- ')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.3.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el Periódico Oficial el 8-10-2003 con última reforma publicada en el P.O.  08-05-2023.')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 16.3.3 ###
    #########################
    tituloTabla16b = doc.add_paragraph()
    dti16b = tituloTabla16b.add_run('\n')
    dti16b_format = tituloTabla16b.paragraph_format
    dti16b_format.line_spacing = 1.15
    dti16b_format.space_after = 0

    dti16b.font.name = 'Bookman Old Style'
    dti16b.font.size = Pt(12)
    tituloTabla16b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.3.3 ###
    #########################
    filas = 40
    columnas = 2
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla16b.cell(rows, cols)
            t16b = cell.paragraphs[0].add_run(' ')
            t16b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 16.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3.4 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.4.- Ley General de Desarrollo Forestal Sustentable.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.3.4 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicada en el diario oficial de la federación, el 25 de febrero de 2003 con la última reforma publicada en el DOF 28-04-2022)')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.3.4 ###
    #########################
    encabezados = [
        'Artículos relacionados con el proyecto',
        'Vinculación',
    ]

    capitulo_seccion = [
            'TÍTULO PRIMERO'
            '\nDe las Disposiciones Generales'
            '\nCapítulo I Objeto y Aplicación de la Ley',
        
            'Título Cuarto '
            '\nDe los Procedimientos en Materia Forestal'
            '\nSección Séptima Del Cambio de Uso de Suelo en Terrenos Forestales'

    ]

    datos_tabla = [
        [
            'Artículo 1. La presente Ley es Reglamentaria del artículo 27 de la Constitución Política de los Estados Unidos Mexicanos, sus disposiciones son de orden e interés público y de observancia general en todo el territorio nacional, y tiene por objeto regular y fomentar el manejo integral y sustentable de los territorios forestales, la conservación, protección, restauración, producción, ordenación, el cultivo, manejo y aprovechamiento de los ecosistemas forestales del país y sus recursos; así como distribuir las competencias que en materia forestal correspondan a la Federación, las Entidades Federativas, Municipios y Demarcaciones Territoriales de la Ciudad de México, bajo el principio de concurrencia previsto en el artículo 73, fracción XXIX-G de la Constitución Política de los Estados Unidos Mexicanos, con el fin de propiciar el desarrollo forestal sustentable. Cuando se trate de recursos forestales cuya propiedad o legítima posesión corresponda a los pueblos y comunidades indígenas se observará lo dispuesto por el artículo 2o. de la Constitución Política de los Estados Unidos Mexicanos. \nARTICULO 3. Son objetivos específicos de esta Ley: \nX. Promover la conservación de los ecosistemas forestales, impulsando su delimitación y manejo sostenible, evitando que el cambio de uso de suelo con fines agropecuarios o de cualquier otra índole afecte su permanencia y potencialidad.',
            'La promovente ingresara el Documento Técnico Unificado para el cambio de uso de suelo en terrenos forestales para su evaluación y ver su viabilidad evitando dañar más área de la autorizada y aplicando las medidas de mitigación que le correspondan de acuerdo con la normatividad ambiental.',
        ],
        [
            'Artículo 7. Para los efectos de esta Ley se entenderá por: \nVI. Cambio de uso del suelo en terreno forestal: La remoción total o parcial de la vegetación de los terrenos forestales para destinarlos a actividades no forestales.',
            'La promovente tiene por entendido la definición del cambio de uso de suelo por lo cual realizará los trámites correspondientes para poder obtener una autorización para por realizar esta actividad de cambio de uso de suelo.'
        ],
        [
            'Artículo 93. La Secretaría autorizará el cambio de uso de suelo en terrenos forestales por excepción, previa opinión técnica de los miembros del Consejo Estatal Forestal de que se trate y con base en los estudios técnicos justificativos cuyo contenido se establecerá en el Reglamento, los cuales demuestren que la biodiversidad de los ecosistemas que se verán afectados se mantenga, y que la erosión de los suelos, el deterioro de la calidad del agua o la disminución en su captación se mitiguen en las áreas afectadas por la remoción de la vegetación forestal. \nEn las autorizaciones de cambio de uso de suelo en terrenos forestales, la Secretaría deberá dar respuesta debidamente fundada y motivada a las opiniones técnicas emitidas por los miembros del Consejo Estatal Forestal de que se trate. \nLas autorizaciones que se emitan deberán integrar un programa de rescate y reubicación de especies de la flora y fauna afectadas y su adaptación al nuevo hábitat conforme se establezca en el Reglamento. Dichas autorizaciones deberán sujetarse a lo que, en su caso, dispongan los programas de ordenamientos ecológicos correspondientes, las Normas Oficiales Mexicanas y demás disposiciones legales y reglamentarias aplicables.',
            'La promovente presenta el Documento Técnico Unificado para el cambio de uso de suelo en terrenos forestales para su evaluación y posterior autorización, en el cual demostrará que no se compromete la biodiversidad, ni se provocará la erosión de los suelos, el deterioro de la calidad del agua o la disminución en su captación; y que los usos alternativos del suelo que se proponen son más productivos a largo plazo que el uso que tiene actualmente, se integra así mismo un programa de rescate de especies de Flora y Fauna, de igual manera se plasman las mejores técnicas y apegándose a la normatividad para la realización del cambio de uso de suelo con medios mecánicos y/o manuales sin el uso de agroquímicos o el uso del fuego.'
        ],
        [
            'Artículo 98. Los interesados en el cambio de uso de suelo en terrenos forestales, deberán comprobar que realizaron el depósito ante el Fondo Forestal Mexicano, por concepto de compensación ambiental, para que se lleven a cabo acciones de restauración de los ecosistemas que se afecten, preferentemente dentro de la cuenca hidrográfica en donde se ubique la autorización del proyecto, en los términos y condiciones que establezca el Reglamento.',
            'Si se solicita el depósito al fondo forestal, la promovente está en la mejor disposición de realizar este depósito para poder obtener la autorización de cambio de uso de suelo en terrenos forestales.',
        ]
    ]

    filas = 7
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla[:2], start=2):
        celda = i  # fila destino en tabla
        
        for cols in range(2):  # solo col 0 y 1
            if cols < len(rows):  # si hay texto para esta columna
                cell = tabla16b.cell(celda, cols)
                t16b = cell.paragraphs[0].add_run(rows[cols])
                t16b.font.size = Pt(12)
                t16b.font.name = 'Arial'
                t16b.bold = False

    for i, rows in enumerate(datos_tabla[-2:], start=5):
        celda = i  # fila destino en tabla
        
        for cols in range(2):  # solo col 0 y 1
            if cols < len(rows):  # si hay texto para esta columna
                cell = tabla16b.cell(celda, cols)
                t16b = cell.paragraphs[0].add_run(rows[cols])
                t16b.font.size = Pt(12)
                t16b.font.name = 'Arial'
                t16b.bold = False

    # Rellenar las celdas fusionadas
    if tabla16b.cell(1, 0) and tabla16b.cell(3, 0):
        for i in range(2):
            celda_fusionada = (i * 3) + 1
            row = tabla16b.rows[celda_fusionada]
            merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
            cell_background_color(merged_cell, '#DBE5F1')

            # Agregar texto a la celda fusionada
            t16b = merged_cell.paragraphs[0].add_run(f'{capitulo_seccion[i]}')
            t16b.font.name = 'Arial'
            t16b.font.size = Pt(12)
            t16b.bold = True
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # Celdas de encabezados
    for cols in range(columnas):
        cell = tabla16b.cell(0, cols)
        t16b = cell.paragraphs[0].add_run(f'{encabezados[cols]}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#007AD1")

    ########################################################################################################################################################################
    # Capitulo 16.3.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.3.5 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.5.- Ley General de Vida Silvestre')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.3.5 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el DOF el 03 de Julio del 2000). Última reforma publicada en el DOF 20-05-2021)')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.3.5 ###
    #########################
    encabezados = [
        'Artículos relacionados con el proyecto',
        'Vinculación',
    ]

    capitulo_seccion = [
            'Ley General De Vida Silvestre \nTítulo I Disposiciones Preliminares',
        
            'Título V. Disposiciones Comunes Para la Conservación y el Aprovechamiento \nSustentable de la Vida Silvestre \nCapítulo I. Disposiciones Preliminares',

            'Capítulo II \nHábitat Crítico Para La Conservación De La Vida Silvestre',

            'Capítulo VIII \nConservación De Las Especies Migratorias',

            'Capitulo II \nDaños',

            'Capítulo V \nInfracciones y Sanciones Administrativas',
    ]

    datos_tabla = [
        [
            'Artículo 1º. La presente Ley es de orden público y de interés social, reglamentario del párrafo tercero del artículo 27 y de la fracción XXIX, inciso G del artículo 73 constitucionales. Su objeto es establecer la concurrencia del Gobierno Federal, de los gobiernos de los Estados y de los Municipios, en el ámbito de sus respectivas competencias, relativa a la conservación y aprovechamiento sustentable de la vida silvestre y su hábitat en el territorio de la República Mexicana y en las zonas en donde la Nación ejerce su jurisdicción.',
            'El presente proyecto no requerirá en ninguna de sus etapas del aprovechamiento de especies de vida silvestre para su comercialización, por lo que no se hará necesario la aplicación de ningún artículo de esta Ley relacionado con ese tipo de consumo.',
        ],
        [
            'Artículo 19. Las autoridades que, en el ejercicio de sus atribuciones, deban intervenir en las actividades relacionadas con la utilización del suelo, agua y demás recursos naturales con fines agrícolas, ganaderos, piscícolas, forestales y otros, observarán las disposiciones de esta Ley y las que de ella se deriven, y adoptarán las medidas que sean necesarias para que dichas actividades se lleven a cabo de modo que se eviten, prevengan, reparen, compensen o minimicen los efectos negativos de las mismas sobre la vida silvestre y su hábitat.',
            'El promovente se acatará a las observación de la secretaria, asi mismo realizara actividades de prevención para mitigar los impactos hacia la fauna silvestre, se evitará realizar ruidos estridentes que causen estrés o ahuyentamiento de la fauna local, se respetara los hábitat de las especies vulnerables, el transito vehiculas será bajo de 40 k/h para evitar colisión, especies de lento desplazamiento.'
        ],
        [
            'Artículo 63. La conservación del hábitat natural de la vida silvestre es de interés público. \nLos hábitats críticos para la conservación de la vida silvestre son áreas específicas terrestres o acuáticas, en las que ocurren procesos biológicos, físicos y químicos esenciales, ya sea para la supervivencia de especies en categoría de riesgo, ya sea para una especie, o para una de sus poblaciones, y que por tanto requieren manejo y protección especial. Son áreas que regularmente son utilizadas para alimentación, depredación, forrajeo, descanso, crianza o reproducción, o rutas de migración.',
            ' '
        ],
        [
            'Artículo 76. La conservación de las especies migratorias se llevará a cabo mediante la protección y mantenimiento de sus hábitats, el muestreo y seguimiento de sus poblaciones, así como el fortalecimiento y desarrollo de la cooperación internacional; de acuerdo con las disposiciones de esta Ley, de la Ley General del Equilibrio Ecológico y la Protección al Ambiente y de las que de ellas se deriven, sin perjuicio de lo establecido en los tratados y otros acuerdos internacionales en los que México sea Parte Contratante.',
            'En caso de que en el area dl proyecto se detecte hábitats estos serán conservados o reubicados dentro del mismas condiciones y será prioridad la conservación de las especies migratorias dejando las especies que sirven como refugio o alimentación y estas áreas serán protegidas o conservación y excluidas de todo aprovechamiento.',
        ],
        [
            'Artículo 106. Sin perjuicio de las demás disposiciones aplicables, toda persona física o moral que ocasione directa o indirectamente un daño a la vida silvestre o a su hábitat, está obligada a repararlo o compensarlo de conformidad a lo dispuesto por la Ley Federal de Responsabilidad Ambiental',
            'La promovente sabe de las sanciones que se puedan dar cuando se ocasione algunos daños a la fauna silvestre y está en la mejor disposición si esto sucede a su remediación. \nAsi mismo estará estrictamente prohibido la colecta o caza de ejemplares de faunas presentes en el area como areas aledañas al proyecto.',
        ],
        [
            'Artículo 122. Son infracciones a lo establecido en esta Ley: \nI.	Realizar cualquier acto que cause la destrucción o daño de la vida silvestre o de su hábitat, en contravención de lo establecido en la presente Ley. \nII.	VII. Presentar información falsa a la Secretaría.',
            'El promovente entiende y se da por enterado de las infracciones que pueden darse con la destrucción o daño de la vida silvestre o su hábitat.',
        ],
    ]

    filas = len(encabezados) + len(datos_tabla)+ len(capitulo_seccion)
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla):
        celda = (i * 2) + 2  # fila destino en tabla
        
        for cols in range(2):  # solo col 0 y 1
            if cols < len(rows):  # si hay texto para esta columna
                cell = tabla16b.cell(celda, cols)
                t16b = cell.paragraphs[0].add_run(rows[cols])
                t16b.font.size = Pt(12)
                t16b.font.name = 'Arial'
                t16b.bold = False

    # Rellenar las celdas fusionadas
    for i, rows in enumerate(capitulo_seccion):
        celda_fusionada = (i * 2) + 1
        row = tabla16b.rows[celda_fusionada]
        merged_cell = row.cells[0].merge(row.cells[1])
        cell_background_color(merged_cell, '#DBE5F1')

        # Agregar texto a la celda fusionada
        t16b = merged_cell.paragraphs[0].add_run(f'{rows}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Celdas de encabezados
    for i, cols in enumerate(encabezados):
        cell = tabla16b.cell(0, i)
        t16b = cell.paragraphs[0].add_run(f'{cols}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#007AD1")

    ########################################################################################################################################################################
    # Capitulo 16.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.4 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.4.- Reglamentos Relacionados con el Proyecto.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.4.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'XVI.6.1.- Constitución Política de los Estados Unidos Mexicanos')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.4.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el Diario Oficial de la Federación el 5 de febrero de 1917, con última reforma del día 28 de mayo de 2021).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.4.1 ###
    #########################
    encabezados = [
        'Artículos relacionados con el proyecto',
        'Vinculación',
    ]

    capitulo_seccion = [
            'Capítulo I. Disposiciones Generales',
        
            'Capítulo II. De las Obras o Actividades que requieren Autorización en Materia de Impacto Ambiental y de Las Excepciones',
    ]

    datos_tabla = [
        [
            'ARTÍCULO 3.- Para los efectos del presente reglamento se considerarán las definiciones contenidas en la ley y las siguientes: \nI Ter. Cambio de uso de suelo: Modificación de la vocación natural o predominante de los terrenos, llevada a cabo por el hombre a través de la remoción total o parcial de la vegetación.',
            'La promovente tiene por entendido la definición del cambio de uso de suelo por lo cual realizará los trámites correspondientes para poder obtener una autorización para poder realizar esta actividad.',
        ],
        [
            'Artículo 5o.- Quienes pretendan llevar a cabo alguna de las siguientes obras o actividades, requerirán previamente la autorización de la Secretaría en materia de impacto ambiental: \nO) Cambios de uso del Suelo de Áreas Forestales, así como en Selvas y Zonas Áridas: \nI. Cambio de uso del suelo para actividades agropecuarias, acuícolas, de desarrollo inmobiliario, de infraestructura urbana, de vías generales de comunicación o para el establecimiento de instalaciones comerciales, industriales o de servicios en predios con vegetación forestal, con excepción de la construcción de vivienda unifamiliar y del establecimiento de instalaciones comerciales o de servicios en predios menores a 1000 metros cuadrados, cuando su construcción no implique el derribo de arbolado en una superficie mayor a 500 metros cuadrados, o la eliminación o fragmentación del hábitat de ejemplares de flora o fauna sujetos a un régimen de protección especial de conformidad con las normas oficiales mexicanas y otros instrumentos jurídicos aplicables; \nII. Cambio de uso del suelo de áreas forestales a cualquier otro uso, con excepción de las actividades agropecuarias de autoconsumo familiar, que se realicen en predios con pendientes inferiores al cinco por ciento, cuando no impliquen la agregación ni el desmonte de más del veinte por ciento de la superficie total y ésta no rebase 2 hectáreas en zonas templadas y 5 en zonas áridas, y  \nIII. Los demás cambios de uso del suelo, en terrenos o áreas con uso de suelo forestal, con excepción de la modificación de suelos agrícolas o pecuarios en forestales, agroforestales o silvopastoriles, mediante la utilización de especies nativas.',
            'La promovente tiene por entendido la definición del cambio de uso de suelo por lo cual realizará los trámites correspondientes para poder obtener una autorización para por realizar esta actividad de cambio de uso de suelo., es por eso que presenta el documento unificado.'
        ],
        [
            'Artículo 14.- Cuando la realización de una obra o actividad que requiera sujetarse al procedimiento de evaluación de impacto ambiental involucre, además, el cambio de uso del suelo de áreas forestales y en selvas y zonas áridas, los promovente podrán presentar una sola manifestación de impacto ambiental que incluya la información relativa a ambos proyectos.',
            'La promovente ingresa el documento técnico unificado modalidad A, para el cambio de uso de suelo en terrenos forestales para su evaluación y ver su viabilidad.'
        ],
    ]

    filas = 1 + len(datos_tabla)+ len(capitulo_seccion)
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla, start=1):
        # calcular fila según la secuencia 2,4,5,7,8,9...
        if i % 2 == 1:  # posición impar → suma de 3
            fila_destino = 2 + 3 * ((i - 1) // 2)
        else:           # posición par → suma de 2
            fila_destino = 4 + 2 * ((i // 2) - 1)
        
        # rellenar columnas 0 y 1 en esa fila
        for cols in range(2):
            if cols < len(rows):
                cell = tabla16b.cell(fila_destino, cols)
                t16b = cell.paragraphs[0].add_run(rows[cols])
                t16b.font.size = Pt(12)
                t16b.font.name = 'Arial'
                t16b.bold = False

    # Rellenar las celdas fusionadas
    for i, rows in enumerate(capitulo_seccion):
        celda_fusionada = (i * 2) + 1
        row = tabla16b.rows[celda_fusionada]
        merged_cell = row.cells[0].merge(row.cells[1])
        cell_background_color(merged_cell, '#DBE5F1')

        # Agregar texto a la celda fusionada
        t16b = merged_cell.paragraphs[0].add_run(f'{rows}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Celdas de encabezados
    for i, cols in enumerate(encabezados):
        cell = tabla16b.cell(0, i)
        t16b = cell.paragraphs[0].add_run(f'{cols}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#007AD1")

    ########################################################################################################################################################################
    # Capitulo 16.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.4.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.6.3.- Reglamento de la Ley General de Desarrollo Forestal Sustentable.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.4.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(Publicado en el diario oficial de la federación el 21 de febrero de 2005 con la última reforma publicada en el DOF 09-12-2020)')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 16.4.2 ###
    #########################
    encabezados = [
        'Articulo',
        'Referencia',
        'Vinculación',
    ]

    capitulo_seccion = [
            'Artículos del Reglamento de la LGDFS aplicables al proyecto',
    ]

    datos_tabla = [
        [
            '1°',
            "Reglamentación de la Ley General de Desarrollo Forestal Sustentable",
            "Aplicación de instrumentos de política forestal para el Estudio Técnico Justificativo para el Cambio de Uso de Suelo en Terrenos con Vegetación Forestal."
        ],
        [
            '2° Fracc. XIX ',
            "Plano Georreferenciados",
            "Levantamiento de poligonal en coordenadas UTM y/o Geográficas ubicadas en su respectiva Cuenca y Subcuenca Hidrológica-Forestal escala 1:50,000 para su localización."
        ],
        [
            '139',
            "Solicitud de Cambio de Uso de Suelo.",
            "Elaboración de Estudio Técnico Justificativo en apego al contenido del Art 141 de la LGDFS."
        ],
        [
            '141',
            "Estudio Técnico Justificativo para el CUSF.",
            "Elaboración de documento en apego a la LGDFS de acuerdo con el contenido según Artículo 141 de dicha ley."
        ],
        [
            '144',
            "Monto de Compensación Ambiental por CUSF.",
            "Apego al Art 152 en base a costos que determine la Comisión para la reforestación, restauración y mantenimiento."
        ]
    ]

    filas = 1 + len(datos_tabla)+ len(capitulo_seccion)
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla):
        celda = i + 2
        for cols in range(columnas):
            cell = tabla16b.cell(celda, cols)
            t16b = cell.paragraphs[0].add_run(f'{rows[cols]}')
            t16b.font.size = Pt(12)
            t16b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rellenar las celdas fusionadas
    for i, rows in enumerate(capitulo_seccion):
        #celda_fusionada = (i * 2) + 1
        row = tabla16b.rows[i]
        merged_cell = row.cells[0].merge(row.cells[2])
        cell_background_color(merged_cell, '#007AD1')

        # Agregar texto a la celda fusionada
        t16b = merged_cell.paragraphs[0].add_run(f'{rows}')
        t16b.font.name = 'Arial'
        t16b.font.size = Pt(12)
        t16b.bold = True
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Celdas de encabezados
    for i, cols in enumerate(encabezados):
        cell = tabla16b.cell(1, i)
        t16b = cell.paragraphs[0].add_run(f'{cols}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#86CDFF")

    ########################################################################################################################################################################
    # Capitulo 16.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.4.3 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\n')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.4.3 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run()
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.4.3 ###
    #########################
    encabezados = [
        'Articulos Relacionados con el proyecto',
        'Vinculación',
    ]

    datos_tabla = [
        [
            "Artículo 2. Además de las definiciones contenidas en el artículo 3o. de la Ley General de Vida Silvestre y la Ley General del Equilibrio Ecológico y la Protección al Ambiente, para efectos del presente Reglamento se entenderá por:\nRemediación. El conjunto de actividades tendentes a resolver, bajo criterios técnicos y mediante medidas de manejo o control, problemas específicos asociados a ejemplares y poblaciones que se tornen perjudiciales, o bien, a la restauración y recuperación del hábitat de las especies silvestres;",
            "En caso de encontrarse especies perjudiciales se realizará medidas de control, así mismo se protegerán los hábitat de las especies que se encuentren dentro del área de estudio."
        ],
        [
            "Artículo 12. Las personas que pretendan realizar cualquier actividad relacionada con hábitat, especies, partes o derivados de vida silvestre y que conforme a la Ley requieran licencia, permiso o autorización de la Secretaría, presentarán la solicitud correspondiente en los formatos que para tal efecto establezca la Secretaría, los cuales deberán contener:",
            " "
        ],
        [
            "Artículo 70. Para los efectos del artículo 63 de la Ley, el Acuerdo Secretarial por el que se establezca el hábitat crítico para la conservación de la vida silvestre se publicará en el Diario Oficial de la Federación y prevendrá la coordinación con las dependencias y entidades de la Administración Pública Federal para que éstas no autoricen proyectos o provean fondos que puedan destruir o amenazar las áreas designadas.\nCuando se establezca un hábitat crítico y se realicen actividades que puedan acelerar los procesos de degradación o destrucción del hábitat, respecto de los cuales se hayan expedido autorizaciones que se encuentren vigentes al momento de su establecimiento, las autoridades que hubiesen expedido dichas autorizaciones promoverán la incorporación de sus titulares a los planes de recuperación previstos en el Acuerdo Secretarial del hábitat crítico de que se trate. Las áreas establecidas como hábitat crítico se definirán por la superficie que ocupaba la distribución de la especie en el momento en que fue listada.",
            " "
        ],
        [
            "Artículo 78. Las medidas de manejo, control y remediación de ejemplares o poblaciones perjudiciales podrán consistir en cualquiera de las siguientes, de acuerdo al orden de prelación que se indica:\nVI. Las acciones o dispositivos para ahuyentar, dispersar, dificultar el acceso de los ejemplares o disminuir el daño que ocasionan, cuando así se justifique.",
            " "
        ],
        [
            "Artículo 80. Cuando en un predio, zona o región sea necesario aplicar medidas de manejo o control de ejemplares o poblaciones perjudiciales, los interesados podrán solicitar autorización a la Secretaría, señalando en el escrito correspondiente la siguiente información:\nI. Especies a controlar, identificadas por nombre común y nombre científico;\nII. Razones para considerar a los ejemplares o poblaciones de la especie o especies de que se trate como perjudiciales;\nIII. Tipo de daño que provocan y su magnitud;\nIV. Método de control que se propone utilizar;\nV. Periodo de tiempo o etapas en que se llevará a cabo el control;\nVI. Responsable técnico que supervisará la ejecución de las medidas propuestas;\nVII. Forma en que se pretende disponer de los ejemplares objeto de las medidas de control, y\nVIII. En su caso, medidas de prevención y control aplicadas con anterioridad para resolver el problema, así como las que se propongan para atender los problemas secundarios que pudieran derivarse de la aplicación del método de control propuesto.",
            " "
        ]
    ]


    filas = 1 + len(datos_tabla)
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla):
        celda = i + 1
        for cols in range(columnas):
            cell = tabla16b.cell(celda, cols)
            t16b = cell.paragraphs[0].add_run(f'{rows[cols]}')
            t16b.font.size = Pt(12)
            t16b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Celdas de encabezados
    for i, cols in enumerate(encabezados):
        cell = tabla16b.cell(0, i)
        t16b = cell.paragraphs[0].add_run(f'{cols}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#86CDFF")

    ########################################################################################################################################################################
    # Capitulo 16.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.4.4 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.4.4.- Reglamentos de la Ley General de Cambio Climático en Materia del Registro Nacional de Emisiones.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.4.4 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('(publicado en el Diario Oficial de la Federación el 28 de octubre de 2014).')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 16.4.3 ###
    #########################
    encabezados = [
        'Articulos Relacionados con el proyecto',
        'Vinculación',
    ]

    datos_tabla = [
        [
            "Artículo 26. Las personas físicas o morales que hayan implementado proyectos o actividades que tengan como resultado la Mitigación, reducción o absorción de Emisiones de Gases o Compuestos de Efecto Invernadero, si éstos se han realizado en el territorio nacional, podrán solicitar la inscripción de dicha información en el Registro, previo Dictamen de Validación expedido por un Organismo acreditado y aprobado para tal efecto, que certifique el resultado de dichos proyectos.",
            " "
        ],
    ]


    filas = 1 + len(datos_tabla)
    columnas = len(encabezados)
    tabla16b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Rellenar los datos de la tabla
    for i, rows in enumerate(datos_tabla):
        celda = i + 1
        for cols in range(columnas):
            cell = tabla16b.cell(celda, cols)
            t16b = cell.paragraphs[0].add_run(f'{rows[cols]}')
            t16b.font.size = Pt(12)
            t16b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # Celdas de encabezados
    for i, cols in enumerate(encabezados):
        cell = tabla16b.cell(0, i)
        t16b = cell.paragraphs[0].add_run(f'{cols}')
        t16b.font.size = Pt(12)
        t16b.font.name = 'Arial'
        t16b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#86CDFF")

    ########################################################################################################################################################################
    # Capitulo 16
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\n')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run()
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.5 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.- Otros')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.5.1 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\nXVI.5.1.- Convenios entre Canadá, Estados Unidos y México para la Protección de las Aves Migratorias y los Mamíferos Cinegéticos.')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.5.1 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Los esfuerzos multilaterales para conservar y proteger de Fauna migratoria y cinegética se iniciaron en el año 1936 entre Estados Unidos y México con la Convención para la Protección de Aves Migratorias y de Mamíferos Cinegéticos, la cual se modificó en 1972 y ratificada en 1997 con una última modificación en el Diario Oficial de la Federación el 21 de julio de 2000. "
        "En el artículo II de este Convenio se tomó acuerdo por ambas partes para dictar leyes, reglamentos y disposiciones que determinaran la fijación de vedas para la captura, caza de aves migratorias o de sus nidos y huevos; la determinación de zonas de refugio en donde se prohíbe la captura de dichas aves; la limitación a cuatro meses de la temporada de caza; la veda del 10 de marzo al 1º de Septiembre, entre otras disposiciones. "
        "El proyecto, dentro de las medidas de prevención y mitigación de los impactos a la Flora y la Fauna silvestre, contempla la capacitación, sensibilización y la prohibición de los trabajadores que participen en el proyecto para no cazar, capturar, comerciar o afectar de algún modo la Fauna silvestre en el área del proyecto. "
        "Por otro lado, hay que mencionar la Convención Internacional de Especies Amenazadas de Flora y Fauna Silvestre (CITES) firmado por Estados Unidos en 1973 y por México en 1991. "
        "Inició sus operaciones formales en 1992, estableciendo un sistema mundial de reglamentaciones de importación y exportación para prevenir la sobreexplotación de las plantas y animales que se encuentran enlistadas en los tres anexos de la Convención."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Vinculación. - El proyecto, toma en consideración las especies listadas en CITES con distribución regional coincidente con el área del proyecto. "
        "Dentro de las medidas de prevención (que se describen en el capítulo 10) se incluyen pláticas de sensibilización al personal que labore en las etapas de preparación del sitio, construcción, operación y/o funcionamiento, así como la prohibición expresa de cazar, capturar y/o comerciar con cualquier especie de Fauna silvestre."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 16.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 16.5.2 ###
    #########################
    capitulo16 = doc.add_paragraph()
    i16 = capitulo16.add_run(f'\n')
    i16_format = capitulo16.paragraph_format
    i16_format.line_spacing = 1.15

    i16.font.name = 'Arial'
    i16.font.size = Pt(12)
    i16.font.bold = True
    capitulo16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 16.5.2 ###
    #########################
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Declaración de la Conferencia de las Naciones Unidas sobre el Medio Humano, aprobada en Estocolmo el 16 de junio de 1972 y,  con el objetivo de establecer una alianza mundial nueva y equitativa mediante la creación de nuevos niveles de cooperación entre los Estados, los sectores claves de las sociedades y las personas, procurando alcanzar acuerdos internacionales en los que se respeten los intereses de todos y se proteja la integridad del sistema ambiental y de desarrollo mundial, reconociendo la naturaleza integral e interdependiente de la Tierra, nuestro hogar, se basa en 27 principios el cual rige a la humanidad en el uso, protección, conservación para la preservación y desarrollo sostenible del medio ambiente. Por ello, el presente proyecto se apegará a los siguientes principios:"
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PRINCIPIO 1
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('PRINCIPIO 1')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Los seres humanos constituyen el centro de las preocupaciones relacionadas con el desarrollo sostenible. Tienen derecho a una vida saludable y productiva en armonía con la naturaleza."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PRINCIPIO 2"
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('PRINCIPIO 2')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "De conformidad con la Carta de las Naciones Unidas y los principios del derecho internacional, los Estados tienen el derecho soberano de aprovechar sus propios recursos según sus propias políticas ambientales y de desarrollo, y la responsabilidad de velar por que las actividades realizadas dentro de su jurisdicción o bajo su control no causen daños al medio ambiente de otros Estados o de zonas que estén fuera de los límites de la jurisdicción nacional."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PRINCIPIO 16"
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('PRINCIPIO 16')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Las autoridades nacionales deberían procurar fomentar la internalización de los costos ambientales y el uso de instrumentos económicos, teniendo en cuenta el criterio de que el que contamina debe, en principio, cargar con los costos de la contaminación, teniendo debidamente en cuenta el interés público y sin distorsionar el comercio ni las inversiones internacionales."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PRINCIPIO 17"
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run('PRINCIPIO 17')
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    descripcionCapitulo16.bold = True
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Deberá emprenderse una evaluación del impacto ambiental, en calidad de instrumento nacional, respecto de cualquier actividad propuesta que probablemente haya de producir un impacto negativo considerable en el medio ambiente y que esté sujeta a la decisión de una autoridad nacional competente."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di16 = doc.add_paragraph()
    descripcionCapitulo16 = di16.add_run(
        "Vinculación. - El presente proyecto cumple con lo establecido en esta declaración al respetar el derecho a la salud, acogiéndose a lo establecido en la legislación mexicana y respondiendo a las necesidades de desarrollo de las generaciones actuales y futuras. "
        "Se considera la internalización de los costos ambientales y la evaluación en materia de impacto ambiental. Además, se cuenta con la certidumbre de que no afectará a ninguna especie animal o vegetal con estatus de protección, o afecte la biodiversidad presente en el área."
    )
    descripcionCapitulo16_format = di16.paragraph_format
    descripcionCapitulo16_format.line_spacing = 1.15
    descripcionCapitulo16.font.name = 'Arial'
    descripcionCapitulo16.font.size = Pt(12)
    di16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 16 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo16() # Crear el documento
