"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos


""" 
    ============================================================
    Creacion del documento
    ============================================================
"""
def capitulo4():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 4
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo IV.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    temasCapitulo4 = [
        "IV.- DESCRIPCIÓN DE LOS ELEMENTOS FISICOS Y BIOLOGICOS DE LA CUENCA HIDROLOGICO-FORESTAL EN DONDE SE UBIQUE EL PREDIO.", [
            "IV.1.- Delimitación del área de estudio donde pretende establecerse el proyecto...",
            "IV.2.- Caracterización y análisis del Sistema Ambiental Hidrológico-Forestal.", [
                "IV.2.1.- Caracterización y análisis retrospectivo de la calidad ambiental del Sistema Ambiental.",
                "IV.2.2.- Medio Físico", [
                    "IV.2.2.1.- Clima.",
                    "IV.2.2.2.- Temperatura.",
                    "IV.2.2.3.- Precipitación.",
                    "IV.2.2.4.- Evapotranspiración.",
                    "IV.2.2.5.- Viento.",
                    "IV.2.2.6.- Riesgos y vulnerabilidad", [
                        "IV.2.2.6.1.- Riesgos Hidrometeorológicos.", [
                            "IV.2.2.6.1.1.- Precipitación.",
                            "IV.2.2.6.1.2.- Tormentas de granizo y nieve.",
                            "IV.2.2.6.1.3.- Heladas.",
                            "IV.2.2.6.1.4.- Ciclo-nes tropicales.",
                            "IV.2.2.6.1.5.- Inundaciones.",
                            "IV.2.2.6.1.6.- Sequía.",
                            "IV.2.2.6.1.7.- Tornados.",
                            "IV.2.2.6.1.8.- Tormentas eléctricas."
                        ]
                    ],
                    "IV.2.2.7.- Suelo.", [
                        "IV.2.2.7.1.- Tipos de erosión presentes en el Sistema Ambiental.",
                        "IV.2.2.7.2.- Estimación de la erosión del suelo.", [
                            "IV.2.2.7.2.1.- Erosión hídrica del suelo en la condición actual del sistema ambiental.",
                            "IV.2.2.7.2.2.-Erosión potencial con el cambio de uso de suelo"
                        ],
                        "IV.2.2.7.3.- Predicción de la erosión eólica.", [
                            "IV.2.2.7.3.1.- Erosión eólica del suelo en la condición actual del sistema ambiental.",
                            "IV.2.2.7.3.2.- Erosión potencial sin vegetación.",
                            "IV.2.2.7.3.3.- Erosión potencial con aplicación de medidas de mitigación"
                        ],
                        "IV.2.2.7.4.- Uso potencial del suelo del Sistema Ambiental."
                    ],
                    "IV.2.2.8.- Geología.",
                    "IV.2.2.9.- Fisiografía.", [
                        "IV.2.2.9.1 Fallas y fracturas"
                    ],
                    "IV.2.2.10.- Sismicidad.",
                    "IV.2.2.11.- Topografía.", [
                        "IV.2.2.11.1.- Elevaciones",
                        "IV.2.2.11.2.- Pendiente",
                        "IV.2.2.11.3.- Exposiciones.",
                        "IV.2.2.11.4.- Toponimias"
                    ],
                    "IV.2.2.12.- Hidrología", [
                        "IV.2.2.12.1.- Acuífero Saltillo-Ramos Arizpe (0510), Estado de Coahuila….",
                        "IV.2.2.12.2.- ANÁLISIS DE LA INFILTRACIÓN.", [
                            "IV.2.2.12.1.1.- Situación actual hidrológica sin proyecto en el Sistema Ambiental.",
                            "IV.2.2.12.1.2.- Con la implementación del proyecto en el área del Sistema Ambiental con cambio de uso de suelo.",
                            "IV.2.2.12.1.3.- Análisis de la infiltración en el sistema ambiental."
                        ]
                    ]
                ],
                "IV.2.3.- Medio biológico.", [
                    "IV.2.3.1.- Vegetación.",
                    "IV.2.3.2.- Análisis de la vegetación Matorral Desértico Rosetófilo (MDR)", 
                    "IV.2.3.3.- Muestreo de la vegetación MDR y MDM en el Sistema Ambiental.", [
                        "IV.2.3.3.1.- Resultados de los sitios de muestreo del MDR en el Sistema Ambiental."
                    ],
                    "IV.2.3.4.- Análisis de diversidad de la vegetación MDR.", [
                        "IV.2.3.4.1.- Análisis de diversidad del estrato de las arbóreas MDR.",
                        "IV.2.3.4.2.- Análisis de diversidad del Estrato de las Arbustivas MDR….",
                        "IV.2.3.4.3.- Análisis de diversidad del Estrato de las Herbáceas MDR…..",
                        "IV.2.3.4.4.- Análisis de diversidad del Estrato de las Gramíneas MDR….",
                        "IV.2.3.4.5.- Análisis de diversidad del Estrato de las Suculentas MDR.",
                        "IV.2.3.4.6.- Análisis de los estratos MDR."
                    ],
                    "IV.2.3.5.- Análisis de la vegetación Matorral Desértico Micrófilo (MDM)…..", [
                        "IV.2.3.5.1.- Resultados de los sitios de muestreo del MDM en el Sistema Ambiental."
                    ],
                    "IV.2.3.6.- Análisis de diversidad de la vegetación MDM.", [
                        "IV.2.3.6.1.- Análisis de diversidad del estrato de las arbóreas MDM.",
                        "IV.2.3.6.2.- Análisis de diversidad del Estrato de las Arbustivas MDM….",
                        "IV.2.3.6.3.- Análisis de diversidad del Estrato de las Herbáceas MDM…..",
                        "IV.2.3.6.4.- Análisis de diversidad del Estrato de las Gramíneas MDM….",
                        "IV.2.3.6.5.- Análisis de diversidad del Estrato de las Suculentas MDM….",
                        "IV.2.3.6.6.- Análisis de los estratos MDM."
                    ],
                    "IV.2.3.7.- Fauna.", [
                        "IV.2.3.7.1.- Metodología para la evaluación de la Fauna Silvestre.",
                        "IV.2.3.7.2.- Resultado de especies faunísticas en el área del sistema ambiental.",
                        "IV.2.3.7.3.- Análisis de información del grupo de las aves en el área del Sistema Ambiental",
                        "IV.2.3.7.4.- Análisis de Información del grupo de los mamíferos en el área del Sistema Ambiental.",
                        "IV.2.3.7.5.- Análisis de información del grupo de los reptiles en el área del sistema Ambiental.",
                        "IV.2.3.7.6.- Análisis de la información del grupo de los lepidópteros en el área del sistema ambiental.",
                        "IV.2.3.7.7.- Análisis de la información de los grupos faunísticos en el área del sistema ambiental."
                    ]
                ]
            ]
        ]
    ]

    ########################################################################################################################################################################
    # Capitulo 4
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 4 ###
    #########################
    capitulo4 = doc.add_paragraph()
    i4 = capitulo4.add_run(f'{temasCapitulo4[0]}')
    i4_format = capitulo4.paragraph_format
    i4_format.line_spacing = 1.15

    i4.font.name = 'Arial'
    i4.font.size = Pt(12)
    i4.font.bold = True
    capitulo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.1 ###
    #########################
    capitulo41 = doc.add_paragraph()
    i41 = capitulo41.add_run(f'{temasCapitulo4[1][0]}')
    i41_format = capitulo41.paragraph_format
    i41_format.line_spacing = 1.15

    i41.font.name = 'Arial'
    i41.font.size = Pt(12)
    i41.font.bold = True
    capitulo41.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4 ###
    #########################
    di41 = doc.add_paragraph()
    descripcionCapitulo41 = di41.add_run('El área sujeta al presente estudio para el establecimiento del proyecto “_______________________________”, está constituida y para su análisis de este capítulo por la microcuenca “_______________________”, el cual será sujeto de análisis y en lo sucesivo se le denominará el Sistema Ambiental o SA, y está inmerso en la Subcuenca _____________________________________________________________________________________. El sistema ambiental se encuentra inmerso en los municipios de __________________________________________________________ esta última perteneciente al Estado de Nuevo León. Las cuales fueron consideradas como una única área de estudio y para la descripción, está constituida en su mayor parte por _______________________________________________________________________________. Esta microcuenca ya están definidas por el programa de Fideicomiso de Riesgo Compartido (FIRCO), en el año 2005, delimito las microcuencas a nivel nacional, la metodología para la delimitación de las microcuencas, para ello utilizaron capas vectoriales a escalas 1:250,000, dichas capas vectoriales o coberturas son las Cuencas y subcuencas Hidrológicas, Carta Topográfica, Red de Caminos, Cartas de Elevaciones, Carta de Pendientes, Carta Fisiográfica, Red de Caminos, Carta de Clima, Carta de uso de Suelo y Vegetación, Carta de toponimias, Red Hidrológica Subterráneas, Red de Hidrología Superficial, Fuente: Capa Delimitación Nacional de Microcuencas, FIRCO 2005, una vez que se tiene la capa de Microcuencas se ubicó el predio o área de cambio de uso de suelo y se selección la microcuenca que inciden en el predio, quedando como el Sistema Ambiental objeto de análisis del presente documento, (Ver anexo Mapa 4.1.- Delimitación del Sistema Ambiental y Área de estudio).')
    descripcionCapitulo41_format = di41.paragraph_format
    descripcionCapitulo41_format.line_spacing = 1.15
    descripcionCapitulo41_format.space_after = 0
    descripcionCapitulo41_format.space_before = 0

    descripcionCapitulo41.font.name = 'Arial'
    descripcionCapitulo41.font.size = Pt(12)
    di41.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2 ###
    #########################
    capitulo42 = doc.add_paragraph()
    i42 = capitulo42.add_run(f'\n{temasCapitulo4[1][1]}')
    i42_format = capitulo42.paragraph_format
    i42_format.line_spacing = 1.15

    i42.font.name = 'Arial'
    i42.font.size = Pt(12)
    i42.font.bold = True
    capitulo42.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.1 ###
    #########################
    capitulo421 = doc.add_paragraph()
    i421 = capitulo421.add_run(f'{temasCapitulo4[1][2][0]}')
    i421_format = capitulo421.paragraph_format
    i421_format.line_spacing = 1.15

    i421.font.name = 'Arial'
    i421.font.size = Pt(12)
    i421.font.bold = True
    capitulo421.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.1 ###
    #########################
    di421 = doc.add_paragraph()
    descripcionCapitulo421 = di421.add_run('Descripción del Capitulo 4.2.1')
    descripcionCapitulo421_format = di421.paragraph_format
    descripcionCapitulo421_format.line_spacing = 1.15
    descripcionCapitulo421_format.space_after = 0
    descripcionCapitulo421_format.space_before = 0

    descripcionCapitulo421.font.name = 'Arial'
    descripcionCapitulo421.font.size = Pt(12)
    di421.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2 ###
    #########################
    capitulo422 = doc.add_paragraph()
    i422 = capitulo422.add_run(f'\n{temasCapitulo4[1][2][1]}')
    i422_format = capitulo422.paragraph_format
    i422_format.line_spacing = 1.15

    i422.font.name = 'Arial'
    i422.font.size = Pt(12)
    i422.font.bold = True
    capitulo422.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.1 ###
    #########################
    capitulo4221 = doc.add_paragraph()
    i4221 = capitulo4221.add_run(f'{temasCapitulo4[1][2][2][0]}')
    i4221_format = capitulo4221.paragraph_format
    i4221_format.line_spacing = 1.15

    i4221.font.name = 'Arial'
    i4221.font.size = Pt(12)
    i4221.font.bold = True
    capitulo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.1 ###
    #########################
    di4221 = doc.add_paragraph()
    descripcionCapitulo4221 = di4221.add_run('En el área del Sistema Ambiental se encuentra en ___________________________________________________________________________________________________________________. En el área que comprende el sistema ambiental se encuentran climas desde __________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo4221_format = di4221.paragraph_format
    descripcionCapitulo4221_format.line_spacing = 1.15
    descripcionCapitulo4221_format.space_after = 0
    descripcionCapitulo4221_format.space_before = 0

    descripcionCapitulo4221.font.name = 'Arial'
    descripcionCapitulo4221.font.size = Pt(12)
    di4221.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4221 = doc.add_paragraph()
    descripcionCapitulo4221 = di4221.add_run("Para representar los tipos de clima presentes, se utilizó la __________________________________, del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el Conjunto de datos vectoriales del Continuo Nacional de Efectos Climáticos Regionales escala 1: 250,000, en formato digital, así como las fórmulas climáticas, se determinó de acuerdo al sistema de clasificación de Köppen modificado por Enriqueta García, encontrando que el clima más dominante es __________________________________________________________________________________________________. A continuación, se enlistan y se describen. (Ver anexo Mapa 4.2.- Tipos de climas del SA).")
    descripcionCapitulo4221_format = di4221.paragraph_format
    descripcionCapitulo4221_format.line_spacing = 1.15
    descripcionCapitulo4221_format.space_after = 0
    descripcionCapitulo4221_format.space_before = 0

    descripcionCapitulo4221.font.name = 'Arial'
    descripcionCapitulo4221.font.size = Pt(12)
    di4221.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.1 ###
    #########################
    tituloTabla42211 = doc.add_paragraph()
    dti42211 = tituloTabla42211.add_run('\nTabla 4.1.- Clasificación de climas del Sistema Ambiental')
    dti42211_format = tituloTabla42211.paragraph_format
    dti42211_format.line_spacing = 1.15
    dti42211_format.space_after = 0

    dti42211.font.name = 'Courier New'
    dti42211.font.size = Pt(12)
    tituloTabla42211.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.1 ###
    #########################
    tabla42211 = doc.add_table(rows=6, cols=5, style='Table Grid')
    tabla42211.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tabla in range(5):
        cell = tabla42211.cell(0, tabla)
        t42211 = cell.paragraphs[0].add_run(f'Columna {tabla + 1}')
        t42211.font.name = 'Arial'
        t42211.font.size = Pt(12)
        t42211.font.bold = True
    
    #########################
    ### Descripcion del capitulo 4.2.2.1 ###
    #########################

    di4221 = doc.add_paragraph()
    descripcionCapitulo4221 = di4221.add_run("\n1A continuación, se describen los diferentes tipos de climas encontrados en el área del sistema ambiental a estudiar.")
    descripcionCapitulo4221_format = di4221.paragraph_format
    descripcionCapitulo4221_format.line_spacing = 1.15
    descripcionCapitulo4221_format.space_after = 0
    descripcionCapitulo4221_format.space_before = 0

    descripcionCapitulo4221.font.name = 'Arial'
    descripcionCapitulo4221.font.size = Pt(12)
    di4221.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.1 ###
    #########################
    tituloTabla42212 = doc.add_paragraph()
    dti42212 = tituloTabla42212.add_run('\nTabla 4.2.- Descripción de los climas del Sistema Ambiental')
    dti42212_format = tituloTabla42212.paragraph_format
    dti42212_format.line_spacing = 1.15
    dti42212_format.space_after = 0

    dti42212.font.name = 'Courier New'
    dti42212.font.size = Pt(12)
    tituloTabla42212.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.1 ###
    #########################
    tabla42212 = doc.add_table(rows=6, cols=3, style='Table Grid')
    tabla42212.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tabla in range(3):
        cell = tabla42212.cell(0, tabla)
        t42212 = cell.paragraphs[0].add_run(f'Columna {tabla + 1}')
        t42212.font.name = 'Arial'
        t42212.font.size = Pt(12)
        t42212.font.bold = True
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.2 ###
    #########################
    capitulo4222 = doc.add_paragraph()
    i4222 = capitulo4222.add_run(f'\n{temasCapitulo4[1][2][2][1]}')
    i4222_format = capitulo4222.paragraph_format
    i4222_format.line_spacing = 1.15

    i4222.font.name = 'Arial'
    i4222.font.size = Pt(12)
    i4222.font.bold = True
    capitulo4222.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.2 ###
    #########################
    di4222 = doc.add_paragraph()
    descripcionCapitulo4222 = di4222.add_run('De acuerdo a la Estación Meteorológica de influencia en el sistema ambiental en estudio, por estar más cerca y presentar datos históricos es la estación ____ ubicada en el municipio ___________ de la Comisión Nacional del Agua (CONAGUA) en un periodo de ___________________, la temperatura máxima fue de ___________________________________. Los meses más cálidos registrados por esta estación fueron los de _____________, con temperaturas superiores a los ____., los meses con temperatura más baja ocurrieron predominantemente en la época de invierno en los meses _________________________________________________________________________________________________________. ')
    descripcionCapitulo4222_format = di4222.paragraph_format
    descripcionCapitulo4222_format.line_spacing = 1.15
    descripcionCapitulo4222_format.space_after = 0
    descripcionCapitulo4222_format.space_before = 0

    descripcionCapitulo4222.font.name = 'Arial'
    descripcionCapitulo4222.font.size = Pt(12)
    di4222.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.2 ###
    #########################
    tituloTabla4222 = doc.add_paragraph()
    dti4222 = tituloTabla4222.add_run('\nTabla 4.3.- Temperatura.')
    dti4222_format = tituloTabla4222.paragraph_format
    dti4222_format.line_spacing = 1.15
    dti4222_format.space_after = 0

    dti4222.font.name = 'Courier New'
    dti4222.font.size = Pt(12)
    tituloTabla4222.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.2 ###
    #########################
    tabla4222 = doc.add_table(rows=14, cols=4, style='Table Grid')
    

    #########################
    # Celda fusionada ""
    row1 = tabla4222.rows[0]
    merged_cell1 = row1.cells[1].merge(row1.cells[1].merge(row1.cells[3]))

    # Agregar texto a la celda fusionada
    t4222 = merged_cell1.paragraphs[0].add_run('Columna Fusionada')
    t4222.font.name = 'Arial'
    t4222.font.size = Pt(12)
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_top = tabla4222.cell(0, 0)
    cell_bottom = tabla4222.cell(1, 0)

    merged_cell = cell_top.merge(cell_bottom)

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('Filas Fusionadas')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tabla in range(3):
        cell = tabla4222.cell(1, tabla + 1)
        t4222 = cell.paragraphs[0].add_run(f'Columna {tabla + 1}')
        t4222.font.name = 'Arial'
        t4222.font.size = Pt(12)
        t4222.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Grafica del capitulo 4.2.2.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo4222 = doc.add_paragraph()
    imagenCapitulo4222.text = ''
    imagenCapitulo4222 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo4222.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo4222.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo4222.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo4222.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.2 ###
    #########################
    tituloGrafico4222 = doc.add_paragraph()
    dgi4222 = tituloGrafico4222.add_run('\nGrafica 4.1.- Temperatura estación CONAGUA')
    dgi4222_format = tituloGrafico4222.paragraph_format
    dgi4222_format.line_spacing = 1.15
    dgi4222_format.space_after = 0

    dgi4222.font.name = 'Bookman Old Style'
    dgi4222.font.size = Pt(12)
    tituloGrafico4222.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.3 ###
    #########################
    capitulo4223 = doc.add_paragraph()
    i4223 = capitulo4223.add_run(f'\n{temasCapitulo4[1][2][2][2]}')
    i4223_format = capitulo4223.paragraph_format
    i4223_format.line_spacing = 1.15

    i4223.font.name = 'Arial'
    i4223.font.size = Pt(12)
    i4223.font.bold = True
    capitulo4223.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.3 ###
    #########################
    di4223 = doc.add_paragraph()
    descripcionCapitulo4223 = di4223.add_run('Las precipitaciones observadas han sido de manera escasas y erráticas a lo largo del año, la precipitación registrada por la estación meteorológica ____ ubicada en el municipio de ____________________________________________________________________________________.')
    descripcionCapitulo4223_format = di4223.paragraph_format
    descripcionCapitulo4223_format.line_spacing = 1.15
    descripcionCapitulo4223_format.space_after = 0
    descripcionCapitulo4223_format.space_before = 0

    descripcionCapitulo4223.font.name = 'Arial'
    descripcionCapitulo4223.font.size = Pt(12)
    di4223.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.3 ###
    #########################
    tituloTabla4223 = doc.add_paragraph()
    dti4223 = tituloTabla4223.add_run('\nTabla 4.4.- Precipitación.')
    dti4223_format = tituloTabla4223.paragraph_format
    dti4223_format.line_spacing = 1.15
    dti4223_format.space_after = 0

    dti4223.font.name = 'Courier New'
    dti4223.font.size = Pt(12)
    tituloTabla4223.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.3 ###
    #########################
    tabla4223 = doc.add_table(rows=2, cols=14, style='Table Grid')

    for tabla in range(14):
        cell = tabla4223.cell(0, tabla)
        t4223 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t4223.font.name = 'Arial'
        t4223.font.size = Pt(12)
        t4223.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Descripcion del capitulo 4.2.2.3 ###
    #########################
    di4223 = doc.add_paragraph()
    descripcionCapitulo4223 = di4223.add_run('\nLos meses de mayor precipitación se encuentran de junio a septiembre, con precipitaciones que sobrepasan __________________________________________________________________________________________.')
    descripcionCapitulo4223_format = di4223.paragraph_format
    descripcionCapitulo4223_format.line_spacing = 1.15
    descripcionCapitulo4223_format.space_after = 0
    descripcionCapitulo4223_format.space_before = 0

    descripcionCapitulo4223.font.name = 'Arial'
    descripcionCapitulo4223.font.size = Pt(12)
    di4223.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.2.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo4223 = doc.add_paragraph()
    imagenCapitulo4223.text = ''
    imagenCapitulo4223 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo4223.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo4223.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo4223.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo4223.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.3 ###
    #########################
    tituloGrafico4223 = doc.add_paragraph()
    dgi4223 = tituloGrafico4223.add_run('\nGrafica 4.2.- Precipitación')
    dgi4223_format = tituloGrafico4223.paragraph_format
    dgi4223_format.line_spacing = 1.15
    dgi4223_format.space_after = 0

    dgi4223.font.name = 'Bookman Old Style'
    dgi4223.font.size = Pt(12)
    tituloGrafico4223.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.4 ###
    #########################
    capitulo4224 = doc.add_paragraph()
    i4224 = capitulo4224.add_run(f'\n{temasCapitulo4[1][2][2][3]}')
    i4224_format = capitulo4224.paragraph_format
    i4224_format.line_spacing = 1.15

    i4224.font.name = 'Arial'
    i4224.font.size = Pt(12)
    i4224.font.bold = True
    capitulo4224.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.4 ###
    #########################
    di4224 = doc.add_paragraph()
    descripcionCapitulo4224 = di4224.add_run('Los valores mensuales de evapotranspiración se calcularon de acuerdo al método de Thornthwaite (1948), este método es basado en la determinación de la evapotranspiración en función de la temperatura media correlacionada con la duración astronómica del día y el número de días. Por lo que cuando más alta es la temperatura, mayor es el valor de evapotranspiración. En el sistema ambiental el valor de evapotranspiración acumulada es de __________, la mayor concentración de valores de evapotranspiración se presentó en el mes de agosto, debido a que es el período de altas temperaturas, teniendo el mes de _________ con menor evapotranspiración, de acuerdo a la estación meteorológica que registra estos datos, a continuación, se muestra la distribución de la evapotranspiración en _______.')
    descripcionCapitulo4224_format = di4224.paragraph_format
    descripcionCapitulo4224_format.line_spacing = 1.15
    descripcionCapitulo4224_format.space_after = 0
    descripcionCapitulo4224_format.space_before = 0

    descripcionCapitulo4224.font.name = 'Arial'
    descripcionCapitulo4224.font.size = Pt(12)
    di4224.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.4 ###
    #########################
    tituloTabla4224 = doc.add_paragraph()
    dti4224 = tituloTabla4224.add_run('\nTabla 4.5.-	Evapotranspiración.')
    dti4224_format = tituloTabla4224.paragraph_format
    dti4224_format.line_spacing = 1.15
    dti4224_format.space_after = 0

    dti4224.font.name = 'Courier New'
    dti4224.font.size = Pt(12)
    tituloTabla4224.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.4 ###
    #########################
    tabla4224 = doc.add_table(rows=2, cols=14, style='Table Grid')

    for tabla in range(14):
        cell = tabla4224.cell(0, tabla)
        t4224 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t4224.font.name = 'Arial'
        t4224.font.size = Pt(12)
        t4224.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Grafica del capitulo 4.2.2.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo4224 = doc.add_paragraph()
    imagenCapitulo4224.text = ''
    imagenCapitulo4224 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo4224.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo4224.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo4224.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo4224.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.4 ###
    #########################
    tituloGrafico4224 = doc.add_paragraph()
    dgi4224 = tituloGrafico4224.add_run('\nGrafica 4.3.- Evapotranspiración.')
    dgi4224_format = tituloGrafico4224.paragraph_format
    dgi4224_format.line_spacing = 1.15
    dgi4224_format.space_after = 0

    dgi4224.font.name = 'Bookman Old Style'
    dgi4224.font.size = Pt(12)
    tituloGrafico4224.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.5 ###
    #########################
    capitulo4225 = doc.add_paragraph()
    i4225 = capitulo4225.add_run(f'\n{temasCapitulo4[1][2][2][4]}')
    i4225_format = capitulo4225.paragraph_format
    i4225_format.line_spacing = 1.15

    i4.font.name = 'Arial'
    i4225.font.size = Pt(12)
    i4225.font.bold = True

    #########################
    ### Descripcion del capitulo 4.2.2.5 ###
    #########################
    di4225 = doc.add_paragraph()
    descripcionCapitulo4225 = di4225.add_run('El valor de velocidad de viendo de obtuvo de la página que maneja los datos del ________________________________________, ubicado en el municipio de ____________, en la cual se obtuvo que la velocidad de viento promedio es de ________________________________________________________________.')
    descripcionCapitulo4225_format = di4225.paragraph_format
    descripcionCapitulo4225_format.line_spacing = 1.15
    descripcionCapitulo4225_format.space_after = 0
    descripcionCapitulo4225_format.space_before = 0

    descripcionCapitulo4225.font.name = 'Arial'
    descripcionCapitulo4225.font.size = Pt(12)
    di4225.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.5 ###
    #########################
    tituloTabla4225 = doc.add_paragraph()
    dti4225 = tituloTabla4225.add_run('\nTabla 4.6.- Velocidad de viento.')
    dti4225_format = tituloTabla4225.paragraph_format
    dti4225_format.line_spacing = 1.15
    dti4225_format.space_after = 0

    dti4225.font.name = 'Courier New'
    dti4225.font.size = Pt(12)
    tituloTabla4225.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.5 ###
    #########################
    tabla4225 = doc.add_table(rows=2, cols=14, style='Table Grid')

    for tabla in range(14):
        cell = tabla4225.cell(0, tabla)
        t4225 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t4225.font.name = 'Arial'
        t4225.font.size = Pt(12)
        t4225.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Grafica del capitulo 4.2.2.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo4225 = doc.add_paragraph()
    imagenCapitulo4225.text = ''
    imagenCapitulo4225 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo4225.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo4225.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo4225.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo4225.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.5 ###
    #########################
    tituloGrafico4225 = doc.add_paragraph()
    dgi4225 = tituloGrafico4225.add_run('Grafica 4.4.- Climograma estación CONAGUA.')
    dgi4225_format = tituloGrafico4225.paragraph_format
    dgi4225_format.line_spacing = 1.15
    dgi4225_format.space_after = 0

    dgi4225.font.name = 'Bookman Old Style'
    dgi4225.font.size = Pt(12)
    tituloGrafico4225.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6 ###
    #########################
    capitulo4226 = doc.add_paragraph()
    i4226 = capitulo4226.add_run(f'\n{temasCapitulo4[1][2][2][5]}')
    print(temasCapitulo4[1][2][2][5])
    i4226_format = capitulo4226.paragraph_format
    i4226_format.line_spacing = 1.15

    i4226.font.name = 'Arial'
    i4226.font.size = Pt(12)
    i4226.font.bold = True
    capitulo4226.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6 ###
    #########################
    di4226 = doc.add_paragraph()
    descripcionCapitulo4226 = di4226.add_run('Basado en sus características fisiográficas, geológicas y morfológicas y la ubicación geográfica del sistema ambiental, está en una zona de bajo riesgo ante la ocurrencia de diferentes fenómenos meteorológicos que pueden alterar estructuralmente las condiciones naturales, del área y el proyecto.')
    descripcionCapitulo4226_format = di4226.paragraph_format
    descripcionCapitulo4226_format.line_spacing = 1.15
    descripcionCapitulo4226_format.space_after = 0
    descripcionCapitulo4226_format.space_before = 0

    descripcionCapitulo4226.font.name = 'Arial'
    descripcionCapitulo4226.font.size = Pt(12)
    di4226.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1 ###
    #########################
    capitulo42261 = doc.add_paragraph()
    i42261 = capitulo42261.add_run(f'\n\n{temasCapitulo4[1][2][2][6][0]}')
    i42261_format = capitulo42261.paragraph_format
    i42261_format.line_spacing = 1.5

    i42261.font.name = 'Arial'
    i42261.font.size = Pt(12)
    i42261.font.bold = True
    capitulo42261.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.1 ###
    #########################
    capitulo422611 = doc.add_paragraph()
    i422611 = capitulo422611.add_run(f'\n{temasCapitulo4[1][2][2][6][1][0]}')
    i422611_format = capitulo422611.paragraph_format
    i422611_format.line_spacing = 1.15

    i422611.font.name = 'Arial'
    i422611.font.size = Pt(12)
    i422611.font.bold = True
    capitulo422611.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.1 ###
    #########################
    di422611 = doc.add_paragraph()
    descripcionCapitulo422611 = di422611.add_run('El área que ocupan el sistema ambiental objeto de estudio la cual se encuentran en una ___________________ ante la ocurrencia de este fenómeno en forma severa, ya que se encuentra, según el mapa de distribución de precipitaciones a nivel nacional, ________________________________, por lo que su afectación no sería considerable de acuerdo a las condiciones generales del terreno. (Ver anexo Mapa 4.3.- Precipitación Media)')
    descripcionCapitulo422611_format = di422611.paragraph_format
    descripcionCapitulo422611_format.line_spacing = 1.15
    descripcionCapitulo422611_format.space_after = 0
    descripcionCapitulo422611_format.space_before = 0

    descripcionCapitulo422611.font.name = 'Arial'
    descripcionCapitulo422611.font.size = Pt(12)
    di422611.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.2 ###
    #########################
    capitulo422612 = doc.add_paragraph()
    i422612 = capitulo422612.add_run(f'\n{temasCapitulo4[1][2][2][6][1][1]}')
    i422612_format = capitulo422612.paragraph_format
    i422612_format.line_spacing = 1.15

    i422612.font.name = 'Arial'
    i422612.font.size = Pt(12)
    i422612.font.bold = True
    capitulo422612.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.2 ###
    #########################
    di422612 = doc.add_paragraph()
    descripcionCapitulo422612 = di422612.add_run('El Sistema ambiental donde se realiza el estudio por su ubicación se encuentra en un ____________________________ de ser afectada por fenómenos, de acuerdo al mapa de Riesgo por municipio de granizadas en México del Centro Nacional de Prevención de Desastres que se muestra a continuación. (Ver anexo Mapa 4.4.- Riesgo por Granizada)')
    descripcionCapitulo422612_format = di422612.paragraph_format
    descripcionCapitulo422612_format.line_spacing = 1.15
    descripcionCapitulo422612_format.space_after = 0
    descripcionCapitulo422612_format.space_before = 0

    descripcionCapitulo422612.font.name = 'Arial'
    descripcionCapitulo422612.font.size = Pt(12)
    di422612.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.3 ###
    #########################
    capitulo422613 = doc.add_paragraph()
    i422613 = capitulo422613.add_run(f'\n{temasCapitulo4[1][2][2][6][1][2]}')
    i422613_format = capitulo422613.paragraph_format
    i422613_format.line_spacing = 1.15

    i422613.font.name = 'Arial'
    i422613.font.size = Pt(12)
    i422613.font.bold = True
    capitulo422613.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.3 ###
    #########################
    di422613 = doc.add_paragraph()
    descripcionCapitulo422613 = di422613.add_run('En el siguiente mapa de riesgos de heladas y nevadas podemos observar que el sistema ambiental en la que se pretende realizar las actividades referentes al proyecto, __________________________________________________ que ocurran dichos fenómenos. (Ver anexo Mapa 4.5.- Riesgo por Bajas Temperaturas)')
    descripcionCapitulo422613_format = di422613.paragraph_format
    descripcionCapitulo422613_format.line_spacing = 1.15
    descripcionCapitulo422613_format.space_after = 0
    descripcionCapitulo422613_format.space_before = 0

    descripcionCapitulo422613.font.name = 'Arial'
    descripcionCapitulo422613.font.size = Pt(12)
    di422613.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.4 ###
    #########################
    capitulo422614 = doc.add_paragraph()
    i422614 = capitulo422614.add_run(f'\n{temasCapitulo4[1][2][2][6][1][3]}')
    i422614_format = capitulo422614.paragraph_format
    i422614_format.line_spacing = 1.15

    i422614.font.name = 'Arial'
    i422614.font.size = Pt(12)
    i422614.font.bold = True
    capitulo422614.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.4 ###
    #########################
    di422614 = doc.add_paragraph()
    descripcionCapitulo422614 = di422614.add_run('Según el mapa que se presenta, el área del sistema ambiental en la que se va a realizar el proyecto, _________________________________________________________________.')
    descripcionCapitulo422614_format = di422614.paragraph_format
    descripcionCapitulo422614_format.line_spacing = 1.15
    descripcionCapitulo422614_format.space_after = 0
    descripcionCapitulo422614_format.space_before = 0

    descripcionCapitulo422614.font.name = 'Arial'
    descripcionCapitulo422614.font.size = Pt(12)
    di422614.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.5 ###
    #########################
    capitulo422615 = doc.add_paragraph()
    i422615 = capitulo422615.add_run(f'\n{temasCapitulo4[1][2][2][6][1][4]}')
    i422615_format = capitulo422615.paragraph_format
    i422615_format.line_spacing = 1.15

    i422615.font.name = 'Arial'
    i422615.font.size = Pt(12)
    i422615.font.bold = True
    capitulo422615.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.5 ###
    #########################
    di422615 = doc.add_paragraph()
    descripcionCapitulo422615 = di422615.add_run('El área del sistema ambiental podemos ver que se encuentra en ____________________________________________________, de acuerdo al mapa de riesgo de inundación (Ver anexo Mapa 4.7.- Riesgo por inundación).')
    descripcionCapitulo422615_format = di422615.paragraph_format
    descripcionCapitulo422615_format.line_spacing = 1.15
    descripcionCapitulo422615_format.space_after = 0
    descripcionCapitulo422615_format.space_before = 0

    descripcionCapitulo422615.font.name = 'Arial'
    descripcionCapitulo422615.font.size = Pt(12)
    di422615.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.6 ###
    #########################
    capitulo422616 = doc.add_paragraph()
    i422616 = capitulo422616.add_run(f'\n{temasCapitulo4[1][2][2][6][1][5]}')
    i422616_format = capitulo422616.paragraph_format
    i422616_format.line_spacing = 1.15

    i422616.font.name = 'Arial'
    i422616.font.size = Pt(12)
    i422616.font.bold = True
    capitulo422616.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.6 ###
    #########################
    di422616 = doc.add_paragraph()
    descripcionCapitulo422616 = di422616.add_run('Uno de los grandes riesgos del área del sistema ambiental son las sequías, que provocan el desabasto de agua y afecta el desarrollo económico del área. Como se observa en el siguiente mapa existe ______________________________________________________________________________________________. (Ver anexo Mapa 4.8.- Riesgo por sequias).')
    descripcionCapitulo422616_format = di422616.paragraph_format
    descripcionCapitulo422616_format.line_spacing = 1.15
    descripcionCapitulo422616_format.space_after = 0
    descripcionCapitulo422616_format.space_before = 0

    descripcionCapitulo422616.font.name = 'Arial'
    descripcionCapitulo422616.font.size = Pt(12)
    di422616.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.7 ###
    #########################
    capitulo422617 = doc.add_paragraph()
    i422617 = capitulo422617.add_run(f'\n{temasCapitulo4[1][2][2][6][1][6]}')
    i422617_format = capitulo422617.paragraph_format
    i422617_format.line_spacing = 1.15

    i422617.font.name = 'Arial'
    i422617.font.size = Pt(12)
    i422617.font.bold = True
    capitulo422617.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.7 ###
    #########################
    di422617 = doc.add_paragraph()
    descripcionCapitulo422617 = di422617.add_run('Otro de los fenómenos naturales que se han estado presentando en los últimos años son los tornados, aun cuando no se tiene está área contemplada en los Atlas de riesgo tanto de Protección Civil Estatal como del CENAPRED (Nacional), ante las condiciones extremas que se presentan y las grandes planicies, existe _______ posibilidad de ocurrencia de este fenómeno, es necesario considerar a este fenómeno si bien no tornados, si fuertes vientos en el área, si las condiciones climatológicas son idóneas para este fenómeno, el proyecto requerirá de un monitoreo continuo. (Ver anexo Mapa 4.9.- Riesgo por Tornados)')
    descripcionCapitulo422617_format = di422617.paragraph_format
    descripcionCapitulo422617_format.line_spacing = 1.15
    descripcionCapitulo422617_format.space_after = 0
    descripcionCapitulo422617_format.space_before = 0

    descripcionCapitulo422617.font.name = 'Arial'
    descripcionCapitulo422617.font.size = Pt(12)
    di422617.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.6.1.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.6.1.8 ###
    #########################
    capitulo422618 = doc.add_paragraph()
    i422618 = capitulo422618.add_run(f'\n{temasCapitulo4[1][2][2][6][1][7]}')
    i422618_format = capitulo422618.paragraph_format
    i422618_format.line_spacing = 1.15

    i422618.font.name = 'Arial'
    i422618.font.size = Pt(12)
    i422618.font.bold = True
    capitulo422618.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.6.1.8 ###
    #########################
    di422618 = doc.add_paragraph()
    descripcionCapitulo422618 = di422618.add_run('Aun cuando no se tiene está área contemplada en los Atlas de riesgo tanto de Protección Civil Estatal como del CENAPRED (Nacional), como áreas de riesgo, debe de considerarse los monitoreos continuos para que no se presentes los incendios por descargas eléctricas y afecta la operación del proyecto, ya que son áreas con _______________. (Ver anexo Mapa 4.10.- Riesgo por tormentas eléctricas).')
    descripcionCapitulo422618_format = di422618.paragraph_format
    descripcionCapitulo422618_format.line_spacing = 1.15
    descripcionCapitulo422618_format.space_after = 0
    descripcionCapitulo422618_format.space_before = 0

    descripcionCapitulo422618.font.name = 'Arial'
    descripcionCapitulo422618.font.size = Pt(12)
    di422618.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7 ###
    #########################
    capitulo4227 = doc.add_paragraph()
    i4227 = capitulo4227.add_run(f'\n{temasCapitulo4[1][2][2][7]}')
    i4227_format = capitulo4227.paragraph_format
    i4227_format.line_spacing = 1.15

    i4227.font.name = 'Arial'
    i4227.font.size = Pt(12)
    i4227.font.bold = True
    capitulo4227.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7 ###
    #########################
    di4227 = doc.add_paragraph()
    descripcionCapitulo4227 = di4227.add_run('Para representar los tipos de suelos presentes en el sistema ambiental se utilizó la __________________________________, del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el conjunto de Datos Vectoriales del Continuo Nacional de Efectos Edafológicos escala 1: 250,000, en formato digital, encontrando lo descrito a continuación.')
    descripcionCapitulo4227_format = di4227.paragraph_format
    descripcionCapitulo4227_format.line_spacing = 1.15
    descripcionCapitulo4227_format.space_after = 0
    descripcionCapitulo4227_format.space_before = 0

    descripcionCapitulo4227.font.name = 'Arial'
    descripcionCapitulo4227.font.size = Pt(12)
    di4227.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.7 ###
    #########################
    tituloTabla4227 = doc.add_paragraph()
    dti4227 = tituloTabla4227.add_run('\nTabla 4.7.- Clasificación de suelos en el Sistema Ambiental.')
    dti4227_format = tituloTabla4227.paragraph_format
    dti4227_format.line_spacing = 1.15
    dti4227_format.space_after = 0

    dti4227.font.name = 'Courier New'
    dti4227.font.size = Pt(12)
    tituloTabla4227.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.7 ###
    #########################
    tabla4227 = doc.add_table(rows=8, cols=5, style='Table Grid')

    for tabla in range(5):
        cell = tabla4227.cell(0, tabla)
        t4227 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t4227.font.name = 'Arial'
        t4227.font.size = Pt(12)
        t4227.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Descripcion del capitulo 4.2.2.7 ###
    #########################
    di4227 = doc.add_paragraph()
    descripcionCapitulo4227 = di4227.add_run('\nEl suelo en el sistema ambiental está directamente vinculado a las condiciones topográficas y geomorfológicas, dentro de esta, objeto del presente estudio encontramos una dominancia de los tipos de suelo ________________________________________ respectivamente ambas ambos con _______________________________________________________________. (Ver anexo Mapa 4.11.- Tipos de Suelos).')
    descripcionCapitulo4227_format = di4227.paragraph_format
    descripcionCapitulo4227_format.line_spacing = 1.15
    #descripcionCapitulo4227_format.space_after = 0
    descripcionCapitulo4227_format.space_before = 0

    descripcionCapitulo4227.font.name = 'Arial'
    descripcionCapitulo4227.font.size = Pt(12)
    di4227.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4227 = doc.add_paragraph()
    descripcionCapitulo4227 = di4227.add_run('A continuación, se presenta la descripción de cada uno de los tipos de suelo que se encontraron en el sistema ambiental.')
    descripcionCapitulo4227_format = di4227.paragraph_format
    descripcionCapitulo4227_format.line_spacing = 1.15
    descripcionCapitulo4227_format.space_after = 0
    #descripcionCapitulo4227_format.space_before = 0

    descripcionCapitulo4227.font.name = 'Arial'
    descripcionCapitulo4227.font.size = Pt(12)
    di4227.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.7 ###
    #########################
    tituloTabla4227 = doc.add_paragraph()
    dti4227 = tituloTabla4227.add_run('\nTabla 4.8.- Tipo de suelo presente en el Sistema Ambiental')
    dti4227_format = tituloTabla4227.paragraph_format
    dti4227_format.line_spacing = 1.15
    dti4227_format.space_after = 0

    dti4227.font.name = 'Courier New'
    dti4227.font.size = Pt(12)
    tituloTabla4227.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.7 ###
    #########################
    tabla4227 = doc.add_table(rows=9, cols=3, style='Table Grid')

    for tabla in range(3):
        cell = tabla4227.cell(0, tabla)
        t4227 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t4227.font.name = 'Arial'
        t4227.font.size = Pt(12)
        t4227.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.1 ###
    #########################
    capitulo42271 = doc.add_paragraph()
    i42271 = capitulo42271.add_run(f'\n{temasCapitulo4[1][2][2][8][0]}')
    i42271_format = capitulo42271.paragraph_format
    i42271_format.line_spacing = 1.15

    i42271.font.name = 'Arial'
    i42271.font.size = Pt(12)
    i42271.font.bold = True
    capitulo42271.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.1 ###
    #########################
    di42271 = doc.add_paragraph()
    descripcionCapitulo42271 = di42271.add_run('La palabra erosión proviene del latín erosio o erosionis que signiﬁca: El desgaste que se produce en la superﬁcie del suelo por la acción de agentes externos como el viento y el agua y que son acelerados por la acción del hombre. Es necesario conocer las características de la erosión del suelo para localizar y delimitar sus distintas formas y grados, ubicar con precisión las áreas más afectadas con criterios de campo homogéneos, apoyados en análisis de laboratorio que nos indiquen la calidad de los suelos que son susceptibles. Entender las causas que están provocando o acelerando este proceso, permitirá implementar las medidas de protección y conservación de este recurso. Para el área del sistema ambiental se realizó la caracterización de acuerdo a la carta de erosión del suelo con una escala de 1:250,000 del Instituto Nacional de Estadística, Geografía e Informática (INEGI), las cuales se enlistan y se describen a continuación:')
    descripcionCapitulo42271_format = di42271.paragraph_format
    descripcionCapitulo42271_format.line_spacing = 1.15
    descripcionCapitulo42271_format.space_after = 0
    descripcionCapitulo42271_format.space_before = 0

    descripcionCapitulo42271.font.name = 'Arial'
    descripcionCapitulo42271.font.size = Pt(12)
    di42271.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.7.1 ###
    #########################
    tituloTabla42271 = doc.add_paragraph()
    dti42271 = tituloTabla42271.add_run('\nTabla 4.9.- Erosión presente en el Sistema Ambiental.')
    dti42271_format = tituloTabla42271.paragraph_format
    dti42271_format.line_spacing = 1.15
    dti42271_format.space_after = 0

    dti42271.font.name = 'Courier New'
    dti42271.font.size = Pt(12)
    tituloTabla42271.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.7.1 ###
    #########################
    tabla42271 = doc.add_table(rows=8, cols=6, style='Table Grid')

    for tabla in range(6):
        cell = tabla42271.cell(0, tabla)
        t42271 = cell.paragraphs[0].add_run(f'Col {tabla + 1}')
        t42271.font.name = 'Arial'
        t42271.font.size = Pt(12)
        t42271.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '4F81BD')

    #########################
    ### Descripcion del capitulo 4.2.2.7.1 ###
    #########################
    di42271 = doc.add_paragraph()
    descripcionCapitulo42271 = di42271.add_run('\nEl grado de erosión se conoce midiendo la capa superficial que queda en un predio después de un evento erosivo determinado ya sea lluvia (Hídrica), viento (Eólica) o por actividades Humanas (Antrópica). (Ver anexo Mapa 4.12.- Tipos de Erosión)')
    descripcionCapitulo42271_format = di42271.paragraph_format
    descripcionCapitulo42271_format.line_spacing = 1.15
    descripcionCapitulo42271_format.space_after = 0
    #descripcionCapitulo42271_format.space_before = 0

    descripcionCapitulo42271.font.name = 'Arial'
    descripcionCapitulo42271.font.size = Pt(12)
    di42271.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42271 = doc.add_paragraph()
    descripcionCapitulo42271 = di42271.add_run('\nA continuación, se describen los tipos de erosión presentes en el área del Sistema Ambiental.')
    descripcionCapitulo42271_format = di42271.paragraph_format
    descripcionCapitulo42271_format.line_spacing = 1.15
    descripcionCapitulo42271_format.space_after = 0
    descripcionCapitulo42271_format.space_before = 0

    descripcionCapitulo42271.font.name = 'Arial'
    descripcionCapitulo42271.font.size = Pt(12)
    di42271.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di42271 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42271 = di42271.add_run(f'Tipo de Erosion {lista + 1}.')
        descripcionCapitulo42271_format = di42271.paragraph_format
        descripcionCapitulo42271_format.line_spacing = 1.15
        descripcionCapitulo42271_format.space_after = 0
        descripcionCapitulo42271_format.space_before = 0

        descripcionCapitulo42271.font.name = 'Arial'
        descripcionCapitulo42271.font.size = Pt(12)
        di42271.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.2 ###
    #########################
    capitulo42272 = doc.add_paragraph()
    i42272 = capitulo42272.add_run(f'\n{temasCapitulo4[1][2][2][8][1]}')
    i42272_format = capitulo42272.paragraph_format
    i42272_format.line_spacing = 1.15

    i42272.font.name = 'Arial'
    i42272.font.size = Pt(12)
    i42272.font.bold = True
    capitulo42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('La susceptibilidad de los suelos a erosionarse depende del tamaño de las partículas del suelo, del contenido de materia orgánica, así como de la estructura, en especial del tamaño de los agregados y de la permeabilidad. Para ello se utiliza la Ecuación Universal de Pérdida de Suelo (EUPS), un modelo que permite estimar en campo, la erosión actual y potencial de los suelos. Esta ecuación constituye un instrumento de planeación para establecer las prácticas y obras de conservación de suelos para que hagan que la erosión actual sea menor que la tasa máxima permisible de erosión.')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Para determinar la erosión actual es necesario determinar la protección que le ofrece la cubierta vegetal y la resistencia que oponen las prácticas mecánicas al suelo, de tal forma que si a la ecuación de erosión potencial le incluimos los factores C y P entonces se puede estimar la erosión actual utilizando dicha fórmula quedando como sigue:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nA= R K L S C P')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.bold = True
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nPara estimar “R” en el ámbito regional, se puede utilizar la precipitación anual y con un modelo lineal muy simple estimarlo')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('De acuerdo con los datos registrados en la Estación meteorológica 5003 _______ en el municipio de _______ perteneciente a la CONAGUA, con los datos de registro efectivo de ___________________________________________________________ anuales y considerando que el Estado de Coahuila entra en la región IV de acuerdo a el Mapa de regiones de erosividad de la lluvia en México por lo que el valor de R para el proyecto sería:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('R = 2.8559P + 0.002983P2')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    tabla42272 = doc.add_table(rows=15, cols=3, style='Table Grid')

    for filasRomanos in range(14):
        numero = filasRomanos + 1
        romano = entero_a_romano(numero)
        cell = tabla42272.cell(filasRomanos + 1, 0)
        t42272 = cell.paragraphs[0].add_run(romano)
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ecuacionRegiones = [
        'R = 1.2078P + 0.002276P2',
        'R = 3.4555P + 0.006470P2',
        'R = 3.6752P - 0.001720P2',
        'R = 2.8559P + 0.002983P2',
        'R = 3.4880P - 0.00088P2',
        'R = 6.6847P + 0.001680P2',
        'R = -0.0334P + 0.006661P2 ',
        'R = 1.9967P + 0.003270P2',
        'R = 7.0458P - 0.002096P2',
        'R = 6.8938P + 0.000442P2',
        'R = 3.7745P + 0.004540P2',
        'R = 2.4619P + 0.006067P2',
        'R = 10.7427P - 0.00108P2',
        'R = 1.5005P + 0.002640P2',
    ]

    ecuaciones = range(len(ecuacionRegiones))

    for ecuacion in ecuaciones:
        cell = tabla42272.cell(ecuacion + 1, 1)
        t42272 = cell.paragraphs[0].add_run(f'{ecuacionRegiones[ecuacion]}')
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    radioCuadrado = [
        0.92,
        0.93,
        0.94,
        0.92,
        0.94,
        0.9,
        0.98,
        0.98,
        0.97,
        0.95,
        0.98,
        0.96,
        0.97,
        0.95,
    ]

    radio = range(len(radioCuadrado))

    for cuadrado in radio:
        cell = tabla42272.cell(cuadrado + 1, 2)
        t42272 = cell.paragraphs[0].add_run(f'{radioCuadrado[cuadrado]}')
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Mapa del capitulo 4.2.2.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    mapaCapitulo42272 = doc.add_paragraph()
    mapaCapitulo42272.text = ''
    mapaCapitulo42272 = doc.add_picture('capitulo4/capitulo42272/capitulo4.png')  # Nombre del archivo, tiene que estar en la parte donde se encuentra el documento
    mapaCapitulo42272.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    mapaCapitulo42272.width = Cm(15.59)  # Ancho de la imagen en centimetros
    mapaCapitulo42272.height = Cm(10.16)  # Alto de la imagen en centimetros
    mapaCapitulo42272.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nErosividad Factor (K):')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.bold = True
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('De acuerdo a cuadro de tipo de suelo tenemos que el más representativo, es suelo ________ y que presenta en su mayoría una textura ___________________________ materia orgánica, su valor sería de acuerdo a la tabla siguiente:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Para conocer el valor de K que se obtiene del siguiente cuadro:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    cuadroCapitulo42272_parrafo = doc.add_paragraph()
    cuadroCapitulo42272_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    cuadroCapitulo42272_run = cuadroCapitulo42272_parrafo.add_run()
    imagen = cuadroCapitulo42272_run.add_picture('capitulo4/capitulo42272/tabla42272.png', width=Cm(9.29), height=Cm(8.77))

    # Opcional: espacio después del párrafo
    cuadroCapitulo42272_parrafo.space_after = Pt(0)

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    tabla42272 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla42272.cell(0, 0)
    t42272 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla42272.cell(0, 1)
    t42272 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla42272.cell(1, 0)
    t42272 = cell.paragraphs[0].add_run('A')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell = tabla42272.cell(2, 0)
    t42272 = cell.paragraphs[0].add_run('B')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla42272.cell(3, 0)
    t42272 = cell.paragraphs[0].add_run('C')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla42272.cell(1, 1)
    t42272 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y lo ess poco compactados')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla42272.cell(2, 1)
    t42272 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla42272.cell(3, 1)
    t42272 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien archillas')
    t42272.font.size = Pt(12)
    t42272.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for widht in range(4):
        cell = tabla42272.cell(widht, 0)
        cell.width = Cm(4.1)

    for widht in range(4):
        cell = tabla42272.cell(widht, 1)
        cell.width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nLs= Longitud y grado de pendiente.')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.bold = True
    descripcionCapitulo42272.font.size = Pt(12)
    #di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Este factor considera la longitud y el grado de pendiente. La pendiente media del terreno se obtiene dividiendo la diferencia de elevación del punto más alto del terreno al más bajo entre la longitud del mismo. Esto se obtiene mediante el uso de la siguiente fórmula:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Primer Formula del capitulo 4.2.2.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    formulaCapitulo42272 = doc.add_paragraph()
    formulaCapitulo42272.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo42272 = formulaCapitulo42272.add_run()
    fCapitulo42272.add_picture('capitulo4/capitulo42272/formula_1.png', width=Cm(5.7), height=Cm(1.59))  # Nombre del archivo, tiene que estar en la parte donde se encuentra el documento
    formulaCapitulo42272.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion de la Formula del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Donde:'
                                               '\nS = Pendiente media del terreno (%)'
                                               '\nHf = Altura más alta del terreno (m).'
                                               '\nHi = Altura más baja del terreno (m).'
                                               '\nL = Longitud del terreno (m).\n')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    tablaCapitulo42272 = doc.add_table(rows=6, cols=2, style='Table Grid')

    for cols in range(6):
        for rows in range (2):
            cell = tablaCapitulo42272.cell(rows, cols)
            t42272 = cell.paragraphs[0].add_run(' ')
            t42272.font.size = Pt(12)
            t42272.font.name = 'Arial'
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nPara calcular LS (el factor de longitud y grado de la pendiente) se puede utilizar la siguiente fórmula:')
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Segunda Formula del capitulo 4.2.2.7.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    formulaCapitulo42272 = doc.add_paragraph()
    formulaCapitulo42272.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo42272 = formulaCapitulo42272.add_run()
    fCapitulo42272.add_picture('capitulo4/capitulo42272/formula_2.png', width=Cm(7.91), height=Cm(1.43))  # Nombre del archivo, tiene que estar en la parte donde se encuentra el documento
    formulaCapitulo42272.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion de la Formula del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Donde:'
                                               '\nLS = Factor de longitud y grado de la pendiente.'
                                               '\nλ = Longitud de la pendiente.'
                                               '\nS = Pendiente media del terreno.'
                                               '\nm = Parámetro cuyo valor es 0.5\n')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)

    #########################
    ### Descripcion de la Formula del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Para determinar P y C es necesaria la aplicación de valores a estas constantes en base a prácticas realizadas desde el punto de vista agrícola con lo cual podremos obtener el valor en base a las características del área. ')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Valores de C que se pueden utilizar para estimar pérdidas de suelo.')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    #descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    #descripcionCapitulo42272.bold = True
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    tabla42272 = doc.add_table(cols=4, rows=11, style='Table Grid')

    for cols in range(4):
        for rows in range(11):
            cell = tabla42272.cell(rows, cols)
            t2272 = cell.paragraphs[0].add_run(' ')
            t2272.font.size = Pt(12)
            t2272.font.name = 'Arial'
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('Fuente SAGARPA')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    descripcionCapitulo42272.bold = True
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nValor P:')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    descripcionCapitulo42272.bold = True
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.2 ###
    #########################
    tabla42272 = doc.add_table(cols=2, rows=9, style='Table Grid')

    practicaValorP = [
        'Surcado al contorno',
        'Surcos rectos',
        'Franjas al contorno',
        'Terrazas (2-7% de pendiente ',
        'Terrazas (7-13% de pendiente)',
        'Terrazas mayores de 13%',
        'Terrazas de banco',
        'Terrazas de banco en contrapendiente',
    ]

    valorP = [
        '0.75-0.90',
        '0.80-0.95',
        '0.60-0.80',
        '0.5',
        '0.6',
        '0.8',
        '0.1',
        '0.05',
    ]

    practicaAndValor = ['Practica', 'Valor de P']

    pValorP = range(len(practicaValorP))
    vP = range(len(valorP))
    pAndV = range(len(practicaAndValor))

    for practica in pValorP:
        cell = tabla42272.cell(practica + 1, 0)
        t42272 = cell.paragraphs[0].add_run(f'{practicaValorP[practica]}')
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'

    for valor in vP:
        cell = tabla42272.cell(valor + 1, 1)
        t42272 = cell.paragraphs[0].add_run(f'{valorP[valor]}')
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'
        cell.paragraphs[0].alignment =  WD_ALIGN_PARAGRAPH.RIGHT
    
    for columnas in pAndV:
        cell = tabla42272.cell(0, columnas)
        t42272 = cell.paragraphs[0].add_run(f'{practicaAndValor[columnas]}')
        t42272.font.size = Pt(12)
        t42272.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla42272.cell(1, 1)
    cell_background_color(cell, '92D050')

    #########################
    ### Descripcion del capitulo 4.2.2.7.2 ###
    #########################
    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nFuente SAGARPA:')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    #descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    descripcionCapitulo42272.bold = True
    di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nCuando no se tiene algún tipo de práctica el valor es igual a 0'
                                               '\nCon las variables anteriores sustituyendo valores en la fórmula:')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    #descripcionCapitulo42272.bold = True
    #di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nA=R*K*LS*C*P')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    descripcionCapitulo42272.bold = True
    di42272.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di42272 = doc.add_paragraph()
    descripcionCapitulo42272 = di42272.add_run('\nSe obtiene el grado de erosión potencial en los siguientes supuestos:\n')
    
    descripcionCapitulo42272_format = di42272.paragraph_format
    descripcionCapitulo42272_format.line_spacing = 1.15
    descripcionCapitulo42272_format.space_after = 0
    descripcionCapitulo42272_format.space_before = 0

    descripcionCapitulo42272.font.name = 'Arial'
    descripcionCapitulo42272.font.size = Pt(12)
    #descripcionCapitulo42272.bold = True
    #di42272.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.2.1 ###
    #########################
    capitulo422721 = doc.add_paragraph()
    i422721 = capitulo422721.add_run(f'\n{temasCapitulo4[1][2][2][8][2][0]}')
    i422721_format = capitulo422721.paragraph_format
    i422721_format.line_spacing = 1.15

    i422721.font.name = 'Arial'
    i422721.font.size = Pt(12)
    i422721.font.bold = True
    capitulo422721.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4 ###
    #########################
    di422721 = doc.add_paragraph()
    descripcionCapitulo422721 = di422721.add_run('Se utilizaron los siguientes factores')
    descripcionCapitulo422721_format = di422721.paragraph_format
    descripcionCapitulo422721_format.line_spacing = 1.15
    descripcionCapitulo422721_format.space_after = 0
    descripcionCapitulo422721_format.space_before = 0

    descripcionCapitulo422721.font.name = 'Arial'
    descripcionCapitulo422721.font.size = Pt(12)
    di422721.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422721 = doc.add_paragraph()
    descripcionCapitulo422721 = di422721.add_run('A= R* K* LS* C')
    descripcionCapitulo422721_format = di422721.paragraph_format
    descripcionCapitulo422721_format.line_spacing = 1.15
    descripcionCapitulo422721_format.space_after = 0
    descripcionCapitulo422721_format.space_before = 0

    descripcionCapitulo422721.font.name = 'Arial'
    descripcionCapitulo422721.font.size = Pt(12)
    descripcionCapitulo422721.bold = True
    di422721.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.2.2 ###
    #########################
    capitulo422722 = doc.add_paragraph()
    i422722 = capitulo422722.add_run(f'\n{temasCapitulo4[1][2][2][8][2][1]}')
    i422722_format = capitulo422722.paragraph_format
    i422722_format.line_spacing = 1.5

    i422722.font.name = 'Arial'
    i422722.font.size = Pt(12)
    i422722.font.bold = True
    capitulo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.2.2 ###
    #########################
    di422722 = doc.add_paragraph()
    descripcionCapitulo422722 = di422722.add_run('Para calcular la pérdida de suelo se aplicará la ecuación potencial de acuerdo a la siguiente fórmula utilizando los valores obtenidos de las variables R, K, LS, quedando como sigue:')
    descripcionCapitulo422722_format = di422722.paragraph_format
    descripcionCapitulo422722_format.line_spacing = 1.15
    descripcionCapitulo422722_format.space_after = 0
    descripcionCapitulo422722_format.space_before = 0

    descripcionCapitulo422722.font.name = 'Arial'
    descripcionCapitulo422722.font.size = Pt(12)
    di422722.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422722 = doc.add_paragraph()
    descripcionCapitulo422722 = di422722.add_run('Ep = R*K*LS')
    descripcionCapitulo422722_format = di422722.paragraph_format
    descripcionCapitulo422722_format.line_spacing = 1.15
    descripcionCapitulo422722_format.space_after = 0
    descripcionCapitulo422722_format.space_before = 0

    descripcionCapitulo422722.font.name = 'Arial'
    descripcionCapitulo422722.font.size = Pt(12)
    descripcionCapitulo422721.bold = True
    di422722.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422722 = doc.add_paragraph()
    descripcionCapitulo422722 = di422722.add_run('\nEn resumen, se tiene lo siguiente:')
    descripcionCapitulo422722_format = di422722.paragraph_format
    descripcionCapitulo422722_format.line_spacing = 1.15
    descripcionCapitulo422722_format.space_after = 0
    descripcionCapitulo422722_format.space_before = 0

    descripcionCapitulo422722.font.name = 'Arial'
    descripcionCapitulo422722.font.size = Pt(12)
    descripcionCapitulo422722.bold = True
    di422722.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.7.2.2 ###
    #########################
    tituloTabla422722 = doc.add_paragraph()
    dti422722 = tituloTabla422722.add_run('\nTabla 4.10.- Erosión Hídrica.')
    dti422722_format = tituloTabla422722.paragraph_format
    dti422722_format.line_spacing = 1.15
    dti422722_format.space_after = 0

    dti422722.font.name = 'Courier New'
    dti422722.font.size = Pt(12)
    tituloTabla422722.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.7.2.2 ###
    #########################

    tabla4227222 = doc.add_table(cols=2, rows=2, style='Table Grid')

    for cols in range(2):
        for rows in range(2):
            cell = tabla4227222.cell(rows, cols)
            t422722 = cell.paragraphs[0].add_run(' ')
            t422722.font.size = Pt(12)
            t422722.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 4.2.2.7.2.2 ###
    #########################
    di422722 = doc.add_paragraph()
    descripcionCapitulo422722 = di422722.add_run('De acuerdo a la tabla anterior, el análisis nos arroja que en las condiciones actuales se puede presentar una pérdida de suelo por acción del agua principalmente ____________________________, sin embargo, con la implementación del proyecto al quedar desnudo el suelo, el factor agua erosiona más rápidamente el sistema ambiental, por lo tanto, esta incrementa esta pérdida de suelo hasta un ____________.')
    descripcionCapitulo422722_format = di422722.paragraph_format
    descripcionCapitulo422722_format.line_spacing = 1.15
    descripcionCapitulo422722_format.space_after = 0
    descripcionCapitulo422722_format.space_before = 0

    descripcionCapitulo422722.font.name = 'Arial'
    descripcionCapitulo422722.font.size = Pt(12)
    di422722.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422722 = doc.add_paragraph()
    descripcionCapitulo422722 = di422722.add_run('Erosividad (K o F):')
    descripcionCapitulo422722_format = di422722.paragraph_format
    descripcionCapitulo422722_format.line_spacing = 1.15
    descripcionCapitulo422722_format.space_after = 0
    descripcionCapitulo422722_format.space_before = 0

    descripcionCapitulo422722.font.name = 'Arial'
    descripcionCapitulo422722.font.size = Pt(12)
    di422722.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     #########################
    ### Tabla (imagen) del capitulo 4.2.2.7.2.2 ###
    #########################

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.3 ###
    #########################
    capitulo42273 = doc.add_paragraph()
    i42273 = capitulo42273.add_run(f'\n\n{temasCapitulo4[1][2][2][8][3]}')
    i42273_format = capitulo42273.paragraph_format
    i42273_format.line_spacing = 1.15

    i42273.font.name = 'Arial'
    i42273.font.size = Pt(12)
    i42273.font.bold = True
    capitulo42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('La metodología que se empleó para obtener dichos resultados es la tomada por SAGARPA la cual es la siguiente:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nLa predicción de erosión eólica se puede llevar a cabo por la ecuación desarrollada por Chepil (1963) similar a la propuesta por Wischmeier (1968).')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Xa = (F, G, R, W, V)')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Donde:'
                                               '\nXa = Promedio potencial de erosión anual.'
                                               '\nF = Erodabilidad del suelo'
                                               '\nG = Factor local geográfico para la erosión por viento.'
                                               '\nR = Rugosidad de la superficie del suelo.'
                                               '\nW = Ancho equivalente del campo.'
                                               '\nV = Cantidad equivalente de cubierta vegetal.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nFactores considerados.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('La velocidad del viento es calculada con la suposición de que la velocidad es superior a la necesaria para mover una partícula del suelo. La humedad del suelo es considerada tratando de encontrar que la erodabilidad del suelo por viento es una función de las fuerzas de cohesión del agua alrededor de las partículas. La máxima erosividad por viento se presenta en suelos que contienen menos de 1/3 de la humedad al punto de marchitamiento permanente (PMP), se considera como un suelo secado al aire, sobre este contenido de humedad la erodabilidad decrece hasta el contenido de PMP, hasta cierto punto en donde la erodabilidad decrece al máximo.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nLa ecuación usa agregados mayores a 0.84 mm obtenidos por tamizado en suelo seco.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nFactor Climatico G:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('G = 1/100 i=1∑12 (V3 / 100) (((PET - P) / PET) * n)'
                                                '\nDónde:' 
                                                '\nG = Promedio de la erosión eólica anual.'    
                                                '\nV = Velocidad media mensual a 2 metros de altura, m/s.' 
                                                '\nP = Precipitación pluvial, mm.'   
                                                '\nPET = Evapotranspiración potencial, mm.'  
                                                '\nn = Número de días erosivos por mes.'  
                                                '\nEl número de días sobre el cual la erosión ocurre es asumido que sea proporcional a (PET - P) / PET por el número de días total al mes.'  
                                                '\nPET, puede ser estimado por Penman, Thornthwaite, Blanney, etc.'  
                                                '\n\nPara estimar la erosión eólica del área se utilizaron los siguientes factores')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    #di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    tabla42273 = doc.add_table(cols=2, rows=6, style='Table Grid')

    for cols in range(2):
        for rows in range(2):
            cell = tabla42273.cell(rows, cols)
            t42273 = cell.paragraphs[0].add_run(' ')
            t42273.font.size = Pt(12)
            t42273.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nLos valores mencionados anteriormente se obtienen de la siguiente manera')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nErosividad (K o F):')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    cuadroCapitulo42273_parrafo = doc.add_paragraph()
    cuadroCapitulo42273_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    cuadroCapitulo42273_run = cuadroCapitulo42273_parrafo.add_run('\n')
    imagen = cuadroCapitulo42273_run.add_picture('capitulo4/capitulo42272/tabla42272.png', width=Cm(9.29), height=Cm(8.77))

    # Opcional: espacio después del párrafo
    cuadroCapitulo42273_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nPara conocer el valor de K que se obtiene del siguiente cuadro:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    tabla42273 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla42273.cell(0, 0)
    t42273 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    t42273.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla42273.cell(0, 1)
    t42273 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    t42273.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla42273.cell(1, 0)
    t42273 = cell.paragraphs[0].add_run('A')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell = tabla42273.cell(2, 0)
    t42273 = cell.paragraphs[0].add_run('B')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla42273.cell(3, 0)
    t42273 = cell.paragraphs[0].add_run('C')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla42273.cell(1, 1)
    t42273 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y lo ess poco compactados')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla42273.cell(2, 1)
    t42273 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla42273.cell(3, 1)
    t42273 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien archillas')
    t42273.font.size = Pt(12)
    t42273.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for widht in range(4):
        cell = tabla42273.cell(widht, 0)
        cell.width = Cm(4.1)

    for widht in range(4):
        cell = tabla42273.cell(widht, 1)
        cell.width = Cm(13.09)
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nFactor G:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    #descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Para obtener el factor G se utilizó información meteorológica de CONAGUA, denominada “5003 del municipio de Ramos Arizpe, Coahuila, con ello y utilizando la fórmula factor climático G = 1/100 i=1∑12 (V3 / 100) (((PET - P) / PET) * n)')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    #descripcionCapitulo42273_format.space_after = 0
    #descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    tabla42273 = doc.add_table(rows=4, cols=3, style='Table Grid')

    for cols in range(3):
        for rows in range(4):
            cell = tabla42273.cell(rows, cols)
            t42273 = cell.paragraphs[0].add_run(' ')
            t42273.font.size = Pt(12)
            t42273.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nSustituyendo la fórmula se obtiene lo siguiente:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    #descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\n\nFactor R:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    #descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Este considera la rugosidad del terreno la cual está, influenciado por el tipo de suelo específicamente en el tamaño granular de las partículas, sabiendo que el tipo de suelo presente en el área de cambio de uso de suelo es _____________________________________________________________________.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\n\nFactor W:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    #descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Este factor contempla la distancia de afectación del área (ancho del terreno en estudio en metros).')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Valores del Factor V o C que se pueden utilizar para estimar pérdidas de suelo.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    #descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    tabla42273 = doc.add_table(rows=15, cols=4, style='Table Grid')

    for cols in range(4):
        for rows in range(15):
            cell = tabla42273.cell(rows, cols)
            t42273 = cell.paragraphs[0].add_run(' ')
            t42273.font.size = Pt(12)
            t42273.font.name = 'Arial'
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Fuente SAGARPA')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    #descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nValor P:')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    #descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.7.3 ###
    #########################
    tabla42273 = doc.add_table(rows=8, cols=2, style='Table Grid')

    for cols in range(2):
        for rows in range(8):
            cell = tabla42273.cell(rows, cols)
            t42273 = cell.paragraphs[0].add_run(' ')
            t42273.font.size = Pt(12)
            t42273.font.name = 'Arial'
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.3 ###
    #########################
    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('Fuente SAGARPA')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    #descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42273 = doc.add_paragraph()
    descripcionCapitulo42273 = di42273.add_run('\nCabe hacer mención que, para estimar la erosión eólica, para escenarios con proyecto y con medidas de mitigación uno de los factores que influye en los resultados es la velocidad del viento y el factor de prácticas de manejo.')
    descripcionCapitulo42273_format = di42273.paragraph_format
    descripcionCapitulo42273_format.line_spacing = 1.15
    descripcionCapitulo42273_format.space_after = 0
    descripcionCapitulo42273_format.space_before = 0

    descripcionCapitulo42273.font.name = 'Arial'
    descripcionCapitulo42273.font.size = Pt(12)
    descripcionCapitulo42273.bold = True
    di42273.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.3.1 ###
    #########################
    capitulo422731 = doc.add_paragraph()
    i422731 = capitulo422731.add_run(f'\n\n{temasCapitulo4[1][2][2][8][4][0]}')
    i422731_format = capitulo422731.paragraph_format
    i422731_format.line_spacing = 1.15

    i422731.font.name = 'Arial'
    i422731.font.size = Pt(12)
    i422731.font.bold = True
    capitulo422731.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.3.1 ###
    #########################
    di422731 = doc.add_paragraph()
    descripcionCapitulo422731 = di422731.add_run('Utilizando la información anterior y la ecuación se tiene lo siguiente')
    descripcionCapitulo422731_format = di422731.paragraph_format
    descripcionCapitulo422731_format.line_spacing = 1.15
    descripcionCapitulo422731_format.space_after = 0
    descripcionCapitulo422731_format.space_before = 0

    descripcionCapitulo422731.font.name = 'Arial'
    descripcionCapitulo422731.font.size = Pt(12)
    di422731.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422731 = doc.add_paragraph()
    descripcionCapitulo422731 = di422731.add_run('Xa = (F, G, R, W, V)')
    descripcionCapitulo422731_format = di422731.paragraph_format
    descripcionCapitulo422731_format.line_spacing = 1.15
    descripcionCapitulo422731_format.space_after = 0
    descripcionCapitulo422731_format.space_before = 0

    descripcionCapitulo422731.font.name = 'Arial'
    descripcionCapitulo422731.font.size = Pt(12)
    descripcionCapitulo422731.bold = True
    di422731.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4 ###
    #########################
    capitulo422732 = doc.add_paragraph()
    i422732 = capitulo422732.add_run(f'\n\n{temasCapitulo4[1][2][2][8][4][1]}')
    i422732_format = capitulo422732.paragraph_format
    i422732_format.line_spacing = 1.15

    i422732.font.name = 'Arial'
    i422732.font.size = Pt(12)
    i422732.font.bold = True
    capitulo422732.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4 ###
    #########################
    di422732 = doc.add_paragraph()
    descripcionCapitulo422732 = di422732.add_run('X a = (F, G, R, W)  ')
    descripcionCapitulo422732_format = di422732.paragraph_format
    descripcionCapitulo422732_format.line_spacing = 1.15
    descripcionCapitulo422732_format.space_after = 0
    descripcionCapitulo422732_format.space_before = 0

    descripcionCapitulo422732.font.name = 'Arial'
    descripcionCapitulo422732.font.size = Pt(12)
    descripcionCapitulo422732.bold = True
    di422732.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.3.3 ###
    #########################
    capitulo422733 = doc.add_paragraph()
    i422733 = capitulo422733.add_run(f'\n{temasCapitulo4[1][2][2][8][4][2]}')
    i422733_format = capitulo422733.paragraph_format
    i422733_format.line_spacing = 1.15

    i422733.font.name = 'Arial'
    i422733.font.size = Pt(12)
    i422733.font.bold = True
    capitulo422733.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.3.3 ###
    #########################
    di422733 = doc.add_paragraph()
    descripcionCapitulo422733 = di422733.add_run('X a = (F, G, R, W, V, P)')
    descripcionCapitulo422733_format = di422733.paragraph_format
    descripcionCapitulo422733_format.line_spacing = 1.15
    descripcionCapitulo422733_format.space_after = 0
    descripcionCapitulo422733_format.space_before = 0

    descripcionCapitulo422733.font.name = 'Arial'
    descripcionCapitulo422733.font.size = Pt(12)
    descripcionCapitulo422733.bold = True
    di422733.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.7.3.3 ###
    #########################
    tituloTabla422733 = doc.add_paragraph()
    dti422733 = tituloTabla422733.add_run('\nTabla 4.11.- Erosión Eolica.')
    dti422733_format = tituloTabla422733.paragraph_format
    dti422733_format.line_spacing = 1.15
    dti422733_format.space_after = 0

    dti422733.font.name = 'Courier New'
    dti422733.font.size = Pt(12)
    tituloTabla422733.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.7.3.3 ###
    #########################

    tabla422733 = doc.add_table(cols=2, rows=2, style='Table Grid')

    for cols in range(2):
        for rows in range(2):
            cell = tabla422733.cell(rows, cols)
            t422733 = cell.paragraphs[0].add_run(' ')
            t422733.font.size = Pt(12)
            t422733.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    
    #########################
    ### Descripcion del capitulo 4.2.2.7.3.3 ###
    #########################
    di422733 = doc.add_paragraph()
    descripcionCapitulo422733 = di422733.add_run('\nAl analizar los datos de la tabla anterior podemos observar que en las condiciones actuales se puede presentar una pérdida de suelo por acción del viento de ___________, sin embargo, al quedar desnuda toda la superficie del área el factor viento erosiona más rápidamente incrementando una pérdida de suelo de hasta ____________, con medidas de mitigación es de ___________.')
    descripcionCapitulo422733_format = di422733.paragraph_format
    descripcionCapitulo422733_format.line_spacing = 1.15
    descripcionCapitulo422733_format.space_after = 0
    descripcionCapitulo422733_format.space_before = 0

    descripcionCapitulo422733.font.name = 'Arial'
    descripcionCapitulo422733.font.size = Pt(12)
    di422733.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.7.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.7.4 ###
    #########################
    capitulo4 = doc.add_paragraph()
    i4 = capitulo4.add_run(f'\n{temasCapitulo4[1][2][2][8][5]}')
    i4_format = capitulo4.paragraph_format
    i4_format.line_spacing = 1.15

    i4.font.name = 'Arial'
    i4.font.size = Pt(12)
    i4.font.bold = True
    capitulo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.7.4 ###
    #########################
    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('Posibilidades de uso Agrícola.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('Describir la posibilidad.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    #descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('\nPosibilidades de uso Pecuario.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('Describir la posibilidad.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    #descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('\nPosibilidades de uso Forestal.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42274 = doc.add_paragraph()
    descripcionCapitulo42274 = di42274.add_run('Describir la posibilidad.')
    descripcionCapitulo42274_format = di42274.paragraph_format
    descripcionCapitulo42274_format.line_spacing = 1.15
    descripcionCapitulo42274_format.space_after = 0
    descripcionCapitulo42274_format.space_before = 0

    descripcionCapitulo42274.font.name = 'Arial'
    descripcionCapitulo42274.font.size = Pt(12)
    #descripcionCapitulo42274.bold = True
    di42274.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.8 ###
    #########################
    capitulo4228 = doc.add_paragraph()
    i4228 = capitulo4228.add_run(f'\n{temasCapitulo4[1][2][2][9]}')
    i4228_format = capitulo4228.paragraph_format
    i4228_format.line_spacing = 1.15

    i4228.font.name = 'Arial'
    i4228.font.size = Pt(12)
    i4228.font.bold = True
    capitulo4228.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.8 ###
    #########################
    di4228 = doc.add_paragraph()
    descripcionCapitulo4228 = di4228.add_run('El área de estudio se ubica dentro de las ____________________________________________________________, las que a su vez están dentro de provincia _____________________, tectono estratigráfico denominado “Coahuila“, para determinar el tipo de roca existente dentro del sistema ambiental utilizamos la __________________________________, del INEGI y el conjunto de datos vectoriales del continuo nacional de efectos geológicos escala 1:250,000, en formato digital, encontrando que los tipos de roca existentes pertenecen a las eras geológicas, ________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo4228_format = di4228.paragraph_format
    descripcionCapitulo4228_format.line_spacing = 1.15
    descripcionCapitulo4228_format.space_after = 0
    descripcionCapitulo4228_format.space_before = 0

    descripcionCapitulo4228.font.name = 'Arial'
    descripcionCapitulo4228.font.size = Pt(12)
    di4228.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4228 = doc.add_paragraph()
    descripcionCapitulo4228 = di4228.add_run('A continuación, en la siguiente tabla se enlistan, las rocas existentes dentro del sistema ambiental:')
    descripcionCapitulo4228_format = di4228.paragraph_format
    descripcionCapitulo4228_format.line_spacing = 1.15
    descripcionCapitulo4228_format.space_after = 0
    descripcionCapitulo4228_format.space_before = 0

    descripcionCapitulo4228.font.name = 'Arial'
    descripcionCapitulo4228.font.size = Pt(12)
    di4228.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.8 ###
    #########################
    tituloTabla4228 = doc.add_paragraph()
    dti4228 = tituloTabla4228.add_run('\nTabla 4.12.- Tipos de Rocas en el Sistema Ambiental.')
    dti4228_format = tituloTabla4228.paragraph_format
    dti4228_format.line_spacing = 1.15
    dti4228_format.space_after = 0

    dti4228.font.name = 'Courier New'
    dti4228.font.size = Pt(12)
    tituloTabla4228.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.8 ###
    #########################

    tabla4228 = doc.add_table(cols=6, rows=7, style='Table Grid')

    for cols in range(6):
        for rows in range(7):
            cell = tabla4228.cell(rows, cols)
            t4228 = cell.paragraphs[0].add_run(' ')
            t4228.font.size = Pt(12)
            t4228.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla4228.cell(0, cols)
        cell_background_color(cell, '4F81BD')
        
    #########################
    ### Descripcion del capitulo 4.2.2.8 ###
    #########################
    di4228 = doc.add_paragraph()
    descripcionCapitulo4228 = di4228.add_run('A continuación, se describen cada uno de los tipos de roca encontrados dentro del área de estudio.')
    descripcionCapitulo4228_format = di4228.paragraph_format
    descripcionCapitulo4228_format.line_spacing = 1.15
    #descripcionCapitulo4228_format.space_after = 0
    #descripcionCapitulo4228_format.space_before = 0

    descripcionCapitulo4228.font.name = 'Arial'
    descripcionCapitulo4228.font.size = Pt(12)
    di4228.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(5):
        di4228 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo4228 = di4228.add_run(f'Elemento {lista + 1}')
        descripcionCapitulo4228_format = di4228.paragraph_format
        descripcionCapitulo4228_format.line_spacing = 1.15
        descripcionCapitulo4228_format.space_after = 0
        descripcionCapitulo4228_format.space_before = 0

        descripcionCapitulo4228.font.name = 'Arial'
        descripcionCapitulo4228.font.size = Pt(12)
        descripcionCapitulo4228.font.bold = True
        di4228.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.9
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.9 ###
    #########################
    capitulo4229 = doc.add_paragraph()
    i4229 = capitulo4229.add_run(f'\n{temasCapitulo4[1][2][2][10]}')
    i4229_format = capitulo4229.paragraph_format
    i4229_format.line_spacing = 1.15

    i4229.font.name = 'Arial'
    i4229.font.size = Pt(12)
    i4229.font.bold = True
    capitulo4229.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.9 ###
    #########################
    di4229 = doc.add_paragraph()
    descripcionCapitulo4229 = di4229.add_run('Las provincias fisiográficas son regiones en que el relieve es el resultado de la acción de un mismo conjunto geológico, un mismo o muy semejante tipo de suelo y de vegetación que sustenta. El área del sistema ambiental objeto del presente estudio se localiza dentro de la Provincia Fisiográfica __________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo4229_format = di4229.paragraph_format
    descripcionCapitulo4229_format.line_spacing = 1.15
    descripcionCapitulo4229_format.space_after = 0
    descripcionCapitulo4229_format.space_before = 0

    descripcionCapitulo4229.font.name = 'Arial'
    descripcionCapitulo4229.font.size = Pt(12)
    di4229.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.9.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.9.1 ###
    #########################
    capitulo42291 = doc.add_paragraph()
    i42291 = capitulo42291.add_run(f'\n{temasCapitulo4[1][2][2][11][0]}')
    i42291_format = capitulo42291.paragraph_format
    i42291_format.line_spacing = 1.15

    i42291.font.name = 'Arial'
    i42291.font.size = Pt(12)
    i42291.font.bold = True
    capitulo42291.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.9.1 ###
    #########################
    di42291 = doc.add_paragraph()
    descripcionCapitulo42291 = di42291.add_run('Descripcion del capitulo 4.2.2.9.1.')
    descripcionCapitulo42291_format = di42291.paragraph_format
    descripcionCapitulo42291_format.line_spacing = 1.15
    descripcionCapitulo42291_format.space_after = 0
    descripcionCapitulo42291_format.space_before = 0

    descripcionCapitulo42291.font.name = 'Arial'
    descripcionCapitulo42291.font.size = Pt(12)
    di42291.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.10
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.10 ###
    #########################
    capitulo422_10 = doc.add_paragraph()
    i422_10 = capitulo422_10.add_run(f'\n{temasCapitulo4[1][2][2][12]}')
    i422_10_format = capitulo422_10.paragraph_format
    i422_10_format.line_spacing = 1.15

    i422_10.font.name = 'Arial'
    i422_10.font.size = Pt(12)
    i422_10.font.bold = True
    capitulo422_10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.10 ###
    #########################
    di422_10 = doc.add_paragraph()
    descripcionCapitulo422_10 = di422_10.add_run('La República Mexicana se encuentra dividida en cuatro zonas sísmicas. Para realizar esta división se utilizaron los catálogos de sismos de la República Mexicana. Estas zonas son un reflejo de que tan frecuentes son los sismos en las diversas regiones y la máxima aceleración del suelo a esperar durante un siglo.')
    descripcionCapitulo422_10_format = di422_10.paragraph_format
    descripcionCapitulo422_10_format.line_spacing = 1.15
    descripcionCapitulo422_10_format.space_after = 0
    descripcionCapitulo422_10_format.space_before = 0

    descripcionCapitulo422_10.font.name = 'Arial'
    descripcionCapitulo422_10.font.size = Pt(12)
    di422_10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_10 = doc.add_paragraph()
    descripcionCapitulo422_10 = di422_10.add_run('Describir el resto del capitulo... .(Ver anexo Mapa 4.16.- Riesgo de Sismos).')
    descripcionCapitulo422_10_format = di422_10.paragraph_format
    descripcionCapitulo422_10_format.line_spacing = 1.15
    descripcionCapitulo422_10_format.space_after = 0
    descripcionCapitulo422_10_format.space_before = 0

    descripcionCapitulo422_10.font.name = 'Arial'
    descripcionCapitulo422_10.font.size = Pt(12)
    di422_10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.11
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.11 ###
    #########################
    capitulo422_11 = doc.add_paragraph()
    i422_11 = capitulo422_11.add_run(f'\n{temasCapitulo4[1][2][2][13]}')
    i422_11_format = capitulo422_11.paragraph_format
    i422_11_format.line_spacing = 1.15

    i422_11.font.name = 'Arial'
    i422_11.font.size = Pt(12)
    i422_11.font.bold = True
    capitulo422_11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.11 ###
    #########################
    di422_11 = doc.add_paragraph()
    descripcionCapitulo422_11 = di422_11.add_run('La representación gráfica de un terreno es importante, el inventario de la infraestructura, Orografía, Hidrografía y de la población de un lugar, así como de su distribución geográfica; en ellas se registra fielmente todos estos factores y las relaciones que guardan entre sí, es así mismo, la base en la cual se sustentan los estudios que se ocupan del inventario de los recursos naturales como los de Geología, Edafología, uso del suelo y vegetación e Hidrología, entre otros. De esta manera se utilizaron las cartas topográficas __________________________________________________, del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el Conjunto Nacional de Datos Vectoriales escala 1: 50,000 y el Continuo Nacional de Elevaciones especificando el área de la misma carta topográfica, la altura mínima registrada en el área del sistema ambiental es de ____________________________________________________________________________, El área de estudio se encuentra en su mayoría en una topografía de ___________________________________________________________________. (Ver anexo Mapa 4.17.- Tipos de Topoformas). Las topoformas más comunes se enlistan continuación.')
    descripcionCapitulo422_11_format = di422_11.paragraph_format
    descripcionCapitulo422_11_format.line_spacing = 1.15
    descripcionCapitulo422_11_format.space_after = 0
    descripcionCapitulo422_11_format.space_before = 0

    descripcionCapitulo422_11.font.name = 'Arial'
    descripcionCapitulo422_11.font.size = Pt(12)
    di422_11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11 ###
    #########################
    tituloTabla422_11 = doc.add_paragraph()
    dti422_11 = tituloTabla422_11.add_run('\nTabla 4.13.- Porcentaje de las topoformas del Sistema Ambiental.')
    dti422_11_format = tituloTabla422_11.paragraph_format
    dti422_11_format.line_spacing = 1.15
    dti422_11_format.space_after = 0

    dti422_11.font.name = 'Courier New'
    dti422_11.font.size = Pt(12)
    tituloTabla422_11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11 ###
    #########################

    tabla422_11 = doc.add_table(cols=4, rows=5, style='Table Grid')

    for cols in range(4):
        for rows in range(5):
            cell = tabla422_11.cell(rows, cols)
            t422_11 = cell.paragraphs[0].add_run(' ')
            t422_11.font.size = Pt(12)
            t422_11.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_11.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    #########################
    ### Descripcion del capitulo 4.2.2.11 ###
    #########################
    di422_11 = doc.add_paragraph()
    descripcionCapitulo422_11 = di422_11.add_run('\nA continuación, se describe la topoforma en la cual se encuentra el área sujeta al estudio del proyecto. ')
    descripcionCapitulo422_11_format = di422_11.paragraph_format
    descripcionCapitulo422_11_format.line_spacing = 1.15
    descripcionCapitulo422_11_format.space_after = 0
    descripcionCapitulo422_11_format.space_before = 0

    descripcionCapitulo422_11.font.name = 'Arial'
    descripcionCapitulo422_11.font.size = Pt(12)
    di422_11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11 ###
    #########################
    tituloTabla422_11 = doc.add_paragraph()
    dti422_11 = tituloTabla422_11.add_run('\nTabla 4.14.- Tipo de topo formas del Sistema Ambiental.')
    dti422_11_format = tituloTabla422_11.paragraph_format
    dti422_11_format.line_spacing = 1.15
    dti422_11_format.space_after = 0

    dti422_11.font.name = 'Courier New'
    dti422_11.font.size = Pt(12)
    tituloTabla422_11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11 ###
    #########################

    tabla422_11 = doc.add_table(cols=2, rows=5, style='Table Grid')

    for cols in range(2):
        for rows in range(5):
            cell = tabla422_11.cell(rows, cols)
            t422_11 = cell.paragraphs[0].add_run(' ')
            t422_11.font.size = Pt(12)
            t422_11.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_11.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.11.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.11.1 ###
    #########################
    capitulo422_111 = doc.add_paragraph()
    i422_111 = capitulo422_111.add_run(f'\n{temasCapitulo4[1][2][2][14][0]}')
    i422_111_format = capitulo422_111.paragraph_format
    i422_111_format.line_spacing = 1.15

    i422_111.font.name = 'Arial'
    i422_111.font.size = Pt(12)
    i422_111.font.bold = True
    capitulo422_111.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.11.1 ###
    #########################
    di422_111 = doc.add_paragraph()
    descripcionCapitulo422_111 = di422_111.add_run('Como podemos observar en la tabla de las pendientes de mayor dominancia son elevaciones de ____________________, mientras que las menos dominantes son las que vas de _____________________________________________________________________________. (Ver anexo Mapa 4.18.- Elevaciones del SA).')
    descripcionCapitulo422_111_format = di422_111.paragraph_format
    descripcionCapitulo422_111_format.line_spacing = 1.15
    descripcionCapitulo422_111_format.space_after = 0
    descripcionCapitulo422_111_format.space_before = 0

    descripcionCapitulo422_111.font.name = 'Arial'
    descripcionCapitulo422_111.font.size = Pt(12)
    di422_111.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11.1 ###
    #########################
    tituloTabla422_111 = doc.add_paragraph()
    dti422_111 = tituloTabla422_111.add_run('\nTabla 4.15.-	Tipo de elevaciones del área en estudio.')
    dti422_111_format = tituloTabla422_111.paragraph_format
    dti422_111_format.line_spacing = 1.15
    dti422_111_format.space_after = 0

    dti422_111.font.name = 'Courier New'
    dti422_111.font.size = Pt(12)
    tituloTabla422_111.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11.1 ###
    #########################

    tabla422_111 = doc.add_table(cols=4, rows=15, style='Table Grid')

    for cols in range(4):
        for rows in range(15):
            cell = tabla422_111.cell(rows, cols)
            t422_111 = cell.paragraphs[0].add_run(' ')
            t422_111.font.size = Pt(12)
            t422_111.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_111.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.11.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.11.2 ###
    #########################
    capitulo422_112 = doc.add_paragraph()
    i422_112 = capitulo422_112.add_run(f'\n{temasCapitulo4[1][2][2][14][1]}')
    i422_112_format = capitulo422_112.paragraph_format
    i422_112_format.line_spacing = 1.15

    i422_112.font.name = 'Arial'
    i422_112.font.size = Pt(12)
    i422_112.font.bold = True
    capitulo422_112.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.11.2 ###
    #########################
    di422_112 = doc.add_paragraph()
    descripcionCapitulo422_112 = di422_112.add_run('El Sistema Ambiental se encuentra en una topoforma de mayoría _______________________________________________________, (Ver anexo Mapa 4.19.- Pendientes del SA), a continuación, se presenta en la siguiente tabla.')
    descripcionCapitulo422_112_format = di422_112.paragraph_format
    descripcionCapitulo422_112_format.line_spacing = 1.15
    descripcionCapitulo422_112_format.space_after = 0
    descripcionCapitulo422_112_format.space_before = 0

    descripcionCapitulo422_112.font.name = 'Arial'
    descripcionCapitulo422_112.font.size = Pt(12)
    di422_112.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11.2 ###
    #########################
    tituloTabla422_112 = doc.add_paragraph()
    dti422_112 = tituloTabla422_112.add_run('\nTabla 4.16.-	Pendientes presentes en el Sistema Ambiental.')
    dti422_112_format = tituloTabla422_112.paragraph_format
    dti422_112_format.line_spacing = 1.15
    dti422_112_format.space_after = 0

    dti422_112.font.name = 'Courier New'
    dti422_112.font.size = Pt(12)
    tituloTabla422_112.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11.2 ###
    #########################

    tabla422_112 = doc.add_table(cols=4, rows=15, style='Table Grid')

    for cols in range(4):
        for rows in range(15):
            cell = tabla422_112.cell(rows, cols)
            t422_112 = cell.paragraphs[0].add_run(' ')
            t422_112.font.size = Pt(12)
            t422_112.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_112.cell(0, cols)
        cell_background_color(cell, '4F81BD')
    
    ########################################################################################################################################################################
    # Capitulo 4.2.2.11.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.11.3 ###
    #########################
    capitulo422_113 = doc.add_paragraph()
    i422_113 = capitulo422_113.add_run(f'\n{temasCapitulo4[1][2][2][14][2]}')
    i422_113_format = capitulo422_113.paragraph_format
    i422_113_format.line_spacing = 1.15

    i422_113.font.name = 'Arial'
    i422_113.font.size = Pt(12)
    i422_113.font.bold = True
    capitulo422_113.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.11.3 ###
    #########################
    di422_113 = doc.add_paragraph()
    descripcionCapitulo422_113 = di422_113.add_run('Para el caso de la exposición en el sistema ambiental, la más abundantes y por las características de las mismas podemos encontrar en mayor proporción la exposición _________________________________________, (Ver anexo Mapa 4.20.- Exposición), como se muestra en la siguiente tabla.')
    descripcionCapitulo422_113_format = di422_113.paragraph_format
    descripcionCapitulo422_113_format.line_spacing = 1.15
    descripcionCapitulo422_113_format.space_after = 0
    descripcionCapitulo422_113_format.space_before = 0

    descripcionCapitulo422_113.font.name = 'Arial'
    descripcionCapitulo422_113.font.size = Pt(12)
    di422_113.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11.3 ###
    #########################
    tituloTabla422_113 = doc.add_paragraph()
    dti422_113 = tituloTabla422_113.add_run('\nTabla 4.17.-	Tipos de exposición dentro del Sistema Ambiental')
    dti422_113_format = tituloTabla422_113.paragraph_format
    dti422_113_format.line_spacing = 1.15
    dti422_113_format.space_after = 0

    dti422_113.font.name = 'Courier New'
    dti422_113.font.size = Pt(12)
    tituloTabla422_113.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11.3 ###
    #########################

    tabla422_113 = doc.add_table(cols=4, rows=7, style='Table Grid')

    for cols in range(4):
        for rows in range(7):
            cell = tabla422_113.cell(rows, cols)
            t422_113 = cell.paragraphs[0].add_run(' ')
            t422_113.font.size = Pt(12)
            t422_113.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_113.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.11.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.11.4 ###
    #########################
    capitulo422_114 = doc.add_paragraph()
    i422_114 = capitulo422_114.add_run(f'\n{temasCapitulo4[1][2][2][14][3]}')
    i422_114_format = capitulo422_114.paragraph_format
    i422_114_format.line_spacing = 1.15

    i422_114.font.name = 'Arial'
    i422_114.font.size = Pt(12)
    i422_114.font.bold = True
    capitulo422_114.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.11.4 ###
    #########################
    di422_114 = doc.add_paragraph()
    descripcionCapitulo422_114 = di422_114.add_run('Dentro de este sistema ambiental podemos encontrar toponimias que van desde lugares, cañones cañadas, sierras, cerros, puertos, así como puntos orográficos que son importantes para su referencia de estudio dentro del predio y área de estudio como se mencionan a continuación')
    descripcionCapitulo422_114_format = di422_114.paragraph_format
    descripcionCapitulo422_114_format.line_spacing = 1.15
    descripcionCapitulo422_114_format.space_after = 0
    descripcionCapitulo422_114_format.space_before = 0

    descripcionCapitulo422_114.font.name = 'Arial'
    descripcionCapitulo422_114.font.size = Pt(12)
    di422_114.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.11.4 ###
    #########################
    tituloTabla422_114 = doc.add_paragraph()
    dti422_114 = tituloTabla422_114.add_run('\nTabla 4.18.-	Principales toponimias del Sistema Ambiental.')
    dti422_114_format = tituloTabla422_114.paragraph_format
    dti422_114_format.line_spacing = 1.15
    dti422_114_format.space_after = 0

    dti422_114.font.name = 'Courier New'
    dti422_114.font.size = Pt(12)
    tituloTabla422_114.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.11.4 ###
    #########################

    tabla422_114 = doc.add_table(cols=3, rows=15, style='Table Grid')

    for cols in range(3):
        for rows in range(7):
            cell = tabla422_114.cell(rows, cols)
            t422_114 = cell.paragraphs[0].add_run(' ')
            t422_114.font.size = Pt(12)
            t422_114.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_114.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.12
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12 ###
    #########################
    capitulo422_12 = doc.add_paragraph()
    i422_12 = capitulo422_12.add_run(f'\n{temasCapitulo4[1][2][2][15]}')
    i422_12_format = capitulo422_12.paragraph_format
    i422_12_format.line_spacing = 1.15

    i422_12.font.name = 'Arial'
    i422_12.font.size = Pt(12)
    i422_12.font.bold = True
    capitulo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12 ###
    #########################
    di422_12 = doc.add_paragraph()
    descripcionCapitulo422_12 = di422_12.add_run('El sistema ambiental objeto del presente estudio se localiza dentro de la Subcuenca .... Descripcion del Capitulo 4.2.2.12 ...... (Ver anexo Mapa 4.22.- Geohidrología del sistema ambiental).')
    descripcionCapitulo422_12_format = di422_12.paragraph_format
    descripcionCapitulo422_12_format.line_spacing = 1.15
    descripcionCapitulo422_12_format.space_after = 0
    descripcionCapitulo422_12_format.space_before = 0

    descripcionCapitulo422_12.font.name = 'Arial'
    descripcionCapitulo422_12.font.size = Pt(12)
    di422_12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.12 ###
    #########################
    tituloTabla422_12 = doc.add_paragraph()
    dti422_12 = tituloTabla422_12.add_run('\nTabla 4.19.-	Geohidrología presente en el Sistema Ambiental.')
    dti422_12_format = tituloTabla422_12.paragraph_format
    dti422_12_format.line_spacing = 1.15
    dti422_12_format.space_after = 0

    dti422_12.font.name = 'Courier New'
    dti422_12.font.size = Pt(12)
    tituloTabla422_12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.12 ###
    #########################

    tabla422_12 = doc.add_table(cols=4, rows=4, style='Table Grid')

    for cols in range(4):
        for rows in range(4):
            cell = tabla422_12.cell(rows, cols)
            t422_12 = cell.paragraphs[0].add_run(' ')
            t422_12.font.size = Pt(12)
            t422_12.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_12.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    ########################################################################################################################################################################
    # Capitulo 4.2.2.12.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12.1 ###
    #########################
    capitulo422_121 = doc.add_paragraph()
    i422_121 = capitulo422_121.add_run(f'\n{temasCapitulo4[1][2][2][16][0]}')
    i422_121_format = capitulo422_121.paragraph_format
    i422_121_format.line_spacing = 1.15

    i422_121.font.name = 'Arial'
    i422_121.font.size = Pt(12)
    i422_121.font.bold = True
    capitulo422_121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.1 ###
    #########################
    di422_121 = doc.add_paragraph()
    descripcionCapitulo422_121 = di422_121.add_run('Descripción del capitulo 4.2.2.12.1')
    descripcionCapitulo422_121_format = di422_121.paragraph_format
    descripcionCapitulo422_121_format.line_spacing = 1.15
    descripcionCapitulo422_121_format.space_after = 0
    descripcionCapitulo422_121_format.space_before = 0

    descripcionCapitulo422_121.font.name = 'Arial'
    descripcionCapitulo422_121.font.size = Pt(12)
    di422_121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.12.1 ###
    #########################
    tituloTabla422_121 = doc.add_paragraph()
    dti422_121 = tituloTabla422_121.add_run('\nTabla 4.X.-	Geohidrología presente en el Sistema Ambiental.')
    dti422_121_format = tituloTabla422_121.paragraph_format
    dti422_121_format.line_spacing = 1.15
    dti422_121_format.space_after = 0

    dti422_121.font.name = 'Courier New'
    dti422_121.font.size = Pt(12)
    tituloTabla422_121.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.12.1 ###
    #########################

    tabla422_121 = doc.add_table(cols=10, rows=4, style='Table Grid')

    for cols in range(10):
        for rows in range(4):
            cell = tabla422_121.cell(rows, cols)
            t422_121 = cell.paragraphs[0].add_run(' ')
            t422_121.font.size = Pt(12)
            t422_121.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

        cell = tabla422_121.cell(0, cols)
        cell_background_color(cell, '4F81BD')

    #########################
    ### Descripcion de la tabla del capitulo 4.2.2.12.1 ###
    #########################
    tituloTabla422_121 = doc.add_paragraph()
    dti422_121 = tituloTabla422_121.add_run('Descripcion de la tabla 4.X.-	Geohidrología presente en el Sistema Ambiental.')
    dti422_121_format = tituloTabla422_121.paragraph_format
    dti422_121_format.line_spacing = 1.15
    dti422_121_format.space_after = 0

    dti422_121.font.name = 'Times New Roman'
    dti422_121.font.size = Pt(10)
    tituloTabla422_121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.1 ###
    #########################
    di422_121 = doc.add_paragraph()
    descripcionCapitulo422_121 = di422_121.add_run('Describir el resto del capitulo (incluir mapas, imagenes, tablas y formulas)')
    descripcionCapitulo422_121_format = di422_121.paragraph_format
    descripcionCapitulo422_121_format.line_spacing = 1.15
    descripcionCapitulo422_121_format.space_after = 0
    descripcionCapitulo422_121_format.space_before = 0

    descripcionCapitulo422_121.font.name = 'Arial'
    descripcionCapitulo422_121.font.size = Pt(12)
    di422_121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.2.12.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12.2 ###
    #########################
    capitulo422_122 = doc.add_paragraph()
    i422_122 = capitulo422_122.add_run(f'\n{temasCapitulo4[1][2][2][16][1]}')
    i422_122_format = capitulo422_122.paragraph_format
    i422_122_format.line_spacing = 1.15

    i422_122.font.name = 'Arial'
    i422_122.font.size = Pt(12)
    i422_122.font.bold = True
    capitulo422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.2 ###
    #########################
    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('Para el análisis de la infiltración del recurso agua en el Sistema Ambiental donde se encuentra el proyecto de cambio de uso de suelo se realizaron los siguientes cálculos.')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nVolumen de Escurrimiento = Precipitación Anual * Área Total * Coeficiente de Escurrimiento\n')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.2.12.2 ###
    #########################

    tabla422_122 = doc.add_table(cols=2, rows=3, style='Table Grid')

    cell = tabla422_122.cell(0, 0)
    t422_122 = cell.paragraphs[0].add_run('COEFICIENTE DE ESCURRIMIENTO ANUAL (Ce)')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    t422_122.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(0, 0)
    cell_background_color(cell, '4F81BD')

    cell = tabla422_122.cell(0, 1)
    t422_122 = cell.paragraphs[0].add_run('K: PARAMETRO QUE DEPENDE DEL TIPO Y USO DE SUELO')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    t422_122.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    
    cell = tabla422_122.cell(0, 1)
    cell_background_color(cell, '4F81BD')

    cell = tabla422_122.cell(1, 0)
    t422_122 = cell.paragraphs[0].add_run('Ce = K(P-250) / 2000')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Times New Roman'
    t422_122.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(1, 1)
    t422_122 = cell.paragraphs[0].add_run('Si K resulta menor o igual que 0.15')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(2, 0)
    t422_122 = cell.paragraphs[0].add_run('Ce = (K(P-250) / 2000) + (K - 0.15) / 1.5')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Times New Roman'
    t422_122.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(2, 1)
    t422_122 = cell.paragraphs[0].add_run('Si K es mayor que 0.15')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in range(2):
        cell = tabla422_122.cell(0, cols)
        cell_background_color(cell, '4F81BD')

        for rows in range(3):
            cell = tabla422_122.cell(rows, cols)
            cell.height = Cm(1.22)
            cell.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    #########################
    ### Descripcion de la tabla del capitulo 4.2.2.12.2 ###
    #########################
    tituloTabla422_122 = doc.add_paragraph()
    dti422_122 = tituloTabla422_122.add_run('Fuente: SAGARPA')
    dti422_122_format = tituloTabla422_122.paragraph_format
    dti422_122_format.line_spacing = 1.15
    dti422_122_format.space_after = 0

    dti422_122.font.name = 'Arial'
    dti422_122.font.size = Pt(12)
    dti422_122.font.bold = True
    tituloTabla422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.2 ###
    #########################
    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nInfiltración = P - ETR - Ve')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nP = Precipitación'
                                                   '\nETR = Evapotranspiración'
                                                   '\nVe = Volumen de escurrimiento')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    #di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nEVATRANSPIRACION POR EL METODO DE COUTAGNE')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    #di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nETR = P - xP2')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nP = Precipitación'
                                                   '\nX = Constante (1/ (0.8 + 0.14 t)'
                                                   '\nt= Temperatura promedio _____')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    #di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nVe= (P) (At) (Ce)')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nP = Precipitación anual en m3'
                                                   '\nAt = Área total del Sistema Ambiental km2'
                                                   '\nCe = Coeficiente de escurrimiento'
                                                   '\nK = Constante de erosividad'
                                                   '\nP = Precipitación')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    #di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nPara obtener los valores se toma en cuenta la precipitación promedio registrada, la temperatura promedio, la superficie total del Sistema Ambiental. ________________________________________________.\n')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.12.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    cuadroCapitulo422_122_parrafo = doc.add_paragraph()
    cuadroCapitulo422_122_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    cuadroCapitulo422_122_run = cuadroCapitulo422_122_parrafo.add_run('\n')
    imagen = cuadroCapitulo422_122_run.add_picture('capitulo4/capitulo42272/tabla42272.png', width=Cm(9.29), height=Cm(8.77))

    # Opcional: espacio después del párrafo
    cuadroCapitulo422_122_parrafo.space_after = Pt(1)

    #########################
    ### Tabla del capitulo 4.2.2.12.2 ###
    #########################
    tabla422_122 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla422_122.cell(0, 0)
    t422_122 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    t422_122.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla422_122.cell(0, 1)
    t422_122 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    t422_122.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla422_122.cell(1, 0)
    t422_122 = cell.paragraphs[0].add_run('A')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    cell = tabla422_122.cell(2, 0)
    t422_122 = cell.paragraphs[0].add_run('B')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(3, 0)
    t422_122 = cell.paragraphs[0].add_run('C')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla422_122.cell(1, 1)
    t422_122 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y lo ess poco compactados')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla422_122.cell(2, 1)
    t422_122 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla422_122.cell(3, 1)
    t422_122 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien archillas')
    t422_122.font.size = Pt(12)
    t422_122.font.name = 'Arial'
    #t42272.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for widht in range(4):
        cell = tabla422_122.cell(widht, 0)
        cell.width = Cm(4.1)

    for widht in range(4):
        cell = tabla422_122.cell(widht, 1)
        cell.width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 4.2.2.12.2 ###
    #########################
    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nCe=K (P-250)/2000+ (K-0.15)/1.5')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nCon los datos necesarios calculados se podrá obtener el grado de infiltración que presenta el Sistema Ambiental, donde se encuentra el área de Cambio de Uso del Suelo desde tres escenarios, tal y como se manifiesta a continuación.')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nPor lo anterior el volumen medio anual de escurrimiento natural se determinó mediante el método indirecto, mediante la siguiente expresión:')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Arial'
    #descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(12)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_122 = doc.add_paragraph()
    descripcionCapitulo422_122 = di422_122.add_run('\nVolumen Anual de Escurrimiento = Precipitación Anual * Área Total km2 * Coeficiente de Escurrimiento')
    descripcionCapitulo422_122_format = di422_122.paragraph_format
    descripcionCapitulo422_122_format.line_spacing = 1.15
    descripcionCapitulo422_122_format.space_after = 0
    descripcionCapitulo422_122_format.space_before = 0

    descripcionCapitulo422_122.font.name = 'Times New Roman'
    descripcionCapitulo422_122.italic = True
    descripcionCapitulo422_122.font.size = Pt(11)
    di422_122.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.2.12.2 ###
    #########################
    tabla422_122 = doc.add_table(rows=2, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla422_122.cell(0, cols)
        cell_background_color(cell, '0070C0')
        for rows in range(2):
            cell = tabla422_122.cell(rows, cols)
            t422_122 = cell.paragraphs[0].add_run(' ')
            t422_122.font.size = Pt(12)
            t422_122.font.name = 'Arial'



    ########################################################################################################################################################################
    # Capitulo 4.2.2.12.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12.2.1 ###
    #########################
    capitulo422_1221 = doc.add_paragraph()
    i422_1221 = capitulo422_1221.add_run(f'\n{temasCapitulo4[1][2][2][16][2][0]}')
    i422_1221_format = capitulo422_1221.paragraph_format
    i422_1221_format.line_spacing = 1.15

    i422_1221.font.name = 'Arial'
    i422_1221.font.size = Pt(12)
    i422_1221.font.bold = True
    capitulo422_1221.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.2.1 ###
    #########################
    di422_1221 = doc.add_paragraph()
    descripcionCapitulo422_1221 = di422_1221.add_run('INFILTRACIÓN')
    descripcionCapitulo422_1221_format = di422_1221.paragraph_format
    descripcionCapitulo422_1221_format.line_spacing = 1.15
    descripcionCapitulo422_1221_format.space_after = 0
    descripcionCapitulo422_1221_format.space_before = 0

    descripcionCapitulo422_1221.font.name = 'Arial'
    descripcionCapitulo422_1221.font.size = Pt(12)
    di422_1221.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_1221 = doc.add_paragraph()
    descripcionCapitulo422_1221 = di422_1221.add_run('Infiltración = P - ETR - Ve')
    descripcionCapitulo422_1221_format = di422_1221.paragraph_format
    descripcionCapitulo422_1221_format.line_spacing = 1.15
    descripcionCapitulo422_1221_format.space_after = 0
    descripcionCapitulo422_1221_format.space_before = 0

    descripcionCapitulo422_1221.font.name = 'Times New Roman'
    descripcionCapitulo422_1221.font.size = Pt(12)
    descripcionCapitulo422_1221.italic = True
    di422_1221.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.12.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12.2.2 ###
    #########################
    capitulo422_1222 = doc.add_paragraph()
    i422_1222 = capitulo422_1222.add_run(f'\n{temasCapitulo4[1][2][2][16][2][1]}')
    i422_1222_format = capitulo422_1222.paragraph_format
    i422_1222_format.line_spacing = 1.15

    i422_1222.font.name = 'Arial'
    i422_1222.font.size = Pt(12)
    i422_1222.font.bold = True
    capitulo422_1222.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.2.12.2.2 ###
    #########################
    di422_1222 = doc.add_paragraph()
    descripcionCapitulo422_1222 = di422_1222.add_run('INFILTRACIÓN')
    descripcionCapitulo422_1222_format = di422_1222.paragraph_format
    descripcionCapitulo422_1222_format.line_spacing = 1.15
    descripcionCapitulo422_1222_format.space_after = 0
    descripcionCapitulo422_1222_format.space_before = 0

    descripcionCapitulo422_1222.font.name = 'Arial'
    descripcionCapitulo422_1222.font.size = Pt(12)
    di422_1222.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di422_1221 = doc.add_paragraph()
    descripcionCapitulo422_1222 = di422_1222.add_run('Infiltración = P - ETR - Ve')
    descripcionCapitulo422_1222_format = di422_1221.paragraph_format
    descripcionCapitulo422_1222_format.line_spacing = 1.15
    descripcionCapitulo422_1222_format.space_after = 0
    descripcionCapitulo422_1222_format.space_before = 0

    descripcionCapitulo422_1222.font.name = 'Times New Roman'
    descripcionCapitulo422_1222.font.size = Pt(12)
    descripcionCapitulo422_1222.italic = True
    di422_1222.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 4.2.2.12.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.2.12.2.3 ###
    #########################
    capitulo422_1223 = doc.add_paragraph()
    i422_1223 = capitulo422_1223.add_run(f'{temasCapitulo4[0]}')
    i422_1223_format = capitulo422_1223.paragraph_format
    i422_1223_format.line_spacing = 1.15

    i422_1223.font.name = 'Arial'
    i422_1223.font.size = Pt(12)
    i422_1223.font.bold = True
    capitulo422_1223.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.2.12.2.3 ###
    #########################
    tituloTabla422_1223 = doc.add_paragraph()
    dti422_1223 = tituloTabla422_1223.add_run('\nTabla 4.x.- Infiltración en el área de estudio.')
    dti422_1223_format = tituloTabla422_1223.paragraph_format
    dti422_1223_format.line_spacing = 1.15
    dti422_1223_format.space_after = 0

    dti422_1223.font.name = 'Courier New'
    dti422_1223.font.size = Pt(12)
    tituloTabla422_1223.alignment = WD_ALIGN_PARAGRAPH.CENTER

     #########################
    ### Tabla del capitulo 4.2.2.12.2.3 ###
    #########################
    tabla422_122 = doc.add_table(rows=3, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla422_122.cell(0, cols)
        cell_background_color(cell, '0070C0')
        for rows in range(3):
            cell = tabla422_122.cell(rows, cols)
            t422_112 = cell.paragraphs[0].add_run(' ')
            t422_112.font.size = Pt(12)
            t422_112.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.2.12.2.3 ###
    #########################
    di422_1223 = doc.add_paragraph()
    descripcionCapitulo422_1223 = di422_1223.add_run('\nDe acuerdo a la tabla anterior se puede observar que de los tres escenarios que se plantean en el factor de infiltración se puede mencionar que en las condiciones actuales del sistema ambiental se infiltra normalmente ________ de agua, en condiciones naturales, al quedar desnudo el sistema ambiental se incrementa la evaporación, por lo tanto, la infiltración es menor, con un ________________.')
    descripcionCapitulo422_1223_format = di422_1223.paragraph_format
    descripcionCapitulo422_1223_format.line_spacing = 1.15
    descripcionCapitulo422_1223_format.space_after = 0
    descripcionCapitulo422_1223_format.space_before = 0

    descripcionCapitulo422_1223.font.name = 'Arial'
    descripcionCapitulo422_1223.font.size = Pt(12)
    di422_1223.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3 ###
    #########################
    capitulo423 = doc.add_paragraph()
    i423 = capitulo423.add_run(f'\n\n{temasCapitulo4[1][2][3]}')
    i423_format = capitulo423.paragraph_format
    i423_format.line_spacing = 1.15

    i423.font.name = 'Arial'
    i423.font.size = Pt(12)
    i423.font.bold = True
    capitulo423.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.1 ###
    #########################
    capitulo4231 = doc.add_paragraph()
    i4231 = capitulo4231.add_run(f'{temasCapitulo4[1][2][4][0]}')
    i4231_format = capitulo4231.paragraph_format
    i4231_format.line_spacing = 1.15

    i4231.font.name = 'Arial'
    i4231.font.size = Pt(12)
    i4231.font.bold = True
    capitulo4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.1 ###
    #########################
    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('Según Rzedowski (1978) ... ... ... ... .. . . . . . . Descripcion.')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)

    di4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('Los elementos integrantes de los recursos básicos. Los cambios en la cobertura y uso del suelo afectan los sistemas globales (por ejemplo atmósfera, clima y nivel del mar), dichos cambios ocurren en un modo localizado que en su conjunto llegan a sumar un total significativo y se reflejan en buena medida en la cobertura vegetal, razón por la cual se toman como referencia para algunas aplicaciones que van desde el monitoreo ambiental, la producción de estadísticas como apoyo en la planeación, evaluación del cambio climático y la evaluación de los procesos de desertificación, entre otros. ')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)
    di4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('Partiendo de este punto se analizan los diferentes ecosistemas para realizar una buena planeación de acuerdo a su uso, por tal motivo para la obtención de esta información se utilizaron las cartas Uso de Suelo y Vegetación, en su serie VII, que son las más actuales, en este caso es la siguientes, _________________, del Instituto Nacional de Estadística, Geografía e Informática (INEGI), se utilizó el Conjunto de Nacional de Datos Vectoriales de Vegetación escala 1: 250 000, los cuales se enlistas y se presentan a continuación.')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)
    di4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('Las Comunidades vegetales que se desarrollan en el sistema ambiental se clasificaron con base al criterio de Henrickson y Johnston (1983), por lo cual ... ... ... ... ... Descripcion ... ... ... ... ..... .')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)
    di4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('Los elementos integrantes de los recursos básicos y los cambios en la cobertura y uso del suelo afectan los sistemas globales, dichos cambios ocurren en un modo localizado que en su conjunto llegan a sumar un total significativo y se reflejan en buena medida en la cobertura vegetal, razón por la cual se toman como referencia para algunas aplicaciones que van desde el monitoreo ambiental, la producción de estadísticas como apoyo a la planeación, evaluación del cambio climático y la evaluación de los procesos de desertificación entre otros, se analizan los diferentes tipos de vegetación para realizar una buena planeación, para la obtención de esta información se utilizaron la cobertura de vegetación del INEGI en las cartas ______________________ ________________________ ____________________ ______________________ _________________ _______, (Ver anexo Mapa 4.23.- Tipos de Vegetación), los tipos de vegetación encontrados se enlistan y se describen a continuación.')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)
    di4231.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.1 ###
    #########################
    tituloTabla4231 = doc.add_paragraph()
    dti4231 = tituloTabla4231.add_run('\nTabla 4.x.- Tipos de vegetacion en el Sistema Ambientals.')
    dti4231_format = tituloTabla4231.paragraph_format
    dti4231_format.line_spacing = 1.15
    dti4231_format.space_after = 0

    dti4231.font.name = 'Courier New'
    dti4231.font.size = Pt(12)
    tituloTabla4231.alignment = WD_ALIGN_PARAGRAPH.CENTER

     #########################
    ### Tabla del capitulo 4.2.3.1 ###
    #########################
    tabla4231 = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla4231.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla4231.cell(rows, cols)
            t4231 = cell.paragraphs[0].add_run(' ')
            t4231.font.size = Pt(12)
            t4231.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.1 ###
    #########################
    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('\nA continuación, se describe cada uno de los tipos de vegetación encontrado dentro del sistema ambiental:')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)

    di4231 = doc.add_paragraph()
    descripcionCapitulo4231 = di4231.add_run('\nTipos de vegetacion')
    descripcionCapitulo4231_format = di4231.paragraph_format
    descripcionCapitulo4231_format.line_spacing = 1.15
    descripcionCapitulo4231_format.space_after = 0
    descripcionCapitulo4231_format.space_before = 0

    descripcionCapitulo4231.font.name = 'Arial'
    descripcionCapitulo4231.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 4.2.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.2 ###
    #########################
    capitulo4232 = doc.add_paragraph()
    i4232 = capitulo4232.add_run(f'\n{temasCapitulo4[1][2][4][1]}')
    i4232_format = capitulo4232.paragraph_format
    i4232_format.line_spacing = 1.5

    i4232.font.name = 'Arial'
    i4232.font.size = Pt(12)
    i4232.font.bold = True
    capitulo4232.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.2 ###
    #########################
    di4232 = doc.add_paragraph()
    descripcionCapitulo4232 = di4232.add_run('A continuación, se presenta el análisis de la información de las especies de flora encontradas durante el levantamiento de los sitios de muestreo en el sistema ambiental, donde se encuentra el área objeto de cambio de uso de suelo, determinando así su importancia y su posición dentro de los estratos encontrados, de acuerdo al inventario realizado.')
    descripcionCapitulo4232_format = di4232.paragraph_format
    descripcionCapitulo4232_format.line_spacing = 1.15
    descripcionCapitulo4232_format.space_after = 0
    descripcionCapitulo4232_format.space_before = 0

    descripcionCapitulo4232.font.name = 'Arial'
    descripcionCapitulo4232.font.size = Pt(12)
    di4232.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.2 ###
    #########################
    tituloTabla4232 = doc.add_paragraph()
    dti4232 = tituloTabla4232.add_run('\nTabla 4.x.- Estatus de las Especies por Estrato.')
    dti4232_format = tituloTabla4232.paragraph_format
    dti4232_format.line_spacing = 1.15
    dti4232_format.space_after = 0

    dti4232.font.name = 'Courier New'
    dti4232.font.size = Pt(12)
    tituloTabla4232.alignment = WD_ALIGN_PARAGRAPH.CENTER

     #########################
    ### Tabla del capitulo 4.2.3.2 ###
    #########################
    tabla4232 = doc.add_table(rows=40, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla4232.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla4232.cell(rows, cols)
            t4232 = cell.paragraphs[0].add_run(' ')
            t4232.font.size = Pt(12)
            t4232.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.3 ###
    #########################
    capitulo4233 = doc.add_paragraph()
    i4233 = capitulo4233.add_run(f'\n{temasCapitulo4[1][2][4][2]}')
    i4233_format = capitulo4233.paragraph_format
    i4233_format.line_spacing = 1.15

    i4233.font.name = 'Arial'
    i4233.font.size = Pt(12)
    i4233.font.bold = True
    capitulo4233.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.3 ###
    #########################
    di4233 = doc.add_paragraph()
    descripcionCapitulo4233 = di4233.add_run('\nDiseño de Muestreo.')
    descripcionCapitulo4233_format = di4233.paragraph_format
    descripcionCapitulo4233_format.line_spacing = 1.15
    descripcionCapitulo4233_format.space_after = 0
    descripcionCapitulo4233_format.space_before = 0

    descripcionCapitulo4233.font.name = 'Arial'
    descripcionCapitulo4233.font.size = Pt(12)
    descripcionCapitulo4233.font.bold = True
    di4233.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4233 = doc.add_paragraph()
    descripcionCapitulo4233 = di4233.add_run('\nPara realizar el muestreo dentro del sistema ambiental y realizar un comparativo de los tipos de vegetación existentes en Sistema Ambiental, así como la diversidad de las especies, para ello se realizó el muestreo con base al tipo de vegetación que se encuentran existentes en el sistema ambiental, __________________________________________. Con el objeto de que la información fuera confiable en la toma de datos en campo, se realizó un sistema de muestreo aleatorio y lo más cercano al área de cambio de uso de suelo y en donde no presentara disturbios en los estratos. Con respeto a la forma y tamaño de los sitios de muestreo se utilizó de forma circular con un radio de 17.84 m., en una superficie de 1,000 m2 esto para el estrato arbóreo, para los estratos arbustivo, y suculento fue con un radio de 8.92 m en una superficie de 250 m2 mientras que, para las herbácea y gramíneas, el muestreo fue de 1 m2 cuadrado, quedando en el centro del sitio. ')
    descripcionCapitulo4233_format = di4233.paragraph_format
    descripcionCapitulo4233_format.line_spacing = 1.15
    descripcionCapitulo4233_format.space_after = 0
    descripcionCapitulo4233_format.space_before = 0

    descripcionCapitulo4233.font.name = 'Arial'
    descripcionCapitulo4233.font.size = Pt(12)
    di4233.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    cuadroCapitulo4233_parrafo = doc.add_paragraph()
    cuadroCapitulo4233_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    cuadroCapitulo4233_run = cuadroCapitulo4233_parrafo.add_run('\n')
    imagen = cuadroCapitulo4233_run.add_picture('capitulo4/capitulo4233/grafico.png', width=Cm(6.04), height=Cm(5.05))

    # Opcional: espacio después del párrafo
    cuadroCapitulo4233_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.3 ###
    #########################
    di4233 = doc.add_paragraph()
    descripcionCapitulo4233 = di4233.add_run('\nCon los muestreos realizados en el sistema ambiental se pudo analizar la condición de la vegetación del área del proyecto, en un área aislada donde no se tiene alteración por algún tipo de uso. Se concluyó que tanto en el área propuesta para el ACUSTF y el sistema ambiental se conserva la diversidad florística de las especies por lo cual al efectuar el procedimiento de eliminación vegetación del área propuesta se mantendrá la Biodiversidad en el sistema ambiental. (Ver anexo Mapa 4.24.- Muestreo del sistema ambiental).')
    descripcionCapitulo4233_format = di4233.paragraph_format
    descripcionCapitulo4233_format.line_spacing = 1.15
    descripcionCapitulo4233_format.space_after = 0
    descripcionCapitulo4233_format.space_before = 0

    descripcionCapitulo4233.font.name = 'Arial'
    descripcionCapitulo4233.font.size = Pt(12)
    #descripcionCapitulo4233.font.bold = True
    di4233.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4233 = doc.add_paragraph()
    descripcionCapitulo4233 = di4233.add_run('\nA continuación, se enlistan las coordenadas de cada sitio de muestreo donde se recabó la información para realizar la comparación, la cuales se encuentran en _______________________________.')
    descripcionCapitulo4233_format = di4233.paragraph_format
    descripcionCapitulo4233_format.line_spacing = 1.15
    descripcionCapitulo4233_format.space_after = 0
    descripcionCapitulo4233_format.space_before = 0

    descripcionCapitulo4233.font.name = 'Arial'
    descripcionCapitulo4233.font.size = Pt(12)
    #descripcionCapitulo4233.font.bold = True
    di4233.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.3 ###
    #########################
    tituloTabla4233 = doc.add_paragraph()
    dti4233 = tituloTabla4233.add_run('\nTabla 4.x.- Coordenadas greográficas de sitios de muestreo del ___')
    dti4233_format = tituloTabla4233.paragraph_format
    dti4233_format.line_spacing = 1.15
    dti4233_format.space_after = 0

    dti4233.font.name = 'Courier New'
    dti4233.font.size = Pt(12)
    tituloTabla4233.alignment = WD_ALIGN_PARAGRAPH.CENTER

     #########################
    ### Tabla del capitulo 4.2.3.3 ###
    #########################
    tabla4233 = doc.add_table(rows=40, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla4233.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla4233.cell(rows, cols)
            t4233 = cell.paragraphs[0].add_run(' ')
            t4233.font.size = Pt(12)
            t4233.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.3.1 ###
    #########################
    capitulo42331 = doc.add_paragraph()
    i42331 = capitulo42331.add_run(f'\n{temasCapitulo4[1][2][4][3][0]}')
    i42331_format = capitulo42331.paragraph_format
    i42331_format.line_spacing = 1.15

    i42331.font.name = 'Arial'
    i42331.font.size = Pt(12)
    i42331.font.bold = True
    capitulo42331.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.3.1 ###
    #########################
    di42331 = doc.add_paragraph()
    descripcionCapitulo42331 = di42331.add_run('\nEn la presente tabla se presenta la vegetación encontrada en el sistema ambiental, por sitios de muestreo, donde se puede observar parámetros como altura, cobertura, así como la cantidad de individuos por especie.')
    descripcionCapitulo42331_format = di42331.paragraph_format
    descripcionCapitulo42331_format.line_spacing = 1.15
    descripcionCapitulo42331_format.space_after = 0
    descripcionCapitulo42331_format.space_before = 0

    descripcionCapitulo42331.font.name = 'Arial'
    descripcionCapitulo42331.font.size = Pt(12)
    descripcionCapitulo42331.font.bold = True
    di42331.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.3.1 ###
    #########################
    tituloTabla4233 = doc.add_paragraph()
    dti4233 = tituloTabla4233.add_run('\nTabla 4.x.- Coordenadas greográficas de sitios de muestreo del ___')
    dti4233_format = tituloTabla4233.paragraph_format
    dti4233_format.line_spacing = 1.15
    dti4233_format.space_after = 0

    dti4233.font.name = 'Courier New'
    dti4233.font.size = Pt(12)
    tituloTabla4233.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.3.1 ###
    #########################
    tabla42331 = doc.add_table(rows=40, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42331.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42331.cell(rows, cols)
            t42331 = cell.paragraphs[0].add_run(' ')
            t42331.font.size = Pt(12)
            t42331.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4 ###
    #########################
    capitulo4234 = doc.add_paragraph()
    i4234 = capitulo4234.add_run(f'\n{temasCapitulo4[1][2][4][4]}')
    i4234_format = capitulo4234.paragraph_format
    i4234_format.line_spacing = 1.15

    i4234.font.name = 'Arial'
    i4234.font.size = Pt(12)
    i4234.font.bold = True
    capitulo4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('La metodología utilizada para determinar la Biodiversidad presente en el sistema ambiental, es a través de la diversidad Alfa, para diferenciarlos en función de las variables biológicas que miden, se dividen en dos grandes grupos:')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    listaCapitulo4234 = [
        'Métodos basados en la cuantificación del número de especies presentes (riqueza específica).',
        'Métodos basados en la estructura de la comunidad, es decir, la distribución proporcional del valor de importancia de cada especie (abundancia relativa de los individuos, su biomasa, cobertura, productividad, etc.).',
    ]

    lista4234 = range(len(listaCapitulo4234))

    for lista in lista4234:
        di4234 = doc.add_paragraph(style='List Number')
        descripcionCapitulo4234 = di4234.add_run(f'{listaCapitulo4234[lista]}')
        descripcionCapitulo4234_format = di4234.paragraph_format
        descripcionCapitulo4234_format.line_spacing = 1.15
        descripcionCapitulo4234_format.space_after = 0
        descripcionCapitulo4234_format.space_before = 0

        descripcionCapitulo4234.font.name = 'Arial'
        descripcionCapitulo4234.font.size = Pt(12)
        di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nCon la información obtenida de los estratos se calcularon los atributos de la vegetación, tales como densidad, dominancia y frecuencia de las especies dentro de la vegetación, consecuentemente se obtuvo el Índice de valor de importancia (IVI).')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    #descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Para realizar los cálculos de los índices y parámetros estructurales se emplearon las siguientes fórmulas:')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    #descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Densidad Absoluta. Está dada por el número de individuos de una especie o de todas las especies dividido por el número de sitios muestreados.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    #descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_1.png', width=Cm(2.01), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nD = Densidad'
                                             '\nN = Número de individuos muestreados por especie'
                                             '\nA = número de sitios muestreados o superficie muestrea según sea (x sito, ha o ACUSTF)')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nDensidad relativa. Está dada por el resultado de la densidad absoluta entre el número total de todos los individuos muestreados expresados en porcentajes ')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_2.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nDer = Densidad Relativa'
                                             '\nNi = Número de individuos de la especie'
                                             '\nNt = Número total de individuos de todas las especies')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nDominancia absoluta. Se define como el porcentaje de biomasa (área basal o superficie horizontal) que aporta una especie.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_3.png', width=Cm(2.75), height=Cm(1.5))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nDa = Densidad absoluta'
                                             '\nABi = Área basal de una especie'
                                             '\nA = Área muestreada (sitios muestreados)')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nLa dominancia relativa. Se calcula como la proporción de una especie en el área total evaluada, expresada en porcentaje.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_4.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nDor = Densidad relativa'
                                             '\nDai = Densidad absoluta de una especie'
                                             '\nDat= Densidad absoluta total de todas las especies')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nFrecuencia absoluta. Permite conocer las veces que se repite una especie en cada sitio de muestreo. ')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_5.png', width=Cm(3.89), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nFa = Frecuencia absoluta'
                                             '\nnsi = sumatoria del número de veces que una especie se observa dentro de todos los sitios de muestreo.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nFrecuencia relativa. Es el resultado de dividir la frecuencia absoluta de cada especie entre el número total de esas especies expresadas en porcentajes.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_6.png', width=Cm(4.74), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nFr = Frecuencia relativa'
                                             '\nFai = Frecuencia absoluta de cada especie'
                                             '\nFat = Frecuencia absoluta de todas las especies')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de valor de importancia (IVI). El índice de valor de importancia define cuáles de las especies presentes contribuyen en el carácter y estructura de una Comunidad. Este valor se obtiene mediante la sumatoria de la frecuencia relativa, la densidad relativa y la dominancia relativa.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_7.png', width=Cm(4.99), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nIVI = Índice de Valor de Importancia'
                                             '\nDer = Densidad relativa'
                                             '\nDor = Dominancia relativa'
                                             '\nFr = Frecuencia relativa')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de Shannon-Wiener (H’). Tiene en cuenta la riqueza de especies y su abundancia. Este índice relaciona el número de especies con la proporción de individuos pertenecientes a cada una de ellas presente en la muestra. Además, mide la uniformidad de la distribución de los individuos entre las especies. ')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_8.png', width=Cm(4.89), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nH’ = índice se Shannon'
                                             '\nS = número de especies'
                                             '\nPi = proporción de individuos de la especie entre todas las especies, A mayor valor de H’ mayor diversidad de especies.'
                                             '\nLn= Logaritmo natural')
    
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una Comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nDmg = Índice de Margalef'
                                             '\nS = Número de especies.'
                                             '\nN = Número total de individuos'
                                             '\nD = Densidad'
                                             '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                             '\nLn= Logaritmo natural')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. El índice de Simpson representa la probabilidad de que dos individuos, dentro de un hábitat, seleccionados al azar pertenezcan a la misma especie. Es decir, cuanto más se acerca el valor de este índice a la unidad existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nƛ = índice de dominancia se Simpson'
                                             '\nID = índice de diversidad'
                                             '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, Que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nDMn = índice de Menhinick'
                                             '\nS= Número total de especies'
                                             '\nN = Numero de total de todos los individuos de todas las especies.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nEl índice de Pielou: se expresa como el grado de uniformidad en la distribución de individuos entre especies. Se puede medir comparando la diversidad observada en una Comunidad contra la diversidad máxima posible de una Comunidad hipotética con el mismo número de especies.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_12.png', width=Cm(4.72), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nê = índice de Pielou'
                                             '\n∑ = es la sumatoria de la proporción de individuos (pi) por la sumatoria del logaritmo natura de la proporción de individuos (lnpi), o el Índice de Shannon – Wiener '
                                             '\nS = es el número de especies presentes')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4234_parrafo = doc.add_paragraph()
    formulaCapitulo4234_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4234_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4234_run.add_picture('capitulo4/capitulo4234/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.4 ###
    #########################
    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde:'
                                             '\nNmax = Es el número de individuos en la especie más abundante.')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('\nRango de escala de 0 - 1')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    descripcionCapitulo4234.bold = True
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph()
    descripcionCapitulo4234 = di4234.add_run('Donde las escalas para la interpretación de los rangos de 0-1 son las siguientes:')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4234 = di4234.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4234 = di4234.add_run('De 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4234 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4234 = di4234.add_run('Mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo4234_format = di4234.paragraph_format
    descripcionCapitulo4234_format.line_spacing = 1.15
    descripcionCapitulo4234_format.space_after = 0
    descripcionCapitulo4234_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4234.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.1 ###
    #########################
    capitulo42341 = doc.add_paragraph()
    i42341 = capitulo42341.add_run(f'\n{temasCapitulo4[1][2][4][5][0]}')
    i42341_format = capitulo42341.paragraph_format
    i42341_format.line_spacing = 1.15

    i42341.font.name = 'Arial'
    i42341.font.size = Pt(12)
    i42341.font.bold = True
    capitulo42341.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.2 ###
    #########################
    capitulo42341 = doc.add_paragraph()
    i42341 = capitulo42341.add_run(f'\n{temasCapitulo4[1][2][4][5][1]}')
    i42341_format = capitulo42341.paragraph_format
    i42341_format.line_spacing = 1.15

    i42341.font.name = 'Arial'
    i42341.font.size = Pt(12)
    i42341.font.bold = True
    capitulo42341.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.4.2 ###
    #########################
    tituloTabla42342 = doc.add_paragraph()
    dti42342 = tituloTabla42342.add_run('\nTabla 4.x.- Índice de diversidad de estrato de las arbbustivas ___')
    dti42342_format = tituloTabla42342.paragraph_format
    dti42342_format.line_spacing = 1.15
    dti42342_format.space_after = 0

    dti42342.font.name = 'Courier New'
    dti42342.font.size = Pt(12)
    tituloTabla42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=40, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.1 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('Descripcion del capitulo')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.15
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42342_parrafo = doc.add_paragraph()
    imagenCapitulo42342_run = imagenCapitulo42342_parrafo.add_run('')
    imagenCapitulo42342_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42342_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.2 ###
    #########################
    tituloGrafico42342 = doc.add_paragraph()
    dgi42342 = tituloGrafico42342.add_run('Grafica 4.5.- Densidad de Estrato Arbustivo ___.')
    dgi42342_format = tituloGrafico42342.paragraph_format
    dgi42342_format.line_spacing = 1.15
    dgi42342_format.space_after = 0

    dgi42342.font.name = 'Bookman Old Style'
    dgi42342.font.size = Pt(12)
    tituloGrafico42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.4.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.2 ###
    #########################
    tituloTabla42342 = doc.add_paragraph()
    dti42342 = tituloTabla42342.add_run('\nTabla 4.x.- Valor de Importancia de las arbustivas ___ en el SA.')
    dti42342_format = tituloTabla42342.paragraph_format
    dti42342_format.line_spacing = 1.15
    dti42342_format.space_after = 0

    dti42342.font.name = 'Courier New'
    dti42342.font.size = Pt(12)
    tituloTabla42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=40, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.15
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42342_parrafo = doc.add_paragraph()
    imagenCapitulo42342_run = imagenCapitulo42342_parrafo.add_run('')
    imagenCapitulo42342_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42342_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.5 ###
    #########################
    tituloGrafico42341 = doc.add_paragraph()
    dgi42342 = tituloGrafico42342.add_run('Grafica 4.6.- Valor de Importancia Estrato Arbustivo ___.')
    dgi42342_format = tituloGrafico42342.paragraph_format
    dgi42342_format.line_spacing = 1.15
    dgi42342_format.space_after = 0

    dgi42342.font.name = 'Bookman Old Style'
    dgi42342.font.size = Pt(12)
    tituloGrafico42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.4.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42341 = doc.add_paragraph()
    descripcionCapitulo42342 = di42341.add_run('ABUNDANCIA')
    descripcionCapitulo42342_format = di42341.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.15
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.2 ###
    #########################
    tituloTabla42342 = doc.add_paragraph()
    dti42342 = tituloTabla42342.add_run('\nTabla 4.x.- Valor de Importancia de las arbustiva ___ en el Sistema Ambiental.')
    dti42342_format = tituloTabla42342.paragraph_format
    dti42342_format.line_spacing = 1.15
    dti42342_format.space_after = 0

    dti42342.font.name = 'Courier New'
    dti42342.font.size = Pt(12)
    tituloTabla42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=40, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.15
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42342_parrafo = doc.add_paragraph()
    imagenCapitulo42342_run = imagenCapitulo42342_parrafo.add_run('')
    imagenCapitulo42342_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42342_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.5 ###
    #########################
    tituloGrafico42342 = doc.add_paragraph()
    dgi42342 = tituloGrafico42342.add_run('Grafica 4.7.- Valor de abundancia absoluta ___.')
    dgi42342_format = tituloGrafico42342.paragraph_format
    dgi42342_format.line_spacing = 1.15
    dgi42342_format.space_after = 0

    dgi42342.font.name = 'Bookman Old Style'
    dgi42342.font.size = Pt(12)
    tituloGrafico42342.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nRIQUEZA DE ESPECIE')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nÍndice de Margalef')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nEl índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nLa riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('Índice de Simpson')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('Índice de Shannon')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.2 ###
    #########################
    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('\nÍndice de Pielou')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    descripcionCapitulo42342.bold = True
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42342 = doc.add_paragraph()
    descripcionCapitulo42342 = di42342.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42342_format = di42342.paragraph_format
    descripcionCapitulo42342_format.line_spacing = 1.5
    descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42342_format.space_before = 0

    descripcionCapitulo42342.font.name = 'Arial'
    descripcionCapitulo42342.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.2 ###
    #########################
    tabla42342 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42342.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42342.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'
    
    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.3 ###
    #########################
    capitulo42343 = doc.add_paragraph()
    i42343 = capitulo42343.add_run(f'\n{temasCapitulo4[1][2][4][5][2]}')
    i42343_format = capitulo42341.paragraph_format
    i42343_format.line_spacing = 1.15

    i42343.font.name = 'Arial'
    i42343.font.size = Pt(12)
    i42343.font.bold = True
    capitulo42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.4.3 ###
    #########################
    tituloTabla42343 = doc.add_paragraph()
    dti42343 = tituloTabla42343.add_run('\nTabla 4.x.- Índice de diversidad de estrato de las arbbustivas ___')
    dti42343_format = tituloTabla42343.paragraph_format
    dti42343_format.line_spacing = 1.15
    dti42343_format.space_after = 0

    dti42343.font.name = 'Courier New'
    dti42343.font.size = Pt(12)
    tituloTabla42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42343.cell(rows, cols)
            t42342 = cell.paragraphs[0].add_run(' ')
            t42342.font.size = Pt(12)
            t42342.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('Descripcion del capitulo')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.15
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42343_parrafo = doc.add_paragraph()
    imagenCapitulo42343_run = imagenCapitulo42343_parrafo.add_run('')
    imagenCapitulo42343_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42343_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.3 ###
    #########################
    tituloGrafico42343 = doc.add_paragraph()
    dgi42343 = tituloGrafico42343.add_run('Grafica 4.8.- Densidad de Estrato Herbaceo ___.')
    dgi42343_format = tituloGrafico42343.paragraph_format
    dgi42343_format.line_spacing = 1.15
    dgi42343_format.space_after = 0

    dgi42343.font.name = 'Bookman Old Style'
    dgi42343.font.size = Pt(12)
    tituloGrafico42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.4.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.3 ###
    #########################
    tituloTabla42343 = doc.add_paragraph()
    dti42343 = tituloTabla42343.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42343_format = tituloTabla42343.paragraph_format
    dti42343_format.line_spacing = 1.15
    dti42343_format.space_after = 0

    dti42343.font.name = 'Courier New'
    dti42343.font.size = Pt(12)
    tituloTabla42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.3 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.15
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42343_parrafo = doc.add_paragraph()
    imagenCapitulo42343_run = imagenCapitulo42343_parrafo.add_run('')
    imagenCapitulo42343_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42343_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.3 ###
    #########################
    tituloGrafico42343 = doc.add_paragraph()
    dgi42343 = tituloGrafico42343.add_run('Grafica 4.9.- Valor de Importancia Estrato Herbácea ___.')
    dgi42343_format = tituloGrafico42343.paragraph_format
    dgi42343_format.line_spacing = 1.15
    dgi42343_format.space_after = 0

    dgi42343.font.name = 'Bookman Old Style'
    dgi42343.font.size = Pt(12)
    tituloGrafico42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nABUNDANCIA')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.15
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.3 ###
    #########################
    tituloTabla42343 = doc.add_paragraph()
    dti42343 = tituloTabla42343.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42343_format = tituloTabla42343.paragraph_format
    dti42343_format.line_spacing = 1.15
    dti42343_format.space_after = 0

    dti42343.font.name = 'Courier New'
    dti42343.font.size = Pt(12)
    tituloTabla42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42343_parrafo = doc.add_paragraph()
    imagenCapitulo42343_run = imagenCapitulo42343_parrafo.add_run('')
    imagenCapitulo42343_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42343_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.3 ###
    #########################
    tituloGrafico42343 = doc.add_paragraph()
    dgi42343 = tituloGrafico42343.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42343_format = tituloGrafico42343.paragraph_format
    dgi42343_format.line_spacing = 1.15
    dgi42343_format.space_after = 0

    dgi42343.font.name = 'Bookman Old Style'
    dgi42343.font.size = Pt(12)
    tituloGrafico42343.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.4.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    #descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.font.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('Índice de Margalef')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('Índice de Simpson')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42342.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42343_format = di42342.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('Índice de Shannon')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.3 ###
    #########################
    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('\nÍndice de Pielou')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    descripcionCapitulo42343.bold = True
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42343 = di42343.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42343_format = di42343.paragraph_format
    descripcionCapitulo42343_format.line_spacing = 1.5
    descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42343_format.space_before = 0

    descripcionCapitulo42343.font.name = 'Arial'
    descripcionCapitulo42343.font.size = Pt(12)
    di42343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.3 ###
    #########################
    tabla42343 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42343.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42343.cell(rows, cols)
            t42343 = cell.paragraphs[0].add_run(' ')
            t42343.font.size = Pt(12)
            t42343.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.3 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.4 ###
    #########################
    capitulo42344 = doc.add_paragraph()
    i42344 = capitulo42344.add_run(f'\n{temasCapitulo4[1][2][4][5][3]}')
    i42344_format = capitulo42344.paragraph_format
    i42344_format.line_spacing = 1.15

    i42344.font.name = 'Arial'
    i42344.font.size = Pt(12)
    i42344.font.bold = True
    capitulo42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.4.4 ###
    #########################
    tituloTabla42344 = doc.add_paragraph()
    dti42344 = tituloTabla42344.add_run('\nTabla 4.x.- Valor de densidad en gramíneas MDR en el Sistema Ambiental. ')
    dti42344_format = tituloTabla42344.paragraph_format
    dti42344_format.line_spacing = 1.15
    dti42344_format.space_after = 0

    dti42344.font.name = 'Courier New'
    dti42344.font.size = Pt(12)
    tituloTabla42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('Descripcion del capitulo')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.15
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42344_parrafo = doc.add_paragraph()
    imagenCapitulo42344_run = imagenCapitulo42344_parrafo.add_run('')
    imagenCapitulo42344_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42344_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.4 ###
    #########################
    tituloGrafico42344 = doc.add_paragraph()
    dgi42344 = tituloGrafico42344.add_run('Grafica 4.10.- Densidad de gramíneas ____.')
    dgi42344_format = tituloGrafico42344.paragraph_format
    dgi42344_format.line_spacing = 1.15
    dgi42344_format.space_after = 0

    dgi42344.font.name = 'Bookman Old Style'
    dgi42344.font.size = Pt(12)
    tituloGrafico42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('VALOR DE IMPORTANCIA')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.15
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.font.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.4 ###
    #########################
    tituloTabla42344 = doc.add_paragraph()
    dti42344 = tituloTabla42344.add_run('\nTabla 4.X.-	Valor de importancia de gramíneas ___ en el Sistema Ambiental.')
    dti42344_format = tituloTabla42344.paragraph_format
    dti42344_format.line_spacing = 1.15
    dti42344_format.space_after = 0

    dti42344.font.name = 'Courier New'
    dti42344.font.size = Pt(12)
    tituloTabla42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.15
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.4 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Grafica del capitulo 4.2.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42344_parrafo = doc.add_paragraph()
    imagenCapitulo42344_run = imagenCapitulo42344_parrafo.add_run('')
    imagenCapitulo42344_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42344_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.4 ###
    #########################
    tituloGrafico42344 = doc.add_paragraph()
    dgi42344 = tituloGrafico42344.add_run('Grafica 4.12.- Abundancia estrato de gramíneas ____.')
    dgi42344_format = tituloGrafico42344.paragraph_format
    dgi42344_format.line_spacing = 1.15
    dgi42344_format.space_after = 0

    dgi42344.font.name = 'Bookman Old Style'
    dgi42344.font.size = Pt(12)
    tituloGrafico42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.4 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nABUNDANCIA')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.15
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.4 ###
    #########################
    tituloTabla42344 = doc.add_paragraph()
    dti42344 = tituloTabla42344.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42344_format = tituloTabla42344.paragraph_format
    dti42344_format.line_spacing = 1.15
    dti42344_format.space_after = 0

    dti42344.font.name = 'Courier New'
    dti42344.font.size = Pt(12)
    tituloTabla42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42344_parrafo = doc.add_paragraph()
    imagenCapitulo42344_run = imagenCapitulo42344_parrafo.add_run('')
    imagenCapitulo42344_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42344_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.4 ###
    #########################
    tituloGrafico42344 = doc.add_paragraph()
    dgi42344 = tituloGrafico42344.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42344_format = tituloGrafico42344.paragraph_format
    dgi42344_format.line_spacing = 1.15
    dgi42344_format.space_after = 0

    dgi42344.font.name = 'Bookman Old Style'
    dgi42344.font.size = Pt(12)
    tituloGrafico42344.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.4 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    #descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.font.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('Índice de Margalef')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('Índice de Simpson')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('Índice de Shannon')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('\nÍndice de Pielou')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    descripcionCapitulo42344.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42344 = doc.add_paragraph()
    descripcionCapitulo42344 = di42344.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42344_format = di42344.paragraph_format
    descripcionCapitulo42344_format.line_spacing = 1.5
    descripcionCapitulo42344_format.space_after = 0
    descripcionCapitulo42344_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.4 ###
    #########################
    tabla42344 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42344.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42344.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.4 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.5 ###
    #########################
    capitulo42345 = doc.add_paragraph()
    i42345 = capitulo42345.add_run(f'\n{temasCapitulo4[1][2][4][5][4]}')
    i42345_format = capitulo42345.paragraph_format
    i42345_format.line_spacing = 1.15

    i42345.font.name = 'Arial'
    i42345.font.size = Pt(12)
    i42345.font.bold = True
    capitulo42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.4.5 ###
    #########################
    tituloTabla42345 = doc.add_paragraph()
    dti42345 = tituloTabla42345.add_run('\nTabla 4.x.- Valor de densidad de suculentas MDR en el Sistema Ambiental. ')
    dti42345_format = tituloTabla42345.paragraph_format
    dti42345_format.line_spacing = 1.15
    dti42345_format.space_after = 0

    dti42345.font.name = 'Courier New'
    dti42345.font.size = Pt(12)
    tituloTabla42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=30, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(30):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('Descripcion del capitulo')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.15
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################5
    ### Grafica del capitulo 4.2.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42345_parrafo = doc.add_paragraph()
    imagenCapitulo42345_run = imagenCapitulo42345_parrafo.add_run('')
    imagenCapitulo42345_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42345_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.5 ###
    #########################
    tituloGrafico42345 = doc.add_paragraph()
    dgi42345 = tituloGrafico42345.add_run('Grafica 4.10.- Densidad de gramíneas ____.')
    dgi42345_format = tituloGrafico42345.paragraph_format
    dgi42345_format.line_spacing = 1.15
    dgi42345_format.space_after = 0

    dgi42345.font.name = 'Bookman Old Style'
    dgi42345.font.size = Pt(12)
    tituloGrafico42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.4.5 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('VALOR DE IMPORTANCIA')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.15
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.font.bold = True
    di42344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.5 ###
    #########################
    tituloTabla42345 = doc.add_paragraph()
    dti42345 = tituloTabla42345.add_run('\nTabla 4.X.-	Valor de importancia de suculentas ___ en el Sistema Ambiental.')
    dti42345_format = tituloTabla42345.paragraph_format
    dti42345_format.line_spacing = 1.15
    dti42345_format.space_after = 0

    dti42345.font.name = 'Courier New'
    dti42345.font.size = Pt(12)
    tituloTabla42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.15
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.5 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Grafica del capitulo 4.2.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42345_parrafo = doc.add_paragraph()
    imagenCapitulo42345_run = imagenCapitulo42345_parrafo.add_run('')
    imagenCapitulo42345_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42345_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.5 ###
    #########################
    tituloGrafico42345 = doc.add_paragraph()
    dgi42345 = tituloGrafico42345.add_run('Grafica 4.12.- Abundancia estrato de gramíneas ____.')
    dgi42345_format = tituloGrafico42345.paragraph_format
    dgi42345_format.line_spacing = 1.15
    dgi42345_format.space_after = 0

    dgi42345.font.name = 'Bookman Old Style'
    dgi42345.font.size = Pt(12)
    tituloGrafico42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.5 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.4.5 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.4.4 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nABUNDANCIA')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.15
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42344.font.name = 'Arial'
    descripcionCapitulo42344.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.5 ###
    #########################
    tituloTabla42345 = doc.add_paragraph()
    dti42345 = tituloTabla42345.add_run('\nTabla 4.x.- Valor de abundancia de suculentas ___ en el Sistema Ambiental.')
    dti42345_format = tituloTabla42345.paragraph_format
    dti42345_format.line_spacing = 1.15
    dti42345_format.space_after = 0

    dti42345.font.name = 'Courier New'
    dti42345.font.size = Pt(12)
    tituloTabla42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=11, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42345_parrafo = doc.add_paragraph()
    imagenCapitulo42345_run = imagenCapitulo42345_parrafo.add_run('')
    imagenCapitulo42345_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42345_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.4.5 ###
    #########################
    tituloGrafico42345 = doc.add_paragraph()
    dgi42345 = tituloGrafico42345.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42345_format = tituloGrafico42345.paragraph_format
    dgi42345_format.line_spacing = 1.15
    dgi42345_format.space_after = 0

    dgi42345.font.name = 'Bookman Old Style'
    dgi42345.font.size = Pt(12)
    tituloGrafico42345.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.5 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    #descripcionCapitulo42343_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.font.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('Índice de Margalef')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42345.cell(rows, cols)
            t42344 = cell.paragraphs[0].add_run(' ')
            t42344.font.size = Pt(12)
            t42344.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('Índice de Simpson')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42345_format = di42344.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    #descripcionCapitulo42342_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('Índice de Shannon')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42345_format = di42344.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.4.5 ###
    #########################
    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('\nÍndice de Pielou')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    descripcionCapitulo42345.bold = True
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42345 = doc.add_paragraph()
    descripcionCapitulo42345 = di42345.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42345_format = di42345.paragraph_format
    descripcionCapitulo42345_format.line_spacing = 1.5
    descripcionCapitulo42345_format.space_after = 0
    descripcionCapitulo42345_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.4.5 ###
    #########################
    tabla42345 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42345.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42345.cell(rows, cols)
            t42345 = cell.paragraphs[0].add_run(' ')
            t42345.font.size = Pt(12)
            t42345.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.5 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.4.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.4.6 ###
    #########################
    capitulo42346 = doc.add_paragraph()
    i42346 = capitulo42346.add_run(f'\n{temasCapitulo4[1][2][4][5][5]}')
    i42346_format = capitulo42346.paragraph_format
    i42346_format.line_spacing = 1.15

    i42346.font.name = 'Arial'
    i42346.font.size = Pt(12)
    i42346.font.bold = True
    capitulo42346.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.4.6 ###
    #########################
    for lista in range(5):
        di42346 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42346 = di42346.add_run(f'Descripcion {lista + 1}')
        descripcionCapitulo42346_format = di42346.paragraph_format
        descripcionCapitulo42346_format.line_spacing = 1.15
        descripcionCapitulo42346_format.space_after = 0
        descripcionCapitulo42346_format.space_before = 0

        descripcionCapitulo42346.font.name = 'Arial'
        descripcionCapitulo42346.font.size = Pt(12)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.4.6 ###
    #########################
    tituloTabla42346 = doc.add_paragraph()
    dti42346 = tituloTabla42346.add_run('\nTabla 4.X.- Rangos y valores de los índices MDR en el sistema ambiental')
    dti42346_format = tituloTabla42346.paragraph_format
    dti42346_format.line_spacing = 1.15
    dti42346_format.space_after = 0

    dti42346.font.name = 'Courier New'
    dti42346.font.size = Pt(12)
    tituloTabla42346.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.4.6 ###
    #########################
    tabla42346 = doc.add_table(rows=19, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla42346.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla42346.cell(rows, cols)
            t42346 = cell.paragraphs[0].add_run(' ')
            t42346.font.size = Pt(12)
            t42346.font.name = 'Arial'
    
    ########################################################################################################################################################################
    # Capitulo 4.2.3.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.5 ###
    #########################
    capitulo4235 = doc.add_paragraph()
    i4235 = capitulo4235.add_run(f'{temasCapitulo4[1][2][4][6]}')
    i4235_format = capitulo4235.paragraph_format
    i4235_format.line_spacing = 1.15

    i4235.font.name = 'Arial'
    i4235.font.size = Pt(12)
    i4235.font.bold = True
    capitulo4235.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.5 ###
    #########################
    di4235 = doc.add_paragraph()
    descripcionCapitulo4235 = di4235.add_run('A continuación, se presenta el análisis de la información de las especies de flora encontradas durante el levantamiento de los sitios de muestreo en el sistema ambiental, donde se encuentra el área objeto de cambio de uso de suelo, determinando así su importancia y su posición dentro de los estratos encontrados, de acuerdo al inventario realizado.')
    descripcionCapitulo4235_format = di4235.paragraph_format
    descripcionCapitulo4235_format.line_spacing = 1.15
    descripcionCapitulo4235_format.space_after = 0
    descripcionCapitulo4235_format.space_before = 0

    descripcionCapitulo4235.font.name = 'Arial'
    descripcionCapitulo4235.font.size = Pt(12)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.5 ###
    #########################
    tituloTabla4235 = doc.add_paragraph()
    dti4235 = tituloTabla4235.add_run('\nTabla 4.X.- Estatus de las especies por estrato ___.')
    dti4235_format = tituloTabla4235.paragraph_format
    dti4235_format.line_spacing = 1.15
    dti4235_format.space_after = 0

    dti4235.font.name = 'Courier New'
    dti4235.font.size = Pt(12)
    tituloTabla4235.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.5 ###
    #########################
    tabla4235 = doc.add_table(rows=19, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla4235.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4235.cell(rows, cols)
            t4235 = cell.paragraphs[0].add_run(' ')
            t4235.font.size = Pt(12)
            t4235.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.5 ###
    #########################
    di4235 = doc.add_paragraph()
    descripcionCapitulo4235 = di4235.add_run('\nA continuación, se enlistan las coordenadas de cada sitio de muestreo donde se recabó la información para realizar la comparación, la cuales se encuentran en UTM zona 14 Norte, Datum WGS-84.')
    descripcionCapitulo4235_format = di4235.paragraph_format
    descripcionCapitulo4235_format.line_spacing = 1.15
    descripcionCapitulo4235_format.space_after = 0
    descripcionCapitulo4235_format.space_before = 0

    descripcionCapitulo4235.font.name = 'Arial'
    descripcionCapitulo4235.font.size = Pt(12)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.5 ###
    #########################
    tituloTabla4235 = doc.add_paragraph()
    dti4235 = tituloTabla4235.add_run('\nTabla 4.X.- Coordenadas geográficas de sitios de muestreo del MDM.')
    dti4235_format = tituloTabla4235.paragraph_format
    dti4235_format.line_spacing = 1.15
    dti4235_format.space_after = 0

    dti4235.font.name = 'Courier New'
    dti4235.font.size = Pt(12)
    tituloTabla4235.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.5 ###
    #########################
    tabla4235 = doc.add_table(rows=19, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla4235.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4235.cell(rows, cols)
            t4235 = cell.paragraphs[0].add_run(' ')
            t4235.font.size = Pt(12)
            t4235.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.5.1 ###
    #########################
    capitulo4235 = doc.add_paragraph()
    i4235 = capitulo4235.add_run(f'{temasCapitulo4[1][2][4][7][0]}')
    i4235_format = capitulo4235.paragraph_format
    i4235_format.line_spacing = 1.15

    i4235.font.name = 'Arial'
    i4235.font.size = Pt(12)
    i4235.font.bold = True
    capitulo4235.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.5.1 ###
    #########################
    di4235 = doc.add_paragraph()
    descripcionCapitulo4235 = di4235.add_run('En la presente tabla se presenta la vegetación encontrada en el sistema ambiental, por sitios de muestreo, donde se puede observar parámetros como altura, cobertura, así como la cantidad de individuos por especie.')
    descripcionCapitulo4235_format = di4235.paragraph_format
    descripcionCapitulo4235_format.line_spacing = 1.15
    descripcionCapitulo4235_format.space_after = 0
    descripcionCapitulo4235_format.space_before = 0

    descripcionCapitulo4235.font.name = 'Arial'
    descripcionCapitulo4235.font.size = Pt(12)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.5.1 ###
    #########################
    tituloTabla4235 = doc.add_paragraph()
    dti4235 = tituloTabla4235.add_run('\nTabla 4.X.- Coordenadas geográficas de sitios de muestreo del ____.')
    dti4235_format = tituloTabla4235.paragraph_format
    dti4235_format.line_spacing = 1.15
    dti4235_format.space_after = 0

    dti4235.font.name = 'Courier New'
    dti4235.font.size = Pt(12)
    tituloTabla4235.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.5.1 ###
    #########################
    tabla4235 = doc.add_table(rows=19, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla4235.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4235.cell(rows, cols)
            t4235 = cell.paragraphs[0].add_run(' ')
            t4235.font.size = Pt(12)
            t4235.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3 ###
    #########################
    capitulo4236 = doc.add_paragraph()
    i4236 = capitulo4236.add_run(f'{temasCapitulo4[1][2][4][8]}')
    i4236_format = capitulo4236.paragraph_format
    i4236_format.line_spacing = 1.15

    i4236.font.name = 'Arial'
    i4236.font.size = Pt(12)
    i4236.font.bold = True
    capitulo423.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3 ###
    #########################
    listaCapitulo4236 = [
        'Métodos basados en la cuantificación del número de especies presentes (riqueza específica).',
        'Métodos basados en la estructura de la comunidad, es decir, la distribución proporcional del valor de importancia de cada especie (abundancia relativa de los individuos, su biomasa, cobertura, productividad, etc.).'
    ]

    lista4236 = range(len(listaCapitulo4236))

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('La metodología utilizada para determinar la Biodiversidad presente en el sistema ambiental, es a través de la diversidad Alfa, para diferenciarlos en función de las variables biológicas que miden, se dividen en dos grandes grupos:')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    for lista in lista4236:
        di4236 = doc.add_paragraph()
        descripcionCapitulo4236 = di4236.add_run(f'{listaCapitulo4236[lista]}')
        descripcionCapitulo4236_format = di4236.paragraph_format
        descripcionCapitulo4236_format.line_spacing = 1.15
        descripcionCapitulo4236_format.space_after = 0
        descripcionCapitulo4236_format.space_before = 0

        descripcionCapitulo4236.font.name = 'Arial'
        descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Con la información obtenida de los estratos se calcularon los atributos de la vegetación, tales como densidad, dominancia y frecuencia de las especies dentro de la vegetación, consecuentemente se obtuvo el Índice de valor de importancia (IVI).')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Para realizar los cálculos de los índices y parámetros estructurales se emplearon las siguientes fórmulas:')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Densidad Absoluta. Está dada por el número de individuos de una especie o de todas las especies dividido por el número de sitios muestreados.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    #descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_1.png', width=Cm(2.01), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nD = Densidad'
                                             '\nN = Número de individuos muestreados por especie'
                                             '\nA = número de sitios muestreados o superficie muestrea según sea (x sito, ha o ACUSTF)')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nDensidad relativa. Está dada por el resultado de la densidad absoluta entre el número total de todos los individuos muestreados expresados en porcentajes ')
    descripcionCapitulo4236_format = di4234.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_2.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4234_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4234.add_run('Donde:'
                                             '\nDer = Densidad Relativa'
                                             '\nNi = Número de individuos de la especie'
                                             '\nNt = Número total de individuos de todas las especies')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4234.font.name = 'Arial'
    descripcionCapitulo4234.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nDominancia absoluta. Se define como el porcentaje de biomasa (área basal o superficie horizontal) que aporta una especie.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4234_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_3.png', width=Cm(2.75), height=Cm(1.5))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nDa = Densidad absoluta'
                                             '\nABi = Área basal de una especie'
                                             '\nA = Área muestreada (sitios muestreados)')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nLa dominancia relativa. Se calcula como la proporción de una especie en el área total evaluada, expresada en porcentaje.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_4.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nDor = Densidad relativa'
                                             '\nDai = Densidad absoluta de una especie'
                                             '\nDat= Densidad absoluta total de todas las especies')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nFrecuencia absoluta. Permite conocer las veces que se repite una especie en cada sitio de muestreo. ')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_5.png', width=Cm(3.89), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nFa = Frecuencia absoluta'
                                             '\nnsi = sumatoria del número de veces que una especie se observa dentro de todos los sitios de muestreo.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nFrecuencia relativa. Es el resultado de dividir la frecuencia absoluta de cada especie entre el número total de esas especies expresadas en porcentajes.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_6.png', width=Cm(4.74), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nFr = Frecuencia relativa'
                                             '\nFai = Frecuencia absoluta de cada especie'
                                             '\nFat = Frecuencia absoluta de todas las especies')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de valor de importancia (IVI). El índice de valor de importancia define cuáles de las especies presentes contribuyen en el carácter y estructura de una Comunidad. Este valor se obtiene mediante la sumatoria de la frecuencia relativa, la densidad relativa y la dominancia relativa.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_7.png', width=Cm(4.99), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4234.add_run('Donde:'
                                             '\nIVI = Índice de Valor de Importancia'
                                             '\nDer = Densidad relativa'
                                             '\nDor = Dominancia relativa'
                                             '\nFr = Frecuencia relativa')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de Shannon-Wiener (H’). Tiene en cuenta la riqueza de especies y su abundancia. Este índice relaciona el número de especies con la proporción de individuos pertenecientes a cada una de ellas presente en la muestra. Además, mide la uniformidad de la distribución de los individuos entre las especies. ')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_8.png', width=Cm(4.89), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nH’ = índice se Shannon'
                                             '\nS = número de especies'
                                             '\nPi = proporción de individuos de la especie entre todas las especies, A mayor valor de H’ mayor diversidad de especies.'
                                             '\nLn= Logaritmo natural')
    
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una Comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nDmg = Índice de Margalef'
                                             '\nS = Número de especies.'
                                             '\nN = Número total de individuos'
                                             '\nD = Densidad'
                                             '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                             '\nLn= Logaritmo natural')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. El índice de Simpson representa la probabilidad de que dos individuos, dentro de un hábitat, seleccionados al azar pertenezcan a la misma especie. Es decir, cuanto más se acerca el valor de este índice a la unidad existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nƛ = índice de dominancia se Simpson'
                                             '\nID = índice de diversidad'
                                             '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, Que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nDMn = índice de Menhinick'
                                             '\nS= Número total de especies'
                                             '\nN = Numero de total de todos los individuos de todas las especies.')
    descripcionCapitulo4236_format = di4234.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nEl índice de Pielou: se expresa como el grado de uniformidad en la distribución de individuos entre especies. Se puede medir comparando la diversidad observada en una Comunidad contra la diversidad máxima posible de una Comunidad hipotética con el mismo número de especies.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_12.png', width=Cm(4.72), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nê = índice de Pielou'
                                             '\n∑ = es la sumatoria de la proporción de individuos (pi) por la sumatoria del logaritmo natura de la proporción de individuos (lnpi), o el Índice de Shannon – Wiener '
                                             '\nS = es el número de especies presentes')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    #########################
    ### Formula del capitulo 4.2.3.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo4236_parrafo = doc.add_paragraph()
    formulaCapitulo4236_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo4236_run = formulaCapitulo4236_parrafo.add_run('')
    imagen = formulaCapitulo4236_run.add_picture('capitulo4/capitulo4234/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo4236_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 4.2.3.6 ###
    #########################
    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde:'
                                             '\nNmax = Es el número de individuos en la especie más abundante.')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('\nRango de escala de 0 - 1')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    descripcionCapitulo4236.bold = True
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4236 = doc.add_paragraph()
    descripcionCapitulo4236 = di4236.add_run('Donde las escalas para la interpretación de los rangos de 0-1 son las siguientes:')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4236 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4236 = di4236.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4236 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4236 = di4236.add_run('De 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4236 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo4236 = di4236.add_run('Mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo4236_format = di4236.paragraph_format
    descripcionCapitulo4236_format.line_spacing = 1.15
    descripcionCapitulo4236_format.space_after = 0
    descripcionCapitulo4236_format.space_before = 0

    descripcionCapitulo4236.font.name = 'Arial'
    descripcionCapitulo4236.font.size = Pt(12)
    #descripcionCapitulo4232.bold = True
    di4236.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.1 ###
    #########################
    capitulo42361 = doc.add_paragraph()
    i42361 = capitulo42361.add_run(f'\n{temasCapitulo4[1][2][4][9][0]}')
    i42361_format = capitulo42361.paragraph_format
    i42361_format.line_spacing = 1.15

    i42361.font.name = 'Arial'
    i42361.font.size = Pt(12)
    i42361.font.bold = True
    capitulo42361.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.2 ###
    #########################
    capitulo42361 = doc.add_paragraph()
    i42361 = capitulo42361.add_run(f'\n{temasCapitulo4[1][2][4][9][1]}')
    i42361_format = capitulo42361.paragraph_format
    i42361_format.line_spacing = 1.15

    i42346.font.name = 'Arial'
    i42346.font.size = Pt(12)
    i42346.font.bold = True
    capitulo42346.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.6.2 ###
    #########################
    tituloTabla42362 = doc.add_paragraph()
    dti42362 = tituloTabla42362.add_run('\nTabla 4.x.- Índice de diversidad de estrato de las arbustivas ___')
    dti42362_format = tituloTabla42362.paragraph_format
    dti42362_format.line_spacing = 1.15
    dti42362_format.space_after = 0

    dti42362.font.name = 'Courier New'
    dti42362.font.size = Pt(12)
    tituloTabla42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=40, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.1 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('Descripcion del capitulo')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.15
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42362_parrafo = doc.add_paragraph()
    imagenCapitulo42362_run = imagenCapitulo42362_parrafo.add_run('')
    imagenCapitulo42362_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42362_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.2 ###
    #########################
    tituloGrafico42362 = doc.add_paragraph()
    dgi42362 = tituloGrafico42362.add_run('Grafica 4.16.- Densidad de Estrato Arbustivo ___.')
    dgi42362_format = tituloGrafico42362.paragraph_format
    dgi42362_format.line_spacing = 1.15
    dgi42362_format.space_after = 0

    dgi42362.font.name = 'Bookman Old Style'
    dgi42362.font.size = Pt(12)
    tituloGrafico42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.6.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.2 ###
    #########################
    tituloTabla42362 = doc.add_paragraph()
    dti42362 = tituloTabla42362.add_run('\nTabla 4.x.- Valor de Importancia de las arbustivas ___ en el SA.')
    dti42362_format = tituloTabla42362.paragraph_format
    dti42362_format.line_spacing = 1.15
    dti42362_format.space_after = 0

    dti42362.font.name = 'Courier New'
    dti42362.font.size = Pt(12)
    tituloTabla42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=40, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.15
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42362_parrafo = doc.add_paragraph()
    imagenCapitulo42362_run = imagenCapitulo42362_parrafo.add_run('')
    imagenCapitulo42362_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42362_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.5 ###
    #########################
    tituloGrafico42361 = doc.add_paragraph()
    dgi42362 = tituloGrafico42362.add_run('Grafica 4.17.- Valor de Importancia Estrato Arbustivo ___.')
    dgi42362_format = tituloGrafico42362.paragraph_format
    dgi42362_format.line_spacing = 1.15
    dgi42362_format.space_after = 0

    dgi42342.font.name = 'Bookman Old Style'
    dgi42342.font.size = Pt(12)
    tituloGrafico42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.6.2 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('ABUNDANCIA')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.15
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.2 ###
    #########################
    tituloTabla42362 = doc.add_paragraph()
    dti42362 = tituloTabla42362.add_run('\nTabla 4.x.- Valor de Importancia de las arbustiva ___ en el Sistema Ambiental.')
    dti42362_format = tituloTabla42362.paragraph_format
    dti42362_format.line_spacing = 1.15
    dti42362_format.space_after = 0

    dti42362.font.name = 'Courier New'
    dti42362.font.size = Pt(12)
    tituloTabla42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=40, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.15
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42362_parrafo = doc.add_paragraph()
    imagenCapitulo42362_run = imagenCapitulo42362_parrafo.add_run('')
    imagenCapitulo42362_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42362_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.2.5 ###
    #########################
    tituloGrafico42362 = doc.add_paragraph()
    dgi42362 = tituloGrafico42362.add_run('Grafica 4.18.- Valor de abundancia absoluta ___.')
    dgi42362_format = tituloGrafico42362.paragraph_format
    dgi42362_format.line_spacing = 1.15
    dgi42362_format.space_after = 0

    dgi42362.font.name = 'Bookman Old Style'
    dgi42362.font.size = Pt(12)
    tituloGrafico42362.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nRIQUEZA DE ESPECIE')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nÍndice de Margalef')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nEl índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nLa riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('Índice de Simpson')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('Índice de Shannon')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.2 ###
    #########################
    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('\nÍndice de Pielou')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    descripcionCapitulo42362.bold = True
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42362 = doc.add_paragraph()
    descripcionCapitulo42362 = di42362.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42362_format = di42362.paragraph_format
    descripcionCapitulo42362_format.line_spacing = 1.5
    descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42362_format.space_before = 0

    descripcionCapitulo42362.font.name = 'Arial'
    descripcionCapitulo42362.font.size = Pt(12)
    di42362.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.2 ###
    #########################
    tabla42362 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42362.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42362.cell(rows, cols)
            t42362 = cell.paragraphs[0].add_run(' ')
            t42362.font.size = Pt(12)
            t42362.font.name = 'Arial'
    
    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.3 ###
    #########################
    capitulo42363 = doc.add_paragraph()
    i42363 = capitulo42363.add_run(f'\n{temasCapitulo4[1][2][4][9][2]}')
    i42363_format = capitulo42363.paragraph_format
    i42363_format.line_spacing = 1.15

    i42363.font.name = 'Arial'
    i42363.font.size = Pt(12)
    i42363.font.bold = True
    capitulo42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.6.3 ###
    #########################
    tituloTabla42363 = doc.add_paragraph()
    dti42363 = tituloTabla42363.add_run('\nTabla 4.x.- Índice de diversidad de estrato de las arbbustivas ___')
    dti42363_format = tituloTabla42363.paragraph_format
    dti42363_format.line_spacing = 1.15
    dti42363_format.space_after = 0

    dti42363.font.name = 'Courier New'
    dti42363.font.size = Pt(12)
    tituloTabla42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('Descripcion del capitulo')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.15
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42363_parrafo = doc.add_paragraph()
    imagenCapitulo42363_run = imagenCapitulo42343_parrafo.add_run('')
    imagenCapitulo42363_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42363_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.3 ###
    #########################
    tituloGrafico42363 = doc.add_paragraph()
    dgi42363 = tituloGrafico42363.add_run('Grafica 4.8.- Densidad de Estrato Herbaceo ___.')
    dgi42363_format = tituloGrafico42363.paragraph_format
    dgi42363_format.line_spacing = 1.15
    dgi42363_format.space_after = 0

    dgi42363.font.name = 'Bookman Old Style'
    dgi42363.font.size = Pt(12)
    tituloGrafico42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.6.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.3 ###
    #########################
    tituloTabla42363 = doc.add_paragraph()
    dti42363 = tituloTabla42363.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42363_format = tituloTabla42363.paragraph_format
    dti42363_format.line_spacing = 1.15
    dti42363_format.space_after = 0

    dti42363.font.name = 'Courier New'
    dti42363.font.size = Pt(12)
    tituloTabla42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.3 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.15
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42363_parrafo = doc.add_paragraph()
    imagenCapitulo42363_run = imagenCapitulo42363_parrafo.add_run('')
    imagenCapitulo42363_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42363_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.3 ###
    #########################
    tituloGrafico42363 = doc.add_paragraph()
    dgi42363 = tituloGrafico42363.add_run('Grafica 4.9.- Valor de Importancia Estrato Herbácea ___.')
    dgi42363_format = tituloGrafico42363.paragraph_format
    dgi42363_format.line_spacing = 1.15
    dgi42363_format.space_after = 0

    dgi42363.font.name = 'Bookman Old Style'
    dgi42363.font.size = Pt(12)
    tituloGrafico42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nABUNDANCIA')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.15
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.3 ###
    #########################
    tituloTabla42363 = doc.add_paragraph()
    dti42363 = tituloTabla42363.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42363_format = tituloTabla42363.paragraph_format
    dti42363_format.line_spacing = 1.15
    dti42363_format.space_after = 0

    dti42363.font.name = 'Courier New'
    dti42363.font.size = Pt(12)
    tituloTabla42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.6.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42363_parrafo = doc.add_paragraph()
    imagenCapitulo42363_run = imagenCapitulo42363_parrafo.add_run('')
    imagenCapitulo42363_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42363_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.3 ###
    #########################
    tituloGrafico42363 = doc.add_paragraph()
    dgi42363 = tituloGrafico42363.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42363_format = tituloGrafico42363.paragraph_format
    dgi42363_format.line_spacing = 1.15
    dgi42363_format.space_after = 0

    dgi42363.font.name = 'Bookman Old Style'
    dgi42363.font.size = Pt(12)
    tituloGrafico42363.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.6.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    #descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.font.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('Índice de Margalef')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('Índice de Simpson')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42343 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('Índice de Shannon')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42363 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.3 ###
    #########################
    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('\nÍndice de Pielou')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    descripcionCapitulo42363.bold = True
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42363 = doc.add_paragraph()
    descripcionCapitulo42363 = di42363.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42363_format = di42363.paragraph_format
    descripcionCapitulo42363_format.line_spacing = 1.5
    descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42363_format.space_before = 0

    descripcionCapitulo42363.font.name = 'Arial'
    descripcionCapitulo42363.font.size = Pt(12)
    di42363.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.3 ###
    #########################
    tabla42343 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42363.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42363.cell(rows, cols)
            t42363 = cell.paragraphs[0].add_run(' ')
            t42363.font.size = Pt(12)
            t42363.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.4.3 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.4 ###
    #########################
    capitulo42364 = doc.add_paragraph()
    i42364 = capitulo42364.add_run(f'\n{temasCapitulo4[1][2][4][9][3]}')
    i42364_format = capitulo42364.paragraph_format
    i42364_format.line_spacing = 1.15

    i42364.font.name = 'Arial'
    i42364.font.size = Pt(12)
    i42364.font.bold = True
    capitulo42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.6.4 ###
    #########################
    tituloTabla42364 = doc.add_paragraph()
    dti42364 = tituloTabla42364.add_run('\nTabla 4.x.- Valor de densidad en gramíneas MDR en el Sistema Ambiental. ')
    dti42364_format = tituloTabla42364.paragraph_format
    dti42364_format.line_spacing = 1.15
    dti42364_format.space_after = 0

    dti42364.font.name = 'Courier New'
    dti42364.font.size = Pt(12)
    tituloTabla42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=10, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('Descripcion del capitulo')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.15
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.6.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42364_parrafo = doc.add_paragraph()
    imagenCapitulo42364_run = imagenCapitulo42364_parrafo.add_run('')
    imagenCapitulo42364_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42364_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.4 ###
    #########################
    tituloGrafico42364 = doc.add_paragraph()
    dgi42364 = tituloGrafico42364.add_run('Grafica 4.10.- Densidad de gramíneas ____.')
    dgi42364_format = tituloGrafico42364.paragraph_format
    dgi42364_format.line_spacing = 1.15
    dgi42364_format.space_after = 0

    dgi42364.font.name = 'Bookman Old Style'
    dgi42364.font.size = Pt(12)
    tituloGrafico42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.6.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('VALOR DE IMPORTANCIA')
    descripcionCapitulo42364_format = di42344.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.15
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.font.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.4 ###
    #########################
    tituloTabla42364 = doc.add_paragraph()
    dti42364 = tituloTabla42364.add_run('\nTabla 4.X.-	Valor de importancia de gramíneas ___ en el Sistema Ambiental.')
    dti42364_format = tituloTabla42364.paragraph_format
    dti42364_format.line_spacing = 1.15
    dti42364_format.space_after = 0

    dti42364.font.name = 'Courier New'
    dti42364.font.size = Pt(12)
    tituloTabla42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.15
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.4 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Grafica del capitulo 4.2.3.6.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42364_parrafo = doc.add_paragraph()
    imagenCapitulo42364_run = imagenCapitulo42364_parrafo.add_run('')
    imagenCapitulo42364_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42364_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.4 ###
    #########################
    tituloGrafico42364 = doc.add_paragraph()
    dgi42364 = tituloGrafico42364.add_run('Grafica 4.12.- Abundancia estrato de gramíneas ____.')
    dgi42364_format = tituloGrafico42364.paragraph_format
    dgi42364_format.line_spacing = 1.15
    dgi42364_format.space_after = 0

    dgi42364.font.name = 'Bookman Old Style'
    dgi42364.font.size = Pt(12)
    tituloGrafico42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.4 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.6.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nABUNDANCIA')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.15
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.4 ###
    #########################
    tituloTabla42364 = doc.add_paragraph()
    dti42364 = tituloTabla42364.add_run('\nTabla 4.x.- Valor de abundancia de herbacéas ___ en el Sistema Ambiental.')
    dti42364_format = tituloTabla42364.paragraph_format
    dti42364_format.line_spacing = 1.15
    dti42364_format.space_after = 0

    dti42364.font.name = 'Courier New'
    dti42364.font.size = Pt(12)
    tituloTabla42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.6.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42364_parrafo = doc.add_paragraph()
    imagenCapitulo42364_run = imagenCapitulo42344_parrafo.add_run('')
    imagenCapitulo42364_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42364_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.4 ###
    #########################
    tituloGrafico42364 = doc.add_paragraph()
    dgi42364 = tituloGrafico42364.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42364_format = tituloGrafico42364.paragraph_format
    dgi42364_format.line_spacing = 1.15
    dgi42364_format.space_after = 0

    dgi42364.font.name = 'Bookman Old Style'
    dgi42364.font.size = Pt(12)
    tituloGrafico42364.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.4 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    #descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.font.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('Índice de Margalef')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('Índice de Simpson')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('Índice de Shannon')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42364_format = di42344.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('\nÍndice de Pielou')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    descripcionCapitulo42364.bold = True
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42364 = doc.add_paragraph()
    descripcionCapitulo42364 = di42364.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42364_format = di42364.paragraph_format
    descripcionCapitulo42364_format.line_spacing = 1.5
    descripcionCapitulo42364_format.space_after = 0
    descripcionCapitulo42364_format.space_before = 0

    descripcionCapitulo42364.font.name = 'Arial'
    descripcionCapitulo42364.font.size = Pt(12)
    di42364.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.4 ###
    #########################
    tabla42364 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42364.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42364.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.4 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.5 ###
    #########################
    capitulo42365 = doc.add_paragraph()
    i42365 = capitulo42365.add_run(f'\n{temasCapitulo4[1][2][4][9][4]}')
    i42365_format = capitulo42365.paragraph_format
    i42365_format.line_spacing = 1.15

    i42365.font.name = 'Arial'
    i42365.font.size = Pt(12)
    i42365.font.bold = True
    capitulo42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.6.5 ###
    #########################
    tituloTabla42365 = doc.add_paragraph()
    dti42365 = tituloTabla42365.add_run('\nTabla 4.x.- Valor de densidad de suculentas MDR en el Sistema Ambiental. ')
    dti42365_format = tituloTabla42365.paragraph_format
    dti42365_format.line_spacing = 1.15
    dti42365_format.space_after = 0

    dti42365.font.name = 'Courier New'
    dti42365.font.size = Pt(12)
    tituloTabla42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=30, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(30):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('Descripcion del capitulo')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.15
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################5
    ### Grafica del capitulo 4.2.3.6.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42365_parrafo = doc.add_paragraph()
    imagenCapitulo42365_run = imagenCapitulo42345_parrafo.add_run('')
    imagenCapitulo42365_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42365_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.5 ###
    #########################
    tituloGrafico42365 = doc.add_paragraph()
    dgi42365 = tituloGrafico42365.add_run('Grafica 4.10.- Densidad de gramíneas ____.')
    dgi42365_format = tituloGrafico42365.paragraph_format
    dgi42365_format.line_spacing = 1.15
    dgi42365_format.space_after = 0

    dgi42365.font.name = 'Bookman Old Style'
    dgi42365.font.size = Pt(12)
    tituloGrafico42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver la tabla del capitulo 4.2.3.6.5 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('VALOR DE IMPORTANCIA')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.15
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.font.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.5 ###
    #########################
    tituloTabla42365 = doc.add_paragraph()
    dti42365 = tituloTabla42365.add_run('\nTabla 4.X.-	Valor de importancia de suculentas ___ en el Sistema Ambiental.')
    dti42365_format = tituloTabla42365.paragraph_format
    dti42365_format.line_spacing = 1.15
    dti42365_format.space_after = 0

    dti42365.font.name = 'Courier New'
    dti42365.font.size = Pt(12)
    tituloTabla42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=11, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.4 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nDescripcion del capitulo')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.15
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.5 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Grafica del capitulo 4.2.3.6.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42365_parrafo = doc.add_paragraph()
    imagenCapitulo42365_run = imagenCapitulo42365_parrafo.add_run('')
    imagenCapitulo42365_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42365_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.5 ###
    #########################
    tituloGrafico42365 = doc.add_paragraph()
    dgi42365 = tituloGrafico42365.add_run('Grafica 4.12.- Abundancia estrato de gramíneas ____.')
    dgi42365_format = tituloGrafico42365.paragraph_format
    dgi42365_format.line_spacing = 1.15
    dgi42365_format.space_after = 0

    dgi42365.font.name = 'Bookman Old Style'
    dgi42365.font.size = Pt(12)
    tituloGrafico42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.5 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el resto del capitulo 4.2.3.6.5 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente codigo muestra como se tiene que insertar la hoja en vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nABUNDANCIA')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.15
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.5 ###
    #########################
    tituloTabla42365 = doc.add_paragraph()
    dti42365 = tituloTabla42365.add_run('\nTabla 4.x.- Valor de abundancia de suculentas ___ en el Sistema Ambiental.')
    dti42365_format = tituloTabla42365.paragraph_format
    dti42365_format.line_spacing = 1.15
    dti42365_format.space_after = 0

    dti42365.font.name = 'Courier New'
    dti42365.font.size = Pt(12)
    tituloTabla42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=11, cols=8, style='Table Grid')

    for cols in range(8):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(11):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Grafica del capitulo 4.2.3.6.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """

    imagenCapitulo42365_parrafo = doc.add_paragraph()
    imagenCapitulo42365_run = imagenCapitulo42365_parrafo.add_run('')
    imagenCapitulo42365_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo42365_run = doc.add_picture('capitulo4/grafico.jpg', width=Cm(15.59), height=Cm(10.16))

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.6.5 ###
    #########################
    tituloGrafico42365 = doc.add_paragraph()
    dgi42365 = tituloGrafico42365.add_run('Grafica 4.9.- Abundancia del estrato Herbáceas ____')
    dgi42365_format = tituloGrafico42365.paragraph_format
    dgi42365_format.line_spacing = 1.15
    dgi42365_format.space_after = 0

    dgi42365.font.name = 'Bookman Old Style'
    dgi42365.font.size = Pt(12)
    tituloGrafico42365.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.5 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nRIQUEZA ESPECÍFICA')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    #descripcionCapitulo42363_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.font.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('Índice de Margalef')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('El índice de Biodiversidad de las __ especies presentes en el Sistema Ambiental nos arroja que tenemos una diversidad del ____ dado que los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como indicativos de alta Biodiversidad, el área del sistema ambiental se contempla en un rango _______ de diversidad de acuerdo al tipo de vegetación y ecosistema donde se desarrolla. ')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla42365.cell(rows, cols)
            t42364 = cell.paragraphs[0].add_run(' ')
            t42364.font.size = Pt(12)
            t42364.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nÍndice de Menhinick')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('La riqueza de las __ especies presentes en el área arroja una diversidad del ____, dado que los rangos van de 2 a 5, donde los rangos inferiores a 2 son considerados como relacionados con zonas de baja Biodiversidad y valores superiores a 5 son considerados como alta Biodiversidad, por lo tanto, el área presenta diversidad _____.')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nINDICE DE DOMINANCIA')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('Índice de Simpson')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('De acuerdo al índice de Simpson, la dominancia en este estrato es de ____, mientras que el índice de diversidad es de ____, por lo que podemos decir que hay _______________, de acuerdo a los rangos que van de 0 a 1.')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nÍndice de Berger - Parker')
    descripcionCapitulo42365_format = di42364.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('De acuerdo al índice de Berger - Parker, tenemos una dominancia de _____ dado que los valores van de 0 a 1, podemos decir que al área tiene una dominancia ______, como podemos ver en el siguiente cuadro')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nINDICE DE EQUIDAD')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    #descripcionCapitulo42362_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('Índice de Shannon')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('El índice de diversidad de las ___ especies presentes en el Sistema Ambiental nos arroja que tenemos una baja diversidad de _____, considerando que los rangos de un valor normal están entre 2 y 3 para los valores inferiores a 2 se consideran bajos y superiores a 3 son altos, podemos decir que la equidad del área ___________.')
    descripcionCapitulo42365_format = di42364.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=4, cols=2, style='Table Grid')

    for cols in range(4):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.6.5 ###
    #########################
    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('\nÍndice de Pielou')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42365.font.name = 'Arial'
    descripcionCapitulo42365.font.size = Pt(12)
    descripcionCapitulo42365.bold = True
    di42365.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42365 = doc.add_paragraph()
    descripcionCapitulo42365 = di42365.add_run('El índice de equidad de acuerdo a Pielou es de _____, __________________________________________________________, considerando que el rango va de 0 a 1, podemos decir que el sistema ambiental está dentro de un área con una equidad ______.')
    descripcionCapitulo42365_format = di42365.paragraph_format
    descripcionCapitulo42365_format.line_spacing = 1.5
    descripcionCapitulo42365_format.space_after = 0
    descripcionCapitulo42365_format.space_before = 0

    descripcionCapitulo42345.font.name = 'Arial'
    descripcionCapitulo42345.font.size = Pt(12)
    di42345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.6.5 ###
    #########################
    tabla42365 = doc.add_table(rows=3, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42365.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla42365.cell(rows, cols)
            t42365 = cell.paragraphs[0].add_run(' ')
            t42365.font.size = Pt(12)
            t42365.font.name = 'Arial'

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.6.5 ###
    #########################
    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 4.2.3.6.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.6.6 ###
    #########################
    capitulo42366 = doc.add_paragraph()
    i42366 = capitulo42366.add_run(f'\n{temasCapitulo4[1][2][4][9][5]}')
    i42366_format = capitulo42366.paragraph_format
    i42366_format.line_spacing = 1.15

    i42366.font.name = 'Arial'
    i42366.font.size = Pt(12)
    i42366.font.bold = True
    capitulo42366.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.6.6 ###
    #########################
    for lista in range(5):
        di42366 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42366 = di42366.add_run(f'Descripcion {lista + 1}')
        descripcionCapitulo42366_format = di42366.paragraph_format
        descripcionCapitulo42366_format.line_spacing = 1.15
        descripcionCapitulo42366_format.space_after = 0
        descripcionCapitulo42366_format.space_before = 0

        descripcionCapitulo42366.font.name = 'Arial'
        descripcionCapitulo42366.font.size = Pt(12)

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.6.6 ###
    #########################
    tituloTabla42366 = doc.add_paragraph()
    dti42366 = tituloTabla42366.add_run('\nTabla 4.X.- Rangos y valores de los índices MDR en el sistema ambiental')
    dti42366_format = tituloTabla42366.paragraph_format
    dti42366_format.line_spacing = 1.15
    dti42366_format.space_after = 0

    dti42366.font.name = 'Courier New'
    dti42366.font.size = Pt(12)
    tituloTabla42366.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.6.6 ###
    #########################
    tabla42366 = doc.add_table(rows=19, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla42366.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla42366.cell(rows, cols)
            t42366 = cell.paragraphs[0].add_run(' ')
            t42366.font.size = Pt(12)
            t42366.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7 ###
    #########################
    capitulo4237 = doc.add_paragraph()
    i4237 = capitulo4237.add_run(f'{temasCapitulo4[1][2][4][10]}')
    i4237_format = capitulo4237.paragraph_format
    i4237_format.line_spacing = 1.15

    i4237.font.name = 'Arial'
    i4237.font.size = Pt(12)
    i4237.font.bold = True
    capitulo4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7 ###
    #########################
    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('México es considerado como un país mega diverso, pues en su territorio se registra la mayor riqueza de especies de reptiles en el mundo, además de que ocupa el segundo lugar en mamíferos y el cuarto lugar en anfibios (Toledo, 1988). Muchas de las especies de vertebrados presentes en territorio mexicano, se distribuyen únicamente en el país, por lo que se cuenta con un alto grado de endemismo, en los anfibios es del 61%, en los reptiles es del 53% y en mamíferos es del 30% (Sélem-Salas C., et. al. 2004). Este alto nivel de endemismos puede ser explicado por factores topográficos y climáticos, los cuales generan condiciones ambientales muy particulares que lo favorecen (Flores-Villela y Gerez, 1994). El grupo de las aves es particularmente importante, pues en México habita el 12% del total de especies del mundo, de las cuales el 10% está considerado como endémico.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('La Fauna Silvestre se caracteriza por especies que habitan de forma natural dentro de un ecosistema las cuales están relacionadas entre sí con el resto de los organismos vivos (vegetación, microorganismos, entre otros), y los no vivos (suelo, clima, agua, radiación solar) que componen los ecosistemas. Así mismo la Ley General del Equilibrio Ecológico y la Protección al Ambiente, publicada el 28 de enero de 1988 (SEMARNAT 1988) y siendo como última reforma el 11-04-2022 define el término Fauna Silvestre como: “Las especies animales que subsisten sujetas a los procesos de selección natural y que se desarrollan libremente, incluyendo sus poblaciones menores que se encuentran bajo control del hombre, así como los animales domésticos que por abandono se tornen salvajes y por ello sean susceptibles de captura y apropiación.”')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Por su extensión territorial el Estado de Coahuila ocupa el tercer lugar a nivel nacional, en la cual se pueden encontrar una gran diversidad de especies propia de la región árida y semiárida del Desierto Chihuahuense, desafortunadamente la riqueza natural se ha ido perdiendo paulatinamente por diversos factores, la Legislación Ambiental estatal señala que se debe salvaguardar la diversidad genética de las especies silvestres de las que depende la continuidad evolutiva, asegurar la preservación y el aprovechamiento sustentable de la biodiversidad del territorio. En lo que representa a la fauna silvestre La Comisión Nacional de la Biodiversidad (CONABIO) reporta que en la entidad existen 275 spp de aves, 107 spp de mamíferos y 24 spp de reptiles. ____________________ ____________________ ____________________ ____________________ ____________________ ____________________ ____________________ ____________________.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Avifauna
    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Avifauna')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.bold = True
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('La avifauna presente de la región varía en cuanto al número y diversidad de especies, de las cuales un cierto número son propias del territorio y otras especies son migratorias. A continuación, se detalla una lista de las especies de Aves que se distribuyen en el sistema ambiental de acuerdo a recorridos de campo y consulta bibliográfica.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7 ###
    #########################
    tituloTabla4237 = doc.add_paragraph()
    dti4237 = tituloTabla4237.add_run('\nTabla 4.X.- Lista de aves presentes en el Sistema Ambiental.')
    dti4237_format = tituloTabla4237.paragraph_format
    dti4237_format.line_spacing = 1.15
    dti4237_format.space_after = 0

    dti4237.font.name = 'Courier New'
    dti4237.font.size = Pt(12)
    tituloTabla4237.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7 ###
    #########################
    tabla4237 = doc.add_table(rows=19, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla4237.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4237.cell(rows, cols)
            t4237 = cell.paragraphs[0].add_run(' ')
            t4237.font.size = Pt(12)
            t4237.font.name = 'Arial'

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('*Pr= Protección especial, A= Amenazada, P= En Peligro de Extinción E=Probablemente Extinta en el medio silvestre, Sc= sin categoría.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(10)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7 ###
    #########################

    ##################################################################
    #Mamiferos
    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('\nMamiferos')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.bold = True
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Entre los mamíferos de distribución en la región se pueden encontrar diferentes especies, entre las cuales podemos mencionar las especies más comunes que se distribuyen en el sistema ambiental depredadores comunes tales como el ______________________________________________________________________.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Especies de Mamíferos que se distribuyen en el sistema ambiental, de acuerdo a recorridos de campo y consulta bibliográfica.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7 ###
    #########################
    tituloTabla4237 = doc.add_paragraph()
    dti4237 = tituloTabla4237.add_run('\nTabla 4.X.- Lista de Mamiferos presentes en el Sistema Ambiental.')
    dti4237_format = tituloTabla4237.paragraph_format
    dti4237_format.line_spacing = 1.15
    dti4237_format.space_after = 0

    dti4237.font.name = 'Courier New'
    dti4237.font.size = Pt(12)
    tituloTabla4237.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7 ###
    #########################
    tabla4237 = doc.add_table(rows=19, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla4237.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4237.cell(rows, cols)
            t4237 = cell.paragraphs[0].add_run(' ')
            t4237.font.size = Pt(12)
            t4237.font.name = 'Arial'

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('*Pr= Protección especial, A= Amenazada, P= En Peligro de Extinción E=Probablemente Extinta en el medio silvestre, Sc= sin categoría.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(10)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Reptiles y Anfibios
    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('\nReptiles y Anfibios')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.bold = True
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Actualmente el Estado de Coahuila alberga 132 especies nativas de anfibios y reptiles las cuales un cierto número de especies se distribuyen en el sistema ambiental, a continuación, se muestra un listado de las especies presentes en el sistema ambiental de acuerdo a los recorridos en campo y consulta bibliográfica.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('Especies de reptiles y anfibios que se distribuyen en el sistema ambiental, de acuerdo ha recorrido de campo y consulta bibliográfica.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7 ###
    #########################
    tituloTabla4237 = doc.add_paragraph()
    dti4237 = tituloTabla4237.add_run('\nTabla 4.X.- Lista de reptiles y anfibios presentes en el Sistema Ambiental.')
    dti4237_format = tituloTabla4237.paragraph_format
    dti4237_format.line_spacing = 1.15
    dti4237_format.space_after = 0

    dti4237.font.name = 'Courier New'
    dti4237.font.size = Pt(12)
    tituloTabla4237.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7 ###
    #########################
    tabla4237 = doc.add_table(rows=19, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla4237.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4237.cell(rows, cols)
            t4237 = cell.paragraphs[0].add_run(' ')
            t4237.font.size = Pt(12)
            t4237.font.name = 'Arial'

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('*Pr= Protección especial, A= Amenazada, P= En Peligro de Extinción E=Probablemente Extinta en el medio silvestre, Sc= sin categoría.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(10)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Lepidópteros
    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('\nLepidópteros')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.bold = True
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('En cuanto a invertebrados, para México se tienen registradas alrededor de 65 mil especies, en su mayoría insectos de los cuales poco menos son alrededor de 48 mil especies. La presencia del orden Lepidóptera de la región varía en cuanto al número y diversidad de especies, de las cuales 14,380 son residentes y migratorias en invierno o verano dentro del territorio nacional y 3,590 especies son exclusivas de México, estas cifran varían dependiendo de los autores como Luis et al., 2000 el cual menciona que en México las mariposas pertenecen al grupo de los insectos más diversos  contando aproximadamente con 1,800 especies  que representan casi el 10% de la fauna de mariposas en el mundo. Para el caso de Coahuila no se cuenta con información en cuanto al número total de especies de mariposas presentes en el estado, mas, sin embargo, se detalla una lista de las especies de mariposas que se distribuyen dentro del sistema ambiental, de acuerdo con recorridos de campo y consulta bibliográfica.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = Pt(12)
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7 ###
    #########################
    tituloTabla4237 = doc.add_paragraph()
    dti4237 = tituloTabla4237.add_run('\nTabla 4.X.- Lista de insectos presentes en el Sistema Ambiental.')
    dti4237_format = tituloTabla4237.paragraph_format
    dti4237_format.line_spacing = 1.15
    dti4237_format.space_after = 0

    dti4237.font.name = 'Courier New'
    dti4237.font.size = Pt(12)
    tituloTabla4237.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7 ###
    #########################
    tabla4237 = doc.add_table(rows=19, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla4237.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(19):
            cell = tabla4237.cell(rows, cols)
            t4237 = cell.paragraphs[0].add_run(' ')
            t4237.font.size = Pt(12)
            t4237.font.name = 'Arial'

    di4237 = doc.add_paragraph()
    descripcionCapitulo4237 = di4237.add_run('*Pr= Protección especial, A= Amenazada, P= En Peligro de Extinción E=Probablemente Extinta en el medio silvestre, Sc= sin categoría.')
    descripcionCapitulo4237_format = di4237.paragraph_format
    descripcionCapitulo4237_format.line_spacing = 1.15
    descripcionCapitulo4237_format.space_after = 0
    descripcionCapitulo4237_format.space_before = 0

    descripcionCapitulo4237.font.name = 'Arial'
    descripcionCapitulo4237.font.size = 12
    di4237.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.1 ###
    #########################
    capitulo42371 = doc.add_paragraph()
    i42371 = capitulo42371.add_run(f'\n{temasCapitulo4[1][2][4][11][0]}')
    i42371_format = capitulo42371.paragraph_format
    i42371_format.line_spacing = 1.15

    i42371.font.name = 'Arial'
    i42371.font.size = Pt(12)
    i42371.font.bold = True
    capitulo42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.1 ###
    #########################
    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('La metodología utilizada para le evaluación de Fauna Silvestre consistió en definir la forma de análisis de trabajo en el área sujeta al cambio de uso de suelo, consistiendo en realizar diversos muestreos, utilizando transectos, colocando estaciones olfativas y de escucha. Se registraron los grupos de vertebrados terrestres representados por reptiles, aves, mamíferos y anfibios para esto se determinaron sus hábitats, se efectuaron observaciones (a simple vista o con binoculares) realizándose de la siguiente forma y orden, todo esto por la cantidad de área que manejaremos en nuestro muestreo.')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Aves
    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('\nAves')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    descripcionCapitulo42371.font.bold = True
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('Para el caso de muestreo de aves se utilizó el método de muestreo en transectos de franja fija, el cual permite estimar la riqueza específica y la abundancia relativa de las especies de fauna silvestre correspondientes a este grupo, el procedimiento en el cual se basó este muestreo cuenta las siguientes etapas: ')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista1Capitulo_42731 = [
        'Elección del transecto: El punto de partida quedo definido por el tipo de hábitat y tipo de especies estableciendo así un transecto de muestreo lineal de ancho variable, además se establecieron cinco cámaras de foto trampeo destituidas en el área de estudio a una distancia de 200 mts entre cámara. Así mismo se colocaron trampas de Sherman las cuales consisten en pequeñas cajas metálicas con una puerta de acceso que se activan al encontrarse algún animal de talla pequeña dentro de ellas como roedores, como atrayente se utilizó una mezcla de avena con crema de cacahuate y vainilla.',
        'Muestreo: El recorrido en transecto se realizó a pie en donde el o los observadores caminan en una línea recta observando a las especies que se avisten dentro del ancho de transecto establecido el cual fue de 25 mts para cada eje del transecto, así mismo durante el recorrido se revisaron las trampas para verificar la captura de algún individuo, además el recorrido se registraron huellas, excretas, restos óseos, pelaje que puedan representar alguna especie de mamífero en el área.',
        'Análisis de datos: como resultado, se confeccionó una lista de especies presentes, con sus respectivas estimaciones lo cual permitió estimar la riqueza específica y la abundancia relativa de las especies de fauna silvestre correspondientes. (Nº de individuos por área).'
    ]

    lista1_42731 = range(len(lista1Capitulo_42731))

    for lista in lista1_42731:
        di42371 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42371 = di42371.add_run(f'{lista + 1}.- {lista1Capitulo_42731[lista]}')
        descripcionCapitulo42371_format = di42371.paragraph_format
        descripcionCapitulo42371_format.line_spacing = 1.15
        descripcionCapitulo42371_format.space_after = 0
        descripcionCapitulo42371_format.space_before = 0

        descripcionCapitulo42371.font.name = 'Arial'
        descripcionCapitulo42371.font.size = Pt(12)
        di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Reptiles
    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('\nReptiles')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('Se utilizó el método de muestreo en transectos, que es el que permite estimar la riqueza específica y la abundancia relativa, el procedimiento se fue desarrollando de acuerdo a las siguientes etapas:')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista2Capitulo_42731 = [
        'Elección del transecto: El punto de partida quedo definido por el tipo de hábitat y tipo de especie, potencialmente presente, en este caso, dado que los hábitats por sitio de estudio, son relativamente homogéneos, el punto de inicio fue seleccionado arbitrariamente, sin embargo, todos los transectos siguieron paralelos al curso transecto inicial.',
        'Longitud del transecto: Cada transecto se realizó en forma lineal y en una extensión de 1661 m de longitud por 12 m de ancho (6 m a cada lado del transecto), en áreas donde se observaron fauna.',
        'Muestreo: el transecto se recorrió a pie, en un tiempo estandarizado para todos los transectos, se registrarán todos los individuos avistados en una franja de 6 metros a cada lado del eje del transecto. se realizó una exhaustiva revisión del área circundante (dentro de la franja) especialmente bajo piedras, remoción somera de sustratos y cerca de las madrigueras anotando en formatos de campo toda especie correspondiente a este grupo.',
        'Análisis de datos: Como resultado, se confeccionó una lista de especies presentes por sitio, con sus respectivas estimaciones de densidad y abundancia (Nº de individuos por área).'
    ]

    lista2_42731 = range(len(lista2Capitulo_42731))

    for lista in lista2_42731:
        di42371 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42371 = di42371.add_run(f'{lista + 1}.- {lista2Capitulo_42731[lista]}')
        descripcionCapitulo42371_format = di42371.paragraph_format
        descripcionCapitulo42371_format.line_spacing = 1.15
        descripcionCapitulo42371_format.space_after = 0
        descripcionCapitulo42371_format.space_before = 0

        descripcionCapitulo42371.font.name = 'Arial'
        descripcionCapitulo42371.font.size = Pt(12)
        di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('Los métodos que se utilizaron en el muestreo de los diferentes tipos de fauna silvestre en el área de estudio son una herramienta básica, que permite al analista por medio de los estudios pertinentes y sus distintos métodos obtener una pequeña idea de las especies que pudieran existir en el área y poder hacer una extrapolación a la superficie que se desee, las tomas muéstrales son sencillamente un procedimiento que empleamos para extraer tan solo una pequeña muestra de una población dentro de una área a lo cual lo llamaremos espacio muestral dentro de una área determinada.')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ##################################################################
    #Lepidópteros
    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('\nLepidópteros')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    descripcionCapitulo42371.font.bold = True
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42371 = doc.add_paragraph()
    descripcionCapitulo42371 = di42371.add_run('Para el grupo de los insectos especialmente para especies de lepidópteros, así como también otras especies de insectos, se utilizó el método de muestreo en transectos, que es el que permite estimar la riqueza específica y la abundancia relativa, el procedimiento se fue desarrollando de acuerdo a las siguientes etapas:')
    descripcionCapitulo42371_format = di42371.paragraph_format
    descripcionCapitulo42371_format.line_spacing = 1.15
    descripcionCapitulo42371_format.space_after = 0
    descripcionCapitulo42371_format.space_before = 0

    descripcionCapitulo42371.font.name = 'Arial'
    descripcionCapitulo42371.font.size = Pt(12)
    di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista3Capitulo_42731 = [
        'Elección del transecto: La elección del método y del transecto quedo definido por el tipo de hábitat y de la especie en cuestión a monitoreo, tomando en cuenta lo anterior, el área de estudio es una área abierta de poca vegetación, esta característica permite al o los observadores tener una visión del área más extensa generando la oportunidad de registrar el mayor número de especies posibles dentro de nuestra superficie de muestreo, una vez analizadas estas variables se optó por implementar el monitoreo estableciendo transectos de franja o de banda, métodos que son adecuados para este tipo de hábitat.',
        'Longitud del transecto: Los transectos se establecieron de manera lineal con una 1661 mts extensión de m de longitud por 20 m de ancho (10 m a cada lado del transecto), en áreas donde se observaron fauna, las dimensiones del transecto de muestreo son determinadas por el observador dependiendo el tipo de hábitat y la superficie del mismo.',
        'Muestreo: el transecto se recorrió a pie, en un tiempo estandarizado para todos los transectos, durante el recorrido se busca registrar todos los individuos avistados en una franja de 10 metros a cada lado del eje del transecto, se realizó una exhaustiva revisión del área circundante (dentro de la franja), para registrar todos los ejemplares que se encuentren dentro de la superficie del transecto, con el transecto establecido es más fácil registrar más fácilmente aquellas especies sedentarias, territoriales y las de vuelo corto así como también permite la identificación rápida al vuelo o la captura en caso necesario para una mejor identificación, el transecto de muestreo de franja, tiene como objetivo registrar a todas las especies que se encuentren dentro de la superficie de muestreo del transecto.',
        'Análisis de datos: Como resultado del recorrido en caso de observaron especies se confecciona una lista de especies presentes, con su identificación y el número de individuos observados para posteriormente realizar los análisis estadísticos utilizando índices de diversidad y riqueza, así como también determinar la densidad de las especies por la superficie de muestreo y la abundancia relativa de las mismas (Nº de individuos por área.).'
    ]

    lista3_42731 = range(len(lista3Capitulo_42731))

    for lista in lista3_42731:
        di42371 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo42371 = di42371.add_run(f'{lista + 1}.- {lista3Capitulo_42731[lista]}')
        descripcionCapitulo42371_format = di42371.paragraph_format
        descripcionCapitulo42371_format.line_spacing = 1.15
        descripcionCapitulo42371_format.space_after = 0
        descripcionCapitulo42371_format.space_before = 0

        descripcionCapitulo42371.font.name = 'Arial'
        descripcionCapitulo42371.font.size = Pt(12)
        di42371.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.2 ###
    #########################
    capitulo42372 = doc.add_paragraph()
    i42372 = capitulo42372.add_run(f'\n{temasCapitulo4[1][2][4][11][1]}')
    i42372_format = capitulo42372.paragraph_format
    i42372_format.line_spacing = 1.15

    i42372.font.name = 'Arial'
    i42372.font.size = Pt(12)
    i42372.font.bold = True
    capitulo42372.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.2 ###
    #########################
    di42372 = doc.add_paragraph()
    descripcionCapitulo42372 = di42372.add_run('Mediante el muestreo en el área del sistema ambiental se recopilo información mediante un listado de las especies observadas durante los recorridos de muestreo, en el listado se conforma de las especies avistadas descritas con su familia, transecto de observación y la cantidad de número de individuos (ni).')
    descripcionCapitulo42372_format = di42372.paragraph_format
    descripcionCapitulo42372_format.line_spacing = 1.15
    descripcionCapitulo42372_format.space_after = 0
    descripcionCapitulo42372_format.space_before = 0

    descripcionCapitulo42372.font.name = 'Arial'
    descripcionCapitulo42372.font.size = Pt(12)
    di42372.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.2 ###
    #########################
    tituloTabla42372 = doc.add_paragraph()
    dti42372 = tituloTabla42372.add_run('\nTabla 4.X.- Fauna presente en el Sistema Ambiental.')
    dti42372_format = tituloTabla42372.paragraph_format
    dti42372_format.line_spacing = 1.15
    dti42372_format.space_after = 0

    dti42372.font.name = 'Courier New'
    dti42372.font.size = Pt(12)
    tituloTabla42372.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.2 ###
    #########################
    tabla42372 = doc.add_table(rows=50, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla42372.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(50):
            cell = tabla42372.cell(rows, cols)
            t42372 = cell.paragraphs[0].add_run(' ')
            t42372.font.size = Pt(12)
            t42372.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.2 ###
    #########################
    di42372 = doc.add_paragraph()
    descripcionCapitulo42372 = di42372.add_run('\nCoordenadas de los transectos de muestreo.')
    descripcionCapitulo42372_format = di42372.paragraph_format
    descripcionCapitulo42372_format.line_spacing = 1.15
    descripcionCapitulo42372_format.space_after = 0
    descripcionCapitulo42372_format.space_before = 0

    descripcionCapitulo42372.font.name = 'Arial'
    descripcionCapitulo42372.font.size = Pt(12)
    di42372.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42372 = doc.add_paragraph()
    descripcionCapitulo42372 = di42372.add_run('\nEn el sistema ambiental de estudio estos transectos y sitios de muestreo fueron realizados en las siguientes coordenadas en geográficas y UTM datum WGS 84 zona 14.')
    descripcionCapitulo42372_format = di42372.paragraph_format
    descripcionCapitulo42372_format.line_spacing = 1.15
    descripcionCapitulo42372_format.space_after = 0
    descripcionCapitulo42372_format.space_before = 0

    descripcionCapitulo42372.font.name = 'Arial'
    descripcionCapitulo42372.font.size = Pt(12)
    di42372.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.2 ###
    #########################
    tituloTabla42372 = doc.add_paragraph()
    dti42372 = tituloTabla42372.add_run('\nTabla 4.X.- Fauna presente en el Sistema Ambiental.')
    dti42372_format = tituloTabla42372.paragraph_format
    dti42372_format.line_spacing = 1.15
    dti42372_format.space_after = 0

    dti42372.font.name = 'Courier New'
    dti42372.font.size = Pt(12)
    tituloTabla42372.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.2 ###
    #########################
    tabla42372 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42372.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42372.cell(rows, cols)
            t42372 = cell.paragraphs[0].add_run(' ')
            t42372.font.size = Pt(12)
            t42372.font.name = 'Arial'

    di42372 = doc.add_paragraph()
    descripcionCapitulo42372 = di42372.add_run('VI; Vértice inicial, V: Vértice final.')
    descripcionCapitulo42372_format = di42372.paragraph_format
    descripcionCapitulo42372_format.line_spacing = 1.15
    descripcionCapitulo42372_format.space_after = 0
    descripcionCapitulo42372_format.space_before = 0

    descripcionCapitulo42372.font.name = 'Arial'
    descripcionCapitulo42372.font.size = Pt(12)
    di42372.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.3 ###
    #########################
    capitulo42373 = doc.add_paragraph()
    i42373 = capitulo42373.add_run(f'\n{temasCapitulo4[1][2][4][11][2]}')
    i42373_format = capitulo42373.paragraph_format
    i42373_format.line_spacing = 1.15

    i42373.font.name = 'Arial'
    i42373.font.size = Pt(12)
    i42373.font.bold = True
    capitulo42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.3 ###
    #########################
    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('Para el análisis de la información del grupo de las aves en el área del Sistema Ambiental el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra  el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al sistema ambiental, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 0
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(12)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.3 ###
    #########################
    tituloTabla42373 = doc.add_paragraph()
    dti42373 = tituloTabla42373.add_run('\nTabla 4.X.- Número de Individuos del grupo de las aves en el SA.')
    dti42373_format = tituloTabla42373.paragraph_format
    dti42373_format.line_spacing = 1.15
    dti42373_format.space_after = 0

    dti42373.font.name = 'Courier New'
    dti42373.font.size = Pt(12)
    tituloTabla42373.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=20, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.3 ###
    #########################
    tituloTabla42373 = doc.add_paragraph()
    dti42373 = tituloTabla42373.add_run('\nTabla 4.X.- Listado de las especies observadas en el SA su categoría de riesgo.')
    dti42373_format = tituloTabla42373.paragraph_format
    dti42373_format.line_spacing = 1.15
    dti42373_format.space_after = 0

    dti42373.font.name = 'Courier New'
    dti42373.font.size = Pt(12)
    tituloTabla42373.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=20, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('Sociabilidad (SOCI.); abundancia (ABU.); residencia (RES.); alimentación (ALIM.) y el tipo de observación (OBS.); Sc: Sociabilidad, R: Residente; C: Común, SL: Solitario, GR: Gregario, PJ: Pareja; Sc: Sin categoría, Pr: Sujeta a protección especial; A: Amenazada; P: En peligro de extinción; E: Extinta en medio silvestre.')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 0
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(10)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.3 ###
    #########################
    tituloTabla42373 = doc.add_paragraph()
    dti42373 = tituloTabla42373.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef, para el grupo de las aves en el área del Sistema ambiental.')
    dti42373_format = tituloTabla42373.paragraph_format
    dti42373_format.line_spacing = 1.15
    dti42373_format.space_after = 0

    dti42373.font.name = 'Courier New'
    dti42373.font.size = Pt(12)
    tituloTabla42373.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=20, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.3 ###
    #########################
    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('\nLos índices de diversidad de las _____ especies del grupo de las aves  presentes en el área del sistema ambiental  muestran que para el índice de Shannon  tenemos una diversidad de _________ lo cual quiere decir que para este grupo los valores resultantes se encuentran en su normalidad  ya que los rangos de valores para este índice  van de 2 a 3 para valores normales, por los contrario para los valores inferiores  serian aquellos inferiores a 2;  para el índice de Simpson resulta una diversidad ___________________ y una dominancia de las especies  baja  _______________________________________________________________________________________________________________________________________________. Por otra parte, el índice de Margalef el cual estima la biodiversidad de una comunidad, muestra valores _______________ ya que los valores de medida   considerados para una baja biodiversidad son para valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores a 3 y muy superiores aquellos con valores de 5. La especie más representativa fue ____________________.')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 0
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(12)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.3 ###
    #########################
    tituloTabla42373 = doc.add_paragraph()
    dti42373 = tituloTabla42373.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para el grupo de las _____________________________________.')
    dti42373_format = tituloTabla42373.paragraph_format
    dti42373_format.line_spacing = 1.15
    dti42373_format.space_after = 0

    dti42373.font.name = 'Courier New'
    dti42373.font.size = Pt(12)
    tituloTabla42373.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=20, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.3 ###
    #########################
    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('\nEl índice de diversidad para el grupo de las aves de las ____ especies presentes en el área del sistema ambiental presenta un índice de dominancia ___________, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de ________; para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del sistema ambiental en estudio nos indica la dominancia de la especie ___________________  como la más representativa, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie lo cual para este índice resulta que la especies _______________________________ saya son las más representativas, tal como se puede observar en las siguiente gráfica.')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 0
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(12)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina en el capitulo 4.2.3.7.3 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Grafica del capitulo 4.2.3.7.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo42373 = doc.add_paragraph()
    imagenCapitulo42373.text = ''
    imagenCapitulo42373 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo42373.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo42373.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo42373.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo42373.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.7.3 ###
    #########################
    tituloGrafico42373 = doc.add_paragraph()
    dgi42373 = tituloGrafico42373.add_run('Grafica 4.27.- Frecuencia y abundancia relativa del grupo de las aves en el área del sistema ambiental.')
    dgi42373_format = tituloGrafico42373.paragraph_format
    dgi42373_format.line_spacing = 1.15
    dgi42373_format.space_after = 0

    dgi42373.font.name = 'Bookman Old Style'
    dgi42373.font.size = Pt(12)
    tituloGrafico42373.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.3 ###
    #########################
    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('\nEn el grupo de las aves dentro del área del sistema ambiental, poseé una riqueza específica de ____ especies las cuales tienen una equidad de __________, con lo cual se puede afirmar que la mayoria de las especies son equitativas. La máxima diversidad que se puede alcanzar en el sistema ambiental de este grupo es de _________ y la diversidad calculada es de _________ lo que indica que este grupo está cerca de alcanzar su máxima diversidad y posee una distribución equitativa, la especie mas representativa para este grupo fue ___________________ con 8 individuos registrados en comparación con las demás especies observadas en el área de estudio, considerando que el grupo tendra un porcentaje de desplazamiento del ____% en el área del sistema ambiental en comparacion con el area  ACUSTF, por lo tanto el grupo de las aves no se vera afectado ya que su distribucion en cuanto a  porcentaje ____________ en comparacion al area del sistema ambiental. ')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 0
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(12)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42373 = doc.add_paragraph()
    descripcionCapitulo42373 = di42373.add_run('\nDe acuerdo con los datos que anteceden por las caracteristicas del area del sistema ambiental el grupo de las aves se presenta en condiciones de ____________ en cuanto a riqueza, para la dominancia de especies se presenta con una calidad ______, en cuanto a la equidad de especies los valores obtenidos fueron medios, la principal actividad que puede degradarlas es la presion de este grupo por degradacion antropogenica, el sobrepastoreo y por el transito vehicular en la zona, acontinuacion se muestra en el cuadro de rangos de valor para el grupo de las aves en el area del sistema ambiental.\n')
    descripcionCapitulo42373_format = di42373.paragraph_format
    descripcionCapitulo42373_format.line_spacing = 1.15
    descripcionCapitulo42373_format.space_after = 1
    descripcionCapitulo42373_format.space_before = 0

    descripcionCapitulo42373.font.name = 'Arial'
    descripcionCapitulo42373.font.size = Pt(12)
    di42373.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.7.3 ###
    #########################
    tabla42373 = doc.add_table(rows=6, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42373.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla42373.cell(rows, cols)
            t42373 = cell.paragraphs[0].add_run(' ')
            t42373.font.size = Pt(12)
            t42373.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.4 ###
    #########################
    capitulo42374 = doc.add_paragraph()
    i42374 = capitulo42374.add_run(f'\n{temasCapitulo4[1][2][4][11][3]}')
    i42374_format = capitulo42374.paragraph_format
    i42374_format.line_spacing = 1.15

    i42374.font.name = 'Arial'
    i42374.font.size = Pt(12)
    i42374.font.bold = True
    capitulo42374.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.4 ###
    #########################
    di42374 = doc.add_paragraph()
    descripcionCapitulo42374 = di42374.add_run('Para el análisis de la información del grupo de los mamíferos en el área del Sistema ambiental se muestra el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, el número de individuos por área de muestreo y  el número de individuos extrapolados a la superficie correspondiente al sistema ambiental además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo42374_format = di42373.paragraph_format
    descripcionCapitulo42374_format.line_spacing = 1.15
    descripcionCapitulo42374_format.space_after = 0
    descripcionCapitulo42374_format.space_before = 0

    descripcionCapitulo42374.font.name = 'Arial'
    descripcionCapitulo42374.font.size = Pt(12)
    di42374.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.4 ###
    #########################
    tituloTabla42374 = doc.add_paragraph()
    dti42374 = tituloTabla42374.add_run('\nTabla 4.X.-	Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para el grupo de los mamíferos en el área del sistema ambiental.')
    dti42374_format = tituloTabla42374.paragraph_format
    dti42374_format.line_spacing = 1.15
    dti42374_format.space_after = 0

    dti42374.font.name = 'Courier New'
    dti42374.font.size = Pt(12)
    tituloTabla42374.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.4 ###
    #########################
    tabla42374 = doc.add_table(rows=10, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42374.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42374.cell(rows, cols)
            t42374 = cell.paragraphs[0].add_run(' ')
            t42374.font.size = Pt(12)
            t42374.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.4 ###
    #########################
    di42374 = doc.add_paragraph()
    descripcionCapitulo42374 = di42374.add_run('El índice de diversidad para el grupo de los mamíferos de las ____ especies presentes en el área del sistema ambiental presenta un índice de ___________, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de __________;, para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del sistema ambiental en estudio nos indica la dominancia de _______________________ como la especie más representativa, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie lo cual para este índice resulta que la especie _____________________ fueron las especies más representativas, tal como se puede observar en las siguiente gráfica.')
    descripcionCapitulo42374_format = di42373.paragraph_format
    descripcionCapitulo42374_format.line_spacing = 1.15
    descripcionCapitulo42374_format.space_after = 0
    descripcionCapitulo42374_format.space_before = 0

    descripcionCapitulo42374.font.name = 'Arial'
    descripcionCapitulo42374.font.size = Pt(12)
    di42374.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.7.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo42374 = doc.add_paragraph()
    imagenCapitulo42374.text = ''
    imagenCapitulo42374 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo42374.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo42374.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo42374.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo42374.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.7.4 ###
    #########################
    tituloGrafico42374 = doc.add_paragraph()
    dgi42374 = tituloGrafico42374.add_run('Grafica 4.28.- Frecuencia y abundancia relativa de mamíferos en el SA.\n')
    dgi42374_format = tituloGrafico42373.paragraph_format
    dgi42374_format.line_spacing = 1.15
    dgi42374_format.space_after = 0

    dgi42374.font.name = 'Bookman Old Style'
    dgi42374.font.size = Pt(12)
    tituloGrafico42374.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.4 ###
    #########################
    tabla42374 = doc.add_table(rows=10, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42374.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42374.cell(rows, cols)
            t42374 = cell.paragraphs[0].add_run(' ')
            t42374.font.size = Pt(12)
            t42374.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.4 ###
    #########################
    di42374 = doc.add_paragraph()
    descripcionCapitulo42374 = di42374.add_run('El grupo de los mamíferos posee una riqueza específica de _____ especies las cuales tienen una distribución de _________ con lo cual se puede afirmar que la presencia de especies es _______. La máxima diversidad que este grupo adquiere dentro del sistema ambiental es de _____________________ para la diversidad calculada lo que quiere decir que este grupo se encuentra cerca de alcanzar su máxima diversidad y pose un distribución equitativa destacando la especie _______________________ representado por ____ individuos dentro del área del sistema ambiental tal y como se manifiesta en la dominancia y la frecuencia, por lo cual la comunidad se considera muy diversa al tener menos especies dominantes. Lo que equivale a que este grupo tendrá un porcentaje de desplazamiento del ___% que se encuentra dentro del área de cambio de uso de suelo en comparación con el sistema ambiental.')
    descripcionCapitulo42374_format = di42373.paragraph_format
    descripcionCapitulo42374_format.line_spacing = 1.15
    descripcionCapitulo42374_format.space_after = 0
    descripcionCapitulo42374_format.space_before = 0

    descripcionCapitulo42374.font.name = 'Arial'
    descripcionCapitulo42374.font.size = Pt(12)
    di42374.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42374 = doc.add_paragraph()
    descripcionCapitulo42374 = di42374.add_run('De acuerdo con los datos que anteceden por las caracteristicas del area del sistema ambiental el grupo de las mamiferos se presenta en condiciones de media calidad en cuanto a riqueza y equidad de especies, para la dominancia de especies los valores presentes fueron ce _____________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________.\n')
    descripcionCapitulo42374_format = di42373.paragraph_format
    descripcionCapitulo42374_format.line_spacing = 1.15
    descripcionCapitulo42374_format.space_after = 0
    descripcionCapitulo42374_format.space_before = 0

    descripcionCapitulo42374.font.name = 'Arial'
    descripcionCapitulo42374.font.size = Pt(12)
    di42374.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.7.4 ###
    #########################
    tabla42374 = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla42374.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla42374.cell(rows, cols)
            t42374 = cell.paragraphs[0].add_run(' ')
            t42374.font.size = Pt(12)
            t42374.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.5 ###
    #########################
    capitulo42375 = doc.add_paragraph()
    i42375 = capitulo42375.add_run(f'\n{temasCapitulo4[1][2][4][11][4]}')
    i42375_format = capitulo42375.paragraph_format
    i42375_format.line_spacing = 1.15

    i42375.font.name = 'Arial'
    i42375.font.size = Pt(12)
    i42375.font.bold = True
    capitulo42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.5 ###
    #########################
    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('Para el análisis de la información del grupo de los reptiles en el área del sistema ambiental el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra  el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al sistema ambiental, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.5 ###
    #########################
    tituloTabla42375 = doc.add_paragraph()
    dti42375 = tituloTabla42375.add_run('\nTabla 4.x.- Listado de las especies de reptiles observadas en el área del sistema ambiental.')
    dti42375_format = tituloTabla42375.paragraph_format
    dti42375_format.line_spacing = 1.15
    dti42375_format.space_after = 0

    dti42375.font.name = 'Courier New'
    dti42375.font.size = Pt(12)
    tituloTabla42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.5 ###
    #########################
    tituloTabla42375 = doc.add_paragraph()
    dti42375 = tituloTabla42375.add_run('\nTabla 4.x.- Listado de las especies de reptiles observadas en el área del sistema ambiental.')
    dti42375_format = tituloTabla42375.paragraph_format
    dti42375_format.line_spacing = 1.15
    dti42375_format.space_after = 0

    dti42375.font.name = 'Courier New'
    dti42375.font.size = Pt(12)
    tituloTabla42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.5 ###
    #########################
    tituloTabla42375 = doc.add_paragraph()
    dti42375 = tituloTabla42375.add_run('\nTabla 4.X.- Listado de reptiles en el área del sistema ambiental con su categoría de riesgo.')
    dti42375_format = tituloTabla42375.paragraph_format
    dti42375_format.line_spacing = 1.15
    dti42375_format.space_after = 0

    dti42375.font.name = 'Courier New'
    dti42375.font.size = Pt(12)
    tituloTabla42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.5 ###
    #########################
    tituloTabla42375 = doc.add_paragraph()
    dti42375 = tituloTabla42375.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef, para el grupo de los reptiles.')
    dti42375_format = tituloTabla42375.paragraph_format
    dti42375_format.line_spacing = 1.15
    dti42375_format.space_after = 0

    dti42375.font.name = 'Courier New'
    dti42375.font.size = Pt(12)
    tituloTabla42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.5 ###
    #########################
    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('Los índices de diversidad de __________ especies del grupo de los reptiles presentes en el área del sistema ambiental muestran que para el índice de Shannon  tenemos una diversidad de ____________ lo cual quiere decir que para este grupo los valores resultantes se encuentran _______________ ya que los rangos de valores para este índice  van de 1.36 a 2, para valores normales por los contrario para los valores inferiores serian aquellos inferiores a 1.35; para el índice de Simpson resulta una diversidad __________________________ y una dominancia de las especies  ____________________ lo cual quiere decir que podemos encontrar especies dominantes para el grupo de los reptiles, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad  muestra  valores ______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.5 ###
    #########################
    tituloTabla42375 = doc.add_paragraph()
    dti42375 = tituloTabla42375.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para el grupo de los reptiles en el sistema ambiental.')
    dti42375_format = tituloTabla42375.paragraph_format
    dti42375_format.line_spacing = 1.15
    dti42375_format.space_after = 0

    dti42375.font.name = 'Courier New'
    dti42375.font.size = Pt(12)
    tituloTabla42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=10, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.5 ###
    #########################
    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('\nEl índice de diversidad para el grupo de los reptiles de las ___ especies presentes en el área del sistema ambiental presenta un índice de ____, para la riqueza de especies que se define como el número de especies presentes en una comunidad se obtiene un total de riqueza de ____________, para la abundancia relativa la cual expresa la representatividad de una especie dentro del conjunto de especies en el área del sistema ambiental en estudio nos indica que la especie más representativa es _________________________, para la frecuencia relativa la cual representa el número de muestras en las que se encuentra una especie lo cual para este índice resulta como especie más representativa ___________________________________________________________________________.')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.7.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo42375 = doc.add_paragraph()
    imagenCapitulo42375.text = ''
    imagenCapitulo42375 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo42375.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo42375.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo42375.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo42375.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.7.5 ###
    #########################
    tituloGrafico42375 = doc.add_paragraph()
    dgi42375 = tituloGrafico42375.add_run('Grafica 4.29.- Frecuencia y abundancia relativa de los reptiles en el SA.\n')
    dgi42375_format = tituloGrafico42375.paragraph_format
    dgi42375_format.line_spacing = 1.15
    dgi42375_format.space_after = 0

    dgi42375.font.name = 'Bookman Old Style'
    dgi42375.font.size = Pt(12)
    tituloGrafico42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.5 ###
    #########################
    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('\nEn el grupo de los reptiles dentro del área del sistema ambiental posee una riqueza específica de ____ especies, con una distribución de ___________ lo que equivale a que la equidad en las especies es ________, la máxima diversidad que se puede alcanzar en este grupo es de _____________________________________________ lo que indica que este grupo se encuentra cerca de alcanzar su ___________________, la especie más representativa fue _______________________________________, considerando que el grupo tendrá un porcentaje de desplazamiento del _____% en el área del sistema ambiental, indicando este el grupo no se afectara ya que su desplazamiento _______________ en comparación al sistema ambiental.')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('\nDe acuerdo con los datos que anteceden por las caracteristicas del area del sistema ambiental el grupo de los reptiles se presenta en condiciones de _______________ en cuanto a riqueza y equidad de especies, para la dominancia de especies la calidad en cuanto a sus valores fue _________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________.\n')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.6 ###
    #########################
    capitulo42376 = doc.add_paragraph()
    i42376 = capitulo42376.add_run(f'\n{temasCapitulo4[1][2][4][11][5]}')
    i42376_format = capitulo42376.paragraph_format
    i42376_format.line_spacing = 1.15

    i42376.font.name = 'Arial'
    i42376.font.size = Pt(12)
    i42376.font.bold = True
    capitulo42376.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.6 ###
    #########################
    di42376 = doc.add_paragraph()
    descripcionCapitulo42376 = di42376.add_run('Para el análisis de la información del grupo de los lepidópteros que se observaron en el área del sistema ambiental se plasma la siguiente información en la cual se muestra el número de individuos (ni) fueron aquellos observados en campo por la metodología aplicada para este grupo, así como también se muestra  el número de individuos por superficie de muestreo y el número de individuos extrapolados a la superficie correspondiente al sistema ambiental, además se plasma el estatus de riesgo en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, la residencia (RES.), la abundancia (ABUN.), la sociabilidad (SOCI.), la alimentación (ALIM.) y el tipo de observación (OBS.).')
    descripcionCapitulo42376_format = di42376.paragraph_format
    descripcionCapitulo42376_format.line_spacing = 1.15
    descripcionCapitulo42376_format.space_after = 0
    descripcionCapitulo42376_format.space_before = 0

    descripcionCapitulo42376.font.name = 'Arial'
    descripcionCapitulo42376.font.size = Pt(12)
    di42376.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################1
    ### Titulo de la tabla del capitulo 4.2.3.7.6 ###
    #########################
    tituloTabla42376 = doc.add_paragraph()
    dti42376 = tituloTabla42376.add_run('\nTabla 4.x.- Listado de especies de lepidópteros observados en el área del sistema ambiental.')
    dti42376_format = tituloTabla42376.paragraph_format
    dti42376_format.line_spacing = 1.15
    dti42376_format.space_after = 0

    dti42376.font.name = 'Courier New'
    dti42376.font.size = Pt(12)
    tituloTabla42376.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.6 ###
    #########################
    tabla42376 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42376.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42376.cell(rows, cols)
            t42376 = cell.paragraphs[0].add_run(' ')
            t42376.font.size = Pt(12)
            t42376.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.6 ###
    #########################
    tituloTabla42376 = doc.add_paragraph()
    dti42376 = tituloTabla42376.add_run('\nTabla 4.x.- Listado de las especies de lepidópteros con su estatus u categorías por especie en el sistema ambiental.')
    dti42376_format = tituloTabla42376.paragraph_format
    dti42376_format.line_spacing = 1.15
    dti42376_format.space_after = 0

    dti42376.font.name = 'Courier New'
    dti42376.font.size = Pt(12)
    tituloTabla42376.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.6 ###
    #########################
    tabla42376 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42376.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42376.cell(rows, cols)
            t42376 = cell.paragraphs[0].add_run(' ')
            t42376.font.size = Pt(12)
            t42376.font.name = 'Arial'

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.6 ###
    #########################
    tituloTabla42376 = doc.add_paragraph()
    dti42376 = tituloTabla42376.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad Shannon, Simpson y Margalef para las especies de lepidópteros observados en el sistema ambiental.')
    dti42376_format = tituloTabla42376.paragraph_format
    dti42376_format.line_spacing = 1.15
    dti42376_format.space_after = 0

    dti42376.font.name = 'Courier New'
    dti42376.font.size = Pt(12)
    tituloTabla42376.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.6 ###
    #########################
    tabla42376 = doc.add_table(rows=10, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla42376.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42376.cell(rows, cols)
            t42376 = cell.paragraphs[0].add_run(' ')
            t42376.font.size = Pt(12)
            t42376.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.6 ###
    #########################
    di42376 = doc.add_paragraph()
    descripcionCapitulo42376 = di42376.add_run('Los índices de diversidad de las 7 especies de lepidópteros presentes en el área del sistema ambiental muestran que para el índice de Shannon tenemos una equidad de 1.4328 lo cual quiere decir que para este grupo los valores resultantes se encuentran a niveles medios ya que los rangos de valores para este índice van de 1.35 a 3 para valores normales por los contrario para los valores inferiores serian aquellos inferiores a 1.35; para el índice de Simpson resulta una diversidad media de 0.666 y una dominancia de las especies media de 0.334, lo cual quiere decir que podemos encontrar especies dominantes, por otra parte el índice de Margalef el cual estima la biodiversidad de una comunidad  muestra valores bajos de 1.7312 ya que los valores de medida considerados para una baja biodiversidad son para valores inferiores a 2 e indicadores de una alta biodiversidad son aquellos con valores superiores a 3 y muy alta con valores de 5; para el grupo de los la especie más representativa fue Zerene cesonia.')
    descripcionCapitulo42376_format = di42376.paragraph_format
    descripcionCapitulo42376_format.line_spacing = 1.15
    descripcionCapitulo42376_format.space_after = 0
    descripcionCapitulo42376_format.space_before = 0

    descripcionCapitulo42376.font.name = 'Arial'
    descripcionCapitulo42376.font.size = Pt(12)
    di42376.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo de la tabla del capitulo 4.2.3.7.6 ###
    #########################
    tituloTabla42376 = doc.add_paragraph()
    dti42376 = tituloTabla42376.add_run('\nTabla 4.X.- Análisis estadístico por índices de diversidad, riqueza de especies, frecuencia y abundancia relativa para las especies de lepidópteros en el área del sistema ambiental.')
    dti42376_format = tituloTabla42376.paragraph_format
    dti42376_format.line_spacing = 1.15
    dti42376_format.space_after = 0

    dti42376.font.name = 'Courier New'
    dti42376.font.size = Pt(12)
    tituloTabla42376.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.6 ###
    #########################
    tabla42376 = doc.add_table(rows=10, cols=11, style='Table Grid')

    for cols in range(11):
        cell = tabla42376.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla42376.cell(rows, cols)
            t42376 = cell.paragraphs[0].add_run(' ')
            t42376.font.size = Pt(12)
            t42376.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.6 ###
    #########################
    di42376 = doc.add_paragraph()
    descripcionCapitulo42376 = di42376.add_run('\nDescripcion.')
    descripcionCapitulo42376_format = di42376.paragraph_format
    descripcionCapitulo42376_format.line_spacing = 1.15
    descripcionCapitulo42376_format.space_after = 0
    descripcionCapitulo42376_format.space_before = 0

    descripcionCapitulo42376.font.name = 'Arial'
    descripcionCapitulo42376.font.size = Pt(12)
    di42376.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 4.2.3.7.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo42375 = doc.add_paragraph()
    imagenCapitulo42375.text = ''
    imagenCapitulo42375 = doc.add_picture('capitulo4/grafico.jpg')  # Ancho de la imagen en centimetros
    imagenCapitulo42375.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo42375.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo42375.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo42375.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Titulo de la grafica del capitulo 4.2.3.7.5 ###
    #########################
    tituloGrafico42375 = doc.add_paragraph()
    dgi42375 = tituloGrafico42375.add_run('Grafica 4.X.- Abundancia y frecuencia relativa para el grupo de los lepidópteros en el área del sistema ambiental.\n')
    dgi42375_format = tituloGrafico42375.paragraph_format
    dgi42375_format.line_spacing = 1.15
    dgi42375_format.space_after = 0

    dgi42375.font.name = 'Bookman Old Style'
    dgi42375.font.size = Pt(12)
    tituloGrafico42375.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=5, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 4.2.3.7.5 ###
    #########################
    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('\nDescripcion Parrafo 1.')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di42375 = doc.add_paragraph()
    descripcionCapitulo42375 = di42375.add_run('\nDescripcion Parrafo 2.\n')
    descripcionCapitulo42375_format = di42375.paragraph_format
    descripcionCapitulo42375_format.line_spacing = 1.15
    descripcionCapitulo42375_format.space_after = 0
    descripcionCapitulo42375_format.space_before = 0

    descripcionCapitulo42375.font.name = 'Arial'
    descripcionCapitulo42375.font.size = Pt(12)
    di42375.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 4.2.3.7.5 ###
    #########################
    tabla42375 = doc.add_table(rows=6, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla42375.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla42375.cell(rows, cols)
            t42375 = cell.paragraphs[0].add_run(' ')
            t42375.font.size = Pt(12)
            t42375.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 4.2.3.7.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 4.2.3.7.7 ###
    #########################
    capitulo42377 = doc.add_paragraph()
    i42377 = capitulo42377.add_run(f'\n{temasCapitulo4[1][2][4][11][6]}')
    i42377_format = capitulo42377.paragraph_format
    i42377_format.line_spacing = 1.15

    i42377.font.name = 'Arial'
    i42377.font.size = Pt(12)
    i42377.font.bold = True
    capitulo42377.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 4.2.3.7.7 ###
    #########################
    di42377 = doc.add_paragraph()
    descripcionCapitulo42377 = di42377.add_run('Descripcion del Capitulo')
    descripcionCapitulo42377_format = di42377.paragraph_format
    descripcionCapitulo42377_format.line_spacing = 1.15
    descripcionCapitulo42377_format.space_after = 0
    descripcionCapitulo42377_format.space_before = 0

    descripcionCapitulo42377.font.name = 'Arial'
    descripcionCapitulo42377.font.size = Pt(12)
    di42377.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO X DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO X DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO X DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 4 DTU EXTRACCION DE MATERIAL PETRO.docx")


"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo4()  # Llamar a la función para ejecutar el código
