"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas

def capitulo2():
    doc = Document()        # -----> Variable global del documento

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 2
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice Capítulo I.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    temasCapitulo2 = ["II.1.- Objetivos del Proyecto", 
                    "II.2.- Naturaleza del Proyecto", 
                    "II.3.- Justificación por que los terrenos son apropiados al nuevo uso", 
                    "II.4.- Programa de Trabajo"]
    
    ########################################################################################################################################################################
    # Comienza Contenido
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    # Capitulo 2
    capitulo2 = doc.add_paragraph()
    ii = capitulo2.add_run("II.- USOS QUE SE PRETENDEN DAR AL TERRENO.")
    ii_format = capitulo2.paragraph_format
    ii_format.line_spacing = 1.5

    ii.font.name = 'Arial'
    ii.font.size = Pt(12)
    ii.font.bold = True

    ### PRIMER DESCRIPCION DEL CAPITULO 2 ###
    dii201 = doc.add_paragraph()
    descripcionCapitulo201 = dii201.add_run('El área en estudio está clasificado como Uso de Terreno Forestal, para el nuevo uso que se tiene proyectado dar a este terreno, es para la ____________________________, del municipio de _______________________, que ostenta vegetación de tipo _________________________ ocupando una superficie de ______ ha la que representa el 100 %, la cual se clasifica  como Uso Forestal, de ahí la necesidad de realizar el cambio de uso de suelo y se contempla llevar a cabo la remoción de la vegetación, razón por la cual se elabora el presente instrumento para su análisis, evaluación y gestionar la autorización correspondiente para llevar a cabo las acciones planeadas que contemplan el Documento Técnico Unificado por Cambio de uso de suelo, bajo lo siguiente:')
    descripcionCapitulo201_format = dii201.paragraph_format
    descripcionCapitulo201_format.line_spacing = 1.15

    descripcionCapitulo201.font.name = 'Arial'
    descripcionCapitulo201.font.size = Pt(12)
    dii201.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ### PRIMER SMARTART DEL CAPITULO 2 ###
    """ 
        No se tiene soporte para poner diagramas de SmartArt, se tienen que poner de manera manual o por imagen.
    """
    imagenCapitulo2 = doc.add_picture('capitulo2.png')  # Ancho de la imagen en centimetros
    imagenCapitulo2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo2.width = Cm(17.07)  # Ancho de la imagen en centimetros
    imagenCapitulo2.height = Cm(6.56)  # Alto de la imagen en centimetros

    ### PEQUEÑA DESCRIPCION DEL CAPITULO 2 ###
    dii2021 = doc.add_paragraph()
    descripcionCapitulo2021 = dii2021.add_run('\nObras complementarias después de las actividades de cambio de uso de suelo.')
    descripcionCapitulo2021_format = dii2021.paragraph_format
    descripcionCapitulo2021_format.line_spacing = 1.15

    descripcionCapitulo2021.font.name = 'Arial'
    descripcionCapitulo2021.font.size = Pt(12)
    dii2021.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ### SEGUNDO SMARTART DEL CAPITULO 2 ###
    """ 
        * No se tiene soporte para poner diagramas de SmartArt, se tienen que poner de manera manual, o por imagen.
        * Se puede poner una imagen de un SmartArt que se haya hecho en Word, PowerPoint o Excel, y guardado como imagen.
        * Se recomienda usar una tabla para representar los datos de manera más clara.
    """
    tablaCapitulo2 = doc.add_table(rows=3, cols=3, style='Table Grid')

    # Fusionar celdas de la primara fila
    row = tablaCapitulo2.rows[0]
    merged_cell = row.cells[0].merge(row.cells[1].merge(row.cells[2]))

    # Agregar texto a la celda fusionada
    t20 = merged_cell.paragraphs[0].add_run('Obras complementarias')
    t20.font.name = 'Arial'
    t20.font.size = Pt(12)
    t20.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, 'D5DBDB')  # Cambiar el color de fondo de la celda fusionada

    # Resto de las celdas
    cell = tablaCapitulo2.cell(1, 0)
    t21 = cell.paragraphs[0].add_run('1. Construccion')
    t21.font.name = 'Arial'
    t21.font.size = Pt(12)
    t21.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '82E0AA')

    cell = tablaCapitulo2.cell(1, 1)
    t21 = cell.paragraphs[0].add_run('2. Operación')
    t21.font.name = 'Arial'
    t21.font.size = Pt(12)
    t21.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '5499C7')

    cell = tablaCapitulo2.cell(1, 2)
    t21 = cell.paragraphs[0].add_run('3. Abandono de Sitio')
    t21.font.name = 'Arial'
    t21.font.size = Pt(12)
    t21.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'D98880')

    terceraFila = range(3)  # Tercera fila de la tabla
    actividades = range(3)  # Actividades de la tercera fila simulacion

    colores3Fila = ['D5F5E3', 'A9CCE3', 'F2D7D5']  # Colores de la tercera fila

    for i in terceraFila:
        cell = tablaCapitulo2.cell(2, i)

        for j in actividades:
            t21 = cell.paragraphs[0].add_run(f'Actividad {j + 1}\n')
            t21.font.name = 'Arial'
            t21.font.size = Pt(12)
            
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cell_background_color(cell, f'{colores3Fila[i]}')

    

    ### SEGUNDA DESCRIPCINO DEL CAPITULO 2 ###
    dii202 = doc.add_paragraph()
    descripcionCapitulo202 = dii202.add_run('\nCon las actividades antes descritas se contempla cubrir todas las necesidades que llevará cada una de las etapas para la implementación en ___________________________.')
    descripcionCapitulo202_format = dii202.paragraph_format
    descripcionCapitulo202_format.line_spacing = 1.15

    descripcionCapitulo202.font.name = 'Arial'
    descripcionCapitulo202.font.size = Pt(12)
    dii202.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.1
    ########################################################################################################################################################################
    ii21 = doc.add_paragraph()
    tituloCapitulo21 = ii21.add_run(f'\n{temasCapitulo2[0]}')
    tituloCapitulo21_format = ii21.paragraph_format
    tituloCapitulo21_format.line_spacing = 1.5
    tituloCapitulo21.bold = True

    tituloCapitulo21.font.name = 'Arial'
    tituloCapitulo21.font.size = Pt(12)
    ii21.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di21 = doc.add_paragraph()
    descripcionCapitulo21 = di21.add_run('Establecer las condiciones ideales para la preparación, construcción y Operación de un _____________________________________________________________________.')
    descripcionCapitulo21_format = di21.paragraph_format
    descripcionCapitulo21_format.line_spacing = 1.15

    descripcionCapitulo21.font.name = 'Arial'
    descripcionCapitulo21.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 2.1.1
    ########################################################################################################################################################################
    capitulo211 = doc.add_paragraph()
    tituloCapitulo211 = capitulo211.add_run('\nII.1.1.- Objetivos Especificos')
    tituloCapitulo211_format = capitulo211.paragraph_format
    tituloCapitulo211_format.line_spacing = 1.15
    tituloCapitulo211.bold = True

    tituloCapitulo211.font.name = 'Arial'
    tituloCapitulo211.font.size = Pt(12)

    objetivosCapitulo211 = ['Elaborar el Documento Técnico Unificado para el Cambio de Uso de Suelo para ________________________.',
                            'Establecer las condiciones ideales para __________________________ en las etapas de preparación del sitio, construcción, operación y abandono del sitio.',
                            'Realizar el cambio de uso de suelo de terrenos forestales a uso ______ (_______________________).']

    for objetivo in objetivosCapitulo211:
        dii211 = doc.add_paragraph(style='ListBullet')
        descripcionCapitulo2111 = dii211.add_run(f'{objetivo}')
        descripcionCapitulo2111_format = dii211.paragraph_format
        descripcionCapitulo2111_format.line_spacing = 1.15
        descripcionCapitulo2111.font.name = 'Arial'
        descripcionCapitulo2111.font.size = Pt(12)
        descripcionCapitulo2111.space_before = Pt(1)
        descripcionCapitulo2111.space_after = Pt(0)
        dii211.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.1.2
    ########################################################################################################################################################################
    capitulo212 = doc.add_paragraph()
    tituloCapitulo212 = capitulo212.add_run('\nII.1.2.- Antecedentes del Proyecto')
    tituloCapitulo212_format = capitulo212.paragraph_format
    tituloCapitulo212_format.line_spacing = 1.15
    tituloCapitulo212.bold = True

    tituloCapitulo212.font.name = 'Arial'
    tituloCapitulo212.font.size = Pt(12)
    capitulo212.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    capitulo212.paragraph_format.space_before = Pt(0)

    di212 = doc.add_paragraph()
    descripcionCapitulo212 = di212.add_run('Descripcion de este capitulo')
    descripcionCapitulo212_format = di212.paragraph_format
    descripcionCapitulo212_format.line_spacing = 1.15

    descripcionCapitulo212.font.name = 'Arial'
    descripcionCapitulo212.font.size = Pt(12)
    di212.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    di212.paragraph_format.space_before = Pt(0)

    ########################################################################################################################################################################
    # Capitulo 2.2
    ########################################################################################################################################################################
    capitulo22 = doc.add_paragraph()
    tituloCapitulo22 = capitulo22.add_run(f'\n{temasCapitulo2[1]}')
    tituloCapitulo22_format = capitulo22.paragraph_format
    tituloCapitulo22_format.line_spacing = 1.5
    tituloCapitulo22.bold = True

    tituloCapitulo22.font.name = 'Arial'
    tituloCapitulo22.font.size = Pt(12)
    capitulo22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        Etapas del capitulo 2.2 (Pueden ser modificados en caso de que cambie el formato del documento)
    """
    etapasCapitulo22 = ['1.- Preparecion de Sitio',
                        '2.- Construcción',
                        '3.- Operación',
                        '4.- Abandono se sitio']
    
    descripcionEtapasCapitulo22 = [
        'El cual consta de la delimitación del área, el rescate de flora y fauna, remoción total de la vegetación y despalme',
        '________________',
        '________________',
        'En cuanto al abandono del sitio esta consistirá en la clausura del área, estabilización de los taludes y actividades de restauración este en apego al concepto de pago por compensación ambiental'
    ]

    dii22 = doc.add_paragraph()

    # Primer Texto del capitulo 2.2
    descripcionCapitulo22_1 = dii22.add_run('El área en estudio se encuentra ubicada _________________,  misma que ostenta vegetación de _____________ la cual se clasifica  como Uso Forestal, de ahí la necesidad de realizar el cambio de uso de suelo, misma que será para realizar el aprovechamiento _____________, para la cual será necesario establecer la condiciones necesaria para su operación, para el establecimiento del nuevo uso, se pretende realizar las actividades, solamente dentro del área autorizada, la cual se clasifica en Cuatro  etapas, ')
    descripcionCapitulo22_1_format = dii22.paragraph_format
    descripcionCapitulo22_1_format.line_spacing = 1.15
    
    descripcionCapitulo22_1.font.name = 'Arial'
    descripcionCapitulo22_1.font.size = Pt(12)

    # Etapas del capitulo 2.2
    """
        Etapas del capitulo 2.2 (Pueden ser modificados en caso de que cambie el formato del documento)
    """
    etapasCapitulo22 = ['1.- Preparecion de Sitio',
                        '2.- Construcción',
                        '3.- Operación',
                        '4.- Abandono se sitio']
    
    descripcionEtapasCapitulo22 = [
        'El cual consta de la delimitación del área, el rescate de flora y fauna, remoción total de la vegetación y despalme',
        '________________',
        '________________',
        'En cuanto al abandono del sitio esta consistirá en la clausura del área, estabilización de los taludes y actividades de restauración este en apego al concepto de pago por compensación ambiental'
    ]
    
    # Etapa 1
    etapa1Capitulo22 = dii22.add_run(f' {etapasCapitulo22[0]}.- ')
    etapa1Capitulo22_format = dii22.paragraph_format
    etapa1Capitulo22_format.line_spacing = 1.15

    etapa1Capitulo22.font.name = 'Arial'
    etapa1Capitulo22.font.size = Pt(12)
    etapa1Capitulo22.bold = True

    descEtapa1Capitulo22 = dii22.add_run(f'{descripcionEtapasCapitulo22[0]}; ')
    descEtapa1Capitulo22_format = dii22.paragraph_format
    descEtapa1Capitulo22_format.line_spacing = 1.15

    descEtapa1Capitulo22.font.name = 'Arial'
    descEtapa1Capitulo22.font.size = Pt(12)

    # Etapa 2
    etapa2Capitulo22 = dii22.add_run(f' {etapasCapitulo22[1]}.- ')
    etapa2Capitulo22_format = dii22.paragraph_format
    etapa2Capitulo22_format.line_spacing = 1.15

    etapa2Capitulo22.font.name = 'Arial'
    etapa2Capitulo22.font.size = Pt(12)
    etapa2Capitulo22.bold = True

    descEtapa2Capitulo22 = dii22.add_run(f'{descripcionEtapasCapitulo22[1]}; ')
    descEtapa2Capitulo22_format = dii22.paragraph_format
    descEtapa2Capitulo22_format.line_spacing = 1.15

    descEtapa2Capitulo22.font.name = 'Arial'
    descEtapa2Capitulo22.font.size = Pt(12)

    # Etapa 3
    etapa3Capitulo22 = dii22.add_run(f' {etapasCapitulo22[2]}.- ')
    etapa3Capitulo22_format = dii22.paragraph_format
    etapa3Capitulo22_format.line_spacing = 1.15

    etapa3Capitulo22.font.name = 'Arial'
    etapa3Capitulo22.font.size = Pt(12)
    etapa3Capitulo22.bold = True

    descEtapa3Capitulo22 = dii22.add_run(f'{descripcionEtapasCapitulo22[2]}; ')
    descEtapa3Capitulo22_format = dii22.paragraph_format
    descEtapa3Capitulo22_format.line_spacing = 1.15

    descEtapa3Capitulo22.font.name = 'Arial'
    descEtapa3Capitulo22.font.size = Pt(12)

    # Etapa 4
    etapa4Capitulo22 = dii22.add_run(f' {etapasCapitulo22[3]}.- ')
    etapa4Capitulo22_format = dii22.paragraph_format
    etapa4Capitulo22_format.line_spacing = 1.15

    etapa4Capitulo22.font.name = 'Arial'
    etapa4Capitulo22.font.size = Pt(12)
    etapa4Capitulo22.bold = True

    descEtapa4Capitulo22 = dii22.add_run(f'{descripcionEtapasCapitulo22[3]}; ')
    descEtapa4Capitulo22_format = dii22.paragraph_format
    descEtapa4Capitulo22_format.line_spacing = 1.15

    descEtapa4Capitulo22.font.name = 'Arial'
    descEtapa4Capitulo22.font.size = Pt(12)

    # Segundo Texto del Capitulo 2.2
    descripcionCapitulo22_2 = dii22.add_run('cualquiera que sea el tipo de actividades siempre es causante de impactos, los factores se pueden poner en riesgo con la implementación del proyecto son principalmente al recurso suelo (propiedades físicas y Erodabilidad), con la ruptura del suelo, vegetación (flora y fauna), topografía (relieve), paisaje (armonía y calidad) y por supuesto el recurso agua (hidrología superficial y subterránea), así como los componentes de la atmósfera (calidad del aire y ruido), estos factores se verán afectados en un extensión de 3.99 ha. Dentro del área de Cambio de uso de suelo no se encuentran cuerpos de agua permanentes que pueda causar desequilibrio en este componente, solamente se manifiestan escurrimientos de tipo superficial en temporadas de lluvia, el cual para no afectar la recarga de los mantos freáticos. De acuerdo a los criterios de regulación aplicables al proyecto se tiene lo siguiente: ')
    descripcionCapitulo22_2_format = dii22.paragraph_format
    descripcionCapitulo22_2_format.line_spacing = 1.15

    descripcionCapitulo22_2.font.name = 'Arial'
    descripcionCapitulo22_2.font.size = Pt(12)

    # Programas de Ordenamiento Ecológico del capitulo 2.2
    """
        Programas de Ordenamiento Ecológico del capitulo 2.2 (Pueden ser modificados en caso de que cambie el formato del documento)
    """

    programasCapitulo22 = ['Programa de Ordenamiento Ecológico General del Territorio Federal',
                           'Programa de Ordenamiento Ecológico General del Territorio del Estado de Coahuila',
                           'Programa De Ordenamiento Ecológico De La Región Cuenca De Burgos']
    
    descripcionProgramasCapitulo22 = ['el presente proyecto realiza acciones de conservación y protección a los ecosistemas y la biodiversidad, así como   recuperación de especies en riesgo, mediante la ejecución de un programa de rescate de flora y fauna. ',
                                      'se tiene lo siguiente: se promueve la conservación de la comunidades vegetales y faunísticas así como el manejo de hábitats, la recuperación de especies que estén en peligro de extinción o amenazadas, con la implementación de un programa de rescate de flora y fauna, se permitirá la continuidad con los predios aledaños al dejar sin aprovechamiento áreas del mismo predio, se disminuirá el riesgo de incendio al actuar la periferia como brecha cortafuego, también se realizaran acciones de conservación esto en apego al pago por compensación ambiental, para el ',
                                      'se tiene lo siguiente: para las acciones de conservación y restauración será en apego al pago por compensación ambiental, para la conservación de la cubierta vegetal se realizara rescate de flora silvestre, se podrá utilizara aguas tratadas para mitigar impactos por emisiones de polvos, se tendrá registro de la generación y descargas de residuos de cualquier tipo, no se modificaran cuerpos de agua de ningún tipo, se realizan platicas de educación ambiental a los trabajadores involucrados en cada una de las etapas del proyecto.']
    

    # Primer Programa de Ordenamiento Ecológico del capitulo 2.2
    programa1Capitulo22 = dii22.add_run(f'{programasCapitulo22[0]}, ')
    programa1Capitulo22_format = dii22.paragraph_format
    programa1Capitulo22_format.line_spacing = 1.15
    
    programa1Capitulo22.font.name = 'Arial'
    programa1Capitulo22.font.size = Pt(12)
    programa1Capitulo22.bold = True

    descPrograma1Capitulo22 = dii22.add_run(f'{descripcionProgramasCapitulo22[0]} ')
    descPrograma1Capitulo22_format = dii22.paragraph_format
    descPrograma1Capitulo22_format.line_spacing = 1.15

    descPrograma1Capitulo22.font.name = 'Arial'
    descPrograma1Capitulo22.font.size = Pt(12)
    dii22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo Programa de Ordenamiento Ecológico del capitulo 2.2
    programa2Capitulo22 = dii22.add_run(f' {programasCapitulo22[1]}, ')
    programa2Capitulo22_format = dii22.paragraph_format
    programa2Capitulo22_format.line_spacing = 1.15
    
    programa2Capitulo22.font.name = 'Arial'
    programa2Capitulo22.font.size = Pt(12)
    programa2Capitulo22.bold = True

    descPrograma2Capitulo22 = dii22.add_run(f'{descripcionProgramasCapitulo22[1]} ')
    descPrograma2Capitulo22_format = dii22.paragraph_format
    descPrograma2Capitulo22_format.line_spacing = 1.15

    descPrograma2Capitulo22.font.name = 'Arial'
    descPrograma2Capitulo22.font.size = Pt(12)
    dii22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Tercer Programa de Ordenamiento Ecológico del capitulo 2.2
    programa3Capitulo22 = dii22.add_run(f' {programasCapitulo22[2]}, ')
    programa3Capitulo22_format = dii22.paragraph_format
    programa3Capitulo22_format.line_spacing = 1.15
    
    programa3Capitulo22.font.name = 'Arial'
    programa3Capitulo22.font.size = Pt(12)
    programa3Capitulo22.bold = True

    descPrograma3Capitulo22 = dii22.add_run(f'{descripcionProgramasCapitulo22[2]} ')
    descPrograma3Capitulo22_format = dii22.paragraph_format
    descPrograma3Capitulo22_format.line_spacing = 1.15

    descPrograma3Capitulo22.font.name = 'Arial'
    descPrograma3Capitulo22.font.size = Pt(12)
    dii22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.3
    ########################################################################################################################################################################
    capitulo23 = doc.add_paragraph()
    tituloCapitulo23 = capitulo23.add_run(f'\n{temasCapitulo2[2]}')
    tituloCapitulo23_format = capitulo23.paragraph_format
    tituloCapitulo23_format.line_spacing = 1.5
    tituloCapitulo23.bold = True
    
    tituloCapitulo23.font.name = 'Arial'
    tituloCapitulo23.font.size = Pt(12)

    dii23 = doc.add_paragraph()
    descripcionCapitulo23 = dii23.add_run('Los terrenos en donde se pretende establecer el proyecto ==Nombre Proyecto y Ubicacion==, Coahuila, en el ==Ubicacion==, Para llegar al área en estudio, ==Como Llegar==, donde se encuentra el área de estudio, al encontrase fuera de las manchas urbanas, se considera adecuado para el nuevo uso de suelo que es la _______________________, además el uso actual del predio se encuentra sin uso y se aplicara un programa de rescate y reubicación de flora y fauna principalmente aquellas que se encuentren dentro de la NOM-059 SEMARNAT-2010, así como aquellas de lento desplazamiento y crecimiento, en cuanto a la vegetación que presenta es del tipo _________________ ocupando una superficie de ____ ha la que representa el 100 %,se encuentra en vegetación primaria en condiciones regulares al ser eliminada no presenta disminución biológica dentro del sistema ambiental que se encuentran representadas las especies, además  en donde se pretende establecer el nuevo proyecto de acuerdo a las características generales del área no afecta a cuerpos de agua importantes como perennes solo se presentan escurrimientos en época de lluvias al encontrase el área en una ladera, el cual no tendrá ningún impacto o que afecte a las corrientes subterráneas, que requieran acciones especiales o que afecten la hidrología del sistema ambiental. En conclusión, el implementar el proyecto __________________, no pone en riesgo la pérdida de Biodiversidad al tener dentro del sistema ambiental la presencia de las especies que serán sujetas a eliminación, se mantendrá la biodiversidad biológica al efectuarse un programa de rescate de flora y fauna silvestre. no se disminuirá la calidad ni cantidad de agua, ni se modificarán cuerpos de agua, Además, durante todos los procesos del proyecto en las emisiones de polvo a la atmósfera, se reducirá con la humectación del suelo utilizando aguas recicladas, así como reducción de smog a la atmosfera, con la aplicación de mantenimientos preventivos a las maquinarias de acuerdo a la NOM-045-SEMARNAT-2006, se realizarán actividades de restauración esto en apego al concepto del pago por compensación ambiental.')
    descripcionCapitulo23_format = dii23.paragraph_format
    descripcionCapitulo23_format.line_spacing = 1.15

    descripcionCapitulo23.font.name = 'Arial'
    descripcionCapitulo23.font.size = Pt(12)
    dii23.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.4
    ########################################################################################################################################################################
    capitulo24 = doc.add_paragraph()
    tituloCapitulo24 = capitulo24.add_run(f'\n{temasCapitulo2[3]}')
    tituloCapitulo24_format = capitulo23.paragraph_format
    tituloCapitulo24_format.line_spacing = 1.5
    tituloCapitulo24.bold = True
    
    tituloCapitulo24.font.name = 'Arial'
    tituloCapitulo24.font.size = Pt(12)

    dii24 = doc.add_paragraph()
    descripcionCapitulo24 = dii24.add_run('Para la ejecución en donde se requiere la ampliación de _____________________________, implica la remoción total de vegetación forestal, para ellos se requieren objetivos bien definidos y planeados para no poner en riesgo los factores ambientales que dan sustento a su entorno, se contempla en primer lugar la visita preliminar para analizar la posibilidad o si es factible realizar el cambio de uso del suelo en el predio en mención posteriormente se lleva a cabo el estudio completo a través del cual se evaluará cada proceso con base a la normatividad establecida y mediante el cual se dará sustento para determinar su autorización o negación de la ejecución siguiendo el orden que se enlista, previo a la ejecución del cambio de uso de suelo se realizará lo siguiente:')
    descripcionCapitulo24_format = dii23.paragraph_format
    descripcionCapitulo24_format.line_spacing = 1.15

    descripcionCapitulo24.font.name = 'Arial'
    descripcionCapitulo24.font.size = Pt(12)
    dii24.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #lista del capitulo 2.4
    """"
        Lista del Capitulo 2.4 (Pueden ser modificados en caso de que cambie el formato del documento)
    """
    listaCapitulo24 = [
        'Elaboración del plano topográfico del área sujeto de estudio _____________________.',
        'Delimitación del predio y área sujeta de estudio.',
        'Estudio de campo (sitios de muestreo, registro de vegetación por estrato, transectos para aplicación de métodos de muestreo de Fauna, toma de fotografías de condición actual y características generales de los medios físicos y biológicos tanto del predio en general como en forma específica del área sujeta a cambio de uso de suelo).,'
        'Elaboración de documento en base a información de campo y revisión bibliográfica.',
        'Revisión y análisis de estudio para autorización del cambio de uso de suelo.'
    ]

    for lista in listaCapitulo24:
        lista24 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo24_1 = lista24.add_run(f'{lista}')
        descripcionCapitulo24_1_format = lista24.paragraph_format
        descripcionCapitulo24_1_format.line_spacing = 1.15

        descripcionCapitulo24_1.font.name = 'Arial'
        descripcionCapitulo24_1.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 2.4.1
    ########################################################################################################################################################################
    capitulo241 = doc.add_paragraph()
    tituloCapitulo241 = capitulo241.add_run('\nII.4.1.- Preparación del Sitio')
    tituloCapitulo241_format = capitulo241.paragraph_format
    tituloCapitulo241_format.line_spacing = 1.5

    tituloCapitulo241.bold = True
    tituloCapitulo241.font.name = 'Arial'
    tituloCapitulo241.font.size = Pt(12)

    # Lista del Capitulo 2.4.1
    """
        Lista del Capitulo 2.4.1 (Pueden ser modificados en caso de que cambie el formato del documento)
    """
    listaCapitulo241 = [
        'Delimitación',
        'Prgrama de Rescate de Flora y Fauna',
        'Remoción de Vegetación (desmonte)',
        'Despalme',
    ]

    for lista in listaCapitulo241:
        lista241 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo241_1 = lista241.add_run(f'{lista}')
        descripcionCapitulo241_1_format = lista241.paragraph_format
        descripcionCapitulo241_1_format.line_spacing = 1.15

        descripcionCapitulo241_1.font.name = 'Arial'
        descripcionCapitulo241_1.font.size = Pt(12)
        lista241.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.4.2 (Opcional)
    ########################################################################################################################################################################
    capitulo242 = doc.add_paragraph()
    tituloCapitulo242 = capitulo242.add_run('\nII.4.2.- Construcción')
    tituloCapitulo242_format = capitulo242.paragraph_format
    tituloCapitulo242_format.line_spacing = 1.5

    tituloCapitulo242.bold = True
    tituloCapitulo242.font.name = 'Arial'
    tituloCapitulo242.font.size = Pt(12)

    # Lista del Capitulo 2.4.2
    """
        Lista del Capitulo 2.4.2 (Pueden ser modificados en caso de que cambie el formato del documento, Son opcionales el texto de este capitulo)
    """
    listaCapitulo242 = [
        'Actvididad 1',
        'Actvididad 2',
        'Actvididad 3',
    ]

    for lista in listaCapitulo242:
        lista242 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo242_1 = lista242.add_run(f'{lista}')
        descripcionCapitulo242_1_format = lista242.paragraph_format
        descripcionCapitulo242_1_format.line_spacing = 1.15

        descripcionCapitulo242_1.font.name = 'Arial'
        descripcionCapitulo242_1.font.size = Pt(12)
        lista242.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.4.3 (Opcional)
    ########################################################################################################################################################################
    capitulo243 = doc.add_paragraph()
    tituloCapitulo243 = capitulo243.add_run('\nII.4.3.- Operación')
    tituloCapitulo243_format = capitulo243.paragraph_format
    tituloCapitulo243_format.line_spacing = 1.5

    tituloCapitulo243.bold = True
    tituloCapitulo243.font.name = 'Arial'
    tituloCapitulo243.font.size = Pt(12)

    # Lista del Capitulo 2.4.3
    """
        Lista del Capitulo 2.4.3 (Pueden ser modificados en caso de que cambie el formato del documento, Son opcionales el texto de este capitulo)
    """
    listaCapitulo243 = [
        'Actvididad 1',
        'Actvididad 2',
        'Actvididad 3',
    ]

    for lista in listaCapitulo243:
        lista243 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo243_1 = lista243.add_run(f'{lista}')
        descripcionCapitulo243_1_format = lista243.paragraph_format
        descripcionCapitulo243_1_format.line_spacing = 1.15

        descripcionCapitulo243_1.font.name = 'Arial'
        descripcionCapitulo243_1.font.size = Pt(12)
        lista243.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.4.4
    ########################################################################################################################################################################
    capitulo244 = doc.add_paragraph()
    tituloCapitulo244 = capitulo244.add_run('\nII.4.4.- Abandono del Sitio')
    tituloCapitulo244_format = capitulo244.paragraph_format
    tituloCapitulo244_format.line_spacing = 1.5

    tituloCapitulo244.bold = True
    tituloCapitulo244.font.name = 'Arial'
    tituloCapitulo244.font.size = Pt(12)

    # Lista del Capitulo 2.4.4
    """
        Lista del Capitulo 2.4.4 (Pueden ser modificados en caso de que cambie el formato del documento)
    """
    listaCapitulo244 = [
        'Clausura del Área',
        'Restaruración del Sitio',
    ]

    for lista in listaCapitulo244:
        lista244 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo244_1 = lista244.add_run(f'{lista}')
        descripcionCapitulo244_1_format = lista244.paragraph_format
        descripcionCapitulo244_1_format.line_spacing = 1.15

        descripcionCapitulo244_1.font.name = 'Arial'
        descripcionCapitulo244_1.font.size = Pt(12)
        lista244.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 2.4.4.1
    ########################################################################################################################################################################
    capitulo2441 = doc.add_paragraph()
    tituloCapitulo2441 = capitulo2441.add_run('\nII.4.4.1.- Calendario de Ejecución del Proyecto')
    tituloCapitulo2441_format = capitulo2441.paragraph_format
    tituloCapitulo2441_format.line_spacing = 1.5

    tituloCapitulo2441.bold = True
    tituloCapitulo2441.font.name = 'Arial'
    tituloCapitulo2441.font.size = Pt(12)

    di2441 = doc.add_paragraph()
    descripcionCapitulo2441 = di2441.add_run(f'Cronograma de Actividades, el presente calendario está en función de la obtención del permiso y empezará a partir de este, bajo lo siguiente, Preparación del sitio (______), Construcción (______), Operación (______), Abandono del sitio (1 año), esta última etapa se realizará en el año ___, ya cuando todas las operaciones se den por concluido en el área.')
    descripcionCapitulo2441_format = di2441.paragraph_format
    descripcionCapitulo2441_format.line_spacing = 1.15

    descripcionCapitulo2441.font.name = 'Arial'
    descripcionCapitulo2441.font.size = Pt(12)
    di2441.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #di2441.add_run().add_break(WD_BREAK.PAGE)  # Salto de página para la tabla
    
    ########################################################################################################################################################################
    # Tabla 2.1
    ########################################################################################################################################################################
    # Agregar nueva sección con orientación horizontal
    # Crear nueva sección
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a horizontal
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    tabla1Capitulo2 = doc.add_table(rows=17, cols=11, style='Table Grid')

    cell =  tabla1Capitulo2.cell(0, 0)
    t21 = cell.paragraphs[0].add_run('PERIODO')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.font.bold = True
    t21._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cell =  tabla1Capitulo2.cell(1, 0)
    t21 = cell.paragraphs[0].add_run('ETAPA')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.font.bold = True
    t21._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    #########################
    # Celda fusionada "Años"
    row1 = tabla1Capitulo2.rows[0]
    merged_cell1 = row1.cells[1].merge(row1.cells[2].merge(row1.cells[10]))

    # Agregar texto a la celda fusionada
    t21 = merged_cell1.paragraphs[0].add_run('Años')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.font.color.rgb = RGBColor(51, 51, 51)  # Cambiar el color de la fuente
    t21.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell1, 'D5DBDB')  # Cambiar el color de fondo de la celda fusionada

    #########################
    # Celda fusionada "Preparacion"
    row1 = tabla1Capitulo2.rows[2]
    merged_cell1 = row1.cells[1].merge(row1.cells[2].merge(row1.cells[10]))

    # Agregar texto a la celda fusionada
    t21 = merged_cell1.paragraphs[0].add_run('Preparación')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Celda fusionada "Construccion"
    row1 = tabla1Capitulo2.rows[7]
    merged_cell1 = row1.cells[1].merge(row1.cells[2].merge(row1.cells[10]))

    # Agregar texto a la celda fusionada
    t21 = merged_cell1.paragraphs[0].add_run('Construcción')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Celda fusionada "Operación"
    row1 = tabla1Capitulo2.rows[10]
    merged_cell1 = row1.cells[1].merge(row1.cells[2].merge(row1.cells[10]))

    # Agregar texto a la celda fusionada
    t21 = merged_cell1.paragraphs[0].add_run('Operación')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Celda fusionada "Abandono del area"
    row1 = tabla1Capitulo2.rows[14]
    merged_cell1 = row1.cells[1].merge(row1.cells[2].merge(row1.cells[10]))

    # Agregar texto a la celda fusionada
    t21 = merged_cell1.paragraphs[0].add_run('Abandono del área')
    t21.font.name = 'Agency FB'
    t21.font.size = Pt(12)
    t21.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Actividades de la Preparacion
    for row in range(3, 7):
        cell = tabla1Capitulo2.cell(row, 0)
        t21 = cell.paragraphs[0].add_run(f'Actividad {row - 2}')
        t21.font.name = 'Agency FB'
        t21.font.size = Pt(12)

    #########################
    # Actividades de la  Construccion
    for row in range(8, 10):
        cell = tabla1Capitulo2.cell(row, 0)
        t21 = cell.paragraphs[0].add_run(f'Actividad {row - 2}')
        t21.font.name = 'Agency FB'
        t21.font.size = Pt(12)
    
    #########################
    # Actividades de la  Operacion
    for row in range(11, 14):
        cell = tabla1Capitulo2.cell(row, 0)
        t21 = cell.paragraphs[0].add_run(f'Actividad {row - 2}')
        t21.font.name = 'Agency FB'
        t21.font.size = Pt(12)
    
    #########################
    # Actividades de la Abandono del area
    for row in range(15, 17):
        cell = tabla1Capitulo2.cell(row, 0)
        t21 = cell.paragraphs[0].add_run(f'Actividad {row - 2}')
        t21.font.name = 'Agency FB'
        t21.font.size = Pt(12)
    

    # Agregar los años a las celdas debajo de la celda fusionada
    años = range(10)

    for year in años:
        cell = tabla1Capitulo2.cell(1, year + 1)
        t21 = cell.paragraphs[0].add_run(f'{year + 1}')
        t21.font.name = 'Agency FB'
        t21.font.size = Pt(12)
        t21.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, 'D5DBDB')  # Cambiar el color de fondo de las celdas de los años

    for row in range (0, 17):
        cell = tabla1Capitulo2.cell(row, 0)
        cell.width = Cm(7.54)
        
    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 2 DTU EXTRACCION DE MATERIAL PETRO.docx")

capitulo2()
