from docx import Document
from docx.shared import Pt  # Para el tamaño en puntos
from docx.oxml.ns import qn  # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

doc = Document()

########################################################################################################################################################################
# Indice Capitulo 1
p = doc.add_paragraph()

# Añadir texto con estilo personalizado
indice = p.add_run("Índice Capítulo I.")

# Cambiar el tipo de letra y tamaño
indice.font.name = 'Bookman Old Style'
indice.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

########################################################################################################################################################################
temasCapitulo1 = ["I.1.- Nombre del proyecto", 
                  "I.2.- Nombre o Razón Social del Promovente", 
                  "I.3.- Ubicación (dirección) del Promovente", 
                  "I.4.- Superficie Solicitada de Cambio de Uso de Suelo y Tipo de Vegetacion",
                  "I.5.- Duración del Proyecto"]

########################################################################################################################################################################
# Comienza Contenido
doc.add_page_break() # Salto de página

# Capitulo 1
capitulo1 = doc.add_paragraph()
i = capitulo1.add_run("I.- DATOS GENERALES DEL PROYECTO Y DEL PROMOVENTE")
i_format = capitulo1.paragraph_format
i_format.line_spacing = 1.5

i.font.name = 'Arial'
i.font.size = Pt(12)
i.font.bold = True

########################################################################################################################################################################
# Capitulo 1.1
capitulo11 = doc.add_paragraph()
i11 = capitulo11.add_run(temasCapitulo1[0])
i11.font.name = 'Arial'
i11.font.size = Pt(12)
i11.font.bold = True

capitulo11.add_run().add_break()  # Salto de línea

# Texto del Capitulo 1.1
di11 = doc.add_paragraph()
descripcionCapitulo11 = di11.add_run('Proyecto "Nombre del proyecto", que para la identificación en este documento se referirá al proyecto como área de extracción, Área de Estudio, extracción de material pétreo.')
descripcionCapitulo11_format = di11.paragraph_format
descripcionCapitulo11_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

descripcionCapitulo11.font.name = 'Arial'       # Tipo de Fuente
descripcionCapitulo11.font.size = Pt(12)        # Tamaño de la Fuente
di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

di11.add_run().add_break()

########################################################################################################################################################################
# Capitulo 1.2
capitulo12 = doc.add_paragraph()
i12 = capitulo12.add_run(temasCapitulo1[1])
i12_format = capitulo12.paragraph_format
i12_format.line_spacing = 1.5

i12.font.name = 'Arial'
i12.font.size = Pt(12)
i12.font.bold = True

# capitulo12.add_run().add_break()  # Salto de línea

# Texto del Capitulo 1.2
naEm = doc.add_paragraph()

nameEmpresa = naEm.add_run('NOMBRE DE LA EMPRESA')
nameEmpresa_format = naEm.paragraph_format
nameEmpresa_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

nameEmpresa.font.name = 'Arial'       # Tipo de Fuente
nameEmpresa.font.size = Pt(12)        # Tamaño de la Fuente
nameEmpresa.font.bold = True
# naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro


di12 = doc.add_paragraph()

descripcionCapitulo12 = di12.add_run('Persona moral legalmente constituida según consta en el instrumento número 552 de fecha 16 de noviembre de 1991, pasada ante la fe del Lic. José María Idumate Acosta, Notario Público número 32 del Distrito de Torreón, Estado de Coahuila. (anexo 1)')
descripcionCapitulo12_format = di12.paragraph_format
descripcionCapitulo12_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

descripcionCapitulo12.font.name = 'Arial'       # Tipo de Fuente
descripcionCapitulo12.font.size = Pt(12)        # Tamaño de la Fuente
di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

di12.add_run().add_break()

########################################################################################################################################################################
# Capitulo 1.2.1
capitulo121 = doc.add_paragraph()
i121 = capitulo121.add_run("1.2.1.- Representante legal")
i121_format = capitulo121.paragraph_format
i121_format.line_spacing = 1.5

i121.font.name = 'Arial'
i121.font.size = Pt(12)
i121.font.bold = True

# capitulo121.add_run().add_break()  # Salto de línea

# Texto del Capitulo 1.2.1
legalRepresentante = doc.add_paragraph()

legalRepresentant = legalRepresentante.add_run('NOMBRE DEL REPRESENTANTE LEGAL')
legalRepresentant_format = legalRepresentante.paragraph_format
legalRepresentant_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

legalRepresentant.font.name = 'Arial'       # Tipo de Fuente
legalRepresentant.font.size = Pt(12)        # Tamaño de la Fuente
legalRepresentant.font.bold = True
# naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro


di121 = doc.add_paragraph()

descripcionCapitulo121 = di121.add_run('Gerente General y con poder según la persona moral legalmente constituida según consta en el instrumento número 552 de fecha 16 de noviembre de 1991, pasada ante la fe del Lic. José María Idumate Acosta, Notario Público número 32 del Distrito de Torreón, Estado de Coahuila. (anexo 1)')
descripcionCapitulo121_format = di121.paragraph_format
descripcionCapitulo121_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

descripcionCapitulo121.font.name = 'Arial'       # Tipo de Fuente
descripcionCapitulo121.font.size = Pt(12)        # Tamaño de la Fuente
di121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

di121.add_run().add_break()

########################################################################################################################################################################
# Capitulo 1.3
capitulo13 = doc.add_paragraph()
i13 = capitulo13.add_run(temasCapitulo1[2])
i13_format = capitulo13.paragraph_format
i13_format.line_spacing = 1.5

i13.font.name = 'Arial'
i13.font.size = Pt(12)
i13.font.bold = True

# capitulo13.add_run().add_break()  # Salto de línea

# Texto del Capitulo 1.3
domEm = doc.add_paragraph()

domicilioEmpresa = domEm.add_run('Domicilio Fiscal')
domicilioEmpresa_format = domEm.paragraph_format
domicilioEmpresa_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

domicilioEmpresa.font.name = 'Arial'       # Tipo de Fuente
domicilioEmpresa.font.size = Pt(12)        # Tamaño de la Fuente
domicilioEmpresa.font.bold = True
# naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro


di13 = doc.add_paragraph()

descripcionCapitulo13 = di13.add_run('Calle Luis F García Jimenéz No. 227\n'
                                     'Colonia: Ciudad Industrial, torreón\n'
                                     'Estado: Coahuila\n'
                                     'C.P.: 27050\n')
descripcionCapitulo13_format = di13.paragraph_format
descripcionCapitulo13_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

descripcionCapitulo13.font.name = 'Arial'       # Tipo de Fuente
descripcionCapitulo13.font.size = Pt(12)        # Tamaño de la Fuente
# di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

di13.add_run().add_break()

########################################################################################################################################################################
# Capitulo 1.3.1
capitulo131 = doc.add_paragraph()
i131 = capitulo13.add_run('1.3.1.- Domicilio para oir y recibir notificaciones')
i131_format = capitulo13.paragraph_format
i131_format.line_spacing = 1.5

i131.font.name = 'Arial'
i131.font.size = Pt(12)
i131.font.bold = True

# capitulo13.add_run().add_break()  # Salto de línea

# Texto del Capitulo 1.3
di131 = doc.add_paragraph()

descripcionCapitulo131 = di131.add_run('Calle: Novena No. 235\n'
                                     'Colonia: Brisas Poniente\n'
                                     'Municipio: Saltillo\n'
                                     'Estado: Coahuila de Zaragoza\n'
                                     'Teléfono: 8444278857\n'
                                     'Correo: mingopaco@yahoo.com\n')
descripcionCapitulo131_format = di131.paragraph_format
descripcionCapitulo131_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

descripcionCapitulo131.font.name = 'Arial'       # Tipo de Fuente
descripcionCapitulo131.font.size = Pt(12)        # Tamaño de la Fuente
# di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

di13.add_run().add_break()

########################################################################################################################################################################
# Guardar Documento
doc.save("CAPITULO 1 DTU EXTRACCION DE MATERIAL PETRO.docx")
