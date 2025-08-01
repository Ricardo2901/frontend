"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo8():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 8
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo VII.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 8
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 8
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 8 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'VIII.-	Estimación del volumen en metros cúbicos por especie y por predio, de las materias primas forestales derivadas del cambio de uso de suelo.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para determinar el volumen forestal que se afectará con la ejecución del proyecto, se requiere conocer la población (tipos de vegetación forestal y especies vegetales presentes) para especificar los parámetros que han de estimarse, por ello, considerando lo estipulado en la legislación forestal, para este aspecto únicamente se obtuvo información de los diferentes estratos arbóreos, arbustivos, gramíneos, herbáceos y suculentos, con lo cual se obtuvo información de las condiciones que presenta el área en estudio durante el ________________________________________________________. Para llevar a cabo lo anterior es necesario realizar una planeación antes de llegar a la fase de campo (inventario forestal), posteriormente con la información obtenida se lleva a cabo el procesamiento de la información en gabinete para obtener finalmente datos o información de la población estudiada.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 8.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 8.1 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.1.- Muestreo')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.1 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para determinar la homogeneidad de las especies presentes en el área de cambio de uso de suelo, se realizó un inventario con un esfuerzo de muestreo del _______________________________________ para arbustivas y suculentas, levantando sitios de muestreo circulares,  y __________ para los estratos gramíneo y herbáceo de forma cuadrangular _________, en el centro del sitios circular, ________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 8.1.A
    ########################################################################################################################################################################
    """
        #########################
        ### A). Diseño e intensidad de muestreo utilizado. ###
        #########################
    """
    #########################
    ### Titulo del capitulo 8.1.A ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\na).- Diseño e intensidad de muestreo utilizado.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.1.A ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para el diseño de acuerdo al tipo de vegetación y la superficie del área de cambio de uso de suelo se realizó un muestreo sistemático, con una separación de _____________________________________________.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capitulo 8.1.A ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo8_parrafo = doc.add_paragraph()
    imagenCapitulo8_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo8_run = imagenCapitulo8_parrafo.add_run('')
    imagen_cap_8 = imagenCapitulo8_run.add_picture('capitulo8/capitulo81A/cap_81A.png', width=Cm(5.29), height=Cm(5.27))

    # Opcional: espacio después del párrafo
    imagenCapitulo8_parrafo.space_after = Pt(1)


    ########################################################################################################################################################################
    # Capitulo 8.1.B
    ########################################################################################################################################################################
    """
        #########################
        ### B) Variables dasométricas (Diámetro normal, altura, total etc.) ###
        #########################
    """
    #########################
    ### Titulo del capitulo 8.1.B ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nb).- Variables dasométricas (Diámetro normal, altura, total etc.)')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.1.B ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('De acuerdo a la composición estructural de las especies observadas en el área de estudio, en el caso de las arbóreas, se midió el diámetro normal del fuste a la altura del pecho (1.30 m), así también la medición de la cobertura (diámetro de copa) y su altura total. El estrato arbustivo, así como de las suculentas, herbáceas y gramíneas se hizo medición de la altura, cobertura y número de individuos.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 8.1.C
    ########################################################################################################################################################################
    """
        #########################
        ### C) Modelo utilizado para la estimación del volumen (m3) ###
        #########################
    """

    #########################
    ### Titulo del capitulo 8.1.C ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nc).- Modelo utilizado para la estimación del volumen (m3)')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.1.C ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para la obtención del volumen para el caso de las arbóreas se utilizó el siguiente modelo Considerando únicamente la rama principal, ya que no tiene un fuste definido:')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 8.1.C ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo8_parrafo = doc.add_paragraph()
    formulaCapitulo8_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo8_run = formulaCapitulo8_parrafo.add_run('')
    formula_cap_8 = formulaCapitulo8_run.add_picture('capitulo8/capitulo81C/formula_1.png', width=Cm(5.29), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo8_parrafo.space_after = Pt(1)
    
    #########################
    ### Descripcion del capitulo 8.1.C ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('\nDónde:'
                                       '\nV = Volumen de la rama'
                                       '\nd = diámetro medio de la rama'
                                       '\nl = Largo de la rama.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)

    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para el caso de las especies arbustivas, herbáceas, gramíneas y suculentas, se utilizó la fórmula para la obtención del volumen de un cono recto de revolución, cuya fórmula es la siguiente:')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formula del capitulo 8.1.C ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo8_parrafo = doc.add_paragraph()
    formulaCapitulo8_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo8_run = formulaCapitulo8_parrafo.add_run('')
    formula_cap_8 = formulaCapitulo8_run.add_picture('capitulo8/capitulo81C/formula_2.png', width=Cm(3.78), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo8_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 8.1.C ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('\nDónde:'
                                       '\nVer = Volumen de la planta'
                                       '\nSo = Área de la sección (π*r2) '
                                       '\nHo = Altura de la planta')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 8.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.2 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.2- Especies Nombre Común y Científico (_____)')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Especies presentes en el área de cambio de uso de suelo en la ___.')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2 ###
    #########################
    tabla8b = doc.add_table(rows=30, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(30):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.2.A
    ########################################################################################################################################################################
    """
        #########################
        ### A) Coordenadas UTM del punto central y de las esquinas de cada uno de los sitios de muestreo siendo estos de forma circular. En este punto se debe incluir un mapa donde se puedan visualizar y ubicar los sitios de muestreo, mismo que tendrá que ser representativos del o los sujetos a cambio de uso de suelo en terrenos forestales. ###
        #########################
    """

    #########################
    ### Titulo del capitulo 8.2.A ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\na).- Coordenadas UTM del punto central y de las esquinas de cada uno de los sitios de muestreo siendo estos de forma circular. En este punto se debe incluir un mapa donde se puedan visualizar y ubicar los sitios de muestreo, mismo que tendrá que ser representativos del o los sujetos a cambio de uso de suelo en terrenos forestales.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.2.A ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Sitios de muestreo dentro del ACUSTF, en coordenadas UTM zona 14 N y Coordenadas geográficas (Ver anexo Mapa 8.1.- Muestreo de vegetación) en la _______________________________')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2.A ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Especies presentes en el área de cambio de uso de suelo en la ___.')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2.A ###
    #########################
    tabla8b = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.2.1 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.2.1.- Número de Individuos por especie que se espera remover en la __________________.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2.1 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Individuos por especie a remover en la __________________')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2.1 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.2.2 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.2.2.- Estimación de existencias volumétricas en la _______________________')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.2.2 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Para la obtención de volúmenes en especies arbóreas en el ecosistemas árido y semiárido las cuales no presenta un fuste definido, para obtener el volumen del estrato de las arbóreas se midió la rama principal de cada especie tomándolo como un único individuo y con su altura promedio para realizar individualmente su volumen utilizando la fórmula de Smalian, para el caso de las arbustivas, herbáceas, suculentas y gramíneas es una tarea difícil ya que no presentan una forma regular como un fuste o forma cilíndrica, por lo que se tomó en consideración realizar el cálculo de este volumen mediante la fórmula de cono recto de revolución, tomando el área de la base de sección transversal de la base (área de la cobertura) y su altura para estimar el volumen.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('De acuerdo al inventario del área sujeta de estudio, los volúmenes obtenidos se muestran en la tabla siguiente:')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2.2 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Estimación de existencias volumétricas en la ________________')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2.2 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 8.2.2 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('\nPara realizar el cálculo de volúmenes se utilizó la cobertura, la altura y el número de individuos (estratos arbustivos, herbáceos, gramíneos y suculentas) y para arbóreos se utilizó el diámetro a la altura 1.30 (DAP) del fuste principal y la altura.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 8.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.2.3 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.2.3.- Resumen de las existencias volumétricas en ________________________________')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2.3 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Resumen de las existencias Volumétricas')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2.3 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.2.4 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.2.4.- Uso principal de las especies ubicadas en los sitios de muestreo.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.2.4 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Resumen de las existencias Volumétricas')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.2.4 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.3 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.3- Especies Nombre Común y Científico (_____)')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Especies presentes en el área de cambio de uso de suelo en la __________________')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.3.A
    ########################################################################################################################################################################
    """
        #########################
        A) Coordenadas UTM del punto central y de las esquinas de cada uno de los sitios de muestreo siendo estos de forma circular. En este punto se debe incluir un mapa donde se puedan visualizar y ubicar los sitios de muestreo, mismo que tendrá que ser representativos del o los sujetos a cambio de uso de suelo en terrenos forestales.
        #########################
    """
    
    #########################
    ### Titulo del capitulo 8.3.A ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\na). - Coordenadas UTM del punto central y de las esquinas de cada uno de los sitios de muestreo siendo estos de forma circular. En este punto se debe incluir un mapa donde se puedan visualizar y ubicar los sitios de muestreo, mismo que tendrá que ser representativos del o los sujetos a cambio de uso de suelo en terrenos forestales.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.3.A ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Sitios de muestreo dentro del ACUSTF, en coordenadas UTM zona 14 N y Coordenadas geográficas (Ver anexo Mapa 8.1.- Muestreo de vegetación) en la ________________________.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3.A ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Coordenadas de los sitios de muestreo en el área de cambio de uso de suelo en la ____________________')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3.A ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.3.1 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.3.1.- Número de Individuos por especie que se espera remover en la ___________________')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3.1 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Individuos por especie a remover en la ______________.')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3.1 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.3.2 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.3.2.- Estimación de existencias volumétricas en la _______________________')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 8.3.2 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('Par la obtención de volúmenes en especies arbóreas en el ecosistemas árido y semiárido las cuales no presenta un fuste definido, para obtener el volumen del estrato de las arbóreas se midió la rama principal de cada especie tomándolo como un único individuo y con su altura promedio para realizar individualmente su volumen utilizando la fórmula de Smalian, para el caso de las arbustivas, herbáceas, suculentas y gramíneas es una tarea difícil ya que no presentan una forma regular como un fuste o forma cilíndrica, por lo que se tomó en consideración realizar el cálculo de este volumen mediante la fórmula de cono recto de revolución, tomando el área de la base de sección transversal de la base (área de la cobertura) y su altura para estimar el volumen.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('De acuerdo al inventario del área sujeta de estudio, los volúmenes obtenidos se muestran en la tabla siguiente:')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3.2 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.10.- Estimación de existencias volumétricas en la ________________')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3.2 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 8.3.2 ###
    #########################
    di8 = doc.add_paragraph()
    descripcionCapitulo8 = di8.add_run('\nPara realizar el cálculo de volúmenes se utilizó la cobertura, la altura y el número de individuos (estratos arbustivos, herbáceos, gramíneos y suculentas) y para arbóreos se utilizó el diámetro a la altura 1.30 (DAP) del fuste principal y la altura.')
    descripcionCapitulo8_format = di8.paragraph_format
    descripcionCapitulo8_format.line_spacing = 1.15
    descripcionCapitulo8_format.space_after = 0
    descripcionCapitulo8_format.space_before = 0

    descripcionCapitulo8.font.name = 'Arial'
    descripcionCapitulo8.font.size = Pt(12)
    di8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 8.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.3.3 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.3.3.- Resumen de las existencias volumétricas en ___________________')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3.3 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.X.- Resumen de las existencias Volumétricas')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3.3 ###
    #########################
    tabla8b = doc.add_table(rows=7, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)
            t8b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 8.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 8.3.4 ###
    #########################
    capitulo8 = doc.add_paragraph()
    i8 = capitulo8.add_run(f'\nVIII.3.4.- Uso principal de las especies ubicadas en los sitios de muestreo.')
    i8_format = capitulo8.paragraph_format
    i8_format.line_spacing = 1.15

    i8.font.name = 'Arial'
    i8.font.size = Pt(12)
    i8.font.bold = True
    capitulo8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 8.3.4 ###
    #########################
    tituloTabla8b = doc.add_paragraph()
    dti8b = tituloTabla8b.add_run('\nTabla 8.x.- Uso principal de las especies encontradas en el ACUSTF')
    dti8b_format = tituloTabla8b.paragraph_format
    dti8b_format.line_spacing = 1.15
    dti8b_format.space_after = 0

    dti8b.font.name = 'Bookman Old Style'
    dti8b.font.size = Pt(12)
    tituloTabla8b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 8.3.4 ###
    #########################
    tabla8b = doc.add_table(rows=20, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla8b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla8b.cell(rows, cols)
            t8b = cell.paragraphs[0].add_run(' ')
            t8b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 8 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo8() # Crear el documento