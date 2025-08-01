"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes, celdas, etc; en pulgadas y en centrimetros
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo13():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 13
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XIII.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 13
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 13
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 13 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'XIII.- Servicios Ambientales que serán afectados por el cambio de uso de suelo propuesto.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.1. Diagnóstico Ambiental.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.1 ###
    #########################
    # Párrafo 1
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        'El concepto “servicio ambiental "se utiliza para designar a cada una de las utilidades que la naturaleza proporciona a la población, desde un punto de vista económico. Este término fue introducido por Robert Constanza y sus colaboradores en trabajos científicos orientados a valorar el medio natural en un lenguaje compatible con el de la economía estándar, que rechaza hablar de valor si no es en términos estrictamente monetarios relativos a transacciones.'
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 2
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        'Los servicios ambientales son directamente dependientes del funcionamiento "saludable" de los ecosistemas y de la biodiversidad que éstos contienen. Cuando los ecosistemas se degradan también lo hacen los servicios que prestan, por ello los beneficios o servicios ambientales, no pueden ser apreciables sin una visión integrada de la naturaleza.'
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 3
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        'Este apartado tiene por interés encontrar los esquemas que permitan estimar y asignar valores objetivos a los bienes y servicios que se pudieran afectar, de tal manera que paralelamente y posteriormente a la ejecución del proyecto se puedan establecer las acciones de conservación y/o restauraciones adecuadas.'
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 4
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        'Los servicios ambientales pueden ser muy concretos tales como la protección del suelo o la conservación de la biodiversidad. En otros casos los servicios ambientales pueden ser algo más abstractos o referirse a un ámbito global como la captura del carbono o belleza escénica, por ejemplo. En general los servicios ambientales son todos aquellos servicios que brindan los ecosistemas y se dividen en cuatro tipos acorde con el Milenium Ecosistema Assessment (MEA 2005):'
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.1.- Servicios ambientales')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.1 ###
    #########################
    tabla13b = doc.add_table(rows=7, cols=4, style='Table Grid')

    # Encabezados
    encabezados = [
        "Servicios de soporte",
        "Servicios de provisión",
        "Servicios de regulación",
        "Servicios culturales"
    ]

    # Filas de contenido
    contenido_filas = [
        ["Biodiversidad", "Alimento", "Regulación de clima", "Belleza escénica"],
        ["Ciclo de nutrientes", "Materias primas", "Captura de CO2", "Recreación"],
        ["Formación de suelo", "Recursos genéticos", "Control de la erosión", "Información cultural y artística"],
        ["Producción primaria", "Recursos mediciones", "Regulación hídrica", "Información histórica"],
        ["Ciclo hidrológico", "Recursos ornamentales", "Provisión de agua", "Ciencia y educación"]
    ]

    # Pintar encabezados con fondo azul
    for cols in range(4):
        cell = tabla13b.cell(1, cols)
        cell_background_color(cell, "#45B1FF")
        t13b = cell.paragraphs[0].add_run(encabezados[cols])
        t13b.font.size = Pt(12)
        t13b.font.name = 'Arial'
        t13b.font.bold = True


    # Agregar contenido de las filas
    for row_idx, fila in enumerate(contenido_filas, start=2):
        for col_idx, texto in enumerate(fila):
            cell = tabla13b.cell(row_idx, col_idx)
            t13b = cell.paragraphs[0].add_run(texto)
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    # ✅ Celdas fusionadas
    row = tabla13b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[3]))

    # Agregar texto a la celda fusionada
    t13b = merged_cell.paragraphs[0].add_run('SERVICIOS AMBIENTALES')
    t13b.font.name = 'Arial'
    t13b.font.size = Pt(12)
    t13b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0095FF")  # Cambiar el color de fondo de la celda fusionada

    ########################################################################################################################################################################
    # Capitulo 13.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.1.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.1.1.-Servicios de soporte: ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.1.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Son aquellos que mantienen los procesos de los ecosistemas y permiten la provisión del resto de los servicios. Estos pueden o no tener implicaciones directas sobre el bienestar humano. Entre ellos se encuentra el mantenimiento de la biodiversidad, el ciclo hidrológico, el ciclo de nutrientes, y la producción primaria.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.1.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.1.2.- Los servicios de provisión: ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.1.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Son recursos tangibles y finitos que se contabilizan y consumen. Además, pueden ser o no renovables. Entre ellos se encuentra la provisión de agua para consumo humano, la provisión de productos como la madera y la producción de comida y medicinas.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.1.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.1.3.-Servicios de regulación: ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.1.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Son lo que mantienen los procesos y funciones naturales de los ecosistemas, a través de las cuales se regulan las condiciones del ambiente humano. Entre ellos encontramos la regulación del clima y gases de efecto invernadero, el control de la erosión o de las inundaciones.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.1.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.1.4.- Servicios culturales: ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.1.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Pueden ser tangibles e intangibles y son producto de percepciones individuales o colectivas; son dependientes del contexto socio-cultural. Intervienen en la forma en que interactuamos con nuestro entorno y con las demás personas. Entre ellos se encuentra la belleza escénica de los ecosistemas como fuente de inspiración y la capacidad recreativa que ofrece el entorno natural a las sociedades humanas.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Acorde a la Ley General de Desarrollo Forestal Sustentable publicado en el Diario Oficial de la Federación el ___________, en el artículo 7 Fracción LXI define los Servicios ambientales: como Beneficios que brindan los ecosistemas forestales de manera natural o por medio del manejo forestal sustentable, que pueden ser servicios de provisión, de regulación, de soporte o culturales, y que son necesarios para la supervivencia del sistema natural y biológico en su conjunto, y que proporcionan beneficios al ser humano;')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo a lo definido anteriormente, se identificó los servicios ambientales que se ponen en riesgo por el cambio de uso de suelo, siendo los más relevantes los que se enumeran a continuación:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista_servicios_riesgo = [
        "1. La provisión del agua en calidad y cantidad.",
        "2. La captura de carbono.",
        "3. El amortiguamiento del impacto de los fenómenos naturales.",
        "4. La modulación o regulación climática.",
        "5. La protección a la biodiversidad, de los ecosistemas y forma de vida.",
        "6. La protección y recuperación de suelo.",
        "7. El paisaje y la recreación."
    ]

    for servicio in lista_servicios_riesgo:
        di13 = doc.add_paragraph()
        descripcionCapitulo13 = di13.add_run(servicio)
        descripcionCapitulo13_format = di13.paragraph_format
        descripcionCapitulo13_format.line_spacing = 1.15
        descripcionCapitulo13.font.name = 'Arial'
        descripcionCapitulo13.font.size = Pt(12)
        di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.- Servicios ambientales que se presume afectar por el cambio de uso de suelo propuesto.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.1.- La provisión del agua en calidad y cantidad.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("El área de cambio de uso de suelo se ubica en la _______________________________ de acuerdo al SIATL (simulador de flujo) del INEGI. Dentro de esta área no existen corrientes fluviales tanto superficiales permanentes, __________________, debido a la pendiente a la permeabilidad en que se encuentra en el área y por el tipo de suelo que domina en el área es de __________, en una pendiente que va de __ a ___%.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("En México existen pocos trabajos sobre estimación de captura de agua en zonas arboladas. Dentro de los trabajos pioneros en esta área se encuentra el de Rivas et al. (1990) y todo el conjunto de modelos de escurrimiento a partir del modelo lluvia-escurrimiento desarrollado por el CENAPRED (Domínguez et al. 1994, citado por Torres, R. J. M. y Guevara, A. S. 2002).")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Para este apartado se eligió la metodología de la NOM-011-CNA-2000, la cual tiene como objetivo establecer el método base para determinar la disponibilidad media anual de las aguas nacionales superficiales y subterráneas, para su explotación, uso o aprovechamiento, y debido a que las especificaciones establecidas en la presente Norma son de observancia obligatoria para la Comisión Nacional del Agua y para los usuarios que realicen estudios para determinar la disponibilidad media anual de aguas nacionales.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Los resultados obtenidos son los siguientes")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("A continuación, se presentan el escurrimiento en el ACUSTF:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.2.- Volumen de Escurrimientos en el ACUSTF.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.1 ###
    #########################
    tabla13b = doc.add_table(rows=2, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("\nConsiderando la información antes señalada, se interrumpe un volumen de escurrimiento de agua de ________ mm en la superficie del ACUSTF de _____ hectáreas.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.3.-	Datos de infiltración en el ACUSTF.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.1 ###
    #########################
    tabla13b = doc.add_table(rows=2, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("\nDerivado del análisis se concluye que en la condición actual con la cobertura que posee, se tiene una infiltración normal de _____ mm anuales, con la implementación del proyecto al quedar sin vegetación esto aumenta la evapotranspiración por lo cual se dejará de captar agua reduciendo su infiltración a _____ mm.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("El costo que tiene la pérdida del servicio se tiene lo siguiente:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Por lo que de acuerdo a las Reglas de Operación emitidas por la Comisión Nacional Forestal el día _____________________________ para el otorgamiento del Programa de SA.1 Pago por Servicios Ambientales en su componente SA1.1. Servicios Ambientales Hidrológicos.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Modalidad SA.1.1. Servicios ambientales, cuyos apoyos son otorgados para conservar los ecosistemas, para mantener los ciclos hidrológicos; y otros beneficios relacionados con los procesos hidrológicos, tales como la recarga de acuíferos y evitar la erosión del suelo, teniendo el criterio de ejecución el siguiente: Son acciones destinadas a implementar una conservación activa de los ecosistemas naturales, con la finalidad de mantener los ciclos hidrológicos y procesos hídricos, tales como: la recarga de acuíferos y evitar la erosión del suelo; así como el uso sustentable de la biodiversidad (flora y fauna silvestre) en ecosistemas forestales para apoyar los compromisos internacionales en materia de biodiversidad suscritos por México, destacando el Convenio de la Diversidad Biológica (CDB); ambos con una visión de manejo de cuenca.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Para la región Norte-Centro según la clasificación de áreas elegibles, el polígono propuesto para el pago deberá tener una cobertura forestal arbórea igual o mayor a 50%, mientras que para la región Centro-Sur deberá contar con una cobertura forestal arbórea igual o mayor a 70%. Estarán exentas de cobertura mínima las áreas propuestas cuando el ecosistema predominante sea de zonas áridas, semiáridas, selva espinosa, vegetación hidrófila, pastizales naturales o sistemas agroforestales.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Se considera que el pago por el Servicio Ambiental es de $________ pesos por hectárea por año, si se toma en consideración los criterios antes mencionados el área de estudio no cumple con estos criterios, sin embargo, en el supuesto que se diera el pago de este servicio en las ______ hectáreas el ingreso anual sería de $ _______ pesos por año, sin embargo, el proyecto tendrá una duración de ____ años por lo que el montón de perdida por esta acción seria de ______________. Por lo que la pérdida del servicio ambiental hidrológica tendrá un costo de $ ___________ en un periodo de ___ años.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.2.- La Captura de Carbono.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Los bosques y selvas capturan, almacenan y liberan carbono como resultado de los procesos fotosintéticos, de respiración y de degradación de materia seca. El saldo es una captura neta positiva cuyo monto depende del manejo que se le dé a la cobertura vegetal, así como de la edad, distribución de tamaños, estructura y composición de ésta. Este servicio ambiental que proveen bosques o selvas como secuestradores de carbono (sumideros) permite equilibrar la concentración de este elemento, misma que se ve incrementada debido a las emisiones producto de la actividad humana (Torres y Guevara, 2002).")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("El concepto de captura de carbono normalmente integra la idea de conservar los inventarios de este elemento que se encuentran en suelos, bosques y otro tipo de vegetación y donde es inminente su desaparición, así como el aumento de los sumideros de carbono a través del establecimiento de plantaciones, sistemas agroforestales y la rehabilitación de bosques degradados (Tipper 2000), sólo por mencionar algunos ejemplos en los que la vegetación es usada como sumidero.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Para el cálculo del contenido de carbono en la superficie forestal donde se efectuará el Cambio de uso de suelo se procedió a la determinación de las existencias reales (m/ha) y existencias reales totales (m3 totales por tipo de vegetación). Se optó por utilizar el método de IPCC esto debido a la precisión de las estimaciones de biomasa es de crítica importancia, porque los modelos determinan la cantidad de carbono que llega a la atmósfera y son muy sensibles a estas estimaciones (Brown y Lugo, 1986) tal como se indica la fórmula y la tabla siguiente:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('CCC = Vr x Fd x FCC')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Times New Roman'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.italic = True
    di13.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Dónde:"
                                         '\nCCC = Coeficiente de Captura de Carbón'
                                         '\nVr = Volumen real en m\u00B3'
                                         '\nFd = Factor de Densidad'
                                         '\nFCC = Factor de Captura de Carbón')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 13.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.4.- Pasos para estimar el carbono almacenado en la superficie forestal.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.2 ###
    #########################
    tabla13b = doc.add_table(rows=7, cols=2, style='Table Grid')

    # Encabezados
    encabezados = [
        "Columna",
        "Concepto y/o descripción del proceso"
    ]

    # Filas de contenido
    contenido_filas = [
        ["1", "Clasificación de los individuos por género."],
        ["2", "Estimación de la superficie total (ha) ocupada por comunidad vegetal"],
        ["3", "Cálculo del volumen en metros cúbicos rollo por hectárea, para cada comunidad vegetal"],
        ["4", "Factor de densidad para coníferas 0,48 y 0,60 para latifoliadas (Toneladas de materia seca / m3)"],
        ["5", "Factor de contenido de carbono 0,45 (Toneladas de Carbono / toneladas de materia seca)"],
        ["6", "Cálculo de biomasa (tonelada de Carbono / ha), mediante la multiplicación de las Columnas 3, 4 y 5"]
    ]

    # Pintar encabezados con fondo azul
    for cols in range(2):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')
        t13b = cell.paragraphs[0].add_run(encabezados[cols])
        t13b.font.size = Pt(12)
        t13b.font.name = 'Arial'
        t13b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar contenido de las filas
    for row_idx, fila in enumerate(contenido_filas, start=1):
        for col_idx, texto in enumerate(fila):
            cell = tabla13b.cell(row_idx, col_idx)
            t13b = cell.paragraphs[0].add_run(texto)
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for rows in tabla13b.rows:
        rows.cells[0].width = Cm(2.20)
        rows.cells[1].width = Cm(16.00)

    #########################
    ### Descripcion del capitulo 13.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En virtud de no encontrarse información referente a la captura de carbono en vegetación diferente a áreas arboladas se tomó como base estudio del INE donde considera que la vegetación de zonas áridas tiene una captura de 0.3 T/ a y considerando la superficie sujeta a cambio de usos de suelo que es de ________ ha se dejaría de capturar lo siguiente:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.5.- Factor de Densidad y contenido de carbono.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.2 ###
    #########################
    columnas = 5
    filas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    cell = tabla13b.cell(0, 2)
    t13b = cell.paragraphs[0].add_run('Factor de contenido de carbono')
    t13b.font.size = Pt(12)
    t13b.font.name = 'Arial'
    t13b.bold = True

    cell = tabla13b.cell(0, 3)
    t13b = cell.paragraphs[0].add_run('Factor de Densidad')
    t13b.font.size = Pt(12)
    t13b.font.name = 'Arial'
    t13b.bold = True

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run('')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 13.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Dato referente a vegetación de zonas áridas (0.3 T/ha. Estudio INE).')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Tomando el total de m3 por comunidad vegetal a intervenir durante la realización del Cambio de Uso de Suelo, así como la superficie de éstas, se tabularon los datos dando como resultado lo presentado a continuación.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Volumen de especies para cálculo de fijación de carbono.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Debido a que se cuenta con poca información para estimar la captura de carbono por año el resultado es de _____ Toneladas, es la cantidad de carbono que se ha almacenado en la vegetación que se pretende remover con el cambio de uso de suelo, si consideramos que de las _____ ha son activas de acuerdo a su cobertura en la captura de carbono con un precio internacional de 10 dlls/Tonelada (PRISMA, 2002), se tiene un ingreso de ______ dólares que de acuerdo a la tasa de cambio es de de $ ______ pesos por dólar al ______ de Banamex, por lo que el pago del servicio ambiental de captura de carbono equivaldría a $ _______ pesos por año, por lo tanto deja de percibir por este concepto en los 10 año de vida útil del proyecto $ _________ pesos, considerando la vegetación de zonas áridas donde  de acuerdo al INE en estudios realizados contemplan que este tipo de vegetación almacena 0.3 T/ ha de carbono.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La pérdida del servicio ambiental de captura de carbono es de $ _________ pesos en el periodo de _____ años.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.3.- El amortiguamiento del impacto de los fenómenos naturales.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Este servicio ambiental considera las funciones que cumple el ecosistema cuando actúa como un regulador de las fluctuaciones y cambios ambientales que se origina en eventos tales como tormentas, inundaciones, sequías y huracanes entre otros, principalmente por la estructura de la vegetación. Es importante señalar que este servicio ambiental depende en gran medida del buen estado de conservación de la vegetación y de la extensión de la misma, ya que la disminución de estos factores disminuyen la capacidad de estos de absorber perturbaciones sin alterar significativamente sus características de estructura y funcionalidad del ecosistema ya que esta depende en forma directa de la riqueza de especies y complejidad de interacciones (es decir, un sistema en el cual sus integrantes tengan más diversidad y número de funciones ecológicas será capaz de soportar de mejor manera una perturbación especifica). En este caso, como fue descrito en los apartados anteriores la superficie por la que se está solicitando el cambio de uso de suelo en terrenos forestales se encuentra en una zona que con el tiempo sea _________________, la _________________________, por lo que se ve disminuida su capacidad de brindar el servicio de amortiguamiento e integridad del ecosistema ante la ocurrencia de los eventos antes mencionados. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.4.- La modulación o regulación climática.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Es evidente que al no existir una cubierta vegetal que refleje los rayos solares, por la realización del proyecto, estos son absorbidos en forma directa al suelo modificando las condiciones del microclima, registrando aumentos de evaporación debido a la radiación directa, así como cambios en los ciclos biogeoquímicos naturales a una escala local. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Sin embargo, considerando que el entorno será modificado en forma parcial donde se llevara a cabo el proyecto, es posible que las alteraciones locales queden marcadas como eventos aislados, en donde la capacidad de resistencia del sistema, junto con medidas de mitigación por el cambio de uso de suelo, disminuyan el efecto negativo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('A manera de conclusión, se puede decir que resulta evidente el cambio de patrones climáticos locales con el desarrollo del proyecto o cualquier otra alteración máxime si se considera que el área sujeta de estudio se ubica en la ____________, donde la radiación es más directa en la mayor parte del día, generando variaciones en la evaporación, radiación en el suelo, escurrimientos, desecación, así como aumento de la temperatura, entre otros. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Dichos cambios no podrán ser cruciales en la dinámica de la cuenca, considerando que la superficie del predio resulta poco significativa en comparación con esta.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.- La protección a la biodiversidad, de los ecosistemas y forma de vida.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.5 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La ejecución del proyecto, no afectará significativamente a la biodiversidad natural del sistema ambiental donde entra el área del proyecto esto debido a que en el área se encuentran  dentro del __________________________________________, donde el uso de suelo actual está ________ y no es óptima para el desarrollo de alguna actividad productiva, de igual forma el proyecto no se encuentra dentro de algún Área Natural Protegida registrada en el Sistema Nacional de Áreas Naturales Protegidas (SINAP), adicionalmente se tiene contemplado la implementación de las medidas de prevención y mitigación de impactos ambientales derivadas del mismo, con el retiro de la vegetación y la influencia del ser humano por considerar la construcción la _____________________, no permitirá mantener un hábitat favorable para la fauna en el transcurso de la ejecución del proyecto y operación del mismo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.1- Estimación económica de los recursos Forestales Maderables y no maderables.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.5.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para los fines de estimación de una aproximación del costo del material forestal presente en el área sujeta al cambio y uso de suelo fue necesario conocer las especies presentes en dicha área, así como también conocer el valor e importancia de cada especie para así realizar una estimación del monto económico del número de individuos totales de acuerdo al inventario proporcionado por la información del levantamiento de datos mediante sitios de muestreo establecidos en el área de estudio.  ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('A continuación, se presentan el valor aproximado de cada especie de acuerdo con el tipo de matorral presente en el área de cambio de uso de suelo. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.5.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valor económico de las especies Maderables de flora silvestre el área de estudio del ________________________')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.1 ###
    #########################
    columnas = 6
    filas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run('')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 13.2.5.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.-	Valor económico de las especies No Maderables de flora silvestre el área de estudio del ______________________')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.1 ###
    #########################
    columnas = 6
    filas = 40
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 13.2.5.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Respecto a la flora que se encuentra en el área se considera la distribución de la cobertura vegetal de ______ ha que corresponde a _____% con respecto al sistema ambiental con una superficie de ________ ha, con un total de individuos a remover de _______ de un total de ____________ individuos en el sistema ambiental, que representa la eliminación _______% de ___________________________________________, ya que es la vegetación que se encuentra en el área de cambio de uso de suelo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Por la pérdida del servicio ambiental de la flora silvestres dentro del área de cambio de uso de suelo será de un valor aproximadamente de $ ______________.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.5.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.-	Valor económico de las especies No Maderables de flora silvestre el área de estudio del ______________________')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.1 ###
    #########################
    columnas = 6
    filas = 40
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 13.2.5.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Respecto a la flora que se encuentra en el área se considera la distribución de la cobertura vegetal de ____ ha que corresponde a _____ % con respecto al sistema ambiental con una superficie de ________ ha, con un total de individuos a remover de ______ de un total de ___________ individuos en el sistema ambiental, que representa la eliminación _____ % de ________________________ ya que es la vegetación que se encuentra en el área de cambio de uso de suelo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Por la pérdida del servicio ambiental de la flora silvestres dentro del área de cambio de uso de suelo será de un valor aproximadamente de $ __________.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.2.- Afectación a la biodiversidad del recurso flora en el Acustf con respecto al sistema ambiental.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.2.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.2.1.- ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE MARGALEF
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Riqueza especifica Índices de Margalef. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo al cuadro y gráfico anterior se observa que, en cuanto a Riqueza de especies para los ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE MENHINICK
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Riqueza especifica Índice Menhinick. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En cuanto a Menhinick se observa en el cuadro y gráfico anterior un __________________________________________, y al realizar el desmonte no pone en riesgo el germoplasma en el ecosistema en que se presenta, ya que es el mismo tipo de vegetación que se desarrolla en su entorno.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE SIMPSON
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Dominancia de especies Índices de Simpson. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo al índice de dominancia de Simpson en el estrato ________________________________________________________________________, la vegetación que se va afectar por el cambio y uso de suelo no se pone en riesgo ya que se desarrolla dentro del Sistema ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE BERGER-PARKER
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Dominancia de especies Índices de Berger-Parker. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para el estrato Arbustivo presenta ________________________________________. Al presentar los valores similares en ambas áreas, la vegetación que se va afectar por el cambio y uso de suelo, no se pone en riesgo ya que se desarrolla dentro del Sistema ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE SHANNON
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Equidad de especies Índices de Shannon. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En el cuadro y gráfico anterior en cuanto a la equidad de Shannon presenta para el estrato ______________________________________________________________. La similitud de valores en ambas áreas se traduce en que todas las especies que se encuentran en el ACUSTF se desarrollan en el Sistema Ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE PIELOU
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Equidad de especies Índices de Pielou. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.1 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En el cuadro y grafico anterior en cuanto a la equidad del índice de Pielou en el estrato ________________________________________________________.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.2.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.2.2.- ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE MARGALEF
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Riqueza especifica Índices de Margalef. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo al cuadro y gráfico anterior se observa que, en cuanto a Riqueza de especies para los ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE MENHINICK
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Riqueza especifica Índice Menhinick. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En cuanto a Menhinick se observa en el cuadro y gráfico anterior un __________________________________________, y al realizar el desmonte no pone en riesgo el germoplasma en el ecosistema en que se presenta, ya que es el mismo tipo de vegetación que se desarrolla en su entorno.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE SIMPSON
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Dominancia de especies Índices de Simpson. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo al índice de dominancia de Simpson en el estrato ________________________________________________________________________, la vegetación que se va afectar por el cambio y uso de suelo no se pone en riesgo ya que se desarrolla dentro del Sistema ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE BERGER-PARKER
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Dominancia de especies Índices de Berger-Parker. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para el estrato Arbustivo presenta ________________________________________. Al presentar los valores similares en ambas áreas, la vegetación que se va afectar por el cambio y uso de suelo, no se pone en riesgo ya que se desarrolla dentro del Sistema ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE SHANNON
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Equidad de especies Índices de Shannon. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En el cuadro y gráfico anterior en cuanto a la equidad de Shannon presenta para el estrato ______________________________________________________________. La similitud de valores en ambas áreas se traduce en que todas las especies que se encuentran en el ACUSTF se desarrollan en el Sistema Ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            INDICE DE PIELOU
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Título de la tabla del capítulo 13.2.5.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Equidad de especies Índices de Pielou. ')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.2.2 ###
    #########################
    filas = 5
    columnas = 5
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(filas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(columnas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En el cuadro y grafico anterior en cuanto a la equidad del índice de Pielou en el estrato ________________________________________________________.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.3.- Estimación económica de los recursos faunístico.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.5.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La Fauna silvestre como componente de estos recursos naturales conforma un gran elemento de suma importancia para el desarrollo de la humanidad, tanto en épocas pasadas como en la actualidad se han utilizado un sin fin de especies en cuanto a usos de medicinales, vestido, calzado y/o para extraer algún tipo de fibra, combustibles o bien establecer algún tipo de comercio, necesarios para satisfacer necesidades del ser humano.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De tal manera podemos enfatizar que la población de la zona no tiene ningún interés por los recursos en cuanto a su valor ambiental, ecológico o económico.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Debemos de considerar el costo de las especies de acuerdo al valor de importancia los cuales puedes ser de carácter cinegético y no cinegético así mismo para el caso de las especies con interés especial como aquellas consideradas por la NOM-059-SEMARNAT-2010, durante los recorridos de campo y los sitios de muestreos realizados en el área de estudio del proyecto, no se localizaron especies con algún estatus.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.5.3 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valor económico de las especies de fauna silvestres en el área sujeta a cambio de uso de suelo forestal.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.3 ###
    #########################
    filas = 9
    columnas = 8
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('El servicio ambiental de la fauna silvestre tendrá una afectación económica de aproximadamente $_________ pesos de los individuos que se distribuyen dentro del área de cambio de uso de suelo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.5.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.5.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.5.4.- Afectación a la biodiversidad del recurso fauna en el Acustf con respecto al sistema ambiental.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.5.4 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Grado de afectación de las especies faunísticas')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.5.4 ###
    #########################
    filas = 9
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.5.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En cuanto a la biodiversidad que se presenta en las áreas tal y como se muestra en la tabla anterior, para determinar grado de afectación que se tendrá, los valores nos arrojan que para el grupo de las aves en el Acustf la afectación será ___________________________, considerando que se tendrá un porcentaje de desplazamiento de especies hacia el área del sistema ambiental, por lo cual estas especies podrán movilizarse con la implementación de ahuyentamiento hacia el área del sistema ambiental.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.6.- Perdida de erosión hídrica y eólica del Suelo.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.2.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Estimación del grado de erosión potencial del suelo.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Metodología para determinar el nivel de degradación potencial del suelo.")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "La degradación del suelo se define como “un grupo de procesos que ocasionan el deterioro del recurso, "
        "los cuales provocan una disminución de la productividad biológica y la pérdida de la biodiversidad”. En este sentido, "
        "el estado de degradación en que se encuentran los suelos de uso agropecuario y forestal, se estima por medio de las pérdidas "
        "de suelo que ocurren en los terrenos, de modo que sea posible determinar si el uso que se está dando a estos es el correcto. "
        "Cuando la tasa de erosión es mayor que la tasa de formación del suelo, es señal de que el manejo está originando su degradación "
        "y se hace necesario realizar prácticas y obras de conservación para de esa forma contribuir al desarrollo sostenible de los recursos "
        "naturales. Para estimar la erosión de los suelos se ha utilizado la Ecuación Universal de Pérdida de Suelo (EUPS), un modelo que permite "
        "estimar la erosión actual en campo y la potencial de dicho recurso. Esta ecuación constituye un instrumento de planeación para establecer "
        "las prácticas y obras de conservación para que hagan que la erosión actual sea menor que la tasa máxima permisible de erosión. La tasa máxima "
        "permisible de pérdidas de suelo es de 10 T/ha (toneladas por hectárea); siendo que mayores pérdidas significan degradación. "
        "(La metodología y resultados se encuentran en el capítulo V)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("En donde se obtuve lo siguiente:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.6.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.6.1.- Erosión hídrica obtenidos en el área de Cambio de Uso de Suelo.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.6.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.23.- Erosión hídrica en el ACUSTF.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.6.1 ###
    #########################
    filas = 2
    columnas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.6.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En las condiciones actuales por efecto de la lluvia se pueden tener pérdidas de ______ mm de suelo/año, con la implementación del proyecto al quedar desnudo el suelo incrementa una pérdida hasta ______ mm de suelo/año.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.2.6.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.2.6.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.2.6.2.- Erosión Eólica obtenidos en el área de Cambio de Uso de Suelo.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.2.6.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.23.- Erosión Eólica en el ACUSTF.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.2.6.2 ###
    #########################
    filas = 2
    columnas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.2.6.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En las condiciones actuales por efecto del viento se tiene una pérdida de suelo de ____ mm de suelo/año, con la implementación del proyecto al quedar desnudo el suelo se incrementa hasta _____ mm de suelo/año, por lo que se recomienda realizar actividades de compensación, logrando tener una erosión potencial solo de _____ mm de suelo/año.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.- El paisaje y la recreación.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3 ###
    #########################

    # Párrafo 1
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "El paisaje puede definirse como la percepción que se posee de un sistema ambiental. Es, por lo tanto, "
        "“el área en el que conviven los rasgos naturales, así como los influenciados por el hombre y que da lugar "
        "a una percepción visual y mental tanto individual como colectiva del conjunto de ese espacio”. "
        "(Abad Soria y García Quiroga 2006)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 2
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "La consideración del paisaje como elemento del medio ambiente implica dos aspectos fundamentales: "
        "el paisaje como elemento aglutinador de una serie de características del medio físico y la capacidad "
        "que tiene un paisaje para absorber, los usos y actuaciones que se desarrollan sobre él."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 3
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Uno de los mayores problemas en el desarrollo de métodos de evaluación cuantitativa de los efectos escénicos "
        "es el de la medición de las contribuciones específicas de los elementos del paisaje, casi todos los modelos "
        "coinciden en tres apartados: la visibilidad, la fragilidad del paisaje y la calidad paisajística. "
        "(Martí Vargas y Pérez González, 2001)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 4
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "La visibilidad o cuenca visual; es la porción de paisaje visualmente auto contenida, que abarca toda el área "
        "de visualización que un observador tiene del paisaje."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 5
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "La fragilidad de un paisaje; es la “susceptibilidad de un paisaje al cambio cuando se desarrolla un uso o actuación sobre él”."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 6
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Calidad paisajística o calidad visual de un paisaje; se entiende “el grado de excelencia de éste, su mérito para no ser "
        "alterado o destruido o de otra manera, su mérito para que su esencia y su estructura actual se conserve” (Blanco, 1979)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Párrafo 7
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "En paisajes naturales, las cuencas hidrográficas constituyen la forma más objetiva para conceptualizar la operatividad "
        "de un geo ecosistema. Esto es así porque forma un sistema discreto, con umbrales bien definidos de entrada y salida de "
        "materia y energía, en el que el agua es el principal elemento funcional (Manzo y López, 1997)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.- Métodos para la evaluación del paisaje:')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.1.- Métodos independientes de los usuarios del paisaje en los que la valoración la realizan los expertos.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Se consideran de “subjetividad aceptada o controlada”, ya que los evaluadores pueden mantener un criterio uniforme. "
        "Se distinguen dos grandes grupos:"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.2.- Métodos directos de valoración de la calidad visual:')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Estos métodos se caracterizan porque la evaluación se realiza por medio de la contemplación del paisaje, "
        "en forma directa o por medios visuales. El paisaje se valora subjetivamente, con calificativos, escalas "
        "de rango o de orden (Fines, 1978)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.3.- Métodos indirectos de valoración de la calidad:')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Estos métodos cualitativos y cuantitativos que evalúan el paisaje analizando y describiendo sus componentes "
        "o a través de categorías estéticas."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo párrafo del mismo capítulo
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Los primeros utilizan la desagregación de las características físicas del paisaje, tales como, topografía, uso del suelo, "
        "agua, etc., a las que se le asigna un valor parcial, el que luego es “sumado” a los demás valores parciales obteniéndose "
        "un valor final de la calidad (Fernández Cañadas, 1977; Gómez Orea, 1979; Ramos, 1979; Wright, 1974)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.4.- Métodos dependientes de los usuarios del paisaje o evaluación observación.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Con estos métodos se pretende obtener una opinión “democrática” de calidad de un área, es decir una opinión representativa. "
        "La esencia de este enfoque es la preferencia de la sentencia del paisaje en su totalidad, por oposición a las técnicas de medición, "
        "que se basan en la definición de los factores para explicar la variación en la calidad del paisaje (Dunn, 1976)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.4.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.4.1.- Modelos psicofísicos:')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.4.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Aquellos que atienden en la valoración del paisaje a las relaciones entre aspectos físicos y los juicios "
        "o respuestas de la percepción de estos estímulos."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.4.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.4.2.- Modelo psicológico relacionado con la teoría de la personalidad.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.4.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "El paisaje es valorado en términos cognitivos de complejidad, legibilidad, misterio, profundidad. "
        "Un paisaje de gran calidad evoca sentimientos positivos, como la seguridad, la relajación, calidez, la alegría o la felicidad, "
        "una baja calidad del paisaje se asocia con el estrés, el miedo, la inseguridad, la dificultad, la oscuridad, u otros sentimientos "
        "negativos (Daniel y Vining 1983)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.1.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.1.4.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.1.4.3.- Método fenomenológico el cual enfatiza en la interpretación del ambiente')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.1.4.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Este modelo representa el extremo de la determinación subjetiva de las características del paisaje. "
        "(Lowenthal, 1972; Lynch, 1960; Burton y Kates, 1974; Seamon, 1979) Una última mención merece la apreciación a partir de la estética "
        "ecológica, dónde el placer es secundario y se deriva de conocer el paisaje y su ajuste ecológico (Gobster 1996)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo párrafo del mismo subcapítulo
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Para la evaluación del paisaje en el área sujeta de estudio se efectuó de acuerdo al siguiente método:"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.- Métodos indirectos de valoración de la calidad:')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Este método que considera aspectos cualitativos y cuantitativos en cada una de las etapas del proyecto que evalúan el paisaje "
        "analizando y describiendo sus componentes o a través de categorías estéticas utilizando para tal fin la evaluación de sus cualidades "
        "(fragilidad y calidad paisajística), armonía (calidad visual), aspectos identificados y valorados en las matrices correspondientes "
        "y la afectación y/o modificación están ligados a las condiciones de una cuenca en la que destacan las características físicas del paisaje, "
        "tales como, topografía, uso del suelo, agua, flora y fauna etc."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "En este apartado se califica la valoración de cada área, siendo el área de cambio de uso de suelo y el sistema ambiental como se muestra a continuación."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Criterios ecológicos a valorar son los siguientes:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 13.3.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.1.-Vegetación (Calidad del paisaje)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Con referencia a la vegetación y uso de suelo; se les asigna mayor calidad a unidades de paisaje con mezcla equilibrada de cultivos, "
        "masas arboladas y vegetación nativa con influencia a 500 m., Mediana calidad cuando presenta mezcla de dos estratos y cercanía a cultivos "
        "o cuerpos de agua a menos de 500 m, baja calidad cuando se presenta un solo matorral dominante y aislado, menor calidad cuando son cultivos "
        "aislados, pastizales y áreas sin vegetación."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.1 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valoración del paisaje factor Vegetación.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.1 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Con los datos anteriores se puede apreciar en la valoración de la calidad del paisaje se puede concluir que el sistema ambiental tiene la capacidad de absorción del proyecto ya que se puede apreciar que existe tan solo la afectación del ____ % la cual corresponde a _______ considerada de ___________ por lo que el impacto es caracterizado como medio. (Ver anexo Mapa 13-1.- Valoración de la vegetación). "
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.2.- Paisaje Agua.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Los cuerpos de agua son altamente ponderados, ya sean estos naturales (lagos y ríos) o artificiales (fuentes y canales).')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.2 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valoración del paisaje factor Agua.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.2 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Como se puede apreciar en la valoración del factor AGUA en el paisaje se puede concluir que el sistema ambiental tiene la capacidad de absorción del proyecto ya que no se modificará ningún cuerpo de agua teniendo el área de cambio de uso de suelo la _______ y el sistema ambiental posee una superficie de _________ ha aproximadamente en esta clase. Por lo que el impacto es considerado bajo de tan solo el ______ %, (Ver anexo Mapa 13-2.- Valoración del Factor agua)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.3.- Suelo y Cubierta vegetal.  (Fragilidad del paisaje)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La fragilidad de la vegetación la definimos como el inverso de la capacidad de ésta para ocultar una actividad que se realice en el territorio.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.3 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valoración del paisaje factor Suelo y Cubierta Vegetal')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.3 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Como se puede apreciar en la valoración del factor Fragilidad del paisaje se puede concluir que a pesar de que el área se encuentra en la _______________ % que es considerada fragilidad _____ el sistema ambiental tiene la capacidad de absorción del proyecto ya que este cuenta con una superficie de __________ ha de esta clase. (Ver anexo Mapa 13-3.- Valoración de Fragilidad3"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.4.- Pendiente. ')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Se considera que a mayor pendiente mayor fragilidad, por producirse una mayor exposición de las acciones. Se ha calculado la pendiente en cada punto del área del proyecto y su influencia en el sistema ambiental en donde, se han establecido tres categorías.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.4 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.X.- Valoración del paisaje factor Pendiente')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.4 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Como se puede apreciar en la valoración del factor PENDIENTE en el paisaje se puede concluir que el sistema ambiental tiene la capacidad de absorción del proyecto ya que se puede apreciar que existe tan solo la afectación del _____ %, ya que se encuentra el área de cambio de uso de suelo en su __________________ considerada de mayor fragilidad a la pendiente, Por lo que el impacto es caracterizado como de impacto bajo al ser área pequeña en consideración con el sistema ambiental. (Ver anexo Mapa 13-4.- Valoración pendiente)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.5 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.5.- Orientación.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.5 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Las laderas soleadas presentan mayor fragilidad por su exposición que las umbrías.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.5 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.X.- Valoración del paisaje factor Orientación')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.5 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.5 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Dentro del sistema ambiental la fragilidad a la que está expuesta se encuentra en su mayoría en baja, de acuerdo a su exposición, el área en estudio de encuentra dentro de la exposición soleado y umbrío, por lo que representa una afectación de ____ %, pero su mayor proporción se encuentra en la ____________ con respecto a la totalidad que existen en el sistema ambiental por lo que se considera ____. (Ver anexo Mapa 13.5.- Valoración Orientación)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.6 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.6.- Tamaño de la cuenca visual. (Visibilidad el paisaje)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Se considera que a mayor extensión de la cuenca visual mayor fragilidad. Para este apartado se tomó en consideración aquellos puntos en que se puede visualizar el proyecto desde cualquier área del sistema ambiental, obteniendo los siguientes valores.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.6 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valoración del paisaje factor cuenca visual')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.6 ###
    #########################
    filas = 7
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Para la fragilidad en el área en estudio es considerada ____ ya que representa el ____ % de afectación con respecto al sistema ambiental, en contraendose en su mayoría en fragilidad ____, sin embargo, se considera muy nula a la afectación de la cuenca visual. (Ver anexo Mapa 13.6.- Cuenca Visual)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.3.2.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.3.2.7 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.3.2.7.- Accesibilidad')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.3.2.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Cuanto mayor es la accesibilidad, mayor es la fragilidad.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.7 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Valoración del paisaje factor Accesibilidad')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.7 ###
    #########################
    tabla13b = doc.add_table(rows=4, cols=4, style='Table Grid')

    for rows in tabla13b.rows:
        rows.cells[0].width = Cm(3.69)
        rows.cells[1].width = Cm(1.75)
        rows.cells[2].width = Cm(7.75)
        rows.cells[3].width = Cm(3.5)

    # Encabezados
    encabezados = [
        "Fragilidad",
        "Clase",
        "Accesibilidad",
        "Valor Asignado"
    ]

    # Filas de contenido
    contenido_filas = [
        ["Menor (baja)", "1", "Sin Acceso", "1"],
        ["Media", "2", "Caminos Vecinales o rutas asfaltadas", "3"],
        ["Mayor (alta)", "3", "Casco Urbano o rutas", "5"]
    ]

    # Pintar encabezados con fondo azul
    for col in range(4):
        cell = tabla13b.cell(0, col)
        cell_background_color(cell, '0070C0')  # asumiendo que tienes esta función definida
        t13b = cell.paragraphs[0].add_run(encabezados[col])
        t13b.font.name = 'Arial'
        t13b.font.size = Pt(12)
        t13b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar contenido de las filas
    for row_idx, fila in enumerate(contenido_filas, start=1):
        for col_idx, texto in enumerate(fila):
            cell = tabla13b.cell(row_idx, col_idx)
            t13b = cell.paragraphs[0].add_run(texto)
            t13b.font.name = 'Arial'
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.3.2.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En cuanto a la fragilidad por accesibilidad se considera _____________________________ a área de estudio y se conecta con carretera asfáltica, estos caminos ya son existentes cercanos al área.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En resumen, se tiene los resultados obtenidos del análisis en cuanto a la fragilidad del paisaje lo siguiente:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.7 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Resultados de la Valoración del paisaje general')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.7 ###
    #########################
    filas = 7
    columnas = 3
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 13.3.2.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Como se puede apreciar el área de cambio de uso de suelo presenta un impacto a la fragilidad en su mayoría catalogada como ____ a media y el sistema ambiental tiene la capacidad de absorción de estos impactos por la implementación del proyecto siempre y cuando se realicen las medidas de mitigación propuestas. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En resumen, la perdida por los servicios ambientales será lo siguientes: ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.3.2.7 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Resumen de los costos de afectación a los servicios ambientales')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.3.2.7 ###
    #########################
    filas = 7
    columnas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 13.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.4.- Estimación económica por la ejecución del proyecto.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Beneficios económicos que traería consigo el proyecto a la sociedad por su puesta en marcha (operación del proyecto proyectada a largo plazo o su vida útil).')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('El proyecto tendrá un periodo de ejecución de aproximadamente de ___ años, ____ meses para el cambio de uso de suelo en las estepas de preparación, construcción y operación. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('El nuevo uso que se pretende dar al área en estudio, se considera es la ____________, mismo que llevará a cabo un proceso de preparación, para la construcción, operación y abandono, dentro del cual se contempla la remoción total de la vegetación (Despalme), de _____ ha., que presenta vegetación forestal, que contempla el proyecto en un periodo de ____ años, así mismo en el cada año se utilizara para la remoción de la vegetación, cada año será para la construcción y operación, el costo económico que tendrá el proyecto para su inversión inicial es de aproximadamente de más de ____________ por lo que traerá beneficios a la región de ___________ en un periodo de __ años, de acuerdo a lo siguiente:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo13_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.4 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Costos de Inversión inicial del proyecto.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.4 ###
    #########################
    filas = 7
    columnas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    tabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    #########################
    ### Título de la tabla del capítulo 13.4 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Costos de Inversión total del proyecto.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.4 ###
    #########################
    filas = 6
    columnas = 8
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    tabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 13.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para el cálculo del análisis financiera para obtener la rentabilidad del proyecto contra los costos que presta los servicios ambientales del predio, se tomaron en consideraran los gastos corrientes en cada una de las etapas del proyecto ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.5 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.5.- Análisis Económico por la ejecución del proyecto.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.5 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Para el análisis económico se realizó y se calculó los costos o flujos de cajas que son los ingresos por venta de ____________, contra los egresos o gastos de operación del proyecto, el tiempo estimado desde el desmonte hasta la venta o renta del último _______________ años, contemplando este periodo para la vida útil del proyecto de ___ años, a continuación, se presenta la metodología y los resultados siguientes:")
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
        # Capitulo 13.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.5.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.5.1.- Cálculo del Valor Actual Neto (VAN)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.5.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run("Consiste en actualizar a valor actual presente los flujos de caja futuros que va a generar el proyecto, descontados de un tipo de interés (tasa de descuento), y compararlos con el importe inicial de la inversión. Se utiliza la tasa de descuento mínima del 6 %.") 
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 13.5.1 ###
    #########################
    filas = 2
    columnas = 14
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    #tabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quitar_bordes_tabla(tabla13b)

    for rows in tabla13b.rows:
        rows.cells[0].width = Cm(2.37)
        rows.cells[1].width = Cm(0.63)
        rows.cells[2].width = Cm(1.65)
        rows.cells[3].width = Cm(0.63)
        rows.cells[4].width = Cm(1.65)
        rows.cells[5].width = Cm(0.63)
        rows.cells[6].width = Cm(1.65)
        rows.cells[7].width = Cm(0.63)
        rows.cells[8].width = Cm(1.65)
        rows.cells[9].width = Cm(0.63)
        rows.cells[10].width = Cm(1.65)
        rows.cells[11].width = Cm(0.63)
        rows.cells[12].width = Cm(1.65)
        rows.cells[13].width = Cm(0.74)

    contenido_filas = [
        ["VAN = -A", "+", "FC\u2081", "+", "FC\u2082", "+", "FC\u2083", "+", "FC\u2084", "+", "FC\u2085", "+", "FC.", ".n",],
        [" ", " ", "(1 + r)\u2081", " ", "(1 + r)\u2082", " ", "(1 + r)\u2083", " ", "(1 + r)\u2084", " ", "(1 + r)\u2085", " ", "(1 + r).", ".n",],
    ]

    # Agregar contenido de las filas
    for row_idx, fila in enumerate(contenido_filas, start=0):
        for col_idx, texto in enumerate(fila):
            cell = tabla13b.cell(row_idx, col_idx)
            t13b = cell.paragraphs[0].add_run(texto)
            t13b.font.name = 'Arial'
            t13b.font.size = Pt(10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    variables = [
        "A: Inversión Inicial",
        "F C = Flujos de caja",
        "n = Número de años",
        "r = Tipo de interés Tasa de descuento mínima)",
        "1/(1+r)^n = Factor de descuento para ese tipo de interés y ese número de año",
        "F C d = Flujos de caja Descontados",
        "A = Inversión Inicial + Capital de Trabajo",
        "Si VAN >0= El proyecto es rentable",
        "Si VAN <0= El proyecto no es rentable"
    ]

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('\n')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    
    for texto in variables:
        di13 = doc.add_paragraph()
        descripcionCapitulo13 = di13.add_run(texto)
        descripcionCapitulo13_format = di13.paragraph_format
        descripcionCapitulo13_format.line_spacing = 1.15
        di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        descripcionCapitulo13.font.name = 'Arial'
        descripcionCapitulo13.font.size = Pt(12)
        di13.style = 'List Bullet'

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            VALOR ACTUAL NETO
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para la obtención de este indicador se realizaron con los siguientes datos:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('FAVOR DE PONER EL VALOR NETO =)')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Si VAN > 0 = El proyecto es rentable, la Van del proyecto fue de $ ________ pesos')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.5.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.5.2.- Cálculo de la tasa interna de retorno (TIR)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.5.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Es la tasa de descuento que iguala la VAN a cero. Se le llama tasa interna de retorno porque supone que el dinero que se gana año con año se reinvierte en su totalidad. Es decir, se trata de la tasa de rendimiento generada en su totalidad en el interior de la empresa por medio de reinversión.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Si TIR > a tasa de descuento (r) = El proyecto es factible')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Si TIR < a tasa de descuento (r) = El proyecto no es factible')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Para el cálculo de la ti se tiene lo siguiente: ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            TASA INTERNA DE RETORNO (TIR)
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('FAVOR DE PONER LA TASA INTERNA DE RETORNO (TIR)')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Si TIR > a tasa de descuento (r) = El proyecto es factible')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La ____________________________________________')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('_______________________________________________')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.5.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.5.3.- Cálculo de la Tasa Mínima Aceptable de Rendimiento (Tmar)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.5.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Es también llamada costo de capital o tasa de descuento. Para formarse, toda empresa debe realizar una inversión inicial. _________________________________________. "
        "Como sea que hayan sido las aportaciones del capital, cada uno de ellos tendrá un costo asociado al capital que aporte y la nueva empresa formada tendrá un costo de capital propio."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Antes de invertir, una persona siempre tiene en mente una tasa mínima de ganancia sobre la inversión propuesta, llamada tasa mínima aceptable de rendimiento (TMAR)."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "TMAR = Índice inflacionario + Premio al riesgo Esto significa que la TMAR que un inversionista le pediría a una inversión debe calcularla sumando dos factores: primero, la inflación. "
        "Cuando un inversionista arriesga su dinero, para él no es atrayente mantener el poder adquisitivo de su inversión, sino más bien que ésta tenga un crecimiento real; es decir, "
        "le interesa un rendimiento que haga crecer su dinero más allá de haber compensado los efectos de la inflación."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "En segundo término, debe ser un premio o sobretasa por arriesgar su dinero en determinada inversión."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Cuando se está evaluando un proyecto en un horizonte de tiempo de _______, la TMAR calculada debe ser válida no sólo en el momento de la evaluación sino durante todos los ______ años. "
        "El índice inflacionario para calcular la TMAR, debe ser el promedio del índice inflacionario promedio. Para este caso se tomó como referencia los últimos __ años desde el ________________, siendo los siguientes:"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 13.5.3 ###
    #########################
    filas = 2
    columnas = 11
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    #tabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.name = 'Arial'
            t13b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 13.5.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Fuente: SIE - Inflación (banxico.org.mx) "
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Se utilizó la media geométrica el cual es de _____ % la tasa de inflación promedio."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Ahora que ya se sabe cómo calcular el primer término de los dos que componen la TMAR, hace falta preguntar ¿cuál debe ser el premio al riesgo que deba ganarse? "
        "En términos generales se considera que un premio al riesgo, considerado ahora como la tasa de crecimiento real invertido, habiendo compensado los efectos inflación, "
        "debe ser de entre  _____________%."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Para ellos se tomó como base los riesgos de acuerdo al tipo de proyecto siendo estos los siguientes riesgos:"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Bajo Riesgo. - Si la demanda de tu producto o servicio es estable y No Existe competencia fuerte de otros productos, el porcentaje de riesgo puede ir de _________ %."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Riesgo Medio. - Son proyectos que tiene una demanda variable y competencia considerable, se estima un porcentaje de __________ %"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Riesgo Alto. - Son Negocios en los que el precio del producto cambia mucho debido a la oferta y demanda, se considera un porcentaje superior a _____ %"
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run(
        "Para el caso de nuestro proyecto se tomó un riesgo medio del ______%."
    )
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            TASA MINIMA DE RENDIMIENTO (TMAR)
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('FAVOR DE PONER LA TASA MINIMA DE RENDIMIENTO (TMAR) =)')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capitulo 13.5.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 13.5.3 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Costos de Ingresos del proyecto.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.1 ###
    #########################
    filas = 20
    columnas = 13
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    #########################
    ### Título de la tabla del capítulo 13.5.3 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Costos de Egresos del proyecto.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.1 ###
    #########################
    filas = 20
    columnas = 13
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 13.6
    ########################################################################################################################################################################

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el contenido del capítulo 13.6 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical:
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Titulo del capitulo 13.6 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.- Estimación de los beneficios sociales')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La generación de empleo es parte fundamental de cualquier región para su desarrollo y a la vez proporcionar a la población una mejor calidad de vida tal y como lo establecen los preceptos de desarrollo social, por las características propias de las actividades de aprovechamiento, se ofrecerán áreas de trabajos directos e indirectos a pobladores del municipio y de municipios cercanos, en todas las etapas del proyecto, desde el desmonte hasta la puesta en marcha, Derivado de esta distribución actividades contempladas para la implementación del proyecto se considera la generación de un total de ___ empleos directos beneficiando a una población de más de ____ personas durante la ejecución del proyecto así mismo se podrán generar más ____ empleos adicionales indirectos que juntos harán una derrama económica de más de $___ millones en la operación del proyecto contemplando la preparación del sitio la construcción y de operación del proyecto, beneficiando en primer lugar a la región de __________________, con la generación de empleos, percibiendo una retribución monetaria en la etapa de preparación y construcción,  lo que se traduce en bienestar social de más de _____ familias el cual incrementara la calidad de vida y mantendrá sin problemas económicos durante ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.6 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Costos de Egresos del proyecto.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.6 ###
    #########################
    filas = 15
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 13.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('* Los costos de los salarios es en base a los costos del salario mínimo del país, el cual es de ________ pesos')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.- Medio socioeconómico.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En este apartado se analiza la manera en la que las poblaciones asentadas en el área del sistema ambiental se relacionan con su entorno. Se presenta toda la información relevante para comprender el contexto socioeconómico en el que se lleva a cabo la realización y operación del proyecto en mención, la evaluación se realiza y se analiza para el Municipio de ______________ en el cual pertenece el Área de Cambio de uso de Suelo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.1 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.1.- Nombre de la grafica')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.1 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Poner la grafica de la CONAPO, usar API')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.2 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.2.- Nombre de la grafica')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.2 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Poner la grafica de la CONAPO, usar API')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.3 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.3.- Nombre de la grafica')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.3 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Poner la grafica de la CONAPO, usar API')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.4 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.4.- Nombre de la grafica')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.4 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Poner la grafica de la CONAPO, usar API')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.5 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.5.- Crecimiento poblacional')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.5 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('El crecimiento poblacional de la ciudad de ___________ ha registrado los siguientes datos: a partir del año 1990 hasta el 2020, siendo en el este ultimo de _______ habitantes:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Poner las imagenes')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.6 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.6.- Población económicamente activa (PEA)')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Se entiende por población económicamente activa al grupo de personas en edad de trabajar que realizan una actividad generadora de bienes y servicios a los que se le imputa valor agregado. ')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.6.1.6 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 13.x.- Población Económicamente activa.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.6.1.6 ###
    #########################
    filas = 6
    columnas = 6
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Fuente: Elaboración propia, con datos de censo de Población y vivienda 2020')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(10)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.6 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('La población económicamente activa del municipio es de _______ personas de las cuales ____ son hombre y _____% son mujeres, mientras que la población no económicamente activa es de 33,855 personas siendo de estas un _____% hombre y _____% son mujeres.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.7 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.7.- Uso del Suelo.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.8 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.8.- Recursos Naturales.')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.8 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.9
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.9 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.9.- Monumentos Históricos')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.9 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.10
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.10 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.10.- Fiestas, Danza y Tradiciones')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.10 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.11
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.11 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.11.- Gastronomía')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.11 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.6.1.12
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.6.1.12 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.6.1.12.- Centros turísticos')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.6.1.12 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Descripcion del capitulo')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 13.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 13.7 ###
    #########################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\nXIII.7.- Conclusiones')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True
    capitulo13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 13.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('De acuerdo a la valoración de los servicios ambientales que presta el área deja de percibir con el cambio de uso de suelo $___________ pesos, como se muestra a continuación:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 13.7 ###
    #########################
    tituloTabla13b = doc.add_paragraph()
    dti13b = tituloTabla13b.add_run('\nTabla 9.1.- Periodo de ejecución por etapa.')
    dti13b_format = tituloTabla13b.paragraph_format
    dti13b_format.line_spacing = 1.15
    dti13b_format.space_after = 0

    dti13b.font.name = 'Bookman Old Style'
    dti13b.font.size = Pt(12)
    tituloTabla13b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 13.7 ###
    #########################
    filas = 6
    columnas = 2
    tabla13b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla13b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla13b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)
            t13b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 13.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Con los datos que anteceden se tiene que el proyecto es más rentable económicamente que los beneficios económicos que prestaría el área de estudio si este genera un bien o servicio económico, teniendo los tres indicadores siguientes:')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Favor de poner los indicadores Financieros =)')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    descripcionCapitulo13.bold = True
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Se tiene que la van es mayor a cero dando un resultado positivo de $ ________ Pesos, para la TIR se tiene que para que la inversión sea igual a cero se requiere una tasa interna de retorno de _______%, mientras que la TMAR se tiene una tasa de riesgo más inflación del ___%, siendo la tasa mínima de aceptación de __%, siendo el proyecto viable., con una relación beneficio costo es de ___.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Grafica del capitulo 13.7 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen, mapa o grafico.
    """
    imagenCapitulo5_parrafo = doc.add_paragraph()
    imagenCapitulo5_run = imagenCapitulo5_parrafo.add_run('')
    imagenCapitulo5_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagenCapitulo5_run = doc.add_picture('capitulo13/grafico.jpg', width=Cm(12.24), height=Cm(8.18))

    #########################
    ### Titulo de la grafica del capitulo 13.7 ###
    #########################
    tituloGrafico5 = doc.add_paragraph()
    dgi5 = tituloGrafico5.add_run('Grafica 1.- Nombre de la graficas')
    dgi5_format = tituloGrafico5.paragraph_format
    dgi5_format.line_spacing = 1.15
    dgi5_format.space_after = 0

    dgi5.font.name = 'Bookman Old Style'
    dgi5.font.size = Pt(12)
    tituloGrafico5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 13.7 ###
    #########################
    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Los costos que dejaría de percibir el área de cambio de uso de suelo considerando solamente los servicios ambientales, captura de carbono, el valor de la flora y la fauna, dejaría de percibir la cantidad de $______________ , Pesos, sin embargo, la ganancia del proyecto con todos los gastos es superior a los ______ millones por lo que es más productivo el nuevo uso que el suelo actual de área de cambo de uso de suelo.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('Bajo una perspectiva de valoración, el cambio de uso de suelo propuesto por la ejecución del proyecto _______________________________________ representa una afectación completa de los servicios ambientales que presta la vegetación forestal. En este caso, la remoción de la cubierta vegetal repercutirá en la disminución de la recarga de los mantos acuíferos, aunado a esto la zona donde se pretende desarrollar el proyecto no cuenta con cualidades para prestar servicios ambientales que tengan un valor cultural, religioso o como espacios importantes para la recreación.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di13 = doc.add_paragraph()
    descripcionCapitulo13 = di13.add_run('En otras palabras, los principales impactos más notorios ocurrirán sobre los recursos naturales Geomorfología, Suelo, Flora y Paisaje.')
    descripcionCapitulo13_format = di13.paragraph_format
    descripcionCapitulo13_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo13.font.name = 'Arial'
    descripcionCapitulo13.font.size = Pt(12)
    di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista_paisaje = [
        "En el caso del paisaje, que durante muchos años será apreciable a la distancia, modificando su topografía; siendo este uno de los factores de mayor alteración.",
        "No se pone en riesgo el conjunto paisajístico de la zona dado que el área del proyecto y en sus alrededores no se encuentran áreas eco turísticas o en las que se promuevan su desarrollo y que pudieran ser impactados con la implementación del proyecto en el sistema ambiental _________________________________________ __________________________________ y aledaños a ella.",
        "El área sujeta a estudio comprende una superficie total de ________ ha, de las cuales serán removidas en su totalidad, donde solicita el cambio de uso de suelo, cuyas características topográficas de ________________________________, actualmente ________________, por lo que se creará el proyecto para la ____________________.",
        "El terreno es rustico, _________________ __________. La mayor cobertura vegetal es de ___________________.",
        "De acuerdo a las evaluaciones realizadas en la flora y fauna, se puede determinar que la ejecución del presente proyecto no afectará a especies ya que estas serán rescatadas si así se requiere. Aunque será retirada la cubierta vegetal de la superficie del proyecto, esto no pondrá en riesgo la biodiversidad en ámbito del sistema ambiental hidrológica, debido a la escasa presencia de especies y a lo común de su distribución. Se hará el rescate y reubicación del estrato suculento.",
        "Los impactos ambientales que genera el desarrollo de las actividades correspondientes al cambio de utilización de terreno forestal para el ___________ _______________, son drásticos debido a la naturaleza del proyecto. Sin embargo, estos son focalizados por lo que únicamente se manifiestan en el área, sin afectar a áreas o poblaciones aledañas.",
        "El polígono del área sujeta de cambio de uso de suelo, cuenta con una estimación de susceptibilidad a la erosión hídrica en las condiciones actuales principalmente de _____ mm/ha, en la erosión eólica presenta _____ mm/ha, con la implementación del proyecto incrementa en ____ mm/ha, en erosión eólica, y hasta ______ mm/ha, en erosión hídrica.",
        "La tasa máxima permisible de pérdidas de suelo es de 10 Ton/ha/año; mayores pérdidas, como en el caso del polígono propuesto, significan degradación (SAGARPA, 2000).",
        "La erosión con proyecto, anteriormente calculada considera que no existiera cobertura del suelo (suelo desnudo) y no se tuvieran prácticas de conservación de suelo y del agua.",
        "De la infiltración podemos deducir que se tiene una infiltración normal en el ACUSTF de ______ mm anuales, al quedar sin vegetación aumenta la evapotranspiración, lo que se dejara de infiltrar en la vigencia del proyecto la cantidad de ______ mm anual, misma que se podrán recuperar con la implementación de las obras de conservación tendiendo una captura anual de _______ mm, mismas que serán filtradas en el cambio de uso de suelo."
    ]

    # Bucle para agregar cada ítem como viñeta al documento
    for item in lista_paisaje:
        di13 = doc.add_paragraph(style='List Bullet')  # Estilo de lista con viñeta
        descripcionCapitulo13 = di13.add_run(item)
        
        # Formato de cada ítem
        descripcionCapitulo13_format = di13.paragraph_format
        descripcionCapitulo13_format.line_spacing = 1.15
        descripcionCapitulo13.font.name = 'Arial'
        descripcionCapitulo13.font.size = Pt(12)
        di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 13 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo13() # Crear el documento
