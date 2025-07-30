"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo7():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 7
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo VII.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 7
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 7
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 7 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'VII ANÁLISIS COMPARATIVO DE LAS TASAS DE EROSIÓN DE LOS SUELOS, ASÍ COMO LA CALIDAD, CAPTACIÓN E INFILTRACIÓN DEL AGUA EN EL ÁREA SOLICITADA RESPECTO A LAS QUE SE TENDRÍAN DESPUÉS DE LA REMOCIÓN DE LA VEGETACIÓN FORESTAL.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.1 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.1.- Erosión hídrica en el área de Cambio de Uso de Suelo.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('En cuanto a la metodología y datos utilizadas se encuentran dentro del numeral V, Subnumeral V.3.1.1.- Metodología para determinar la erosión hídrica en el área de Cambio de Uso de Suelo, con los datos siguientes:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Estimación del valor del Factor de longitud y grado de la pendiente del ACUSTF ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Estimación del valor del Factor de longitud y grado de la pendiente del ACUSTF')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.bold = True
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Para la cual se tiene los siguientes resultados.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.1.1 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\n')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.1.1 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.1.- Porcentaje de la cubierta vegetal en el ACUSTF.')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.1.1 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 7.1.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nDe la misma manera, para estimar el valor de LS se hace necesario tomar en cuenta las características topográficas del polígono de afectación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Estimación del valor del Factor de longitud y grado de la pendiente del ACUSTF ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Estimación del valor del Factor de longitud y grado de la pendiente del ACUSTF')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.1.1 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.2.- Valor de longitud y grado de la pendiente del ACUSTF.')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.1.1 ###
    #########################
    factor711 = [
        'Altura más alta del terreno',
        'Altura más baja del terreno',
        'Longitud del Pendiente',
        'Pendiente (S)',
        'Factor de grado y longitud de la pendiente (L S)'
    ]

    factor711Rango = range(len(factor711))
    cols711 = len(factor711) + 1

    tabla7b = doc.add_table(rows=cols711, cols=2, style='Table Grid')
    tabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for rows in tabla7b.rows:
        rows.cells[0].width = Cm(10.11)
        rows.cells[1].width = Cm(3.13)

    cell = tabla7b.cell(0, 0)
    t7b = cell.paragraphs[0].add_run('Factor')
    t7b.font.size = Pt(12)
    t7b.font.name = 'Arial'
    t7b.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '0070C0')

    cell = tabla7b.cell(0, 1)
    t7b = cell.paragraphs[0].add_run('Valor')
    t7b.font.size = Pt(12)
    t7b.font.name = 'Arial'
    t7b.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '0070C0')

    for factor in factor711Rango:
        celda = factor + 1
        cell = tabla7b.cell(celda, 0)
        t7b = cell.paragraphs[0].add_run(factor711[factor])
        t7b.font.size = Pt(12)
        t7b.font.name = 'Arial'
        t7b.font.bold = True

    for valor in factor711Rango:
        celda = valor + 1
        cell = tabla7b.cell(celda, 1)
        t7b = cell.paragraphs[0].add_run(' ')
        t7b.font.size = Pt(12)
        t7b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 7.1.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nSacar lo de las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nFinalmente, se estima la Erosión Potencial (Ep) sustituyendo estos valores en la ecuación, obtenidos en la fórmula: Ep = R*K*LS*C. Los resultados se presentan en la tabla siguiente:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.1.1 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.3.- Erosión potencial para el Cambio de Uso de Suelo.')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.1.1 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 7.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.1.2 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.1.2.- Erosión hídrica en la condición actual del ACUSTF.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.1.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('La Erosión Potencial (Ep) sustituyendo estos valores en la ecuación, obtenidos en la fórmula: Ep = R*K*LS*C, dado lo siguiente el valor de R sería de ___________ de acuerdo a la tabla de ecuaciones para estimar la erosión de la lluvia y el mapa de regiones de la erosión de la lluvia, en la tabla de valores del factor K de acuerdo al tipo de vegetación y el porcentaje de la cobertura de la misma arroja un valor de ______ y el factor de longitud de grado de pendiente que es igual a ______ como resultado tenemos que la erosión potencial sería de ___________________________________________________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.1.2 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.1.3.- Erosión potencial con el cambio de uso de suelo.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.1.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Para calcular la pérdida de Suelo se aplicará la ecuación potencial de acuerdo a la siguiente fórmula utilizando los valores obtenidos de las variables R, K, LS, quedando como sigue: ')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Ep = R*K*LS')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir la formula anterior')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.1.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.1.4 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.1.5. Resultados de erosión hídrica en el área del ACUSTF.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.1.4 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.4.- Erosión hídrica en el ACUSTF')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.1.4 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 7.1.4 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('En las condiciones actuales por efecto de la lluvia se pueden tener pérdidas de _____________________, con la implementación del proyecto al quedar desnudo el suelo incrementa una pérdida hasta ______________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('La pérdida total de suelo en un periodo de _________ que estará el suelo sin vegetación o desnudo por las actividades de remoción será de ____________, al igual que, durante el periodo de extracción de material pétreo será de _______________________, como se muestra a continuación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir el periodo')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\n')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 7.1.4 ###
    #########################
    tabla7b = doc.add_table(rows=3, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Se aclara que, la duración del proyecto será de ________________________________________ has, por lo que la _____________________________ serán un total de ________, de las cuales el aprovechamiento pro año será de ___________, por año quedando sin vegetación __________ para la remoción de la vegetación y __________ en cada etapa para la extracción, en total durante los _____________________________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('En el área del proyecto con los datos que anteceden y las observaciones realizadas en el recorrido de campo se considera que la calidad de las condiciones del suelo se encuentra de ____________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.2 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.2.- Erosión eólica en el área de Cambio de Uso de Suelo.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Para la estimación de la erosión eólica en el área de cambio de uso de suelo se realizaron con la metodología que se encuentra en el numeral V, y Subnumeral V.3.2.- Metodología para determinar la erosión eólica en CUSTF, con los datos siguientes: ')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.2 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.5.- Valor de factores para cálculo de erosión en el área de CUSTF.')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.2 ###
    #########################
    tabla7b = doc.add_table(rows=7, cols=2, style='Table Grid')

    for cols in range(2):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### FACTOR 'G' ###
    #########################
    ### Descripcion del capitulo 7.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Para obtener el factor G se utilizó información meteorológica de CONAGUA, en su estación, _________________________________, para obtener información sobre la velocidad del viento, se obtuvo de la página ___________________________________________, para obtener la información mensual del año inmediato anterior, con ello y utilizando la fórmula Factor climático:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nG.C = 1/100 i =1∑12 (Vel / 100) (((PET - P) / PET) * n)\n')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Título de la tabla del capítulo 7.2 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.6.- Datos para el factor G de erosión eólica ')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.2 ###
    #########################
    tabla7b = doc.add_table(rows=4, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(4):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 7.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nSustituyendo la fórmula se obtiene lo siguiente:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir la formula')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nUtilizando la información anterior y la ecuación se tiene lo siguiente:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.2.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.2.1 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.2.1.- Estimación de la erosión eólica actual en el ACUSTF ')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.2.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.2.2 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.2.2.- Erosión eólica con la implementación del proyecto')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.2.2 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.2.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.2.3 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.2.3.- Erosión potencial con el cambio de uso de suelo')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.2.3 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.2.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.2.4 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.2.4.- Resultados de erosión eólica del ACUSTF')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.2.4 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.7.- Erosión eólica en el ACUSTF')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.2.4 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'


    #########################
    ### Descripcion del capitulo 7.2.4 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('En las condiciones actuales por efecto del viento se tiene una pérdida de suelo de _____________________, con la implementación del proyecto al quedar desnudo el suelo se incrementa hasta ____________________, por lo que se recomienda realizar actividades de compensación, logrando tener una erosión potencial solo de ____________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('En el área del proyecto con los datos que anteceden y las observaciones realizadas en el recorrido de campo se considera que la calidad de las condiciones del suelo se encuentra de _____________________ con tendencia a degradación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('La pérdida total de erosión eólica que tendrá por acciones del viento durante el tiempo que el área quedara sin vegetación es de ________________________ y el tiempo que durará la extracción de material pétreo por etapa de extracción será de __________, la pérdida por cambio de uso de suelo será de __________, como se muestra a continuación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir el periodo')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 7.2.4 ###
    #########################
    tabla7b = doc.add_table(rows=3, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### Título de la tabla del capítulo 7.2.4 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.8.- Grado de afectación de la erosión hídrica y eólica en el ACUSTF')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.2.4 ###
    #########################
    tabla7b = doc.add_table(rows=7, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(7):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 7.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\n')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.3.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3.1 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.3.1.- Metodología para el cálculo de infiltración.')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.3.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Para la cuantificación del volumen medio anual de escurrimiento natural se determinó indirectamente, mediante la siguiente expresión:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nVolumen Anual de Escurrimiento = Precipitación Anual * Área Total * Coeficiente de Escurrimiento\n')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(11)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Área Total ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Área Total')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.bold = True
    descripcionCapitulo7.underline = True
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir las formulas')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Coeficiente de escurrimiento ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nCoeficiente de Escurrimiento')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.bold = True
    descripcionCapitulo7.underline = True
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('El cual se calcula mediante las fórmulas siguientes:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del Capítulo 7.3.1 ###
    #########################
    tabla7 = doc.add_table(cols=2, rows=3, style='Table Grid')

    cell = tabla7.cell(0, 0)
    t7 = cell.paragraphs[0].add_run('COEFICIENTE DE ESCURRIMIENTO ANUAL (Ce)')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    t7.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_background_color(cell, '4F81BD')

    cell = tabla7.cell(0, 1)
    t7 = cell.paragraphs[0].add_run('K: PARÁMETRO QUE DEPENDE DEL TIPO Y USO DE SUELO')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    t7.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_background_color(cell, '4F81BD')

    cell = tabla7.cell(1, 0)
    t7 = cell.paragraphs[0].add_run('Ce = K(P-250) / 2000')
    t7.font.size = Pt(12)
    t7.font.name = 'Times New Roman'
    t7.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(1, 1)
    t7 = cell.paragraphs[0].add_run('Si K resulta menor o igual que 0.15')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(2, 0)
    t7 = cell.paragraphs[0].add_run('Ce = (K(P-250) / 2000) + (K - 0.15) / 1.5')
    t7.font.size = Pt(12)
    t7.font.name = 'Times New Roman'
    t7.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(2, 1)
    t7 = cell.paragraphs[0].add_run('Si K es mayor que 0.15')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in range(2):
        cell = tabla7.cell(0, cols)
        cell_background_color(cell, '4F81BD')

        for rows in range(3):
            cell = tabla7.cell(rows, cols)
            cell.height = Cm(1.22)
            cell.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    #########################
    ### Descripción del capítulo 7.3.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run(
        "\nDónde:"
        '\nCe = Coeficiente de escurrimiento para diferentes superficies'
        '\nP = Precipitación media anual'
        '\nK = Factor que depende de la cobertura arbolada y del tipo de suelo la cual se describe en el siguiente cuadro:'
    )
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nValores del Factor K')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 7.3.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    formulaCapitulo7 = doc.add_paragraph()
    formulaCapitulo7.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    fCapitulo7 = formulaCapitulo7.add_run()
    fCapitulo7.add_picture('capitulo7/capitulo731/tabla731.png', width=Cm(9.27), height=Cm(9.36))  # Nombre del archivo, debe estar en la carpeta correcta
    formulaCapitulo7.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Tabla del capitulo 7.3.1 ###
    #########################
    tabla7 = doc.add_table(rows=4, cols=2, style='Table Grid')

    cell = tabla7.cell(0, 0)
    t7 = cell.paragraphs[0].add_run('TIPO DE SUELO')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    t7.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla7.cell(0, 1)
    t7 = cell.paragraphs[0].add_run('CARACTERISTICAS')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    t7.font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, '4F81BD')

    cell = tabla7.cell(1, 0)
    t7 = cell.paragraphs[0].add_run('A')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(2, 0)
    t7 = cell.paragraphs[0].add_run('B')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(3, 0)
    t7 = cell.paragraphs[0].add_run('C')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla7.cell(1, 1)
    t7 = cell.paragraphs[0].add_run('Suelos permeables, tales como arenas profundas y loess poco compactados')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla7.cell(2, 1)
    t7 = cell.paragraphs[0].add_run('Suelos medianamente permeables, tales como arenas de mediana profundidad: loess algo más compactos que los correspondientes a los suelos A; terrenos migajosos')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cell = tabla7.cell(3, 1)
    t7 = cell.paragraphs[0].add_run('Suelos casi impermeables, tales como arenas o loess muy delgados sobre una capa impermeable, o bien arcillas')
    t7.font.size = Pt(12)
    t7.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for width in range(4):
        tabla7.cell(width, 0).width = Cm(4.1)

    for width in range(4):
        tabla7.cell(width, 1).width = Cm(13.09)

    #########################
    ### Descripcion del capitulo 7.3.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Fuente: NOM-011-CNA-2000')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Tomando en consideración la condición del suelo presente en el área y de acuerdo a INEGI se determinó que es un tipo de suelo _____ y de acuerdo a la información recabada en campo se cuenta con una cobertura vegetal de _________________________________________ de materia orgánica por lo que nos da un factor de K de ___________ por lo anterior que el Coeficiente de Escurrimiento Anual (Ce) se determinará a través de la siguiente fórmula:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nCe= K ((P-250)/2000)) + (K-0.15)1.5')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nDescribir el resto')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nPor lo anterior el volumen medio anual de escurrimiento natural se determinó mediante el método indirecto, mediante la siguiente expresión:')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nVe= (P) (At) (Ce)')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER


    ########################################################################################################################################################################
    # Capitulo 7.3.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3.1.1 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.3.1.1.- Evapotranspiración por el método de Coutagne')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.3.1.1 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('ETR = P-xP2')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run(
        'Donde'
        '\nETR= Evapotranspiración m/año'
        '\nP= Precipitación en m/año'
        '\nX= 1/ (0.8 + 0.14 t)'
        '\nDescribir el resto'
    )
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nCon los datos necesarios calculados se podrá obtener el grado de infiltración en el área sujeta a Cambio de Uso del Suelo desde tres escenarios tal y como se manifiesta a continuación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3.2 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.3.2.- Infiltración sin proyecto en el ACUSTF')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.3.2 ###
    #########################

    #########################
    ### INFILTRACION ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('INFILTRACION')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.bold = True
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Infiltración = P – ETR – Ve')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run(
        'Donde'
        '\nDescribir el resto'
    )
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 7.3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3.3 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.3.3.- Infiltración con la implementación del proyecto')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 7.3.3 ###
    #########################
    
    #########################
    ### INFILTRACION ###
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('INFILTRACION')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.bold = True
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Infiltración = P – ETR – Ve')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Times New Roman'
    descripcionCapitulo7.font.size = Pt(12)
    descripcionCapitulo7.italic = True
    di7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run(
        'Donde'
        '\nDescribir el resto'
    )
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 7.3.3 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.9.- Volumen de escurrimiento en el ACUSTF ')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.3.3 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('CE: Coeficiente de escurrimiento; VE: Volumen de escurrimiento')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alingnment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nConsiderando la información antes señalada, se interrumpe un volumen de escurrimiento de agua de __________ a la superficie del ACUSTF de _______________. La cual se puede capturar con la implementación de obras de conservación.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alingnment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 7.3.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 7.3.4 ###
    #########################
    capitulo7 = doc.add_paragraph()
    i7 = capitulo7.add_run(f'\nVII.3.4.- Resultados obtenidos de la Infiltración')
    i7_format = capitulo7.paragraph_format
    i7_format.line_spacing = 1.15

    i7.font.name = 'Arial'
    i7.font.size = Pt(12)
    i7.font.bold = True
    capitulo7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 7.3.4 ###
    #########################
    tituloTabla7b = doc.add_paragraph()
    dti7b = tituloTabla7b.add_run('\nTabla 7.10.- Infiltración en el ACUSTF para los tres escenarios.')
    dti7b_format = tituloTabla7b.paragraph_format
    dti7b_format.line_spacing = 1.15
    dti7b_format.space_after = 0

    dti7b.font.name = 'Bookman Old Style'
    dti7b.font.size = Pt(12)
    tituloTabla7b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 7.3.4 ###
    #########################
    tabla7b = doc.add_table(rows=2, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(2):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 7.3.4 ###
    #########################
    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Derivado del análisis se concluye que en la condición actual con la cobertura que posee, se tiene una infiltración normal de _________ anuales, con la implementación del proyecto al quedar sin vegetación esto aumenta la evapotranspiración por lo cual se dejará de captar agua reduciendo su infiltración a _________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('La pérdida de infiltración total que sufrirá el área durante el periodo que durará sin vegetación es de ________ en un periodo de ____________, así mismo durante el periodo de __________________________________________ la pérdida será de ___________, posteriormente a este periodo el área quedará intacta para su recuperación de manera natural.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('Describir el periodo')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\n')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 7.3.4 ###
    #########################
    tabla7b = doc.add_table(rows=3, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla7b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(3):
            cell = tabla7b.cell(rows, cols)
            t7b = cell.paragraphs[0].add_run(' ')
            t7b.font.size = Pt(12)
            t7b.font.name = 'Arial'

    di7 = doc.add_paragraph()
    descripcionCapitulo7 = di7.add_run('\nSe aclara que, la duración del proyecto será de _________ en una superficie total de _________, por lo que la __________________________________________ ha por año quedando sin vegetación _________ para la remoción de la vegetación y __________________________________________.')
    descripcionCapitulo7_format = di7.paragraph_format
    descripcionCapitulo7_format.line_spacing = 1.15
    descripcionCapitulo7_format.space_after = 0
    descripcionCapitulo7_format.space_before = 0

    descripcionCapitulo7.font.name = 'Arial'
    descripcionCapitulo7.font.size = Pt(12)
    di7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 7 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo7() # Crear el documento