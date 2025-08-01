"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo9():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 9
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo IX.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 9
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 9
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 9 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'IX PLAZO Y FORMA DE EJECUCIÓN DEL CAMBIO DE USO DE SUELO.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'IX.1.- Plazo de ejecución.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.1.- Periodo de ejecución por etapa.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=3, style='Table Grid')

    for cols in range(3):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 9.1 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Describir en este parrafo.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('La elaboración del Estudio Técnico Justificativo es con el propósito de solicitar a la autoridad normativa la autorización para poder llevar a cabo el cambio de uso de suelo en esta área que sustenta vegetación clasificada como forestal (______________________) _____________________.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('El período considerado para llevar a cabo las actividades necesarias para el cambio de uso del suelo (desmonte y despalme) ________________________________, y ello en virtud de que serán aprovechado paulatinamente y cada año se pretende realizar rescate de algunas especies de interés y colocarlas dentro del área del predio en las mismas condiciones para un mayor éxito en la reubicación, para el desarrollo total del proyecto se contempla la preparación del sitio y construcción e inicio de operación en un período de __ años, aproximadamente período en el cual se incluye la realización de los estudios correspondientes con base a la Normatividad y de ser autorizados para iniciar los procesos de operación mediante las acciones planteadas para iniciar con __________________________________ de acuerdo a las necesidades del promovente y la demanda en los mercados locales y regionales. Proceso que se tiene contemplado iniciar en caso de ser aprobado en el __________________. El cual se ejecutará de la siguiente manera: ')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.- Formas de ejecución del cambio de uso de suelo en terrenos forestales.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('De acuerdo a las observaciones y estudios geológicos realizados, se determinó la viabilidad y factibilidad para ___________________________________ y necesidad de contar con el permiso para realizar el cambio de uso de suelo el nuevo uso será __________________________________, considerando que en la superficie solicitada se tienen las condiciones requeridas para tal fin es necesario efectuar el documento técnico unificado para Cambio de Uso de Suelo, así como la manifestación de impacto ambiental y aprovechar la disponibilidad de la superficie planteada.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Ante ello en el presente documento se manifiestan los procesos, tiempos calendario acorde a la proyección que tiene la promovente así como su secuencia para la ejecución del proyecto una vez que se diagnostica y elabora el documento correspondiente donde se plasman las condiciones actuales del área de cambio de uso de suelo, en lo  referente a factores ambientales Bióticos y Abióticos que puedan ser afectados como agua y suelo como es el uso actual del terreno y la justificación para su cambio de uso desde el punto de vista económico y social para la región y el Estado.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Previamente para efectuar el estudio para el ACUSTF se realiza lo siguiente')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista92 = [
        'Visita preliminar al área de Estudio',
        'Análisis General del área en el Sistema de Información Geográfica',
        'Obtención de información de campo mediante la aplicación de muestreos a la flora y fauna y observación general de las condiciones del área sujeta de estudio y sistema ambiental',
        'Desarrollo del Estudio con base a la Normatividad aplicable',
        'Gestión de trámite de autorización de permisos correspondientes',
    ]

    

    lista92Rango = range(len(lista92))
    

    for lista in lista92Rango:
        di9 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo9 = di9.add_run(f'{lista92[lista]}')
        descripcionCapitulo9_format = di9.paragraph_format
        descripcionCapitulo9_format.line_spacing = 1.15

        descripcionCapitulo9.font.name = 'Arial'
        descripcionCapitulo9.font.size = Pt(12)
        di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('\n')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 9.2 ###
    #########################
    columnas92 = [
        'Preparación del Sitio',
        'Construcción',
        'Operación',
        'Abandono del Sitio',
    ]

    columnas92Rango = range(len(columnas92))
    columnas92Len = len(columnas92)

    tabla9b = doc.add_table(rows=3, cols=columnas92Len, style='Table Grid')

    #########################
    ### Celda fusionada ###
    row = tabla9b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[1].merge(row.cells[3]))

    # Agregar texto a la celda fusionada
    t9b = merged_cell.paragraphs[0].add_run('Secuencia de las Etapas del Proyecto')
    t9b.font.name = 'Arial'
    t9b.font.size = Pt(12)
    t9b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, 'D5DBDB')  # Cambiar el color de fondo de la celda fusionada

    for cols in columnas92Rango:
        cell = tabla9b.cell(1, cols)
        t9b = cell.paragraphs[0].add_run(f'{cols + 1}. {columnas92[cols]}')
        t9b.font.name = 'Arial'
        t9b.font.size = Pt(12)
        t9b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell_background_color(cell, '5499C7')

        cell = tabla9b.cell(2, cols)
        cell_background_color(cell, 'A9CCE3')

        for actividades in columnas92Rango:
            cell = tabla9b.cell(2, actividades)
            t9b = cell.paragraphs[0].add_run(f'Actividad {actividades + 1}\n')
            t9b.font.name = 'Arial'
            t9b.font.size = Pt(12)

        for rows in range(3):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(f'')
            t9b.font.name = 'Arial'
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.1- ACTIVIDADES A DESARROLLAR PARA EL CAMBIO DE USO DE SUELO')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.1 ###
    #########################

    #########################
    ### ACTIVIDADES INHERENTES AL CAMBIO DE USO DE SUELO ###
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('ACTIVIDADES INHERENTES AL CAMBIO DE USO DE SUELO')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    descripcionCapitulo9.bold = True
    di9.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Imagen del capítulo 9.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo9_parrafo = doc.add_paragraph()
    imagenCapitulo9_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo9_run = imagenCapitulo9_parrafo.add_run('')
    imagen_cap_9 = imagenCapitulo9_run.add_picture('capitulo9/capitulo921/cap_921.png', width=Cm(17.68), height=Cm(3.17))

    # Opcional: espacio después del párrafo
    imagenCapitulo9_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 9.2.1 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripción de las Actividades: En este punto se describen en forma secuencial cada una de las actividades proyectadas en cada una de las etapas para la implementación del proyecto como resultado del estudio técnico para el cambio de uso de suelo.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('\nEquipos que se utilizarán para las etapas del cambio de uso de suelo.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    descripcionCapitulo9.bold = True
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista921 = [
        'Bulldozer',
        'Mononiveladoras',
        'Cargador Frontal',
        'Camiones de volteo',
        'Pipas de Riego (Camion cisterna)'
    ]

    lista921Rango = range(len(lista921))

    for lista in lista921Rango:
        di9 = doc.add_paragraph()
        descripcionCapitulo9 = di9.add_run(f'{lista + 1}.- {lista921[lista]}')
        descripcionCapitulo9_format = di9.paragraph_format
        descripcionCapitulo9_format.line_spacing = 1.15
        descripcionCapitulo9_format.space_after = 0
        descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo9.font.name = 'Arial'
        descripcionCapitulo9.font.size = Pt(12)
        di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2- Formas de Ejecución')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nX.2.2.1.- Etapa: Preparación del Sitio.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.1 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Para la ejecución del cambio de uso de suelo se hace mención que durante los primeros ________ de cada año se realizará la remoción de _________ aproximadamente durante los ___________________________________________')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.2.2.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.2.- Calendario de actividades Etapa de Preparación el sitio')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.2.2.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=17, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(17):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.2.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2.2.- Delimitación del Área para CUSTF: ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.2 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Esta actividad consiste en la ubicación de los vértices del área en estudio, la poligonal que delimita en campo la superficie donde se pretende llevar a cabo el cambio de uso de suelo, dicha actividad requiere un plazo de tiempo muy corto ya que la superficie del área es pequeña permitiendo la fácil movilidad para el traslado interno y para los desplazamientos, se marcan los limites geo referenciándolos tomando como base  tanto las escrituras del predio como el área propuesta para la implementación del proyecto, con base a ello se determina el área sujeta de estudio para el ACUSTF, posterior a este proceso se determina el diseño y sitios de muestreo en referencia a las condiciones de vegetación existente. ')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.2.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2.2.1.- Delimitación de áreas para cada proceso proyectado.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.2.1 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('De acuerdo al proyecto ejecutivo se tiene proyectado efectuar la _____________________________')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2.3.- Programa para Rescate y Conservación de Flora y Fauna Silvestre.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.3 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('En el proceso de obtención de datos para los sitios de muestreo que se llevó a cabo dentro del área de cambio de uso de suelo, en los cuales se establecieron __________________________________________________________________________________________ las especies fueron ubicadas dentro del área de estudio durante los levantamientos de sitios, antes de cualquier actividad se recorrerá a detalle el área para ubicar más especies de interés ecológico, antes de realizar el desmonte. Se ha considerado que durante las diferentes etapas del proyecto y en cada uno de los procesos en el período de ejecución contemplado  si es ubicada algún tipo de vegetación que requiera ser protegido y que durante el rescate se haya obviado sea rescatada y ubicada dentro de  sitios  cercanos al área del proyecto con las mismas características para lo cual se ha designado específicamente, dicha área será designada por las mismas características del ecosistema donde se ubica el ACUSTF sujeto de análisis. Por otra parte, y de igual manera para las especies faunísticas, las especies consideradas para su rescate, serán aquellas especies de lenta movilidad y por consiguiente vulnerables, como lo son las especies de reptiles estas serán rescatadas y reubicadas dentro de las áreas del predio y en el cual también se describen cada una de las actividades a realizar en el Programa de Rescate de Fauna Silvestre (capitulo 10).')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2.3.- Remoción de Vegetación (Desmonte). ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.4 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Consiste en la remoción de la vegetación en el área, misma que se pretende realizar durante un periodo de ________________________________________________________________________ se aprovechara por año, propuesta para cambio de uso de suelo, con el objeto de que la superficie restante no sufra más impactos por erosión eólica e hídrica,  la remoción se llevará a cabo de acuerdo al proyecto del promovente tomando como base al período de ejecución, para esta actividad se utilizará maquinaria pesada tipo Bulldozer y será en una forma inmediata una vez obtenidas las autorizaciones correspondientes con el objeto de ir integrando las áreas al proceso y no dejar expuesto el suelo a la erosión hídrica y eólica; en caso de que en esta etapa se lleguen a identificar algunas especies de Flora y Fauna de interés, serán rescatadas. ')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.2.5 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.2.4.- Despalme.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.2.5 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Esta actividad se llevará a cabo con el uso de maquinaria del tipo Bulldozer y será con la finalidad de eliminar los residuos del desmonte eliminando la capa arable y/o suelo orgánico para en forma posterior iniciar cada uno de los procesos proyectados __________________________________________________________. Los residuos del desmonte serán depositados en áreas que no serán afectadas con la finalidad de poder ser reutilizados en los procesos de restauración al término del proyecto, si así se requiere.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\n')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.2.3 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.3.- Calendario de actividades Etapa de Construcción')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.2.3 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=17, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(17):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.2.3.x
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.3.x ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.3.X.- Rampa:')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.3.x ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripcion del Capitulo 9.2.3.x')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.3.x
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.3.x ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.3.x.- Área de Cargaderos.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.3.x ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripcion del Capitulo 9.2.3.x')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.4.- Etapa de Operación.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.2.4 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.4.- Calendario de actividades Etapa de Operación.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.2.4 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=17, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(17):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.2.4.x
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.4.x ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.4.x.- Trituración.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.4.x ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripcion del Capitulo 9.2.4.x')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.4.x
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.4.x ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.4.x.- Carga y Acarreo.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.4.x ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripcion del Capitulo 9.2.4.x')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.- Etapa de Abandono del Sitio.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.2.5 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.5.- Calendario de actividades Etapa de Abandono del Sitio.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.2.5 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=17, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(17):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 9.2.5 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('\nUna vez que el proyecto ha llegado a su vida útil se procederá a realizar las actividades siguientes:')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.1.- Desmantelamiento de infraestructura de apoyo.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.5.1 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripción')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.2.- Para el cierre de accesos al área.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.5.2 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripción')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.2.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.3.- Plan de Cierre de Terreros Estériles.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.5.3 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Descripcion')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.5.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5.5 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.5.- Plan de Cierre de Infraestructuras.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.5.5 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.2.5.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.2.5.6 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.2.5.6.- Plan de Cierre de Manejo de residuos y otros.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.2.5.6 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\n')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.3 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Son aquellas acciones que se pueden anticipar a los efectos o modificaciones que se pudieran generar durante el desarrollo del proyecto contemplando cada una de sus etapas, plasmando su forma de aplicación a fin de minimizar los impactos ambientales de carácter Moderado, Severo y Critico para cada etapa del proyecto y en cada una de las acciones donde se estarán manifestando, para ello se efectúa la adecuación pertinente incluyendo los objetivos principales de dichas medidas de acuerdo a la siguiente tabla:')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.3.1- Preparación del Sitio.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.3.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.1.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.3.1.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.1.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.6.- Impactos generados en la _________________________. Etapa de preparación del sitio')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.1.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.3.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.1.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI..3.1.2')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.1.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.7.- Impactos generados en el ___________. Etapa de preparación del sitio')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.1.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.3.2.- Etapa de Construcción')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.3.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.2.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.2.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.2.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.8.- Impactos generados en ________________, etapa de Construcción.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.2.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.2.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.2.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.2.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.9.- Impactos generados en _________________, etapa de Construcción.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.2.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.2.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.2.3.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.2.3 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.10.- Impactos generados en los ________________________, etapa de Construcción')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.2.3 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.2.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.2.4.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.2.4 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.11.- Impactos generados en las ____________________, etapa de Construcción')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.2.4 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.3.3.- Etapa de operación')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.3.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.3.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.3.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.3.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.12.- Impactos generados en ________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.3.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.3.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.3.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.3.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.13.- Impactos generados en el _____________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.3.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.3.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.3.3.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.3.3 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.14.- Impactos generados en la _______________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.3.3 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.3.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.3.4.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.3.4 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.15.- Impactos generados en la carga y transporte de material, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.3.4 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.3.4.- Etapa de Abandono')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.3.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.4.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.4.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.4.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.16.- Impactos generados en la __________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.4.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.3.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.3.4.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.3.4.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.3.4.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.17.- Impactos generados en la _____________________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.3.4.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.4.- Medidas de Mitigación:')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 9.4 ###
    #########################
    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('Es el conjunto de acciones que se implementan una vez que se identifica el impacto y la magnitud del mismo, con la finalidad de minimizar en lo posible los efectos de dicho impacto sobre todo aquellos que prevalecerán aun con la aplicación de las medidas preventivas, sobre todo en aquellas acciones del proyecto que son inevitables y de carácter severo, Moderado, Severos y Críticos para algunos factores ambientales y sus componentes que lo integran.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph()
    descripcionCapitulo9 = di9.add_run('\nObjetivos de las medidas:')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    descripcionCapitulo9.bold = True
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    objetivosMedidas = [
        'Reducir los impactos a través de la limitación de su magnitud',
        'Rectificar el impacto a través de la reparación, rehabilitación o restauración del componente ambiental afectado',
        'Minimizar o eliminar el impacto a travése del tiempo con la implementación de actividades resultado de los ananlisis aplicados mediante la organización establecida para la conservación y mantenimiento durante la vida del proyecto.'
    ]

    objetivosMedidasRango = range(len(objetivosMedidas))

    for lista in objetivosMedidasRango:
        di9 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo9 = di9.add_run(f'{objetivosMedidas[lista]}')
        descripcionCapitulo9_format = di9.paragraph_format
        descripcionCapitulo9_format.line_spacing = 1.15
        descripcionCapitulo9_format.space_after = 0
        descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo9.font.name = 'Arial'
        descripcionCapitulo9.font.size = Pt(12)
        di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di9 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo9 = di9.add_run(f'Para ello se describen y se adecuan las principales acciones que serán sujetas a aplicar y dichas medidas de mitigación con los ajustes adecuados de acuerdo a lo siguiente.')
    descripcionCapitulo9_format = di9.paragraph_format
    descripcionCapitulo9_format.line_spacing = 1.15
    descripcionCapitulo9_format.space_after = 0
    descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo9.font.name = 'Arial'
    descripcionCapitulo9.font.size = Pt(12)
    di9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.4.1- Preparación del Sitio.')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

     ########################################################################################################################################################################
    # Capitulo 9.4.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.1.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.4.1.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.1.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.18.- Impactos generados en _________________________. Etapa de preparación del sitio')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.1.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.4.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.1.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI..4.1.2')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.1.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.19.- Impactos generados en el ___________. Etapa de preparación del sitio')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.1.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        cell = tabla9b.cell(1, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

     ########################################################################################################################################################################
    # Capitulo 9.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nIX.4.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.4.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.2.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.2.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.2.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.20.- Impactos generados en ________________, etapa de Construcción.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.2.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.2.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.2.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.2.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.21.- Impactos generados en _________________, etapa de Construcción.')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.2.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.2.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.2.3.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.2.3 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.22.- Impactos generados en los ________________________, etapa de Construcción')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.2.3 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.2.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.2.4.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.2.4 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.23.- Impactos generados en las ____________________, etapa de Construcción')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.2.4 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.4.3.- Etapa de operación')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.4.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.3.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.3.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.3.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.24.- Impactos generados en ________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.3.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.3.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.3.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.3.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.25.- Impactos generados en el _____________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.3.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.3.3 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.3.3.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.3.3 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.26.- Impactos generados en la _______________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.3.3 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.3.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.3.4.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.3.4 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.27.- Impactos generados en la carga y transporte de material, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.3.4 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.4 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.4.4.- Etapa de Abandono')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 9.4.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.4.1 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.4.1.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.4.1 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.28.- Impactos generados en la __________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.4.1 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 9.4.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 9.4.4.2 ###
    #########################
    capitulo9 = doc.add_paragraph()
    i9 = capitulo9.add_run(f'\nXI.9.4.4.2.- ')
    i9_format = capitulo9.paragraph_format
    i9_format.line_spacing = 1.15

    i9.font.name = 'Arial'
    i9.font.size = Pt(12)
    i9.font.bold = True
    capitulo9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 9.4.4.2 ###
    #########################
    tituloTabla9b = doc.add_paragraph()
    dti9b = tituloTabla9b.add_run('\nTabla 9.29.- Impactos generados en la _____________________________________, etapa de Operación')
    dti9b_format = tituloTabla9b.paragraph_format
    dti9b_format.line_spacing = 1.15
    dti9b_format.space_after = 0

    dti9b.font.name = 'Bookman Old Style'
    dti9b.font.size = Pt(12)
    tituloTabla9b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 9.4.4.2 ###
    #########################
    tabla9b = doc.add_table(rows=6, cols=6, style='Table Grid')

    for rows in tabla9b.rows:
        rows.cells[0].width = Cm(6.65)

    for cols in range(6):
        cell = tabla9b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in range(6):
            cell = tabla9b.cell(rows, cols)
            t9b = cell.paragraphs[0].add_run(' ')
            t9b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 9 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo9() # Crear el documento