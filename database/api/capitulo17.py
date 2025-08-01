"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
import os
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes, celdas, etc; en pulgadas y en centrimetros
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo17(nombre_proyecto, carpeta_base="proyectos_guardados"):
    ruta_proyecto = os.path.join(carpeta_base, nombre_proyecto.replace(" ", "_").upper())
    os.makedirs(ruta_proyecto, exist_ok=True)

    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 17
    ########################################################################################################################################################################
    """
        p = doc.add_paragraph()

        # Añadir texto con estilo personalizado
        indice = p.add_run("Índice de Contenido Capitulo XVI.")

        # Cambiar el tipo de letra y tamaño
        indice.font.name = 'Bookman Old Style'      # Tipo de letra
        indice.font.size = Pt(12)                   # Tamaño de la letra
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
        indice.bold = True

        doc.add_page_break() # Salto de página
    """

    ########################################################################################################################################################################
    # Capitulo 17
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 17 ###
    #########################
    capitulo17 = doc.add_paragraph()
    i17 = capitulo17.add_run(f'XVII.- IDENTIFICACIÓN DE LOS INSTRUMENTOS METODOLÓGICOS Y ELEMENTOS TÉCNICOS QUE SUSTENTAN LA INFORMACIÓN SEÑALADA EN LAS FRACCIONES ANTERIORES ')
    i17_format = capitulo17.paragraph_format
    i17_format.line_spacing = 1.15

    i17.font.name = 'Arial'
    i17.font.size = Pt(12)
    i17.font.bold = True
    capitulo17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 17.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 17.1 ###
    #########################
    capitulo17 = doc.add_paragraph()
    i17 = capitulo17.add_run(f'XVII.1 Presentación de la información formato Word y PDF en digital e impresa')
    i17_format = capitulo17.paragraph_format
    i17_format.line_spacing = 1.15

    i17.font.name = 'Arial'
    i17.font.size = Pt(12)
    i17.font.bold = True
    capitulo17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 17.1 ###
    #########################
    di17 = doc.add_paragraph()
    descripcionCapitulo17 = di17.add_run(
        ""
    )
    descripcionCapitulo17_format = di17.paragraph_format
    descripcionCapitulo17_format.line_spacing = 1.15
    descripcionCapitulo17.font.name = 'Arial'
    descripcionCapitulo17.font.size = Pt(12)
    di17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 17.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 17.2 ###
    #########################
    capitulo17 = doc.add_paragraph()
    i17 = capitulo17.add_run(f'XVII.2.- ANEXOS')
    i17_format = capitulo17.paragraph_format
    i17_format.line_spacing = 1.15

    i17.font.name = 'Arial'
    i17.font.size = Pt(12)
    i17.font.bold = True
    capitulo17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 17.2 ###
    #########################
    anexos = [
        'Documentos Legal del Promovente',
        'Documento Legal del Predio',
        'Mapas',
    ]

    for i, lista in enumerate(anexos):
        di17 = doc.add_paragraph()
        descripcionCapitulo17 = di17.add_run(f'ANEXO {1 + i}.- {lista}')
        descripcionCapitulo17_format = di17.paragraph_format
        descripcionCapitulo17_format.line_spacing = 1.15
        descripcionCapitulo17.font.name = 'Arial'
        descripcionCapitulo17.font.size = Pt(12)
        di17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Glosario de Terminos
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            GLOSARIO DE TÉRMINOS
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    #########################
    ### Titulo del Glosario ###
    #########################
    capitulo17 = doc.add_paragraph()
    i17 = capitulo17.add_run(f'GLOSARIO DE TÉRMINOS')
    i17_format = capitulo17.paragraph_format
    i17_format.line_spacing = 1.15

    i17.font.name = 'Arial'
    i17.font.size = Pt(12)
    i17.font.bold = True
    capitulo17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del Glosario ###
    #########################
    glosario = [
        "ACUSTF.- se refiere al área de estudio o el área de cambio de suelo donde se pretende llevar a cabo las actividades de desmonte.",
        "Acuífero: Cualquier formación geológica o conjunto de formaciones geológicas hidráulicamente conectadas entre sí, por las que circula o se almacenan aguas del subsuelo que pueden ser extraídas para su explotación, uso o aprovechamiento.",
        "Anticlinales: En las formas geológicas plegadas producidas por orogenias, sería el pliegue convexo hacia arriba. Es decir, es la ondulación de una capa de amplitud y forma variable, en la que los estratos más antiguos se encuentran en el núcleo del pliegue.",
        "Autodepuración: Es el proceso de recuperación de un cuerpo de agua después de un proceso de contaminación orgánica.",
        "Eutrofización: Proceso natural y/o antropogénico que consiste en el enriquecimiento de las aguas con nutrientes, a un ritmo tal que no puede ser compensado por la mineralización total, de manera que la descomposición del exceso de materia orgánica produce una disminución del oxígeno en las aguas profundas. Sus efectos pueden interferir de modo importante con los distintos usos que el hombre puede hacer de los recursos acuáticos (abastecimiento de agua potable, riego, recreación, etc.).",
        "Eco hidrológico: Es la relación directa que hay entre la vegetación y la accesibilidad al agua.",
        "Evacuación del suelo: Sistema individual para el tratamiento de aguas residuales.",
        "Biodiversidad: Son las variadas formas de vida que se pueden desarrollar en un ambiente natural pudiendo ser plantas, animales, microorganismos y el material genético que lo conforma.",
        "Biocenosis: Es una comunidad o conjunto de poblaciones de distintas especies, las cuales habitan en un lugar geográfico determinado y están influenciadas por factores físicos como la luz, la humedad, la temperatura, etc.",
        "Causes: Es la concavidad que sirve de piso firme a una corriente de agua natural en su curso normal.",
        "Compensación: Es un proceso mediante el cual se aplican diferentes actividades encaminadas a restaurar algunos factores alterados por modificaciones al medio natural.",
        "Contingencia: Hecho que es probable que ocurra aunque no se tiene certeza al respecto, es considerado como espontáneo o provocado.",
        "Deslizamiento: Es el movimiento del suelo, generalmente por acción de una falla o debilidad del terreno y se puede presentar de dos formas.",
        "Desmonte: El desmonte es un tipo de laboreo extraordinario que consiste en manipular mecánicamente el suelo para extraer vegetación arbustiva y herbácea.",
        "Despalme: Es la remoción de las capas superficiales del terreno natural.",
        "Difusión: Es un proceso físico irreversible en el que partículas materiales se introducen en un medio que inicialmente estaba ausente.",
        "Dispersión: Es la capacidad que tiene una población de colonizar nuevos hábitats por pequeños desplazamientos al azar de sus individuos.",
        "Erosión: Desagregación, desprendimiento y arrastre de sólidos desde la superficie terrestre por la acción del agua, viento, gravedad, hielo u otro proceso por el cual el sustrato resquebrajado y acarreado lejos de un área.",
        "Escorrentía: Es la cantidad de lluvia que excede la capacidad de infiltración en el suelo.",
        "Especies endémicas: Son aquellas especies de plantas o animales que tienen su distribución restringida a un territorio determinado. Puede ser endémica de algún estado, de alguna montaña, cueva, lago, río o manantial etc.",
        "Escenario: es un conjunto de acciones formado por la descripción de una situación futura y un camino de sucesos que permiten pasar de una situación actual a la futura. Entre los escenarios, se distinguen aquellos que se consideran posibles, los realizables, los deseables, y los tendenciales. También se pueden trabajar los llamados escenarios de contraste y los horizontes normativos.",
        "Fallas: En Geología una falla es una fractura en el terreno a lo largo de la cual hubo movimiento de uno de los lados respecto del otro. Las fallas se forman por esfuerzos tectónicos o gravitorios actuantes en la corteza. La zona de ruptura tiene una superficie generalmente bien definida denominada plano de falla, aunque puede hablarse de banda de falla cuando la fractura y la deformación asociada tienen una cierta anchura.",
        "Impactos: Son la modificación del ambiente ocasionada por la acción del hombre o de la naturaleza. Un huracán o un sismo pueden provocar impactos ambientales, sin embargo el instrumento Evaluación de Impacto Ambiental (EIA) se orienta a los impactos ambientales que eventualmente podrían ser provocados por obras o actividades que se encuentran en etapa de proyecto (impactos potenciales), o sea que no han sido iniciadas.",
        "Impactos ambientales: Es el efecto de la modificación del medio ambiente causado por actividades humanas o de la naturaleza.",
        "Impacto residual: es aquel que persistirá en el ámbito donde se haya efectuado un cambio de condición aun después de aplicar las medidas de mitigación.",
        "IVI: se refiere al Índice de Valor de Importancia del análisis de las especies de cada estrato.",
        "Malpaís: Terreno muy erosionado en el que suele faltar agua por lo que no es apropiado para cultivar.",
        "Mantos freáticos: El manto freático es el agua subterránea llamado acuífero.",
        "Matrices: Son métodos cuantitativos de evaluación que se utiliza para identificar el impacto inicial de un proyecto en un entorno natural. El sistema consiste en una matriz de información donde las columnas representan varias actividades que se hacen durante el proyecto (p. ej.: desbroce, extracción de tierras, incremento del tráfico, ruido, polvo), y en las filas se representan varios factores ambientales que son considerados (aire, agua, geología).",
        "Método GOD: Método utilizado para la determinación del riesgo de contaminación de aguas subterráneas con el fin de establecer prioridades, a través del cual se determina la vulnerabilidad intrínseca por lo que no toma en cuenta el tipo de contaminante.",
        "Mitigación: Es un conjunto de medidas que se pueden tomar para contrarrestar o minimizar los impactos negativos que pudieran tener algunas intervenciones antrópicas.",
        "N.O.M.- Normas Oficiales Mexicanas.",
        "Pronóstico Ambiental: es una técnica a través de la cual se pueden predecir las características futuras del ambiente derivadas de la ejecución de acciones antropogénicas o naturales que modifican el medio natural.",
        "Programa de Manejo Ambiental: análisis de las condiciones de un determinado ecosistema mediante la observación y la evaluación realizada a los impactos ambientales que se pudiesen generar por la ejecución de un proyecto.",
        "Programa de monitoreo: herramienta destinada a verificar el cumplimiento de las medidas planteadas en el Plan de Manejo Ambiental.",
        "Relieve: Conjunto de formas que resaltan sobre un plano o superficie. En Geografía, el relieve hace referencia a los diferentes desniveles o irregularidades que presenta la superficie terrestre, y es fundamental en los estudios del clima y en la distribución de la vegetación.",
        "Sequía: Considerada uno de los fenómenos medioambientales que más afectan al desarrollo del ser humano y de todas las formas posibles de vida, la sequía puede ser descripta como la ausencia de riego o de agua en la tierra o superficie. La sequía es usualmente causada por la falta de lluvias en una región, y mientras en algunos casos puede tener que ver con el ciclo común de los eventos (es decir, en zonas que son proclives a la sequía), en otros puede suceder de manera inesperada.",
        "Servicios de soporte: Son aquellos que mantienen los procesos de los ecosistemas y permiten la provisión del resto de los servicios.",
        "Servicios de provisión: Son recursos tangibles y finitos que se contabilizan y consumen.",
        "Servicios de regulación: Son los que mantienen los procesos y funciones naturales de los ecosistemas.",
        "Sismicidad: Se denomina sismicidad al análisis del número de sismos que se suceden en una región geográfica determinada. Tal estudio registra en un mapa a los diversos epicentros existentes, además de tomar en cuenta la frecuencia con que se suceden estos fenómenos.",
        "Valoración: Conjunto de técnicas y métodos que permiten medir las expectativas de beneficios y costos derivadas de las acciones. Término utilizado para asignar valor a procesos indirectos tales como las cuencas hidrográficas y el abastecimiento de agua; los bosques en el secuestro de carbono y el control de la erosión; conservación de los ecosistemas y el mantenimiento de material genético.",
        "Vulnerabilidad: Es la incapacidad de resistencia cuando se presenta un fenómeno amenazante, o la incapacidad para reponerse después de que ha ocurrido un desastre."
    ]

    for i, lista in enumerate(glosario):
        di17 = doc.add_paragraph()
        descripcionCapitulo17 = di17.add_run(f'{lista}')
        descripcionCapitulo17_format = di17.paragraph_format
        descripcionCapitulo17_format.line_spacing = 1.15
        descripcionCapitulo17.font.name = 'Arial'
        descripcionCapitulo17.font.size = Pt(12)
        di17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Bibliografía
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            BIBLIOGRAFÍA
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    #########################
    ### Titulo de la Bibliografia ###
    #########################
    capitulo17 = doc.add_paragraph()
    i17 = capitulo17.add_run(f'BIBLIOGRAFÍA')
    i17_format = capitulo17.paragraph_format
    i17_format.line_spacing = 1.15

    i17.font.name = 'Arial'
    i17.font.size = Pt(12)
    i17.font.bold = True
    capitulo17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion de la Bibliografia ###
    #########################

    referencias = [
        "Boletín de la Sociedad Geológica Mexicana 2006",
        "Casas Andreu, Valenzuela López G.G., y Ramírez Bautista, A. (1990) Como hacer una colección de anfibios y reptiles. Cuadernos del Instituto de Biología. Serie 10. Universidad Nacional Autónoma de México.",
        "Drews, C. (2003). Conceptos y panorama del rescate de fauna en el Neotrópico.",
        "Dirección General de Fomento Editorial Benemérita Universidad Autónoma de Puebla Merriam, C. H. 1898.",
        "En Manejo de Fauna Silvestre en Amazonía y Latinoamérica. Selección de Trabajos del V Congreso Internacional, ed. R. Polanco‐Ochoa. CITES, Fundación Natura. Bogotá, Colombia. P. 351 ‐ 356.Eliosa‐Leon, H.R. y Castillo Salazar, A. (2006). Recolecta de Anfibios. En Manual de Métodos de Colecta de Plantas y Animales",
        "Dirección General de Fomento Editorial Benemérita Universidad Autónoma de Puebla Merriam, C. H. 1898.",
        "Plan estatal Cuenca de Burgos para el estado de Coahuila",
        "Plan estatal de desarrollo para el estado de Coahuila 2013-2018",
        "Identificación de material vegetativo Dr. José Ángel Villarreal Escamilla.",
        "Libro de Botánica de las especies de Dr. José Ángel Villareal Escamilla",
        "Ley General del Equilibrio Ecológico y la Protección al Ambiente (LGEEPA).",
        "Ley General de Desarrollo Forestal Sustentable y su Reglamento.",
        "Ley General de Vida Silvestre",
        "Ley del equilibrio Ecológico y la Protección al Ambiente del Estado de Coahuila",
        "Libro de Botánica de las especies de Dr. José Ángel Villareal Escamilla",
        "Informe 2011. Instituto Nacional de Ecología, Instituto Mexicano de Tecnología del Agua.",
        "Manejo de Fauna Silvestre en Amazonía y Latinoamérica. Selección de Trabajos del V Congreso Internacional, ed. R. Polanco‐Ochoa. CITES, Fundación Natura. Bogotá, Colombia. P. 351 ‐ 356.Eliosa‐Leon, H.R. y Castillo Salazar, A. (2006). Recolecta de Anfibios. En Manual de Métodos de Colecta de Plantas y Animales.",
        "Manual de Técnicas para Estudios de la Fauna.- Dra Sonia Gallina Tessaro. Instituto Nacional de Ecología; Dr. Carlos A. López González Universidad Autónoma de Querétaro; Salvador Mandujano Rodríguez. 2011",
        "Norma Oficial Mexicana NOM-059-SEMARNAT-2010, Protección ambiental-Especies nativas de México de flora y fauna silvestres-Categorías de riesgo y especificaciones para su inclusión, exclusión o cambio-Lista de especies en riesgo, publicada en el Diario Oficial de la Federación el 30 de diciembre de 2010.",
        "NOM-011- CONAGUA 2000 Conservación del Recurso Agua - Que establece las especificaciones y el método para determinar la disponibilidad media anual de las aguas nacionales”.",
        "Pedro Linares Llamas; Universidad Pontificia Comillas de Madrid. Economía y Medio Ambiente; Herramientas de valoración Ambiental.",
        "Programa de Ordenamiento Ecológico General del Territorio",
        "Normas Oficiales Mexicanas en materia de seguridad",
        "Plan Nacional de Desarrollo (2013-2018)",
        "Plan Estatal de Desarrollo (2011-2017).",
        "REGLAS de Operación del Programa Nacional Forestal 2015.",
        "Reglamento de la Ley General del Equilibrio Ecológico y la Protección al Ambiente en Materia de Evaluación del Impacto Ambiental.",
        "Regulo León Arteta.- Grados de Erosión o severidad y capacidad del uso del suelo. Marzo 2007.",
        "Síntesis Geográfica del Estado de Coahuila.",
        "TNC The Nature Conservancy.",
        "Convenios entre Canadá, Estados Unidos y México para la protección de las aves migratorias y los mamíferos cinegéticos",
        "Convención RAMSAR para la protección de humedales de importancia internacional.",
        "Declaración de Rio sobre el medio ambiente y el desarrollo",
        "Secretaría de Agricultura, Ganadería, Desarrollo Rural, Pesca y Alimentación (SAGARPA). http://www.gob.mx/sagarpa",
        "Instituto Nacional de Investigaciones Forestales y Agropecuarias (INIFAP). http://www.inifap.gob.mx/SitePages/Inicio.aspx",
        "Comisión Nacional de Áreas Naturales Protegidas (CONANP). http://www.conanp.gob.mx/",
        "Comisión Nacional para el Conocimiento y Uso de la Biodiversidad (CONABIO). http://www.conabio.gob.mx/",
        "Secretaría De Medio Ambiente Y Recursos Naturales (SEMARNAT). http://www.gob.mx/semarnat",
        "Comisión Nacional del Agua (CONAGUA). http://www.conagua.gob.mx/",
        "Cartografía INEGI 2012",
        "Instituto Nacional de Estadística y Geografía (INEGI). Cartografía, Datos Hidrológicos de Aguas Superficiales, Informe 2011."
    ]

    for i, lista in enumerate(referencias):
        di17 = doc.add_paragraph()
        descripcionCapitulo17 = di17.add_run(f'{lista}')
        descripcionCapitulo17_format = di17.paragraph_format
        descripcionCapitulo17_format.line_spacing = 1.15
        descripcionCapitulo17.font.name = 'Arial'
        descripcionCapitulo17.font.size = Pt(12)
        di17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """

    nombre_archivo = f"CAPITULO 17 EXTRACCION_ DE MATERIAL PETREO DTU {nombre_proyecto.replace(' ', '_').upper()}.docx"
    ruta_completa = os.path.join(ruta_proyecto, nombre_archivo)
    doc.save(ruta_completa)
    print(f"✅ Guardado: {ruta_completa}")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""

