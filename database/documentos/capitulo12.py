"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo12():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 12
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XI.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 12
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 12
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 12 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'XII.- MEDIDAS DE PREVENCIÓN Y MITIGACIÓN DE IMPACTOS SOBRE LOS RECURSOS FORESTALES, LA FLORA Y FAUNA SILVESTRE, APLICABLES DURANTE LAS DISTINTAS ETAPAS DEL DESARROLLO DEL CAMBIO DE USO DE SUELO.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En el presente estudio de acuerdo a su proyección, durante su proceso de evaluación para su desarrollo se han detectado posibles impactos ambientales adversos que afectarán en tres etapas (Preparación del Sitio, Construcción y Operación), principales consideradas para la implementación del proyecto, con el propósito de revertir dichos impactos se han establecido medidas de prevención y/o mitigación y de ser posible algún proceso para la restauración del sitio. Estas medidas deberán de tomarse como una responsabilidad no solo como complemento del estudio, en el sentido de que de su aplicación dependerá la condición que pueda acoger el comportamiento de la biodiversidad de este tipo de ecosistema que se verá alterado en la superficie propuesta para el cambio de uso de suelo y notificación del giro natural que ostenta el Sistema Ambiental.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.- Clasificación de Medidas.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.- Medidas Preventivas:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.1.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Son aquellas acciones que se pueden anticipar a los efectos o modificaciones que se pudieran generar durante el desarrollo del proyecto contemplando cada una de sus etapas, plasmando su forma de aplicación a fin de minimizar los impactos ambientales de carácter Moderado, Severo y Critico para cada etapa del proyecto y en cada una de las acciones donde se estarán manifestando, para ello se efectúa la adecuación pertinente incluyendo los objetivos principales de dichas medidas de acuerdo a la siguiente tabla: ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1- Preparación del Sitio.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1.1- PS-04 Desmonte (Remoción de la vegetación)')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.1.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.1.- Impactos generados en la remoción de vegetación. etapa de preparación del sitio')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.1.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.1.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1.1- PS-05.- Despalme')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.1.1.2 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.1.- Impactos generados en la remoción de vegetación. etapa de preparación del sitio')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.1.1.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.- Etapa de Construcción')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.1.- CO-01')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.1.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.3.- Impactos generados en ________________, etapa de Construcción.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.2.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.2.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.2.2.- CO-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.2.2 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.4.- Impactos generados en Rampas de acceso, etapa de Construcción')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.2.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.- Etapa de Construcción')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.1.- CO-01')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.3.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.3.- Impactos generados en ________________, etapa de Construcción.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.3.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.2.2.- CO-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.3.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.4.- Impactos generados en Rampas de acceso, etapa de Construcción')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.3.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.4.- Etapa de Abandono')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.4.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.4.1.- AB-01.- ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Impactos generados en la clausura del sitio, etapa de Operación')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.4.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.1.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.1.4.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.4.2.- AB-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.1.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Impactos generados en la reforestación y restauración del área, etapa de Operación')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.1.4.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.2.- Medidas de Mitigación:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.1.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Es el conjunto de acciones que se implementan una vez que se identifica el impacto y la magnitud del mismo, con la finalidad de minimizar en lo posible los efectos de dicho impacto sobre todo aquellos que prevalecerán aun con la aplicación de las medidas preventivas, sobre todo en aquellas acciones del proyecto que son inevitables y de carácter severo, Moderado, Severos y Críticos para algunos factores ambientales y sus componentes que lo integran, ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Objetivos de las medidas.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Primer ítem
    di12 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo12 = di12.add_run("Reducir los impactos a través de la limitación de su magnitud")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo ítem
    di12 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo12 = di12.add_run("Rectificar el impacto a través de la reparación, rehabilitación o restauración del componente ambiental afectado.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Tercer ítem
    di12 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo12 = di12.add_run(
        "Minimizar o eliminar el impacto a través del tiempo con la implementación de actividades resultado de los análisis aplicados mediante la organización establecida para la conservación y mantenimiento durante la vida del proyecto."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Para ello se describen y se adecuan las principales acciones que serán sujetas a aplicar y dichas medidas de mitigación con los ajustes adecuados de acuerdo a lo siguiente.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1- Preparación del Sitio.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.2.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1.1- PS-04 Desmonte (Remoción de la vegetación)')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.1.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.1.- Impactos generados en la remoción de vegetación. etapa de preparación del sitio')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.1.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.1.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.1.1- PS-05.- Despalme')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.1.2 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.1.- Impactos generados en la remoción de vegetación. etapa de preparación del sitio')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.1.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.- Etapa de Construcción')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.1.- CO-01')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.2.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.3.- Impactos generados en ________________, etapa de Construcción.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.2.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.2.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.2.2.- CO-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.2.2 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.4.- Impactos generados en Rampas de acceso, etapa de Construcción')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.2.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.- Etapa de Construcción')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.2.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.2.1.- CO-01')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.3.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.3.- Impactos generados en ________________, etapa de Construcción.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.3.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.1.1.2.2.- CO-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.3.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.4.- Impactos generados en Rampas de acceso, etapa de Construcción')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.3.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.1.4.- Etapa de Abandono')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.1.2.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.4.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'X.1.1.4.1.- AB-01.- ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Impactos generados en la clausura del sitio, etapa de Operación')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.4.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.1.2.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.1.2.4.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.1.2.4.2.- AB-02.- ____________________')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.1.2.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Impactos generados en la reforestación y restauración del área, etapa de Operación')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.1.2.4.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.- Verificación de las medidas de prevención y mitigación planteadas.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.1.- Etapa de aplicación: Preparación del Sitio.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.1.1.- Factor Ambiental Afectado: Flora')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Cobertura y Abundancia.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    #########################
    ### Descripcion del capitulo 12.2.1.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Cobertura y Abundancia.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.1.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medida de Mitigación del factor Flora.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.1.1 ###
    #########################
    filas = 6  # 1 encabezado + 5 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medidas",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Datos de la tabla
    datos_tabla = [
        [
            "Realizar registro del rescate de especies de lento crecimiento o especiales para preservar la especie y biodiversidad en la zona previo al desmonte georreferenciando ubicación.",
            "Revisión del área del proyecto, verificación de especies densidad, georreferenciación registro de aquellas que se logren visualizar rescatar y reubicar en los sitios de reserva en otras etapas del proyecto.",
            "Ubicación de especies sujetas a rescate en las áreas de extracción en las diferentes etapas del proyecto"
        ],
        [
            "Delimitar el área consideradas de conservación y reubicación o de franjas de amortiguamiento con las áreas aledañas y protegerlas.",
            "Verificar que se mantengan siempre en buen estado las especies rescatadas.",
            "Estrés de las plantas manifestación de daños."
        ],
        [
            "Previo a la remoción de la vegetación delimitar el área sujeta al ACUSTF autorizado para que no afectar áreas no autorizadas de acuerdo al plano del proyecto.",
            "Verificar en campo que se apeguen a las áreas estipuladas para remoción de vegetación.",
            "Límites fuera del área autorizada."
        ],
        [
            "Determinar los sitios donde se acumulará el residuo del desmonte seleccionando sitios estratégicos ya sea para uso posterior o disposición final.",
            "Vigilar que los residuos se tengan dentro del sitio del proyecto y posterior reúso o destino final, así como verificar que no se pongan en sitios no contemplados como cause de arroyos, escorrentías o reservas.",
            "Dispersión de residuos de vegetación"
        ],
        [
            "Prohibir e impedir la extracción furtiva de especies vegetales de interés del sitio del proyecto o áreas aledañas por el personal que labore en el proyecto, capacitación sobre la importancia de las especies.",
            "Mantener vigilancia continua para evitar el furtivismo a través del monitoreo que se establezca. Verificar capacitación aplicada para protección de especies de flora.",
            "Detección de incidentes de saqueo."
        ]
    ]

    # Llenar las celdas
    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            t12b = cell.paragraphs[0].add_run(datos_tabla[fila-1][col])
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.2.- Etapa de Preparación del Sitio y Construcción.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.2.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.2.1.- Factor Ambiental Afectado: Atmósfera.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Calidad del Aire y Calidad Sonora.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    #########################
    ### Descripcion del capitulo 12.2.2.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Calidad del Aire y Calidad Sonora.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #########################
    ### Título de la tabla del capítulo 12.2.2.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación planteadas en Preparación del Sitio, Construcción y operación con afectación a la Atmósfera.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.2.1 ###
    #########################
    filas = 7  # 1 fila encabezado + 6 filas datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medidas",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos de la tabla
    datos_tabla = [
        [
            "Aplicar el mantenimiento continuo a los equipos para reducir emisiones de contaminantes a la atmosfera.",
            "Vigilancia constante con el fin de verificar el cumplimiento de los equipos que laboren en el área del proyecto.",
            "Opacidad del aire"
        ],
        [
            "Humectar el sustrato que se rescate del despalme y se fije en algún sitio para su uso posterior así mismo mantenerse humectados.",
            "Efectuar monitoreos constantes a fin de verificar el cumplimiento de que los materiales y sitios de operación se mantengan húmedos.",
            "Cuando se observe acumulamiento de polvos en la vegetación aledaña al sitio."
        ],
        [
            "Colocar letreros de identificación de velocidades permitidas no mayores a 20 Km/hora. Determinar horarios para aplicación de humedad para reducir desprendimiento de partículas de polvo y dispersión a la atmósfera.",
            "Monitoreo continuo e instalación de al menos 2 letreros en cada una de los trayectos que se utilicen e instalar un aviso de horario de aplicación de líquidos para humectar los suelos.",
            "Mantener visible los avisos para cumplimiento de medida."
        ],
        [
            "Efectuar los mantenimientos preventivos de la maquinaria cada 200 hr y correctivo cuando sea requerido fuera del área del proyecto.",
            "Seguimiento periódico a bitácoras de revisión mantenimiento de vehículos y maquinaria.",
            "Opacidad del aire con presencia de contaminantes."
        ],
        [
            "Efectuar mantenimiento a los vehículos de carga utilizados en los procesos y cumplan con las normas vigentes respecto a las emisiones permitidas de combustión y niveles de ruido.",
            "Verificar bitácoras en apego a la NOM-045-SEMARNAT-2017 que establece los límites máximos permisibles de emisión de gases contaminantes provenientes del escape de los vehículos automotores en circulación. Así también los sistemas de escape y silenciadores se mantendrán acordes a los niveles máximos permisibles de emisión de ruido proveniente del escape de vehículos automotor que están establecidos en la norma NOM-080-SEMARNAT-1994.",
            "Opacidad del aire con presencia de contaminantes e incremento de nivel sonoro."
        ],
        [
            "Aplicar durante el día, al menos 2 riegos en los frentes de avance de las actividades tanto de preparación del sitio como de construcción para reducir la volatilidad de partículas al ambiente.",
            "Verificar el registro de agua aplicado en los procesos de preparación del sitio y construcción, registrando el lugar de procedencia del agua utilizada y el tipo, así como la cantidad diaria aplicada.",
            "Opacidad del aire con presencia de partículas en el ambiente y observación de polvo en la vegetación aledaña."
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            t12b = cell.paragraphs[0].add_run(datos_tabla[fila-1][col])
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.3.- Etapa de aplicación: Construcción y Operación.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.3.1.- Factor Ambiental Afectado: Hidrología')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Escurrimientos y Recarga de mantos.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.2.3.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Escurrimientos y Recarga de mantos.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.3.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación en la Construcción y Operación con afectación a la Hidrología.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.3.1 ###
    #########################
    filas = 4  # 1 encabezado + 3 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medida",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)

    # Datos de la tabla
    datos_tabla = [
        [
            "No afectar cuerpos de agua o afloramiento naturales en el área del proyecto y vialidades de acceso y utilizar solo el área autorizada.",
            "Monitorea y respetar las áreas destinadas para reserva sin alterar y verificara la no afectación de causes.",
            "Desvío de cuerpos de agua e invasión de áreas de conservación."
        ],
        [
            "Evitarse el vertido de aguas orgánicas en áreas aledañas o cauces (utilizar letrinas secas) 2 por cada 15 trabajadores.",
            "Verificar con respecto a documentación el uso del tipo de letrinas 2 en el área del proyecto y verificar el manejo de residuos.",
            "Mal funcionamiento o indicios de uso al aire libre."
        ],
        [
            "Mantener inalteradas las áreas dispuestas como franjas de amortiguamiento y obras de conservación de suelo y agua para retención de suelo e infiltración de agua y procesos biológicos de flora y fauna",
            "Verificar la condición conservación del área para los fines de retención de agua, nicho y hábitat de flora y fauna, manteniéndolas inalterables.",
            "Indicios de alteración de cualquiera de los factores para lo cual se contempló la conservación."
        ],
        [
            "",
            "",
            ""
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            texto = datos_tabla[fila-1][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.4.- Etapa de aplicación: Construcción y Operación.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.4.1.- Factor Ambiental Afectado: Topografía')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Relieve
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.2.4.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Relieve')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación en la construcción y operación, con afectación a la Topografía.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.4.1 ###
    #########################
    filas = 4  # 1 encabezado + 3 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medidas",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)

    # Datos de la tabla
    datos_tabla = [
        [
            "Efectuar el desmonte y despalme solo en el área autorizada.",
            "Verificar que solo el área autorizada sea modificada.",
            "Rebase de límites permitidos"
        ],
        [
            "Establecer solo las vialidades necesarias con aplicación de obras de drenajes",
            "Verificar obras de drenajes como drenes y canales de desagüe en su trayectoria.",
            "Formación de canalillos por escurrimientos."
        ],
        [
            "Modificar el relieve solo en el área autorizada y establecer bancos viables de restaurar.",
            "Con la modificación del relieve verificar la formación de terraplenes que puedan ser utilizados para reforestación.",
            "Observar apilamientos fuera del área autorizada."
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            texto = datos_tabla[fila-1][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.5 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.5.- Etapa de aplicación: Todas las Etapas.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.5.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.2.5.1- Factor Ambiental Afectado: Suelo')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Propiedades, Erodabilidad.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.2.5.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Propiedades, Erodabilidad.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.5.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación en todas las etapas del proyecto con afectación al Suelo.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.5.1 ###
    #########################
    filas = 4  # 1 encabezado + 3 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medida",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)

    # Datos de la tabla
    datos_tabla = [
        [
            "Mantener humectado el área de trabajo donde se genere volatilidad de partículas a la atmósfera al menos dos riegos diarios.",
            "Verificar que se aplique la humectación al menos dos veces al día en las áreas de operación",
            "Dispersión del suelo en el predio."
        ],
        [
            "Señalización y protección de áreas sujetas de resguardo de suelo orgánico registro de humectación.",
            "Verificación de letreros en el sitio del proyecto alusivos a áreas especiales para almacén de suelo orgánico y bitácoras de aplicación de humedad.",
            "Dispersión del suelo por acciones erosivas hídricas y eólicas."
        ],
        [
            "Evitar la Afectación fuera de lo autorizado y áreas de reserva establecidas en el proyecto estableciendo señalización.",
            "Verificar la señalización de sitios para resguardo de sustrato orgánico en los sitios designados.",
            "Dispersión de acumulamientos de sustrato."
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            texto = datos_tabla[fila-1][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.5.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.2.5.2.- Factor Ambiental Afectado: Fauna')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Abundancia
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.2.5.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Abundancia')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.5.2 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación en todas las Etapas del Proyecto con afectación a la Fauna.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.5.2 ###
    #########################
    filas = 6  # 1 encabezado + 5 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medidas",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)

    # Datos de la tabla
    datos_tabla = [
        [
            "Aplicar programa de rescate y reubicación de especies de lenta movilidad",
            "Registro de especies rescatadas y reubicadas con datos de georreferenciación.",
            "Detección de individuos"
        ],
        [
            "Se efectuarán acciones de ahuyentamiento de fauna propiciando y facilitando su libre desplazamiento durante el desmonte y despalme.",
            "Verificación de registro de acciones y fauna avistada en el proceso.",
            "Detección de Individuos."
        ],
        [
            "Prohibir e impedir la cacería o extracción de fauna que se localice en el sitio del proyecto máxime al no haber tenido avistamiento",
            "Efectuar vigilancia continua para evitar acciones en deterioro de la fauna caza y extracción por trabajadores o visitantes",
            "Indicios de extracción."
        ],
        [
            "Capacitación en el manejo de fauna documentándose.",
            "Verificación documental de capacitación de personal de operación y visitantes.",
            "Indicio de mal manejo de individuos rescatados."
        ],
        [
            "Evitar la afectación del hábitat de la fauna en los sitios considerados de reserva.",
            "Vigilar y verificar que se respeten los espacios considerados como reserva para generación de hábitat posible para las especies que se ubiquen en el área.",
            "Indicios de alteración en área de conservación."
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            texto = datos_tabla[fila-1][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.size = Pt(12)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.2.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.2.5.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.2.5.3.- Factor Ambiental Afectado: Paisaje')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Componentes Ambientales Afectados: Armonía y Calidad Paisajística.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.2.5.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Componentes Ambientales Afectados: Armonía y Calidad Paisajística.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.2.5.3 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Medidas de mitigación y prevención en todas las Etapas del Proyecto, con afectación al Paisaje.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.2.5.3 ###
    #########################
    filas = 4  # 1 encabezado + 3 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    encabezados = [
        "Medidas",
        "Verificación de Medidas",
        "Umbral de Alerta"
    ]

    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        cell_background_color(cell, '0070C0')
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t12b.font.bold = True
        t12b.font.color.rgb = RGBColor(255, 255, 255)

    # Datos de la tabla
    datos_tabla = [
        [
            "Evitar la diseminación de residuos de tipo doméstico, aguas sanitarias, residuos de manejo especial y/o peligroso.",
            "Supervisar a las empresas que dan servicio de mantenimiento de equipos que cumplan con los procesos para evitar que se dejen residuos abandonados que den mal aspecto al paisaje.",
            "Detección de Residuos."
        ],
        [
            "Recolección y disposición de residuos generados en contenedores y uso de letrinas portátiles.",
            "Verificar que no se dispersen residuos en el sitio del proyecto y la funcionalidad de contenedores y letrinas",
            "Indicios de dispersión de residuos."
        ],
        [
            "Recomendar a las empresas prestadoras de servicios apego a la normatividad en el manejo de residuos generando registro de residuos, manejo y disposición final.",
            "Verificar bitácoras de la recolección temporal de residuos como el retiro hasta su disposición final se cumpla y no se dejen abandonados residuos que afecte el paisaje.",
            "Indicios de residuos fuera de contenedores."
        ]
    ]

    for fila in range(1, filas):
        for col in range(columnas):
            cell = tabla12b.cell(fila, col)
            texto = datos_tabla[fila-1][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            t12b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.3.- Impactos Residuales.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Un impacto residual es aquel que persistirá en el ámbito donde se haya efectuado un cambio de condición aun después de aplicar las medidas de mitigación.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Como resultado de la evaluación se considera que en un escenario en el cual se llevan a cabo este tipo de acciones aun cuando se apliquen todas las medidas de prevención, de mitigación y de compensación que se plantean en el capítulo correspondiente, se han identificado impactos residuales que aun con la aplicación de las medidas estos impactos residuales persistirán.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('De acuerdo a ello se han identificado los siguientes:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    impactos_residuales = [
        'El impacto de carácter erosivo por acciones eólicas e hídricas y la modificación estructural del suelo, así como el uso continuo de las vialidades dentro del área autorizada, aun con su mantenimiento prevalecerá de forma moderada aun con las acciones de mitigación que se apliquen.',
        'El impacto acumulado sobre la flora se verá reflejado en la abundancia aun cuando es moderado por la cobertura, densidad y tipo de vegetación se mantiene aún con las medidas de mitigación por su eliminación en el desmonte y despalme.',
        'El impacto residual sobre el paisaje aun cuando no es de alto valor económico si tiene valor ecológico al ser el característico de la zona y su afectación persistirá aun con las medidas aplicables por la eliminación de su cubierta vegetal y por las oquedades realizadas en la extracción del material pétreo.',
        'El impacto sobre la fauna se mantendrá durante la vida útil del proyecto al no tener las condiciones adecuadas para su desarrollo y sobrevivencia por la movilidad de equipos y seres humanos en la operación para extraer material patero, objeto del proyecto.',
        'Tanto para los efectos de los impactos flora, suelo, agua y fauna entre otros no de menor importancia se contempla tener un área exclusiva sin alterar para la conservación de la biodiversidad existente.'
    ]

    for impacto in impactos_residuales:
        di12 = doc.add_paragraph(style='List Bullet')  # Esto genera viñetas automáticas
        descripcionCapitulo12 = di12.add_run(impacto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.4.- Impactos Ambientales Acumulativos:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.4 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Como se define en el reglamento de la LGEEPA En Materia de Impacto Ambiental, un impacto ambiental acumulativo es el efecto en el ambiente que resulta del incremento de los impactos de acciones particulares ocasionadas por la interacción con otros que se efectuaron con anterioridad y que están ocurriendo en el presente.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("De acuerdo a la metodología empleada para la valoración de la importancia de los impactos ambientales desarrollada, se consideraron diversos impactos ambientales, aun cuando el área no tiene grado extremo de impacto por su uso al que se encuentra el área sujeta de estudio, siendo estos sobre los cuales se pudo llevar a cabo un análisis de interacción acumulativa.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Por tal motivo se enumeran los impactos ambientales significativos que presentan la capacidad de acumularse con otros impactos.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    impactos_acumulativos = [
        "La acumulación del mayor impacto es por la modificación del área que se adhiere a las actividades que actualmente se realizan en el sistema ambiental, como son el demás aprovechamiento de material pétreo de la zona, el relleno sanitario de Ramos Arizpe.",
        "La acumulación del impacto para la operación del proyecto al tener mayor cantidad de personas y equipos, ocasionara el ahuyentamiento de la poca fauna del lugar y evitar el regreso a su hábitat natural, al menos en el mediano plazo.",
        "El impacto acumulativo de la eliminación de la vegetación dará como resultado mayor cantidad de área desprovistas de vegetación con lo cual se incrementará la erosión eólica e hídrica y los incrementos de temperatura por no existir el amortiguamiento natural.",
        "Otro impacto acumulativo con referencia a la vegetación lo será la reducción de densidades de poblaciones de especies que existan en el área del proyecto.",
        "La relación de impacto acumulativo en lo referente al suelo en el sitio del proyecto es la pérdida de suelo orgánico que puede darle sustento a la vegetación existente aun cuando es baja su densidad y cobertura pero que sirve como regulador de impactos atmosféricos.",
        "El impacto acumulativo con respecto a la operación del proyecto modificará el paisaje natural aun cuando no tiene una proyección sustentable para algún otro tipo de actividad de la región su modificación será notoria e irreversible, que se acumulará con los demás ___________________________________________.",
        "Para disminuir los impactos causados en la vegetación y fauna del área de cambio de uso de suelo se proponte un programa de rescate de flora y fauna silvestre (capitulo 10)."
    ]

    for texto in range(len(impactos_acumulativos)):
        di12 = doc.add_paragraph()
        descripcionCapitulo12 = di12.add_run(f"{texto + 1}.- {impactos_acumulativos[texto]}")
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.5 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.5.- Impactos Ambientales Sinérgicos.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.5 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Un impacto ambiental sinérgico está definido por el Reglamento de la LGEEPA EN Materia de Impacto Ambiental como aquel que se produce cuando el efecto conjunto de la presencia simultánea de varias acciones supone una incidencia ambiental mayor que la suma de las incidencias individuales contempladas en forma aislada.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Es decir, la propiedad de interactuar con otros impactos ambientales distintos generando un efecto mayor comparados con el resultado de los impactos cuando fueran analizados en forma individual o independiente, bajo esta consideración las sinergias encontradas fueron las siguientes.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Los diferentes tipos de interacción sinérgica identificadas son las siguientes:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sinergias = [
        "La sinergia entre los impactos del desmonte y la fauna silvestre aun cuando es muy reducida la presencia, la reducción de la fauna puede tener una disminución de la dispersión de semillas que como consecuencia se tendría poca abundancia de flora.",
        "La sinergia del impacto del desmonte y despalme con el paisaje es que una vez eliminada la vegetación y al retirar los residuos pueda aparecer otro impacto visual no contemplado o se sumara a los impactos por las _____________________.",
        "Así mismo dentro de esta misma sinergia puede ser que al retirar los residuos derivados de las acciones del desmonte y despalme se puedan encontrar otros impactos que hayan permanecido ocultos.",
        "Otra sinergia de impactos derivado del acumulamiento de materiales y uso de combustibles puede provocar derramamientos accidentales y por ende contaminación en el sitio.",
        "La modificación de la topografía en su relieve se sumará a la modificación ya existente por la operación actual en la _____________________________.",
        "La modificación al paisaje al menos en el área del proyecto se modificará su visibilidad y fragilidad al extraer al _______________ que se sumará a los aprovechamientos en el sistema ambiental.",
        "Los impactos de la topografía, suelo, fauna y paisaje serán sinérgicos al encontrarse dentro del sistema ambiental ___________________________________ cercana al área y la operación del _________________ teniendo lo siguiente:"
    ]

    for i, texto in enumerate(sinergias, 1):
        di12 = doc.add_paragraph()
        descripcionCapitulo12 = di12.add_run(f'{i}.- {texto}')
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Los Factores Ambientales donde se genera impacto sinérgico y acumulativo:")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Atmósfera: Calidad del Aire y Calidad Sonora. Se presentará volatilidad de partículas de polvo y smog a la atmósfera por el uso de maquinaria y movimientos de sustratos, siendo desde poco moderado hasta crítico ocurriendo este, en la etapa de desmonte y despalme. Este factor será mitigable y será temporal y el grado de será bajo.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Hidrología: Componente ambiental (Escurrimiento). Se observa una valoración de poco significativa a moderados en virtud de la afectación de los escurrimientos al eliminar la vegetación incrementando la velocidad del flujo reduciendo la posibilidad de infiltración y ocasionando arrastre de sustrato. De acuerdo a los recorridos dentro del proyecto no se encuentran cuerpos de agua ni escurrimientos permanentes solo se manifiesta escurrimientos intermitentes, en cuanto a la calidad y cantidad de agua estos serán mitigables con actividades monitoreo ambiental no derramando combustibles, aceites o basura urbana dentro del área y sobre todo del escurrimiento, infiltrante el agua a los mantos freáticos teniendo un grado mínimo o bajo a moderado de afectación.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Suelo: Componente ambiental (Erodabilidad, Erosión). Ante la eliminación de vegetación se tendrá exposición del sustrato generando erosión del tipo hídrica y eólica desde el punto de vista crítico por la modificación a los perfiles, la topografía y el paisaje. En cuanto a este componente la afectación será con un grado crítico, y no se puede ser compensable o mitigable en el área, el daño es irreversible, sin embargo, con el la erosión de las áreas de cambio de uso de suelo se puede mitigar riegos de mitigación, que se proponen para recuperar esa pérdida de suelo por acción del agua y el aire.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Flora Silvestre: Componente ambiental (Densidad y Cobertura). Como parte del proceso al eliminar la vegetación se tendrá pérdida de densidad de individuos y especies afectando la cobertura generando un impacto crítico para efecto de conservar la Biodiversidad se contempla efectuar un programa de rescate de las especies consideradas como de lento desarrollo y difícil regeneración y/o aquellas que se adapten a este proceso. El grado de afectación es crítica, con la eliminación de la vegetación, sin embargo, se realiza un rescate para mitigar o compensar la pérdida.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before =

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Fauna Silvestre: Componente ambiental (Abundancia y Hábitat). Ante la presencia de maquinaria, equipo y seres humanos se tendrá un impacto crítico con la presencia de maquinaria, el uso de explosivos de bajo impacto, en referencia a que las especies serán ahuyentadas de su entorno, al igual que en la Flora se podrá aplicar un programa de rescate para aquellas especies de lenta movilidad en el área del proyecto. El impacto es considerado moderado a crítico, ya que la fauna será ahuyentada y no eliminada, para mitigar se realizar rescate de individuos de baja movilidad.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Paisaje: Componente Ambiental (Armonía y Calidad Paisajística). Resultado de esta acción del proyecto se tendrá una modificación que generará un impacto desde moderado hasta críticos al modificar su entorno ante la eliminación de la vegetación y ocasionar oquedades por las características propias del proyecto, en el establecimiento del proyecto se modificará el paisaje considerando el impacto de moderado a critico por la infraestructura de los bancos de aprovechamiento, rampas de acceso, lo cual se podrá mitigar realizando reforestación al abandono del proyecto, asi como no dejar residuos domésticos en el predio.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.6 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.6.- Análisis del Escenario sin Proyecto.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.6 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('El área sujeta de estudio como se menciona en el uso actual del terreno está clasificado como _______________________________, sin embargo, existen especies de alto valor ecológico como las cactáceas, especies de Asparagaceae como las principales que deberán ser rescatadas para mantener el germoplasma y ocasionar su reproducción que mantenga la biodiversidad en el área, de acuerdo a los recorridos realizados durante los muestreos de vegetación y fauna.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('De acuerdo al análisis efectuado en el área de estudio se determinó la condición que deberá aplicarse para mantener la biodiversidad consistiendo en mantener una superficie sin alteración y que sirva para amortiguar los efectos de los impactos en la atmosfera, agua, suelo topografía, flora y fauna y paisaje, aunque en algunas actividades serán más notorios se podrán mitigar y difícilmente volverlos al estado de origen. ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.6.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.6.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.6.1.- Recurso Flora')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.6.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('La vegetación que se encuentra dentro del área del proyecto, de acuerdo a la carta de uso de suelo y vegetación del INEGI, encontramos _________________________________________, por lo tanto, el muestreo y análisis de datos se realizó por tipo de vegetación teniendo lo siguientes resultados:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nTipo de Vegetación')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        Todo esto tiene que estar en una base de datos
    """

    for lista in range(5):
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(f'Actividad {lista + 1}')
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.6.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.6.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.6.2.- Recurso Fauna.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.6.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Analizando los resultados obtenidos en el área de estudio para el cambio de uso de suelo, se tiene lo siguiente:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo faunístico de las aves en el ACUSTF se avistaron ___________________ de las cuales la más abundante fue ___________________ con ___________________ individuos avistados, las ___________________ menos abundantes fueron ___________________, ___________________, ___________________ con ___________________ individuo avistado.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de mamíferos se avistaron ___________________ con ___________________ individuos registrados, la especie más abundante es ___________________ con ___________________ individuos avistados, en cuanto a la especie menos abundante para este grupo fueron ___________________ y ___________________ en el ACUSTF con ___________________ individuos avistados.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el caso del grupo de los reptiles se registraron ___________________ y un total de ___________________ individuos la especie más representativa fue ___________________ con ___________________ individuos y la especie menos representativa fue ___________________ con ___________________ individuo registrado,")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de los lepidópteros presentes en el ACUSTF se registraron en el area un total de ___________________ de las cuales suman ___________________ individuos avistados, la especie más abundante en el área fue ___________________ con ___________________ individuos y la especie menos representativa fue ___________________ con ___________________ individuos avistado.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("En conclusión, se puede apreciar que en el área ACUSTF la presencia de especies faunísticas son ________ por lo cual la afectación en el mismo será _____ permitiendo que la biodiversidad no disminuya ya que de las especies presentes serán rescatadas y otras tendrán la capacidad de desplazarse hacia el área del sistema ambiental al momento de la ejecución del proyecto.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.6.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.6.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.6.3.- Recurso Suelo')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Erosión hídrica
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Erosión hídrica')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.6.3 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Erosión Hídrica en el ACUSTF sin implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.6.3 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En las condiciones actuales por efecto de la lluvia se pueden tener pérdidas de ______ mm de suelo/año.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Erosión eólica 
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Erosión eólica ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.6.3 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Erosión Eólica en el ACUSTF sin implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.6.3 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En las condiciones actuales por efecto del viento en el área de Cambio de Uso de Suelo, se pueden tener pérdidas de ____ss mm de suelo/año.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            * Recurso del agua
            * Infiltración
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Recurso agua'
                                         '\nInfiltración')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    #di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.6.3 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Infiltración en el ACUSTF sin implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.6.3 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.6.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Derivado del análisis se concluye que en la condición actual con la cobertura que tiene el área de Cambio de Uso de Suelo se tiene una infiltración normal de _____ mm.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.- Análisis del Escenario con Proyecto.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.1.- Factor Ambiental Atmósfera:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('El efecto que se tendrá es en sinergia con la eliminación de la vegetación que afectara con la emisión de partículas de polvo a la atmósfera, modificación de las condiciones climáticas al no existir amortiguamiento para la reducción de la presencia de estas afectaciones, así como se incrementara el ruido por la presencia de más maquinaria tanto por la ____________________________.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.2.- Factor Ambiental Hidrología Superficial y Subterránea.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('El área de cambio de uso de suelo, con referencia a los escurrimientos dentro del área considerada para la implementación del proyecto estos tomarán mayor velocidad pudiendo ocasionar en forma sinérgica perdida de suelo por arrastres y a la vez afectará la infiltración hacia los mantos freáticos al no existir retención del escurrimiento, no se afectaran cuerpos de aguas o afloramientos naturales de estos al no existir dentro del área, solamente se presenta escurrimiento intermitente superficiales que por la naturaleza del proyecto  ____________________, fuera de estos y aledaños a ellos estos no serán alterados y de ser posible en la ___________________ _________________________ para canalizar el agua de escurrimiento al flujo natural para sus procesos de infiltración y abastecimiento al sistema ambiental. Al eliminarse la vegetación aumenta la evapotranspiración por lo cual se dejará de captar la poca agua, reduciendo su infiltración hasta de _____ mm anual para el caso del área ACUSTF, requiriéndose aplicar medidas de compensación para la recuperación de dicha perdida.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.3.- Recurso Suelo')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Ante los movimientos de suelo por acciones del despalme posterior a la eliminación de vegetación el suelo quedará expuesto y se podrá perder por acciones erosivas de tipo eólico e hídrico al ser eliminada la cubierta vegetal que modificara su cobertura y densidad y por ende las condiciones ambientales desde el punto de vista climática, así mismo por las propias actividades de _________________________ se tendrán modificaciones a la estructura del suelo y aunado a ello habrá emisiones de polvo a la atmosfera en forma constante conforme avance el proyecto. ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('De acuerdo al análisis efectuado tanto en el predio y el área para CUSTF, se obtuvo el siguiente resultado para la erosión hídrica.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.3.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\n')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.7.3.1 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Erosión Hídrica en el ACUSTF con la implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.7.3.1 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.7.3.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con la implementación del proyecto al quedar desnudo el suelo, el factor agua erosiona más rápidamente el área, en el área de cambio de uso de suelo es de _______ mm/ha.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.3.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\n')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.3.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con referencia a la perdida de suelo por acciones eólicas derivado del análisis se obtuvo el siguiente resultado:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    #########################
    ### Título de la tabla del capítulo 12.7.3.2 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Erosión Eólica en el ACUSTF con la implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.7.3.2 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.7.3.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con la implementación del proyecto al quedar el suelo desnudo el factor viento erosiona más rápidamente las áreas en donde se incrementa esta pérdida en el área de cambio de uso de suelo es de _____ mm/año.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.3.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.3.- Recurso Agua')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.3.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con referencia al agua que escurre dentro del área considerada para la implementación del proyecto estos tomarán mayor velocidad pudiendo ocasionar en forma sinérgica perdida de suelo por arrastre y a la vez afectará la infiltración hacia los mantos freáticos al no existir retención del escurrimiento. Al eliminarse la vegetación aumenta la evapotranspiración por lo cual se dejará de captar agua reduciendo su infiltración, requiriéndose aplicar medidas de compensación para la recuperación de dicha perdida.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    #########################
    ### Título de la tabla del capítulo 12.7.3.3 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Infiltración en ACUSTF con la implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.7.3.3 ###
    #########################
    filas = 2
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.7.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.4.- Factor Ambiental Topografía.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.4 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Las modificaciones que sufrirá el área de cambio de uso de suelo serán en la topografía, en el factor relieve esto en referencia tanto en la construcción y operación del proyecto, en la construcción de___________________, que harán sinergia con la modificación al paisaje, mismo que _______________ y al ___________________ este impacto visual será crítico y característico de la zona.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.7.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.5 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.5.- Recurso Flora ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.5 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con la implementación del proyecto se tendrá una remoción de vegetación de acuerdo a lo siguiente:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Tipo de vegetacion
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    """
        Se tiene que describir el reseto del capitulo con bases de datos
    """

    ########################################################################################################################################################################
    # Capitulo 12.7.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.7.6 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.7.6.- Recurso Fauna.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.7.6 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("De acuerdo a los registros obtenidos dentro del área de cabio de uso de suelo al eliminarse la vegetación y se ahuyente los pocos individuos se tendrá un desplazamiento en los grupos siguientes:")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de las aves se cuenta con el registro de __ especies y un total de __ individuos, este grupo contará con un porcentaje de desplazamiento del __% es decir que las especies que se encuentren se podrán desplazar sin ningún inconveniente hacia el área del sistema ambiental con una similitud entre áreas del __ al __%, por lo cual estas especies _____________________________________.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de los mamíferos se registraron __ especies con __ individuos observados, para este grupo se tendrá un porcentaje de desplazamiento de especies del ____% por lo cual las especies se podrán desplazar hacia el área del sistema ambiental sin ningún problema siendo esta área similar en un _______% por lo que las especies podrán establecerse, __________________________________.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de los reptiles se presentaron __ especies con un total de ___ individuos, este grupo contará con un porcentaje de desplazamiento de especies del ___% en donde las especies se podrán desplazar hacia el área del sistema ambiental que comparte con el ACUSTF una similitud del ______% por lo que al desplazarse las especies no tendrán ningún problema, además este grupo debido a que las especies son de lento desplazamiento se implementarán actividades de rescate para ayudar con su desplazamiento y no se vean afectadas de manera significativa.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el grupo de los insectos en el cual solo se registraron especies de lepidópteros se cuenta para este grupo un total de __ especies con un total de ____ individuos avistados, de tal manera este grupo presentará un porcentaje de desplazamiento del ______% haciendo que las especies que se encuentren en el área de estudio se desplacen hacia el área del sistema ambiental con una similitud del _______% entre áreas de tal manera que este grupo no se vea afectado.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.8 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.8.- Paisaje')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.8 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Para este factor se contempla modificar solo el área autorizada, así mismo eliminar los residuos que se generen durante los procesos siguiendo las leyes normativas para su disposición final, además de la extracción de equipos fuera de uso.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Como resultado del análisis para la implementación del proyecto, en cuanto a la afectación que sufrirá el paisaje se tiene lo siguientes resultados con respecto al sistema ambiental.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.8 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Impacto a la fragilidad del Paisaje con la implementación del proyecto en el Sistema Ambiental y ACUSTF.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.8 ###
    #########################
    filas = 9
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.8 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Como se puede apreciar el área de cambio de uso de suelo presenta un impacto a la fragilidad de ___________, sin embargo, tomando en cuenta que dicha superficie es muy pequeño comparado a la superficie total del sistema ambiental, se concluye que el sistema ambiental tiene la capacidad de absorción de estos impactos por la implementación del proyecto siempre y cuando se realicen las medidas de mitigación propuestas. ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Por lo que se presume con base a los resultados que la ejecución del proyecto el paisaje tendrá alteraciones temporales considerando que se aplicarán medidas de prevención, mitigación y/o compensación para minimizar o anular los impactos ambientales que se presentarán por el cambio de uso de suelo de terrenos forestales a excepción de la topografía en su componente suelo que por la naturaleza del proyecto será inevitable permaneciendo de manera permanente. ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.9
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.9 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.9.- Pronóstico Ambiental.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.9 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Los impactos ambientales que sufrirá el área por afectar, son indiscutibles aun cuando se pretende que su magnitud sea menor, por la naturaleza misma del proyecto se observarán afectaciones que serán temporales y otras que permanecerán por el nuevo tipo de uso que se pretende dar al área, incidiendo en los elementos naturales como son: la hidrología, topografía, el suelo, y por otro lado en la vegetación forestal y fauna, la modificación al paisaje; el amortiguamiento de los efectos adversos será como ya se ha descrito, además de las medidas o regulaciones que al respecto emita la autoridad.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Por otra parte, es necesario destacar que el elemento social y económico es un parámetro a considerar por lo que se busca tener un equilibrio con el impacto ambiental que sufrirá esta área, por lo que al hacer una ponderación sensata se llegó a determinar que el uso propuesto para ésta área constituye una opción viable generadora de fuentes de empleos y beneficios económicos al generar empleos, durante las diferentes etapas del proyecto siendo empleos directos e indirectos para la región que sigue fortaleciendo fuentes de empleo, así mismo contribuir ____________________ que es fortaleza para la región de _________________.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("En el área sujeto de estudio para ACUSTF, tal y como se manifiesta en capítulos anteriores por sus características Bióticas y Abióticas solo es viable para aprovechamiento _____________, y la otra parte de _______________________, de acuerdo al análisis y con base al muestreo y la observación en campo el tipo de vegetación que sustenta ___________.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Dentro de la superficie considerada para llevar a cabo el cambio de uso de suelo el tipo de vegetación que predomina es el ________________________________. De acuerdo al registro de vegetación en los sitios muestreados y derivado de las observaciones se ubicaron diferentes especies para su rescate y conservación dentro del predio.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Con respecto a la Fauna en el área sujeta de estudio se ____________________ de algunos, lagomorfos y pequeños mamíferos pudiéndose apreciar diferentes tipos de aves mismas que se registraron en el rubro de fauna dentro del área.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Referente a la hidrología en el área, este se encuentra dentro de los parámetros de precipitación de 125-400 mm media anual según (SMN de CNA) con registro en la estación meteorológica más cercana, dentro del área ______________________, no se localizan cuerpos de agua siendo que no se modificarán los que se encuentren aledaños a los márgenes del área ni sus lechos para que cumpla con su función abastecedora hacia el sistema ambiental en las partes bajas.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("La profundidad del suelo en lo general tiene una profundidad menos a 1.0 m, el cual es de ___________, de acuerdo a los perfiles observados en el sitio, con base a lo observado presenta erosión ________________________ derivado de las lluvias anuales sin que se observaran acarreos importantes a las partes bajas del área.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Bajo este contexto y considerando el nuevo uso que pretende dársele al área sujeto de estudio las afectaciones serán severas y críticas ante la eliminación de vegetación en el área sujeta a cambio de uso de suelo. Para asegurar la biodiversidad dentro del área en estudio se contempla rescatar y reubicar especies de lento crecimiento y difícil regeneración como las cactáceas y asparagáceas las cuales se reubicarán en el sitio considerado como área de reserva o conservación.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("El impacto en el paisaje se modificará en forma paulatina sin poder revertir su modificación en el corto y mediano plazo, sin embargo, con acciones de conservación, restauración y mitigación se podrá simular los efectos negativos, aunque sin llegar a su estado original.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.10. Programa de Manejo Ambiental. ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Un programa de manejo se deriva del análisis de las condiciones de un determinado ecosistema mediante la observación y la evaluación realizada a los impactos ambientales que se pudiesen generar por la ejecución de un proyecto.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("El objetivo principal del programa es presentar medidas destinadas a la prevención, control y mitigación de los potenciales impactos negativos generados sobre los componentes físicos, biológicos, socioeconómicos, como consecuencia de la ejecución del Proyecto de _____________________________, mismo que se apegará a la normatividad vigente en la materia.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Este programa contiene diferentes medidas, que pueden ser de manejo, prevención, mitigación, control, protección, vigilancia o compensación, y la forma, momento y lugar donde deben ser aplicadas, para controlar los impactos identificados.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para el caso del área de estudio, en el cual se detectaron los factores afectados biótico y abiótico derivado de las acciones que implican la ejecución del proyecto visto y analizada desde los siguientes aspectos.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo12_format.space_after = 0
    #descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    factores = [
        "Atmósfera (Calidad del aire, Ruido)",
        "Hidrología Superficial (Escorrentías)",
        "Hidrología Subterránea (Acuíferos)",
        "Suelo (Erodabilidad)",
        "Topografía (Relieve)",
        "Flora (Cobertura, Abundancia)",
        "Fauna (Abundancia)",
        "Paisaje (Armonía, Calidad Paisajística)"
    ]

    for factor in factores:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(factor)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo12_format.space_after = 0
        #descripcionCapitulo12_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.10.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.10.1.-Medidas de ubicación y de diseño.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.1 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'XII.10.1.1.- Atmósfera: ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.1 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En lo referente a calidad del aire y ruido que son los impactos que pudiesen presentarse se consideran como severos y críticos en el desmonte, despalme, ___________________________, para reducir los impactos se recomienda lo siguiente:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    medidas = [
        "Humectar las áreas donde se lleve a cabo esta actividad en forma constante para reducir las emisiones a la atmosfera.",
        "Instalar avisos visibles de tránsito con velocidad de 20-30 Km., por hora para reducir la volatilidad de partículas de suelo a la atmosfera y evitar impactos con la escasa fauna que se pueda presentar dentro del radio de influencia del proyecto.",
        "Verificar que los vehículos que operen en el área estén en condiciones óptimas de funcionamiento y no emitan contaminantes a la atmosfera."
    ]

    for medida in medidas:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(medida)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo12_format.space_after = 0
        #descripcionCapitulo12_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 12.10.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.2 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\n')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con referencia a las escorrentías que pudiesen ser impactadas durante el proceso de preparación del sitio y construcción estas se consideran severos a críticas por el proceso derivado de la eliminación de vegetación donde los escurrimientos incrementarán su velocidad pudiendo generar erosión en el suelo.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Derivado del análisis de las condiciones del área, con la información obtenida y calculada para este factor ambiental se obtuvieron los siguientes resultados:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.10.1.2 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Análisis de Infiltración en el ACUSTF.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.10.1.2 ###
    #########################
    filas = 2
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.10.1.2 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Ante ello y considerando el impacto que genera el cambio de uso del suelo en lo referente a los escurrimientos generando pérdidas de suelo por acciones de erosión hídrica, incremento de velocidad de escurrimientos y baja posibilidad de infiltración hacia los mantos freáticos.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.3 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.10.1.3.- Hidrología Subterránea:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.3 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Respecto a la afectación de acuíferos estos solamente se verán afectados en forma no significativa en la superficie donde se lleven a cabo el cambio de uso del suelo _______________________ en una superficie de ____ ha, para reducir los impactos se considera definir un área del predio sin alteración con sus escurrimientos en forma natural y sin alterar la vegetación para que cumplan la función de infiltración.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.4 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.10.1.4.- Suelo: ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.4 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En lo referente a este factor se considera critico en la superficie solicitada para efectuar el cambio de uso de suelo ya que al eliminar la vegetación se perderá este factor, para ello se contempla mantener un área de conservación en la cual se podrá mantener parte del suelo del despalme para _______________________________________.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.5 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.11.1.5.- Topografía:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.5 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('En lo referente a la topografía donde se modificará el relieve del suelo, las modificaciones serán bastantes considerables ya que se realizan ___________________, para la implementación de la _____________________, siendo crítico en su relieve como se menciona anteriormente será de tipo ______________, por lo que al final de _____________________ se propondrá ______________________________ tratando de llegar lo más cerca posible a su estado natural en cuanto al desarrollo de vegetación.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.6 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.10.1.6.- Flora:')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.6 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('De acuerdo al manejo de la información basada en la condición que presenta el área de cambio de uso de suelo, como en todo tipo de cambio de uso de suelo la vegetación se verá afectada en cuanto a cobertura y abundancia considerándose critico en el primer proceso que es la preparación del sitio, ante ello con el fin de reducir este impacto se contempla la eliminación de la vegetación solo en el área de estudio que es de _____ ha., con lo cual se conservará la biodiversidad y el proceso evolutivo de cada especie, así mismo se manifestaran avisos donde se prohíba el saqueo de especies del área para mantener la mayor cobertura y abundancia posible y mantener la biodiversidad en el sitio en las partes aledañas a ellas.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.10.1.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.7 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.11.1.7.- Fauna: ')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.7 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "En el área de estudio se observaron algunas especies mismas que están registradas en el capítulo de fauna del área en estudio, "
        "se consideran críticos en la preparación del sitio donde se afectaría la densidad y la abundancia ocasionando movilidad de especies "
        "del área aun cuando se __________________ pudiendo ser por la superficie del área que es ________ _____ ha, así como las condiciones "
        "del lugar que no presenta vegetación abundante, al ser analizado durante el mes de ________, _________________________."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Como medida para reducir los impactos se contempla establecer letreros de prohibición de cacería para mantener las escasas especies "
        "en el área además de rescatar y reubicar aquellas de lenta movilidad."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "En un ahuyentamiento de la Fauna se emplean diferentes metodologías y técnicas como estímulos visuales (siluetas o globos), "
        "estímulos auditivos (reproducción de sonidos que alerten del peligro), estímulos mecánicos (movimiento de la vegetación arbórea y arbustiva), "
        "los cuales generan un cierto grado de alerta a la Fauna por lo que se desplazará del lugar. En el área de cambio de uso de suelo el ahuyentamiento "
        "se emplearán diferentes herramientas dependiendo del grupo de especies que se desee ahuyentar, entre los cuales se recomienda la forma de sonidos y "
        "el tipo de ahuyentador que son silbatos, bocinas y claxon de automóviles, varas para mover las ramas de árboles y arbustos, con la finalidad de propiciar "
        "la migración de individuos de especies de Fauna Silvestre, que se encuentren en el área de cambio de uso de suelo y se desplaza a áreas aledañas. "
        "Estas actividades se realizarán unas horas antes de que la maquinaria pesada inicie las actividades de desmontes, el cual consistirá en realizar recorridos "
        "en el terreno de manera sistemática a manera de ir espantando y en su caso capturando mamíferos, haciendo el mayor número de ruido o sonidos."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Para la protección de refugios o áreas de exclusión de hábitats solamente se propone realizar actividades dentro del área solamente, "
        "fuera de este serán protegidas las zonas de madrigueras que pudiera existir, en caso de encontrase madrigueras de mamíferos dentro del área "
        "de cambio de uso de suelo estas serán rescatadas y reubicadas en condiciones similares a donde se encontraron."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "En cuanto a espacios y refugios dentro del área de cambio de uso de suelo _________________ de ellos, sin embargo, si se llegó avistar aves, "
        "________________, para desincentivar los hábitats, una de las actividades que se toman en cuenta inmediatamente es que solamente se realizará "
        "sobre las áreas autorizadas por la autoridad correspondiente, con esto se evita dañar más allá del hábitat de las aves, anudado a esto se prohíbe "
        "estrictamente la movilización, caza o captura de aves, mamíferos o reptiles así como el uso de sus nidos, de encontrase refugio de aves estas serán "
        "reubicados en las condiciones similares aledañas al proyecto."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Así también se tendrá en consideración lo siguiente:"
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    medidas_fauna = [
        "Implementar estrategias de pláticas de educación ambiental (manejo y conservación de fauna silvestre), participación social y cultural para generar una comunicación y difusión que resalten la importancia ecológica de la especie, entre los operadores tanto en el área de trabajo como en áreas de trasporte.",
        
        "Informar a las personas involucradas en el proyecto, sobre la importancia de la especie y su conservación para el medio ambiente y a su vez se tomen medidas drásticas para la protección, cuidados y tratos especiales para la especie.",
        
        "Disminuir la velocidad de vehículos, maquinas, etc., en carreteras con la finalidad de reducir y evitar el impacto ocasionado por viajar a velocidades altas y por consecuencia pequeños invertebrados como los lepidópteros terminen impactados en los cristales de los vehículos, así mismo no solo se beneficiará a estos individuos si no también será beneficiada la fauna en general disminuyendo la mortandad de las especies durante el ciclo de vida del proyecto.",
        
        "Queda estrictamente prohibido la caza, captura y destrucción de refugios de anidamiento en el área de extracción, predio y área de influencia del proyecto."
    ]

    for lista in range(len(medidas_fauna)):
        di12 = doc.add_paragraph()
        descripcionCapitulo12 = di12.add_run(f'{lista + 1}.- {medidas_fauna[lista]}')
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "La ubicación de este proceso será en toda el área contemplada en el estudio sujeta a modificación (ver programa de rescate de fauna capítulo 10)."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Así mismo de las especies de invertebrados de importancia como lo es la mariposa monarca en el área se cuenta con distribución de esta especie esto quiere decir que durante la migración de esta especie es probable avistarla, de tal manera y al considerar el posible avistamiento futuro de esta especie se tomaran las medidas siguientes con la finalidad de salvaguardar a la especie."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    medidas_proteccion = [
        "Reducción de la velocidad en el tránsito de los vehículos en el área a no más de 20 km/h, colocando anuncios alusivos para que se respete los niveles de velocidad tanto para maquinaria como para vehículos particulares, así como también se verificará el cumplimiento de los límites establecidos.",
        "Conservar las especies de flora que sirvan como fuente de alimento, refugio y de percha para la especie durante su paso por el área.",
        "Utilizar los equipos y maquinaria solo cuando sea necesario para disminuir el impacto por ruido.",
        "Prohibir la manipulación y/o manejo de la especie cuando se observe por el área para permitirle el libre paso hacia su destino.",
        "Impartir platicas de divulgación sobre la información científica y de importancia de la especie.",
        "Impartir pláticas de concientización al personal que se encontrará trabajando en el área del proyecto, así puedan identificarla más fácilmente si se llegará a tener presencia durante las actividades del proyecto."
    ]

    for medida in range(len(medidas_proteccion)):
        di12 = doc.add_paragraph()
        descripcionCapitulo12 = di12.add_run(f'{medida + 1}.- {medidas_proteccion[medida]}')
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con respecto a las especies de quirópteros en el área de estudio no se obtuvieron registro de especies de quirópteros, sin embargo, si se llegaran a tener futuros avistamientos se tomarán las medidas siguientes:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    medidas_murcielagos = [
        "Fomentar la divulgación de información de la importancia de los murciélagos para el ecosistema para así contrarrestar falsas supersticiones, mitos y temores, que existe sobre estos individuos y ayudar a la conservación de las especies que se distribuyen en el área.",
        "Para el caso de encontrar murciélagos con algún daño o que entren a casas habitación es recomendable llamar a los siguientes números para su atención, 844 2 93 79 80 para cualquier asesoría y atención de incidentes con murciélagos.",
        "Se conservarán las especies vegetales que sirvan como refugio, hábitat y/o alimentación para las especies de quirópteros.",
        "Queda prohibido la manipulación, extracción, caza, hostigamiento hacia las especies que se pudieran avistar."
    ]

    for medida in medidas_murcielagos:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(medida)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    

    ########################################################################################################################################################################
    # Capitulo 12.10.1.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.10.1.8 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\n')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.10.1.8 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Con referencia a la modificación que sufrirá el paisaje en lo referente a Armonía y Calidad Paisajística; este ante la eliminación de vegetación aun cuando sea mínima altera su condición por lo que es considerada de severa a critica, para reducir dichos impactos se plantea aplicar procesos de ________________, para resarcir en parte los daños o modificaciones sin llegar a alcanzar su condición original.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Las obras y/o actividades descritas se contemplan como medida de mitigación para reducir los impactos generados al paisaje, _____________________________, en donde se establecerán especies de la región.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Procedimientos de construcción y operación (en su caso)
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nProcedimientos de construcción y operación (en su caso)')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Dentro de los procesos que se establezcan para el control de aquellos impactos que se generen en los diferentes procesos motivo del estudio, se construirán e instalaran los avisos necesarios para la protección principalmente de la flora y la fauna así mismo si es necesario establecer sitios para el acumulamiento  de residuos temporales de suelo y vegetación para incorporarse al mismo suelo en el área de obras de conservación y/o traslado a deposito final en caso requerido, estos procedimientos serán aplicados antes de iniciar cualquier acción posterior a la autorización si así lo determina la autoridad competente.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.11
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.11 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.12.- Programa de monitoreo Ambiental.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.11 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "El Programa de Monitoreo Ambiental constituye una herramienta destinada a verificar el cumplimiento de las medidas planteadas."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "La ejecución de este Programa de Monitoreo estará a cargo del responsable del área por parte de la empresa bajo la asesoría del técnico responsable mediante los mecanismos que se establezcan para tal efecto, dando un seguimiento puntual a las medidas tanto preventivas como correctivas basado en una bitácora de seguimiento."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "El monitoreo se efectuará durante las etapas de Preparación del sitio y Construcción mediante la observación y seguimiento al listado de factores y componentes o indicadores ambientales registrados y descritos en este estudio en cuanto a las medidas a adoptar para mitigar los impactos durante estos procesos y así mismo de ser necesario reordenar o aplicar medidas preventivas y/o correctivas adicionales a las propuestas al manifestarse impactos que no fueron observados durante el análisis del predio."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Este programa está ligado al calendario de ejecución del proyecto siempre y cuando sea autorizado por la autoridad competente en la materia. Por ello, el Programa de Monitoreo Ambiental servirá como una herramienta que retroalimente al Programa de Prevención Corrección y Mitigación, de tal modo que los impactos ambientales se atenúen, eliminen o se reduzcan a la mínima afectación."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.12
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.12 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nX.12.- Planes de contingencia y respuesta de emergencia.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.12 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Para todo tipo de acción es requerido implementar un plan de carácter preventivo que se enlace con los planes de protección de carácter gubernamental, máxime en este caso donde se pretende llevar a cabo la operación de equipos y seres humanos considerando los riesgos propios de la operación ya que tendrá en forma intermitente seres humanos y uso de maquinaria de acuerdo a como se vaya avanzando en los procesos en caso de ser autorizado su cambio de uso."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Para tal efecto es necesario considerar desde los procesos iníciales un plan que contenga lo siguiente:"
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "El Plan de Contingencia está dividido en dos partes: Plan Estratégico y Plan de Acción; el primero definirá la estructura y la organización para la atención de emergencias, las funciones y las responsabilidades de las personas encargadas de ejecutar el Plan; los recursos necesarios y, las estrategias preventivas y operativas que deben aplicarse en cada uno de los posibles escenarios, definidos a partir de la evaluación de los riesgos asociados a la implementación del proyecto."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "El Plan de Acción por su parte, establecerá los procedimientos a seguir en caso de Emergencia para la aplicación a cada una de las fases de respuesta establecidas en el Plan Estratégico."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Plan Estratégico.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nPlan Estratégico.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items = [
        "Estrategias de prevención y control de contingencias.",
        "Estrategias preventivas.",
        "Definición de responsabilidades.",
        "Estrategias operativas.",
        "Acciones generales para el control de contingencias.",
        "Plan de evacuación.",
        "Control de contingencias por incendio.",
        "Acciones en caso de huracanes.",
        "Acciones en caso de inundaciones.",
        "Equipos para la prevención y control de contingencias.",
        "Organización y recursos.",
        "Funciones y responsabilidades del personal durante la contingencia.",
        "Programa de capacitación."
    ]

    for i, texto in enumerate(items):
        letra = chr(ord('a') + i)
        di12 = doc.add_paragraph()
        descripcionCapitulo12 = di12.add_run(f"{letra}) {texto}")
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1

        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Como medida inicial deberá de establecerse el siguiente mecanismo, mismo que se podrá aplicar en caso de cualquier incidente de acuerdo a la magnitud y según la lista anterior.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PLAN DE MANEJO DE INCIDENTE
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nPLAN DE MANEJO DE INCIDENTE')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.12 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.x.- Plan de manejo de incidente.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.12 ###
    #########################
    filas = 5  # 1 encabezado + 5 filas de datos
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # Encabezados
    componentes = [
        'Mando de Incidente',
        'Operaciones',
        'Planeación',
        'Logística',
        'Finanzas',
    ]

    for rows in range(filas):
        cell = tabla12b.cell(rows, 1)
        t12b = cell.paragraphs[0].add_run(componentes[rows])
        t12b.font.size = Pt(12)
        t12b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_top = tabla12b.cell(0, 0)
    cell_bottom = tabla12b.cell(4, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('Componentes')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 12.12 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Para cada una de estas fases se deberá de contar con un protocolo de actuación ligado con aquellos que apliquen las instancias gubernamentales en los periodos de mayor riesgo para el caso de ocurrencia de eventos extraordinarios tales como, nevadas, bajas temperaturas extremas, lluvias extraordinarias, incendios forestales entre otros, por estar el predio en una región donde se pueden presentar dichos eventos."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.13
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.13 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.13.- Medidas socioeconómicas.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.13 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Dentro de los impactos considerados con la implementación del proyecto sujeto de estudio en caso de proceder este aspecto se considera positivo y benéfico desde diferentes puntos de vista."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Primero, al considerar el predio como __________________________________________ por la propia condición estructural que ostenta en cuanto al uso actual y/o potencial del suelo, el establecimiento del proyecto conllevara a la generación empleos directos e indirectos que activarán la economía local y regional mejorando el aspecto socioeconómico y mejorando la calidad de vida."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Derivado de esta distribución actividades contempladas para la implementación del proyecto se considera la generación de un total de ___ empleos _______________ beneficiando a una población de _____ personas durante la ejecución del proyecto así mismo se podrán generar ____ empleos adicionales indirectos que juntos harán una derrama económica de más de $ ___________ de pesos en la operación del proyecto, beneficiando en primer lugar a la región de _____________________, con la generación de empleos,  percibiendo una retribución monetaria alta no solo en la etapa de construcción  sino también en la operación con el establecimiento de la ______________________,  lo que se traduce en bienestar social de más de _________ el cual incrementara la calidad de vida y mantendrá sin problemas económicos durante la operación del proyecto con la generación de empleos por la inversión."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.13 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x- Derrama económica empleada para la generación de empleos.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.13 ###
    #########################
    filas = 5  # 1 encabezado + 5 filas de datos
    columnas = 2
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '#0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)
            t10b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 12.13 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "*Los costos de los salarios es en base a los costos del salario mínimo del país, el cual es de ________ pesos"
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ########################################################################################################################################################################
    # Capitulo 12.14
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.14 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.14.-Compensación por pérdida o daños.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.14 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Con respecto a las acciones que se puedan implementar para compensar los daños ocasionados por la posible ejecución del proyecto, para este mecanismo la autoridad normativa emite el costo por ha de modificación de acuerdo al tipo de ecosistema con lo cual se podrían llevar actividades como reforzamiento a aquellos impactos que se mantendrán aun con las medidas de mitigación en los elementos del factor biótico y abiótico."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Las medidas de compensación definidas en este contexto como las acciones dirigidas a resarcir o retribuir la biodiversidad por los impactos negativos generados en los procesos del proyecto de acuerdo a las obras o actividades definidas serán aplicables en caso que ya no sea posible minimizar o restaurar los impactos."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "De acuerdo a los sitios de afectación es recomendable que los recursos que se obtengan por el cambio de uso de suelo se apliquen en el mismo ecosistema semejante a donde se hizo la modificación."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Ante ello en cada uno de los rubros se menciona el grado de afectación en cada elemento, aire, agua, suelo, flora y fauna y las posibles medidas con las cuales se puede mitigar el efecto."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "El área para compensación derivado del mecanismo administrativo para aplicación de recursos y resarcir los daños deberá estar sujeto prioritariamente a lo siguiente:"
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items = [
        "Que se trate del mismo ecosistema natural afectado.",
        "Que sea de igual o mayor tamaño lo que determina la aplicación del factor compensatorio.",
        "Que sea de igual o mayor contexto paisajístico al fragmento del ecosistema impactado.",
        "Que tenga igual o mayor riqueza de especies al fragmento del ecosistema impactado."
    ]

    for item in enumerate(items):
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(f"{item}")
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Dentro de las acciones contempladas para la implementación del proyecto se contemplan los siguientes programas con sus montos aproximados para su ejecución."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.14 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Programas contemplados con la implementación del proyecto.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.14 ###
    #########################
    filas = 7  # 1 encabezado + 5 filas de datos
    columnas = 3
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '#0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)
            t10b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 12.15
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.15 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.15.- Acciones de supervisión.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.15 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Es necesario establecer un programa de supervisión para verificar que las acciones recomendadas se apliquen conforme a cada uno de los procesos que implica el desarrollo del proyecto para ello se tomara como base el periodo que se contempla para ejecución de la obra, considerando que este proceso se llevará a cabo en mediano y largo plazo en virtud de que se efectuara de acuerdo a la oferta y demanda de los materiales, aspecto que facilitará las acciones de supervisión mismas que serán de acuerdo a las medidas de mitigación que se establecen en cada una de las acciones que implica ________________________________.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.16
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.16 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.16.- Evaluación del Desempeño Ambiental.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.16 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Es una herramienta que permite medir y evaluar el desempeño ambiental de acuerdo a las medidas de mitigación propuestas para cada una de las actividades inherentes al proceso del proyecto, de los resultados que se observen durante los procesos se podrán modificar o mantener según sea el caso.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Este proceso deberá de ser continuo y acorde al desarrollo o implementación de cada acción del proyecto en los cuales se han detectado los posibles impactos, con ello se podrá determinar su magnitud e implementar alguna acción con la tendencia de minimizar en lo mayor posible sus efectos.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run("Para los procesos de revisión será necesario documentar tipo de impacto mecanismo de remediación y resultados alcanzados.")
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 12.17
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.17 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.17.-Procedimientos para instrumentar medida de mejora.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.17 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Derivado del análisis o evaluación por las autoridades en la materia al presente estudio en el cual se emitirán recomendaciones para fortalecer acciones de mitigación hacia aquellos impactos que permanecerán por la propia condición del proyecto y tomando con base la causa y efecto de cada impacto, así como las medidas de mitigación aplicadas para reducir dichos impactos estos deberán de evaluarse constantemente en cada uno de los procesos registrándose de la siguiente forma:"
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "ACCIÓN del proyecto en la etapa correspondiente (Preparación del sitio, Construcción), al considerar el EFECTO que se siga manifestando aun con las medidas de mitigación serán resultados de una CAUSA que obliga a evaluar la aplicación de las MEDIDAS DE MITIGACIÓN, si estas se siguen manifestando se tendrán que implementar MEJORA de las medidas para contrarrestar los efectos y eliminar las causas que generen la manifestación del impacto para que este sea de menor afectación o desaparezca del entorno del proyecto."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.17 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Bitácora de registro para acciones de mejora en mitigación de impactos.')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.17 ###
    #########################
    encabezados = [
        'Acción del Proyecto',
        'Efecto',
        'Causa',
        'Medida de Mitigación',
    ]

    datos_tabla = [
        ['Preparación del sitio', '\u2714', '\u2714', '\u2714'],
        ['Construccion', '\u2714', '\u2714', '\u2714'],
        ['Operación', '\u2714', '\u2714', '\u2714']
    ]
    filas = len(datos_tabla) + 1  # 1 fila para encabezados
    columnas = len(encabezados)   # 4 columnas
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # ✅ Encabezados
    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.name = 'Arial'
        t12b.font.size = Pt(12)
        t12b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#0070C0')

    # ✅ Filas de datos
    for fila in range(len(datos_tabla)):            # 0..2
        for col in range(len(datos_tabla[fila])):   # 0..3
            cell = tabla12b.cell(fila + 1, col)     # +1 porque fila 0 es encabezado
            texto = datos_tabla[fila][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            t12b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
                Matriz de planeación
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 12.17 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nMatriz de planeación')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.17 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Costo de medidas de mitigación despalme y construcción')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.17 ###
    #########################
    encabezados = [
        'Acción del Proyecto',
        'Medida de Prevención y Mitigación',
        'Duración o Tiempo en el que se instrumentará',
        'Recursos necesarios ($): costo, equipos, obras, instrumentos, etc.',
        'Supervisión y grado de cumplimiento, eficiencia y eficacia',
    ]

    datos_tabla = [
        [
            'Rescate y reubicación de especies de flora y fauna del área previo al desmonte con mayor énfasis en aquellas que se encuentren listadas en la NOM-059-SEMARNAT-2010.', 
            'Inicialmente se aplicará el rescate en un periodo de 3-6 meses, anualmente de acuerdo al desarrollo se llevará a cabo un monitoreo adicional para rescate de especies que se detecten. ', 
            ' ', 
            'Verificación del registro de especies rescatadas manteniendo una supervivencia del 85 % por un periodo de 5 años.'
        ],
        [
            'Medidas de Prevención y Mitigación', 
            'Establecimiento de letrinas, riego de mitigación de polvos', 
            ' ', 
            'Aplicación de 2 riesgo por semana o cuando se rebase el umbral de alerta, establecimiento de letrinas 1 por cada 15 personas.'
        ],
    ]
    filas = len(datos_tabla) + 3  # 1 fila para encabezados
    columnas = len(encabezados)   # 4 columnas
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(3, 0)
    cell_bottom = tabla12b.cell(4, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('Desmonte y Despalme')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Linea Estrategica')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Conservación de los Ecosistemas y su Biodiversidad')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Etapa del Proyecto')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Preparación del Sitio')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Encabezados
    for col in range(columnas):
        cell = tabla12b.cell(2, col)
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.name = 'Arial'
        t12b.font.size = Pt(12)
        t12b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#79C9FF")

    # ✅ Filas de datos
    for fila in range(len(datos_tabla)):            # 0..2
        for col in range(len(datos_tabla[fila])):   # 0..3
            cell = tabla12b.cell(fila + 3, col + 1)     # +1 porque fila 0 es encabezado
            texto = datos_tabla[fila][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            t12b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 12.17 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Costo de medidas de mitigación caminos y construcciones complementarias')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.17 ###
    #########################
    encabezados = [
        'Acción del Proyecto',
        'Medida de Prevención y Mitigación',
        'Duración o Tiempo en el que se instrumentará',
        'Recursos necesarios ($): costo, equipos, obras, instrumentos, etc.',
        'Supervisión y grado de cumplimiento, eficiencia y eficacia',
    ]

    datos_tabla = [
        [
            'Despalme',
            'Establecer un programa de monitoreo ambiental para cada medida planteada', 
            'Diariamente en ejecución del proyecto', 
            ' ', 
            'Verificación y cumplimiento de cada una de las medidas propuestas con el establecimiento del proyecto'
        ],
        [
            'Construcción de Rampas y Bancos de Extraccion',
            'Humectación de áreas para evitar dispersión de polvos', 
            ' ', 
            ' ', 
            'Verificación en bitácora la aplicación de humectación, mensual, la eficiencia será la nula volatilidad de partículas y acarreo de sustratos por escurrimientos.'
        ],
    ]
    filas = len(datos_tabla) + 3  # 1 fila para encabezados
    columnas = len(encabezados)   # 4 columnas
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Linea Estrategica')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Conservación de los Ecosistemas y su Biodiversidad')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Etapa del Proyecto')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Preparación del Sitio')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Encabezados
    for col in range(columnas):
        cell = tabla12b.cell(2, col)
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.name = 'Arial'
        t12b.font.size = Pt(12)
        t12b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#79C9FF")

    # ✅ Filas de datos
    for fila in range(len(datos_tabla)):            # 0..2
        for col in range(len(datos_tabla[fila])):   # 0..3
            cell = tabla12b.cell(fila + 3, col)     # +1 porque fila 0 es encabezado
            texto = datos_tabla[fila][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            t12b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 12.17 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Costo de medidas de mitigación caminos y construcciones complementarias')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.17 ###
    #########################
    encabezados = [
        'Acción del Proyecto',
        'Medida de Prevención y Mitigación',
        'Duración o Tiempo en el que se instrumentará',
        'Recursos necesarios ($): costo, equipos, obras, instrumentos, etc.',
        'Supervisión y grado de cumplimiento, eficiencia y eficacia',
    ]

    datos_tabla = [
        [
            'Preparación y construcción y operación ',
            ' ', 
            'Diaria 1 x cada 15 personas', 
            'Letrinas $ 15,000.00 anuales', 
            'Verificación de residuos residuales de aguas mediante bitácoras semanales.'
        ],
        [
            'Vigilancia de medidas de mitigación',
            'Establecer un programa de monitoreo ambiental para cada medida planteada',
            'Diariamente en ejecución del proyecto', 
            ' ', 
            'Verificación y cumplimiento de cada una de las medidas propuestas con el establecimiento del proyecto.'
        ],
    ]
    filas = len(datos_tabla) + 3  # 1 fila para encabezados
    columnas = len(encabezados)   # 4 columnas
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Linea Estrategica')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Conservación de los Ecosistemas y su Biodiversidad')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#2569FA")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Etapa del Proyecto')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Preparación del Sitio')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, "#0099FF")  # Cambiar el color de fondo de la celda fusionada

    # ✅ Encabezados
    for col in range(columnas):
        cell = tabla12b.cell(2, col)
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.name = 'Arial'
        t12b.font.size = Pt(12)
        t12b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, "#79C9FF")

    # ✅ Filas de datos
    for fila in range(len(datos_tabla)):            # 0..2
        for col in range(len(datos_tabla[fila])):   # 0..3
            cell = tabla12b.cell(fila + 3, col)     # +1 porque fila 0 es encabezado
            texto = datos_tabla[fila][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            t12b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 12.17 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nSe pone a consideración de la autoridad el monto por las medidas de mitigación, prevención y/o compensación para fijar la fianza de acuerdo a lo siguiente:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 12.17 ###
    #########################
    tituloTabla12b = doc.add_paragraph()
    dti12b = tituloTabla12b.add_run('\nTabla 12.x.- Costo de medidas de mitigación fianza')
    dti12b_format = tituloTabla12b.paragraph_format
    dti12b_format.line_spacing = 1.15
    dti12b_format.space_after = 0

    dti12b.font.name = 'Bookman Old Style'
    dti12b.font.size = Pt(12)
    tituloTabla12b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.17 ###
    #########################
    columnas = 2
    filas = 8
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla12b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla12b.cell(rows, cols)
            t12b = cell.paragraphs[0].add_run(' ')
            t12b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 12.18
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 12.18 ###
    #########################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\nXII.18.- Seguimiento y Control.')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True
    capitulo12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 12.18 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Considerando el objetivo para lo cual se está solicitando el cambio de uso del suelo en terrenos del ____________________________________, con el fin de llevar a cabo la ________________________, es necesario llevar a cabo un seguimiento y control puntual sobre las acciones que se lleven a cabo, así como la atención expedita de las recomendaciones que se emitan una vez aprobado el presente estudio para los procesos solicitados."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Mediante la aplicación de una bitácora de registro a través de la cual se puedan detectar situaciones de riesgo de carácter ambiental social y económico, ello para facilitar los planes de contingencia necesarios al momento de incidentes de cualquier orden dado que se pretende darle uso en el que está involucrado el ser humano, maquinaria y equipo, ello además por los fenómenos naturales que pueden presentarse en la zona."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(
        "Desde el punto de vista ambiental el registro servirá para tener un control sobre las manifestaciones de los impactos aun con las medidas de mitigación para aplicar mejoras y a la vez aplicarse los procedimientos tal y como se han venido plasmando en el documento aspecto que evitará llegar a ser sancionado."
    )
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capitulo 12.18 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    filas = 17
    columnas = 9
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[8]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('BITÁCORA DE SEGUIMIENTO Y CONTROL DE MITIGACIÓN DE IMPACTOS AMBIENTALES')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('ETAPA: ')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[5].merge(row.cells[5].merge(row.cells[8]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('FECHA DE ELABORACION:')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True

    # ✅ Celdas fusionadas
    row = tabla12b.rows[2]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('SEMANA Y MES: ')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True

    # ✅ Celdas fusionadas
    row = tabla12b.rows[2]
    merged_cell = row.cells[5].merge(row.cells[5].merge(row.cells[8]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('DIA:')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 0)
    cell_bottom = tabla12b.cell(5, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('FACTORES A REVISAR')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 1)
    cell_bottom = tabla12b.cell(5, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('COMPONENTES A REVISAR')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    row = tabla12b.rows[4]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[4]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('CONDICIÓN EN LA QUE SE ENCUENTRA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(5, 2)
    t12b = cell.paragraphs[0].add_run('Buena')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(5, 3)
    t12b = cell.paragraphs[0].add_run('Regular')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(5, 4)
    t12b = cell.paragraphs[0].add_run('Mala')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 5)
    cell_bottom = tabla12b.cell(5, 5)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('MEDIDA DE APLICACIÓN')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 6)
    cell_bottom = tabla12b.cell(5, 6)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('ACCIÓN')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 7)
    cell_bottom = tabla12b.cell(5, 7)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('MEJORA DE LA ACCIÓN')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(4, 8)
    cell_bottom = tabla12b.cell(5, 8)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('OBSERVACIÓN')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(6, 0)
    cell_bottom = tabla12b.cell(7, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('ATMOSFERA')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    cell_top = tabla12b.cell(12, 0)
    cell_bottom = tabla12b.cell(13, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('PAISAJE')
    t12b.font.name = 'Arial'
    t12b.bold = True
    t12b.font.size = Pt(12)

    # ✅ Celdas independiente
    cell = tabla12b.cell(6, 1)
    t12b = cell.paragraphs[0].add_run('CALIDAD AIRE')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(7, 1)
    t12b = cell.paragraphs[0].add_run('RUIDO')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(8, 0)
    t12b = cell.paragraphs[0].add_run('TOPOGRAFÍA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(9, 0)
    t12b = cell.paragraphs[0].add_run('H. SUPERFICIAL')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(10, 0)
    t12b = cell.paragraphs[0].add_run('SUELO')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(11, 0)
    t12b = cell.paragraphs[0].add_run('FAUNA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(8, 1)
    t12b = cell.paragraphs[0].add_run('RELIEVE')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(9, 1)
    t12b = cell.paragraphs[0].add_run('ESCURRIMIENTOS')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(10, 1)
    t12b = cell.paragraphs[0].add_run('ERODABILIDAD')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(11, 1)
    t12b = cell.paragraphs[0].add_run('ABUNDANCIA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(12, 1)
    t12b = cell.paragraphs[0].add_run('ARMONIA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas independiente
    cell = tabla12b.cell(13, 1)
    t12b = cell.paragraphs[0].add_run('CALIDAD PAISAJÍSTICA')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    t12b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[16]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('RESPONSABLE DE LA ELABORACIÓN')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[16]
    merged_cell = row.cells[6].merge(row.cells[6].merge(row.cells[8]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('RESPONSABLE DE LA EJECUCION DEL PROYECTO')
    t12b.font.name = 'Arial'
    t12b.font.size = Pt(12)
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    ########################################################################################################################################################################
    ### Hoja en Vertical para ver contenido del capitulo 12.18 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripcion del capitulo 12.18 ###
    #########################
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Derivado de los resultados de la matriz de identificación en el subnumeral XI.4.5- Resultados Obtenidos en las matrices del capítulo XI., se considera como la herramienta que proporcionará información para medir y evaluar las medidas preventivas y de mitigación aplicadas, con ello se determinan las modificaciones mejoras y/o correcciones mismas que se realizaran de acuerdo al efecto registrado y que no se estén cumpliendo los fines de aplicación o que requieran alguna modificación con el propósito de dar cumplimiento a las medidas propuestas o alguna otra determinación que manifieste la autoridad para dar cumplimiento en cada acción del proyecto en sus diferentes etapas de implementación; las mediciones o instrumentación se realizara bajo lo siguiente:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor a Afectar, Atmosfera 
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor a Afectar, Atmosfera')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Calidad del Aire. - se medirá la calidad del aire en función de la contaminación principalmente por monóxido de carbono que arrojan la maquinara que estará trabajando según sea la etapa y se clasificara según la norma oficial mexicana NOM-045-SEMARNAT-2017.",
        "Visibilidad. - esto es consecuencia de las condiciones en que se encuentre el aire, por lo que al tener mala visibilidad o regular se aplicara las medidas necesarias como son humectación del sustrato, para que este no siga contaminando con partículas de polvo.",
        "Ruido. - se aplicará la medición diaria de ruido clasificándolo de acuerdo a lo NOM-080-SEMARNAT-1994 Y NOM-081-SEMARNT-1994, de acuerdo a el peso del vehículo."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor Topografía 
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor Topografía')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Relieve. - se revisará que en el área de trabajo no se dejen acumulaciones de sustrato que causen encharcamiento o desvió de las corrientes de agua y que este sea dentro de lo autorizado."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor suelo
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor suelo')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Erodabilidad. - Verificar si las actividades que se realizan se encuentran a dentro de los vértices para realizar el despalme nivelación y cortes, así como los rellenos autorizados.",
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor Ambiental Hidrología
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor Ambiental Hidrología')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Escurrimiento superficial. - verificar que no exista obstrucción alguna en los cauces naturales de escurrimientos y verificar que se lleve a cabo las actividades de desagüe por los canales y conducirlos a sus cauces naturales.",
        "Escurrimientos subterráneos. - evitar la contaminación del agua, no tirar residuos o basura, utilizando letricas secas portátiles, verificando el funcionamiento y mantenimiento de estas."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor Ambiental Flora
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor Ambiental Flora')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Densidad. - para conservar este recurso se verificará que no se eliminen, ni se maltraten los árboles existentes en pie para que favorezca a la proliferación de semilla y la presencia de arbustivas, así como de gramíneas.",
        "Diversidad. - conservar las especies que se encuentran en el área sin eliminarlas y darles su mantenimiento respectivo como podas de formación, prohibir la extracción de individuos.",
        "Especies enlistadas. - de ser encontrada alguna especie que este en la NOM-059-SEMARNAT-2010 se realizara su respectivo rescate y reubicación."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor ambiental Fauna
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor ambiental Fauna')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Hábitat. - se realizará el monitoreo y conteo del centro de anidación si se llega a detectar y que no sean eliminados, así como de animales de lento movimiento serán reubicados se llegan a detectar dentro del área de cambio de uso de suelo.",
        "Densidad. - monitorear la poca o nulas especies que se localice en el área y para que esta sea mayor la densidad proponer poner comederos para aves y mamíferos para mantener el equilibrio de las especies.",
        "Especies enlistadas. - Monitorear el área y predio para su detección y de ser encontradas rescatarlas reubicarlas y realizar su registro correspondiente."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Factor Ambiental Paisaje
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Factor Ambiental Paisaje')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    items_lista = [
        "Calidad. - evitar la diseminación de residuos domésticos, aguas residuales y de algún residuo de manejo especial.",
        "Visibilidad. - realizar recolección de residuos mediante control con recipientes adecuados y por empresas autorizadas en el manejo de los mismos.",
        "Fragilidad. - realizar calendario de limpiezas de las áreas de trabajo en el área de cambio de uso de suelo y recomendar a las empresas que presten algún servicio apegarse al manejo de la normatividad correspondiente.",
        "Medidas de Aplicación. - En este apartado se verificará si se está llevando a cabo las Medidas de Mitigación, que se propusieron en el documento técnico.",
        "Acción. - Revisar si se está llevando a cabo las medidas de aplicación y como se está realizando.",
        "Mejora de la acción. - En este apartado se trata si alguna Medida de Mitigación no está cumpliendo aun con las acciones de verificación se propondrá una nueva medida y una nueva acción para corregirla."
    ]

    for texto in items_lista:
        di12 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo12 = di12.add_run(texto)
        descripcionCapitulo12_format = di12.paragraph_format
        descripcionCapitulo12_format.line_spacing = 1.15
        descripcionCapitulo12.font.name = 'Arial'
        descripcionCapitulo12.font.size = Pt(12)
        di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capitulo 12.18 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            CONTROL DE USO DE AGUA PARA CONTROL DE POLVO
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('CONTROL DE USO DE AGUA PARA CONTROL DE POLVO')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Contratista y Fecha de emision
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=2, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('CONTRATISTA: ______________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('FECHA DE EMISIÓN: _________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('VOLUMEN m\u00B3 AUTORIZADO')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\n')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'FECHA',
        'No. ECONOMICO',
        'PLACA',
        'CAPACIDAD',
        'OPERADOR',
        'FRENTE',
        'VOLUMEN DISPUESTO m\u00B3',
    ]

    tabla12b = doc.add_table(rows=5, cols=len(encabezados), style='Table Grid')

    for i, cols in enumerate(encabezados):
        cell = tabla12b.cell(0, i)
        t12b = cell.paragraphs[0].add_run(cols)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(16)
        t12b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#CFCFCF')

    """
        *************************************
        Quitar los bordes de las celdas
        *************************************
    """
    cell = tabla12b.cell(4, 0)
    quitar_borde_especifico(cell, 'left')
    quitar_borde_especifico(cell, 'right')
    quitar_borde_especifico(cell, 'bottom')

    cell = tabla12b.cell(4, 1)
    quitar_borde_especifico(cell, 'left')
    quitar_borde_especifico(cell, 'right')
    quitar_borde_especifico(cell, 'bottom')

    cell = tabla12b.cell(4, 2)
    quitar_borde_especifico(cell, 'left')
    quitar_borde_especifico(cell, 'right')
    quitar_borde_especifico(cell, 'bottom')

    """
        *************************************
        Celda fusionada
        *************************************
    """
    # ✅ Celda fusionada
    row = tabla12b.rows[4]
    merged_cell = row.cells[3].merge(row.cells[3].merge(row.cells[5]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('VOLUMEN TOTAL SEMANA m\u00B3')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    t12b.bold = True
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\n')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        -------------------------------------
            Responsables
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=2, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('ELABORÓ')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('APROBÓ')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('\n________________________________________________\nSUPERVISOR DE OBRA')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('\n________________________________________________\nSUPERVISOR AMBIENTAL')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PROGRAMA ANUAL DE MANTENIMIENTO A MAQUINARIA Y EQUIPO
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('PROGRAMA ANUAL DE MANTENIMIENTO A MAQUINARIA Y EQUIPO')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Datos del Programa de Mantenimiento
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=6, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Contratista: __________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('No. de Contrato: _______________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(2, 0)
    t12b = cell.paragraphs[0].add_run('Equipo o maquinaria: ___________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(3, 0)
    t12b = cell.paragraphs[0].add_run('Modelo: ______________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(4, 0)
    t12b = cell.paragraphs[0].add_run('Fecha de Ingreso: _____________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(5, 0)
    t12b = cell.paragraphs[0].add_run('Observación: _________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Periodo de generación: __________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('Fecha de elaboración de registro: _________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(2, 1)
    t12b = cell.paragraphs[0].add_run('Marca: ________________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(3, 1)
    t12b = cell.paragraphs[0].add_run('No. de serie: ___________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(4, 1)
    t12b = cell.paragraphs[0].add_run('Actividad: ______________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'No.',
        'Actividad',
        'Enero',
        'Febrero',
        'Marzo',
        'Abril',
        'Mayo',
        'Junio',
        'Julio',
        'Agosto',
        'Septiembre',
        'Octubre',
        'Noviembre',
        'Diciembre',
    ]

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla12b = doc.add_table(rows=13, cols=len(encabezados), style='Table Grid')

    datos_tabla = [
        'Cambio de aceite',
        'Cambio de aceite de transmisión',
        'Cambio de filtro de diésel',
        'Cambio de filtro de aceite',
        'Cambio de filtro de aire',
        'Verificación de niveles de electrolitos',
        'Limpieza del equipo',
        'Verificación del nivel de aceite',
        'Verificación del nivel de aceite de transmisión hidráulica',
        'Verificación de nivel de anticongelante',
        'Aplicación de grasa de articulaciones a maquinaria',
        'Inspección visual de maquinaria y vehículos',

    ]

    for i, cols in enumerate(encabezados):
        cell = tabla12b.cell(0, i)
        t12b = cell.paragraphs[0].add_run(cols)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(12)
        t12b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#CFCFCF')

    for numero in range(12):
        cell = tabla12b.cell(numero + 1, 0)
        t12b = cell.paragraphs[0].add_run(f'{numero + 1}')
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, actividad in enumerate(datos_tabla):
        cell = tabla12b.cell(i + 1, 1)
        t12b = cell.paragraphs[0].add_run(f'{actividad}')
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for rows in tabla12b.rows:
        rows.cells[0].width = Cm(0.84)
        rows.cells[1].width = Cm(5.73)
        rows.cells[2].width = Cm(1.22)
        rows.cells[3].width = Cm(1.54)
        rows.cells[4].width = Cm(1.26)
        rows.cells[5].width = Cm(1.07)
        rows.cells[6].width = Cm(1.12)
        rows.cells[7].width = Cm(1.17)
        rows.cells[8].width = Cm(1.08)
        rows.cells[9].width = Cm(1.37)
        rows.cells[10].width = Cm(2.05)
        rows.cells[11].width = Cm(1.54)
        rows.cells[12].width = Cm(1.94)
        rows.cells[13].width = Cm(1.85)

    """
        -------------------------------------
            Firmas de los responsables
        -------------------------------------
    """
    tabla12b = doc.add_table(rows=3, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Elaboró')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 0)
    t12b = cell.paragraphs[0].add_run('Supervisor de Obra')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Aprobó')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 1)
    t12b = cell.paragraphs[0].add_run('Supervisor Ambiental')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            PROGRAMA DE VERIFICACIÓN Y MANTENIMIENTO VEHICULAR
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('PROGRAMA DE VERIFICACIÓN Y MANTENIMIENTO VEHICULAR')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Datos del Programa de Mantenimiento
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=2, cols=1, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Reporte número: _____________________________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('PERIODO DE: ______________________ A ______________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'Marca',
        'Año',
        'Placa',
        'Tipo de combustible',
        'Operador',
        'Km recorridos',
        'Fecha de lectura',
        'Km de ingreso',
        'Próximo mantenimiento',
        'Última fecha de servicio',
        'Servicio relizado',
        'Nombre taller o agencia',
        'Registro de verificación vehicular',
        'Observaciones',
    ]

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla12b = doc.add_table(rows=13, cols=len(encabezados), style='Table Grid')

    for i, cols in enumerate(encabezados):
        cell = tabla12b.cell(0, i)
        t12b = cell.paragraphs[0].add_run(cols)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(11)
        t12b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#CFCFCF')

    """
        -------------------------------------
            Firmas de los responsables
        -------------------------------------
    """
    tabla12b = doc.add_table(rows=3, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Elaboró')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 0)
    t12b = cell.paragraphs[0].add_run('Supervisor de Obra')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Aprobó')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 1)
    t12b = cell.paragraphs[0].add_run('Supervisor Ambiental')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            BITÁCORA SEMANAL DE RESIDUOS NO PELIGROSOS
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('BITÁCORA SEMANAL DE RESIDUOS NO PELIGROSOS')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Datos de la bitacora
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=1, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Contratista: ____________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Fecha de registro: _______________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'Tipo',
        'Cantidad kg',
        'Sitio de almacenamiento temporal',
        'Sitio o centro de acopio final',
        'Observaciones',
    ]

    datos_tabla = [
        'Plástico',
        'Cartón',
        'Madera',
        'Metal',
        'Orgánicos',
        'Residuos domésticos',
        'Otros (Especificar)',
        'Total',
    ]

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla12b = doc.add_table(rows=9, cols=len(encabezados), style='Table Grid')

    for i, cols in enumerate(encabezados):
        cell = tabla12b.cell(0, i)
        t12b = cell.paragraphs[0].add_run(cols)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(16)
        t12b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#CFCFCF')

    for i, datos in enumerate(datos_tabla):
        cell = tabla12b.cell(i + 1, 0)
        t12b = cell.paragraphs[0].add_run(datos)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(16)

    for i in range(3):
        cell = tabla12b.cell(8, i + 2)
        quitar_bordes_celda(cell)

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        -------------------------------------
            Firmas de los responsables
        -------------------------------------
    """
    tabla12b = doc.add_table(rows=3, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Elaboró')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 0)
    t12b = cell.paragraphs[0].add_run('Supervisor de Obra')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Aprobó')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 1)
    t12b = cell.paragraphs[0].add_run('Supervisor Ambiental')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            BITÁCORA PARA CONTROL DE AGUA RESIDUAL
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('BITÁCORA PARA CONTROL DE AGUA RESIDUAL')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Datos de la bitacora
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=2, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Contratista: ____________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Fecha: _______________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Nombre de personal control de aguas residuales: ___________________________________________________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)

    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'Fecha',
        'Fuente de Trabajo',
        'Sistema de captación',
        'Estatus',
        'Volumen extraído (lt)',
        'Volumen dispueseto (lt)',
        'Sitio de disposición',
        'Observaciones',
    ]

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Arial'
    descripcionCapitulo12.font.size = Pt(12)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla12b = doc.add_table(rows=7, cols=len(encabezados), style='Table Grid')

    for i, cols in enumerate(encabezados):
        cell = tabla12b.cell(0, i)
        t12b = cell.paragraphs[0].add_run(cols)
        t12b.font.name = 'Agency FB'
        t12b.font.size = Pt(16)
        t12b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, '#CFCFCF')

    for i in range(2):
        cell = tabla12b.cell(6, i)
        quitar_bordes_celda(cell)

    for i in range(3):
        cell = tabla12b.cell(6, i + 5)
        quitar_bordes_celda(cell)

    # ✅ Celdas fusionadas
    row = tabla12b.rows[6]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Volumen total mensual')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)
    t12b.bold = True

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run(' ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(12)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    """
        -------------------------------------
            Firmas de los responsables
        -------------------------------------
    """
    tabla12b = doc.add_table(rows=3, cols=2, style='Table Grid')
    quitar_bordes_tabla(tabla12b)

    cell = tabla12b.cell(0, 0)
    t12b = cell.paragraphs[0].add_run('Elaboró')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 0)
    t12b = cell.paragraphs[0].add_run('\n_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 0)
    t12b = cell.paragraphs[0].add_run('Supervisor de Obra')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(0, 1)
    t12b = cell.paragraphs[0].add_run('Aprobó')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 1)
    t12b = cell.paragraphs[0].add_run('\n_________________________________')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(2, 1)
    t12b = cell.paragraphs[0].add_run('Supervisor Ambiental')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(16)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nNOTA:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    descripcionCapitulo12 = di12.add_run(' Estatus: A: servicio de limpieza y succión realizado B: servicio de limpieza no realizado')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('Observaciones:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    descripcionCapitulo12 = di12.add_run(' 1: caseta fuera de lugar, 2: Caseta dañada, 3: caseta obstruida, 4: sin material (sin papel).  ')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            BITÁCORAS DE RESIDUOS PELIGROSOS Y SITIOS CONTAMINADOS.\n
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    """
        =================================================================================
            Salto de pagina
        =================================================================================
    """
    doc.add_page_break() # Salto de página

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('BITÁCORAS DE RESIDUOS PELIGROSOS Y SITIOS CONTAMINADOS.\n'
                                        'Formato SEMARNAT-07-027-A')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 12.18 ###
    #########################
    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    tabla12b = doc.add_table(rows=6, cols=18, style='Table Grid')

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[12]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('GENERACIÓN')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)
    t12b.bold = True
    cell_background_color(merged_cell, '#AFAEAE')
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[13].merge(row.cells[13].merge(row.cells[14]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('ALMACENAMIENTO TEMPORAL\n'
                                             '\nArt. 71 fracción I inciso (d)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)
    t12b.bold = True
    cell_background_color(merged_cell, '#AFAEAE')
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[0]
    merged_cell = row.cells[15].merge(row.cells[15].merge(row.cells[17]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('MANEJO')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(14)
    t12b.bold = True
    cell_background_color(merged_cell, "#AFAEAE")
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ 2DA FILA

    # ✅ Celda fusionada
    cell_top = tabla12b.cell(1, 0)
    cell_bottom = tabla12b.cell(2, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('Nombre de los residuos peligrosos'
                                '\nArt. 71 fracción I inciso (a)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)

    # ✅ Celda fusionada
    cell_top = tabla12b.cell(1, 1)
    cell_bottom = tabla12b.cell(2, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    t12b = paragraph.add_run('Cantidad generada Ton.')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[11]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Características de peligrosidad del residuo – Código de peligrosidad de los residuos (CPR)'
                                            'Art. 71 fracción I inciso (b)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(1, 12)
    t12b = cell.paragraphs[0].add_run('Área o proceso de generación'
                                        'Art. 71 fracción I inciso (c)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(1, 13)
    t12b = cell.paragraphs[0].add_run('Fecha de ingreso')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(1, 14)
    t12b = cell.paragraphs[0].add_run('Fecha de salida')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(1, 15)
    t12b = cell.paragraphs[0].add_run('Fase de manejo siguiente a la salida del almacén\n'
                                        'Art. 71 fracción I inciso (e)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 2)
    t12b = cell.paragraphs[0].add_run('C')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 3)
    t12b = cell.paragraphs[0].add_run('R')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 4)
    t12b = cell.paragraphs[0].add_run('E')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 5)
    t12b = cell.paragraphs[0].add_run('T')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 6)
    t12b = cell.paragraphs[0].add_run('Te')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 7)
    t12b = cell.paragraphs[0].add_run('Th')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 8)
    t12b = cell.paragraphs[0].add_run('Tt')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 9)
    t12b = cell.paragraphs[0].add_run('T')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 10)
    t12b = cell.paragraphs[0].add_run('B')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 11)
    t12b = cell.paragraphs[0].add_run('M')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 16)
    t12b = cell.paragraphs[0].add_run('Nombre, denominación o razón social')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla12b.cell(2, 17)
    t12b = cell.paragraphs[0].add_run('Número de autorización')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[1]
    merged_cell = row.cells[16].merge(row.cells[16].merge(row.cells[17]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Prestador de servicio\n'
                                            'Art. 71 fracción I inciso ( f)')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla12b.cell(5, 0)
    t12b = cell.paragraphs[0].add_run('Total')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Celdas fusionadas
    row = tabla12b.rows[5]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[17]))

    # Agregar texto a la celda fusionada
    t12b = merged_cell.paragraphs[0].add_run('Nombre del responsable técnico de la bitácora')
    t12b.font.name = 'Agency FB'
    t12b.font.size = Pt(12)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    di12 = doc.add_paragraph()
    descripcionCapitulo12 = di12.add_run('\nNOTA:')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    descripcionCapitulo12.bold = True
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    descripcionCapitulo12 = di12.add_run(' C: Corrosivo, R: Reactivo, E: Explosivo, T: Toxico, Te: Tóxico ambiental, Th, Tóxico agudo, Tt: Tóxico crónico, I: Inflamable, B: Biológico Infeccioso, M: mezcla de residuo.')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.line_spacing = 1.15
    descripcionCapitulo12_format.space_after = 0
    descripcionCapitulo12_format.space_before = 0

    descripcionCapitulo12.font.name = 'Agency FB'
    descripcionCapitulo12.font.size = Pt(16)
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 12 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo12() # Crear el documento
