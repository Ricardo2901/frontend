"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes, celdas, etc; en pulgadas y en centrimetros
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo14():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 14
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XIV.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 14
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 14
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 14 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'XIV.- Análisis que demuestre que la biodiversidad de los ecosistemas que se verán afectados por el cambio de uso de suelo se mantenga.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 14 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run(
        "La actividad humana ha llevado a la degradación de los recursos naturales, debido al mal manejo que de ellos se ha hecho y a la intensidad con que se han explotado, de manera desproporcionada y sin algún programa de manejo que contemple medidas de mitigación, lo cual ha traído consigo la destrucción, desgaste y desperdicio de dichos recursos; por ello urge buscar principios que orienten el progreso tecnológico y la vida en armonía con la naturaleza, aprovechando los recursos naturales bajo un esquema de conservación, aunque se sabe, que todo proyecto causa impactos al medio ambiente, sobre todo al suelo, vegetación y fauna es por ello que se pretende minimizar estos impactos y causar el menor daño sobre ellos, sin destruirlos; aquí radica la importancia de un estudio técnico unificado de cambio de uso de suelo en terrenos forestales, modalidad “A” servirá, además como base para realizar las labores tendientes al cambio de uso de suelo en terrenos forestales; así como también las labores y avances de las mismas."
    )
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run(
        "Con el objeto de demostrar lo que establece el artículo 93, párrafo primero, de la Ley General de Desarrollo Forestal Sustentable, de cuyo cumplimiento depende la autorización de cambio de uso de suelo de terreno forestales solicitada, se evocó al estudio de la información y documentación que obra en el Documento Técnico Unificado considerando lo siguiente:"
    )
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run(
        "El artículo 93, párrafo primero, de la LGDFS, establece:"
    )
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run(
        "Artículo 93. La Secretaría autorizará el cambio de uso de suelo en terrenos forestales por excepción, previa opinión técnica de los miembros del Consejo Estatal Forestal de que se trate y con base en los estudios técnicos justificativos cuyo contenido se establecerá en el Reglamento, los cuales demuestren que la biodiversidad de los ecosistemas que se verán afectados se mantenga, y que la erosión de los suelos, el deterioro de la calidad del agua o la disminución en su captación se mitiguen en las áreas afectadas por la remoción de la vegetación forestal."
    )
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run(
        "A continuación, se da una justificación de la obra desde tres puntos de vista importantes."
    )
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.- Justificación Técnica')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 14.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('De acuerdo a las características del área en estudio que no presenta o reditúa un valor económico el promovente, opta por la ______________, ya que se dispone de ________________ y para satisfacer la demanda de la región, por lo que requiere de la autorización en materia de Cambio de Uso de Suelo y Manifestación de Impacto Ambiental en una superficie de _______ ha., con base en el estudio técnico justificativo se demuestra que no se compromete la biodiversidad a la vegetación presente y como consecuencia a los grupos de fauna silvestre asociada, las características más particulares del proyecto son los siguientes criterios:')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    justificacion_tecnica = [
        'La superficie total del Sistema Ambiental es de ________ hectáreas, de las cuales serán afectadas _____ ha. contempladas para llevar a cabo el ACUSTF lo cual tendrá una afectación del _____% con respecto al Sistema Ambiental.',
        'La vegetación a remover en el proyecto es de un total de _________ plantas en los ___ tipos de vegetación __________________________________________________ en el ACUSTF, siendo este un porcentaje muy inferior si se compara con el número total de plantas que estarían presente en el Sistema Ambiental teniendo un número de ______________ plantas, que representa una pérdida de cobertura del ______%.',
        'En cuanto a la fauna por la propia actividad del proyecto, con el paso de las personas y vehículos, se desplazan temporalmente y al no estar permitido su caza y captura por parte del personal de la empresa estos pueden desplazarse en forma libre en el Sistema Ambiental.',
    ]

    for lista in range(len(justificacion_tecnica)):
        di14 = doc.add_paragraph()
        descripcionCapitulo14 = di14.add_run(f'{1 + lista}.- {justificacion_tecnica[lista]}')
        descripcionCapitulo14_format = di14.paragraph_format
        descripcionCapitulo14_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo14.font.name = 'Arial'
        descripcionCapitulo14.font.size = Pt(12)
        di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.- Para Demostrar que no se compromete la biodiversidad en el área de cambio de uso de suelo y se mantenga dentro del sistema ambiental	')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.- Recurso Flora')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 14.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Comparativo de índices de biodiversidad de la vegetación dentro del ecosistema presente en el ACUSTF (_________________________________________________________) en comparación con el sistema ambiental, para ello se utilizaron los índices Menhinick y Simpson, donde se midió la biodiversidad de cinco estratos. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.1- ')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.1.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.1.1- índices de diversidad Estrato Arbustivo _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.1.-	Valores de diversidad (Estructura) Simpson estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato Arbustivo el área de cambio de uso de suelo presenta una Estructura ______________________________________________________________________________________________________________________________. Consecuente a ello las especies que se encuentran en ambas áreas son poco dominantes. Es decir, hay mayor Diversidad que dominancia.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.2.-	Valores de diversidad (estructura) Berger-Parker estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con grado bajo.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.3.-	Valores de diversidad (Riqueza) Margalef estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con _____________. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.4.-	Valores de riqueza Menhinick estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________________, por lo que, con la implementación de las actividades, el grado de afectación al sistema ambiental se considera ____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.5.-	Valores de diversidad (Equidad) Shannon estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.6.-	Valores de Equidad pielou estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo que, con la implementación de las actividades, el grado de afectación al sistema ambiental se considera ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.1.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.1.2- índices de diversidad Estrato Gramíneo _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.7.- Valores de diversidad Simpson estrato gramíneo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________, por lo tanto, con las actividades que implica el proyecto la afectación a la diversidad del sistema ambienta se considera de _______________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.8.- Valores de diversidad (estructura) Berger-Parker estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con _________________. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.9.-	Valores de riqueza Menhinick estrato gramíneo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental es considera con _______________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.10.- Valores de riqueza Margalef estrato gramíneo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental es considera con _________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.11.- Valores de Equidad Shannon estrato gramíneo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato __________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental es considera con ____________. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.12.- Valores de Equidad Pielou estrato gramíneo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental es considera con _____________ .')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.1.3 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.1.3- índices de diversidad Estrato Herbáceo _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.13.- Valores de diversidad (Estructura) Simpson estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________________, por lo tanto, con las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se considera de ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.14.- Valores de diversidad (Estructura) Berger Parker estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ____________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.15.- Valores de riqueza Menhinick estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.16.- Valores de riqueza Margalef estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de _____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.17.- Valores de Equidad Shannon estrato herbáceo MDR _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ______________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.18.- Valores de Equidad Pielou estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de __________________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.1.4 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.1.4- índices de diversidad Estrato Suculento _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.19.- Valores de diversidad Simpson estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo tanto se concluye que este estrato está muy equilibrado en ambas áreas y con las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se considera de impacto _____.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.20.- Valores de Dominancia Berger-Parker estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ____________________________________________ por lo que el grado de afectación al sistema ambiental se considera ____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.21.- Valores de riqueza Menhinick estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________ por lo que el grado de afectación al sistema ambiental se considera __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.22.- Valores de riqueza Margalef estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________________________ por lo que el grado de afectación al sistema ambiental se considera _________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.23.- Valores de Equidad Shannon estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ________________________________________ por lo que el grado de afectación al sistema ambiental se considera __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.24.- Valores de Equidad Pielou estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.1.4 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________________ por lo que el grado de afectación al sistema ambiental se considera _________________. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            CONCLUSION DE LA VEGETACION
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Conclusion de la Vegetacion =)')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    descripcionCapitulo14.bold = True
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Describir la Conclusion de la Vegetacion =)')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 14.1.1.1.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.2- ')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.2.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.2.1- índices de diversidad Estrato Arbustivo _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.25.-	Valores de diversidad (Estructura) Simpson estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ________________________________________________________. Por ser muy similares. Consecuente a ello las especies que se encuentran en ambas áreas son poco dominantes. Es decir, hay mayor Diversidad que dominancia.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.26.-	Valores de diversidad (estructura) Berger-Parker estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.27.-	Valores de diversidad (Riqueza) Margalef estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con _____________. ')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.28.- Valores de riqueza Menhinick estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ______________________________________________________, por lo que, con la implementación de las actividades, el grado de afectación al sistema ambiental se considera ____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.29.- Valores de diversidad (Equidad) Shannon estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________________, por lo que, con la implementación de las actividades, grado de afectación al sistema ambiental ser considera con ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.30.- Valores de Equidad pielou estrato arbustivo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.1 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _______________________________________________, por lo que, con la implementación de las actividades, el grado de afectación al sistema ambiental se considera ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.2.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.2.2- índices de diversidad Estrato Herbáceo _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.31.- Valores de diversidad Simpson estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo tanto, con las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se considera de __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.32.- Valores de riqueza Berger - Parker estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ___________________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.33.- Valores de riqueza Margalef estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de _____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.34.- Valores de riqueza Menhinick estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _____________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.35.- Valores de Equidad Shannon estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ___________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.36.- Valores de riqueza Pielou estrato herbáceo _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.2 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo tanto, las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se consideran de ____________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.1.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.1.2.3 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.1.2.3- índices de diversidad Estrato Suculento _____')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.37.- Valores de diversidad Simpson estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato __________________________________________________________, por lo tanto, por lo tanto, se concluye que este estrato está muy equilibrado en ambas áreas y con las actividades que implica el proyecto la afectación a la diversidad del sistema ambiental se considera de impacto ____.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.38.- Valores de riqueza Berger - Parker estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________________________, por lo que el impacto de las actividades del proyecto se considera __________.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.39.- Valores de riqueza Menhinick estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________ por lo que el grado de afectación al sistema ambiental ser considera _________. Debido a que es similar las dos áreas en comparación.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.40.- Valores de riqueza Margalef estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________ por lo que el grado de afectación al sistema ambiental ser considera _________. Debido a que es similar las dos áreas en comparación.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.41.- Valores de Equidad estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato ____________________________________________ por lo que el grado de afectación al sistema ambiental ser considera ___________. Debido a que es similar las dos áreas en comparación.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.42.- Valores de Equidad Pielou estrato suculento _____')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.1.2.3 ###
    #########################
    filas = 3
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el estrato _________________________________________ por lo que el grado de afectación al sistema ambiental ser considera _________. Debido a que es similar las dos áreas en comparación.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            CONCLUSION DE LA VEGETACION
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 14.1.1.1.1.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Conclusion de la Vegetacion =)')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    descripcionCapitulo14.bold = True
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Describir la Conclusion de la Vegetacion =)')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 14.1.1.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'XIV.1.1.2.- Recurso Fauna.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 14.1.1.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('La biodiversidad se conoce como la riqueza o el número de especies en un área dada, y para medirla se pueden utilizar diferentes métodos de medición por niveles ya sea a nivel alfa, beta y gamma (Escalante, Tania; et al, 2002), y estos niveles se diferencian dependiendo de la superficie del área en la que se encuentran, para este caso la superficie del proyecto se encuentra determinada como un área de nivel local en la cual la riqueza de las especies que se pudieran presentar se considera ___________, la riqueza especifica es una forma sencilla de poder medir la biodiversidad, ya que esta se basa en el número de especies presentes sin tomar en cuenta el valor de importancia de las mismas de tal manera que para medir la biodiversidad de esta área se utilizaron dos índices de biodiversidad alfa tales como 1) el índice de Menhinick el cual mide la biodiversidad de un área, sus rangos de valoración va de menor o igual a 1 para valores bajos, valores menores o igual a 2 se consideran medios y valores de 2 en adelantes se consideran altos de biodiversidad, este índice se basa en la relación entre el número de especies y el número total de los individuos observados, y 2) el índice de Simpson el cual se utiliza también para cuantificar la biodiversidad de un hábitat. Sus rangos de valoración van de 0 a 1 para el índice de Simpson donde de 0 a 0.33 son para valores bajos; 0.34 a 0.66 para valores medios y mayores a 0.67 valores altos. En este apartado se plasmará un análisis comparativo entre sus valores por área de muestreo, además, se realizó un comparativo entre abundancias absolutas de las especies avistadas entre ambos sitios de muestreo.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.2.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.2.1.- Índice de Biodiversidad para el grupo de las aves.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.2.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.43.- Valores de biodiversidad para el grupo de las aves.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.2.1 ###
    #########################
    filas = 6
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.2.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el grupo de las aves en el _____________________________________________, por lo cual considerando ambos índices la afectación para este grupo el grado de afectación se considera ________, sin embargo, el porcentaje de desplazamiento que se tendrá en el área será de ___% en comparación con el área del sistema ambiental, por lo tanto, al efectuarse el cambio de uso de suelo la diversidad del grupo de las aves se mantendrá en el sistema ambiental.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.2.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.2.2.- Índice de Biodiversidad para el grupo de los mamíferos.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.2.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.44.- Valores de biodiversidad para el grupo de los mamíferos.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.2.2 ###
    #########################
    filas = 6
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.2.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el grupo de los mamíferos en ____________________________________________, sin embargo, el porcentaje de desplazamiento que tendrá este grupo será de ___% en comparación con el sistema ambiental, por lo tanto, al efectuarse el cambio de uso de suelo la diversidad del grupo de los mamíferos se mantendrá en el sistema ambiental.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.2.3 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.2.3.- Índice de Biodiversidad para el grupo de los reptiles.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.2.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.45.- Valores de biodiversidad para el grupo de los reptiles.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.2.3 ###
    #########################
    filas = 6
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.2.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el grupo de los reptiles en el _____________________________________________, sin embargo, el porcentaje de desplazamiento que tendrá este grupo será de ___% en comparación con el sistema ambiental, por lo tanto, al efectuarse el cambio de uso de suelo la diversidad del grupo de los reptiles se mantendrá en el sistema ambiental.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.2.4 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.2.4.- Índice de Biodiversidad para el grupo de los lepidópteros.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.2.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.46.- Valores de biodiversidad para el grupo de los lepidópteros.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.2.4 ###
    #########################
    filas = 6
    columnas = 6
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.2.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('\nPara el grupo de los lepidópteros se obtuvo que para el ________________________________________________________________________, sin embargo, el porcentaje de desplazamiento que tendrá este grupo será de _____% en comparación con el sistema ambiental, por lo tanto, al efectuarse el cambio de uso de suelo la diversidad del grupo de los lepidópteros se mantendrá en el sistema ambiental.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.3 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.3.- Análisis comparativo por áreas de estudio.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.3.1 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.3.1- Análisis comparativo para el grupo de las aves.')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.3.1 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.47.- Comparativo para el grupo de las aves')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.3.1 ###
    #########################
    filas = 10 #Va dependiendo de la base de datos
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.3.1 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Para el grupo de las aves en el ACUSTF cuenta con un total de __ especies presentes en el área y un numero de ___ individuos, la especie más abundante fue la especie __________________________ individuos observados. Para el área del sistema se presentan ___ especies y ___ individuos, la especie más representativa fue ________________ individuos avistados en la superficie de muestreo. Lo cual el porcentaje de desplazamiento que tendrán las especies en el ACUSTF hacia el área del sistema ambiental es de ___%.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.3.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.3.2.- Análisis comparativo para el grupo de los mamíferos.')
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.3.2 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.48.- Comparativo para el grupo de los mamíferos.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.3.2 ###
    #########################
    filas = 10 #Va dependiendo de la base de datos
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.3.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('Para el grupo de los mamíferos en el ACUSTF cuenta con un total de __ especies y __ individuos de los cuales la especie que presento una mayor abundancia fue _____________ con un total de __ individuos. Para el área del sistema ambiental cuenta con 6 especies presentes con un total de ___ individuos, la especie más abundante _________________ con __ individuos avistados. El porcentaje de desplazamiento que tendrán las especies del ACUSTF hacia el área del sistema ambiental es de __%.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.3.3 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.3.3.- Análisis comparativo para el grupo de los reptiles.')
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.3.3 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.49.- Comparativo para el grupo de los reptiles.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.3.3 ###
    #########################
    filas = 10 #Va dependiendo de la base de datos
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.3.3 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('El grupo de los reptiles en el ACUSTF cuenta con un numero de __ especies presentes en el área con un total de __ individuos, siendo la especie más representativa ______________ con __ individuos; para el área del sistema ambiental se registraron __ especies con un total de ___ individuos siendo la especie más representativa __________________ con __ individuos, para este grupo el porcentaje de desplazamiento es de ___%, por lo cual las especies que se encuentran dentro del área ACUSTF se podrán desplazar hacia el área del sistema ambiental.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.1.1.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 14.1.1.3.4 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.1.1.3.4.- Análisis comparativo para el grupo de los Lepidópteros.')
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 14.1.1.3.4 ###
    #########################
    tituloTabla14b = doc.add_paragraph()
    dti14b = tituloTabla14b.add_run('\nTabla 14.50.- Comparativo para el grupo de los lepidópteros.')
    dti14b_format = tituloTabla14b.paragraph_format
    dti14b_format.line_spacing = 1.15
    dti14b_format.space_after = 0

    dti14b.font.name = 'Bookman Old Style'
    dti14b.font.size = Pt(12)
    tituloTabla14b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 14.1.1.3.4 ###
    #########################
    filas = 10 #Va dependiendo de la base de datos
    columnas = 5
    tabla14b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla14b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(filas):
            cell = tabla14b.cell(rows, cols)
            t13b = cell.paragraphs[0].add_run(' ')
            t13b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 14.1.1.3.4 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('El grupo de los lepidópteros en el ACUSTF cuenta para el ACUSTF un total de __ especies con un total de ___ individuos registrados, la especie más representativa fue _________________ con ___ individuos; para el área del sistema ambiental se presentaron __ especies con un total de ___ individuos, la especie más representativa fue _______________ con ___ individuos, el porcentaje de desplazamiento que tendrá este grupo será del _____ % por lo cual las especies que se encuentren en el ACUSTF se podrán desplazar hacia el área del sistema ambiental sin ningún inconveniente.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 14.2
    ########################################################################################################################################################################

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            CONCLUSION DEL CAPITULO 14
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """

    #########################
    ### Titulo del capitulo 14.2 ###
    #########################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(f'\nXIV.2.- Conclusiones')
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True
    capitulo14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 14.2 ###
    #########################
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('En cuanto a la flora se menciona que los estratos: ________________________________, representan un ______________. Solamente resaltar que en el estrato ____________ presenta una dominancia _______ debido a las ________ especies que se encontraron y al tener mayor cobertura las hace dominantes de todos los estratos estudiados, lo cual no se verán afectados debido a que presentaron índices iguales del ACUSTF y SA. Por último, la remoción de la vegetación en todos los estratos es de _____%, la cual se podrá compensar con el rescate y reubicación de las especies principalmente Cactáceas y Asparagaceas.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('En cuanto a la fauna Silvestre se puede mencionar que para el grupo de las aves y mamíferos la diversidad fue más representativa obteniendo valores ___________ en comparación con el ajuste que sus valores fueron ________. Por otra parte, para el grupo de los reptiles en ambas áreas la diversidad de este grupo fue __________, para el grupo de los lepidópteros se registró una riqueza de especies más representativa en el ACUSTF con valores ____________. De tal manera y en base a los resultados obtenidos se puede determinar que en que en ambas áreas cuentan con biodiversidad en valores ____________ siendo aún más representativa para el sistema ambiental por lo cual no es equiparable a que se tenga una afectación mayor ya que las especies serán rescatadas y reubicadas para su conservación, realizando una acción de desplazamiento de especies entre áreas de muestreo.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 14 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo14() # Crear el documento
