"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas

def capitulo3():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 3
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo III.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True                          # Negrita

    temasCapitulo3 = [
        'III.- UBICACIÓN Y SUPERFICIE DEL PREDIO O CONJUNTO DE PREDIOS Y DELIMITACIÓN DE LA PORCIÓN DE DONDE SE PRETENDE REALIZAR EL CAMBIO DE USO DE SUELO A TRAVÉS DE PLANOS GEORREFERENCIADOS.', [
            "III.1.- Ubicación del Predio o Conjunto de Predios Donde se Ubica el Proyecto.", [
                'III.1.1.- Datos Generales de Ubicación:',
                'III.1.2.- Itinerario de acceso:',
                'III.1.3.- Superficie del predio:',
                'III.1.4.- Situación legal del predio.', [
                    'III.1.4.1- Colindancias del predio solicitado para cambio de uso de suelo.'
                ],
                'III.1.5.- Situación legal del promovente.', [
                    'III.1.5.1.- Promovente',
                    'III.1.5.2.- Representante Legal'
                ]
            ],
            'III.2. Ubicación y Delimitación Física de la Superficie del Proyecto.', [
                'III.2.1.- Ubicación y Delimitación Física de la Superficie del Proyecto.',
                'III.2.2.- Ubicación y Delimitación Física de la Superficie del Proyecto.'
            ],
            'III.3. Representación Gráfica de la Ubicación Geográfica y Geopolítica.',
            'III.4. Indicar si el Proyecto se Ubica dentro de alguna Modalidad de Área Natural Protegida (ANP).', [
                'III.4.1.- Referente al Área Natural Protegida.',
                'III.4.2.- El Predio se encuentra dentro de alguna AICA.',
                'III.4.3.- El Predio está dentro de alguna RHP.',
                'III.4.4.- El predio está dentro de una RTP.',
                'III.4.5.- El predio está dentro de un sitio RAMSAR.'
            ]
        ]
    ]

    ########################################################################################################################################################################
    # Capitulo 3.1
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 3 ###
    #########################
    capitulo3 = doc.add_paragraph()
    i3 = capitulo3.add_run(f'{temasCapitulo3[0]}')  # Añadir un salto de línea después del título
    i3_format = capitulo3.paragraph_format
    i3_format.line_spacing = 1.5

    i3.font.name = 'Arial'
    i3.font.size = Pt(12)
    i3.font.bold = True
    capitulo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Titulo del capitulo 3.1 ###
    #########################
    capitulo31 = doc.add_paragraph()
    i31 = capitulo31.add_run(f'{temasCapitulo3[1][0]}')  # Añadir un salto de línea después del título
    print(f'{temasCapitulo3[1][0]}')  # Imprimir el título en la consola
    i31_format = capitulo31.paragraph_format
    i31_format.line_spacing = 1.5

    i31.font.name = 'Arial'
    i31.font.size = Pt(12)
    i31.font.bold = True

    #########################
    ### Descripcion del capitulo 3.1 ###
    #########################
    di31 = doc.add_paragraph()
    descripcionCapitulo31 = di31.add_run('El Proyecto “________________”, se encuentra ubicado en el ________________________________, cuya superficie para realizar el cambio de uso de suelo es de ____ ha., como se manifiesta en el Anexo mapa No. 3-1. Ubicación del área en estudio.')
    descripcionCapitulo31_format = di31.paragraph_format
    descripcionCapitulo31_format.line_spacing = 1.5

    descripcionCapitulo31.font.name = 'Arial'
    descripcionCapitulo31.font.size = Pt(12)
    di31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di31 = doc.add_paragraph()
    descripcionCapitulo31 = di31.add_run('En apego a la legislación vigente y considerando lo referente al Reglamento de la Ley General de Desarrollo Forestal Sustentable, publicado en el Diario Oficial de La Federación el 09 de diciembre de 2020, definiendo en el artículo 2 en su fracción XIX  como Planos georreferenciados, aquél que se presenta en coordenadas UTM o geográficas, con precisión a décimas de segundo de cada punto de la poligonal de los predios, ubicándolos dentro de su respectiva cuenca y subcuenca hidrológico-forestal, con una escala mínima de 1:50,000, a fin de identificar su localización por entidad federativa y municipio.')
    descripcionCapitulo31_format = di31.paragraph_format
    descripcionCapitulo31_format.line_spacing = 1.5

    descripcionCapitulo31.font.name = 'Arial'
    descripcionCapitulo31.font.size = Pt(12)
    di31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    #   Capitulo 3.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.1 ###
    #########################
    capitulo311 = doc.add_paragraph()
    i311 = capitulo311.add_run(f'\n{temasCapitulo3[1][1][0]}')  # Añadir un salto de línea después del título
    print(f'{temasCapitulo3[1][0]}')  # Imprimir el título en la consola
    i311_format = capitulo311.paragraph_format
    i311_format.line_spacing = 1.5

    i311.font.name = 'Arial'
    i311.font.size = Pt(12)
    i311.font.bold = True

    ### Lista del capitulo 3.1.1 ###
    listaCapitulo311 = [
        'Estado:',
        'Región:',
        'Municipio:',
        'Localidad o Ejido:',
    ]

    #########################
    ### Descripcion del capitulo 3.1.1 ###
    #########################

    for lista in listaCapitulo311:
        di311 = doc.add_paragraph(style='ListBullet')
        descripcionCapitulo311 = di311.add_run(f'{lista}')
        descripcionCapitulo311_format = di311.paragraph_format
        descripcionCapitulo311_format.line_spacing = 1.5

        descripcionCapitulo311.font.name = 'Arial'
        descripcionCapitulo311.font.size = Pt(12)

    ########################################################################################################################################################################
    #   Capitulo 3.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.2 ###
    #########################
    capitulo312 = doc.add_paragraph()
    i312 = capitulo312.add_run(f'\n{temasCapitulo3[1][1][1]}')  # Añadir un salto de línea después del título
    i312_format = capitulo312.paragraph_format
    i312_format.line_spacing = 1.5

    i312.font.name = 'Arial'
    i312.font.size = Pt(12)
    i312.font.bold = True

    #########################
    ### Descripcion del capitulo 3.1.2 ###
    #########################
    di312 = doc.add_paragraph()
    descripcionCapitulo312 = di312.add_run('El área donde se pretende _____________________, se encuentra en el municipio de ______________, Coahuila. Para llegar al área en estudio es _______________________....._________....._________....___________..........______......_______.....________....._______ donde se encuentra el área de estudio.')
    descripcionCapitulo312_format = di312.paragraph_format
    descripcionCapitulo312_format.line_spacing = 1.5

    descripcionCapitulo312.font.name = 'Arial'
    descripcionCapitulo312.font.size = Pt(12)
    di312.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.1.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo312 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo312.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo312.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo312.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo312.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.1.2 ###
    #########################
    diMap312 = doc.add_paragraph()
    descripcionCapituloMapa312 = diMap312.add_run('Mapa 3.1.-	Ubicación de acceso al área de estudio.')
    descripcionCapituloMapa312_format = diMap312.paragraph_format
    descripcionCapituloMapa312_format.line_spacing = 1.15
    descripcionCapituloMapa312.space_before = Pt(0)  # Espacio después del texto

    descripcionCapituloMapa312.font.name = 'Bookman Old Style'
    descripcionCapituloMapa312.font.size = Pt(12)
    descripcionCapituloMapa312.font.italic = True
    diMap312.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.3 ###
    #########################
    capitulo313 = doc.add_paragraph()
    i313 = capitulo313.add_run(f'\n{temasCapitulo3[1][1][2]}')  # Añadir un salto de línea después del título
    i313_format = capitulo313.paragraph_format
    i313_format.line_spacing = 1.5

    i313.font.name = 'Arial'
    i313.font.size = Pt(12)
    i313.font.bold = True

    #########################
    ### Descripcion de la Tabla del capitulo 3.1.3 ###
    #########################
    tituloTabla313 = doc.add_paragraph()
    i313 = tituloTabla313.add_run('Tabla 3.1.- Distribución de Superficies')
    i313_format = tituloTabla313.paragraph_format
    i313_format.line_spacing = 1.15

    i313.font.name = 'Bookman Old Style'
    i313.font.size = Pt(12)
    i313.font.italic = True
    tituloTabla313.alignment = WD_ALIGN_PARAGRAPH.CENTER


    #########################
    ### Tabla del capitulo 3.1.3 ###
    #########################
    tabla313 = doc.add_table(rows=4, cols=3, style='Table Grid')  # Crear una tabla con 4 filas y 3 columnas

    columnasTabla313 = [
        'DISTRIBUCION DE SUPERFICIES',
        'SUPERFICIE (ha)',
        '%'
    ]

    filasTabla313 = [
        'Superficie Ejidal',
        'Superficie de Contrato',
        'Superficie sujeta a ACUSTF'
    ]    

    colt313 = range(len(columnasTabla313))  # Obtener el rango de las columnas
    rowt313 = range(len(filasTabla313))  # Obtener el rango de las filas


    for columnas in colt313:  # Iterar sobre las columnas
        cell = tabla313.cell(0, columnas)  # Obtener la celda de la primera fila
        t313 = cell.paragraphs[0].add_run(f'{columnasTabla313[columnas]}')  # Añadir texto a la celda
        t313.font.name = 'Arial'  # Tipo de letra
        t313.font.size = Pt(12)  # Tamaño de la letra
        t313.font.bold = True  # Negrita
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro
        cell_background_color(cell, 'A6A6A6')  # Cambiar el color de fondo de la celda a amarillo

    for filas in rowt313:  # Iterar sobre las filas
        cell = tabla313.cell(filas + 1, 0)  # Obtener la celda de la primera columna
        t313 = cell.paragraphs[0].add_run(f'{filasTabla313[filas]}')  # Añadir texto a la celda
        t313.font.name = 'Arial'  # Tipo de letra
        t313.font.size = Pt(12)  # Tamaño de la letra
        t313.font.bold = True  # Negrita
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro

    for widthColumna1 in colt313:  # Iterar sobre las columnas de la primera columna
        cell = tabla313.cell(widthColumna1, 0)  # Obtener la celda de la primera columna
        cell.width = Cm(8.16)  # Ancho de la celda en centimetros

    dt313 = doc.add_paragraph()
    descripcionTabla313 = dt313.add_run('* Segun ADDATE')
    descripcionTabla313_format = dt313.paragraph_format
    descripcionTabla313_format.space_before = Pt(0)
    descripcionTabla313_format.space_after = Pt(0)

    descripcionTabla313.font.name = 'Arial'
    descripcionTabla313.font.size = Pt(10.5)

    dt313 = doc.add_paragraph()
    descripcionTabla313 = dt313.add_run('* Segun contrato de arrendamiento')
    descripcionTabla313.font.name = 'Arial'
    descripcionTabla313.font.size = Pt(10.5)


    ########################################################################################################################################################################
    #   Capitulo 3.1.4
    ########################################################################################################################################################################
    
     #########################
    ### Titulo del capitulo 3.1.4 ###
    #########################
    capitulo314 = doc.add_paragraph()
    i314 = capitulo314.add_run(f'\n{temasCapitulo3[1][1][3]}')
    i314_format = capitulo314.paragraph_format
    i314_format.line_spacing = 1.5

    i314.font.name = 'Arial'
    i314.font.size = Pt(12)
    i314.font.bold = True

     #########################
    ### Descripcion del Capitulo del capitulo 3.1.3 ###
    #########################
    di314 = doc.add_paragraph()
    descripcionCapitulo314 = di314.add_run('Describir la situacion legal de predio (Anexo 2).')
    descripcionCapitulo314_format = di314.paragraph_format
    descripcionCapitulo314_format.line_spacing = 1.5

    descripcionCapitulo314.font.name = 'Arial'
    descripcionCapitulo314.font.size = Pt(12)
    di314.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    #   Capitulo 3.1.4.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.4.1 ###
    #########################
    capitulo3141 = doc.add_paragraph()
    i3141 = capitulo3141.add_run(f'\n{temasCapitulo3[1][1][4][0]}')
    i3141_format = capitulo3141.paragraph_format
    i3141_format.line_spacing = 1.5

    i3141.font.name = 'Arial'
    i3141.font.size = Pt(12)
    i3141.font.bold = True

    #########################
    ### Titulo de la Tabla del capitulo 3.1.4 ###
    #########################
    tituloTabla3141 = doc.add_paragraph()
    i3141 = tituloTabla3141.add_run('Tabla 3.2.- Colindancia del ejido y área de arrendamiento')
    i3141_format = tituloTabla3141.paragraph_format
    i3141_format.line_spacing = 1.15

    i3141.font.name = 'Bookman Old Style'
    i3141.font.size = Pt(12)
    i3141.font.italic = True
    tituloTabla3141.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 3.1.4 ###
    #########################
    tabla3141 = doc.add_table(rows=10, cols=4, style='Table Grid')
    
    #########################
    # Celda fusionada "De acuerdo al ADATTE Ampara una superficie de 740-92-61.45 Has"
    row1 = tabla3141.rows[0]
    merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[3]))

    # Agregar texto a la celda fusionada
    t3141 = merged_cell1.paragraphs[0].add_run('De acuerdo al ADATTE Ampara una superficie de 740-92-61.45 Has')
    t3141.font.name = 'Arial'
    t3141.font.size = Pt(12)
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    #########################
    # Columnas de la Tabla del Capitulo 3.1.4.1
    columnasTabla3141 = [
        'SUPERFICIE TOTAL, PREDIO ha',
        'SUPERFICIE ACUSTF ha',
        'RUMBOS Y DISTANCIAS',
        'COLINDANCIA'
    ]

    rumbosTabla3141 = [
        'Norte',
        'Sur',
        'Este',
        'Oeste',
    ]

    colt3141 = range(len(columnasTabla3141))
    rumbos3141 = range(len(rumbosTabla3141))

    for columnas in colt3141:   #Crea las columnas de la tabla de forma directa sin escribirlas manualmente
        cell = tabla3141.cell(1, columnas)
        t3141 = cell.paragraphs[0].add_run(f'{columnasTabla3141[columnas]}')
        t3141.font.name = 'Arial'
        t3141.font.size = Pt(12)
        t3141.bold = True
        tabla3141.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, 'BFBFBF')

    
    #########################
    # Celdas fusionadas
    # SUPERFICIE TOTAL, PREDIO ha
    for superficiesHa in range(2):
        fila_inicio = 2 + (superficiesHa * 4)
        fila_final = 3 + fila_inicio
        cell_top = tabla3141.cell(fila_inicio, 0)
        cell_bottom = tabla3141.cell(fila_final, 0)

        merged_cell = cell_top.merge(cell_bottom)

        # Agregar texto (opcional)
        paragraph = merged_cell.paragraphs[0]
        run = paragraph.add_run('Texto fusionado en filas 2 y 3')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Celdas fusionadas
    # SUPERFICIE ACUSTF ha
    for superficies in range(2):
        fila_inicio = 2 + (superficies* 4)
        fila_final = 3 + fila_inicio

        cell_top = tabla3141.cell(fila_inicio, 1)
        cell_bottom = tabla3141.cell(fila_final, 1)

        merged_cell = cell_top.merge(cell_bottom)

        # Agregar texto (opcional)
        paragraph = merged_cell.paragraphs[0]
        run = paragraph.add_run('Texto fusionado en filas 2 y 3')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Celdas combinadas
    # RUMBOS Y DISTANCIAS
    for i in range(2):
        for rumbos in rumbos3141:
            celda = (i * 3) + rumbos + (i + 2)
            cell = tabla3141.cell(celda, 2)
            t3141 = cell.paragraphs[0].add_run(f'{rumbosTabla3141[rumbos]}')
            t3141.font.name = 'Arial'
            t3141.font.size = Pt(12)

    #########################
    # Celdas combinadas
    # COLINDANCIAS
    for x in range(2):
        for colindancia in range(4):
            celda = (x * 3) + colindancia + (x + 2)
            cell = tabla3141.cell(celda, 3)
            t3141 = cell.paragraphs[0].add_run(f'Colindancia {colindancia + 1}')
            t3141.font.name = 'Arial'
            t3141.font.size = Pt(12)

    ########################################################################################################################################################################
    #   Capitulo 3.1.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.5 ###
    #########################
    capitulo315 = doc.add_paragraph()
    i315 = capitulo315.add_run(f'\n{temasCapitulo3[1][1][6][0]}')  # Añadir un salto de línea después del título
    i315_format = capitulo315.paragraph_format
    i315_format.line_spacing = 1.5

    i315.font.name = 'Arial'
    i315.font.size = Pt(12)
    i315.font.bold = True

    #########################
    ### Descripcion del capitulo 3.1.5 ###
    #########################
    di315 = doc.add_paragraph()
    descripcionCapitulo315 = di315.add_run(f'Describir la situacion legal del promovente')
    descripcionCapitulo315_format = di315.paragraph_format
    descripcionCapitulo315_format.line_spacing = 1.5

    descripcionCapitulo315.font.name = 'Arial'
    descripcionCapitulo315.font.size = Pt(12)
    di315.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    #   Capitulo 3.1.5.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.5.1 ###
    #########################
    capitulo3151 = doc.add_paragraph()
    i3151 = capitulo3151.add_run(f'\n{temasCapitulo3[1][1][6][0]}')  # Añadir un salto de línea después del título
    i3151_format = capitulo3151.paragraph_format
    i3151_format.line_spacing = 1.5

    i3151.font.name = 'Arial'
    i3151.font.size = Pt(12)
    i3151.font.bold = True

    #########################
    ### Descripcion del capitulo 3.1.5.1 ###
    #########################
    di3151 = doc.add_paragraph()
    descripcionCapitulo3151 = di3151.add_run(f'Nombre del promovente')
    descripcionCapitulo3151_format = di3151.paragraph_format
    descripcionCapitulo3151_format.line_spacing = 1.5

    descripcionCapitulo3151.font.name = 'Arial'
    descripcionCapitulo3151.font.size = Pt(12)
    di3151.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    #   Capitulo 3.1.5.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.1.5.2 ###
    #########################
    capitulo3152 = doc.add_paragraph()
    i3152 = capitulo3152.add_run(f'\n{temasCapitulo3[1][1][6][1]}')  # Añadir un salto de línea después del título
    i3152_format = capitulo3152.paragraph_format
    i3152_format.line_spacing = 1.5

    i3152.font.name = 'Arial'
    i3152.font.size = Pt(12)
    i3152.font.bold = True

    #########################
    ### Descripcion del capitulo 3.1.5.2 ###
    #########################
    di3152 = doc.add_paragraph()
    descripcionCapitulo3152 = di3152.add_run(f'Nombre del representante legal del promovente')
    descripcionCapitulo3152_format = di3152.paragraph_format
    descripcionCapitulo3152_format.line_spacing = 1.5
    descripcionCapitulo3152_format.space_after = 0

    descripcionCapitulo3152.font.name = 'Arial'
    descripcionCapitulo3152.font.size = Pt(12)
    descripcionCapitulo3152.bold = True
    di3152.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di3152 = doc.add_paragraph()
    descripcionCapitulo3152 = di3152.add_run(f'Descripcion del representante legal del promovente')
    descripcionCapitulo3152_format = di3152.paragraph_format
    descripcionCapitulo3152_format.line_spacing = 1.15

    descripcionCapitulo3152.font.name = 'Arial'
    descripcionCapitulo3152.font.size = Pt(12)
    di3152.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    #  Capitulo 3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.2 ###
    #########################
    capitulo32 = doc.add_paragraph()
    i32 = capitulo32.add_run(f'\n{temasCapitulo3[1][2]}')
    i32_format = capitulo32.paragraph_format
    i32_format.line_spacing = 1.5

    i32.font.name = 'Arial'
    i32.font.size = Pt(12)
    i32.font.bold = True

    #########################
    ### Descripcion del capitulo 3.2 ###
    #########################
    di32 = doc.add_paragraph()
    descripcionCapitulo32 = di32.add_run('Localización del área del Proyecto con Coordenadas UTM con Datum WGS 84. Zona 14N.')
    descripcionCapitulo32_format = di32.paragraph_format
    descripcionCapitulo32_format.line_spacing = 1.5

    descripcionCapitulo32.font.name = 'Arial'
    descripcionCapitulo32.font.size = Pt(12)
    di32.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capitulo 3.2 ###
    #########################
    tabla32 = doc.add_table(rows=11, cols=3, style='Table Grid')

    cell = tabla32.cell(0, 0)
    t32 = cell.paragraphs[0].add_run('Vértices')
    t32.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla32.cell(0, 1)
    t32 = cell.paragraphs[0].add_run('X')
    t32.font.name = 'Arial'

    cell = tabla32.cell(0, 2)
    t32 = cell.paragraphs[0].add_run('Y')
    t32.font.name = 'Arial'

    for vertices in range(10):
        cell = tabla32.cell(vertices + 1, 0)
        t32 = cell.paragraphs[0].add_run(f'{vertices + 1}')
        t32.font.name = 'Arial'
        t32.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for x in range(11):
        cell = tabla32.cell(x, 1)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t32 = cell.paragraphs[0].add_run(' ')
        t32.font.name = 'Arial'
        t32.font.size = Pt(12)
        
    for y in range(11):
        cell = tabla32.cell(y, 2)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t32 = cell.paragraphs[0].add_run(' ')
        t32.font.name = 'Arial'
        t32.font.size = Pt(12)

    for verticesWidth in range(11):
        cell = tabla32.cell(verticesWidth, 0)
        cell.width = Cm(1.78)
    
    for xWidth in range(11):
        cell = tabla32.cell(xWidth, 1)
        cell.width = Cm(4.05)

    for yWidth in range(11):
        cell = tabla32.cell(yWidth, 2)
        cell.width = Cm(4.05)

    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla32.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Mapa del capitulo 3.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo32 = doc.add_paragraph()
    imagenCapitulo32.text = ''
    imagenCapitulo32 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo32.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo32.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo32.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo32.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2 ###
    #########################
    diMap33 = doc.add_paragraph()
    descripcionCapituloMapa33 = diMap33.add_run('Mapa 3.2.-	Ubicación de vertices.')
    descripcionCapituloMapa33_format = diMap33.paragraph_format
    descripcionCapituloMapa33_format.line_spacing = 1.15
    descripcionCapituloMapa33.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa33.font.name = 'Bookman Old Style'
    descripcionCapituloMapa33.font.size = Pt(12)
    descripcionCapituloMapa33.font.italic = True
    diMap33.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.2.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.2.1 ###
    #########################
    capitulo321 = doc.add_paragraph()
    i321 = capitulo321.add_run(f'\n{temasCapitulo3[1][3][0]}')
    i321_format = capitulo321.paragraph_format
    i321_format.line_spacing = 1.5

    i321.font.name = 'Arial'
    i321.font.size = Pt(12)
    i321.font.bold = True

    #########################
    ### Descripcion del capitulo 3.2.1 ###
    #########################
    di321 = doc.add_paragraph()
    descripcionCapitulo32 = di321.add_run('Localización del área del Proyecto con Coordenadas UTM con Datum WGS 84. Zona 14N.')
    descripcionCapitulo32_format = di321.paragraph_format
    descripcionCapitulo32_format.line_spacing = 1.15

    descripcionCapitulo32.font.name = 'Arial'
    descripcionCapitulo32.font.size = Pt(12)
    di312.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion de la Tabla del capitulo 3.2.1 ###
    #########################
    tituloTabla33 = doc.add_paragraph()
    i33 = tituloTabla33.add_run('Tabla 3.4.- Coordenadas de las Parcelas')
    i33_format = tituloTabla33.paragraph_format
    i33_format.line_spacing = 1.15

    i33.font.name = 'Bookman Old Style'
    i33.font.size = Pt(12)
    i33.font.italic = True
    tituloTabla33.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 3.2.1 ###
    #########################

    tabla33 = doc.add_table(rows=13, cols=5, style='Table Grid')

    for columnas in range(5):
        cell = tabla33.cell(0, columnas)
        t33 = cell.paragraphs[0].add_run(f'Columna {columnas + 1}')
        t33.font.name = 'Arial'
        t33.font.size = Pt(12)
        t33.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, 'BFBFBF')

    for columna1 in range(3):
        fila_inicio = 2 + (columna1 * 4) - 1
        fila_final = 3 + fila_inicio
        cell_top = tabla33.cell(fila_inicio, 0)
        cell_bottom = tabla33.cell(fila_final, 0)

        merged_cell = cell_top.merge(cell_bottom)

        # Agregar texto (opcional)
        paragraph = merged_cell.paragraphs[0]
        run = paragraph.add_run(f'{columna1 + 1}')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for columna2 in range(3):
        fila_inicio = 2 + (columna2 * 4) - 1
        fila_final = 3 + fila_inicio
        cell_top = tabla33.cell(fila_inicio, 1)
        cell_bottom = tabla33.cell(fila_final, 1)

        merged_cell = cell_top.merge(cell_bottom)

        # Agregar texto (opcional)
        paragraph = merged_cell.paragraphs[0]
        run = paragraph.add_run(f'{columna2 + 1}')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Mapa del capitulo 3.2.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo321 = doc.add_paragraph()
    imagenCapitulo321.text = ''
    imagenCapitulo321 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo321.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo321.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo321.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo321.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2.1 ###
    #########################
    diMap321 = doc.add_paragraph()
    descripcionCapituloMapa321 = diMap321.add_run('Mapa 3.3.-	Ubicación de vertices de las parcelas.')
    descripcionCapituloMapa321_format = diMap321.paragraph_format
    descripcionCapituloMapa321_format.line_spacing = 1.15
    descripcionCapituloMapa321.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa321.font.name = 'Bookman Old Style'
    descripcionCapituloMapa321.font.size = Pt(12)
    descripcionCapituloMapa321.font.italic = True
    diMap321.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 3.2.2 ###
    #########################
    capitulo322 = doc.add_paragraph()
    i322 = capitulo322.add_run(f'\n{temasCapitulo3[1][3][1]}')
    i322_format = capitulo322.paragraph_format
    i322_format.line_spacing = 1.5

    i322.font.name = 'Arial'
    i322.font.size = Pt(12)
    i322.font.bold = True

    #########################
    ### Descripcion del capitulo 3.2.2 ###
    #########################
    di322 = doc.add_paragraph()
    descripcionCapitulo322 = di322.add_run('Localización del área del Proyecto con Coordenadas UTM con Datum WGS 84. Zona 14N.')
    descripcionCapitulo322_format = di322.paragraph_format
    descripcionCapitulo322_format.line_spacing = 1.15

    descripcionCapitulo322.font.name = 'Arial'
    descripcionCapitulo322.font.size = Pt(12)
    di322.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion de la Tabla del capitulo 3.2.2 ###
    #########################
    tituloTabla322 = doc.add_paragraph()
    i322 = tituloTabla322.add_run('Tabla 3.5.- Coordenadas del _____')
    i322_format = tituloTabla322.paragraph_format
    i322_format.line_spacing = 1.15

    i322.font.name = 'Bookman Old Style'
    i322.font.size = Pt(12)
    i322.font.italic = True
    tituloTabla322.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capitulo 3.2.2 ###
    #########################
    
    tabla322 = doc.add_table(rows=6, cols=3, style='Table Grid')

    for columnas in range(3):
        cell = tabla322.cell(0, columnas)
        t322 = cell.paragraphs[0].add_run(f'Columna {columnas + 1}')
        t322.font.name = 'Arial'
        t322.font.size = Pt(12)
        t322.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, 'BFBFBF')

    for columna1 in range(6):
        cell = tabla322.cell(columna1, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.width = Cm(1.80)


    for columna2 in range(6):
        cell = tabla322.cell(columna2, 1)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.width = Cm(3.63)

    for columna3 in range(6):
        cell = tabla322.cell(columna3, 2)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.width = Cm(3.63)

    tabla322.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #########################
    ### Mapa del capitulo 3.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo322 = doc.add_paragraph()
    imagenCapitulo322.text = ''
    imagenCapitulo322 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo322.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo322.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo322.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo322.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2.1 ###
    #########################
    diMap322 = doc.add_paragraph()
    descripcionCapituloMapa322 = diMap322.add_run('Mapa 3.4.-	Ubicación de vertices de _________.')
    descripcionCapituloMapa322_format = diMap322.paragraph_format
    descripcionCapituloMapa322_format.line_spacing = 1.15
    descripcionCapituloMapa322.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa322.font.name = 'Bookman Old Style'
    descripcionCapituloMapa322.font.size = Pt(12)
    descripcionCapituloMapa322.font.italic = True
    diMap322.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

     ########################################################################################################################################################################
    #  Capitulo 3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.3 ###
    #########################
    capitulo33 = doc.add_paragraph()
    i33 = capitulo33.add_run(f'\n{temasCapitulo3[1][4]}')
    i33_format = capitulo33.paragraph_format
    i33_format.line_spacing = 1.5

    i33.font.name = 'Arial'
    i33.font.size = Pt(12)
    i33.font.bold = True

    #########################
    ### Descripcion del capitulo 3.3 ###
    #########################
    listaCapitulo33 = [
        'Región Hidrológica: ',
        'Cuenca Hidrográfica: ',
        'Subcuenca Hidrográfica: ',
        'Microcuenca: '
    ]

    listaCap33 = range(len(listaCapitulo33))

    di33 = doc.add_paragraph()
    descripcionCapitulo33 = di33.add_run('Para la ubicación del predio a nivel de la cuenca forestal se tiene lo siguiente:')
    descripcionCapitulo33_format = di33.paragraph_format
    descripcionCapitulo33_format.line_spacing = 1.5
    descripcionCapitulo33_format.space_after = 0

    descripcionCapitulo33.font.name = 'Arial'
    descripcionCapitulo33.font.size = Pt(12)
    
    for lista in listaCap33:
        di33 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo33 = di33.add_run(f'{listaCapitulo33[lista]}')
        descripcionCapitulo33_format = di33.paragraph_format
        descripcionCapitulo33_format.line_spacing = 1.5

        descripcionCapitulo33.font.name = 'Arial'
        descripcionCapitulo33.font.size = Pt(12)

    #########################
    ### Mapa del capitulo 3.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo33 = doc.add_paragraph()
    imagenCapitulo33.text = ''
    imagenCapitulo33 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo33.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo33.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo33.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo33.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.3 ###
    #########################
    diMap33 = doc.add_paragraph()
    descripcionCapituloMapa33 = diMap33.add_run('Mapa 3.5.-	Ubicación con respecto a la Región Hidraulica.')
    descripcionCapituloMapa33_format = diMap33.paragraph_format
    descripcionCapituloMapa33_format.line_spacing = 1.15
    descripcionCapituloMapa33.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa33.font.name = 'Bookman Old Style'
    descripcionCapituloMapa33.font.size = Pt(12)
    descripcionCapituloMapa33.font.italic = True
    diMap33.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    #########################
    ### Segundo Mapa del capitulo 3.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo33 = doc.add_paragraph()
    imagenCapitulo33.text = ''
    imagenCapitulo33 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo33.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo33.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo33.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo33.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Segunda Descripcion del mapa del capitulo 3.3 ###
    #########################
    diMap33 = doc.add_paragraph()
    descripcionCapituloMapa33 = diMap33.add_run('Mapa 3.6.-	Ubicación del predio con respecto a la subcuenca y microcuenca.')
    descripcionCapituloMapa33_format = diMap33.paragraph_format
    descripcionCapituloMapa33_format.line_spacing = 1.15
    descripcionCapituloMapa33.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa33.font.name = 'Bookman Old Style'
    descripcionCapituloMapa33.font.size = Pt(12)
    descripcionCapituloMapa33.font.italic = True
    diMap33.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.4
    ########################################################################################################################################################################
    capitulo34 = doc.add_paragraph()
    i34 = capitulo34.add_run(f'\n{temasCapitulo3[1][5]}')
    i34_format = capitulo34.paragraph_format
    i34_format.line_spacing = 1.5
    i34_format.space_after = 0

    i34.font.name = 'Arial'
    i34.font.size = Pt(12)
    i34.font.bold = True

    ########################################################################################################################################################################
    #   Capitulo 3.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 3.4.1 ###
    #########################
    capitulo341 = doc.add_paragraph()
    i341 = capitulo341.add_run(f'\n{temasCapitulo3[1][6][0]}')
    i341_format = capitulo341.paragraph_format
    i341_format.line_spacing = 1.5

    i341.font.name = 'Arial'
    i341.font.size = Pt(12)
    i341.font.bold = True

    #########################
    ### Descripcion del capitulo 3.4.1 ###
    #########################
    di341 = doc.add_paragraph()
    descripcionCapitulo341 = di341.add_run('Cabe destacar que el sitio propuesto para el proyecto inmerso en el municipio de __________, Coahuila se encuentran ______ del área natural protegida, denominada _______________, (ver anexo mapa 3.7.- Ubicación del proyecto en ANP). El desarrollo del mismo no afectará o no causará daños a esta área, ya que durante el desarrollo del proyecto se considera la aplicación de las medidas preventivas (control de polvo, ruido y mantenimiento preventivo y correctivo de la maquinaria para evitar la contaminación del suelo), por lo que el proyecto se apegará a lo indicado en esta acción. del ANP Federal "_________" lineamientos y estrategias del, Ordenamiento Territorial y Desarrollo Urbano del Estado de Coahuila de Zaragoza, así como al Plan director de Desarrollo Urbano del municipio de _________, Coahuila. ')
    descripcionCapitulo341_format = di341.paragraph_format
    descripcionCapitulo341_format.line_spacing = 1.15
    descripcionCapitulo341_format.space_after = 0

    descripcionCapitulo341.font.name = 'Arial'
    descripcionCapitulo341.font.size = Pt(12)
    di341.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.4.1 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo341 = doc.add_paragraph()
    imagenCapitulo341.text = ''
    imagenCapitulo341 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo341.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo341.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo341.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo341.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.4.1 ###
    #########################
    diMap341 = doc.add_paragraph()
    descripcionCapituloMapa341 = diMap341.add_run('Mapa 3.7.-	Ubicación en el ANP.')
    descripcionCapituloMapa341_format = diMap341.paragraph_format
    descripcionCapituloMapa341_format.line_spacing = 1.15
    descripcionCapituloMapa341.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa341.font.name = 'Bookman Old Style'
    descripcionCapituloMapa341.font.size = Pt(12)
    descripcionCapituloMapa341.font.italic = True
    diMap341.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.4.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 3.4.2 ###
    #########################
    capitulo342 = doc.add_paragraph()
    i342 = capitulo342.add_run(f'\n{temasCapitulo3[1][6][1]}')
    i342_format = capitulo342.paragraph_format
    i342_format.line_spacing = 1.5

    i342.font.name = 'Arial'
    i342.font.size = Pt(12)
    i342.font.bold = True

    #########################
    ### Descripcion del capitulo 3.4.2 ###
    #########################
    di342 = doc.add_paragraph()
    descripcionCapitulo342 = di342.add_run('El área en estudio del municipio Ramos Arizpe _____ encuentra inmerso dentro de ninguna AICA, pero es importante señalar que la AICA más cercana se encuentra a ___ m hacia el sureste, la cual es la AICA MEX-1 (_____) __________________, sin embargo, de acuerdo al monitoreo de especies registradas y observadas dentro del área y del predio, no se ubicó ninguna que este considerada con algún tipo de protección siendo el hábitat de esta sierra por sus características abióticas así que dan sustento a las especies registradas dentro de esta AICA.')
    descripcionCapitulo342_format = di342.paragraph_format
    descripcionCapitulo342_format.line_spacing = 1.15
    descripcionCapitulo342_format.space_after = 0

    descripcionCapitulo342.font.name = 'Arial'
    descripcionCapitulo342.font.size = Pt(12)
    di342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di342 = doc.add_paragraph()
    descripcionCapitulo342 = di342.add_run('La posible influencia con la implementación del proyecto se considera temporal y a la vez se mantendrá una superficie sin alteración donde podrá haber anidación de las aves que tengan su hábitat en esta zona o de aquellas que sean migratorias mismas que no son comunes por la falta de agua ya que solo se generan encharcamientos posteriores a la temporada de lluvias, siendo estas solamente de paso hacia esta área. (ver mapa 3.82')
    descripcionCapitulo342_format = di342.paragraph_format
    descripcionCapitulo342_format.line_spacing = 1.15
    descripcionCapitulo342_format.space_after = 0

    descripcionCapitulo342.font.name = 'Arial'
    descripcionCapitulo342.font.size = Pt(12)
    di342.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo342 = doc.add_paragraph()
    imagenCapitulo342.text = ''
    imagenCapitulo342 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo342.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo342.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo342.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo342.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.4.2 ###
    #########################
    diMap342 = doc.add_paragraph()
    descripcionCapituloMapa342 = diMap342.add_run('Mapa 3.8.-	Ubicación con respecto a las AICAS.')
    descripcionCapituloMapa342_format = diMap342.paragraph_format
    descripcionCapituloMapa342_format.line_spacing = 1.15
    descripcionCapituloMapa342.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa342.font.name = 'Bookman Old Style'
    descripcionCapituloMapa342.font.size = Pt(12)
    descripcionCapituloMapa342.font.italic = True
    diMap342.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 3.4.3 ###
    #########################
    capitulo343 = doc.add_paragraph()
    i343 = capitulo343.add_run(f'\n{temasCapitulo3[1][6][2]}')
    i343_format = capitulo343.paragraph_format
    i343_format.line_spacing = 1.5

    i343.font.name = 'Arial'
    i343.font.size = Pt(12)
    i343.font.bold = True

    #########################
    ### Descripcion del capitulo 3.4.3 ###
    #########################
    di343 = doc.add_paragraph()
    descripcionCapitulo343 = di343.add_run('Las Regiones Hidrológicas Prioritarias (RHP) incluyen una rica variedad de ecosistemas, muchos de los cuales están física y biológicamente conectados o articulados por el flujo del agua y el movimiento de las especies. En estas se incluyen lagos, ríos, estanques, corrientes, aguas subterráneas, manantiales, cavernas sumergidas, planicies de inundación, charcos e incluso el agua acumulada en las cavidades de los árboles, actualmente se contempla un total 110 regiones para el territorio nacional.')
    descripcionCapitulo343_format = di343.paragraph_format
    descripcionCapitulo343_format.line_spacing = 1.15
    descripcionCapitulo343_format.space_after = 0

    descripcionCapitulo343.font.name = 'Arial'
    descripcionCapitulo343.font.size = Pt(12)
    di343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di343 = doc.add_paragraph()
    descripcionCapitulo343 = di343.add_run('El proyecto en mención se encuentra ________________________, (Ver anexo Mapa 3.9.- Ubicación en RHP)')
    descripcionCapitulo343_format = di343.paragraph_format
    descripcionCapitulo343_format.line_spacing = 1.15
    descripcionCapitulo343_format.space_after = 0

    descripcionCapitulo343.font.name = 'Arial'
    descripcionCapitulo343.font.size = Pt(12)
    di343.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.4.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo343 = doc.add_paragraph()
    imagenCapitulo343.text = ''
    imagenCapitulo343 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo343.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo343.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo343.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo343.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.4.3 ###
    #########################
    diMap343 = doc.add_paragraph()
    descripcionCapituloMapa343 = diMap343.add_run('Mapa 3.9.-	Ubicación dentro de RHP.')
    descripcionCapituloMapa343_format = diMap343.paragraph_format
    descripcionCapituloMapa343_format.line_spacing = 1.15
    descripcionCapituloMapa343.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa343.font.name = 'Bookman Old Style'
    descripcionCapituloMapa343.font.size = Pt(12)
    descripcionCapituloMapa343.font.italic = True
    diMap343.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 3.4.4 ###
    #########################
    capitulo344 = doc.add_paragraph()
    i344 = capitulo344.add_run(f'\n{temasCapitulo3[1][6][3]}')
    i344_format = capitulo344.paragraph_format
    i344_format.line_spacing = 1.5

    i344.font.name = 'Arial'
    i344.font.size = Pt(12)
    i344.font.bold = True

    #########################
    ### Descripcion del capitulo 3.4.4 ###
    #########################
    di344 = doc.add_paragraph()
    descripcionCapitulo344 = di344.add_run('Las Regiones Terrestres Prioritarias (RTP), corresponden a unidades físico-temporales estables en la parte continental del territorio nacional, que destacan por la presencia de una riqueza ecosistémica y específica y una presencia de especies endémicas comparativamente mayor que en el resto del país. En el territorio nacional se dispone de 152 RTP para la conservación de la biodiversidad que cubren una superficie de 515,558 km2.')
    descripcionCapitulo344_format = di344.paragraph_format
    descripcionCapitulo344_format.line_spacing = 1.15
    descripcionCapitulo344_format.space_after = 0

    descripcionCapitulo344.font.name = 'Arial'
    descripcionCapitulo344.font.size = Pt(12)
    di344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di344 = doc.add_paragraph()
    descripcionCapitulo344 = di344.add_run('El área destinada para el proyecto se encuentra ________............_________________................, ya que el proyecto contempla la aplicación de medidas preventivas (control de polvo, ruido y contaminación de suelo), mediante la humectación del área para reducir la emisión de partículas a la atmósfera, así como el mantenimiento preventivo y correctivo de la maquinaria, y en caso de algún derrame la extracción del suelo contaminado para su posterior manejo y disposición, así también se aplicara un programa de rescate de flora y fauna silvestre dando prioridad aquellas de lento crecimiento y desplazamiento, así como aquellas que pudieran estar enlistadas en la NOM-059-SEMARNAT-2010.')
    descripcionCapitulo344_format = di344.paragraph_format
    descripcionCapitulo344_format.line_spacing = 1.15
    descripcionCapitulo344_format.space_after = 0

    descripcionCapitulo344.font.name = 'Arial'
    descripcionCapitulo344.font.size = Pt(12)
    di344.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.4.4 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo344 = doc.add_paragraph()
    imagenCapitulo344.text = ''
    imagenCapitulo344 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo344.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo344.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo344.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo344.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.4.4 ###
    #########################
    diMap344 = doc.add_paragraph()
    descripcionCapituloMapa344 = diMap344.add_run('Mapa 3.10.-	Ubicación dentro de RTP.')
    descripcionCapituloMapa344_format = diMap344.paragraph_format
    descripcionCapituloMapa344_format.line_spacing = 1.15
    descripcionCapituloMapa344.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa344.font.name = 'Bookman Old Style'
    descripcionCapituloMapa344.font.size = Pt(12)
    descripcionCapituloMapa344.font.italic = True
    diMap344.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    #   Capitulo 3.4.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 3.4.5 ###
    #########################
    capitulo345 = doc.add_paragraph()
    i345 = capitulo345.add_run(f'\n{temasCapitulo3[1][6][4]}')
    i345_format = capitulo345.paragraph_format
    i345_format.line_spacing = 1.5

    i345.font.name = 'Arial'
    i345.font.size = Pt(12)
    i345.font.bold = True

    #########################
    ### Descripcion del capitulo 3.4.5 ###
    #########################
    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('Por primera vez los gobiernos, las ONG y los expertos en humedales piden la creación de un tratado internacional sobre humedales y una lista de humedales de importancia internacional. Por lo que se da una conferencia del 12 al 16 de noviembre de 1962 “La Conferencia MAR” (de MARshes, MARécages, MARismas) organizada por el Dr. Luc Hoffmann se celebra en Les Saintes Maries-de-la-Mer en la Camarga francesa.')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('Del año 1963 a 1970 el texto de una convención sobre los humedales se negocia en una serie de reuniones internacionales con el apoyo del International Waterfowl and Wetlands Research Bureau (IWRB, actualmente Wetlands International), el Profesor G.V.T. Matthews y el Gobierno de los Países Bajos.')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('Para el año 1970, los días 2 y 3 de febrero el Departamento de Caza y Pesca del Irán organiza una conferencia que se celebra a orillas del mar Caspio, en el balneario de Ramsar (Irán), donde los representantes de 18 naciones acuerdan la “Convención relativa a los Humedales de Importancia Internacional especialmente como Hábitat de Aves Acuáticas”.')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('La Convención RAMSAR de Humedales posee relevancia internacional, especialmente como hábitat de aves acuáticas, de la cual México es uno de los integrantes. Este convenio se firmó el 2 de febrero de 1971 en la Ciudad de RAMSAR (Irán) y entró en vigor en 1975. En el año 2010, 159 países se han sumado a dicho acuerdo, protegiendo (hasta el momento) un total de 1888 humedales con una superficie total de 185.2 millones de ha. Bajo este esquema, cada país miembro propone los humedales presentes en su territorio que puedan ser considerados de importancia internacional.')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('La Convención RAMSAR define a los humedales como extensiones de marismas, pantanos, turberas o aguas de régimen natural o artificial, permanente o temporal, estancado o corriente, dulce, salobre o salado, incluyendo las extensiones de agua marina cuya profundidad en marea baja no exceda de seis metros, de igual manera, define a las aves acuáticas como aquellas que dependen ecológicamente de las zonas húmedas (humedales).')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di345 = doc.add_paragraph()
    descripcionCapitulo345 = di345.add_run('El área sujeta de estudio para cambio de uso de suelo _______________________________________________________________________________________________ (ver anexo mapa 3.11 Ubicación del área en RAMSAR), _____________________. Se hace mención que el proyecto en cuestión no representa un riesgo al encontrarse a una distancia considerable, considerando además que el proyecto contempla aplicar las medidas preventivas de control de ruido y emisión de gases a la atmósfera mediante el mantenimiento preventivo y correctivo de la maquinaria y equipo, así como la prohibición de caza, captura o movilización de fauna silvestre.')
    descripcionCapitulo345_format = di345.paragraph_format
    descripcionCapitulo345_format.line_spacing = 1.15

    descripcionCapitulo345.font.name = 'Arial'
    descripcionCapitulo345.font.size = Pt(12)
    di345.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 3.4.5 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo345 = doc.add_paragraph()
    imagenCapitulo345.text = ''
    imagenCapitulo345 = doc.add_picture('capitulo3.png')  # Ancho de la imagen en centimetros
    imagenCapitulo345.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo345.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo345.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo345.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del mapa del capitulo 3.4.5 ###
    #########################
    diMap345 = doc.add_paragraph()
    descripcionCapituloMapa345 = diMap345.add_run('Mapa 3.11.-	Ubicación en el RAMSAR.')
    descripcionCapituloMapa345_format = diMap345.paragraph_format
    descripcionCapituloMapa345_format.line_spacing = 1.15
    descripcionCapituloMapa345.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa345.font.name = 'Bookman Old Style'
    descripcionCapituloMapa345.font.size = Pt(12)
    descripcionCapituloMapa345.font.italic = True
    diMap345.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO X DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO X DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO X DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 3 DTU EXTRACCION DE MATERIAL PETRO.docx")


capitulo3()  # Llamar a la función para ejecutar el código
