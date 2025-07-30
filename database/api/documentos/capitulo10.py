"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo10():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 10
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo VII.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 10
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 10
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 10 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'X.- Propuesta de programa de rescate y reubicación de especies de flora y fauna que pudieran resultar afectadas y su adaptación al nuevo hábitat, en caso de autorizarse el cambio de uso de suelo.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'X.1.- Programa de Rescate y Reubicación de Flora Silvestre')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'X.1.1.- Introducción')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('En México, las cactáceas se han distribuido extensamente en territorios principalmente áridos y semiáridos del norte y centro del país, llegando a ser, probablemente, las plantas más características del paisaje mexicano, en nuestro país desempeñan un papel muy importante desde el punto de vista biológico, social y económico. Se calcula que la familia Cactaceae incluye aproximadamente 110 géneros y cerca de 1,500 especies en total, de los cuales 52 géneros y 850 especies se encuentran presente en México, por lo cual es considerado el país con mayor diversidad de cactáceas, el estado de Coahuila ocupa el segundo lugar a nivel estado con mayor número de especies con 126.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El Estado de Coahuila cuenta con una gran biodiversidad prueba de ello es que la vegetación de esta región donde está inmerso el área del proyecto forma parte de una amplia distribución de especies que requieren de procesos de protección para mantener la biodiversidad de México, de ahí la importancia de llevar a cabo la ejecución del programa de rescate de Flora, hasta establecerlo en áreas dentro del proyecto o bien en las zonas aledañas, con la finalidad de establecer al menos un 85 %.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El Gobierno Federal a través del Decreto publicado en el diario oficial de la federación el día 9 ____________________, adiciona el artículo al Artículo 141, fracción IX del reglamento de la Ley General de Desarrollo Forestal Sustentable, en donde se aluce que la SEMARNAT incluirá en su resolución de autorización de cambio de uso del suelo en terrenos forestales, un programa de rescate y reubicación de especies de la vegetación forestal afectadas y su adaptación al nuevo hábitat, mismo que estará obligado a cumplir el titular de la autorización.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El objetivo del programa de rescate de Flora es la conservación de la biodiversidad, refiriendo esto a las presiones que han estado expuestos a lo largo de la historia los ecosistemas; la actividad antropogénica ha generado una gran variedad de cambios y debilitamientos al mismo, los cuales han ocasionado el deterioro de los distintos compartimentos ambientales, incluyendo el agua, el aire, el suelo, así como de la biota asociada y por ende de los ecosistemas.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las asociaciones vegetativas que integran un ecosistema se mantienen en equilibrio de acuerdo a su fenología, distribución y desarrollo en forma natural en las regiones donde tenga su hábitat fisiológico, el aprovechamiento de este tipo de ecosistemas en ocasiones desmedido en su aprovechamiento y el sobre pastoreo ha ocasionado que las especies se vallan extinguiendo o en su defecto su reproducción sea muy reducida motivada por las condiciones meteorológicas tan adversas que se han manifestado en las últimas décadas así como la aparición de especies diferentes que han venido modificando la condición natural de los ecosistemas.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para conservar la biodiversidad en alguna región, es necesario mantener las especies que lo integran lo menos alterable posible, cuando por necesidades de cambio de uso de suelo como es el caso, el rescatar especímenes para mantener la biodiversidad de la región donde se tendrán alteraciones fuera de lo ordinario es de vital importancia ya que en determinado tiempo se podrá volver a repoblar un sitio con las especies nativa.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.1.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.1.1.- Justificación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.1.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El presente Programa de Rescate de Flora; deberá considerar para su desarrollo a la totalidad de los ejemplares de las especies de flora silvestre respectivamente que estén consideradas en la NOM-059-SEMARNAT-2010, además de las consideradas de lento crecimiento y difícil regeneración y que potencialmente podrían localizarse en los sitios destinados a la realización del proyecto; así como aquellas consideradas económica y ecológicamente importantes por su interrelación con las demás especies.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.2.- Objetivos ')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.2.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.2.1.- Objetivo General')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.2.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Realizar un programa de rescate de Flora que nos conduzca a la conservación y protección de especies de interés ecológico, considerando principalmente aquellas que sean de lento crecimiento y difícil desarrollo, y las especies de interés ecológico o económico así como otras que se adapten a este proceso y que puedan ser utilizadas en actividades de restauración y conservación que se localicen dentro del área propuesta en las 17.41 ha para el cambio de uso de suelo para el establecimiento del proyecto, _______________________________________________________________________________ en el Estado de Coahuila.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    objetivoGeneral10_121 = [
        'Atender la Legislación vigente en el párrafo tercero del Artículo 93 de la Ley General de Desarrollo Forestal Sustentable, publicado en el diario oficial el ____________________, así como acorde al Artículo 141, fracción IX, del Reglamento de la Ley General de Desarrollo Forestal Sustentable  de acuerdo a su última reforma publicado el ________________________s, donde se manifiesta que se deberá de presentar un programa de rescate y reubicación de especies de la vegetación forestal y su adaptación al nuevo hábitat enfocado a especies de mayor representación en el área de cambio de uso de suelo y de menor representatividad en la unidad de análisis, de importancia ecológica, así como las especies con alguna categoría en la NOM-059-SEMARNAT-2010.',
        '_________________________________________________ especies, observadas y registradas en los sitios de muestreo que son susceptibles al proceso de extracción y reubicación en condiciones óptimas para su plantación buscando el establecimiento en sitios seleccionados en forma previa dentro de los terrenos del mismo predio.'
    ]

    objetivoGeneral10_121Rango = range(len(objetivoGeneral10_121))

    for lista in objetivoGeneral10_121Rango:
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{objetivoGeneral10_121[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.2.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.2.2.- Objetivos Específicos')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.2.2 ###
    #########################
    objetivosEspecificos10_122 = [
        'Efectuar recorrido minucioso dentro del área sujeta de estudio para ubicar las especies a rescatar que no se localizaron en los sitios de muestreo.',
        'Seleccionar las plantas de mejor calidad de acuerdo a los datos obtenidos en el muestreo para asegurar su establecimiento.',
        'Registrar mediante georreferenciación los sitios de extracción de planta de las especies consideradas.',
        'Ubicar y delimitar los sitios para la reubicación de las plantas rescatadas.',
    ]

    objetivosEspecificos10_122Rango = range(len(objetivosEspecificos10_122))

    for lista in objetivosEspecificos10_122Rango:
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{objetivosEspecificos10_122[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.3 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.3.- Metas')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Antes de realizar las actividades de desmonte y despalme del área sujeta a Cambio de Uso de Suelo se llevará a cabo un recorrido a efecto de detectar las especies propuestas para rescate y reubicación, poniendo especial atención en las especies que se encuentren listadas en la NOM-059-SEMARNAT-2010, especies de lento crecimiento y difícil regeneración, así como especies de interés ecológico o económico.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('A continuación, se presenta el listado de especies registradas y observadas en el recorrido del área sujeto a Cambio de Uso de Suelo para la implementación del Proyecto, los siguientes datos que se enlistan son el resultado de la sumatoria de los ____________________________________________________________________________________________________________________, correspondientes al área de cambio de uso de suelo, por lo que para tener una certeza del número de individuos a rescatar se utilizarán este número de individuos y no los extrapolados, por lo que se tiene lo siguiente:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.1.3 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 1.- Especies y número de individuos a rescatar.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.1.3 ###
    #########################
    columnas = 5
    filas = 10

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 10.1.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nDe acuerdo al recorrido realizado nos arroja que existen _________________________________, las cuales son especies catalogadas con parámetros como: lento crecimiento y difícil regeneración, así como de importancia en valor ecológica, las cuales serán rescatadas y reubicadas en el sitio seleccionado previamente con las características similares a su sitio de desarrollo natural buscando que su establecimiento en el nuevo sitio sea favorable debido a que las condiciones ambientales están muy alteradas y por ende la calidad de la planta está en mal estado.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Dichas especies por sus propias características son las más adaptables al proceso de extracción y en condiciones se adapten a las nuevas áreas logrando una óptima plantación, buscando dentro del predio el establecimiento en sitios seleccionados en forma previa.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.4s ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.4.- Metodología para Rescate y reubicación de especies')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.4s ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para justificar las especies a rescatar dentro del área sujeta de estudio serán estas las de valor ecológico, lento crecimiento y difícil regeneración además de aquellas que estén en algún estatus en la Norma Oficial Mexicana NOM-059- SEMARNAT-2010 y especies que no fueron observadas durante los recorridos de muestreos y que se adapten al proceso, principalmente las especies del estrato suculento.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Una vez seleccionadas las plantas de cada especie que se extraerán del área en estudio donde se llevará a cabo el Cambio de Uso de Suelo y con el propósito de aprovechar parte de las plantas que serán removidas, se contempla rescatar los especímenes y establecerlos en lugares específicos que se tienen designados donde no se tendrá alteración de la cubierta vegetal, _________________________________________________.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('De acuerdo al registro de plantas arrojada en los sitios muestreados en los dos tipos de vegetación de acuerdo al INEGI y considerando la superficie de __________, sujeto de estudio para el cambio de uso de suelo en proceso de gestión se contempla el rescate del 100 % de las especies del cuadro anterior.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las especies enlistadas se contemplan como prioritarias y con el propósito de mantener el germoplasma en el sistema ambiental al considerarse como especies de lento crecimiento y difícil regeneración y de alta prioridad con lo cual se busca mantener la Biodiversidad.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Parámetros considerados para el rescate de plantas:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El rescate de plantas propuesto está fundamentado bajo los siguientes criterios:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    metodologiaRescate10_14 = [
        'Especies de lento crecimiento.',
        'Especies de difícil regeneración.',
        'Especies con algún estatus en la NOM-059-SEMARNAT-2010.',
        'Especies de interés ecológico y económico.',
    ]

    metodologiaRescate10_14Rango = range(len(metodologiaRescate10_14))

    for lista in metodologiaRescate10_14Rango:
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{metodologiaRescate10_14[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.4.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.4.1.- Método de extracción')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.4.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Este proceso será manual con herramientas comunes, mediante la excavación circular, con el fin de no lastimar el sistema radicular, se contempla su extracción con la mayor parte del cepellón posible para asegurar su establecimiento, al momento de su extracción deberá ser cubierta para evitar la aireación de sus raíces, sin embargo las del género Opuntia, estas se aprovecharán solo las pencas (raqueta); el traslado se realizará en vehículo al sitio de plantación definitiva donde tendrán su proceso de cicatrización de ser necesario, las especies a rescatar por sus características se adaptan a este proceso.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Se llevará a cabo un registro en bitácora con georreferenciación tanto del lugar de extracción como el lugar de su reubicación, para así tener una disposición geográfica en el terreno para el momento de su plantación este en la misma posición respecto a los puntos cardinales igual a su desarrollo natural. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.5 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.5.- Lugar de acopio y extracción de especies ')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Resultado del análisis dentro del área contemplada para el establecimiento del proyecto con _________, sujeto de estudio para cambio de uso de suelo y con base a la información de campo obtenida de las especies vegetales que integran la Biodiversidad en el predio se obtuvo información para reubicación del material vegetativo con las características idóneas para su rescate y viabilidad para su establecimiento. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('De acuerdo al levantamiento de los sitios en los cuales se muestrearon las diferentes especies considerando la calidad y cantidad, en este aspecto las más idóneas para llevar a cabo el proceso serán las medianas y chicas en virtud de que se lastimarían menos durante la extracción y podrán tener mayores posibilidades de sobrevivencia.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Considerando las características del área sujeta para el establecimiento del rescate esta corresponde al mismo ecosistema por lo que su adaptabilidad al cambio no generaría procesos de adaptación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Otros aspectos que se tomarán en cuenta serán la condición sanitaria, características de desarrollo de la planta, turgencia y condición fisiológica para no alterar su proceso durante la extracción.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.6 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.6.- Localización de los sitios de reubicación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.6 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para establecer la planta rescatada en sitios aledaños al área sujeto de estudio será necesario en primera instancia ubicar el sitio para su óptimo desarrollo considerando la asociación vegetativa natural para que tenga las condiciones idóneas para su establecimiento dentro del mismo predio, como del área que fueron extraídas, tanto en el rescate como en la plantación deberá de efectuarse un registro mediante bitácora georreferenciando su ubicación. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Al tener la autorización del Cambio de Uso de Suelo se podrá iniciar el proceso de preparación del sitio tal y como se menciona con respecto al cronograma de actividades en la superficie contemplada para tal efecto, cabe señalar que las especies a rescatar son _________________________, así como las demás que se localicen dentro del área en estudio, por lo que antes de realizar este proceso se deberá tener el área de reubicación ya definida.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### COORDENADAS UTM (Universal Transversal de Mercator) ÁREA PARA REUBICACIÓN DE ESPECIES RESCATADAS ###
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('COORDENADAS UTM (Universal Transversal de Mercator) ÁREA PARA REUBICACIÓN DE ESPECIES RESCATADAS')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Se le señala a la autoridad que debido a que el área en donde se pretende realizar el proyecto de la extracción de material pétreo, no cuenta con más superficie idónea para la reubicación de las plantas, se propone la siguiente área que cuenta con una superficie de _________, las cuales se ubican en las coordenadas UTM zona 14R siguientes:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.1.6 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 2.- Coordenadas del área de rescate')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.1.6 ###
    #########################
    columnas = 3
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Mapa del capitulo 10.1.6 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo10 = doc.add_paragraph()
    imagenCapitulo10.text = '\n'
    imagenCapitulo10 = doc.add_picture('capitulo10/mapa.png')  # Ancho de la imagen en centimetros
    imagenCapitulo10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo10.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo10.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo10.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2 ###
    #########################
    diMap10 = doc.add_paragraph()
    descripcionCapituloMapa10 = diMap10.add_run('Describir el Mapa =)')
    descripcionCapituloMapa10_format = diMap10.paragraph_format
    descripcionCapituloMapa10_format.line_spacing = 1.15
    descripcionCapituloMapa10.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa10.font.name = 'Bookman Old Style'
    descripcionCapituloMapa10.font.size = Pt(12)
    descripcionCapituloMapa10.font.italic = True
    diMap10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    # Capitulo 10.1.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.- Acciones a realizar para el establecimiento, mantenimiento y supervivencia')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Una vez seleccionada el área para la reubicación de las especies rescatadas se procederá a seleccionar el sitio para plantación debiendo considerar las características más semejantes al sitio de su desarrollo natural; para este proceso se realizará la plantación de las especies rescatadas de acuerdo al desarrollo natural de cada especie, ello en virtud del área donde se pretende realizar estas actividades presenta vegetación típica del ______________________________________________________________, por lo que para no dañar las plantas ya establecidas o que se encuentren en el área se buscará un espacio para el establecimiento de las mismas condiciones, así también esta área deberá estar aislada del resto del área del proyecto y que no será impactada por lo cual se estima la sobrevivencia de las especies rescatadas.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.- Actividades para realizar la Reforestación de especies rescatadas')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para llevar a cabo este proceso será necesario seleccionar los sitios más idóneos para la plantación con base a la asociación donde se desarrolla en forma natural las especies a extraer con el propósito de proporcionar las condiciones adecuadas a la planta en el terreno donde se contempla su desarrollo definitivo y no será alterada la vegetación, con base a ello se marcarán los sitios donde se establecerá la planta, con ello se prosigue a llevar a cabo las actividades necesarias.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.1.- Requerimiento de personal y equipo.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.1.7.1.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 3.- Personal requerido para las actividades de reforestación de especies de vegetación rescatada.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.1.7.1.1 ###
    #########################
    columnas = 4
    filas = 10

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.2.- Herramientas Manuales y Equipo Necesario para Rescate y Plantación.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.2 ###
    #########################
    for lista in range(7):
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'Herramienta y/o manual {lista + 1}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Estos materiales, herramientas y equipo son los mínimos requeridos para llevar a cabo el proceso en forma manual tanto para la extracción como en la preparación del terreno antes y durante la plantación y seguimiento a su establecimiento. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.3 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.3.- Apertura de cepas')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Este proceso se recomienda llevarlo a cabo antes de la plantación con el fin de que se aire el suelo, si se contempla aplicar riego es recomendable para que el suelo inicie su recuperación de humedad y asentamiento, caso contrario preparar el terreno cercano a los períodos de lluvia de acuerdo al calendario que se tenga considerado. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las dimensiones de las cepas serán de diferentes medidas de acuerdo al tipo de planta y su volumen y contemplándose la plantación en armonía a su condición lo más cercano a su origen.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('De acuerdo a las especies registradas y su cantidad de individuos a rescatar y reubicar no se pude determinar una densidad de plantación ya que estas estarán sujetas a las condiciones de los sitios que se seleccionen solo se observará que haya correlación con la vegetación donde naturalmente se ha venido desarrollando.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.4 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.4.- Diseño de plantación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.4 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Este no será viable, de acuerdo a las especies contempladas para el proceso, debido al número de individuos a extraer y de acuerdo a su condición de desarrollo estas no cumplen con el diseño de plantación, lo recomendado es utilizar un espaciamiento mínimo de 4 m x 4 m, con método de siembra cuadrado en zonas planas, al considerar que las especies que serán rescatadas son del tipo suculentas que no requieren espacios abiertos, el sitio determinará la distribución buscando que dichas plantas se interrelacionen con las naturales para que reciban protección y facilite su establecimiento, sin embargo, tomando en cuenta el número de individuos  a rescatar este diseño no se cumplirá y el diseño será irregular.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.5 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.5.- Densidad')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Por sus características de las especies estas no se adecuan a ninguna densidad debido a sus características ya que son especies suculentas que no todas se adaptas a estas características, estas se establecerán irregularmente dentro del área designada buscando espacios que reúnan las características similares al sitio de extracción donde se desarrollaban naturalmente, según el manual de la CONAFOR para los ecosistemas de zonas áridas recomienda el establecimiento de plantaciones de 625 plantas por has.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.6 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.1.6.- Transporte de plántulas')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.6 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El transporte del material vegetal, se hará con los cuidados necesarios para evitar el maltrato a las plantas y protegiendo la planta en los trayectos aun cuando sean cortos, para mantener su condición lo menos alterable posible, en caso de ser necesario se realizarán en vehículo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.1.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.1.7 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\n')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.1.7 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El proceso más recomendado es llevarlo a cabo en días previos a la época de lluvia, para garantizar el completo establecimiento de las plantas y tener una alta sobrevivencia en la reubicación, sin embargo, como alternativa se plantea aplicar de uno a dos riegos de auxilio a la planta para favorecer su establecimiento en virtud de estar el sitio en una región donde la precipitación es baja.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('_______________________________________________________________________________________________________________________________________________, para ello se determinará un periodo según la condición de humedad para el proceso de cicatrización antes de ser llevado el material vegetal al sitio definitivo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('En forma general para lograr un desarrollo radicular óptimo y una adecuada adaptación se deberán efectuar riegos de auxilio al menos dos antes del periodo ordinario de las precipitaciones para asegurar el éxito del establecimiento de la reubicación de las especies propuestas de al menos en un 85 %. En el proceso de plantación se propone adherir un enraizador para facilitar el establecimiento y en los riegos de auxilio se recomienda aplicar algún fertilizante, con ello se tendrá mayor seguridad de sobrevivencia de la plantación y cumplir con el objetivo y mandato de ley para mantener la Biodiversidad de aquellas especies de difícil regeneración y lento crecimiento y poder establecer su adaptación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.7.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.7.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.7.2.- Protección del área de reubicación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.7.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Considerando que el área seleccionada para la reubicación de las plantas rescatadas deberá contar con las condiciones adecuadas para su desarrollo y establecimiento es recomendable proteger y aislar dicha área de cualquier tipo de actividad hasta tener bien establecido el material rescatado para su óptimo desarrollo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.8 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.8.- Acciones a realizar para el mantenimiento y supervivencia')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.8 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Con el fin de asegurar que la reubicación de especies se desarrolle de acuerdo a la fisiología de cada una de las especies se recomienda que se le apliquen riegos de auxilio antes de la temporada de lluvias, esto de ser necesario considerando el tipo de vegetación que no requiere de mucha humedad, el propósito de considerar la plantación cercana a los períodos de lluvia es para tener mejores resultados en su establecimiento. Así también el mantenimiento constará de reposición de plantas que por alguna razón no sobreviva a este tipo de proceso ya que se necesita una sobrevivencia de al menos el 85 % de las plantas establecidas, al igual que establecer y mantener una mejor cobertura las Microcuencas para captación de humedad y mantener libre de malezas el cajete para su óptimo desarrollo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.9
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.9 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.9.- Programa de actividades')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.9 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Se contempla el programa de actividades en virtud de que el período de adaptación y establecimiento de las plantas descritas sujetas a rescatar en un período de seis meses se podrá obtener su establecimiento definitivo, el siguiente cronograma se realizará en relación a la fecha de autorización del cambio de uso de suelo el cual contempla las siguientes actividades.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.1.9 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 4.- Programa de actividades para el programa de rescate y reubicación de flora silvestre.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.1.9 ###
    #########################
    columnas = 17
    filas = 7

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 10.1.9 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nPara el caso del mantenimiento posterior al establecimiento del área a rescatar se contempla un monitoreo y vigilancia de al menos 5 años posteriores a este para asegurar el establecimiento y sobre vivencia propuesto de acuerdo a lo siguiente:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.1.9 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 5.- Programa de mantenimiento de las especies rescatadas.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.1.9 ###
    #########################
    columnas = 7
    filas = 7

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 10.1.10
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.10 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.10.- Evaluación del rescate y reubicación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.10 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Con el procedimiento para la extracción y el manejo que se le dará a cada especie se considera que se podrá tener una adaptación al nuevo sitio máxime si se considera que estará reubicada en el mismo ecosistema y con las condiciones similares a su desarrollo natural.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Se evaluará mensualmente su desarrollo tiempo en el cual se podrán realizar actividades de mantenimiento para fortalecer su establecimiento, en este proceso se analizará el comportamiento de adaptación, sanidad de las plantas y posible crecimiento al ser de lento desarrollo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.11
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.11 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.11.- Informe de avances y resultados')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.11 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Al término del proceso se deberá de efectuar un informe en el cual se mencione la cantidad de planta final que se recuperó con la caracterización del hábitat natural, así como las características del sitio donde se estableció para su desarrollo definitivo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('De igual forma se deberá describir la condición en la cual se encuentra la plantación mencionando todo el proceso desde la ubicación, selección, plantación, mantenimiento y la condición de desarrollo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Durante cada uno de los procesos se obtendrá material fotográfico que sustente cada actividad para enriquecer el documento y se mantenga como prueba del seguimiento e interés por mejorar el entorno ambiental una vez obtenida el producto para el cual se solicita el documento técnico unificado, las actividades de monitoreo se establecerán seis meses después de su establecimiento y un monitoreo cada año para asegurar su sobrevivencia el cual se reportara bajo el siguiente formato.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.12
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 10.1.12 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.12.- Formato de Muestreo en la Reforestación')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Formato del capitulo 10.1.12 ###
    #########################
    """
        =============================================================================================
        --- Aqui se tienen que rellenar los datos manualmente, cuando arroje el archivo ---
        =============================================================================================
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nFormato de muestreo de vegetación en la reforestación')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nNo. de sitio: ___________                                                  Fecha: ______/______/______')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Coordenadas UTM y Altitud del Sitio de Muestreo:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('X: ____________________        Y: _________________    Altitud: ________________')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nFotografías del Sitio:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('De la fotografía: ____________ A la fotografía: ______________\n')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True

    """
        =============================================================================================
        --- Aqui se mostrará la tabla del formato de muestreo de reforestación ---
        =============================================================================================
    """

    columnasFormato = [
        'Especie Establecida',
        'Número de Individuos',
        'Cobertura (Cm)',
        'Altura (Cm)',
        'Estado fitosanitario'
    ]

    columnas = len(columnasFormato)
    filas = 7

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        t10b = cell.paragraphs[0].add_run(columnasFormato[cols])
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)
            t10b.font.name = 'Arial'

    """
        =============================================================================================
        --- Aqui se mostrará el apartado de 'Observaciones' ---
        =============================================================================================
    """

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nObservaciones:__________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    descripcionCapitulo10_format.space_after = 0
    descripcionCapitulo10_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.1.13
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.1.13 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.1.13.- Especies a rescatar')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.1.13 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Insertar las imagenes de las especies a rescatar.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2
    ########################################################################################################################################################################

    #########################
    ### Salto de Página ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 10.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'X.2.- Programas de Rescate de Fauna Silvestre')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\n')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El constante cambio del uso del suelo y la creciente demanda de los recursos naturales resulta en la pérdida de los paisajes naturales, trayendo consigo perturbación en los ciclos de vida y funciones generales de la flora y fauna.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('La pérdida del hábitat por la modificación de las coberturas vegetales para el desarrollo de proyectos, trae consigo alteraciones en las dinámicas de especies faunísticas, tanto es su estructura como en su función, ya que los hábitats son utilizados como refugio y proveedor principal de alimento, permitiendo el desarrollo espacial de cada especie para llevar a cabo sus funciones y procesos básicos, en especial para las especies de lento desplazamiento como los anfibios, reptiles y algunos pequeños mamíferos, siendo estos grupos los más susceptibles al limitar su movilidad sufriendo atropellamientos por vehículos o maquinaria.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El área del sistema ambiental en estudio aloja una gran variedad de especies, de las cuales se mencionan algunas como ______________________________________________________________________________________________ etc.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Considerando la descripción del hábitat con vegetación de regiones desérticas y semiáridas, dentro de la región del área de estudio se localizaron en su mayoría especies __________________________________________________. Por lo tanto, se realiza el presente programa de rescate y conservación de las especies en el sitio que se pretende destinar para efectuar el proyecto, las medidas, obras y actividades que se desempeñaran para preservar el hábitat y las poblaciones que se detecten.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('En caso de que durante el proceso sea localizada una especie del género Crotalus, se llevará a cabo el rescate de la misma dentro del área de cambio de uso de suelo, sin afectar su hábitat y el desarrollo de esta.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('La ejecución de este programa es una medida de conservación de las especies silvestres y para el mantenimiento de la biodiversidad local.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.2.- Objetivos')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.2.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.2.1.- Objetivo General.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.2.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Realizar un programa de rescate de Fauna Silvestre encaminado a la conservación y protección de las especies de interés ecológico, principalmente de aquellas de lento desplazamiento, al igual que aquellas especies que se encuentren en estatus de protección en la Norma Oficial Mexicana NOM-059-SEMARNAT-2010, para las actividades contempladas para la ejecución de cambio de uso de suelo para el establecimiento del _____________________________________________________________________..2.2.2')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.2.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.2.2.- Objetivos Específicos.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.2.2 ###
    #########################
    objetivosEspecificos10_222 = [
        'Conforme al listado de fauna obtenido de la información recabada en campo, serán seleccionadas especies prioritarias para llevar a cabo el rescate, principalmente para aquellas de lento desplazamiento.',
        'Ubicación y delimitación del área destinada para la reubicación y liberación de las especies.',
        'Captura, registro y geo-referenciación de las especies de fauna sujetas a rescate.',
        'Manipulación de las especies capturadas, brindando un trato digno y adecuado para no provocar estrés que conduzca a la muerte de las especies capturadas.',
        'Liberación de las especies faunísticas capturadas en el área propuesta para reubicación.',
    ]

    objetivosEspecificos10_222Rango = range(len(objetivosEspecificos10_222))

    for lista in objetivosEspecificos10_222Rango:
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run('Describir en este parrafo.')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.3 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.3.- Área de estudio')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('______________________________________________________________________, como se manifiesta en el Anexo mapa No. 3-1. Ubicación del área en estudio. Para llegar al área en estudio es por la vía de acceso de carretera pavimentada partiendo _________________________________________________________________. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 10.2.3 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo10 = doc.add_paragraph()
    imagenCapitulo10.text = '\n'
    imagenCapitulo10 = doc.add_picture('capitulo10/mapa.png')  # Ancho de la imagen en centimetros
    imagenCapitulo10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo10.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo10.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo10.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2 ###
    #########################
    diMap10 = doc.add_paragraph()
    descripcionCapituloMapa10 = diMap10.add_run('Describir el Mapa =)')
    descripcionCapituloMapa10_format = diMap10.paragraph_format
    descripcionCapituloMapa10_format.line_spacing = 1.15
    descripcionCapituloMapa10.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa10.font.name = 'Bookman Old Style'
    descripcionCapituloMapa10.font.size = Pt(12)
    descripcionCapituloMapa10.font.italic = True
    diMap10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    # Capitulo 10.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.- Actividades de Rescate de Fauna Silvestre.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las actividades de rescate de las especies faunísticas comprenderán como inicio la fase de preparación del sitio y antes de la fase de desmonte y despalme, en las cuales se realizaran labores de ahuyentamiento de especies de los grupo de fauna registrados en los muestreos realizados, si alguna de las especies no se desplaza al predio o sistema ambiental, se llevará a cabo la búsqueda y captura de los especímenes, siguiendo las medidas establecidas por la autoridad ambiental conforme a la autorización de cambio de uso de suelo en terrenos forestales para la protección de las especies de fauna que se distribuyen dentro de la zona donde se llevarán a cabo las actividades del proyecto, proponiendo las siguientes estrategias:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    actividades10_24 = [
        'Detectar la presencia de ejemplares y marcar, delimitar y excluir las actividades dentro del área que puedan afectar, hasta que se evalúen las condiciones del hábitat y de las especies, analizando las alternativas de conservación y preservación posible o necesarias.',
        'Aplicar medidas preventivas o correctivas que mejoren las condiciones de las poblaciones de la especie.',
        'Evaluar los indicadores de éxito del programa con el objeto de tener la certeza del cumplimiento de los objetivos.',
        'Elaborar y presentar oportunamente los reportes e informes técnicos periódicos y finales del proyecto y conservación de la especie.',
    ]

    actividades10_24Rango = range(len(actividades10_24))

    for lista in actividades10_24Rango:
        di10 = doc.add_paragraph()
        descripcionCapitulo10 = di10.add_run(f'{actividades10_24[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.1.- Metodología.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Describir en este parrafo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.2.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 6.- Listado de las especies de faunísticas (aves, mamíferos y reptiles).')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.4.1 ###
    #########################
    columnas = 6
    filas = 7

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Especies observadas en el ACUSTF ya sea mediante huellas, excretas, capturas fotográficas y observación directa. Sc: sin categoría, Pr: sujeta a protección especial; A: Amenazada, ni: Número de individuos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(10)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Es de suma importancia tener el conocimiento previo sobre la fauna existente dentro del área donde se realizará el ahuyentamiento y rescate de las especies, realizando investigación previa de la zona y la región en medios en línea, proyectos previos, bibliografía e inventarios de fauna realizados con anterioridad; con el fin de tener la información general de las especies con distribución potencial en el área de trabajo, para lograr facilitar la identifican de las especies faunísticas en campo. ')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Como parte del programa es importante que el personal que lleve a cabo las actividades de captura, transporte y liberación de fauna, posea conocimientos sobre el manejo de fauna, para reducir los riesgos de mortalidad de los individuos a reubicar. Por ello se sugiere que un mismo equipo de personas se encargue de dichas actividades, lo que a su vez permitirá la adquisición de experiencia y habilidades que aumentara el éxito de sobrevivencia de los individuos con esto se pretenden evitar riesgos de daño a la fauna.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('_________________________________________________________________________________________, Como resultado del inventario realizado para los grupos faunísticos en el área ACUSTF, arrojaron como resultado que existen 7 individuos de 4 especies para el grupo de aves, 4 individuos de 3 especies para el grupo de mamíferos y 9 individuos de 3 especies para el grupo de los reptiles, estas especies observadas y registradas en el ACUSTF, así mismo siendo este último grupo el más vulnerable y conveniente para labores de rescate debido a su lento desplazamiento, cabe señalar ante la autoridad que en los inventarios solo se registraron las especies antes mencionadas por medio de observación directa, excreta, y huellas, ya que nos enfrentamos a especies con movilidad dentro del área y son susceptibles al ahuyentamiento, siendo innecesaria su captura e incluso se pueden ya no observar en el área como puede ser el caso de las especies de aves y algunos mamíferos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las especies que sean sujetas a rescate son prioritariamente las que estén dentro de alguna categoría de riesgo de acuerdo a la NOM-059-SEMARNAT-2010, especies de importancia ecológica y de desplazamiento lento, por ello y con la información obtenida por los inventarios y análisis de las especies a rescate serán los del grupo de los reptiles, debido a la susceptibilidad al cambio de uso de suelo. A demás no solo se rescatarán aquellas especies pertenecientes al grupo de los reptiles si no también se considerará para esta acción aquellas especies de aves y mamíferos que así lo requieran.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.2.- Especies propuestas para su rescate')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las especies propuestas a rescate serán aquellas especies pertenecientes al grupo de los reptiles debido a que son especies de lento desplazamiento y por lo cual son más susceptibles a impactos de las diferentes etapas que conlleva el proyecto denominado __________________________________________________________________________________.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.2.4.1 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 7.- Especies propuestas para su rescate.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.4.1 ###
    #########################
    columnas = 4
    filas = 7

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 10.2.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.3 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.3.- Requerimientos de Personal y Equipo')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Dentro de las actividades que se llevarán a cabo para el rescate de fauna silvestre, se requiere que el personal tenga conocimientos previos en captura, manejo y manipulación de las especies, por lo que será necesario para el personal de material para captura, personal encargado y personal que realice actividades de rescate, las cuales serán enlistadas a continuación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.2.4.3 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 8.- Personal requerido para las actividades de rescate.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.4.3 ###
    #########################
    columnas = 4
    filas = 10

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 10.2.4.3 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 9.- Equipo requerido para las actividades de rescate.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.15
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.4.3 ###
    #########################
    columnas10_243 = [
        'MATERIA',
        'EQUIPO',
    ]

    materia10_243 = [
        'Libreta de campo',
        'Claves taxonómicas',
        'Cajas de plástico',
        'Cintas de colores',
        'GPS',
        'Formatos de identificación',
        'Guía de campo',
        'Palas',
        'Mallas de plástico o de tela',
        'Redes para captura',
        'Vehículos para trasporte',
        'Lámparas',
    ]

    equipo10_243 = [
        'Cámaras fotográficas',
        'Pinzas herpetológicas',
        'Ganchos herpetológicos',
        'Cebos para fauna',
        'Frascos de plástico',
        'Guantes de carnaza',
        'Cinta de papel',
        'Saco de tela',
        'Chaparreras',
        'Ganchos fijos para lagartijas',
        'Cámaras nocturnas ',
        'Formatos de campo',
    ]

    columnas = len(columnas10_243)
    filas = len(materia10_243) + 1

    rangoColumnas = range(columnas)
    rangoFilas = range(len(materia10_243))

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        t10b = cell.paragraphs[0].add_run(columnas10_243[cols])
        cell_background_color(cell, 'D9D9D9')
        t10b.font.name = 'Arial'
        t10b.font.size = Pt(12)
        t10b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for materia in rangoFilas:
        cell = tabla10b.cell(materia + 1, 0)
        t10b = cell.paragraphs[0].add_run(materia10_243[materia])
        t10b.font.name = 'Arial'
        t10b.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for equipo in rangoFilas:
        celda_actual = equipo + 1
        cell = tabla10b.cell(celda_actual, 1)
        t10b = cell.paragraphs[0].add_run(equipo10_243[equipo])
        t10b.font.name = 'Arial'
        t10b.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 10.2.4.3 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nEstos materiales herramientas y equipo son los mínimos requeridos para llevar a cabo el proceso en forma manual tanto para la captura, traslado y seguimiento a su establecimiento y reubicación de las especies.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.4 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.4.- Técnicas aplicadas en el Rescate de Fauna Silvestre')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.4 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El manejo y la manipulación de la fauna silvestre, implica técnicas específicas para el uso y aplicación adecuada de métodos para la captura y ahuyentamiento para cada grupo faunístico respectivamente, salvaguardando la integridad de los individuos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.4 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_244/cap_10.244.png', width=Cm(15.19), height=Cm(13.45))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Figura 1.- Plan de rescate de Fauna Silvestre por fases.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 10.2.4.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.4.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.4.1.- Técnicas de Ahuyentamiento')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.4.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El ahuyentamiento de la fauna se concentra en técnicas que generan las condiciones que permita el desplazamiento de los animales que se encuentran dentro de la zona que será intervenida para un proyecto, combinando con el rescate y la reubicación de los individuos. En un ahuyentamiento de fauna se emplean diferentes metodologías y técnicas, como estímulos visuales, estímulos Auditivos, estímulos mecánicos y estímulos químicos, para el rescate de especies se seleccionan aquellas técnicas que sean las más apropiadas para el ahuyentamiento ya que algunas de las técnicas generan un cierto grado de estrés a los animales.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para esto la principal técnica que se utilizará para el _________________________________________________________________________ serán aquellas que generen estímulos visuales, los cuales con la sola presencia de una o varias personas, provocará que los individuos se muevan hacia sitios o área aledañas al área de estudio. Las especies más susceptibles a esta técnica son el grupo de las aves siguiendo por el grupo de los mamíferos de tamaño mediano como coyotes, liebres y conejos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.4.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.4.2.- Técnicas Para la Captura.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.4.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Existen varios métodos para capturar animales de los diferentes grupos, los cuales implementan estrategias (trampas, cebos, captura manual, etc.), que llevan a la captura de los individuos, están diseñadas para minimizar el estrés a los animales en todo momento, tomando en cuenta el comportamiento y características físicas de la especie.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las técnicas que a continuación se describen ya han sido utilizadas en actividades de rescate, así como también son técnicas utilizadas para el estudio de la fauna silvestre para las cuales se requiere de la captura de los individuos apareciendo en diferentes manuales de técnicas para la investigación de fauna silvestre, por ello se optó por la elección de dichas técnicas de captura.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
    ############################
    ### AVIFAUNA ###
    ############################
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nAVIFAUNA')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    descripcionCapitulo10.italic = True
    di10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El grupo faunístico de las aves cuentan con un alto rango de movilidad, por lo que abarcan gran área de distribución dentro del ACUSTF y hacia otros lugares en comparación a otros grupos, este grupo se verá perturbado durante el proceso de cambio de uso de suelo por lo cual no se capturarán aves para su reubicación ya que estas migrarán hacia áreas circundantes con características esenciales al área de ACUSTF, por ello este grupo comprende varios métodos y técnicas de captura. Para este grupo se recurrirá en primera instancia a técnicas de ahuyentamiento por su susceptibilidad a estas actividades, en conjunto con el establecimiento de transectos para la observación de los individuos que no reaccionen a esta técnica, en caso de ser necesario se colocaran redes ornitológicas para poder llevar a cabo su captura durante horarios establecidos. El área de cambio de uso de suelo no es utilizada como área de anidación, por lo cual no será necesario manipular nidos. En dado caso de que se encuentren nidos en el área se verificará si está activo o inactivo, si este se encuentra activo el nido se reubicará removiendo la base de la rama que lo sostiene, el nido se manipulara con guantes para evitar la impregnación de olores que puedan hacer desistir del nido a la pareja y provoquen su abandono, por lo contrario, en caso de que el nido se encuentre inactivo se procederá a la remoción del nido para evitar que alguna otra especie anide en él.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen1.png', width=Cm(12.64), height=Cm(7.2))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Figura 1.- Plan de rescate de Fauna Silvestre por fases.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
    ############################
    ### MAZTOFAUNA ###
    ############################
    """
    #########################
    ### Descripcion del capitulo 10.2.4.4.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nMAZTOFAUNA')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    descripcionCapitulo10.italic = True
    di10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Los mamíferos son especies las cuales tienen por lo general un solo pico de actividad, presentando en algunos un rango de movilidad alto y otros más reducido, en general responden de forma positiva a estímulos Visuales, Auditivos, Químicos y Mecánicos, dicha característica dará paso a que las especies de este grupo se muevan hacia otros lugares con la presencia de la maquinaria de trabajo que sea utilizada, de igual manera, se realizará actividades de ahuyentamiento tanto en las horas del día como en la noche, sin embargo, se tendrá una mayor dedicación durante la última. Consecuentemente, se debe realizar una actividad de ahuyentamiento en las horas de la mañana, entre las 7:00 y las 9:00 y otra en la tarde entre las 6:00 pm.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las capturas de los individuos se realizan de forma manual mediante el uso de aparatos mecánicos como trampas de tipo Sherman y Tomahawk, previamente cebadas con olores atractivos y adecuados para las diferentes especies, con esto se facilitará la captura de los ejemplares ya sea para roedores, liebres, coyotes, felinos, etc., que se puedan encontrar dentro del área.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Además, se instalarán cámaras trampa con la finalidad de tener información del tránsito de especies en el área y de conocer las especies que no se observan, con esta información se ayuda a la elección de las técnicas y de las estrategias para la captura de las especies.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen2.png', width=Cm(7.34), height=Cm(5.71))
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen3.png', width=Cm(6.97), height=Cm(5.51))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 2.- Trampas tipo Sherman (izquierda) y trampas tipo Tomahawk (derecha).')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
    ############################
    ### HERPETOFAUNA ###
    ############################
    """
    #########################
    ### Descripcion del capitulo 10.2.4.4.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nHERPETOFAUNA')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    descripcionCapitulo10.italic = True
    di10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Los reptiles y anfibios se conocen en conjunto como herpetofauna, este tipo de individuos se encuentran en diferentes ambientes y son considerados especies de lento desplazamiento y difícil observar cómo especie sobre todo aquellos de talla pequeña, su avistamiento también dependerá de la temporada en la que se encuentre, así como la actividad de las especies ya que algunas son de actividades diurnas, nocturnas o crepusculares. Los métodos que se describen a continuación aplican para las dos clases de vertebrados (anfibios y reptiles) para su captura y manipulación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo10 = di10.add_run('Captura directa: Consiste en la captura de los individuos manualmente o con ayuda de instrumentos, los cuales pueden ser redes con cabo de madera o metal, lanzas de cuerda delgada sujeta al extremo de una vara o caña de pescar para el caso de lagartijas; y gancho o pinzas herpetológicas, etc., por mencionar algunos de los instrumentos. Esta actividad se llevará a cabo mediante recorridos de 1000 a 2000 mts, se revisarán caminos, veredas, encharcamientos, lugares con vegetación muerta acumulada.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen4.png', width=Cm(13.46), height=Cm(6.32))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 3.- Ejemplo de captura directa con redes. Fuente: Casas-Andreu et al. 1991.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen5.png', width=Cm(12.72), height=Cm(5.68))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 4.- Uso de lanzas.  Fuente: Vanzolini y Nelson 1990.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen6.png', width=Cm(7.76), height=Cm(6.32))
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen7.png', width=Cm(6.13), height=Cm(6.64))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 5.- Uso de Ganchos herpetológicos. Fuente: Ferri 1992')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen8.png', width=Cm(12.72), height=Cm(6.98))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 6.- Embolsado de una serpiente para transporte. Fuente: Knudsen 1972.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 10.2.4.4.2 ###
    #########################
    di10 = doc.add_paragraph(style='List Bullet')
    descripcionCapitulo10 = di10.add_run('Captura con trampas y barreras de desvío: Consiste de una barrera física de metal o de plástico de 50 cm hasta 1 m de ancho que impide el libre tránsito de los animales que deambulan en el área, atrapándolos al entrar a las trampas de foso (cubetas de 20 lts) con tapa de embudo enterradas al ras del suelo y de embudo colocadas junto a la barrera. Cada barrera tendrá una longitud de 300 mts, las bayas se verificarán cada dos horas despues de su instalación, los organismos capturados seran colectados para su verificación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.4.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_2442/imagen9.png', width=Cm(13.34), height=Cm(6.05))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.4.4.2 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 7.- Captura con barreras y trampas terrestres. Fuente: Modificado de Heyer et al 2001.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 10.2.4.4.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Todas las especies sean aves, mamíferos, reptiles y anfibios principalmente estos dos últimos grupos, de los cuales se requiera de su manipulación manual o directa sin la utilización de alguna herramienta de captura se realizará de forma en que todos los individuos serán considerados como peligrosos, agresivos, venenosos (tóxicos); esto con la finalidad de aumentar la seguridad e integridad del personal y de los individuos capturados.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Antes de cada captura y manipulación el personal utilizará obligatoriamente guantes de látex ya que algunas especies pueden ser susceptibles a reacciones con el contacto de sudor, protector solar; etc., por mencionar algunas. La manipulación será breve y únicamente para extraer a los individuos del área de estudio (ACUSTF), identificarlos y reubicarlos en el área propuesta para su reubicación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.5 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.5. Transporte de Individuos.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Los organismos capturados deberán ser depositados de forma separada, utilizando el material correspondiente y adecuado a cada grupo faunístico.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    transporteIndividuosTitulo = [
        'Transporte de individuos para el grupo de las aves.',
        'Transporte de individuos para el grupo de los mamíferos.',
        'Transporte de individuos para el grupo de los reptiles y anfibios.',
    ]

    transporteIndividuosDescripcion = [
        'Para el grupo de las aves en el área sujeta a cambio de uso de suelo no se requiere de acciones de captura y transportación, pero en dado caso de que este tipo de acciones se requiera se utilizará para su transporte jaulas y se recurrirá a la resguardo de los árboles o arbustos con presencia de nido hasta que los individuos salgan del nido, en caso de que no sea factible esta actividad se efectuará a la reubicación de nidos, se utilizarán guantes para esta acción para evitar la impregnación de olor propio al nido ya que puede ser un factor para que los padres abandonen el nido. La liberación de estos individuos será de forma inmediata.',
        'Para el grupo de los mamíferos estos serán transportados en las jaulas en las que se encuentren capturados sin retirarlos de las mismas, en caso de que la captura sea manual se colocarán en las trampas tipo Tomahawk, Sherman o Pet Carrier, el tamaño dependerá de la talla del individuo. Durante su transporte se evitará que las trampas se encuentren al contacto directo con la luz solar, calor o frio extremo, se colocara una tela sobre de ellas para minimizar el estrés de los individuos, se evitara también que los individuos permanezcan mucho tiempo en las trampas, posteriormente se procederá a su liberación.',
        'Para el caso del grupo de los reptiles y anfibios  estos serán transportados en sacos o bolsas de manta completamente sellados, contenedores transparentes, recipientes de plástico, contenedores tipo cubeta de 20 lts de capacidad, bolsas tipo ziploc con orificios para su ventilación, el material utilizado dependerá del individuo capturado, se tendrá sumo cuidando  a la hora de su captura y transporte, se revisara que los individuos, para el  caso de los anfibios, cuenten con  la humedad suficiente para evitar que estos se deshidraten (sustrato húmedo “peat moss” ) y mueran, también se debe mantener un flujo adecuado del aire.',
    ]

    transporteIndividuosTituloRango = range(len(transporteIndividuosTitulo))
    transporteIndividuosDescripcionRango = range(len(transporteIndividuosDescripcion))

    for cap10_244 in transporteIndividuosTituloRango:
        di10 = doc.add_paragraph()
        descripcionCapitulo10 = di10.add_run(f'{transporteIndividuosTitulo[cap10_244]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        descripcionCapitulo10.bold = True
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        di10 = doc.add_paragraph()
        descripcionCapitulo10 = di10.add_run(f'{transporteIndividuosDescripcion[cap10_244]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para cada individuo capturado se deberán tomar las medidas morfológicas y registrar cada una de sus características principales, así como también se le tomarán una serie de fotografías para su pronta identificación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Imagen del capítulo 10.2.4.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Favor de Poner la Imagen, remplazando este texto =)')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la Imagen del capítulo 10.2.4.5 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 8.- Registro, transporte e identificación de individuos.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 10.2.4.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.6 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.5.- Liberación y Reubicación de los Individuos.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.6 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Previo a la liberación, se debe analizar el estado de salud (si se encuentran lastimados, heridos etc.,) en el que se encuentran los individuos. Se deberá tratar de identificar la especie a la que pertenece o bien fotografiar el ejemplar, en el caso en que se desconozca su identidad específica.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Las especies de mamíferos deberán ser liberadas durante el crepúsculo o durante la noche, ya que algunas especies requieren de su pronta liberación tales como los pequeños roedores (ratones), por lo contrario, las especies de reptiles deberán de ser liberados durante el día o durante el crepúsculo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El área donde se llevará a cabo la reubicación será en un área con las condiciones similares a las del área de estudio ACUSTF, considerando como factores importantes en el sitio de destino:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    listaCap10_245 = [
        'Vegetación.',
        'Disponibilidad de agua.',
        'Altitud (msnm).',
        'Grado de conservación.',
        'Seguridad de los ejemplares y las personas.',
    ]

    listaCap10_245Rango = range(len(listaCap10_245))

    for lista in listaCap10_245Rango:
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{listaCap10_245[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Dichos factores deberán tener condiciones similares a las del sitio original, evitando en la medida de lo posible, la sobrecarga (tolerancia de un ecosistema al uso de sus componentes sin rebasar su capacidad de recuperación) del sitio. Otro punto importante a ser considerado será que los sitios para relocalización no se encuentren muy distantes del sitio de captura, evitando largos periodos de confinamiento, disminuyendo el estrés resultante de la manipulación del ejemplar. A tal fin se utilizarán las zonas aledañas a la “huella” del Proyecto.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.4.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.7 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.7.- Área de Reubicación de las Especies Rescatadas.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.7 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Describir en este parrafo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.2.4.7 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 10.- Coordenadas del área de rescate')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.5
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.4.7 ###
    #########################
    columnas = 3
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Mapa del capitulo 10.2.4.7 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo10 = doc.add_paragraph()
    imagenCapitulo10.text = '\n'
    imagenCapitulo10 = doc.add_picture('capitulo10/mapa.png')  # Ancho de la imagen en centimetros
    imagenCapitulo10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo10.width = Cm(15.59)  # Ancho de la imagen en centimetros
    imagenCapitulo10.height = Cm(10.16)  # Alto de la imagen en centimetros
    imagenCapitulo10.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 10.2.4.7 ###
    #########################
    diMap10 = doc.add_paragraph()
    descripcionCapituloMapa10 = diMap10.add_run('Describir el Mapa =)')
    descripcionCapituloMapa10_format = diMap10.paragraph_format
    descripcionCapituloMapa10_format.line_spacing = 1.15
    descripcionCapituloMapa10.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa10.font.name = 'Bookman Old Style'
    descripcionCapituloMapa10.font.size = Pt(12)
    descripcionCapituloMapa10.font.italic = True
    diMap10.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    ########################################################################################################################################################################
    # Capitulo 10.2.4.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.4.8 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.4.8.- Indicadores de Éxito.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.4.8 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Con el fin de dar seguimiento y asegurar el cumplimiento de los objetivos del Programa se evaluarán periódicamente los indicadores de éxito siguiente:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    indicadoresCap10_248 = [
        'Se realizan recorridos periódicos en las áreas de liberación y conservación de las especies.',
        'Monitoreo de las especies rescatadas para su verificación.'
    ]

    for lista in range(len(indicadoresCap10_248)):
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{indicadoresCap10_248[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('A fin de cumplir con el programa de trabajo, se aplicará cada una de las actividades señaladas en la calendarización poniendo especial cuidado en el seguimiento de la ejecución a través de los indicadores de éxito establecidos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para el estudio y conocimiento de la especie se aplicarán y replicarán las metodologías- técnicas que han permitido a la fecha incrementar el acervo y conocimiento de la especie.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Así también se tendrá en consideración lo siguiente:')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    consideracionCap10_248 = [
        'Implementación de pláticas de educación ambiental, participación social y cultural para generar una comunicación y difusión que resalten la importancia ecológica de las especies, entre los operadores tanto en el área de trabajo como en áreas de trasporte y predios aledaños.',
        'Informar a las personas involucradas en las etapas del proyecto, sobre la importancia de la o las especies y su conservación para el medio ambiente y a su vez se tomen medidas drásticas para la protección, cuidados y tratos especiales para la especie.',
        'Disminuir la velocidad  de  vehículos, máquinas, etc., en vías de entrada y salida cercanas y pertenecientes al área del proyecto, reduciendo y evitando el impacto ocasionado por viajar a velocidades altas, en donde las especies de  reptiles, mamíferos y aves puedan ser atropellados, así mismo no solo se beneficiará a estos individuos si no también será beneficiada la fauna en general, disminuyendo la mortandad de las especies durante el ciclo de vida del proyecto, tales como pequeños y medianos mamíferos y algunas aves, incluyendo la especie en cuestión.',
        'Quedará estrictamente prohibida la caza, captura y destrucción de refugios de anidamiento en el área de extracción, predio y área de influencia del proyecto.',
    ]

    for lista in range(len(consideracionCap10_248)):
        di10 = doc.add_paragraph()
        descripcionCapitulo10 = di10.add_run(f'{lista + 1}.- {consideracionCap10_248[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.5 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\n')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Como parte del programa de trabajo a estudios posteriores dentro del áreas del proyecto y de las acciones aplicadas conforme a las etapas correspondientes será necesario cumplir con la Normatividad Ambiental para prevenir y mitigar los daños causados al medio ambiente derivado a las actividades que implica el proyecto _____________________________, ejecutando programas de rescate de especies de fauna silvestre en el tiempo que conlleva la preparación del sitio.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El programa de rescate deberá de realizarse de manera previa y durante las actividades de cambio de uso de suelo y antes de los trabajos de desmonte y despalme de área correspondiente al ACUSTF.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 10.2.5 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 11.- Periodo de duración de las actividades del proyecto.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.5
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.5 ###
    #########################
    columnas = 3
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 10.2.5 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 12.- Calendario de actividades Etapa de Preparación, construcción y operación el sitio del proyecto.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.5
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.5 ###
    #########################
    columnas = 17
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 10.2.5 ###
    #########################
    tituloTabla10b = doc.add_paragraph()
    dti10b = tituloTabla10b.add_run('\nTabla 13.- Calendario de actividades correspondientes al programa de rescate de fauna silvestre.')
    dti10b_format = tituloTabla10b.paragraph_format
    dti10b_format.line_spacing = 1.5
    dti10b_format.space_after = 0

    dti10b.font.name = 'Bookman Old Style'
    dti10b.font.size = Pt(12)
    tituloTabla10b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 10.2.5 ###
    #########################
    columnas = 17
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 10.2.5 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Describir este parrafo (Opcional).')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.6 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.6.- Resultados.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.6 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Los resultados derivados del rescate de las especies de fauna silvestre se obtendrán después de ejecutar las actividades propuestas dentro del calendario de actividades, se entregará un informe en el cual se muestren los resultados obtenidos durante ese periodo de tiempo. En el informe se plasmarán las actividades que se realizaron y los individuos que fueron capturados durante el periodo de calendarización respaldados por un anexo fotográfico.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('El informe final o de finiquito se entregará en el periodo de plazo otorgado en la autorización para la remoción de la vegetación forestal, en este documento se plasmará información más detallada sobre todas las actividades realizadas, así como también se plasmará el resultado del total de los individuos capturados, gráficas y fotografías que demuestren lo descrito en dicho documento.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.7
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.7 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.7.- Formato de Rescate de Fauna Silvestre.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.7 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Formato de registro que se utilizará para las actividades de rescate de fauna silvestre.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Salto de Linea ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Imagen del capítulo 10.2.7 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    imagenCapitulo10_parrafo = doc.add_paragraph()
    imagenCapitulo10_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    imagenCapitulo10_run = imagenCapitulo10_parrafo.add_run('')
    imagen_cap_10 = imagenCapitulo10_run.add_picture('capitulo10/capitulo10_27/formato.png', width=Cm(16.17), height=Cm(21.73))

    # Opcional: espacio después del párrafo
    imagenCapitulo10_parrafo.space_after = Pt(1)

    #########################
    ### Título de la Imagen del capítulo 10.2.7 ###
    #########################
    tituloGrafico10 = doc.add_paragraph()
    dgi10 = tituloGrafico10.add_run('Imagen 9.- Formato de registro de captura de las especies de fauna silvestre.')
    dgi10_format = tituloGrafico10.paragraph_format
    dgi10_format.line_spacing = 1.15
    dgi10_format.space_after = 0

    dgi10.font.name = 'Bookman Old Style'
    dgi10.font.size = Pt(12)
    tituloGrafico10.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 10.2.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.8 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.8.- Evaluación y monitoreo de especies rescatadas.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.8.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.8.1 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.8.1.- Tipo de muestreo.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    """
        #########################
        ### Reptiles ###
        #########################
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Reptiles')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||
        || Captura Oportunista ||
        |||||||||||||||||||||||||
    """

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Colecta oportunista: Consiste en una búsqueda sistemática de organismos a diferentes horas del día o estaciones del año. se pueden hacer recorridos diurnos y nocturnos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||
        || Captura Directa ||
        |||||||||||||||||||||
    """

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Captura directa: Colectas nocturnas, se pueden capturar con redes en el caso de aves y murciélagos')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    capturaDirecta10_281 = [
        'Captura oportunista',
        'Abundancia relativa, riqueza de especies',
        'Bajo',
        'Bajo',
        'Bajo a Medio',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{capturaDirecta10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    """
        ||||||||||||||||||||||||
        ||  Encuentro Visual  ||
        ||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nEncuentro visual: Consiste en la observación y conteo de los organismos a lo largo de trayectos de distancia fija o aleatorios durante un periodo de tiempo fijo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    encuentroVisual10_281 = [
        'Encuentro Visual',
        'Abundancia relativa, riqueza de especies',
        'Bajo',
        'Bajo',
        'Bajo',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{encuentroVisual10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    """
        ||||||||||||||||||||||||||||||||||
        ||  Colecta de tiempo limitado  ||
        ||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nColecta de tiempo limitado: es la búsqueda para la captura de organismos incidiendo en un ambiente o micro ambiente especifico, en un tiempo determinado.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    limitedTime10_281 = [
        'Colecta de tiempo limitado',
        'Abundancia relativa, riqueza de especies',
        'Medio',
        'Bajo',
        'Medio',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{limitedTime10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    """
        ||||||||||||||||||
        ||  Transectos  ||
        ||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nTransectos: Son recorridos de longitud previamente establecidos que permiten evaluar diferencias faunísticas entre varias áreas (gradientes topográficos, gradientes de hábitat, zonas con diferentes tipos de vegetación, etc.).')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    transectos10_281 = [
        'Transectos',
        'Abundancia relativa, riqueza de especies',
        'Medio a Alto',
        'Bajo',
        'Bajo a Medio',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{transectos10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ||||||||||||||||||
        ||  Cuadrantes  ||
        ||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nCuadrantes: Son áreas delimitadas del terreno de tamaño conocido para identificar y contar a todos los individuos que estén presentes. los resultados que se puedan obtener dependen del tamaño, forma, y numero de cuadrantes utilizados y si el hábitat es homogéneo o heterogéneo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    cuadrantes10_281 = [
        'Cuadrantes',
        'Abundancia relativa, riqueza de especies',
        'Alto',
        'Bajo a Medio',
        'Medio a Alto',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{cuadrantes10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ||||||||||||||||||||||||||||||
        ||  Remoción de Individuos  ||
        ||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nRemoción de individuos: Extracción o marcado de los individuos para contabilizarlos en un área dada por lo menos una hectárea de superficie, método efectivo para especies muy visibles o de fácil captura.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 5
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo de Evaluación',
        'Informacion a obtener',
        'Inversión de tiempo',
        'Costo',
        'Requerimiento de Personal',
    ]

    remocionIndividuos10_281 = [
        'Remoción de Individuos',
        'Abundancia relativa, riqueza de especies',
        'Medio a Alto',
        'Medio',
        'Medio a Alto',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{remocionIndividuos10_281[content]}')
        t10b.font.size = Pt(12)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
        ||  Materiales que se pueden utilizar para la captura y manipulación de los individuos recolectados. ||  
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nMateriales que se pueden utilizar para la captura y manipulación de los individuos recolectados.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    materialesRecolectados = [
        'Redes',
        'Lanza de cuerda delgada',
        'Ganchos herpetológicos',
        'Trampas y barreras de desvios',
        'Trampas de fosos y de cilindro',
        'Guías para la identificación',
    ]

    for lista in range(len(materialesRecolectados)):
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{materialesRecolectados[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    """
        #########################
        ### Aves ###
        #########################
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nAves')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ///////////////////////////////////////////
        // Recuento en punto (puntos estáticos). //
        ///////////////////////////////////////////
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Recuento en punto (puntos estáticos).')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||
        || Sin estimado de distancia ||
        |||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Sin estimado de distancia: Las aves detectadas se cuentan sin tomar en consideración su distancia del observador, cuenta con un radio limitado, no puede usarse para estimar densidad.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    cell = tabla10b.cell(1, 0)
    t10b = cell.paragraphs[0].add_run('Recuento en punto\n')
    t10b.font.size = Pt(10)
    t10b.font.name = 'Arial'
    t10b.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    estimadoDistancia10_281 = [
        'Sin estimado de distancia',
        'X',
        'X',
        'X',
        ' ',
        'X',
        ' ',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{estimadoDistancia10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ||||||||||||||||||||
        || Radio variable ||
        ||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nRadio variable: Se estima y se mide la distancia que lo separa del ave detectada.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    radioVariable10_281 = [
        'Sin estimado de distancia',
        'X',
        'X',
        'X',
        ' ',
        'X',
        ' ',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{radioVariable10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    """
        ||||||||||||||||||
        ||  Radio fijo  ||
        ||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Radio fijo: Se detectan aves en un círculo de radio fijo, el tamaño del radio dependerá de la cantidad de la vegetación. un radio estándar es entre 25 y 30 metros y durante un periodo entre 5 y 10 minutos.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    radioFijo10_281 = [
        'Sin estimado de distancia',
        'X',
        'X',
        'R',
        ' ',
        'X',
        ' ',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{radioFijo10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ////////////////
        ////////////////
        // Transectos //
        ////////////////
        ////////////////
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nTransectos')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    """
        ||||||||||||||||||||||||||||||||||
        ||  Sin estimado de distancias  ||
        ||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Sin estimado de distancias: Es la forma más sencilla de censo, permite al observador generar una lista de especies que se encuentren presentes en el hábitat. no se puede estimar densidades.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    estimadoDistancia10_281 = [
        'Sin estimado de distancia',
        'X',
        'X',
        'X',
        ' ',
        ' ',
        ' ',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{estimadoDistancia10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    """
        |||||||||||||||||||||||||||
        ||  Transecto de franja  ||
        |||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Transecto de franja: Se registran a las aves detectadas mediante se camina a través de una línea recta y estableciendo franjas de ancho fijo (w). equidistancia entre 25-50 metros dependiendo de la densidad de la vegetación.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    transectoFranja10_281 = [
        'Sin estimado de distancia',
        'X',
        'X',
        'X',
        'X',
        ' ',
        ' ',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{transectoFranja10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ||||||||||||||||||||||||||||||||||||||||||
        ||  Representación en mapa estadístico  ||
        ||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Representación en mapa estadístico: Se basa en la conducta territorial de las aves, consiste en marcar sobre un mapa la posición de las aves observadas en visitas consecutivas a la parcela.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    mapaEstadistico10_281 = [
        'Sin estimado de distancia',
        'E',
        'E',
        'E',
        ' ',
        'R',
        'X',
        ' ',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{mapaEstadistico10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        ||||||||||||||||||||||||||||||||||||||||
        ||  Capturas con redes ornitológicas  ||
        ||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Capturas con redes ornitológicas: También conocidas como redes de niebla o redes japonesa han sido utilizadas para la captura de aves durante años y se han convertido en herramientas efectivas para el monitoreo de las poblaciones. permite generar información sobre datos demográficos de la población.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.8.1 ###
    #########################
    columnas = 8
    filas = 2

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    columnas10_281 = [
        'Metodo',
        'Presencia de Especie',
        'Abundancia Relativa',
        'Tendencia poblacional',
        'Densidad',
        'Uso de habitat',
        'Condición',
        'Supervivencia',
    ]

    mapaEstadistico10_281 = [
        'Sin estimado de distancia',
        ' ',
        ' ',
        'X',
        ' ',
        ' ',
        'R',
        'X',
    ]

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')
        t10b = cell.paragraphs[0].add_run(f'{columnas10_281[cols]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        t10b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for content in rangoColumnas:
        cell = tabla10b.cell(1, content)
        t10b = cell.paragraphs[0].add_run(f'{mapaEstadistico10_281[content]}')
        t10b.font.size = Pt(10)
        t10b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        |||||||||||||||||||||||||||||||||||||||||||||||
        ||  Grabación de sonidos o muestreo acústico ||  
        |||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Grabación de sonidos o muestreo acústico: Técnica utilizada para atraer las aves al observador de tal manera que pueden ser identificadas visualmente o como técnica de monitoreo y grabar a las especies que no son detectadas visualmente o especies raras difíciles de observar.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.8.1 ###
    #########################
    """
        #########################
        ### Mamiferos ###
        #########################
    """
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('\nMamiferos')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    descripcionCapitulo10.bold = True
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    herramientasMamiferos = [
        'Trampas de acero y resorte: se utilizan trampas de acero de fabricación comercial o también conocidos como cepos.',
        'Trampas caja: Se pueden capturar a pequeños y medianos mamíferos dependiendo del tamaño de la trampa.',
        'Trampas olfativas: Área limpia de hojas, rocas en un área de 1 m2, en la cual se cierne arena o cal para que quede un sustrato donde se pueda registrar huellas, dentro de la trampa lleva un cebo utilizado como atrayente.', 
        'Trampas corral: Se utiliza para la captura de caza mayor, en donde los animales son conducidos hacia estructuras permanentes construidas con madera o alambre utilizando cebos para la atracción.',
        'Trampas cámara: Técnica más recomendable para obtener tendencias y estimaciones confiables, estas pueden ser activas o pasivas',
        'Redes trampas: se pueden capturar venados con una especie de red carpa con mecanismo central y disparador, otros mamíferos como los murciélagos son capturados con redes de niebla.',
        'Cebos para trampas: Tipos de alimentos preparados o comerciales que atraen hacia la trampa a diferentes especies dependiendo del cebo o animal que se quiera capturar.',
    ]

    cebosTrampas = [
        'Esencias ',
        'Señuelos y otros medios de atracción'
    ]

    for lista in range(len(herramientasMamiferos)):
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{herramientasMamiferos[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lista in range(len(cebosTrampas)):
        di10 = doc.add_paragraph()
        descripcionCapitulo10 = di10.add_run(f'     o {cebosTrampas[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 10.2.8.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.8.2 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.8.2.2.- Resultados.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.8.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Describir en este parrafo.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 10.2.5 ###
    #########################
    columnas = 17
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in rangoColumnas:
        cell = tabla10b.cell(0, cols)
        cell_background_color(cell, 'D9D9D9')

        for rows in rangoFilas:
            cell = tabla10b.cell(rows, cols)
            t10b = cell.paragraphs[0].add_run(' ')
            t10b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 10.2.8.2 ###
    #########################
    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Los resultados derivados del Monitoreo de las especies de fauna silvestre se obtendrán después de ejecutar las actividades propuestas dentro del calendario de actividades, se entregará un informe en el cual se muestren los resultados obtenidos durante ese periodo de tiempo. En el informe se plasmarán las actividades que se realizaron y los individuos que fueron observados y/o capturados durante el periodo de calendarización respaldados por un anexo fotográfico. Para ello se utilizará la bitácora siguiente.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di10 = doc.add_paragraph()
    descripcionCapitulo10 = di10.add_run('Para el monitoreo de las especies ahuyentadas y/o rescatadas se realizarán las bitácoras de acuerdo a los periodos de monitoreo, cabe destacar que este formato está sujeto a cambio de acuerdo a las actividades a realizar.')
    descripcionCapitulo10_format = di10.paragraph_format
    descripcionCapitulo10_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo10.font.name = 'Arial'
    descripcionCapitulo10.font.size = Pt(12)
    di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Bitacora del capitulo 10.2.8.2 ###
    #########################
    """
        ============================================================
            El siguiente codigo muestra como hacer la bitacora
        ============================================================
    """

    columnas = 2
    filas = 5

    rangoColumnas = range(columnas)
    rangoFilas = range(filas)

    tabla10b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    """
        ***********************
        Ancho de las columnas
        ***********************
    """
    for width in tabla10b.rows:
        width.cells[0].width = Cm(12.5)
        width.cells[1].width = Cm(5.5)

    """
        ***********************
        Encabezados de la tabla
        ***********************
    """
    cell = tabla10b.cell(0, 0)
    t10b = cell.paragraphs[0].add_run('BITACORA DE MONITOREA DE FAUNA SILVESTRE')
    t10b.font.size = Pt(12)
    t10b.font.bold = True
    t10b.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_background_color(cell, 'EAF1DD')

    cell = tabla10b.cell(0, 1)
    t10b = cell.paragraphs[0].add_run()
    t10b.add_picture('capitulo10/capitulo10_2822/logo.png', width=Cm(4.3), height=Cm(2.15))
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for width in tabla10b.rows:
        width.cells[0].width = Cm(12.5)
        width.cells[1].width = Cm(5.5)

    """
        ***********************
        Datos del Animal
        ***********************
    """
    cell = tabla10b.cell(1, 0)
    t10b = cell.paragraphs[0].add_run('Localizacion general (poblado más cercano):'
                                      '\n_______________________________________________________\n'
                                      '_______________________________________________________\n'
                                      '\nLatitud: ________________           Longitud: __________________\n'
                                      '\nAltitud: ________________'
                                      '\n\nClave o No. asignado en campo: ________________________'
                                      '\n\nColector y/o observador:'
                                      '\n_______________________________________________________'
                                      '\n\nFamilia: ______________           Especie: ___________________'
                                      '\n\nDeterminó: ___________________________________________________'
                                      '\n\nCondiciones de captura y/o observado:'
                                      '\n_______________________________________________________'
                                      '\n_______________________________________________________'
                                      '\n\nMétodo de captura o avistamiento:'
                                      '\n_______________________________________________________'
                                      '\n')
    t10b.font.size = Pt(11)
    t10b.font.bold = True
    t10b.font.name = 'Arial'
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #cell_background_color(cell, 'EAF1DD')

    """
        ***********************
        Fecha y Hora
        ***********************
    """
    cell = tabla10b.cell(1, 1)
    t10b = cell.paragraphs[0].add_run('\nFecha: _______________'
                                      '\n\nHora: ________________')
    t10b.font.size = Pt(11)
    t10b.font.bold = True
    t10b.font.name = 'Arial'
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #cell_background_color(cell, 'EAF1DD')

    """
        ***********************
        Altura de celdas
        ***********************
    """
    fila = tabla10b.rows[2]
    fila.height = Cm(4.20)
    fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    fila = tabla10b.rows[3]
    fila.height = Cm(4.20)
    fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


    """
        ***********************
        Descripcion del Habitat
        ***********************
    """
    cell = tabla10b.cell(2, 0)
    t10b = cell.paragraphs[0].add_run('Descripción del hábitat:')
    t10b.font.size = Pt(11)
    t10b.font.bold = True
    t10b.font.name = 'Arial'
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #cell_background_color(cell, 'EAF1DD')

    """
        ***********************
        N° de Foto y Camara
        ***********************
    """
    cell = tabla10b.cell(2, 1)
    t10b = cell.paragraphs[0].add_run('\nN° de Foto:____________\n'
                                      '\nCámara: ______________')
    t10b.font.size = Pt(11)
    t10b.font.bold = True
    t10b.font.name = 'Arial'
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #cell_background_color(cell, 'EAF1DD')

    """
        ***********************
        Descripcion del Animal
        ***********************
    """
    ### Celda fusionada ###
    row = tabla10b.rows[3]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t9b = merged_cell.paragraphs[0].add_run('Descripción del animal: (Color, tamaño, etc.)')
    t9b.font.name = 'Arial'
    t9b.font.size = Pt(11)
    t9b.bold = True

    """
        ***********************
        Observaciones
        ***********************
    """
    row = tabla10b.rows[4]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t9b = merged_cell.paragraphs[0].add_run('Observaciones: (Número total de animales de la misma especie capturados y/o observados etc.)')
    t9b.font.name = 'Arial'
    t9b.font.size = Pt(11)
    t9b.bold = True

    ########################################################################################################################################################################
    # Capitulo 10.2.8
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 10.2.8 ###
    #########################
    capitulo10 = doc.add_paragraph()
    i10 = capitulo10.add_run(f'\nX.2.8.- Bibliografía.')
    i10_format = capitulo10.paragraph_format
    i10_format.line_spacing = 1.15

    i10.font.name = 'Arial'
    i10.font.size = Pt(12)
    i10.font.bold = True
    capitulo10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 10.2.8 ###
    #########################
    bibliografia = [
        'Casas–Andreu, G., G. Valenzuela–López y A. Ramírez–Bautista. (1991). Como hacer una colección de Anfibios y Reptiles. Instituto de Biología, UNAM, Cuadernos No. 10, México.',
        'Ferri, V. 1992. El libro de las serpientes de todo el mundo. Editorial de Vecchi, S.A., Barcelona.',
        'Flores-Villela, O., H. M. Smith y D. Chiszar. (2004). The history of herpetological exploration in Mexico. Bonner Zoologische Beiträge 3/4:311-335.',
        'Flores-Villela, O. & Garcia-Vázquez, Uri, O. (2014). Biodiversidad de reptiles en México. Revista Mexicana de Biodiversidad, Supl. 85: S467 S475, 2014.DOI: 10.7550/rmb.43236 ',
        'Gallina-Tessaro, S. & López-González, C. (2011). Manual de Técnicas para el Estudio de la fauna. Universidad Autónoma de Querétaro Instituto de Ecología, A. C.',
        'Heyer, E.R., M.A. Donnelly, R.W. McDiarmid, L.A.C. Hayek y M.S. Foster. (Eds.). (2001). Medición y monitoreo de la diversidad biológica. Métodos estandarizados para anfibios. Smithsonian Institution Press/ Editorial Universitaria de la Patagonia.',
        'Leopold, A.S. (1959). Fauna Silvestre de México. Pax México. Segunda edición. México.',
        'Knudsen, J. W. (1966). Biological techniques; collecting, preserving, and illustrating plants and animals. New York, Harper & Row.',
        'Oscar Sánchez y Ella Vázquez - Domínguez (editores). (1999). Diplomado en manejo de vida silvestre. Conservación y manejo de vertebrados del norte árido y semiárido de México. Comisión Nacional para el Conocimiento y Uso de la Biodiversidad, Dirección General de Vida Silvestre (INE-SEMARNAP), Servicio de Pesca y Vida Silvestre de los Estados Unidos de América (USFWS), Facultad de Ciencias Forestales (UANL). México.',
        'Valdez. R. y Ortega-S. A. (2014). Ecología y Manejo de fauna silvestre en México. Colegio de Postgraduados. Montecillo, Texcoco. Estado de México. (1 er. Ed.). Pp. 557.',
        'Vanzolini, P.E. y P. Nelson. (1990). Manual de recolección y preparación de animales (2da. ed.). Facultad de Ciencias, UNAM, México.',
    ]
    
    for lista in range(len(bibliografia)):
        di10 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo10 = di10.add_run(f'{bibliografia[lista]}')
        descripcionCapitulo10_format = di10.paragraph_format
        descripcionCapitulo10_format.line_spacing = 1.5
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo10.font.name = 'Arial'
        descripcionCapitulo10.font.size = Pt(12)
        di10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 10 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo10() # Crear el documento
