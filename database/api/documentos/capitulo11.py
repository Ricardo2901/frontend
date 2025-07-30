"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo11():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 11
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XI.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Indice de Tablas del Capitulo 11
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("ÍNDICE DE TABLA.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    # Capitulo 11
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 11 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'XI.- Identificación, Descripción y Evaluación de los Impactos ambientales.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En el presente estudio durante su desarrollo se han detectado impactos ambientales adversos que afectarán durante el desarrollo y la implementación del proyecto, con el propósito de revertir dichos impactos se han establecido medidas de prevención y/o mitigación y de ser posible algún proceso para la restauración del sitio. Estas medidas deberán de tomarse como una responsabilidad no solo como complemento del estudio en el sentido de que de su aplicación dependerá la vida útil del proyecto y la incorporación del área a su condición lo más cercana a su origen.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nX.1.- Programa de Rescate y Reubicación de Flora Silvestre')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'XI.1.1.- Metodología utilizada para la identificación de los impactos.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.1.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para evaluar e identificar los impactos ambientales derivados de las acciones implícitas del cambio de uso del suelo se agruparán en dos facetas principalmente que nos podrán proporcionar elementos para analizar y definir la afectación, pero lo más importante se podrá proyectar el cómo remediar o reducir dichos impactos.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.1.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.1.1.- Criterios.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    criterios11_111 = [
        'Caracterización ambiental y descripción del proyecto.',
        'Predicción y evaluación de impactos.',
    ]

    #########################
    ### Descripcion del capitulo 11.1.1.1 ###
    #########################
    for lista in range(len(criterios11_111)):
        di11 = doc.add_paragraph()
        descripcionCapitulo11 = di11.add_run(f'{lista + 1}) {criterios11_111[lista]}')
        descripcionCapitulo11_format = di11.paragraph_format
        descripcionCapitulo11_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo11.font.name = 'Arial'
        descripcionCapitulo11.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('La etapa de caracterización incluye la descripción del proyecto en cada una de las diferentes facetas mismas que se vincularán con la observación de campo, así como sus posibles efectos para su evaluación y caracterización ambiental.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para ello es necesario utilizar información actualizada y veraz del sitio donde se detectan los elementos que serán afectados, por lo que los recorridos de campo proporcionarán información misma que bajo el análisis de diferentes opiniones arrojarán los resultados más adecuados y confiables que reduzcan la afectación de los impactos que serán ocasionados bajo la objetividad de dar un uso diferente que mejora algunos aspectos pero que en determinado momento se regrese el sitio a la sustentabilidad y armonía ambiental del sitio con los elementos del sistema ambiental determinada donde se aplique el proyecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.1.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.1.2.- Descripción de la Obra.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.1.1.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('El principal objetivo es dar a conocer todas las actividades que serán necesarias para el desarrollo del proyecto desde sus inicios durante el proceso y hasta su término y/o abandono del sitio de acuerdo al calendario de actividades, recursos físicos y humanos que se requieran y apliquen, así como las condiciones que se tengan en el sitio antes durante y posterior al término ya que será de acuerdo a su proyección el punto de inicio para buscar su recuperación parcial más nunca la original. Dada la importancia del proyecto que se pretende establecer aun con el deterioro de ciertos factores (suelo, vegetación) desde el corto hasta mediano plazo modificará algunas condiciones ambientales en los predios pretendiéndose que estos se puedan mitigar con acciones de restauración al término de la operación, en un largo plazo existirán para la biodiversidad al eliminar vegetación y posibles sitios para nichos de la fauna, estos últimos de carácter biótico pueden ser atenuados con el rescate de especies y la restauración al término del proyecto. ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.1.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.1.3.- Caracterización Ambiental.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.1.1.3 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En este punto se describen  los elementos físicos, biológicos perceptuales y socioeconómicos en términos generales que pueden ser modificados dando a conocer sus características y el nivel de afectación en cada una de las etapas y el factor que será alterado, teniendo como objetivo principal hacer del conocimiento sobre las condiciones actuales del área sujeta de estudio, las modificaciones que va a sufrir el área donde se desarrollará el proyecto; sin embargo la relevancia será el beneficio desde el punto de vista económico a través del cual se contempla mejorar las condiciones para la población en lo referente a calidad de vida en la localidad dando un uso más productivo al terreno sujeto de estudio en virtud de que actualmente no proporciona ningún tipo de beneficio a los usufructuarios que generen economías para el desarrollo, con la aplicación de medidas de restauración al término de su aprovechamiento se podrán crear condiciones ambientales similares a las actuales o aquellas que se puedan mejorar desde el punto de vista ambiental.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.5
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.1.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.1.4.- Predicción y evaluación de impactos.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.1.1.4 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Esta información se obtiene de diferentes trabajos de campo cotejado con trabajos similares al respecto fortaleciéndose con observaciones efectuadas en el entorno del proyecto, esto y con base a la experiencia con lo cual se obtienen elementos para proyectar la remediación visualizando la recuperación de algunos factores del sitio objeto de modificación.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Dentro del aspecto socioeconómico se identifican los intereses de los sectores sociales, esto permite acentuar las problemáticas ambientales que puedan ser ocasionadas por el proyecto, para el caso en particular se contempla puedan ser más los beneficios con los cuales la población pueda obtener mejores condiciones de bienestar para su desarrollo, reflejándose en la generación de empleo, mediante la transformación u obtención de diferentes minerales.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En forma mediática en este aspecto en la economía se verá reflejada desde el punto de vista de generación de empleos directos e indirectos mejorando las condiciones de vida para la población aledaña, así como la reducción de costos en los aspectos productivos.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Por tal motivo es recomendable efectuar los muestreos y análisis necesarios dependiendo de las características del proyecto y sus atributos ambientales.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Al identificar y evaluar los impactos se incorporan y analizan los resultados obtenidos en la fase de caracterización ambiental y características del proyecto. ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Los objetivos en esta fase son:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    objetivos11_114 = [
        'Identificar los posibles impactos asociados con el proyecto. ',
        'Proporcionar de ser posible, algunas predicciones cuantitativas y cualitativas de los efectos de los impactos identificados.',
        'Estructurar medidas preventivas y de mitigación.',
        'Revertir los impactos generados con el proyecto.',
    ]

    for lista in range(len(objetivos11_114)):
        di11 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo11 = di11.add_run(f'{objetivos11_114[lista]}')
        descripcionCapitulo11_format = di11.paragraph_format
        descripcionCapitulo11_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo11.font.name = 'Arial'
        descripcionCapitulo11.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.1.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.2.- Justificación de la Metodología.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.1.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('El mayor peso específico aportado para la utilización de esta metodología, es la identificación de los indicadores ambientales y su interacción real con las actividades del proyecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Por otra parte, permite diversificar las opiniones que representen un respaldo confiable a este método, utilizando diversos criterios, pertenecientes a personas con diferentes disciplinas que participaron en el presente estudio. No obstante, para que estas matrices tengan validez, fue necesario que se basaran en el análisis del sistema ambiental utilizado, debiéndose exponer en una explicación de los impactos identificados, el valor de los mismos, y las medidas de mitigación y control.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Una herramienta de gran utilidad para definir un problema es el checklist  ya que a través de esta se utiliza para identificar información específica y para el caso del impacto ambiental complementar la descripción de un problema al requerirse respuesta a diferentes preguntas tales como: ¿Cuál es el problema, ¿qué afectará, ¿dónde afectará, ¿por qué se afectará ¿con qué frecuencia y ¿qué posibilidades hay de recuperación?, de tal o cual factor ambiental que se vea alterado en el desarrollo del proyecto. Estas interrogantes son utilizadas en forma complementaria para identificar los impactos ambientales que se ocasionarán en las diferentes etapas del proyecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Una vez analizadas las actividades que comprende el proyecto y las que estén asociadas a este dentro del sistema ambiental, se observa que no hay proyectos que puedan sumarse o sean acumulativos con las actividades del proyecto, para lo cual se utilizó para la identificación y valoración la Matriz de Conessa  que es la utilizada, cuenta con varias ventajas ya que puede ayudar a identificar impactos positivos y negativos, puede usarse para identificar impactos en varias fases temporales del proyecto y para describir los impactos asociados a varios ámbitos identificados.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Tiene la ventaja que permite la estimación de los impactos mediante una escala numérica, la comparación de alternativas, la determinación de interacciones, la identificación de acciones del proyecto que causan impactos de menor o mayor impacto e importancia, además es la que más se adecua a las condiciones del proyecto que combinadas con los juicios técnicos del personal participante basados en las observaciones de campo y experiencia con lo cual se da una adecuada interpretación de los impactos identificados para dictar las medidas de mitigación necesarias.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.2.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.2.1.- Metodología de Evaluación.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.2.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para identificar los posibles impactos ambientales que pudiera ocasionar el desarrollo del proyecto se elabora un listado simple de factores y componentes ambientales, así como de acciones causales de impacto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('De acuerdo a los resultados se elaboró y utilizó la Matriz de Conessa, que toma en cuenta procedimientos paralelos analizando el proyecto y su entorno, el cruce de ambos análisis nos proporciona la identificación de los impactos que nos arroja información para catalogar los efectos en un nivel cualitativo y cuantitativo en forma porcentual y que a la vez conlleva a definir acciones para mitigar dichos efectos como resultado de la aplicación del proyecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.2.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.1.2.2.- Listado Simple.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.2.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('A través de verificación se identificaron los factores y componentes ambientales susceptibles de ser impactados, así como las acciones causales de impacto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Se identificaron 20 componentes agrupados en 10 factores ambientales con susceptibilidad de afectación por las acciones o actividades que involucra la obra. Así mismo se identificaron ____________ agrupadas en ________ para el desarrollo del proyecto las cuales son la Preparación del Sitio, Construcción, Operación, Adicionalmente se incluyen ___ acciones al ___________________ considerando los impactos positivos más de carácter benéfico aun cuando hay componentes sobre los que se manifiesta impacto por acciones derivadas de la restauración del área afectada. ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.2.- Identificación de Impactos.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para identificar los posibles impactos ambientales que podría ocasionar el desarrollo del proyecto, se registra un listado simple derivado de cada acción del proyecto, esto considerando el análisis (in situ) y la proyección de lo que se pretende con el cambio de uso del suelo.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.1.- Componentes ambientales a afectar.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.2 ###
    #########################
    columnasCap11_2 = [
        'SUBSISTEMA',
        'FACTOR AMBIENTAL',
        'COMPONENTE A AFECTAR'
    ]

    componenteCap11_2 = [
        'Calidad del aire',
        'Visibilidad',
        'Ruido',
        'Escorrentías',
        'Calidad',
        'Capacidad de Recarga',
        'Propiedades físico-químicas',
        'Erodabilidad',
        'Relieve',
        'Hábitat',
        'Abundancia',
        'Spp. En Estatus.',
        'Hábitat',
        'Abundancia',
        'Spp. En Estatus',
        'Calidad',
        'Visibilidad',
        'Fragilidad',
        'Nivel de Ingresos',
        'Calidad de vida',
    ]

    filas = len(componenteCap11_2) + 1
    columnas = len(columnasCap11_2)

    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    #########################
    ### Celdas fusionadas de Subsistema ###
    cell_top = tabla11b.cell(1, 0)
    cell_bottom = tabla11b.cell(9, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('1. Abiotico')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True

    ###############################
    cell_top = tabla11b.cell(10, 0)
    cell_bottom = tabla11b.cell(15, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('2. Biotico')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True

    ###############################
    cell_top = tabla11b.cell(16, 0)
    cell_bottom = tabla11b.cell(18, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('3. Perceptual')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True

    ###############################
    cell_top = tabla11b.cell(19, 0)
    cell_bottom = tabla11b.cell(20, 0)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('4. Socioeconomicos')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True


    #########################
    ### Celdas fusionadas de Factor Ambiental ###

    cell_top = tabla11b.cell(1, 1)
    cell_bottom = tabla11b.cell(3, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('1. Atmosfera')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    ###############################
    cell_top = tabla11b.cell(4, 1)
    cell_bottom = tabla11b.cell(5, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('2. Hidrologia Superficial')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    ###############################
    cell = tabla11b.cell(6, 1)
    t11b = cell.paragraphs[0].add_run('3. Hidrologia Subterranea')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ###############################
    cell_top = tabla11b.cell(7, 1)
    cell_bottom = tabla11b.cell(8, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('4. Suelo')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    ###############################
    cell = tabla11b.cell(9, 1)
    t11b = cell.paragraphs[0].add_run('4. Topografía')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ###############################
    cell_top = tabla11b.cell(10, 1)
    cell_bottom = tabla11b.cell(12, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('6. Fauna Silvestre')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    ###############################
    cell_top = tabla11b.cell(13, 1)
    cell_bottom = tabla11b.cell(15, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('7. Flora Silvestre')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    ###############################
    cell_top = tabla11b.cell(16, 1)
    cell_bottom = tabla11b.cell(18, 1)

    merged_cell = cell_top.merge(cell_bottom)
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Agregar texto (opcional)
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('8. Paisaje')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    ###############################
    cell = tabla11b.cell(19, 1)
    t11b = cell.paragraphs[0].add_run('9. Economia')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ###############################
    cell = tabla11b.cell(20, 1)
    t11b = cell.paragraphs[0].add_run('10. Poblacion')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    #########################
    ### Celdas de componente a afectar ###
    for componente in range(len(componenteCap11_2)):
        cell = tabla11b.cell(componente + 1, 2)
        t11b = cell.paragraphs[0].add_run(f'{componente + 1}. {componenteCap11_2[componente]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    
    #########################
    ### Columnas de la tabla ###
    for cols in range(len(columnasCap11_2)):
        cell = tabla11b.cell(0, cols)
        t11b = cell.paragraphs[0].add_run(f'{columnasCap11_2[cols]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        t11b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '#3498db')

    #########################
    ### Descripcion del capitulo 11.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Una vez identificados los posibles impactos con base al Subsistema, Factor Ambiental y Componente se registran las diferentes etapas del proyecto con las acciones que se tienen contempladas implementar para el desarrollo del proyecto de acuerdo al siguiente cuadro:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.3.- Caracterización de los Impactos.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.3 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('DescripEn este aspecto se analiza  cada una de las etapas del proyecto con las acciones requeridas aplicables identificando los principales subsistemas en los cuales están inmersos los factores y componentes ambientales susceptibles de ser impactados, así mismo se relacionan las acciones causales de impacto, con este análisis se aplican metodologías con el uso de matrices que determina la interacción de las acciones en cada etapa del proyecto, obteniendo una estimación subjetiva de los impactos mediante la aplicación de escala numérica, con ello se puede lograr la comparación de alternativas determinando las interacciones que facilitan la interpretación de los impactos.cion')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para ello se elaboró y utilizó la Matriz de Conessa, la cual toma en cuenta procedimientos paralelos, analizando el proyecto, por una parte y por el otro su entorno, el cruce de ambos análisis nos proporciona la identificación de los impactos y su valoración respectiva.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.2.- Etapas y acciones o Indicadores Ambientales.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.3 ###
    #########################
    columnas = 3
    filas = 8
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.- Valoración de los impactos generados.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Obtenida la matriz de identificación de impactos, se realiza una revisión y valoración de los mismos. En esta etapa del estudio, se medirá el impacto, sobre una base del grado de manifestación cualitativa que determina la magnitud del efecto, mismo que quedará reflejado en la que se define como importancia del efecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('La importancia (IM) es un indicador, que mide cualitativamente el impacto ambiental, con relación al grado de incidencia o intensidad de la alteración producida y de la caracterización del efecto, el cual responde a su vez a una serie de atributos de tipo cualitativo, tales como: Naturaleza, Intensidad, Extensión, Momento, Persistencia, Reversibilidad, Sinergia, Acumulación, Efecto, Periodicidad, Recuperabilidad.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Magnitud: Se considera como la extensión del impacto y es precedido por el signo más (+) o menos (-) que se refiere al carácter del impacto. Se asigna un valor numérico que varía de 1 a 10, donde 10 representa mayor magnitud respecto a 1.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 11.4 ###
    #########################
    columnasCap11_4 = [
        'Magnitud',
        'Valor',
    ]

    magnitudCap11_4 = [
        'Muy Baja Magnitud',
        'Baja Magnitud',
        'Mediana Magnitud',
        'Alta Magnitud',
        'Muy Alta Magnitud',
    ]

    valorCap11_4 = [
        1,
        3,
        5,
        7,
        10,
    ]

    columnas = len(columnasCap11_4)
    filas = len(magnitudCap11_4) + 1
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')
        t11b = cell.paragraphs[0].add_run(f'{columnasCap11_4[cols]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(10)
        t11b.font.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for maginitud in range(len(magnitudCap11_4)):
        cell = tabla11b.cell(maginitud + 1, 0)
        t11b = cell.paragraphs[0].add_run(f'{magnitudCap11_4[maginitud]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for valor in range(len(valorCap11_4)):
        cell = tabla11b.cell(valor + 1, 1)
        t11b = cell.paragraphs[0].add_run(f'{valorCap11_4[valor]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 11.4 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Importancia: está relacionada con la intensidad o el grado de alteración de la acción impactante sobre el componente ambiental. La escala va de 1 a 10, siendo asignado su valor con base al juicio de la persona encargada de hacer la evaluación')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 11.1 ###
    #########################
    columnasCap11_4 = [
        'Importancia',
        'Valor',
    ]

    importanciaCap11_4 = [
        'Sin Importancia',
        'Poco Importante',
        'Medianamente Importante',
        'Importante',
        'Muy Importante',
    ]

    valorCap11_4 = [
        1,
        3,
        5,
        7,
        10,
    ]

    columnas = len(columnasCap11_4)
    filas = len(importanciaCap11_4) + 1
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')
        t11b = cell.paragraphs[0].add_run(f'{columnasCap11_4[cols]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(10)
        t11b.font.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for importancia in range(len(importanciaCap11_4)):
        cell = tabla11b.cell(importancia + 1, 0)
        t11b = cell.paragraphs[0].add_run(f'{importanciaCap11_4[importancia]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for valor in range(len(valorCap11_4)):
        cell = tabla11b.cell(valor + 1, 1)
        t11b = cell.paragraphs[0].add_run(f'{valorCap11_4[valor]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 11.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.1.- Obtención del valor de importancia.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('El valor de la importancia es un parámetro que mide cualitativamente el impacto ambiental, su medición se realiza en función del grado de incidencia e intensidad de la alteración como resultado de una acción, así como de las características del efecto, que responden a una serie de atributos tipo cualitativo que son:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    listaCap11_41 = [
        'Naturaleza',
        'Intensidad',
        'Extensión',
        'Momento',
        'Reversibilidad',
        'Sinergia',
        'Acumulación',
        'Efecto',
        'Periodicidad',
        'Recuperabilidad',
    ]

    for lista in range(len(listaCap11_41)):
        di11 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo11 = di11.add_run(f'{listaCap11_41[lista]}')
        descripcionCapitulo11_format = di11.paragraph_format
        descripcionCapitulo11_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo11.font.name = 'Arial'
        descripcionCapitulo11.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Cada impacto identificado se caracterizó en función de los atributos anteriores, cada uno con su propia escala ordinal.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.2.- Criterios para el cálculo de la importancia.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Con base en estos criterios, de acuerdo con los rangos que se muestran en la tabla adjunta, se obtiene la importancia (I) de las consecuencias ambientales del impacto aplicando el siguiente algoritmo: ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('I = (3IN+2EX+MO+PE+RV+SI+AC+EF+PR+MC) ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Times New Roman'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    descripcionCapitulo11.italic = True
    di11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 11.4.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.2.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.2.1.- Los criterios de evaluación. ')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.2.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Los criterios utilizados por el método Conessa para la evaluación de los impactos ambientales son los siguientes: ')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.3.- Criterios de evaluación.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.2.1 ###
    #########################
    columnasCap11_421 = [
        'CRITERIOS',
        'SIGNIFICADO',
    ]

    criteriosCap11_421 = [
        'Signo',
        'Intensidad',
        'Extensión',
        'Momento',
        'Persistencia',
        'Reversibilidad',
        'Recuperabilidad',
        'Sinergia',
        'Acumulación',
        'Efecto',
        'Periodicidad',
    ]

    signosCap11_421 = [
        '+/-'
        'IN',
        'EX',
        'MO',
        'PE',
        'RV',
        'MC',
        'SI',
        'AC',
        'EF',
        'PR',
    ]

    significadoCap11_421 = [
        'Hace alusión al carácter benéfico (+) o perjudicial (-) de las distintas acciones que van a actuar sobre los distintos factores considerados',
        'Grado de incidencia de la acción sobre el factor en el ámbito específico en el que actúa.  Varía entre 1 y 12, siendo 12 la expresión de la destrucción total del factor en el área en la que se produce el efecto y 1 una mínimo afectación.',
        'Área de influencia teórica del impacto en relación con el entorno de la actividad (% de área, respecto al entorno, en que se manifiesta el efecto). Si la acción produce un efecto muy localizado, se considera que el impacto tiene un carácter puntual (1).  Si, por el contrario, el impacto no admite una ubicación precisa del entorno de la actividad, teniendo una influencia generalizada en todo él, el impacto será Total (8). Cuando el efecto se produce en un lugar crítico, se le atribuirá un valor de cuatro unidades por encima del que le correspondía en función del % de extensión en que se manifiesta.',
        'Alude al tiempo entre la aparición de la acción que produce el impacto y el comienzo de las afectaciones sobre el factor considerado. Si el tiempo transcurrido es nulo, el momento será Inmediato, y si es inferior a un año, Corto plazo, asignándole en ambos casos un valor de cuatro (4).  Si es un período de tiempo mayor a cinco años, Largo Plazo (1).',
        'Tiempo que supuestamente permanecerá el efecto desde su aparición y, a partir del cual el factor afectado retornaría a las condiciones iniciales previas a la acción por los medios naturales o mediante la introducción de medidas correctoras.',
        'Se refiere a la posibilidad de reconstrucción del factor afectado como consecuencia de la acción acometida, es decir, la posibilidad de retornar a las condiciones iniciales previas a la acción, por medios naturales, una vez aquella deje de actuar sobre el medio.',
        'Se refiere a la posibilidad de reconstrucción, total o parcial, del factor afectado como consecuencia de la actividad acometida, es decir, la posibilidad de retornar a las condiciones iniciales previas a la acción, por medio de la intervención humana (medidas de manejo ambiental). Cuando el efecto es irrecuperable (alteración imposible de reparar, tanto por la acción natural, como por la humana) le asignamos el valor de ocho (8).  En caso de ser irrecuperable, pero existe la posibilidad de introducir medidas compensatorias, el valor adoptado será cuatro (4).',
        'Este atributo contempla el reforzamiento de dos o más efectos simples.  La componente total de la manifestación de los efectos simples, provocados por acciones que actúan simultáneamente, es superior a la que cabría de esperar cuando las acciones que las provocan actúan de manera independiente, no simultánea.',
        'Este atributo da idea del incremento progresivo de la manifestación del efecto cuando persiste de forma continuada o reiterada la acción que lo genera.  Cuando un acción no produce efectos acumulativos (acumulación simple), el efecto se valora como uno (1); si el efecto producido es acumulativo el valor se incrementa a cuatro (4).',
        'Este atributo se refiere a la relación causa-efecto, o sea, a la forma de manifestación del efecto sobre un factor, como consecuencia de una acción.  Puede ser directo o primario, siendo en este caso la repercusión de la acción consecuencia directa de ésta, o indirecto o secundario, cuando la manifestación no es consecuencia directa de la acción, sino que tiene lugar a partir de un efecto primario, actuando este como una acción de segundo orden.',
        'Se refiere a la regularidad de manifestación del efecto, bien sea de manera cíclica o recurrente (efecto periódico), de forma impredecible en el tiempo (efecto irregular) o constante en el tiempo (efecto continuo)',

    ]

    columnas = len(columnasCap11_421) + 1
    filas = len(criteriosCap11_421) + 1
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for rows in tabla11b.rows:
        rows.cells[0].width = Cm(3.74)
        rows.cells[1].width = Cm(1.25)
        rows.cells[2].width = Cm(14)

    row = tabla11b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('CRITERIOS')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    t11b.bold = True
    cell_background_color(merged_cell, '#3498db')
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla11b.cell(0, 2)
    cell_background_color(cell, '#3498db')
    t11b = cell.paragraphs[0].add_run(f'SIGNIFICADO')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    t11b.font.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for importancia in range(len(criteriosCap11_421)):
        cell = tabla11b.cell(importancia + 1, 0)
        t11b = cell.paragraphs[0].add_run(f'{criteriosCap11_421[importancia]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for valor in range(len(signosCap11_421)):
        cell = tabla11b.cell(valor + 1, 1)
        t11b = cell.paragraphs[0].add_run(f'{signosCap11_421[valor]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for valor in range(len(significadoCap11_421)):
        cell = tabla11b.cell(valor + 1, 2)
        t11b = cell.paragraphs[0].add_run(f'{significadoCap11_421[valor]}')
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.2.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En Base a la evaluación mediante la aplicación del algoritmo se determina que:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    listaCap11_421 = [
        'En el rango 17 a 25 son considerados irrelevantes o no significativos los Impactos',
        'Desde el 26 hasta 50 son impactos Moderados.',
        'Del rango de 51 al 75 los impactos son considerados severos',
        'Superiores a 76 los impactos son considerados como Críticos.',
    ]

    for lista in range(len(listaCap11_421)):
        di11 = doc.add_paragraph(style='List Bullet')
        descripcionCapitulo11 = di11.add_run(f'{listaCap11_421[lista]}')
        descripcionCapitulo11_format = di11.paragraph_format
        descripcionCapitulo11_format.line_spacing = 1.15
        #descripcionCapitulo9_format.space_after = 0
        #descripcionCapitulo9_format.space_before = 0

        descripcionCapitulo11.font.name = 'Arial'
        descripcionCapitulo11.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.4.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.3.- Rangos para el cálculo de importancia para cada atributo.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.4.- Rangos de valoración de la importancia')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.2.1 ###
    #########################
    columnas = 4
    filas = 12
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for rows in tabla11b.rows:
        rows.cells[0].width = Cm(8.16)
        rows.cells[1].width = Cm(1.14)
        rows.cells[2].width = Cm(8.16)
        rows.cells[3].width = Cm(1.14)

    """
        ****************************
        * Naturaleza *
        ****************************
    """
    row = tabla11b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('NATURALEZA')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(1, 0)
    t11b = cell.paragraphs[0].add_run(
        'Impacto benéfico'
        '\nImpacto perjudicial'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '##85c1e9')

    cell = tabla11b.cell(1, 1)
    t11b = cell.paragraphs[0].add_run(
        '+'
        '\n-'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Intensidad *
        ****************************
    """
    row = tabla11b.rows[0]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('INTENSIDAD'
                                             '\n(Grado de Destrucción)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(1, 2)
    t11b = cell.paragraphs[0].add_run(
        'Baja'
        '\nMedia'
        '\nAlta'
        '\nMuy alta'
        '\nTotal'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(1, 3)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        '\n8'
        '\n12'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Extension *
        ****************************
    """
    row = tabla11b.rows[2]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('EXTENSIÓN (EX)'
                                             '\n(Área de influencia)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(3, 0)
    t11b = cell.paragraphs[0].add_run(
        'Puntual'
        '\nParcial'
        '\nExtensa'
        '\nTotal'
        '\nCrítica'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(3, 1)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        '\n8'
        '\n(+4)'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Momento *
        ****************************
    """
    row = tabla11b.rows[2]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('MOMENTO (MO)'
                                             '\n(Plazo de manifestación)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(3, 2)
    t11b = cell.paragraphs[0].add_run(
        'Largo plazo'
        '\nMedio Plazo'
        '\nInmediato'
        '\nCrítico'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(3, 3)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        '\n(+4)'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Persistencia *
        ****************************
    """
    row = tabla11b.rows[4]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('PERSISTENCIA (PE)'
                                             '\n(Permanencia del efecto)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(5, 0)
    t11b = cell.paragraphs[0].add_run(
        'Fugaz'
        '\nTemporal'
        '\nPermanente'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(5, 1)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Reversibilidad *
        ****************************
    """
    row = tabla11b.rows[4]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('REVERSIBILIDAD (RV)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(5, 2)
    t11b = cell.paragraphs[0].add_run(
        'Corto plazo'
        '\nMedio plazo'
        '\nIrreversible'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(5, 3)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Sinergia *
        ****************************
    """
    row = tabla11b.rows[6]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('EXTENSIÓN (EX)'
                                             '\n(Área de influencia)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(7, 0)
    t11b = cell.paragraphs[0].add_run(
        'Sin sinergismo (simple)'
        'Sinérgico '
        'Muy sinérgico '
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(7, 1)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Acumulacion *
        ****************************
    """
    row = tabla11b.rows[6]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('ACUMULACIÓN (AC)'
                                             '\n(Incremento progresivo)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(7, 2)
    t11b = cell.paragraphs[0].add_run(
        'Simple'
        '\nAcumulativo'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(7, 3)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Efecto *
        ****************************
    """
    row = tabla11b.rows[8]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('EEFECTO (EF)'
                                             '\n(Relación causa-efecto)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(9, 0)
    t11b = cell.paragraphs[0].add_run(
        'Indirecto (secundario)'
        '\nDirecto'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(9, 1)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Periodicidad *
        ****************************
    """
    row = tabla11b.rows[8]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('PERIODICIDAD (PR)'
                                             '\n(regularidad de la manifestación)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(9, 2)
    t11b = cell.paragraphs[0].add_run(
        'Irregular o aperiódico o descontinuo'
        '\nPeriódico'
        '\nContinuo'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(9, 3)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Recuperabilidad *
        ****************************
    """
    row = tabla11b.rows[10]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[1]))
    t11b = merged_cell.paragraphs[0].add_run('RECUPERABILIDAD (MC)'
                                             '\n(Reconstrucción por medios humanos)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(11, 0)
    t11b = cell.paragraphs[0].add_run(
        'Recuperable inmediato'
        '\nRecuperable a medio plazo'
        '\nMitigable o compensable'
        '\nIrrecuperable'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(11, 1)
    t11b = cell.paragraphs[0].add_run(
        '1'
        '\n2'
        '\n4'
        '\n8'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    cell_background_color(cell, '#85c1e9')

    """
        ****************************
        * Importancia *
        ****************************
    """
    row = tabla11b.rows[10]
    merged_cell = row.cells[2].merge(row.cells[2].merge(row.cells[3]))
    t11b = merged_cell.paragraphs[0].add_run('IMPORTANCIA (I)')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    cell = tabla11b.cell(11, 2)
    t11b = cell.paragraphs[0].add_run(
        'I=(3IN+2EX+MO+PE+RV+SI+AC+EF+PR+MC)'
        )
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(12)
    t11b.italic = True
    t11b.bold = True
    cell_background_color(cell, '#85c1e9')

    cell = tabla11b.cell(11, 3)
    cell_background_color(cell, '#85c1e9')

    ########################################################################################################################################################################
    # Capitulo 11.4.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.4.- Desarrollo de la Técnica.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.4 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En cada casilla de cruce en la matriz o elemento tipo, otorga la idea del efecto de cada acción impactada sobre cada factor impactado al determinar la importancia del impacto de cada elemento considerado en la siguiente expresión.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('I = (3IN+2EX+MO+PE+RV+SI+AC+EF+PR+MC)')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Times New Roman'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.italic = True
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Los valores de importancia que se obtienen con el modelo propuesto, tomaran valores positivos o negativos, entre 13 y 100')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('De acuerdo a las condiciones que afectaran al medio ambiente en cada una de sus etapas y por cada acción se contempla la siguiente afectación de acuerdo a los siguientes valores.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Finalmente, en esta fase del estudio, se estandarizó la importancia del impacto a cada uno de los valores, dividiendo todos los valores de importancia, entre el máximo valor de importancia que es posible obtener para todos los impactos ambientales identificados, utilizándose para tal fin, se utiliza la siguiente fórmula:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Iij = Iij / 100')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Times New Roman'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.italic = True
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Siendo'
                                         '\nlij = Representa el impacto de acción sobre cada indicador de impacto.'
                                         '\nI = Indicador de impacto.'
                                         '\nJ = Acción del proyecto.'
                                         '\n88 = Número máximo de valoración cuando el atributo se manifiesta al máximo.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Los valores de importancia que se obtienen con el modelo anterior, pueden tomar valores diferentes para este proyecto en sus diferentes etapas de acuerdo al tipo de las acciones y factores sobre los que se tiene efecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En este proyecto en cada una de sus etapas altera el medio ambiente en diferentes dimensiones de acuerdo al siguiente esquema:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.5.- Categorización de los impactos')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    columnasCap11_44 = [
        'Valor de la Importancia',
        'Significancia',
        'Magnitud',
    ]

    valorImportanciaCap11_44 = [
        'Rango de 13 hasta 25',
        'Rango de 26 hasta 50',
        'Rango de 51 hasta 75',
        'Mayor de 75',
    ]

    significanciaCap11_44 = [
        'No significativo',
        'Moderado',
        'Severo',
        'Criticos',
    ]

    magnitudCap11_44 = [
        'Baja y Muy Baja Magnitud',
        'Mediana Magnitud',
        'Alta Magnitud',
        'Muy Alta Magnitud',
    ]

    filas = len(valorImportanciaCap11_44) + 2
    columnas = len(columnasCap11_44)
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    row = tabla11b.rows[0]
    merged_cell = row.cells[0].merge(row.cells[0].merge(row.cells[2]))
    t11b = merged_cell.paragraphs[0].add_run('Categorización de los Impactos')
    t11b.font.name = 'Arial'
    t11b.font.size = Pt(11)
    t11b.bold = True
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell, '#3498db')

    for cols in range(columnas):
        cell = tabla11b.cell(1, cols)
        t11b = cell.paragraphs[0].add_run(columnasCap11_44[cols])
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        t11b.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(cell, '#3498db')

    for rows in range(len(valorImportanciaCap11_44)):
        cell = tabla11b.cell(rows + 2, 0)
        t11b = cell.paragraphs[0].add_run(valorImportanciaCap11_44[rows])
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for rows in range(len(significanciaCap11_44)):
        cell = tabla11b.cell(rows + 2, 1)
        t11b = cell.paragraphs[0].add_run(significanciaCap11_44[rows])
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for rows in range(len(magnitudCap11_44)):
        cell = tabla11b.cell(rows + 2, 2)
        t11b = cell.paragraphs[0].add_run(magnitudCap11_44[rows])
        t11b.font.name = 'Arial'
        t11b.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    ########################################################################################################################################################################
    # Capitulo 11.4.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.5 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.5- Resultados Obtenidos en las matrices. ')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.4.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.5.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.5.1.- Resultados de la identificación de impactos.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.4.5.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('De acuerdo a la valoración de los impactos se obtuvieron la matriz de Conessa los siguientes resultados.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.6.-	Cantidad de conceptos obtenidos.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 7
    columnas = 2
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 11.4.5.1 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nImpactos Identificados por etapas del proyecto')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.7.- Identificación de posibles impactos a generar por etapas')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 7
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ########################################################################################################################################################################
    # Capitulo 11.4.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.4.5.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.4.5.2.- Resultado de la valoración de impactos según la etapa del proyecto')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.8.-	Clasificación de los impactos.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 7
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capitulo 11.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los resultados obtenidos son los siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.9.- Matriz de identificación de impactos')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 10
    columnas = 15
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el contenido del capítulo 11.4.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical:
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.10.- Valoración de impactos etapa de preparación de sitio')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 40
    columnas = 15
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.11.- Matriz de identificación de impactos etapa de construcción')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 40
    columnas = 15
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.12.- Valoración de impactos, etapa de operación')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 40
    columnas = 15
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Título de la tabla del capítulo 11.4.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.13.- Valoración de impactos etapa de abandono')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.4.4 ###
    #########################
    filas = 40
    columnas = 15
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ########################################################################################################################################################################
    # Capitulo 11.5
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.-Caracterización de los impactos negativos por Etapa del Proyecto.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.5.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.1.- Preparación del Sitio.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.5.1.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.1.1.- PS-01 Delimitación del área.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.1.1 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.14.- Impactos generados en la delimitación del área, etapa preparación de sitio')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.1.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.1.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.1.2.- PS-02 Rescate de Flora y Fauna.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.1.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.15.- Impactos generados en el rescate de flora y fauna, etapa preparación de sitio')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.1.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.1.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.1.3.- PS-03 Despalme (Remoción de la capa superficial del suelo). ')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.1.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.16.- Impactos generados en el Desmonte, preparación del sitio.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.1.3 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.1.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.1.4.- PS-03 Despalme (Remoción de la capa superficial del suelo). ')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.1.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.17.- Impactos generados en el Despalme, preparación del sitio.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.1.4 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.- Etapa de Construcción.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.5.2.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.2.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.1.- CO-01.- __________________')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.2.1 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.18.- Impactos generados en ________________, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.2.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.2.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.2.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.2.- CO-02.- _________________.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.2.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.19.- Impactos generados en rampas de acceso, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.2.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.2.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.2.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.3.- CO-03 _____________________________________.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.2.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.20.- Impactos generados en los ____________________, etapa de Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.2.3 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.2.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.2.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.4.- CO-04 __________________')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.2.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.21.- Impactos generados en el _____________________, etapa de Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.2.4 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.3.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.3.- Etapa de Operación.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.5.3.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.3.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.3.1.- OP-01.- __________________')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.3.1 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.18.- Impactos generados en ________________, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.3.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.3.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.3.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.2.2.- CO-02.- _________________.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.3.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.19.- Impactos generados en rampas de acceso, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.3.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.3.3
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.3.3 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.3.3.- OP-03.- _____________________________________.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.3.3 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.20.- Impactos generados en los ____________________, etapa de Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.3.3 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.3.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.3.4 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.3.4.- OP-04.- __________________')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.3.4 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.21.- Impactos generados en el _____________________, etapa de Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.3.4 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.4.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.4.- Abandono del sitio.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 11.5.4.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.4.1 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.4.1.- AB.01.- __________________')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.4.1 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.18.- Impactos generados en ________________, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.4.1 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.5.4.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.5.4.2 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nXI.5.4.2.- AB.02.- _________________.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.4.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.19.- Impactos generados en rampas de acceso, Construcción.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.4.2 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    """
        ############################
        Modelos de dispersion atmosferica
    """
    #########################
    ### Descripcion del capitulo 11.5.4.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Modelos de dispersión atmosférica.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ============================
        Metodo de Pasquill
        ============================
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Método de Pasquill')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Pasquill propuso un sistema para clasificar la dispersión turbulenta. Definió entonces Categorías de Estabilidad, de la A la F con la siguiente definición:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 11.5.4.2 ###
    #########################
    columnasCap11_542 = [
        'Categoria de Estabilidad',
        'Definición',
    ]

    col1Cap11_542 = [
        'A',
        'B',
        'C',
        'D',
        'E',
        'F',
    ]

    col2Cap11_542 = [
        'Extremadamente Inestable',
        'Moderadamente Inestable',
        'Ligeramente Inestable',
        'Neutra',
        'Ligeramente Estable',
        'Moderadamente Estable',
    ]

    filas = len(col1Cap11_542) + 1
    columnas = len(columnasCap11_542)
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        t11b = cell.paragraphs[0].add_run(columnasCap11_542[cols])
        t11b.font.size = Pt(12)
        t11b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t11b.font.bold = True
        cell_background_color(cell, '#3498db')

    for col1 in range(len(col1Cap11_542)):
        cell = tabla11b.cell(col1 + 1, 0)
        t11b = cell.paragraphs[0].add_run(f'{col1Cap11_542[col1]}')
        t11b.font.size = Pt(12)
        t11b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for col2 in range(len(col2Cap11_542)):
        cell = tabla11b.cell(col2 + 1, 1)
        t11b = cell.paragraphs[0].add_run(f'{col2Cap11_542[col2]}')
        t11b.font.size = Pt(12)
        t11b.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 11.5.4.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Dichas categorías están asociadas a las condiciones dispersivas de la atmósfera, función de la turbulencia, y no al estado del tiempo (lluvias, viento, tormentas, etc.) Como se mencionó, la turbulencia depende esencialmente de la radiación que recibe la superficie del suelo y del viento. Cuando no hay nubosidad la primera depende de la altura del sol sobre el horizonte. La nubosidad disminuye la radiación durante el día, pero la aumenta durante la noche por reemisión. Por lo tanto, la cobertura nubosa modifica la cantidad de radiación incidente. Es así que cada una de estas categorías está asociada a condiciones de velocidad del viento y radiación imperante para lo cual se postuló la siguiente tabla:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.5.4.2 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.x.- Radiación solar')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.5.4.2 ###
    #########################
    colsCap11_542 = [
        'm/s',
        'Fuerte',
        'Moderado',
        'Debil',
        'Cubierto',
        'Liger Cub',
        'Despejado',
    ]

    filas_datos = [
        ["-2",  "A",   "A-B", "B", "D", "E", "F"],
        ["2-3", "A-B", "B",   "C", "D", "E", "E"],
        ["3-5", "B",   "B-C", "C", "D", "D", "E"],
        ["5-6", "C",   "C-D", "D", "D", "D", "D"],
        ["+ 6", "C",   "D",   "D", "D", "D", "D"]
    ]


    filas = len(filas_datos) + 2
    columnas = len(colsCap11_542)
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        t11b = cell.paragraphs[0].add_run(colsCap11_542[cols])
        t11b.font.size = Pt(12)
        t11b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t11b.font.bold = True
        cell_background_color(cell, '#3498db')

    for fila_idx, fila_datos in enumerate(filas_datos, start=2):
        for col_idx, dato in enumerate(fila_datos):
            cell = tabla11b.cell(fila_idx, col_idx)
            run = cell.paragraphs[0].add_run(dato)
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 11.5.4.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Para la radiación solar se utiliza una clasificación simplificada que tiene en cuenta el ángulo de elevación del sol desde el horizonte, como sigue:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 11.5.4.2 ###
    #########################
    colsCap11_542 = [
        "ALTURA DEL SOL H",
        "RADIACIÓN SOLAR"
    ]

    filas_datos = [
        ["H - 60°", "FUERTE"],
        ["H 35-40°", "MODERADA"],
        ["H-35", "DEBIL"]
    ]

    filas = len(filas_datos) + 2
    columnas = len(colsCap11_542)
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        t11b = cell.paragraphs[0].add_run(colsCap11_542[cols])
        t11b.font.size = Pt(12)
        t11b.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        t11b.font.bold = True
        cell_background_color(cell, '#3498db')

    for fila_idx, fila_dato in enumerate(filas_datos, start=2):
        for col_idx, dato in enumerate(fila_dato):
            cell = tabla11b.cell(fila_idx, col_idx)
            t11b = cell.paragraphs[0].add_run(dato)
            t11b.font.size = Pt(12)
            t11b.font.name = 'Arial'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capítulo 11.5.4.2 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Existen modificaciones para esta tabla que contemplan los distintos grados de cobertura nubosa durante el día. En el caso de cielos completamente cubiertos, se considera una categoría Pasquill D."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Este esquema presentado es sencillo y de fácil aplicación. En la actualidad hay una gran variedad de métodos alternativos para la categorización de la estabilidad atmosférica. La mayoría usa la misma definición que las categorías de Pasquill de la A, a la F."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Todos se basan en indicadores de turbulencia y poseen distintos grados de resolución y precisión. Si bien existen diferencias menores, cualquier sistema de clasificación es válido y permite utilizar los correspondientes parámetros de difusión turbulenta que se aplicarán en las fórmulas de cálculo."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Tomando en cuenta este método se puede decir que en el área en estudio se presenta  este tipo de clasificaciones durante el día y en diferentes épocas del año, durante las observaciones efectuadas la velocidad máxima de los vientos correspondía a la medida de 0 .5-1m/s con radiación de fuerte a moderada durante el periodo de toma de datos donde la condición ambiental se encontraba desde moderadamente inestable hasta ligeramente inestable y en esta época el cielo durante la noche permanece ligeramente cubierto con condición atmosférica de neutra a ligeramente estable; la radiación está determinada por el ángulo del suelo con respecto a la posición del sol considerando la exposición del predio  90° por lo tanto la radiación es fuerte en forma ordinaria tal y como se demuestra en el tipo de vegetación que hay en el área. Este tipo de condiciones van a estar presentes en el área de estudio, por lo que el proceso que se plantea llevar a cabo afectará en algunos aspectos en forma significativa y en la mayoría será no significativa, esta irradiación en el predio se presenta entre los 4.8 y los 5.1 kW h/m2, como se puede apreciar en la siguiente figura."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Mapa del capitulo 11.5.4.2 ###
    #########################
    """ 
        El siguiente codigo muestra como se tiene que insertar la imagen o mapa.
    """
    imagenCapitulo11 = doc.add_paragraph()
    imagenCapitulo11.text = '\n'
    imagenCapitulo11 = doc.add_picture('capitulo11/mapa.png')  # Ancho de la imagen en centimetros
    imagenCapitulo11.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro de la imagen
    imagenCapitulo11.width = Cm(17.09)  # Ancho de la imagen en centimetros
    imagenCapitulo11.height = Cm(11.06)  # Alto de la imagen en centimetros
    imagenCapitulo11.space_after = Pt(0)  # Espacio después de la imagen

    #########################
    ### Descripcion del Mapa del capitulo 3.2 ###
    #########################
    diMap11 = doc.add_paragraph()
    descripcionCapituloMapa11 = diMap11.add_run('Fuente: Instituto de Investigaciones Eléctricas')
    descripcionCapituloMapa11_format = diMap11.paragraph_format
    descripcionCapituloMapa11_format.line_spacing = 1.15
    descripcionCapituloMapa11.space_before = Pt(0)  # Espacio antes del texto

    descripcionCapituloMapa11.font.name = 'Arial'
    descripcionCapituloMapa11.font.size = Pt(12)
    descripcionCapituloMapa11.font.italic = True
    diMap11.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro del texto

    #########################
    ### Descripcion del capítulo 11.5.4.2 ###
    #########################
    """
        ############################
        Modelos de vulnerabilidad a la contaminación de aguas de mantos freáticos.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos de vulnerabilidad a la contaminación de aguas de mantos freáticos.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "La vulnerabilidad es un concepto que se refiere a la potencial pérdida de calidad del agua subterránea debido al grado de exposición natural de los acuíferos. Los riesgos son la probabilidad de ocurrencia de algo nocivo o dañino que depende tanto de la intensidad de la amenaza como de los niveles de vulnerabilidad del acuífero, el riesgo es el grado de pérdida de calidad y cantidad de agua en el subsuelo debido a las amenazas de contaminación a la modificación de los flujos, a la sobre extracción y el cambio climático entre otros de acuerdo al nivel de vulnerabilidad."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Los suelos poseen una capacidad natural para amortiguar el paso de los contaminantes y depurar el agua, por esta razón se debe considerar este elemento del ambiente en el análisis de vulnerabilidad, la cual no se pone en riesgo ya que no se encuentran cuerpos de agua permanentes que se puedan contaminar."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Si se considera a la vulnerabilidad como la potencial perdida de la calidad de agua subterránea debido al grado de exposición natural se deberá de analizar la función ambiental que cumple cada elemento del medio físico como son:"
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Relieve: Factor que regula los flujos del agua superficial, los disipa y los concentra."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    di11.style = 'List Bullet'

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Suelo y subsuelo: Al ser las capas protectoras y funcionan como filtros naturales debido a los procesos de retención y descomposición."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    di11.style = 'List Bullet'

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Clima: Es la vía de transporte de posibles contaminantes."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    di11.style = 'List Bullet'

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Hidrología: Con referencia a los flujos de agua superficial o subterránea para determinar el posible destino de los contaminantes."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    di11.style = 'List Bullet'

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En relación a lo anterior en el predio sujeto de estudio y considerando las características del sistema ambiental en la cual está inmerso el área en estudio, el relieve del área sujeta de estudio presenta pendientes de diversos porcentajes de oscilación, por lo cual los grados de inclinación van de __, ya que se encuentra en topoformas como __ en la mayor parte de la superficie y __, en donde presenta escurrimientos intermitentes, los cuales obtienen agua de la sierra donde este se encuentra el flujo de agua precipitada en temporada va a dar a escorrentías que se encuentra en áreas aguas abajo en donde no será sujeto a recibir aguas contaminadas."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En el área en estudio, predomina el tipo de clima __: Este grupo pertenece a los climas ______________________________________________."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Al estar el área en estudio dentro del clima mencionado descrito donde se tiene registro de Las precipitaciones totales anuales para la estación meteorológica van de _ a los _ mm, de acuerdo a los registros históricos."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Las precipitaciones observadas han sido de manera escasas y erráticas a lo largo de los años tal es el caso de la distribución de la precipitación registrada por la estación meteorológica, la cual pertenece a la Red Meteorológica de CONAGUA ubicada en el municipio de _ muestra que la precipitación anual acumulada es de _ mm anuales, durante un periodo de _ años."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Con respecto a la Hidrología en el área del proyecto, los períodos de precipitación se encuentran en los meses de _ a _, de manera oficial, sin embargo, los mese con mayor precipitación fueron _ y _, en cuyos meses las precipitaciones son altas superando los _ mm."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos de difusión y dispersión en causes.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos de difusión y dispersión en causes.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En la modelación del transporte de sustancias es importante simular los procesos de mezcla debidos a la dispersión."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En el análisis comparativo de algunos métodos permitió determinar que no se puede obtener información que dé certeza sobre el tipo y grado de sustancias que se conducen en el cauce al no encontrarse corrientes perennes que arrojen resultados para evaluar la posible contaminación."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Las diferentes metodologías que se han aplicado para determinar la difusión y dispersión en cauces se han efectuado en corrientes perennes para determinar mediante logaritmos numéricos el grado de difusión y dispersión de sustancias que ocurren en un cauce, para el caso del área en estudio donde se tienen corrientes intermitentes la aplicación de estos modelos sería errónea por su inconsistencia de lecturas así mismo al ser aplicado en aquellos lugares donde las corrientes o los causes reciben corrientes artificiales que integrarían contaminantes, dado que en el área sujeto de estudio no se tiene esta condición no se pude aplicar ni analizar por no estar en sitio de agua estacionaria y al no existir actividad alguna en los terrenos del predio, el flujo natural solo contendrá microorganismos propios del ecosistema."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos para determinar la capacidad de autodepuración.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos para determinar la capacidad de autodepuración.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Los ríos (y en general cualquier sistema natural) se analizan y estudian como reactores biogeoquímicos en los que, al igual que una estación depuradora de aguas residuales o una estación de tratamiento agua potable, el agua entra con una composición o calidad determinada y sale con otra composición distinta."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Los cambios en la composición del agua que se producen en los sistemas naturales tienen lugar en virtud de una serie de procesos físicos, químicos y biológicos que son los mismos en sistemas naturales o diseñados por el hombre."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "El modelo de Streeter y Phelps ignoran otras fuentes y sumideros de oxígeno que existen en ríos y sistemas naturales en general, los cuales incluyen:"
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista_fuentes_oxigeno = [
        "Producción de oxígeno durante la fotosíntesis de algas y otras plantas acuáticas",
        "Demanda de oxígeno desde el sedimento",
        "Consumo de oxígeno por respiración de plantas acuáticas.",
        "Oxidación de compuestos nitrogenados en el agua residual."
    ]

    # For con numeración automática en el add_run
    for idx, item in enumerate(lista_fuentes_oxigeno, start=1):
        di11 = doc.add_paragraph()
        descripcionCapitulo11 = di11.add_run(f"{idx}.- {item}")
        descripcionCapitulo11_format = di11.paragraph_format
        descripcionCapitulo11_format.line_spacing = 1.15
        descripcionCapitulo11.font.name = 'Arial'
        descripcionCapitulo11.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("A la capacidad de los sistemas naturales para modificar la composición del agua que reciben, y eliminar contaminantes (materia orgánica, sustancias en suspensión, etc.) se conoce como capacidad de autodepuración.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Considerando que en el área es estudio no se cuentan con cauces de ríos, arroyos o cuerpos de agua que presentan esta actividad no se registran estas actividades que puedan someterse al manejo a efecto de que se depure al agua de sustancias contaminantes mediante la oxidación de la materia orgánica y sustancias en suspensión, el manejo de acuerdo a las normas establecidas se realizara a las aguas solamente de tipo residuales con empresas certificadas para que les dé mantenimientos a las aguas provenientes de los baños portátiles, para someter a los tratamientos para su reciclamiento y uso en otros procesos.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos para evaluar el riesgo de eutrofización
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos para evaluar el riesgo de eutrofización')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Las corrientes que son afectados por captaciones, presentan tramos en donde las condiciones físicas y biológicas del ecosistema cambian en relación a las condiciones naturales, estos tramos son el objeto de la implementación de los regímenes de caudales ecológicos. Las variables físicas o hidráulicas y las variables biológicas son el insumo eco hidráulico de los modelos de caudales ecológicos que se pueden integrar a la hidrología local.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Fundamento Eco hidrológico
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nFundamento Eco hidrológico')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Los regímenes de caudales históricos han determinado el tipo de hábitats que se pueden encontrar en el flujo de una escorrentía y con ellos los organismos que se han adaptado a estas condiciones. El comportamiento de la escorrentía lo caracteriza su régimen hidrológico, es por ello que al evaluar los hábitats viables es necesario evaluar la estacionalidad y la temporalidad del régimen hidrológico. La ecología de los organismos no es estática por lo que el hábitat está sujeto al flujo de energía del caudal.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Este fundamento es el que nos obliga a entender que el concepto de caudal ecológico se traduce en un régimen de caudales ecológicos similar al régimen hidrológico natural.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("El fundamento eco hidrológico de los caudales ecológicos es la variabilidad natural de la escorrentía no afectada, la misma que se busca mantener con el régimen de caudales ecológicos. Estimar esta respuesta se puede lograr con la temporalidad de la información hidrológica disponible e información ecológica correspondiente a las estaciones.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Con estos antecedentes, el concepto de caudal ecológico se sustenta en la combinación de los criterios ecológicos, hidráulicos e hidrológicos para reconocer los hábitats viables y mantener los regímenes de caudales óptimos para las comunidades clave del ecosistema acuático.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Para el caso del área en estudio no cuenta dentro de sus áreas cuerpos de agua o corrientes intermitentes que conducen agua en la temporada de lluvias, se puede decir que de acuerdo al tipo de vegetación existente esta se ha adaptado al régimen pluviométrico y es característica de la zona el tipo de asociación vegetativa y su deterioro lo determina el régimen hidrológico ya que al presentarse sequias prolongadas se pierden especímenes y aparecen otras con menor requerimiento de flujos de humedad para su sobrevivencia que pueden alterar el entorno ecológico.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos de evacuación del suelo y de sus diferentes funciones.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos de evacuación del suelo y de sus diferentes funciones.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Sistema individual para el tratamiento de aguas residuales producidas por familias que habitan en zonas residenciales poco pobladas.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Este sistema puede recibir tanto el agua con los excrementos humanos como aquella proveniente de cocinas y baños.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("El material sedimentado (los sólidos) forma en el fondo del depósito una capa de lodos o fango, degradado biológicamente con el tiempo y que debe extraerse periódicamente.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("En este aspecto  en el área en estudio, al contemplar las letrinas secas se podrá tener un control sobre la evacuación con el fin de mantener el suelo libre de contaminantes que a la vez pudiesen percollarse hacia los mantos freáticos, aun cuando el suelo por su tipo de estructura es de alto grado de infiltración actúa como filtro eliminando aquellos que alteran tanto la estructura como las corrientes naturales, ante ello considerando que en el plazo de la operación del proyecto se deberá de establecer y mantener un sistema de control para evitar infiltraciones al subsuelo de aguas residuales al efectuar un cambio estructural del suelo y contaminación de mantos freáticos por la operación propia del proyecto.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos de evacuación del suelo y de sus diferentes funciones.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nModelos de evacuación del suelo y de sus diferentes funciones.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("La vulnerabilidad a la contaminación es una característica de los acuíferos difícil de determinar y depende de la interacción entre diferentes factores, como profundidad del nivel freático o techo del acuífero (mismo que en la zona es un aproximado entre 150 m), la capacidad de atenuación de las capas litológicas sobrepuestas al acuífero, la tasa de recarga y otros factores.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("La vulnerabilidad puede entenderse como la sensibilidad en la calidad del agua subterránea ante una carga contaminante impuesta, la cual es determinada por las características intrínsecas del acuífero. Por lo tanto, la vulnerabilidad es inversa a la capacidad de atenuación de contaminantes del acuífero.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("En el ámbito de las aguas subterráneas el riesgo de contaminación está formado por la interacción de dos partes:")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lista_vulnerabilidad = [
        "La pasiva, representada por la vulnerabilidad, que no depende de la actividad humana y no cambia perceptiblemente con el tiempo.",
        "La activa, representada por la amenaza, que depende directamente de la actividad humana en la superficie o subsuelo y puede cambiar con el tiempo."
    ]

    for item in lista_vulnerabilidad:
        di11 = doc.add_paragraph(style='List Bullet')
        run = di11.add_run(f"{item}")
        di11_format = di11.paragraph_format
        di11_format.line_spacing = 1.15
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En el área sujeto de estudio por sus características litológicas está dentro del tipo cosechadoras de agua y con capacidad para la infiltración, dentro del proyecto no se tienen registrados pozos de importancia hídrica en los cuales se vea afectada o vulnerada la condición hidráulica con la implementación del proyecto en virtud de que solamente se presenta escurrimientos intermitentes típicos de lomeríos.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Métodos para proyectar alteraciones en la biocenosis y en general en los ecosistemas
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nMétodos para proyectar alteraciones en la biocenosis y en general en los ecosistemas')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Los ecosistemas son sistemas complejos como el bosque, el río o el lago, formados por una trama de elementos físicos (el biotopo) y biológicos (la biocenosis o comunidad de organismos). Hay que insistir en que la vida humana se desarrolla en estrecha relación con la naturaleza y que su funcionamiento nos afecta totalmente. Es un error considerar que nuestros avances tecnológicos: coches, grandes casas, industria, etc. nos permiten vivir al margen del resto de la biosfera y el estudio de los ecosistemas, de su estructura y de su funcionamiento, mismos que nos demuestra la profundidad de estas relaciones')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En este sentido, la evaluación de impacto ambiental (EIA) como método constituye una de las herramientas de protección ambiental que fortalece la toma de decisiones a nivel de políticas, planes, programas y proyectos, ya que incorpora variables que tradicionalmente no han sido consideradas durante su planificación, diseño o implementación.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('La evaluación de impacto ambiental, en el contexto actual, se entiende como un proceso de análisis que anticipa los futuros impactos ambientales negativos y positivos de acciones humanas permitiendo seleccionar las alternativas que, cumpliendo con los objetivos propuestos, maximicen los beneficios y disminuyan los impactos no deseados. La experiencia de diversos países permite su aplicación no sólo para grandes proyectos de inversión, sino también a actividades de desarrollo que involucren planes y programas de ordenamiento territorial, políticas y alternativas de acción, entre otras, que requieren de una variedad de proyectos individuales, evitando de esta forma los efectos acumulativos a nivel regional.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('El fin de una evaluación de impacto ambiental es identificar, predecir, valorar, prevenir o corregir y comunicar los efectos y los impactos ambientales producidos por una obra, discriminando entre las distintas alternativas. La selección de los factores ambientales y de las acciones de la obra conducen a identificar los posibles impactos ambientales y para evaluar estos se tienen los indicadores.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('En este aspecto en el área de estudio se han identificado los principales impactos y sus posibles efectos para cada una de las etapas del proyecto de acuerdo a las acciones que se implementarán para el desarrollo del proyecto, para ello se determina causa y efecto mismo que mediante matrices nos dará el resultado cuantitativo y cualitativo que servirán de base para determinar el grado de significancia y las medidas aplicables para mitigar los efectos en la implementación del proyecto.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 11.6
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 11.6 ###
    #########################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(f'\nIX.6. Conclusiones.')
    i11_format = capitulo11.paragraph_format
    i11_format.line_spacing = 1.15

    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True
    capitulo11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 11.6 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "La actividad contemplada para la ________________________ en el área en estudio es un polo importante de desarrollo con lo cual se proporcionan elementos para el crecimiento armónico de la actividad productiva que proporciona empleos conllevando a una mejor calidad de vida a la población aledaña (______________________) ya que se contratara personal de la región, el área sujeto de estudio de acuerdo a su vocación potencial es de tipo _____________ y ante lo errático de su productividad la implementación del proyecto como actividad alterna a la de uso actual, afectara en parte a la biodiversidad, por ser ______________________________________________________________________."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Con la información que se ha plasmado en el presente documento, dispone la autoridad normativa de los elementos técnicos suficientes que le permitan realizar la dictaminaciòn de la propuesta para llevar a cabo el cambio de utilización de terrenos forestales para la implementación del proyecto para la ______________________________."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Desde el punto de vista técnico y una vez analizados los elementos biológicos, (flora y fauna silvestre) geológicos, geofísicos y climáticos en el área de estudio citado se tuvieron los siguientes resultados."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "La superficie propuesta para el Cambio de Uso de Suelo, de acuerdo al análisis de campo y mapas correspondientes de acuerdo a INEGI que se anexa, se clasifica como _______________________________________________________________________________; aunque es de resaltar que actualmente la condición de la vegetación está dentro de la _______________ y cobertura. La superficie del proyecto se sitúa geográficamente en _________________ respectivos los cuales son parte del expediente que contiene el estudio."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Cabe destacar que, en el presente estudio, se dio cumplimiento a todos los elementos técnicos que de manera pormenorizada señalan las legislaciones forestal y ambiental vigentes, por lo que, en el mismo, se describen todas y cada una de las medidas de amortiguamiento y de mitigación de impactos en el área."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Con base al análisis y valoración de los impactos identificados se resumen de la siguiente forma: "
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Modelos de evacuación del suelo y de sus diferentes funciones.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nResultado de la valoración de impactos.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.6 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.8.- Resultados de valoración de los impactos.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.6 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    #########################
    ### Descripcion del capitulo 11.6 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Considerando dentro de la valoración mediante la aplicación del logaritmo se resume que los impactos se clasifican de acuerdo a la siguiente tabla:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 11.6 ###
    #########################
    tituloTabla11b = doc.add_paragraph()
    dti11b = tituloTabla11b.add_run('\nTabla 11.9.- Clasificación de los impactos.')
    dti11b_format = tituloTabla11b.paragraph_format
    dti11b_format.line_spacing = 1.15
    dti11b_format.space_after = 0

    dti11b.font.name = 'Bookman Old Style'
    dti11b.font.size = Pt(12)
    tituloTabla11b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 11.6 ###
    #########################
    filas = 8
    columnas = 6
    tabla11b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for cols in range(columnas):
        cell = tabla11b.cell(0, cols)
        cell_background_color(cell, '#3498db')

        for rows in range(filas):
            cell = tabla11b.cell(rows, cols)
            t11b = cell.paragraphs[0].add_run(' ')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 11.6 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo11_parrafo = doc.add_paragraph()
    imagenCapitulo11_run = imagenCapitulo11_parrafo.add_run('\n')
    imagenCapitulo11_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo11_run.add_picture('capitulo11/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 11.6 ###
    #########################
    tituloGrafico11 = doc.add_paragraph()
    dgi11 = tituloGrafico11.add_run('Grafica 1.- Resultados de la valoración de los impactos')
    dgi11_format = tituloGrafico11.paragraph_format
    dgi11_format.line_spacing = 1.15
    dgi11_format.space_after = 0

    dgi11.font.name = 'Bookman Old Style'
    dgi11.font.size = Pt(12)
    tituloGrafico11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 11.6 ###
    #########################
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Durante la implementación del proyecto en sus diferentes etapas se presentarán impactos de diferentes magnitudes e importancia de acuerdo a cada etapa del proceso.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Etapa de Preparación del Sitio.
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('\nEtapa de Preparación del Sitio.')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Acción: Desmonte y Despalme:")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Los Factores Ambientales donde se genera impacto:")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Atmósfera: Calidad del Aire y Calidad Sonora.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Se presentará volatilidad de partículas de polvo y smog a la atmosfera por el uso de maquinaria y movimientos de sustratos, siendo desde poco significativo hasta crítico ocurriendo este, en la etapa de desmonte y despalme.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Hidrología: Componente ambiental (Escurrimiento).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Se observa una valoración de poco significativa a moderados en virtud de la afectación de los escurrimientos al eliminar la vegetación incrementando la velocidad del flujo reduciendo la posibilidad de infiltración y ocasionando arrastre de sustrato.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Suelo: Componente ambiental (Erodabilidad).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Ante la eliminación de vegetación se tendrá exposición del sustrato generando erosión del tipo hídrica y eólica desde el punto de vista significativo hasta crítico.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Flora Silvestre: Componente ambiental (Densidad y Cobertura).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Como parte del proceso al eliminar la vegetación se tendrá perdida de densidad de individuos y especies afectando la cobertura generando un impacto crítico para efecto de conservar la Biodiversidad se contempla efectuar un programa de rescate de las especies consideradas como de lento desarrollo y difícil regeneración y/o aquellas que se adapten a este proceso.")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Fauna Silvestre: Componente ambiental (Abundancia y Hábitat).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Ante la presencia de maquinaria, equipo y seres humanos se tendrá un impacto significativo en referencia a que las especies serán ahuyentadas de su entorno, "
        "al igual que en la Flora se podrá aplicar un programa de rescate para aquellas especies de lenta movilidad en el área del proyecto."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Paisaje: Componente Ambiental (Armonía y Calidad Paisajística).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Resultado de esta acción del proyecto se tendrá una modificación que generará un impacto desde poco significativo hasta severos al modificar su entorno ante la eliminación "
        "de la vegetación y ocasionar oquedades por las características propias del proyecto."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        ############################
        Etapa de Construcción y operación:
    """
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Etapa de Construcción y operación:')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    #descripcionCapitulo9_format.space_after = 0
    #descripcionCapitulo9_format.space_before = 0

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    descripcionCapitulo11.bold = True
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En esta etapa se consideran impactos de carácter moderados, severos y críticos en sus diferentes acciones de acuerdo a las acciones que se implementaran:"
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Atmosfera. Componentes Ambientales (Calidad del aire y ruido).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Se tendrá un impacto moderado derivado del uso de maquinaria generando volatilidad de partículas y smog en el ambiente, así mismo se incrementará el ruido por el uso de los equipos afectando solo al personal que labora en la actividad al estar alejado de las poblaciones."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Hidrología. Superficial y Subterránea: Componentes Ambientales (escurrimientos, acuíferos)."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "De acuerdo al tipo de actividad está, modificara el escurrimiento natural en el área del proyecto, afectando en forma crítica inclusive la infiltración a los mantos freáticos, en las acciones de construcción de rampas de acceso y bancos de extracción."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Suelo. Componente ambiental (Erodabilidad).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "En esta actividad con los cortes y extracciones necesarios para habilitar las rampas y bancos de extracción se tendrán pérdidas de sustrato por acción erosiva de tipo hídrica y eólica generando un impacto crítico."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Fauna. Componente ambiental (Abundancia).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Ante la presencia de maquinaria, equipo y seres humanos se mantendrá el impacto de tipo significativo critico ante el desplazamiento que se genera por la alteración del medio donde se desarrollan las especies registradas."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run("Paisaje. Componentes Ambientales (Armonía y Calidad paisajística).")
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run(
        "Ante las modificaciones que se perfilan para la implementación del proyecto será crítico en virtud de generar una oquedad en la vegetación alterando la armonía y la calidad de paisaje al modificar su entorno."
    )
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.line_spacing = 1.15
    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 11 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo11() # Crear el documento