from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from utils import cell_background_color         # Importa la función para establecer el color de fondo de las celdas
from docx.enum.table import WD_ALIGN_VERTICAL   # Para alinear verticalmente el contenido de las celdas
from docx.enum.table import WD_TABLE_ALIGNMENT  # Para alinear la tabla en el centro de la página


def capitulo1():
    doc = Document()

    ########################################################################################################################################################################
    # Este es un comentario, por favor, si hay algo asi como este, no lo borres, sirve para especificar, que es lo que hace el código y vienen en todas partes
    ########################################################################################################################################################################

    """
        Si encuentras este comentario de este color, son notas y sugerencias que se pueden usar en el codigo en caso de hacer unos cambios
        que requieras.
    """
    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 1
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice Capítulo I.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    ########################################################################################################################################################################
    """
        La siguiente, variable muestra una lista de los temas del capitulo.

        Nota: La siguiente lista se puede editar para agregar o eliminar temas dependiendo de los formatos oficiales

        Ejemplo de como usar:

        -> Variable:

            temasCapituloX = ["Y.1.- Nombre del proyecto", 
                            "Y.2.- Nombre o Razón Social del Promovente", 
                            "Y.3.- Ubicación (dirección) del Promovente", 
                            "Y.4.- Superficie Solicitada de Cambio de Uso de Suelo y Tipo de Vegetacion Forestal",
                            "Y.5.- Duración del Proyecto"]

        -> Contexto del codigo anterior:
            * temasCapituloX es la variable: Se tiene que remplazar la X por el numero de capitulo, p. ej: temasCapitulo2.
            
            * "Y.X.- Nombre del tema": La Y se tiene que remplazar por el numero de capitulo (p. ej: I.X por el capitulo 1, II.X por el capitulo 2, etc.), 
                la X por el numero del tema (p. ej: I.1, I.2, etc) y el nombre del tema.
            
            * Para agregar un nuevo tema se tiene que poner una coma despues del texto entre comillas, p. ej: "Y.6.- Nombre del tema", Y.7.- Nombre del Tema, etc;
                no se puede poner fuera de [] tiene que estar dentro de los corchetes; y se tiene que poner entre comillas.
            
            * Para eliminar un elemento de la lista se tiene que poner un comentario, o bien eliminar el elemento de la lista, p. ej: "Y.X.- Nombre del proyecto",
                /* Se tiene que eliminar el texto junto con las comillas y con la coma que tiene despues.
                /* Se puede comentar si no se necesita por el momento, por ejemplo:
                    # "Y.X.- Nombre del proyecto",
            
            * Cuando se utilicen variables con listas, se tienen que usar corchetes [].
            
            * Si hay un numero de elementos en una lista, no se empieza desde el numero 1, sino desde el 0, por lo que el primer elemento es 0, el segundo es 1, etc.
                Se le tiene que restar siempre 1, ya que cuando queremos que solo me imprima el elemento poniendo el numero de elemento coloquial (si es primero es 1),
                va a imprimir el 2 elemento, en programacion el primer elemento o numero siempre es el cero y asi sucesivamente.
            
        -> Ejemplo de como usar la variable temasCapituloX en el codigo:
            * Si se quiere que se imprima el elemento de la lista, se tiene que poner el numero de elemento, por ejemplo:
                
                capitulo11 = doc.add_paragraph()
                i11 = capitulo11.add_run(temasCapitulo1[0])     ---------------------------> Esta linea es la que va imprimir el primer elemento, asi que su valor es 0.
                i11.font.name = 'Arial'
                i11.font.size = Pt(12)
                i11.font.bold = True

                capitulo12 = doc.add_paragraph()
                i12 = capitulo11.add_run(temasCapitulo1[1])     ---------------------------> Esta linea es la que va imprimir el segundo elemento, asi que su valor es 1.
                i12.font.name = 'Arial'
                i12.font.size = Pt(12)
                i12.font.bold = True
                
                A lo anterior es asi como se va hacer para lo de los capitulos.
    """

    temasCapitulo1 = ["I.1.- Nombre del proyecto", 
                    "I.2.- Nombre o Razón Social del Promovente", 
                    "I.3.- Ubicación (dirección) del Promovente", 
                    "I.4.- Superficie Solicitada de Cambio de Uso de Suelo y Tipo de Vegetacion Forestal",
                    "I.5.- Duración del Proyecto"]

    ########################################################################################################################################################################
    # Comienza Contenido
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    # Capitulo 1
    capitulo1 = doc.add_paragraph()
    i = capitulo1.add_run("I.- DATOS GENERALES DEL PROYECTO Y DEL PROMOVENTE")
    i_format = capitulo1.paragraph_format
    i_format.line_spacing = 1.5

    i.font.name = 'Arial'
    i.font.size = Pt(12)
    i.font.bold = True

    """
        =========================
            Contenido
        =========================
        Para el contenido de cada parrafo que se encuentra en el documento se muestra lo siguiente:
        
        # Contenido de Parrafo con texto #
            di11 = doc.add_paragraph()                                                      ---------------------> Esta linea es la que va agregar un nuevo parrafo.
            descripcionCapitulo11 = di11.add_run('Aqui va el contenido del parrafo')        ---------------------> Esta linea es la que va agregar el contenido del parrafo.
            descripcionCapitulo11_format = di11.paragraph_format                            ---------------------> Esta linea es la que va configurar el formato del parrafo.
            descripcionCapitulo11_format.line_spacing = 1.5                                 ---------------------> Esta linea es la que va configurar el interlineado del parrafo.

        # Contenido de Parrafo con variables #
            variable = 'El profe me roba las donitas (Contenido de parrafo)'    --------------> Esta es una variable con el contenido del parrafo 
            capitulo11 = doc.add_paragraph()                                    --------------> Esta linea es la que va agregar un nuevo parrafo.
            i11 = capitulo11.add_run(variable)                                  --------------> Esta linea es la que va agregar el contenido del parrafo desde la variable.
            i11.font.name = 'Arial'                                             --------------> Esta linea es la que va configurar el tipo de letra del parrafo.
            i11.font.size = Pt(12)                                              --------------> Esta linea es la que va configurar el tamaño de letra.
            i11.font.bold = True                                                --------------> Esta linea es la que va configurar el tipo de letra en negrita (Opcional).

        # Contenido si se tiene una lista (Ver nota anterior) #
            temasCapituloX = ["Tema 1", "Tema 2", "Tema 3"]                 -----------------> Esta es una variable que tiene una lista de temas.
            capituloX1 = doc.add_paragraph()                                -----------------> Esta linea es la que va agregar un nuevo parrafo.
            iX1 = capituloX1.add_run(temasCapituloX[Numero_de_elemento])    -----------------> Esta linea va mostrar el elemento, si es 1 es igual a 0, si es 2 es igual a 1 y asi sucesivamente.
            iX1.font.name = 'Arial'                                         -----------------> Esta linea es la que va configurar el tipo de letra del parrafo.
            iX1.font.size = Pt(12)                                          -----------------> Esta linea es la que va configurar el tamaño de letra.
            iX1.font.bold = True                                            -----------------> Esta linea es la que va a configurar el tipo de letra en negrita (Opcional).

        
        =========================
            Variables
        =========================
        Para las variables que se encuentran en el documento se muestra lo siguiente:
        -> Variables de temas de capitulo:
            * capituloXY - Representa la variable del capitulo, la X es el numero de capitulo, la Y es el tema, p. ej. capitulo11; esta variable significa que es el capitulo 1.1
            * capituloXYZ - Representa la variable del capitulo, la X es el numero de capitulo, la Y es el tema, la Z es el subtema, p. ej. capitulo121; esta variable significa que es el capitulo 1.2.1
            * i - Representa la variable del capitulo
            * iXYZ - Ver puntos anteriores, pero son para tipo de letra, tamaño de fuente y estilo de fuente del titulo de capitulo.
            * diXYZ - Representa que se va a agregar un nuevo parrafo del capitulo y la edicion de la fuente
            * descripcionCapituloXYZ - Representa la variable del contenido que va a tener el capitulo
            * descripcionCapituloXYZ_format - Representa la variable del formato del contenido del capitulo
    """

    ########################################################################################################################################################################
    # Capitulo 1.1
    ########################################################################################################################################################################
    capitulo11 = doc.add_paragraph()
    i11 = capitulo11.add_run(temasCapitulo1[0])
    i11.font.name = 'Arial'
    i11.font.size = Pt(12)
    i11.font.bold = True

    # Texto del Capitulo 1.1
    di11 = doc.add_paragraph()
    descripcionCapitulo11 = di11.add_run('Proyecto "Nombre del Proyecto", que para la identificación en este documento se referirá al proyecto como "tipo de material".')
    descripcionCapitulo11_format = di11.paragraph_format
    descripcionCapitulo11_format.space_after = Pt(0)
    descripcionCapitulo11_format.line_spacing = 1.5

    descripcionCapitulo11.font.name = 'Arial'
    descripcionCapitulo11.font.size = Pt(12)
    di11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #di11.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.2
    ########################################################################################################################################################################
    capitulo12 = doc.add_paragraph()
    i12 = capitulo12.add_run(f'\n{temasCapitulo1[1]}')
    i12_format = capitulo12.paragraph_format
    i12_format.line_spacing = 1.15

    i12.font.name = 'Arial'
    i12.font.size = Pt(12)
    i12.font.bold = True

    # capitulo12.add_run().add_break()  # Salto de línea

    # ====================================================================== #
    # Texto del Capitulo 1.2
    # ====================================================================== #
    naEm = doc.add_paragraph()

    nameEmpresa = naEm.add_run('NOMBRE DE LA EMPRESA')
    nameEmpresa_format = naEm.paragraph_format
    nameEmpresa_format.space_after = Pt(0)
    nameEmpresa_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

    nameEmpresa.font.name = 'Arial'       # Tipo de Fuente
    nameEmpresa.font.size = Pt(12)        # Tamaño de la Fuente
    nameEmpresa.font.bold = True
    # naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro


    di12 = doc.add_paragraph()

    descripcionCapitulo12 = di12.add_run('Descripcion de la empresa')
    descripcionCapitulo12_format = di12.paragraph_format
    descripcionCapitulo12_format.space_after = Pt(0)
    descripcionCapitulo12_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

    descripcionCapitulo12.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo12.font.size = Pt(12)        # Tamaño de la Fuente
    di12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al centro

    # di12.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.2.1
    ########################################################################################################################################################################
    capitulo121 = doc.add_paragraph()
    i121 = capitulo121.add_run("\n1.2.1.- Representante legal")
    i121_format = capitulo121.paragraph_format
    i121_format.line_spacing = 1.15

    i121.font.name = 'Arial'
    i121.font.size = Pt(12)
    i121.font.bold = True

    # capitulo121.add_run().add_break()  # Salto de línea

    # ====================================================================== #
    # Texto del Capitulo 1.2.1
    # ====================================================================== #
    legalRepresentante = doc.add_paragraph()

    legalRepresentant = legalRepresentante.add_run('NOMBRE DEL REPRESENTANTE LEGAL')
    legalRepresentant_format = legalRepresentante.paragraph_format
    legalRepresentant_format.space_after = Pt(0)
    legalRepresentant_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

    legalRepresentant.font.name = 'Arial'       # Tipo de Fuente
    legalRepresentant.font.size = Pt(12)        # Tamaño de la Fuente
    legalRepresentant.font.bold = True
    # naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación justificada


    di121 = doc.add_paragraph()

    descripcionCapitulo121 = di121.add_run('Descripcion del Representante Legal')
    descripcionCapitulo121_format = di121.paragraph_format
    descripcionCapitulo121_format.line_spacing = 1.5  # Interlineado de 1.5 líneas

    descripcionCapitulo121.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo121.font.size = Pt(12)        # Tamaño de la Fuente
    di121.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación justificada

    # di121.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.3
    ########################################################################################################################################################################
    capitulo13 = doc.add_paragraph()
    i13 = capitulo13.add_run(f'\n{temasCapitulo1[2]}')
    i13_format = capitulo13.paragraph_format
    i13_format.line_spacing = 1.15

    i13.font.name = 'Arial'
    i13.font.size = Pt(12)
    i13.font.bold = True

    # capitulo13.add_run().add_break()  # Salto de línea

    # ====================================================================== #
    # Texto del Capitulo 1.3
    # ====================================================================== #
    domEm = doc.add_paragraph()

    domicilioEmpresa = domEm.add_run('Domicilio Fiscal')
    domicilioEmpresa_format = domEm.paragraph_format
    domicilioEmpresa_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

    domicilioEmpresa.font.name = 'Arial'       # Tipo de Fuente
    domicilioEmpresa.font.size = Pt(12)        # Tamaño de la Fuente
    domicilioEmpresa.font.bold = True
    # naEm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación jusrificada


    di13 = doc.add_paragraph()

    descripcionCapitulo13 = di13.add_run('Calle XXXXX, No. XYZ\n'
                                        'Colonia: XXXXXX\n'
                                        'Municipio: XXXXXX\n'
                                        'Estado: Coahuila\n'
                                        'C.P.: XXXXX\n'
                                        'Telefono: XXX-XXX-XXXX\n'
                                        'Correo: example@email.com')
    descripcionCapitulo13_format = di13.paragraph_format
    # descripcionCapitulo13_format.space_after = Pt(0)
    descripcionCapitulo13_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

    descripcionCapitulo13.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo13.font.size = Pt(12)        # Tamaño de la Fuente
    # di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación al justificada

    di13.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.3.1
    ########################################################################################################################################################################
    capitulo131 = doc.add_paragraph()
    i131 = capitulo131.add_run('1.3.1 Domicilio para oír y recibir notificaciones')
    i131_format = capitulo131.paragraph_format
    i131_format.line_spacing = 1.15

    i131.font.name = 'Arial'
    i131.font.size = Pt(12)
    i131.font.bold = True

    # capitulo13.add_run().add_break()  # Salto de línea

    # ====================================================================== #
    # Texto del Capitulo 1.3.1
    # ====================================================================== #
    di131 = doc.add_paragraph()

    descripcionCapitulo131 = di131.add_run('Calle XXXXX, No. XYZ\n'
                                        'Colonia: XXXXXX\n'
                                        'Municipio: XXXXXX\n'
                                        'Estado: Coahuila\n'
                                        'C.P.: XXXXX\n'
                                        'Telefono: XXX-XXX-XXXX\n'
                                        'Correo: example@email.com')
    descripcionCapitulo131_format = di131.paragraph_format
    # descripcionCapitulo131_format.space_after = Pt(0)
    descripcionCapitulo131_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

    descripcionCapitulo131.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo131.font.size = Pt(12)        # Tamaño de la Fuente
    # di13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY     # Alineación justificada

    di131.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.4
    ########################################################################################################################################################################
    capitulo14 = doc.add_paragraph()
    i14 = capitulo14.add_run(temasCapitulo1[3])
    i14_format = capitulo14.paragraph_format
    i14_format.line_spacing = 1.15

    i14.font.name = 'Arial'
    i14.font.size = Pt(12)
    i14.font.bold = True

    # capitulo13.add_run().add_break()  # Salto de línea

    # ====================================================================== #
    # Nombre de la tabla 1.1
    # ====================================================================== #
    tituloTabla11 = doc.add_paragraph()
    t11 = tituloTabla11.add_run('Tabla 1.1.- Distribución de Superficies.')
    t11_format = tituloTabla11.paragraph_format
    t11_format.space_after = Pt(0)

    t11.font.name = 'Courier New'
    t11.font.size = Pt(12)
    # t11.font.bold = True

    # ====================================================================== #
    # Tabla 1.1
    # ====================================================================== #

    # Distribucion de Superficies
    tabla11 = doc.add_table(rows=3, cols=3, style='Table Grid')
    tabla11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in tabla11.rows:
        tabla11.columns[0].width = Cm(8.28)
        tabla11.columns[1].width = Cm(4.94)
        tabla11.columns[2].width = Cm(2.12)

    cell = tabla11.cell(0, 0)
    t11 = cell.paragraphs[0].add_run('DISTRIBUCION DE SUPERFICIES')
    t11.font.name = 'Arial'
    t11.font.size = Pt(12)
    t11.font.bold = True
    t11._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Superficies (ha)
    cell = tabla11.cell(0, 1)
    t11 = cell.paragraphs[0].add_run('SUPERFICIE (ha)')
    t11.font.name = 'Arial'
    t11.font.size = Pt(12)
    t11.font.bold = True
    t11._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Porcentaje
    cell = tabla11.cell(0, 2)
    t11 = cell.paragraphs[0].add_run('%')
    t11.font.name = 'Arial'
    t11.font.size = Pt(12)
    t11.font.bold = True
    t11._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Superficies para ACUSTFF
    cell = tabla11.cell(1, 0)
    t11 = cell.paragraphs[0].add_run('Superficie para ACUSTF')
    t11.font.name = 'Arial'
    t11.font.size = Pt(12)
    t11.font.bold = True
    t11._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Superficies total
    cell = tabla11.cell(2, 0)
    t11 = cell.paragraphs[0].add_run('Superficie Total')
    t11.font.name = 'Arial'
    t11.font.size = Pt(12)
    t11.font.bold = True
    t11._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columan de Distribucion de Superficies
    for row in range(3):
        cell = tabla11.cell(row, 0)
        cell.width = Cm(8.28)  # Ancho de la columna 1

    # Columna de Superficies (ha)
    for row in range(3):
        cell = tabla11.cell(row, 1)
        cell.width = Cm(4.94)  # Ancho de la columna 2
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    # Columna de Porcentaje
    for row in range(3):
        cell = tabla11.cell(row, 2)
        cell.width = Cm(2.12)  # Ancho de la columna 3
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación al centro
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    # ====================================================================== #
    # Texto del capitulo 1.4
    # ====================================================================== #

    #Viñeta despues de la tabla 1.1
    dit14 = doc.add_paragraph(style='ListBullet') # Estilo de la lista
    descripcionCapitulo14table = dit14.add_run('Segun contrato de arrendamiento (Texto Opcional)') # Texto de parrafo
    descripcionCapitulo14table_format = dit14.paragraph_format                    #Impresion de texto
    descripcionCapitulo14table_format.space_before = Pt(0)                      # Espacio antes del cuadro
    descripcionCapitulo14table_format.space_after = Pt(0)                      # Espacio despues
    descripcionCapitulo14table_format.line_spacing = 1.15                         # Interlineado de 1.15 líneas

    descripcionCapitulo14table.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo14table.font.size = Pt(12)        # Tamaño de la Fuente
    # dit14.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro

    # dit14.add_run().add_break() # Salto de linea

    #Texto despues de la viñeta
    di14 = doc.add_paragraph()
    descripcionCapitulo14 = di14.add_run('La vegetación que se presenta dentro del area de cambio de usos de suelo es el "______________" ocupando una superficie de XXX ha la que representa el XX %, el "____________" el cual ocupa una superficie de XXX ha representado el XXX % de la superficie según la carta de vegetación en su serie VII del INEGI está determinado que su principal uso es de carácter pecuario y forestal, dado, por las características de la comunidad vegetativa, donde se tienen especies adecuadas para el ramoneo del ganado caprino en forma extensiva. Dentro de la superficie contemplada para la implementación del proyecto se tiene suelo dominante de "_______________" de "______________", ocasionando que la vegetación tenga baja cobertura ante la poca disponibilidad de sustrato para su desarrollo.')
    descripcionCapitulo14_format = di14.paragraph_format
    descripcionCapitulo14_format.space_before = Pt(0)
    descripcionCapitulo14_format.space_after = Pt(0)
    descripcionCapitulo14_format.line_spacing = 1.15

    descripcionCapitulo14.font.name = 'Arial'
    descripcionCapitulo14.font.size = Pt(12)
    di14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # di14.add_run().add_break()

    ########################################################################################################################################################################
    # Capitulo 1.5
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'\n{temasCapitulo1[4]}')
    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True

    # ====================================================================== #
    # Texto del Capitulo 1.5
    # ====================================================================== #
    di151 = doc.add_paragraph()
    descripcionCapitulo151 = di151.add_run('El proyecto objetivo del estudio que se ubica "______________", contemplado para implementar el proyecto de “_________________________________” se tiene proyectado llevar a cabo la implementación del mismo en un periodo de _____ años donde se pretende aprovechar _____ ha por año, considerando para ello 4 etapas adicionales y actividades que se desarrollarán en forma paulatina, siendo las que a continuación se describen:')
    descripcionCapitulo151_format = di151.paragraph_format
    descripcionCapitulo151_format.line_spacing = 1.15  # Interlineado de 1.5 líneas

    descripcionCapitulo151.font.name = 'Arial'       # Tipo de Fuente
    descripcionCapitulo151.font.size = Pt(12)        # Tamaño de la Fuente
    di151.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY     # Alineación al centro

    #di151.add_run().add_break()

    # ====================================================================== #
    # Titulo de la tabla 1.2
    # ====================================================================== #
    tituloTabla12 = doc.add_paragraph()
    t12 = tituloTabla12.add_run('Tabla 1.2.- Etapas y actividades del proyecto.')
    t12_format = tituloTabla12.paragraph_format
    t12_format.space_after = Pt(0)
    t12_format_line_spacing = 1

    t12.font.name = 'Courier New'
    t12.font.size = Pt(12)
    # t11.font.bold = True

    # ====================================================================== #
    # Tabla 1.2
    # ====================================================================== #

    # tabla15columnas = 15

    # Distribucion de Superficies
    tabla15 = doc.add_table(rows=5, cols=4, style='Table Grid')
    tabla15.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Etapas
    cell = tabla15.cell(0, 0)
    t15 = cell.paragraphs[0].add_run('Etapa')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Actividad
    cell = tabla15.cell(0, 1)
    t15 = cell.paragraphs[0].add_run('Actividad')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Periodo
    cell = tabla15.cell(0, 2)
    t15 = cell.paragraphs[0].add_run('Periodo')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Superficie en ha
    cell = tabla15.cell(0, 3)
    t15 = cell.paragraphs[0].add_run('Superficie en ha')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(cell, 'B8CCE4')

    # Primera
    cell = tabla15.cell(1, 0)
    t15 = cell.paragraphs[0].add_run('Primera')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Segunda
    cell = tabla15.cell(2, 0)
    t15 = cell.paragraphs[0].add_run('Segunda')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tercera
    cell = tabla15.cell(3, 0)
    t15 = cell.paragraphs[0].add_run('Tercera')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cuarta
    cell = tabla15.cell(4, 0)
    t15 = cell.paragraphs[0].add_run('Cuarta')
    t15.font.name = 'Arial'
    t15.font.size = Pt(12)
    t15.font.bold = True
    t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Establecer el ancho de la columna de etapas
    for row in range(5):
        cell = tabla15.cell(row, 0)
        cell.width = Cm(2.20)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    # Establecer el ancho de la columna de actividades
    for row in range(5):
        cell = tabla15.cell(row, 1)
        cell.width = Cm(5.70)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    # Establecer el ancho de la columna de periodos
    for row in range(5):
        cell = tabla15.cell(row, 2)
        cell.width = Cm(4.15)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación horizontal al centro
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    # Establecer el ancho de la columna de superficie en ha
    for row in range(5):
        cell = tabla15.cell(row, 3)
        cell.width = Cm(3.00)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación horizontal al centro
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Alineación vertical al centro

    """
    for i in range (0, tabla15columnas):
        cell = tabla15.cell(i, 0)
        t15 = cell.paragraphs[0].add_run('')
        t15.font.name = 'Arial'
        t15.font.size = Pt(12)
        t15.font.bold = True
        t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial)
        
        cell = tabla15.cell(i, 1)
        t15 = cell.paragraphs[0].add_run('')
        t15.font.name = 'Arial'
        t15.font.size = Pt(12)
        t15.font.bold = True
        t15._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial)
    """

    # ====================================================================== #
    # Nota del Capitulo 1.5 para la tabla 1.2
    # ====================================================================== #

    nota15 = doc.add_paragraph()
    n15 = nota15.add_run('El área destinada para el proyecto contempla una superficie total de XXXX ha, pero solo el XXXX ha serán destinadas para la extraccion de material xxxxxxx, mientras que el XXX ha restantes serán destinadas para otras infraestructuras.')
    n15_format = nota15.paragraph_format
    n15_format.space_before = Pt(0)
    n15_format.space_after = Pt(0)
    n15_format.line_spacing = 1.15

    n15.font.name = 'Arial'
    n15.font.size = Pt(9)
    nota15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ====================================================================== #
    # Texto del Capitulo 1.5
    # ====================================================================== #
    di152 = doc.add_paragraph()
    descripcionCapitulo152 = di152.add_run('\nEn relación al cuadro anterior sería la vigencia del citado estudio, así también previo a estas etapas se contempla el proceso de elaboración del estudio, evaluación y gestión hasta su resolutivo e implementación del proyecto de acuerdo a s etapas descritas, los periodos de ejecución para cada una de ellas se determina el plazo en virtud de que serán desarrolladas en forma paulatina y de acuerdo a la necesidad de _________________________ y la demanda de la región.')
    descripcionCapitulo152_format = di152.paragraph_format
    descripcionCapitulo152_format.space_before = Pt(0)
    descripcionCapitulo152_format.space_after = Pt(0)
    descripcionCapitulo152_format.line_spacing = 1.15

    descripcionCapitulo152.font.name = 'Arial'
    descripcionCapitulo152.font.size = Pt(12)
    di152.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 1 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo1()
