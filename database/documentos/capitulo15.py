"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes, celdas, etc; en pulgadas y en centrimetros
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos
from utils import quitar_borde_especifico       # Importar la función para quitar el borde de una celda
from utils import quitar_bordes_tabla           # Importar la función para quitar los bordes de una tabla
from utils import quitar_bordes_celda           # Importar la función para quitar los bordes de una celda

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""

def capitulo15():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 15
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo XV.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    doc.add_page_break() # Salto de página

    ########################################################################################################################################################################
    # Capitulo 15
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 15 ###
    #########################
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'XV.-	Datos de Inscripción en el Registro del Prestador de Servicios Forestales que Haya Elaborado el Estudio y, del Que Estará a Cargo de la Ejecución del Cambio De Uso De Suelo.')
    i15_format = capitulo15.paragraph_format
    i15_format.line_spacing = 1.15

    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True
    capitulo15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 15.1
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 15.1 ###
    #########################
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'XV.1 Prestador de servicios técnicos.')
    i15_format = capitulo15.paragraph_format
    i15_format.line_spacing = 1.15

    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True
    capitulo15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Nombre
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 15.1 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "a) Nombre:"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    descripcionCapitulo15.bold = True
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   Ing. Francisco Mancilla Barboza"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Domicilio
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 15.1 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "b) Domicilio:"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    descripcionCapitulo15.bold = True
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   Calle Novena # 235 Col. Brisas Poniente C.P. 25225, Saltillo, Coahuila, tel. 844 2307257. Correo: mingopago@yahoo.com, multiambientales@hotmail.com"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Número de inscripción en Registro Forestal Nacional (Número, libro, tipo y volumen).
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 15.1 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "c) Número de inscripción en Registro Forestal Nacional (Número, libro, tipo y volumen).:"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    descripcionCapitulo15.bold = True
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   REGISTRO FORESTAL NACIONAL EN EL LIBRO COAH, TUI, "
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   CVOL 6, NÚM. 4 AÑO 12."
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Copia de su inscripción en el registro forestal del prestador de servicios técnicos forestales.
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 15.1 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "d) Copia de su inscripción en el registro forestal del prestador de servicios técnicos forestales.:"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    descripcionCapitulo15.bold = True
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   Se Anexa Copia Del Registro Forestal"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
            Copia simple de identificación oficial (Credencial de Elector, Pasaporte, Cartilla del SMN, Cedula Profesional).
        |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
    """
    #########################
    ### Descripcion del capitulo 15.1 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "e) Copia simple de identificación oficial (Credencial de Elector, Pasaporte, Cartilla del SMN, Cedula Profesional)."
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    descripcionCapitulo15.bold = True
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "   Se anexa identificación oficial expedida por el Instituto Nacional Electoral."
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 15.2
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 15.2 ###
    #########################
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'\nXV.2 Colaboradores o participantes')
    i15_format = capitulo15.paragraph_format
    i15_format.line_spacing = 1.15

    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True
    capitulo15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 15.2 ###
    #########################
    """
        -------------------------------------
            Datos a rellenar
        -------------------------------------
    """

    encabezados = [
        'NOMBRE',
        'CEDULA PROFESIONAL',
        'PROFESIÓN',
    ]

    datos_tabla = [
        ['ALMA GUADALUPE LUNA CASANOVA', '9350486', 'INGENIERO FORESTAL'],
        ['FREDDY SÁNCHEZA AGUILAR', '7374879', 'INGENIERO FORESTAL'],
        ['GABRIEL AIN HERRERA MARTINEZ', '10348194', 'INGENIERO FORESTAL'],
    ]

    filas = len(datos_tabla) + 1  # 1 fila para encabezados
    columnas = len(encabezados)   # 4 columnas
    tabla12b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')

    for rows in tabla12b.rows:
        rows.cells[0].width = Cm(8.19)
        rows.cells[1].width = Cm(3.39)
        rows.cells[2].width = Cm(5.61)

    # ✅ Encabezados
    for col in range(columnas):
        cell = tabla12b.cell(0, col)
        t12b = cell.paragraphs[0].add_run(encabezados[col])
        t12b.font.name = 'Arial'
        t12b.font.size = Pt(12)
        t12b.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ✅ Filas de datos
    for fila in range(len(datos_tabla)):            # 0..2
        for col in range(len(datos_tabla[fila])):   # 0..3
            cell = tabla12b.cell(fila + 1, col)     # +1 porque fila 0 es encabezado
            texto = datos_tabla[fila][col]
            t12b = cell.paragraphs[0].add_run(texto)
            t12b.font.name = 'Arial'
            t12b.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    ########################################################################################################################################################################
    # Capitulo 15.3
    ########################################################################################################################################################################

    #########################
    ### Salto de Pagina ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 15.3 ###
    #########################
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'XV.3 Firmas')
    i15_format = capitulo15.paragraph_format
    i15_format.line_spacing = 1.15

    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True
    capitulo15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    """
        -------------------------------------
            Firma del prestador de servicios forestales
        -------------------------------------
    """
    #########################
    ### Descripcion del capitulo 15.3 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "\n\n\n\n____________________________________"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "Ing. Francisco Mancilla Barboza"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.bold = True
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "Prestador de servicios Técnicos Forestales"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.CENTER

    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "Responsable de la Elaboración del Proyecto\n\n\n\n\n\n\n"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1
    descripcionCapitulo15_format.space_after = 0
    descripcionCapitulo15_format.space_before = 0
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.CENTER

    """
        -------------------------------------
            Firma del cliente o Representante legal del cliente
        -------------------------------------
    """
    #########################
    ### Tabla del capítulo 15.3 ###
    #########################
    filas = 3
    columnas = 1
    tabla15b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    quitar_bordes_tabla(tabla15b)
    tabla15b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for rows in tabla15b.rows:
        rows.cells[0].width = Cm(8.46)

    cell = tabla15b.cell(0, 0)
    t15b = cell.paragraphs[0].add_run('____________________________________')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla15b.cell(1, 0)
    t15b = cell.paragraphs[0].add_run('Nombre del representante legal')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    t15b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla15b.cell(2, 0)
    t15b = cell.paragraphs[0].add_run('Representante Legal Jibe Construcciones y Pavimentos')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ########################################################################################################################################################################
    # Capitulo 15.4
    ########################################################################################################################################################################

    #########################
    ### Titulo del capitulo 15.4 ###
    #########################
    capitulo15 = doc.add_paragraph()
    i15 = capitulo15.add_run(f'\n\n\n\nXV.4.- Nombre del responsable de dirigir la ejecución del cambio de uso del suelo autorizado.')
    i15_format = capitulo15.paragraph_format
    i15_format.line_spacing = 1.15

    i15.font.name = 'Arial'
    i15.font.size = Pt(12)
    i15.font.bold = True
    capitulo15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 15.4 ###
    #########################
    di15 = doc.add_paragraph()
    descripcionCapitulo15 = di15.add_run(
        "\n\n\n\n\n\n\n"
    )
    descripcionCapitulo15_format = di15.paragraph_format
    descripcionCapitulo15_format.line_spacing = 1.15
    descripcionCapitulo15.font.name = 'Arial'
    descripcionCapitulo15.font.size = Pt(12)
    di15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 15.3 ###
    #########################
    filas = 3
    columnas = 1
    tabla15b = doc.add_table(rows=filas, cols=columnas, style='Table Grid')
    quitar_bordes_tabla(tabla15b)
    tabla15b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for rows in tabla15b.rows:
        rows.cells[0].width = Cm(8.46)

    cell = tabla15b.cell(0, 0)
    t15b = cell.paragraphs[0].add_run('____________________________________')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla15b.cell(1, 0)
    t15b = cell.paragraphs[0].add_run('Nombre del representante legal')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    t15b.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell = tabla15b.cell(2, 0)
    t15b = cell.paragraphs[0].add_run('Representante Legal Jibe Construcciones y Pavimentos')
    t15b.font.size = Pt(12)
    t15b.font.name = 'Arial'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 15 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo15() # Crear el documento
