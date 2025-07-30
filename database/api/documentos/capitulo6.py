"""
    ============================================================
    Archivos de las dependencias del proyecto
    ============================================================
"""
from docx import Document                       # Es el modulo principal para crear documentos word
from docx.shared import Pt                      # Para el tamaño en puntos
from docx.oxml.ns import qn                     # Para nombres de fuentes asiáticas
from docx.enum.text import WD_ALIGN_PARAGRAPH   # Para alinear texto
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Cm              # Para el tamaño de las imagenes
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK             # Para los saltos de linea
from docx.enum.section import WD_ORIENT         # Para la orientación de la página
from docx.enum.section import WD_SECTION        # Para secciones del documento
from docx.shared import RGBColor                # Para colores RGB
from docx.enum.table import WD_ROW_HEIGHT_RULE  # Para el alto de las filas de las tablas


""" 
    ============================================================
    Archivos locales desarrollados por el usuario
    ============================================================
"""
from utils import cell_background_color         # Importar la función para cambiar el color de fondo de las celdas
from utils import entero_a_romano               # Importar la función para convertir números a romanos

""" 
    ============================================================
    Creacion del documento
    ============================================================
"""
def capitulo6():
    doc = Document()  # Crear un nuevo documento, tambien es una variable global

    ########################################################################################################################################################################
    # Establecer los margenes del documento
    ########################################################################################################################################################################
    margin = doc.sections
    inch = 2.54 # Una pulgada es igual a 2.54 centimetros

    # Los numero de las siguientes variables se pueden cambiar dependiendo de las necesidades del usuario tienen que estar en centimetros
    top_cm = 2.5 # Margen superior en centimetros
    bottom_cm = 2.5 # Margen inferior en centimetros
    left_cm = 2.5 # Margen izquierdo en centimetros
    right_cm = 2 # Margen derecho en centimetros

    # Convierte los centimetros en pulgadas; Nota: Las siguientes variables no se pueden editar
    top_inch = top_cm / inch
    bottom_inch = bottom_cm / inch
    left_inch = left_cm / inch
    right_inch = right_cm / inch

    """
        Nota: Los margenes se pueden establecer de la siguiente manera:
        - top_margin: Margen superior
        - bottom_margin: Margen inferior
        - left_margin: Margen izquierdo
        - right_margin: Margen derecho

        Python no tiene una función para establecer los margenes en centimetros, por lo que se convierten a pulgadas
    """
    for section in margin:
        section.top_margin = Inches(top_inch)           # Margen superior
        section.bottom_margin = Inches(bottom_inch)     # Margen inferior
        section.left_margin = Inches(left_inch)         # Margen izquierdo
        section.right_margin = Inches(right_inch)       # Margen derecho

    ########################################################################################################################################################################
    # Indice Capitulo 6
    ########################################################################################################################################################################
    p = doc.add_paragraph()

    # Añadir texto con estilo personalizado
    indice = p.add_run("Índice de Contenido Capitulo VI.")

    # Cambiar el tipo de letra y tamaño
    indice.font.name = 'Bookman Old Style'      # Tipo de letra
    indice.font.size = Pt(12)                   # Tamaño de la letra
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER     # Alineación al centro
    indice.bold = True

    ########################################################################################################################################################################
    # Capitulo 6
    ########################################################################################################################################################################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.-	Análisis comparativo de la composición florística y faunística del área sujeta a Cambio de uso de suelo en Terrenos Forestales con relación a los tipos de vegetación del ecosistema de la cuenca, subcuenca o microcuenca hidrográfica, que permita determinar el grado de afectación por el Cambio de Uso de Suelo en Terrenos forestales.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.- Comparativos de composición de la vegetación presente dentro del ACUSTF y Sistema Ambiental.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('En este capítulo se analizarán _______________________________________________________________ con sus cuatro comparativos en composición florística: Comparativo de individuos e índice de valor de importancia; es decir, total de individuos extrapolados y su porcentaje de afectación al Sistema Ambiental por el Cambio y uso de Suelo, Comparativo por índices de biodiversidad, Comparativo por valor densidad de especies y Comparativos de Índices de similitud/disimilitud.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.1- Comparativo de individuos e Índice de Valor de Importancia por estrato del Sistema Ambiental –ACUSTF en el ______________________')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Según Aguirre (1999) el índice de valor de importancia (IVI), indica que tan importante es una especie dentro de la comunidad. Las especies que tienen el IVI más alto significa entre otras cosas que es dominante ecológicamente: que absorbe muchos nutrientes, que ocupa mayor espacio físico, que controla en un porcentaje alto la energía que llega a este sistema. Este índice sirve para comparar el peso ecológico de cada especie dentro del ecosistema. Para calcular este parámetro se utiliza la Densidad relativa, Frecuencia relativa y Dominancia relativa.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('A continuación, se describen las fórmulas que se utilizaron para la estimación del Índice de Valor de Importancia.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDensidad relativa. Está dada por el resultado de la densidad absoluta entre el número total de todos los individuos muestreados expresados en porcentajes ')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Formula del Capitulo 6.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_2.png', width=Cm(4.79), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del Capitulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDer = Densidad Relativa'
                                    '\nNi = Número de individuos de la especie'
                                    '\nNt = Número total de individuos de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nFrecuencia relativa. Es el resultado de dividir la frecuencia absoluta de cada especie entre el número total de esas especies expresadas en porcentajes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_6.png', width=Cm(4.74), height=Cm(1.50))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nFr = Frecuencia relativa'
                                    '\nFai = Frecuencia absoluta de cada especie'
                                    '\nFat = Frecuencia absoluta de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nLa dominancia relativa. Se calcula como la proporción de una especie en el área total evaluada, expresada en porcentaje.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_4.png', width=Cm(4.79), height=Cm(1.50))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDor = Densidad relativa'
                                    '\nDai = Densidad absoluta de una especie'
                                    '\nDat= Densidad absoluta total de todas las especies')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de valor de importancia (IVI). El índice de valor de importancia define cuáles de las especies presentes contribuyen en el carácter y estructura de una Comunidad. Este valor se obtiene mediante la sumatoria de la frecuencia relativa, la densidad relativa y la dominancia relativa.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo611/formula_7.png', width=Cm(4.99), height=Cm(1.20))
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nIVI = Índice de Valor de Importancia'
                                    '\nDer = Densidad relativa'
                                    '\nDor = Dominancia relativa'
                                    '\nFr = Frecuencia relativa')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capítulo 6.1.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 6.1.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.1.- Comparativo por total de individuos e Índice de Valor de Importancia en el ____')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.1 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=9, style='Table Grid')

    for cols in range(9):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescribir los del cuadro.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver  resto del contenido del capítulo 6.1.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripción del capítulo 6.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nEn general las especies de lento crecimiento y las enlistadas en la NOM- 059- SEMARNAT 2010 se rescatarán y reubicarán a una superficie que tenga las mismas condiciones donde se distribuyen actualmente para que no pierdan su germoplasma.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescripcion.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina en el capitulo 6.1.2 ###
    #########################
    doc.add_page_break() # Salto de página

    #########################
    ### Titulo del capitulo 6.1.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'VI.1.2.- Comparativo por índices de biodiversidad del Sistema Ambiental –ACUSTF del _____________________________.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices basados en la dominancia son parámetros inversos al concepto de uniformidad o equidad de la comunidad. Toman en cuenta la representatividad de las especies con mayor valor de importancia sin evaluar la contribución del resto de las especies. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para medir la dominancia de las especies los índices de biodiversidad más comunes son: Simpson y Berger Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 6.1.2 ###
    #########################
    listaValores612 = [
        '0 - 0.33',
        '0.34 - 0.66',
        '> 0.67',
    ]

    diversidadSignificancia612 = [
        'Diversidad Baja',
        'Diversidad Media',
        'Diversidad Alta',
    ]

    heterogeneoHomogeneoSignificancia612 = [
        'Heterogéneo en abundancia',
        'Ligeramente Heterogéneo en en abundancia',
        'Homogéneo en abundancia',
    ]

    valores612 = range(len(listaValores612))
    diversidad612 = range(len(diversidadSignificancia612))
    heterogeneoHomogeneo612 = range(len(heterogeneoHomogeneoSignificancia612))

    filasCap612 = len(valores612) + 2
    
    tabla6 = doc.add_table(rows=filasCap612, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(3.46)
        rows.cells[1].width = Cm(3.66)
        rows.cells[2].width = Cm(9.66)

    #########################
    # Celda fusionada "Escalas de interpretación de significancia 0-1"
    row1 = tabla6.rows[0]
    merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Escalas de interpretación de significancia 0-1')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell1, '0070C0')

    #########################
    # Celda fusionada "Significancia"
    row1 = tabla6.rows[1]
    merged_cell1 = row1.cells[1].merge(row1.cells[1].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Significancia')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla6.cell(1, 0)
    t6 = cell.paragraphs[0].add_run('Valores')
    t6.font.size = Pt(12)
    t6.font.name = 'Arial'
    t6.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in valores612:
        cell = tabla6.cell(cols + 2, 0)
        t6 = cell.paragraphs[0].add_run(f'{listaValores612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        t6.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in diversidad612:
        cell = tabla6.cell(cols + 2, 1)
        t6 = cell.paragraphs[0].add_run(f'{diversidadSignificancia612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in heterogeneoHomogeneo612:
        cell = tabla6.cell(cols + 2, 2)
        t6 = cell.paragraphs[0].add_run(f'{heterogeneoHomogeneoSignificancia612[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el Índice de Margalef el criterio es de 2-5, donde, sus escalas de interpretación son: de 0-2 se considera diversidad baja, de 2-5 se considera diversidad media y mayor de 5 se considera diversidad alta y el Índice se Shannon tiene un criterio de 2-3 donde su escala de interpretación es: 0-2 se considera diversidad baja, de 2-3 se considera diversidad media y mayor de 3 se considera diversidad alta. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDe acuerdo al análisis realizado en el área de cambio de uso de suelo y sistema ambiental se tiene lo siguiente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.2.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.2.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.2.1.- Riqueza específica') 
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('La riqueza específica (S) es la forma más sencilla de medir la biodiversidad, ya que se basa únicamente en el número de especies presentes, sin tomar en cuenta el valor de importancia de las mismas. La forma ideal de medir la riqueza específica es contar con un inventario completo que nos permita conocer el número total de especies (S) obtenido por un censo de la comunidad. Esto es posible únicamente para ciertas taxas bien conocidos y de manera puntual en tiempo y en espacio. La mayoría de las veces tenemos que recurrir a índices de riqueza específica obtenidos a partir de un muestreo de la comunidad. A continuación, se describen los índices más comunes para medir la riqueza de especies de acuerdo a (Moreno 2001)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una Comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDmg = Índice de Margalef'
                                    '\nS = Número de especies.'
                                    '\nN = Número total de individuos'
                                    '\nD = Densidad'
                                    '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                    '\nLn= Logaritmo natural')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.2.- Riqueza de especies (Índice de Margalef)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.1.- Riqueza de especies (Índice de Margalef)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De acuerdo al cuadro y gráfico anterior se observa que, en cuanto a Riqueza de especies para los estratos, arbustivo y suculento presentan un valor medio para las dos áreas (ACUSTF y Sistema Ambiental) para el estrato gramíneo ambas áreas presentan valores bajos y para el estrato herbáceo en el ACUSTF valor medio mientras que en el SA el valor es bajo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.1 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.3.- Riqueza de especies (Índice de Menhinick)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.2.- Riqueza de especies (Índice de Menhinick)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.1.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.1.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.2.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.2.2.- Dominancia de especies ')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices basados en la dominancia son parámetros inversos al concepto de uniformidad o equidad de la comunidad. Toman en cuenta la representatividad de las especies con mayor valor de importancia sin evaluar la contribución del resto de las especies. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para medir la dominancia de las especies los índices de biodiversidad más comunes son: Simpson y Berger Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Índice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. Manifiesta la probabilidad de que dos individuos tomados al azar de una muestra sean de la misma especie. Está fuertemente influido por la importancia de las especies más dominantes (Magurran, 1988; Peet, 1974). Es decir, cuanto más se acerca el valor de este índice a la unidad existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.1.2.2 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6122/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde,'
                                        '\nƛ = índice de dominancia se Simpson'
                                        '\nID=índice de diversidad'
                                        '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra'
    )
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia, de 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia y mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.4.- Dominancia de especies (Índice de Simpson)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.3.- Dominancia de especies (Índice de Simpson)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('FAVOR DE DESCRIBIR EL RESTO DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia (Magurran, 1988).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.2 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.5.- Dominancia de especies (Berger Parker)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.4.- Dominancia de especies (Berger Parker)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.1.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.1.2.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.2.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.2.3.- Equidad de especies')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices más comunes para medir la equidad de las especies son Shannon y Pielou.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Índice de Shannon-Wiener (H’). Tiene en cuenta la riqueza de especies y su abundancia. Este índice relaciona el número de especies con la proporción de individuos pertenecientes a cada una de ellas presente en la muestra. Además, mide la uniformidad de la distribución de los individuos entre las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.1.2.3 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6123/formula_8.png', width=Cm(4.89), height=Cm(1.20))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde,'
                                        '\nH’ = índice se Shannon'
                                        '\nS = número de especies'
                                        '\nPi = proporción de individuos de la especie entre todas las especies, A mayor valor de H’ mayor diversidad de especies.'
                                        '\nLn = Logaritmo natural '
    )
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De 0-2 se considera diversidad baja, de 2-3 se considera diversidad media y mayor de 3 se considera diversidad alta')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.6.- Equidad de especies (Índice de Shannon)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.5.- Equidad de especies (Índice de Shannon)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.1.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('FAVOR DE DESCRIBIR EL RESTO DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nEl índice de Pielou: se expresa como el grado de uniformidad en la distribución de individuos entre especies. Se puede medir comparando la diversidad observada en una Comunidad contra la diversidad máxima posible de una comunidad hipotética con el mismo número de especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.1.2.3 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6123/formula_12.png', width=Cm(4.72), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.1.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                        '\nê = índice de Pielou'
                                        '\n∑ = es la sumatoria de la proporción de individuos (pi) por la sumatoria del logaritmo natura de la proporción de individuos (lnpi), o el Índice de Shannon – Wiener '
                                        '\nS = es el número de especies presentes')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDonde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDe 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia, de 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia y mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.1.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.7.- Equidad de especies (Índice de Pielou)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.1.2.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.1.2.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.6.- Equidad de especies (Índice de Pielou)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.1.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.1.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.3.- Comparativo por valor densidad de especies en el Sistema Ambiental -ACUSTF en el ______________________________.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para realizar este comparativo se extrapolaron los individuos de las especies a hectáreas, y se calificó de acuerdo a los siguientes cuadros que mencionan los valores de densidad. Es decir, la densidad de individuos por hectáreas y su respectiva calificación si es vegetación Rala, Semidensa y Densa. Estos cuadros fueron extraídos de la Guía de Métodos para medir la biodiversidad de la revista Área Agropecuaria y de Recursos Naturales Renovables de Ecuador.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.8.- Valores de densidad para estimar la densidad de la vegetación de la Sistema Ambiental- ACUSTF')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.3 ###
    #########################
    # Densidad de los arboles #
    densidadVegetacion = [
        'los árboles', 'los arbustos', 'las hierbas'
    ]

    densidadArboles = [
        '0-300 Individuos/hectárea',
        '301-600 Individuos/hectárea',
        'más de 600 Individuos/hectárea',
    ]

    densidadArbustos = [
        '0-500 Individuos/hectárea',
        '501-1000 Individuos/hectárea',
        'más de 1000 Individuos/hectárea',
    ]

    densidadHierbas = [
        '0-1000 Individuos/hectárea',
        '1001-2000 Individuos/hectárea',
        'más de 2000 Individuos/hectárea',
    ]

    valorPonderado = [
        1.67,
        3.33,
        5,
    ]

    calificacionCap613 = [
        'Vegetación Rala (R)',
        'Vegetación Semidensa (SD)',
        'Vegetación Densa (D)',
    ]

    columnasCap613 = [
        'Valor Calculado de Densidad',
        'Valor Ponderado',
        'Clasificación',
    ]

    densidadVegetacionRango = range(len(densidadVegetacion))
    densidadArbolesRango = range(len(densidadArboles))
    densidadArbustosRango = range(len(densidadArbustos))
    valorPonderadoRango = range(len(valorPonderado))
    densidadHierbasRango = range(len(densidadHierbas))
    calificacionCap613Rango = range(len(calificacionCap613))
    columnasCap613Rango = range(len(columnasCap613))

    VegetacionRango = len(densidadVegetacion)
    ArbolesRango = len(densidadArboles)
    ArbustosRango = len(densidadArbustos)
    ponderadoRango = len(valorPonderado)
    HierbasRango = len(densidadHierbas)
    cap613Rango = len(calificacionCap613)
    cap613ColumnasRango = len(columnasCap613)
    
    tabla6 = doc.add_table(rows=15, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(7.77)
        rows.cells[1].width = Cm(4.06)
        rows.cells[2].width = Cm(6.62)

    #########################
    # Celdas fusionadas "Valores de densidad para estimar la densidad de ..."
    for celda_fusionada in densidadVegetacionRango:
        i = celda_fusionada * 5

        row1 = tabla6.rows[i]
        merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

        # Agregar texto a la celda fusionada
        t6 = merged_cell1.paragraphs[0].add_run(f'Valores de densidad para estimar la densidad de {densidadVegetacion[celda_fusionada]}')
        t6.font.name = 'Arial'
        t6.font.size = Pt(12)
        t6.bold = True
        merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(merged_cell1, '4F81BD')

    for rows in densidadVegetacionRango:
        i = (rows * 5) + 1

        for cols in densidadVegetacionRango:
            cell = tabla6.cell(i, cols)
            t6 = cell.paragraphs[0].add_run(f'{columnasCap613[cols]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            t6.bold = True

    for cols in densidadVegetacionRango:
        i = (cols * 5) + 2

        for valor in densidadVegetacionRango:
            k = i + valor
            cell = tabla6.cell(k, 1)
            t6 = cell.paragraphs[0].add_run(f'{valorPonderado[valor]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for calif in densidadVegetacionRango:
            k = i + calif
            cell = tabla6.cell(k, 2)
            t6 = cell.paragraphs[0].add_run(f'{calificacionCap613[calif]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for arboles in densidadArbolesRango:
        i = arboles + 2
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadArboles[arboles]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    for arbustos in densidadArbustosRango:
        i = arbustos + 7
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadArbustos[arbustos]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    for hierbas in densidadHierbasRango:
        i = arbustos + 12
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadHierbas[hierbas]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 6.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Fuente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Times New Roman'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.font.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('MENDOZA, Zhofre Aguirre. Guía de métodos para medir la biodiversidad. Área Agropecuaria y de Recursos Naturales Renovables. Carrera de Ingeniería Forestal, Universidad Nacional de Loja. Loja-Ecuador, 2013, vol. 37, no 6, p. 82.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Times New Roman'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.italic = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capítulo 6.1.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 6.1.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.9.- Comparativo para la calificación de la densidad de individuos de la Sistema Ambiental- ACUSTF en el ___.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.3 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver  resto del contenido del capítulo 6.1.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripción del capítulo 6.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion de este capitulo en una cuartilla.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.1.4
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina en el capitulo 6.1.4 ###
    #########################
    doc.add_page_break() # Salto de página
    
    #########################
    ### Titulo del capitulo 6.1.4 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.1.4.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.1.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Expresan el grado en el que dos muestras son semejantes por las especies presentes en ellas, por lo que son una medida inversa de la diversidad beta, que se refiere al cambio de especies entre dos muestras (Magurran, 1988; Baev y Penev, 1995; Pielou, 1975). Sin embargo, a partir de un valor de similitud (s) se puede calcular fácilmente la disimilitud (d) entre las muestras: d=1_s (Magurran, 1988). Estos índices pueden obtenerse con base en datos cualitativos o cuantitativos directamente o a través de métodos de ordenación o clasificación de las comunidades (Baev y Penev, 1995).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Se utilizó los índices cualitativos es decir se utiliza presencia y ausencia de especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nCoeficiente de similitud de Sørensen (Czekanovski-Dice-Sørensen)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.1.4 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo614/formula_1.png', width=Cm(2.97), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 6.1.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Relaciona el número de especies en común con la media aritmética de las especies enambos sitios (Magurran, 1988).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nCoeficiente de similitud de Jaccard')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.1.4 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo614/formula_2.png', width=Cm(3.53), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 6.1.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\na = número de especies presentes en el sitio A (Cambio de uso de suelo)'
                                    '\nb = número de especies presentes en el sitio B (Sistema Ambiental)'
                                    '\nc = número de especies presentes en ambos sitios A y B')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nFuente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- POLO URREA, Claudia Sofía. Índices más comunes en biología. Segunda parte, similaridad y riqueza beta y gama. 2008. Facultad de Ciencias Básicas Vol. 4(1): 135-142.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- MORENO, Claudia E. Métodos para medir la biodiversidad. M&T–Manuales y Tesis SEA, vol. 1. Zaragoza, 2001, vol. 84, no 922495, p. 2.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- MAGURRAN, Anne E. Ecological diversity and its measurement. Princeton university press, 1988.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido a partir del capítulo 6.1.3 al capitulo 6.2.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los resultados obtenidos son los siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.1.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.X.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___________________..')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.1.3 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Describir el resto del capitulo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.2.1
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            Todavia continua las hojas en Horizontal.
        ==================================================================================================================================================================
    """
    #########################
    ### Titulo del capitulo 6 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nV.8.3.6.- Análisis de la información de la fauna en el ACUSTF.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.X.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___________________..')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDescribir el resto del capitulo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver  resto del contenido del capítulo 6.2.1 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            * Aqui terminan las hojas en Horizontal. (Comienzan desde del capitulo 6.1.3 al capitulo 6.2.1)
            * El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripción del capítulo 6.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion de este capitulo de acuerdo a lo de la tabla anterior.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('En general las especies de lento crecimiento y las enlistadas en la NOM- 059- SEMARNAT 2010 se rescatarán y reubicarán a una superficie que tenga las mismas condiciones donde se distribuyen actualmente para que no pierdan su germoplasma.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Describir el resto del capitulo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.2.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.2.- Comparativo por índices de biodiversidad del Sistema Ambiental – ACUSTF ______')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para realizar este comparativo se utilizó la Metodología para el estudio de las Comunidades vegetales. También se menciona que estos índices fueron estimados con individuos reales, es decir con individuos muestreados en los 14 sitios de muestreo, tanto del ACUSTF como Sistema Ambiental. Recalcando que los resultados se interpretan usando la siguiente escala de significancia entre 0-1 para los índices de Simpson, Berger Parker y Pielou, donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 6.2.2 ###
    #########################
    listaValores622 = [
        '0 - 0.33',
        '0.34 - 0.66',
        '> 0.67',
    ]

    diversidadSignificancia622 = [
        'Diversidad Baja',
        'Diversidad Media',
        'Diversidad Alta',
    ]

    heterogeneoHomogeneoSignificancia622 = [
        'Heterogéneo en abundancia',
        'Ligeramente Heterogéneo en en abundancia',
        'Homogéneo en abundancia',
    ]

    valores622 = range(len(listaValores622))
    diversidad622 = range(len(diversidadSignificancia622))
    heterogeneoHomogeneo622 = range(len(heterogeneoHomogeneoSignificancia622))

    filasCap622 = len(valores622) + 2
    
    tabla6 = doc.add_table(rows=filasCap622, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(3.46)
        rows.cells[1].width = Cm(3.66)
        rows.cells[2].width = Cm(9.66)

    #########################
    # Celda fusionada "Escalas de interpretación de significancia 0-1"
    row1 = tabla6.rows[0]
    merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Escalas de interpretación de significancia 0-1')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_background_color(merged_cell1, '0070C0')

    #########################
    # Celda fusionada "Significancia"
    row1 = tabla6.rows[1]
    merged_cell1 = row1.cells[1].merge(row1.cells[1].merge(row1.cells[2]))

    # Agregar texto a la celda fusionada
    t6 = merged_cell1.paragraphs[0].add_run('Significancia')
    t6.font.name = 'Arial'
    t6.font.size = Pt(12)
    t6.bold = True
    merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell = tabla6.cell(1, 0)
    t6 = cell.paragraphs[0].add_run('Valores')
    t6.font.size = Pt(12)
    t6.font.name = 'Arial'
    t6.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in valores622:
        cell = tabla6.cell(cols + 2, 0)
        t6 = cell.paragraphs[0].add_run(f'{listaValores622[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        t6.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in diversidad622:
        cell = tabla6.cell(cols + 2, 1)
        t6 = cell.paragraphs[0].add_run(f'{diversidadSignificancia622[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    for cols in heterogeneoHomogeneo622:
        cell = tabla6.cell(cols + 2, 2)
        t6 = cell.paragraphs[0].add_run(f'{heterogeneoHomogeneoSignificancia622[cols]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

    #########################
    ### Descripcion del capitulo 6.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el Índice de Margalef el criterio es de 2-5, donde, sus escalas de interpretación son: de 0-2 se considera diversidad baja, de 2-5 se considera diversidad media y mayor de 5 se considera diversidad alta y el Índice se Shannon tiene un criterio de 2-3 donde su escala de interpretación es: 0-2 se considera diversidad baja, de 2-3 se considera diversidad media y mayor de 3 se considera diversidad alta. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDe acuerdo al análisis realizado en el área de cambio de uso de suelo y sistema ambiental se tiene lo siguiente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.2.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.2.2.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.2.1.- Riqueza específica') 
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('La riqueza específica (S) es la forma más sencilla de medir la biodiversidad, ya que se basa únicamente en el número de especies presentes, sin tomar en cuenta el valor de importancia de las mismas. La forma ideal de medir la riqueza específica es contar con un inventario completo que nos permita conocer el número total de especies (S) obtenido por un censo de la comunidad. Esto es posible únicamente para ciertas taxas bien conocidos y de manera puntual en tiempo y en espacio. La mayoría de las veces tenemos que recurrir a índices de riqueza específica obtenidos a partir de un muestreo de la comunidad. A continuación, se describen los índices más comunes para medir la riqueza de especies de acuerdo a (Moreno 2001)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Margalef. - Es utilizado para estimar la biodiversidad de una Comunidad con base en la distribución numérica de los individuos de las diferentes especies en función del número de individuos existentes en los sitios de muestreo. Valores inferiores a dos son considerados como zonas de baja biodiversidad y valores superiores a cinco son indicativos de alta biodiversidad.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.2.2.1 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_9.png', width=Cm(3.54), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDmg = Índice de Margalef'
                                    '\nS = Número de especies.'
                                    '\nN = Número total de individuos'
                                    '\nD = Densidad'
                                    '\nValores cercanos a 1 representan condiciones hacia especies igualmente abundantes y aquellos cercanos a 0 la dominancia de una sola especie.'
                                    '\nLn= Logaritmo natural')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.2.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Riqueza de especies (Índice de Margalef)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.7.- Riqueza de especies (Índice de Margalef)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.2.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De acuerdo al cuadro y gráfico anterior se observa que, en cuanto a Riqueza de especies para los estratos, arbustivo y suculento presentan un valor medio para las dos áreas (ACUSTF y Sistema Ambiental) para el estrato gramíneo ambas áreas presentan valores bajos y para el estrato herbáceo en el ACUSTF valor medio mientras que en el SA el valor es bajo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de diversidad de Menhinick. - Se basa en la relación entre el número de especies y el número total de individuos observados, que aumenta al aumentar el tamaño de la muestra.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.2.2.1 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_11.png', width=Cm(2.91), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.2.2.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Riqueza de especies (Índice de Menhinick)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.8.- Riqueza de especies (Índice de Menhinick)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.2.2.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.2.2.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.2.2.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.2.2.- Dominancia de especies ')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices basados en la dominancia son parámetros inversos al concepto de uniformidad o equidad de la comunidad. Toman en cuenta la representatividad de las especies con mayor valor de importancia sin evaluar la contribución del resto de las especies. (Moreno, 2001).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para medir la dominancia de las especies los índices de biodiversidad más comunes son: Simpson y Berger Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Índice de diversidad de Simpson. - Se obtiene de un determinado número de especies presentes en el hábitat y su abundancia absoluta expresado al cuadrado. Manifiesta la probabilidad de que dos individuos tomados al azar de una muestra sean de la misma especie. Está fuertemente influido por la importancia de las especies más dominantes (Magurran, 1988; Peet, 1974). Es decir, cuanto más se acerca el valor de este índice a la unidad existe una mayor posibilidad de dominancia de una especie en una población.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.2.2.2 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6122/formula_10.png', width=Cm(3.25), height=Cm(1.50))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde,'
                                        '\nƛ = índice de dominancia se Simpson'
                                        '\nID=índice de diversidad'
                                        '\npi = es la abundancia relativa de la especie (pi), es decir, el número de individuos de la especie (p), i dividido entre el número total de individuos de la muestra'
    )
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia, de 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia y mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Dominancia de especies (Índice de Simpson)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.9.- Dominancia de especies (Índice de Simpson)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.2.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('FAVOR DE DESCRIBIR EL RESTO DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Berger-Parker Es un índice que interpreta un aumento en la equidad y una disminución en la dominancia (Magurran, 1988).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.2.2.2 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6121/formula_13.png', width=Cm(2.88), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\nDMn = Índice de Menhinick'
                                    '\nS = Número total de especies'
                                    '\nN = Número total de todos los individuos de todas las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nBajo los índices anteriormente descritos se realizó una comparación de índices de vegetación para biodiversidad, que de acuerdo a (Moreno 2001) estiman la riqueza de especies, señalando que para poder compararlos se realizaron las estimaciones con datos de muestreo reales (datos de los sitios de muestreo) para no sobreestimar a la hora de extrapolarlos a las áreas correspondientes.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nPara el índice de Menhinick el criterio de evaluación es de 1-2, donde la escala de interpretación es menor a 1 se considera diversidad baja, de 1-2 se considera diversidad media y mayor de 2 se considera diversidad alta.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.2.2.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Dominancia de especies (Berger Parker)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.10.- Dominancia de especies (Berger Parker)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.2.2.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.2.2.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.2.2.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.2.3.- Equidad de especies')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).Algunos de los índices más reconocidos sobre diversidad se basan principalmente en el concepto de equidad, por lo que se describen en esta sección. Al respecto se pueden encontrar discusiones profundas en Peet (1975), Camargo (1995), Smith y Wilson (1996) y Hill (1997).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices más comunes para medir la equidad de las especies son Shannon y Pielou.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Índice de Shannon-Wiener (H’). Tiene en cuenta la riqueza de especies y su abundancia. Este índice relaciona el número de especies con la proporción de individuos pertenecientes a cada una de ellas presente en la muestra. Además, mide la uniformidad de la distribución de los individuos entre las especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.2.2.3 ###
    #########################
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6123/formula_8.png', width=Cm(4.89), height=Cm(1.20))

    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde,'
                                        '\nH’ = índice se Shannon'
                                        '\nS = número de especies'
                                        '\nPi = proporción de individuos de la especie entre todas las especies, A mayor valor de H’ mayor diversidad de especies.'
                                        '\nLn = Logaritmo natural '
    )
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('De 0-2 se considera diversidad baja, de 2-3 se considera diversidad media y mayor de 3 se considera diversidad alta')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Equidad de especies (Índice de Shannon)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.11.- Equidad de especies (Índice de Shannon)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripcion del capitulo 6.2.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('FAVOR DE DESCRIBIR EL RESTO DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nEl índice de Pielou: se expresa como el grado de uniformidad en la distribución de individuos entre especies. Se puede medir comparando la diversidad observada en una Comunidad contra la diversidad máxima posible de una comunidad hipotética con el mismo número de especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Fórmula del capítulo 6.2.2.3 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo6123/formula_12.png', width=Cm(4.72), height=Cm(1.50))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripción del capítulo 6.2.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                        '\nê = índice de Pielou'
                                        '\n∑ = es la sumatoria de la proporción de individuos (pi) por la sumatoria del logaritmo natura de la proporción de individuos (lnpi), o el Índice de Shannon – Wiener '
                                        '\nS = es el número de especies presentes')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDonde las escalas para la interpretación de los rangos son las siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDe 0 – 0.33 se considera diversidad baja o Heterogéneo en abundancia, de 0.34 – 0.66 se considera diversidad media o Ligeramente Heterogéneo en abundancia y mayor de 0.67 se considera diversidad alta o Homogéneo en abundancia.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Título de la tabla del capítulo 6.2.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Equidad de especies (Índice de Pielou)')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.2.2.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.2.2.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.12.- Equidad de especies (Índice de Pielou)')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.2.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('DESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.2.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.2.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.3.- Comparativo por valor densidad de especies en el Sistema Ambiental -ACUSTF en el ______________________________.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para realizar este comparativo se extrapolaron los individuos de las especies a hectáreas, y se calificó de acuerdo a los siguientes cuadros que mencionan los valores de densidad. Es decir, la densidad de individuos por hectáreas y su respectiva calificación si es vegetación Rala, Semidensa y Densa. Estos cuadros fueron extraídos de la Guía de Métodos para medir la biodiversidad de la revista Área Agropecuaria y de Recursos Naturales Renovables de Ecuador.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Valores de densidad para estimar la densidad de la vegetación de la Sistema Ambiental- ACUSTF')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.3 ###
    #########################
    # Densidad de los arboles #
    densidadVegetacion = [
        'los árboles', 'los arbustos', 'las hierbas'
    ]

    densidadArboles = [
        '0-300 Individuos/hectárea',
        '301-600 Individuos/hectárea',
        'más de 600 Individuos/hectárea',
    ]

    densidadArbustos = [
        '0-500 Individuos/hectárea',
        '501-1000 Individuos/hectárea',
        'más de 1000 Individuos/hectárea',
    ]

    densidadHierbas = [
        '0-1000 Individuos/hectárea',
        '1001-2000 Individuos/hectárea',
        'más de 2000 Individuos/hectárea',
    ]

    valorPonderado = [
        1.67,
        3.33,
        5,
    ]

    calificacionCap623 = [
        'Vegetación Rala (R)',
        'Vegetación Semidensa (SD)',
        'Vegetación Densa (D)',
    ]

    columnasCap623 = [
        'Valor Calculado de Densidad',
        'Valor Ponderado',
        'Clasificación',
    ]

    densidadVegetacionRango = range(len(densidadVegetacion))
    densidadArbolesRango = range(len(densidadArboles))
    densidadArbustosRango = range(len(densidadArbustos))
    valorPonderadoRango = range(len(valorPonderado))
    densidadHierbasRango = range(len(densidadHierbas))
    calificacionCap623Rango = range(len(calificacionCap623))
    columnasCap623Rango = range(len(columnasCap623))

    VegetacionRango = len(densidadVegetacion)
    ArbolesRango = len(densidadArboles)
    ArbustosRango = len(densidadArbustos)
    ponderadoRango = len(valorPonderado)
    HierbasRango = len(densidadHierbas)
    cap623Rango = len(calificacionCap613)
    cap623ColumnasRango = len(columnasCap613)
    
    tabla6 = doc.add_table(rows=15, cols=3, style='Table Grid')
    tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # Ancho de Celdas #
    for rows in tabla6.rows:
        rows.cells[0].width = Cm(7.77)
        rows.cells[1].width = Cm(4.06)
        rows.cells[2].width = Cm(6.62)

    #########################
    # Celdas fusionadas "Valores de densidad para estimar la densidad de ..."
    for celda_fusionada in densidadVegetacionRango:
        i = celda_fusionada * 5

        row1 = tabla6.rows[i]
        merged_cell1 = row1.cells[0].merge(row1.cells[0].merge(row1.cells[2]))

        # Agregar texto a la celda fusionada
        t6 = merged_cell1.paragraphs[0].add_run(f'Valores de densidad para estimar la densidad de {densidadVegetacion[celda_fusionada]}')
        t6.font.name = 'Arial'
        t6.font.size = Pt(12)
        t6.bold = True
        merged_cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_background_color(merged_cell1, '4F81BD')

    for rows in densidadVegetacionRango:
        i = (rows * 5) + 1

        for cols in densidadVegetacionRango:
            cell = tabla6.cell(i, cols)
            t6 = cell.paragraphs[0].add_run(f'{columnasCap623[cols]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            t6.bold = True

    for cols in densidadVegetacionRango:
        i = (cols * 5) + 2

        for valor in densidadVegetacionRango:
            k = i + valor
            cell = tabla6.cell(k, 1)
            t6 = cell.paragraphs[0].add_run(f'{valorPonderado[valor]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for calif in densidadVegetacionRango:
            k = i + calif
            cell = tabla6.cell(k, 2)
            t6 = cell.paragraphs[0].add_run(f'{calificacionCap623[calif]}')
            t6.font.size = Pt(12)
            t6.font.name = 'Arial'
            tabla6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for arboles in densidadArbolesRango:
        i = arboles + 2
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadArboles[arboles]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    for arbustos in densidadArbustosRango:
        i = arbustos + 7
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadArbustos[arbustos]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    for hierbas in densidadHierbasRango:
        i = arbustos + 12
        cell = tabla6.cell(i, 0)
        t6 = cell.paragraphs[0].add_run(f'{densidadHierbas[hierbas]}')
        t6.font.size = Pt(12)
        t6.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 6.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Fuente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Times New Roman'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.font.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('MENDOZA, Zhofre Aguirre. Guía de métodos para medir la biodiversidad. Área Agropecuaria y de Recursos Naturales Renovables. Carrera de Ingeniería Forestal, Universidad Nacional de Loja. Loja-Ecuador, 2013, vol. 37, no 6, p. 82.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Times New Roman'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.italic = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capítulo 6.2.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Título de la tabla del capítulo 6.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo para la calificación de la densidad de individuos de la Sistema Ambiental- ACUSTF en el ___.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=7, style='Table Grid')

    for cols in range(7):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capituo en base a lo de la tabla anterior =).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Vertical para ver  resto del contenido del capítulo 6.2.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Vertical.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Descripción del capítulo 6.2.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion de este capitulo en base a los datos de la tabla anterior, Resto del capitulo.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.2.4
    ########################################################################################################################################################################
    
    #########################
    ### Salto de Pagina en el capitulo 6.2.4 ###
    #########################
    doc.add_page_break() # Salto de página
    
    #########################
    ### Titulo del capitulo 6.2.4 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.4.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.2.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Expresan el grado en el que dos muestras son semejantes por las especies presentes en ellas, por lo que son una medida inversa de la diversidad beta, que se refiere al cambio de especies entre dos muestras (Magurran, 1988; Baev y Penev, 1995; Pielou, 1975). Sin embargo, a partir de un valor de similitud (s) se puede calcular fácilmente la disimilitud (d) entre las muestras: d=1_s (Magurran, 1988). Estos índices pueden obtenerse con base en datos cualitativos o cuantitativos directamente o a través de métodos de ordenación o clasificación de las comunidades (Baev y Penev, 1995).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Se utilizó los índices cualitativos es decir se utiliza presencia y ausencia de especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nCoeficiente de similitud de Sørensen (Czekanovski-Dice-Sørensen)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.2.4 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo614/formula_1.png', width=Cm(2.97), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 6.2.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Relaciona el número de especies en común con la media aritmética de las especies enambos sitios (Magurran, 1988).')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nCoeficiente de similitud de Jaccard')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Fórmula del capítulo 6.2.4 ###
    #########################
    # Agregar un párrafo para contener la imagen
    formulaCapitulo6_parrafo = doc.add_paragraph()
    formulaCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen

    # Insertar la imagen dentro del párrafo
    formulaCapitulo6_run = formulaCapitulo6_parrafo.add_run('')
    imagen = formulaCapitulo6_run.add_picture('capitulo6/capitulo614/formula_2.png', width=Cm(3.53), height=Cm(1.20))

    # Opcional: espacio después del párrafo
    formulaCapitulo6_parrafo.space_after = Pt(1)

    #########################
    ### Descripcion del capitulo 6.2.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Donde:'
                                    '\na = número de especies presentes en el sitio A (Cambio de uso de suelo)'
                                    '\nb = número de especies presentes en el sitio B (Sistema Ambiental)'
                                    '\nc = número de especies presentes en ambos sitios A y B')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nFuente:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- POLO URREA, Claudia Sofía. Índices más comunes en biología. Segunda parte, similaridad y riqueza beta y gama. 2008. Facultad de Ciencias Básicas Vol. 4(1): 135-142.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- MORENO, Claudia E. Métodos para medir la biodiversidad. M&T–Manuales y Tesis SEA, vol. 1. Zaragoza, 2001, vol. 84, no 922495, p. 2.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('- MAGURRAN, Anne E. Ecological diversity and its measurement. Princeton university press, 1988.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(10)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nLos resultados obtenidos son los siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    ### Hoja en Horizontal para ver contenido del capitulo 6.2.4 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            El siguiente código muestra cómo se tiene que insertar la hoja en Horizontal.
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los resultados obtenidos son los siguientes:')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.X.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___________________.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Describir el resto del capitulo en base de los datos de la tabla anterior.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.3
    ########################################################################################################################################################################
    
    ########################################################################################################################################################################
    ### Hoja en Vertical para ver el contenido del capítulo 6.3 ###
    ########################################################################################################################################################################
    """
        ==================================================================================================================================================================
            * Capitulo 6.2: Hoja en Horizontal anteriormente.
            * Capitulo 6.3: Se cambia a hora Vertical por el contenido del capitulo.
            * El siguiente código muestra cómo se tiene que insertar la hoja en Vertical:
        ==================================================================================================================================================================
    """

    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Cambiar orientación a vertical
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2.5)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)

    #########################
    ### Titulo del capitulo 6.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3. Fauna ')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('La medición de la composición faunística muestra el estado de una población en base a los resultados obtenidos mediante la utilización de distintos índices que miden la diversidad, de los cuales y en base a los resultados obtenido se realizaron análisis comparativos entre ambas áreas de estudio para conocer de manera más precisa las diferencias que existen entre una y otra.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.2.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.X.- Comparativos de Índices de similitud/disimilitud en el Sistema Ambiental y ACUSTF en el ___________________.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.2.3 ###
    #########################
    tabla6b = doc.add_table(rows=40, cols=10, style='Table Grid')

    for cols in range(10):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(40):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    ########################################################################################################################################################################
    # Capitulo 6.3.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.1.1.-Grupo de Aves.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.3.1.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.1.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.1.1.-Grupo de Aves.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('La abundancia absoluta (Ai) representa el número de individuos avistados por especies que se encuentran en el área de estudio en cuestión, por otra parte, la abundancia relativa se destaca en la relación porcentual del número de individuo de una especie con respecto al total de individuos que se observan en la parcela o área de estudio.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.1.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de abundancias ACUSTF y sistema ambiental.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.1.1 ###
    #########################
    tabla6b = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.3.1.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.1.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.13.- Análisis comparativo de abundancias para el grupo de las aves en el ACUSTF y sistema ambiental.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.3.1.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.3.1.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.1.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.1.2.-Grupo de Mamíferos.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion (Opcional)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.1.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.X.- Comparativo de abundancias para el grupo de los mamíferos.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.1.2 ###
    #########################
    tabla6b = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.3.1.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.1.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.14.- Análisis comparativo de abundancias para el grupo de los mamíferos en el ACUSTF y sistema ambiental.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.3.1.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.3.1.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.1.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.1.3.-Grupo de Reptiles.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion (Opcional)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.1.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de abundancias del grupo de los reptiles.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.1.3 ###
    #########################
    tabla6b = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.3.1.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.1.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.15.- Análisis de comparativo de densidades para el grupo de los reptiles en ACUSTF y sistema ambiental.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.3.1.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.3.1.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.1.4 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.1.4.-Grupo de los Lepidópteros.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.1.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion (Opcional)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.1.4 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de abundancia para el grupo de los insectos.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.1.4 ###
    #########################
    tabla6b = doc.add_table(rows=10, cols=6, style='Table Grid')

    for cols in range(6):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(10):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Gráfica del capítulo 6.3.1.4 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.1.4 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.16.- Analisis de comparativo de denisdades para el grupo de los insectos en el ACUSTF y sistema ambinetal.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Descripción del capítulo 6.3.1.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    ########################################################################################################################################################################
    # Capitulo 6.3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.2.- Análisis comparativo por densidades en el ACUSTF y sistema ambiental.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.25.- Análisis comparativo de los grupos faunísticos en el ACUSTF y sistema ambiental.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.2 ###
    #########################
    tabla6b = doc.add_table(rows=20, cols=4, style='Table Grid')

    for cols in range(4):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.17.- Comparativo por densidades para el grupo de las aves.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Gráfica del capítulo 6.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.18.- Comparativo por densidades para el grupo de los mamíferos.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Gráfica del capítulo 6.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.19.- Comparativo por densidades para el grupo de los reptiles.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Gráfica del capítulo 6.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.20.- Comparativo de densidades para el grupo de los insectos.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6.3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.3.- Análisis comparativo por índices de diversidad.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Capitulo 6.3.3.1
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.3.1 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.3.1.- Índice de Equidad.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.3.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para determinar los valores de equidad de especies en las áreas de estudio se utilizaron dos índices que nos permitirán conocer el tipo de equidad que se presenta en dichas áreas de estudio, los índices utilizados fueron el índice de Shannon-Wiener el cual es un índice que asume que todas las especies están representadas en las muestras y que todos los individuos muestreados fueron al azar, midiendo así el grado de incertidumbre en predecir a que especie pertenecerá un individuo elegido al azar de una muestra. Sus valores van de 0 a 1.35 son correspondientes a un valor bajo; 1.36 a 3.5 son valores medios y mayor de 3.5 son valores altos; y el índice de Pielou el cual mide la proporción de la diversidad observada en relación a la máxima diversidad esperada los valores de este índice van de 0 a 1 en donde aquel valor más cercano en donde hace referencia a que los valores resultantes a 1 o más cercanos a este hace inferencia a que las especies son igualmente abundantes, en cambio para valores a cero se señala ausencia de uniformidad de especies.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Índice de Shannon ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Shannon.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.3.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de equidad por el índice de Shannon.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.3.1 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.21.- Comparativo de equidad por el índice de Shannon.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Índice de Pielou ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Pielou.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.3.1 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de equidad por el Índice de Pielou.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.1 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.1 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.1 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.22.- Comparativo de equidad por el Índice de Pielou.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6.3.3.2
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.3.2 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.3.2.- Índice de Dominancia.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.3.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para el cálculo de la dominancia de especies se utilizaron dos índices, índice de Simpson e índice de Berger-Parker, estos índices buscan determinar qué tan representativas son las especies que se encuentran en un área o sitio. Siendo que para el índice de dominancia mediante el índice de Simpson determina la probabilidad que existe entre la selección de individuos al azar en una muestra y estos corresponden a la misma especie, siendo influido por la importancia de las especies dominantes, las interpretaciones de los valores del índice de Simpson van de 0 a 0.33 para valores bajos; 0.34 a 0.66 para valores medios y mayores a 0.67 para valores altos. El índice de Berger-Parker, mide la proporción de las especies más comunes en una comunidad o muestra, sus valores van de 0 para valores bajo y de 1 o cercanos a este para valores altos.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Índice de Simpson. ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Simpson.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.3.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de dominancia por el índice de Simpson.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.3.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.23.- Comparativo de dominancia por el índice de Simpson.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Índice de Berger-Parker. ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Berger-Parker.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.2 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de equidad por el Índice de Pielou.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.3.2 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.2 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.2 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.2 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.24.- Análisis comparativo de dominancia por el índice de Berger-Parker.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6.3.3.3
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.3.3 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.3.3.- Índice de Riqueza.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.3.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Para la medición de la riqueza específica la cual es la manera más adecuada para conocer la biodiversidad de las especies, se basa en el número de especies presentes sin tomar en cuenta el valor de importancia de las mismas; para ello se utilizaron dos índices para medir la biodiversidad del ACUSTF y el sistema ambiental tales índices fueron, el índice de Margalef el cual se encuentra relacionado con el número de especies de acuerdo con el número total de individuos y el índice de Menhinick que se basa en la relación entre el número de especies y el número total de individuos observados, los rangos de valor para Margalef van de valores menores a 2 corresponden a una riqueza baja y valores a 5 son valores de una riqueza alta, los valores para Menhinick van de 0 a 1 para valores bajos , valores medios de 1 a 2 y valores altos aquellos superiores a 2.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Índice de Margalef ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Margalef')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.3.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de dominancia por el Índice de Margalef.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.3.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.25.- Comparativo de dominancia por el Índice de Margalef.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Índice de Menhinick. ###
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nÍndice de Menhinick.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    descripcionCapitulo6.bold = True
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.3.3 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de equidad por el Índice de Menhinick.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.3.3 ###
    #########################
    tabla6b = doc.add_table(rows=6, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(6):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripción del capítulo 6.3.3.3 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\nDESCRIPCION DEL CAPITULO en base a los datos anteriores')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)

    #########################
    ### Gráfica del capítulo 6.3.3.3 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.3.3 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.26.- Análisis comparativo de dominancia por el Índice de Menhinick.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6.3.4
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.4 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.2.4.- Análisis comparativo por índices de similitud.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Los índices de similitud expresan el grado de semejanza entre dos áreas calculado por las especies presentes en cada área; para este caso se analizaron dos áreas ACUSTF y sistema ambiental determinando, los valores de medición van de 0 a 0.33 como áreas diferentes, 0.34 a 0.66 para medianamente similares y 0.67 a 1 para áreas similares.')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Título de la tabla del capítulo 6.3.4 ###
    #########################
    tituloTabla6b = doc.add_paragraph()
    dti6b = tituloTabla6b.add_run('\nTabla 6.x.- Comparativo de similitud para los grupos faunísticos en el ACUSTF y sistema ambiental.')
    dti6b_format = tituloTabla6b.paragraph_format
    dti6b_format.line_spacing = 1.15
    dti6b_format.space_after = 0

    dti6b.font.name = 'Bookman Old Style'
    dti6b.font.size = Pt(12)
    tituloTabla6b.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    ### Tabla del capítulo 6.3.4 ###
    #########################
    tabla6b = doc.add_table(rows=20, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(20):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'
    
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('\n')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Tabla del capítulo 6.3.4 ###
    #########################
    tabla6b = doc.add_table(rows=5, cols=5, style='Table Grid')

    for cols in range(5):
        cell = tabla6b.cell(0, cols)
        cell_background_color(cell, '0070C0')

        for rows in range(5):
            cell = tabla6b.cell(rows, cols)
            t6b = cell.paragraphs[0].add_run(' ')
            t6b.font.size = Pt(12)
            t6b.font.name = 'Arial'

    #########################
    ### Descripcion del capitulo 6.3.4 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Describir el resto del capitulo de acuerdo a los datos anteriores =)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Gráfica del capítulo 6.3.4 ###
    #########################
    """ 
        El siguiente código muestra cómo se tiene que insertar la imagen, mapa o gráfico.
    """
    imagenCapitulo6_parrafo = doc.add_paragraph()
    imagenCapitulo6_run = imagenCapitulo6_parrafo.add_run('\n')
    imagenCapitulo6_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grafico = imagenCapitulo6_run.add_picture('capitulo6/grafico.jpg', width=Cm(7.48), height=Cm(5))

    #########################
    ### Título de la gráfica del capítulo 6.3.4 ###
    #########################
    tituloGrafico6 = doc.add_paragraph()
    dgi6 = tituloGrafico6.add_run('Grafica 6.27.- Comparativo de similitud por grupos faunísticos en el ACUSTF y sistema ambiental.')
    dgi6_format = tituloGrafico6.paragraph_format
    dgi6_format.line_spacing = 1.15
    dgi6_format.space_after = 0

    dgi6.font.name = 'Bookman Old Style'
    dgi6.font.size = Pt(12)
    tituloGrafico6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ########################################################################################################################################################################
    # Capitulo 6.3.5
    ########################################################################################################################################################################
    
    #########################
    ### Titulo del capitulo 6.3.5 ###
    #########################
    capitulo6 = doc.add_paragraph()
    i6 = capitulo6.add_run(f'\nVI.3.5.- Análisis de la información de la fauna en el ACUSTF y Sistema Ambiental.')
    i6_format = capitulo6.paragraph_format
    i6_format.line_spacing = 1.15

    i6.font.name = 'Arial'
    i6.font.size = Pt(12)
    i6.font.bold = True
    capitulo6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #########################
    ### Descripcion del capitulo 6.3.5 ###
    #########################
    di6 = doc.add_paragraph()
    descripcionCapitulo6 = di6.add_run('Descripcion del capitulo (1 cuartilla y media)')
    descripcionCapitulo6_format = di6.paragraph_format
    descripcionCapitulo6_format.line_spacing = 1.15
    descripcionCapitulo6_format.space_after = 0
    descripcionCapitulo6_format.space_before = 0

    descripcionCapitulo6.font.name = 'Arial'
    descripcionCapitulo6.font.size = Pt(12)
    di6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ########################################################################################################################################################################
    # Guardar Documento
    ########################################################################################################################################################################
    """
        Nota: Para guardar los documentos tienen que ser gardados de dos formas:
            1. Si no se quiere cambiar el titulo, por ejemplo:
                doc.save('CAPITULO 1 DTU NOMBRE DEL PROYECTO')  ---> Si no se necesita cambiar nada

            2. Si se quiere poner variables, por ejemplo

                nombreProyecto = 'NOMBRE DEL PROYECTO'          ---> Variable que se va a utilizar
                doc.save('CAPITULO 1 DTU ' + nombreProyecto)    ---> Puede ser de esta manera

                o bien:

                nombreProyecto = 'NOMBRE PROYECTO'             ---> Variable que se va a utilizar
                doc.save(`CAPITULO 1 DTU ${nombreProyecto}`)   ---> Puede ser de esta manera

            3. Una vez guardado; editar el indice y los titulos, ya que Python no puede o no tiene soporte para hacer un indice de tabla de contenido

            4. Tener consideracion para editar el documentos, debido a que Python no tiene soporte para editar documentos word, por lo que se debe editar manualmente
                -> Editar el diseño de las tablas.
                -> Editar el diseño de los titulos.
                -> Editar el diseño de los indices.
                -> Editar el numero de paginas.

            5. Una parte del documentos se tiene que editar con el encabezado y el pie de pagina, si se pueden editar, pero por consideracion del cliente, se debe editar manualmente
                -> Editar el encabezado.
                    * El encabezado debe que tener el nombre del proyecto

                -> Editar el pie de pagina.
                    * El pie de pagina debe tener el nombre del cliente
    """
    doc.save("CAPITULO 6 DTU EXTRACCION DE MATERIAL PETRO.docx")

"""
    ============================================================
        Llamar la funcion para poder crear el documento
    ============================================================
"""
capitulo6() # Crear el documento